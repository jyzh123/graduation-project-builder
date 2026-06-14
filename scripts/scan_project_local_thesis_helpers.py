#!/usr/bin/env python3
"""Preflight scan for risky project-local thesis helper scripts.

The scanner is intentionally conservative for thesis lanes: project-local
scripts may be thin wrappers, but they must not become their own DOCX rewrite
engines for headings, abstracts, captions, tables, references, images, or
pagination.
"""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
import json
import shlex
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from validate_thesis_local_adapter import validate_adapter_file


SCRIPT_EXTENSIONS = {".py", ".ps1", ".bat", ".cmd", ".sh", ".js", ".mjs", ".cjs", ".ts", ".ipynb"}
ADAPTER_EXTENSIONS = {".json"}
ADAPTER_NAME_MARKERS = (
    "thesis-local-adapter",
    "thesis_template_profile",
    "thesis-template-profile",
    "thesis_project_manifest",
    "thesis-project-manifest",
    "thesis-run-manifest",
    "thesis_adapter",
    "thesis-adapter",
)
SKIP_DIR_NAMES = {
    ".git",
    ".hg",
    ".svn",
    ".idea",
    ".vscode",
    ".venv",
    "venv",
    "env",
    "node_modules",
    "vendor",
    "site-packages",
    "dist-packages",
    "archive",
    "backups",
    "__pycache__",
    ".pytest_cache",
    ".mypy_cache",
    ".codex_tmp_pylibs",
}
SKIP_DIR_PREFIXES = (
    "quarantine_high_risk_thesis_scripts",
)

DOCX_ENGINE_MARKERS = (
    "from docx import Document",
    "docx import Document",
    "Document(",
    "word/document.xml",
    "word/footer",
    "word/header",
    "[Content_Types].xml",
    "sectPr",
    "pgNumType",
    "w:pStyle",
    "w:rPr",
    "JSZip",
    "PizZip",
    "adm-zip",
    "docxtemplater",
)
DOCX_PACKAGE_SURFACE_MARKERS = (
    ".docx",
    "word/document.xml",
    "word/footer",
    "word/header",
    "[Content_Types].xml",
    "sectPr",
    "pgNumType",
    "w:pStyle",
    "w:rPr",
)
DESTRUCTIVE_PARAGRAPH_MARKERS = (
    "paragraph.clear(",
    "clear_paragraph(",
    "set_paragraph_text(",
    "replace_paragraph_text(",
    "add_run(",
    "p.remove(child)",
    "._p.remove(",
    "writestr(",
    "writeFileSync(",
    ".replace(",
    ".replaceAll(",
)
PROTECTED_SURFACE_MARKERS = (
    "heading",
    "caption",
    "abstract",
    "keyword",
    "toc",
    "bibliograph",
    "reference",
    "table",
    "figure",
    "header",
    "footer",
    "pagination",
    "标题",
    "摘要",
    "关键词",
    "目录",
    "参考文献",
    "图题",
    "表题",
    "页眉",
    "页脚",
)
POWERPOINT_WORD_COM_MARKERS = (
    "Range.Style =",
    "InlineShapes.AddPicture",
    "ParagraphFormat.LineSpacingRule",
    "ParagraphFormat.FirstLineIndent",
    "Set-ParagraphText",
    "Find-Paragraph",
    "New-Object -ComObject Word.Application",
    "TablesOfContents.Item",
    "Fields.Update",
    "Repaginate",
)
LIBREOFFICE_INDEX_MARKERS = (
    "import uno",
    "loadComponentFromURL",
    "UpdateDocMode",
    "getDocumentIndexes",
    "getTextFields",
    "doc.store()",
    "storeToURL",
)
ZIP_XML_REWRITE_MARKERS = (
    "word/document.xml",
    "word/footer",
    "word/header",
    "[Content_Types].xml",
    "writestr(",
    "JSZip",
    "PizZip",
    "adm-zip",
    "writeFileSync(",
)
ZIP_WRITE_MARKERS = (
    "writestr(",
    "writeFileSync(",
    "zip.file(",
    ".replace(",
    ".replaceAll(",
    "os.replace(",
    "Move-Item",
    "Copy-Item",
    '"w"',
    "'w'",
    '"a"',
    "'a'",
    '"Update"',
    "'Update'",
)
LOCAL_DRAWING_MARKERS = (
    "from PIL import Image",
    "ImageDraw",
    "Image.new(",
)
MIXED_SCRIPT_POLICY_MARKERS = (
    "replace_mixed_identifier_paragraph(",
    "set_paragraph_fonts(",
    "append_text_run(",
)
KEYWORD_RUN_POLICY_MARKERS = (
    "replace_english_keywords_paragraph(",
    "key words:",
    "keyword label",
    "keyword content",
)
IMAGE_INSERT_MARKERS = (
    "add_picture(",
    "InlineShapes.AddPicture",
)
PASS_SHAPED_ACCEPTANCE_WRITER_MARKERS = (
    "final-acceptance",
    "final_acceptance",
    "final-verification",
    "final_verification",
    "acceptance record",
    "validation result: pass",
    "handoff status: pass",
    "known caveats: none",
    "officecli validate: pass",
    "officecli validate passed",
    "rendered pdf pages",
    "rendered page evidence count",
    "required phrase checks",
    "old terms count",
    "media count",
)
SMOKE_ONLY_ACCEPTANCE_SIGNAL_MARKERS = (
    "officecli validate",
    "rendered pdf",
    "page images",
    "old terms",
    "required phrase",
    "media count",
    "image count",
    "phrase checks",
)


@dataclass(frozen=True)
class ScriptRisk:
    path: Path
    reasons: tuple[str, ...]

    def as_dict(self, project_root: Path) -> dict[str, object]:
        try:
            display_path = str(self.path.relative_to(project_root))
        except ValueError:
            display_path = str(self.path)
        return {"path": display_path, "reasons": list(self.reasons)}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def is_in_ignored_graduation_archive(path: Path, project_root: Path, active_run_dir: Path | None) -> bool:
    try:
        relative_parts = path.relative_to(project_root).parts
    except ValueError:
        return False
    if len(relative_parts) < 3:
        return False
    if relative_parts[0].lower() != ".codex":
        return False
    if relative_parts[1].lower() != "graduation-project-builder":
        return False
    if active_run_dir is None:
        return False
    try:
        path.relative_to(active_run_dir)
        return False
    except ValueError:
        return True


def iter_candidate_scripts(project_root: Path, active_run_dir: Path | None = None) -> list[Path]:
    candidates: list[Path] = []
    generated_run_stage = (project_root / "records" / "local-thesis-adapter.json").exists()
    for path in project_root.rglob("*"):
        if is_in_ignored_graduation_archive(path, project_root, active_run_dir):
            continue
        if generated_run_stage:
            try:
                relative_parts = path.relative_to(project_root).parts
            except ValueError:
                relative_parts = ()
            if relative_parts and relative_parts[0].lower() == "stage":
                continue
        lowered_parts = [part.lower() for part in path.parts]
        if any(part in SKIP_DIR_NAMES for part in lowered_parts):
            continue
        if any(part.startswith(prefix) for part in lowered_parts for prefix in SKIP_DIR_PREFIXES):
            continue
        if path.is_file() and path.suffix.lower() in SCRIPT_EXTENSIONS:
            candidates.append(path)
    return sorted(candidates)


def iter_candidate_adapter_manifests(project_root: Path, active_run_dir: Path | None = None) -> list[Path]:
    candidates: list[Path] = []
    for path in project_root.rglob("*"):
        if is_in_ignored_graduation_archive(path, project_root, active_run_dir):
            continue
        lowered_parts = [part.lower() for part in path.parts]
        if any(part in SKIP_DIR_NAMES for part in lowered_parts):
            continue
        if any(part.startswith(prefix) for part in lowered_parts for prefix in SKIP_DIR_PREFIXES):
            continue
        lowered_name = path.name.lower()
        if (
            path.is_file()
            and path.suffix.lower() in ADAPTER_EXTENSIONS
            and any(marker in lowered_name for marker in ADAPTER_NAME_MARKERS)
        ):
            candidates.append(path)
    return sorted(candidates)


def contains_any(text: str, tokens: tuple[str, ...]) -> bool:
    lowered = text.lower()
    return any(token.lower() in lowered for token in tokens)


def count_any(text: str, tokens: tuple[str, ...]) -> int:
    lowered = text.lower()
    return sum(1 for token in tokens if token.lower() in lowered)


def is_hardcoded_docx_paragraph_index(text: str) -> bool:
    return bool(re.search(r"\bdoc\.paragraphs\[\s*\d+\s*\]", text))


def is_thin_wrapper(text: str) -> bool:
    lowered = text.lower()
    delegates = (
        "delegated canonical helper" in lowered
        or "canonical skill" in lowered
        or ".agents\\skills\\graduation-project-builder\\scripts" in lowered
        or ".agents/skills/graduation-project-builder/scripts" in lowered
    )
    destructive = (
        contains_any(text, DESTRUCTIVE_PARAGRAPH_MARKERS)
        or is_hardcoded_docx_paragraph_index(text)
        or contains_any(text, POWERPOINT_WORD_COM_MARKERS)
    )
    return delegates and not destructive


def is_read_only_project_quality_gate(path: Path, text: str) -> bool:
    """Allow project quality gates that only inspect DOCX output and delegate acceptance.

    Some projects keep a local `scripts/run_quality_gate.py` that reads DOCX XML to
    triage OfficeCLI false positives before delegating the thesis acceptance decision
    to the canonical skill gate. That file is not a thesis rewrite helper and should
    not force a clean-source restart unless it contains a real write primitive.
    """
    if path.name != "run_quality_gate.py":
        return False
    lowered = text.lower()
    delegates_to_canonical_gate = (
        "validate_skill_gate.py" in lowered
        and "thesis_gate_record_path" in lowered
        and "officecli" in lowered
    )
    if not delegates_to_canonical_gate:
        return False
    forbidden_write_markers = (
        "paragraph.clear(",
        "clear_paragraph(",
        "set_paragraph_text(",
        "replace_paragraph_text(",
        "add_run(",
        "p.remove(child)",
        "._p.remove(",
        "writestr(",
        "writefilesync(",
        "doc.save(",
        "document(",
    )
    return not contains_any(lowered, forbidden_write_markers)


def writes_pass_shaped_acceptance_without_canonical_gate(text: str) -> bool:
    lowered = text.lower()
    writes_acceptance_named_file = (
        "final-acceptance" in lowered
        or "final_acceptance" in lowered
        or "final-verification" in lowered
        or "final_verification" in lowered
    )
    if not writes_acceptance_named_file:
        return False
    pass_shaped_count = count_any(text, PASS_SHAPED_ACCEPTANCE_WRITER_MARKERS)
    smoke_signal_count = count_any(text, SMOKE_ONLY_ACCEPTANCE_SIGNAL_MARKERS)
    delegates_to_canonical_gate = "validate_skill_gate.py" in lowered and "--gate-record" in lowered
    if delegates_to_canonical_gate:
        return False
    return pass_shaped_count >= 3 and smoke_signal_count >= 2


def classify_script(path: Path, text: str) -> ScriptRisk | None:
    if is_thin_wrapper(text):
        return None
    if is_read_only_project_quality_gate(path, text):
        return None

    reasons: list[str] = []
    docx_package_surface = contains_any(text, DOCX_PACKAGE_SURFACE_MARKERS)
    docx_engine = contains_any(text, DOCX_ENGINE_MARKERS)
    destructive_paragraph = contains_any(text, DESTRUCTIVE_PARAGRAPH_MARKERS)
    hardcoded_paragraph_index = is_hardcoded_docx_paragraph_index(text)
    protected_surface_count = count_any(text, PROTECTED_SURFACE_MARKERS)
    power_com = contains_any(text, POWERPOINT_WORD_COM_MARKERS)
    libreoffice_index = count_any(text, LIBREOFFICE_INDEX_MARKERS) >= 2
    zip_xml_surface_access = docx_package_surface and count_any(text, ZIP_XML_REWRITE_MARKERS) >= 2
    zip_xml_write_primitive = contains_any(text, ZIP_WRITE_MARKERS)
    zip_xml_rewrite = zip_xml_surface_access and zip_xml_write_primitive
    local_drawing = contains_any(text, LOCAL_DRAWING_MARKERS)
    image_insert = contains_any(text, IMAGE_INSERT_MARKERS)
    mixed_script_policy = count_any(text, MIXED_SCRIPT_POLICY_MARKERS) >= 2
    keyword_run_policy = count_any(text, KEYWORD_RUN_POLICY_MARKERS) >= 2

    if docx_engine and destructive_paragraph:
        reasons.append("local script clears/rebuilds DOCX paragraphs")
    if hardcoded_paragraph_index:
        reasons.append("local script mutates hardcoded doc.paragraphs[n] indexes")
    if docx_engine and protected_surface_count >= 3 and destructive_paragraph:
        reasons.append("local script owns multiple protected thesis surfaces")
    if power_com and protected_surface_count >= 2:
        reasons.append("local PowerShell/COM script owns protected thesis formatting")
    if libreoffice_index:
        reasons.append("local LibreOffice/UNO script refreshes DOCX indexes, fields, or stores outside canonical renderer")
    if zip_xml_rewrite and (protected_surface_count >= 1 or ".docx" in text.lower()):
        reasons.append("local script rewrites DOCX package XML outside canonical helpers")
    if local_drawing and image_insert:
        reasons.append("local script generates and inserts thesis figures outside canonical draw.io route")
    if docx_engine and "tblBorders" in text and destructive_paragraph:
        reasons.append("local script reimplements table/caption formatting policy")
    if docx_engine and mixed_script_policy:
        reasons.append("local script owns mixed-script body-text run splitting or font policy")
    if docx_engine and keyword_run_policy:
        reasons.append("local script owns abstract/keyword label-content run policy")
    if docx_engine and "ParagraphFormat.LineSpacingRule" in text:
        reasons.append("local script controls paragraph line spacing outside canonical helpers")
    if writes_pass_shaped_acceptance_without_canonical_gate(text):
        reasons.append("local script writes pass-shaped thesis acceptance/verification without canonical validate_skill_gate --gate-record")

    if reasons:
        return ScriptRisk(path=path.resolve(), reasons=tuple(dict.fromkeys(reasons)))
    return None


def scan_project_local_helper_scripts(project_root: Path, active_run_dir: Path | None = None) -> list[ScriptRisk]:
    root = project_root.resolve()
    active = active_run_dir.resolve() if active_run_dir else None
    risks: list[ScriptRisk] = []
    if not root.exists():
        return risks
    for path in iter_candidate_scripts(root, active_run_dir=active):
        try:
            text = read_text(path)
        except OSError:
            continue
        risk = classify_script(path, text)
        if risk is not None:
            risks.append(risk)
    skill_root = Path(__file__).resolve().parents[1]
    for path in iter_candidate_adapter_manifests(root, active_run_dir=active):
        issues = validate_adapter_file(path, skill_root=skill_root)
        if issues:
            risks.append(ScriptRisk(path=path.resolve(), reasons=tuple(issues)))
    return risks


def risk_summary(risks: list[ScriptRisk]) -> str:
    if risks:
        return "failed project-local thick thesis rewrite scripts detected"
    return "passed clean project-local thesis helper scanner"


def render_report(
    project_root: Path,
    risks: list[ScriptRisk],
    *,
    active_run_dir: Path | None = None,
    command: str = "",
    exit_status: int | None = None,
) -> str:
    resolved_root = project_root.resolve()
    generated_at = datetime.now(timezone.utc)
    if exit_status is None:
        exit_status = 2 if risks else 0
    lines = [
        "# Project-Local Thesis Helper Preflight",
        "",
        "- report schema: graduation-project-builder.project-local-helper-preflight.v2",
        f"- generated at UTC: {generated_at.isoformat().replace('+00:00', 'Z')}",
        f"- generated at unix: {generated_at.timestamp():.6f}",
        f"- project root: {resolved_root}",
        f"- active run dir: {active_run_dir.resolve() if active_run_dir else 'not specified'}",
        f"- summary: {risk_summary(risks)}",
        f"- risky script count: {len(risks)}",
        "- scanner: scripts/scan_project_local_thesis_helpers.py",
        f"- scanner command: {command or 'scan_project_local_thesis_helpers.py --project-root ' + str(resolved_root)}",
        f"- scanner exit status: {exit_status}",
        "",
        "## Risky Scripts And Adapters",
    ]
    if not risks:
        lines.append("- none")
    else:
        for risk in risks:
            try:
                display_path = risk.path.relative_to(project_root.resolve())
            except ValueError:
                display_path = risk.path
            lines.append(f"- {display_path}")
            for reason in risk.reasons:
                lines.append(f"  - {reason}")
    lines.append("")
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", required=True, help="Project directory to scan")
    parser.add_argument("--report-out", help="Optional markdown report path")
    parser.add_argument(
        "--active-run-dir",
        help="Optional current graduation-project-builder run directory; sibling historical run directories under .codex/graduation-project-builder are treated as archives.",
    )
    parser.add_argument("--json", action="store_true", help="Print JSON instead of markdown")
    parser.add_argument("--fail-on-risk", action="store_true", help="Return non-zero when risky scripts are found")
    args = parser.parse_args()

    project_root = Path(args.project_root).resolve()
    active_run_dir = Path(args.active_run_dir).resolve() if args.active_run_dir else None
    risks = scan_project_local_helper_scripts(project_root, active_run_dir=active_run_dir)
    exit_status = 2 if args.fail_on_risk and risks else 0
    command = " ".join(
        shlex.quote(part)
        for part in [
            "scan_project_local_thesis_helpers.py",
            "--project-root",
            str(project_root),
            *(["--fail-on-risk"] if args.fail_on_risk else []),
            *(["--active-run-dir", str(active_run_dir)] if active_run_dir else []),
            *(["--json"] if args.json else []),
            *(["--report-out", str(Path(args.report_out).resolve())] if args.report_out else []),
        ]
    )
    if args.report_out:
        report_path = Path(args.report_out)
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(
            render_report(project_root, risks, active_run_dir=active_run_dir, command=command, exit_status=exit_status),
            encoding="utf-8",
        )

    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    if args.json:
        generated_at = datetime.now(timezone.utc)
        print(json.dumps(
            {
                "schema": "graduation-project-builder.project-local-helper-preflight.v2",
                "generated_at_utc": generated_at.isoformat().replace("+00:00", "Z"),
                "generated_at_unix": generated_at.timestamp(),
                "project_root": str(project_root),
                "active_run_dir": str(active_run_dir) if active_run_dir else None,
                "summary": risk_summary(risks),
                "risky_script_count": len(risks),
                "scanner": "scripts/scan_project_local_thesis_helpers.py",
                "scanner_command": command,
                "scanner_exit_status": exit_status,
                "risks": [risk.as_dict(project_root) for risk in risks],
            },
            ensure_ascii=True,
            indent=2,
        ))
    else:
        print(render_report(project_root, risks, active_run_dir=active_run_dir, command=command, exit_status=exit_status))
    return exit_status


if __name__ == "__main__":
    raise SystemExit(main())
