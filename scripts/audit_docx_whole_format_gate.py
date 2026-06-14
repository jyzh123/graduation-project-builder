#!/usr/bin/env python3
"""Fail-closed structural format audit for whole-thesis DOCX handoff."""

from __future__ import annotations

import argparse
import base64
import hashlib
import json
import re
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from audit_docx_font_encoding import (
        CHINESE_FONT_SIZE_HALF_POINTS,
        bibliography_font_slot_hits,
        collect_bibliography_entries,
        validate_wps_named_size_evidence,
    )
except Exception:  # pragma: no cover - optional only for isolated fixture use
    CHINESE_FONT_SIZE_HALF_POINTS = {}  # type: ignore[assignment]
    bibliography_font_slot_hits = None  # type: ignore[assignment]
    collect_bibliography_entries = None  # type: ignore[assignment]
    validate_wps_named_size_evidence = None  # type: ignore[assignment]

try:
    from audit_docx_list_pollution import audit_docx_list_pollution
except Exception:  # pragma: no cover - optional only for isolated fixture use
    audit_docx_list_pollution = None  # type: ignore[assignment]


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
W = f"{{{W_NS}}}"
R = f"{{{R_NS}}}"
NS = {"w": W_NS, "r": R_NS, "a": A_NS, "v": V_NS}
PDF_BLANK_INK_RATIO_MAX = 0.00005
PDF_NEAR_EMPTY_INK_RATIO_MAX = 0.006


def qn(local: str) -> str:
    return f"{W}{local}"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().lower()


def audit_rendered_pdf_page_occupancy(
    pdf_path: Path | None,
    *,
    allowed_near_empty_pages: set[int] | None = None,
) -> dict[str, object]:
    if pdf_path is None:
        return {"provided": False, "passed": True, "issues": []}
    allowed = set(allowed_near_empty_pages or set())
    issues: list[str] = []
    if not pdf_path.exists():
        return {
            "provided": True,
            "path": str(pdf_path),
            "passed": False,
            "issues": [f"rendered PDF not found: {pdf_path}"],
        }
    try:
        import fitz  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency guard
        return {
            "provided": True,
            "path": str(pdf_path),
            "passed": False,
            "issues": [f"PyMuPDF/fitz unavailable for rendered PDF occupancy audit: {exc}"],
        }

    page_metrics: list[dict[str, object]] = []
    blank_pages: list[int] = []
    near_empty_pages: list[int] = []
    with fitz.open(str(pdf_path)) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(1.0, 1.0), alpha=False)
            samples = pix.samples
            components = max(1, pix.n)
            ink_pixels = 0
            for offset in range(0, len(samples), components):
                rgb = samples[offset : offset + min(3, components)]
                if any(channel < 245 for channel in rgb):
                    ink_pixels += 1
            total_pixels = max(1, pix.width * pix.height)
            ratio = round(ink_pixels / total_pixels, 8)
            text = normalize_text(page.get_text("text") or "")
            blank = ratio <= PDF_BLANK_INK_RATIO_MAX
            near_empty = ratio <= PDF_NEAR_EMPTY_INK_RATIO_MAX
            metric = {
                "page": page_index,
                "width": pix.width,
                "height": pix.height,
                "ink_pixels": ink_pixels,
                "ink_ratio": ratio,
                "text_length": len(text),
                "text_preview": text[:80],
                "image_count": len(page.get_images(full=True)),
                "drawing_count": len(page.get_drawings()),
                "blank": blank,
                "near_empty": near_empty,
                "allowed": page_index in allowed,
            }
            page_metrics.append(metric)
            if blank:
                blank_pages.append(page_index)
            elif near_empty:
                near_empty_pages.append(page_index)

    unexpected_blank_pages = [page for page in blank_pages if page not in allowed]
    unexpected_near_empty_pages = [page for page in near_empty_pages if page not in allowed]
    if unexpected_blank_pages:
        issues.append(f"rendered PDF contains unexpected blank pages: {unexpected_blank_pages}")
    if unexpected_near_empty_pages:
        issues.append(f"rendered PDF contains unexpected near-empty pages: {unexpected_near_empty_pages}")
    return {
        "provided": True,
        "path": str(pdf_path),
        "sha256": sha256_file(pdf_path),
        "page_count": len(page_metrics),
        "blank_ink_ratio_max": PDF_BLANK_INK_RATIO_MAX,
        "near_empty_ink_ratio_max": PDF_NEAR_EMPTY_INK_RATIO_MAX,
        "allowed_near_empty_pages": sorted(allowed),
        "blank_pages": blank_pages,
        "near_empty_pages": near_empty_pages,
        "unexpected_blank_pages": unexpected_blank_pages,
        "unexpected_near_empty_pages": unexpected_near_empty_pages,
        "page_metrics": page_metrics,
        "passed": not issues,
        "issues": issues,
    }


def text_of(element: ET.Element) -> str:
    return "".join(node.text or "" for node in element.findall(".//w:t", NS))


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def attr(element: ET.Element | None, local: str = "val") -> str:
    if element is None:
        return ""
    return element.attrib.get(qn(local), "")


def compact(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", text or "").strip().lower()


def paragraph_style_id(paragraph: ET.Element) -> str:
    return attr(paragraph.find("./w:pPr/w:pStyle", NS))


def style_ppr_summary(style: ET.Element) -> dict[str, object]:
    ppr = style.find("./w:pPr", NS)
    ind = ppr.find("./w:ind", NS) if ppr is not None else None
    spacing = ppr.find("./w:spacing", NS) if ppr is not None else None
    jc = ppr.find("./w:jc", NS) if ppr is not None else None
    return {
        "style_id": attr(style, "styleId"),
        "name": attr(style.find("./w:name", NS)),
        "based_on": attr(style.find("./w:basedOn", NS)),
        "alignment": attr(jc),
        "indent": {
            "left": attr(ind, "left"),
            "right": attr(ind, "right"),
            "leftChars": attr(ind, "leftChars"),
            "rightChars": attr(ind, "rightChars"),
            "firstLine": attr(ind, "firstLine"),
            "firstLineChars": attr(ind, "firstLineChars"),
            "hanging": attr(ind, "hanging"),
            "hangingChars": attr(ind, "hangingChars"),
        },
        "spacing": {
            "before": attr(spacing, "before"),
            "after": attr(spacing, "after"),
            "line": attr(spacing, "line"),
            "lineRule": attr(spacing, "lineRule"),
        },
        "numPr": ppr.find("./w:numPr", NS) is not None if ppr is not None else False,
    }


def style_record_map(styles_root: ET.Element | None) -> dict[str, dict[str, object]]:
    if styles_root is None:
        return {}
    records: dict[str, dict[str, object]] = {}
    for style in styles_root.findall("./w:style", NS):
        if attr(style, "type") != "paragraph":
            continue
        style_id = attr(style, "styleId")
        if style_id:
            records[style_id] = style_ppr_summary(style)
    return records


def default_paragraph_style_id(styles_root: ET.Element | None, records: dict[str, dict[str, object]]) -> str:
    if styles_root is not None:
        for style in styles_root.findall("./w:style", NS):
            if attr(style, "type") == "paragraph" and attr(style, "default").lower() in {"1", "true"}:
                style_id = attr(style, "styleId")
                if style_id:
                    return style_id
    for style_id, record in records.items():
        if str(record.get("name", "")).lower() == "normal":
            return style_id
    return ""


def paragraph_style_chain(style_id: str, records: dict[str, dict[str, object]]) -> list[dict[str, object]]:
    chain: list[dict[str, object]] = []
    seen: set[str] = set()
    current = style_id
    while current and current in records and current not in seen:
        seen.add(current)
        record = records[current]
        chain.append(record)
        current = str(record.get("based_on", "") or "")
    return chain


def inherited_value(
    direct: dict[str, object],
    chain: list[dict[str, object]],
    key: str,
) -> str:
    value = str(direct.get(key, "") or "")
    if value != "":
        return value
    for record in chain:
        indent = record.get("indent") if isinstance(record.get("indent"), dict) else {}
        value = str(indent.get(key, "") or "")
        if value != "":
            return value
    return ""


def effective_indent_summary(
    direct: dict[str, object],
    chain: list[dict[str, object]],
) -> dict[str, str]:
    keys = ("left", "right", "leftChars", "rightChars", "firstLine", "firstLineChars", "hanging", "hangingChars")
    return {key: inherited_value(direct, chain, key) for key in keys}


def has_page_break(paragraph: ET.Element) -> bool:
    return any(attr(br, "type") == "page" for br in paragraph.findall(".//w:br", NS))


def has_page_break_before(paragraph: ET.Element) -> bool:
    return paragraph.find("./w:pPr/w:pageBreakBefore", NS) is not None


def has_section_break(paragraph: ET.Element) -> bool:
    return paragraph.find("./w:pPr/w:sectPr", NS) is not None


def field_char_types(paragraph: ET.Element) -> list[str]:
    return [
        attr(node, "fldCharType")
        for node in paragraph.findall(".//w:fldChar", NS)
        if attr(node, "fldCharType")
    ]


def has_field_surface(paragraph: ET.Element) -> bool:
    return bool(paragraph.findall(".//w:fldChar", NS) or paragraph.findall(".//w:instrText", NS))


def paragraph_toc_instr(paragraph: ET.Element) -> str:
    pieces = [
        *(node.text or "" for node in paragraph.findall(".//w:instrText", NS)),
        *(attr(node, "instr") for node in paragraph.findall(".//w:fldSimple", NS)),
    ]
    return " ".join(item for item in pieces if item).strip()


def has_toc_instruction(paragraph: ET.Element) -> bool:
    return bool(re.search(r"(^|\s)TOC(\s|$)", paragraph_toc_instr(paragraph), flags=re.IGNORECASE))


def inspect_live_toc_fields(root: ET.Element) -> dict[str, int]:
    count = 0
    locked_count = 0
    field_stack: list[dict[str, object]] = []
    for node in root.iter():
        if node.tag == f"{W}fldSimple":
            instr = node.attrib.get(f"{W}instr", "")
            if re.search(r"(^|\s)TOC(\s|$)", instr, re.IGNORECASE):
                count += 1
                if node.attrib.get(f"{W}fldLock", "").lower() == "true":
                    locked_count += 1
            continue
        if node.tag == f"{W}fldChar":
            field_type = node.attrib.get(f"{W}fldCharType", "")
            if field_type == "begin":
                field_stack.append(
                    {
                        "instr": "",
                        "locked": node.attrib.get(f"{W}fldLock", "").lower() == "true",
                    }
                )
            elif field_type == "end" and field_stack:
                field = field_stack.pop()
                instr = str(field.get("instr", ""))
                if re.search(r"(^|\s)TOC(\s|$)", instr, re.IGNORECASE):
                    count += 1
                    if bool(field.get("locked")):
                        locked_count += 1
            continue
        if node.tag == f"{W}instrText" and field_stack:
            field_stack[-1]["instr"] = str(field_stack[-1].get("instr", "")) + (node.text or "")
    return {"count": count, "locked_count": locked_count}


def previous_visible_paragraph_index(texts: list[str], index: int) -> int | None:
    for previous in range(index - 1, 0, -1):
        if normalize_text(texts[previous - 1]):
            return previous
    return None


def opens_on_new_page(paragraphs: list[ET.Element], texts: list[str], index: int) -> bool:
    if index <= 0 or index > len(paragraphs):
        return False
    paragraph = paragraphs[index - 1]
    if has_page_break_before(paragraph) or has_section_break(paragraph):
        return True
    previous = previous_visible_paragraph_index(texts, index)
    if previous is None:
        return False
    for separator_index in range(previous + 1, index):
        separator = paragraphs[separator_index - 1]
        if has_page_break(separator) or has_page_break_before(separator) or has_section_break(separator):
            return True
    previous_paragraph = paragraphs[previous - 1]
    return has_page_break(previous_paragraph) or has_section_break(previous_paragraph)


def has_tab(paragraph: ET.Element) -> bool:
    return paragraph.find(".//w:tab", NS) is not None


def has_dotted_tab(paragraph: ET.Element) -> bool:
    return paragraph.find(".//w:tabs/w:tab[@w:leader='dot']", NS) is not None


def paragraph_tabs(paragraph: ET.Element) -> list[dict[str, str]]:
    return [
        {
            "val": attr(tab),
            "pos": attr(tab, "pos"),
            "leader": attr(tab, "leader"),
        }
        for tab in paragraph.findall(".//w:tabs/w:tab", NS)
    ]


def has_right_tab(paragraph: ET.Element) -> bool:
    return any(item.get("val") in {"right", "end"} and item.get("pos") for item in paragraph_tabs(paragraph))


def run_font_sizes(paragraph: ET.Element) -> list[int]:
    sizes: list[int] = []
    for size in paragraph.findall(".//w:rPr/w:sz", NS):
        raw = attr(size)
        if re.fullmatch(r"\d+", raw or ""):
            sizes.append(int(raw))
    return sizes


def paragraph_font_size_summary(paragraph: ET.Element) -> dict[str, object]:
    sizes = run_font_sizes(paragraph)
    return {
        "has_explicit_size": bool(sizes),
        "sizes_half_points": sorted(set(sizes)),
        "min_half_points": min(sizes) if sizes else None,
        "max_half_points": max(sizes) if sizes else None,
    }


def paragraph_by_body_child_index(document_root: ET.Element) -> dict[int, ET.Element]:
    body = document_root.find("./w:body", NS)
    if body is None:
        return {}
    return {
        index: child
        for index, child in enumerate(list(body))
        if child.tag == qn("p")
    }


def paragraph_max_font_size(paragraph: ET.Element) -> int:
    sizes = run_font_sizes(paragraph)
    return max(sizes) if sizes else 0


def paragraph_has_right_dot_tab(paragraph: ET.Element) -> bool:
    return any(
        item.get("val") in {"right", "end"}
        and item.get("leader") == "dot"
        and item.get("pos")
        for item in paragraph_tabs(paragraph)
    )


def is_toc_style(style_id: str, style_name: str = "") -> bool:
    tokens = {style_id.strip().lower(), style_name.strip().lower()}
    return any(token.startswith("toc") or token.startswith("toc ") for token in tokens if token)


PLACEHOLDER_TEXT_PATTERNS = [
    re.compile(r"(?i)\b(?:x{2,}|todo|tbd|placeholder)\b"),
    re.compile(r"[\u00d7\u25a1_]{2,}"),
    re.compile(r"[\[\(\uff08\u3010<]\s*(?:\u59d3\u540d|\u5b66\u53f7|\u4e13\u4e1a|\u73ed\u7ea7|\u5b66\u9662|\u6307\u5bfc\u6559\u5e08|\u9898\u76ee)\s*[\]\)\uff09\u3011>]"),
    re.compile(r"(?:\u8bf7\u586b\u5199|\u5f85\u586b\u5199|\u5f85\u5b9a|\u8bf7\u8f93\u5165|\u8bf7\u5220\u9664\u672c\u884c|\u5982\u9898\u76ee\u53ea\u6709\u4e00\u884c)"),
    re.compile(r"(?:\u4e2d\u6587\u9898\u76ee|\u82f1\u6587\u9898\u76ee|\u5916\u6587\u9898\u76ee)"),
]
# The whole-format gate is a broad structural detector, not the owner of a
# school-specific footer point size. Report-equivalent audits enforce the exact
# requested size when an external checker says "五号" or "小五"; this gate accepts
# the two recurring Chinese thesis footer sizes to avoid false failures.
EXPECTED_FOOTER_PAGE_SIZE_HALF_POINTS = {18, 21}
ACKNOWLEDGEMENT_TITLE_EXACT = {
    "\u81f4\u8c22",
    "\u81f4 \u8c22",
    "\u81f4\u3000\u8c22",
    "\u8c22\u8f9e",
    "Acknowledgements",
    "Acknowledgments",
}
HEADER_FOOTER_BODY_LEAK_PATTERNS = [
    re.compile(r"^\s*\u7b2c\s*[0-9\u4e00-\u9fa5]+\s*\u7ae0"),
    re.compile(r"^\s*[1-9]\d*(?:\.\d+){1,3}\s+\S"),
    re.compile(r"\bTOC_PLACEHOLDER\b", re.IGNORECASE),
]


def has_unresolved_placeholder_text(text: str) -> bool:
    value = normalize_text(text)
    if not value:
        return False
    return any(pattern.search(value) for pattern in PLACEHOLDER_TEXT_PATTERNS)


def cover_field_key(text: str) -> str:
    value = re.sub(r"[:\uff1a\s\u3000]+", "", text or "")
    fields = {
        "title": ("\u8bba\u6587\u9898\u76ee", "\u9898\u76ee"),
        "college": ("\u5b66\u9662", "\u9662\u7cfb", "\u7cfb\u90e8"),
        "major": ("\u4e13\u4e1a",),
        "class": ("\u73ed\u7ea7",),
        "student_id": ("\u5b66\u53f7",),
        "student_name": ("\u59d3\u540d", "\u5b66\u751f\u59d3\u540d"),
        "advisor": ("\u6307\u5bfc\u6559\u5e08", "\u5bfc\u5e08"),
    }
    for key, labels in fields.items():
        if any(label in value for label in labels):
            return key
    return ""


def is_unresolved_cover_value(text: str) -> bool:
    value = normalize_text(text)
    compact_value = compact(value)
    if not compact_value:
        return True
    if has_unresolved_placeholder_text(value):
        return True
    return compact_value in {
        "\u59d3\u540d",
        "\u5b66\u53f7",
        "\u4e13\u4e1a",
        "\u73ed\u7ea7",
        "\u5b66\u9662",
        "\u6307\u5bfc\u6559\u5e08",
        "\u8bba\u6587\u9898\u76ee",
        "\u9898\u76ee",
    }


def has_cover_title_label(text: str) -> bool:
    return bool(re.search(r"(?:^|[\s\u3000])(?:\u8bba\s*\u6587\s*)?\u9898\s*\u76ee\s*[:\uff1a]", text or ""))


def cover_placeholder_contract(zf: zipfile.ZipFile) -> dict[str, object]:
    root = load_xml(zf, "word/document.xml")
    if root is None:
        return {"passed": False, "issues": ["word/document.xml cannot be parsed for cover placeholder audit"], "rows": []}
    body = root.find("./w:body", NS)
    if body is None:
        return {"passed": False, "issues": ["word/document.xml has no body for cover placeholder audit"], "rows": []}
    rows: list[dict[str, object]] = []
    issues: list[str] = []
    title_label_rows: list[dict[str, object]] = []
    top_level_paragraph_index = 0
    for child in list(body):
        child_text = text_of(child).strip()
        if child.tag == qn("p"):
            top_level_paragraph_index += 1
            if is_zh_abstract(child_text) or is_toc_title(child_text) or is_body_heading(child_text):
                break
            if has_cover_title_label(child_text):
                title_label_rows.append(
                    {
                        "kind": "cover_title_label_paragraph",
                        "paragraph_index": top_level_paragraph_index,
                        "text": normalize_text(child_text)[:160],
                    }
                )
            if has_unresolved_placeholder_text(child_text):
                rows.append(
                    {
                        "kind": "cover_paragraph_placeholder",
                        "paragraph_index": top_level_paragraph_index,
                        "text": normalize_text(child_text)[:120],
                    }
                )
                issues.append(f"cover paragraph {top_level_paragraph_index} still contains template placeholder/instruction text")
            continue
        if child.tag != qn("tbl"):
            continue
        table_index = sum(1 for row in rows if row.get("kind") == "cover_table_value") + 1
        for row_index, table_row in enumerate(child.findall("./w:tr", NS), start=1):
            cells = table_row.findall("./w:tc", NS)
            if len(cells) < 2:
                continue
            for cell_index, label_cell in enumerate(cells[:-1]):
                label = normalize_text(text_of(label_cell))
                key = cover_field_key(label)
                if has_cover_title_label(label):
                    title_label_rows.append(
                        {
                            "kind": "cover_title_label_cell",
                            "table_index": table_index,
                            "row_index": row_index,
                            "cell_index": cell_index + 1,
                            "text": label[:160],
                        }
                    )
                if not key:
                    continue
                value_cell = cells[cell_index + 1]
                value = normalize_text(text_of(value_cell))
                unresolved = is_unresolved_cover_value(value)
                rows.append(
                    {
                        "kind": "cover_table_value",
                        "table_index": table_index,
                        "row_index": row_index,
                        "label": label,
                        "field": key,
                        "value": value[:120],
                        "unresolved": unresolved,
                    }
                )
                if unresolved:
                    issues.append(f"cover identity field {key} is blank or still contains a placeholder")
    if not title_label_rows:
        issues.append("cover title label `题目：` is missing")
    return {
        "passed": not issues,
        "issues": issues,
        "rows": rows,
        "title_label_rows": title_label_rows,
    }


def cover_layout_profile(zf: zipfile.ZipFile) -> dict[str, object]:
    root = load_xml(zf, "word/document.xml")
    if root is None:
        return {"available": False, "issues": ["word/document.xml cannot be parsed for cover layout profile"]}
    body = root.find("./w:body", NS)
    if body is None:
        return {"available": False, "issues": ["word/document.xml has no body for cover layout profile"]}
    children = list(body)

    def find_index(predicate) -> int | None:
        for index, child in enumerate(children):
            if child.tag == qn("p") and predicate(text_of(child), child):
                return index
        return None

    def starts_with(label: str):
        label_compact = compact(label)
        return lambda text, _child: compact(text).startswith(label_compact)

    def contains_label(label: str):
        return lambda text, _child: label in text

    def contains_all(*tokens: str):
        return lambda text, _child: all(token in compact(text) for token in tokens)

    declaration_index = find_index(contains_all("学位论文", "原创性声明"))
    section_break_indices = [
        index
        for index, child in enumerate(children)
        if child.tag == qn("p") and has_section_break(child)
    ]
    cover_section_end_index = None
    if declaration_index is not None:
        before_declaration = [index for index in section_break_indices if index < declaration_index]
        if before_declaration:
            cover_section_end_index = before_declaration[-1]
    date_index = None
    if declaration_index is not None:
        for index, child in enumerate(children[:declaration_index]):
            text = compact(text_of(child))
            if "年" in text and "月" in text and "日" in text and len(text) <= 24:
                date_index = index
    return {
        "available": True,
        "attachment8_prefix_present": (
            len(children) >= 2
            and compact(text_of(children[0])) == compact("附件8")
            and "河北北方学院学士学位论文模板" in text_of(children[1])
        ),
        "prefix_texts": [normalize_text(text_of(child)) for child in children[:2]],
        "first_logo_index": find_index(lambda _text, child: child.find(".//w:drawing", NS) is not None or child.find(".//w:pict", NS) is not None),
        "degree_title_index": find_index(lambda text, _child: "学士学位论文" in compact(text) and "设计" in compact(text)),
        "title_label_index": find_index(starts_with("题目：")),
        "name_index": find_index(contains_label("姓    名")),
        "student_id_index": find_index(contains_label("学    号")),
        "department_index": find_index(contains_label("院    系")),
        "major_index": find_index(contains_label("专    业")),
        "advisor_index": find_index(starts_with("指导教师")),
        "date_index": date_index,
        "declaration_index": declaration_index,
        "cover_section_end_index": cover_section_end_index,
        "section_break_indices": section_break_indices[:12],
    }


def cover_template_layout_contract(
    zf: zipfile.ZipFile,
    reference_layout: dict[str, object] | None,
) -> dict[str, object]:
    actual = cover_layout_profile(zf)
    issues: list[str] = []
    if not reference_layout or not reference_layout.get("available"):
        return {
            "passed": True,
            "required": False,
            "reason": "no reference cover layout profile",
            "actual": actual,
            "reference": reference_layout,
            "issues": issues,
        }
    required = bool(reference_layout.get("attachment8_prefix_present"))
    tolerated_index_drift: dict[str, str] = {}
    for key in (
        "title_label_index",
        "student_id_index",
        "major_index",
        "advisor_index",
    ):
        if reference_layout.get(key) is not None and actual.get(key) is None:
            issues.append(f"cover template field disappeared for {key}: reference={reference_layout.get(key)} actual=missing")
    if required:
        if actual.get("attachment8_prefix_present") is not True:
            issues.append("Attachment 8 cover prefix is missing or no longer matches the locked template")
        for key in (
            "first_logo_index",
            "degree_title_index",
            "title_label_index",
            "name_index",
            "student_id_index",
            "department_index",
            "major_index",
            "advisor_index",
            "date_index",
            "cover_section_end_index",
        ):
            if actual.get(key) != reference_layout.get(key):
                if (
                    key == "cover_section_end_index"
                    and isinstance(actual.get(key), int)
                    and isinstance(reference_layout.get(key), int)
                    and int(actual.get(key)) + 1 == int(reference_layout.get(key))
                ):
                    tolerated_index_drift[key] = (
                        "empty cover spacer section removed to prevent a rendered blank page after the cover"
                    )
                    continue
                issues.append(
                    f"cover Attachment 8 skeleton index mismatch for {key}: "
                    f"actual={actual.get(key)} reference={reference_layout.get(key)}"
                )
        if actual.get("declaration_index") is not None and actual.get("cover_section_end_index") is None:
            issues.append("cover section boundary is not before the declaration/front-matter page")
    return {
        "passed": not issues,
        "required": required,
        "actual": actual,
        "reference": reference_layout,
        "tolerated_index_drift": tolerated_index_drift,
        "issues": issues,
    }


def paragraph_format_summary(paragraph: ET.Element) -> dict[str, object]:
    ppr = paragraph.find("./w:pPr", NS)
    ind = ppr.find("./w:ind", NS) if ppr is not None else None
    spacing = ppr.find("./w:spacing", NS) if ppr is not None else None
    jc = ppr.find("./w:jc", NS) if ppr is not None else None
    return {
        "style_id": paragraph_style_id(paragraph),
        "alignment": attr(jc),
        "indent": {
            "left": attr(ind, "left"),
            "right": attr(ind, "right"),
            "leftChars": attr(ind, "leftChars"),
            "rightChars": attr(ind, "rightChars"),
            "firstLine": attr(ind, "firstLine"),
            "firstLineChars": attr(ind, "firstLineChars"),
            "hanging": attr(ind, "hanging"),
            "hangingChars": attr(ind, "hangingChars"),
        },
        "spacing": {
            "before": attr(spacing, "before"),
            "after": attr(spacing, "after"),
            "line": attr(spacing, "line"),
            "lineRule": attr(spacing, "lineRule"),
        },
        "tabs": paragraph_tabs(paragraph),
        "font_size": paragraph_font_size_summary(paragraph),
        "numPr": ppr.find("./w:numPr", NS) is not None if ppr is not None else False,
    }


def int_attr_value(value: str) -> int | None:
    if value == "":
        return None
    try:
        return int(value)
    except ValueError:
        return None


def has_direct_paragraph_metrics(summary: dict[str, object]) -> bool:
    indent = summary.get("indent") if isinstance(summary.get("indent"), dict) else {}
    spacing = summary.get("spacing") if isinstance(summary.get("spacing"), dict) else {}
    return bool(summary.get("alignment") or any(indent.values()) or any(spacing.values()))


def has_style_or_direct_format(summary: dict[str, object]) -> bool:
    return bool(summary.get("style_id")) or (
        has_direct_paragraph_metrics(summary)
        and bool(summary.get("font_size", {}).get("has_explicit_size"))
    )


def trailing_page_number(text: str) -> str:
    match = re.search(r"(\d+|[ivxlcdmIVXLCDM]+)\s*$", text or "")
    return match.group(1) if match else ""


def has_toc_page_number_text(text: str) -> bool:
    page = trailing_page_number(text)
    if not page:
        return False
    label = re.sub(rf"{re.escape(page)}\s*$", "", text or "").strip()
    if not label:
        return False
    if "\t" in (text or "") or "." in (text or "") or " " in (text or ""):
        return True
    # Fanyu/WPS TOC caches may flatten the right-tab run and page number into
    # one visible string, for example "摘要I" or "AbstractIII". Treat those as
    # TOC rows while the scan is inside the TOC block so they are not mistaken
    # for the real abstract title or first body heading.
    return bool(
        is_zh_abstract(label)
        or is_en_abstract(label)
        or is_body_heading(label)
        or is_reference_title(label)
        or is_acknowledgement_title(label)
    )


def toc_entry_has_visible_label(text: str) -> bool:
    page = trailing_page_number(text)
    label = re.sub(rf"{re.escape(page)}\s*$", "", text or "").strip() if page else (text or "").strip()
    return bool(label and not re.fullmatch(r"[\dIVXLCDMivxlcdm]+", label))


def toc_heading_match_key(text: str) -> str:
    value = normalize_text(text)
    value = value.replace("\t", " ")
    page_match = re.search(r"(\d+|[ivxlcdmIVXLCDM]+)\s*$", value or "")
    if page_match:
        page = page_match.group(1)
        prefix = value[: page_match.start(1)]
        separated_page = page.isdigit() or not prefix or bool(
            re.search(r"[\s\u3000\.\u2026\u00b7\u2022\u30fb]$", prefix)
        )
    else:
        page = ""
        separated_page = False
    if page and separated_page:
        value = re.sub(
            rf"[\s\u3000\.\u2026\u00b7\u2022\u30fb]*{re.escape(page)}\s*$",
            "",
            value,
        )
    value = value.replace("\uff0e", ".")
    value = re.sub(r"[\u2026\u00b7\u2022\u30fb]+$", "", value)
    value = re.sub(r"[\s\u3000\u25a1]+", "", value)
    return value.lower()


def body_heading_toc_coverage(
    body_heading_rows: list[dict[str, object]],
    toc_entry_rows: list[dict[str, object]],
) -> dict[str, object]:
    toc_keys = {
        toc_heading_match_key(str(row.get("text") or ""))
        for row in toc_entry_rows
        if toc_heading_match_key(str(row.get("text") or ""))
    }
    required_rows: list[dict[str, object]] = []
    missing_rows: list[dict[str, object]] = []
    for row in body_heading_rows:
        level = row.get("level")
        if not isinstance(level, int) or level < 1 or level > 3:
            continue
        text = str(row.get("text") or "")
        key = toc_heading_match_key(text)
        if not key:
            continue
        record = {
            "paragraph_index": row.get("paragraph_index"),
            "level": level,
            "text": normalize_text(text),
            "match_key": key,
        }
        required_rows.append(record)
        if key not in toc_keys:
            missing_rows.append(record)
    return {
        "passed": not missing_rows,
        "required_heading_count": len(required_rows),
        "toc_entry_count": len(toc_entry_rows),
        "required_headings": required_rows,
        "missing_headings": missing_rows,
    }


def is_reference_entry(text: str) -> bool:
    return bool(re.match(r"^\s*(?:\[\d+\]|\d+[\.\u3001])\s*\S", text or ""))


def is_toc_title(text: str) -> bool:
    return compact(text) in {"\u76ee\u5f55", "contents", "tableofcontents"}


def is_cover_date_text(text: str) -> bool:
    return bool(re.fullmatch(r"\d{4}\u5e74\d{1,2}\u6708", compact(text)))


def is_zh_abstract(text: str) -> bool:
    value = compact(re.sub(r"[:\uff1a].*$", "", text or ""))
    return value == "\u6458\u8981"


def is_en_abstract(text: str) -> bool:
    return compact(re.sub(r"[:\uff1a].*$", "", text or "")).startswith("abstract")


def is_zh_keyword(text: str) -> bool:
    return compact(text).startswith("\u5173\u952e\u8bcd")


def is_en_keyword(text: str) -> bool:
    return compact(text).startswith(("keywords", "keyword"))


def is_body_heading(text: str) -> bool:
    stripped = (text or "").strip()
    return bool(
        re.match(r"^\u7b2c\s*[0-9\u4e00-\u9fa5]+\s*\u7ae0", stripped)
        or re.match(r"^[1-9]\d*(?:\.\d+){0,3}\s+\S", stripped)
    )


def is_chapter_heading_text(text: str) -> bool:
    return bool(re.match(r"^\s*\u7b2c\s*[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*\u7ae0(?:\s+\S.*)?\s*$", text or ""))


def body_heading_level(text: str) -> int | None:
    stripped = (text or "").strip()
    if is_chapter_heading_text(stripped):
        return 1
    if re.match(r"^[1-9]\d*(?:\.\d+){3}\s+\S", stripped):
        return 4
    if re.match(r"^[1-9]\d*\.\d+\.\d+\s+\S", stripped):
        return 3
    if re.match(r"^[1-9]\d*\.\d+\s+\S", stripped):
        return 2
    return None


def collect_body_heading_format_rows(
    paragraphs: list[ET.Element],
    texts: list[str],
    styles_root: ET.Element | None = None,
) -> list[dict[str, object]]:
    indices = first_surface_indices(texts)
    start = indices["first_body"] or 1
    end = min(
        [item for item in (indices["references"], indices["appendix"], indices["acknowledgement"]) if item is not None and item > start],
        default=len(paragraphs) + 1,
    )
    rows: list[dict[str, object]] = []
    style_records = style_record_map(styles_root)
    default_style_id = default_paragraph_style_id(styles_root, style_records)
    for index in range(start, end):
        text = texts[index - 1]
        level = body_heading_level(text)
        if level is None:
            continue
        summary = paragraph_format_summary(paragraphs[index - 1])
        style_id = str(summary.get("style_id", "") or default_style_id)
        style_chain = paragraph_style_chain(style_id, style_records)
        direct_indent = summary.get("indent") if isinstance(summary.get("indent"), dict) else {}
        rows.append(
            {
                "paragraph_index": index,
                "level": level,
                "text": normalize_text(text),
                "style_id": summary.get("style_id", ""),
                "alignment": summary.get("alignment", ""),
                "indent": direct_indent,
                "effective_indent": effective_indent_summary(direct_indent, style_chain),
                "style_chain": style_chain,
                "numPr": summary.get("numPr") is True,
                "spacing": summary.get("spacing", {}),
                "font_size": summary.get("font_size", {}),
            }
        )
    return rows


def body_heading_direct_format_issues(rows: list[dict[str, object]]) -> list[str]:
    issues: list[str] = []
    for row in rows:
        level = row.get("level")
        indent = row.get("indent") if isinstance(row.get("indent"), dict) else {}
        effective_indent = row.get("effective_indent") if isinstance(row.get("effective_indent"), dict) else indent
        spacing = row.get("spacing") if isinstance(row.get("spacing"), dict) else {}
        left = int_attr_value(str(effective_indent.get("left", "")))
        right = int_attr_value(str(effective_indent.get("right", "")))
        left_chars = int_attr_value(str(effective_indent.get("leftChars", "")))
        right_chars = int_attr_value(str(effective_indent.get("rightChars", "")))
        first_line = int_attr_value(str(effective_indent.get("firstLine", "")))
        first_line_chars = int_attr_value(str(effective_indent.get("firstLineChars", "")))
        hanging = int_attr_value(str(effective_indent.get("hanging", "")))
        hanging_chars = int_attr_value(str(effective_indent.get("hangingChars", "")))
        alignment = str(row.get("alignment", ""))
        text = str(row.get("text", ""))
        paragraph_index = row.get("paragraph_index")
        style_chain = row.get("style_chain") if isinstance(row.get("style_chain"), list) else []
        style_chain_summary = [
            {
                "style_id": str(item.get("style_id", "")),
                "name": str(item.get("name", "")),
                "indent": item.get("indent", {}),
            }
            for item in style_chain[:4]
            if isinstance(item, dict)
        ]
        for key, value in (
            ("left", left),
            ("right", right),
            ("leftChars", left_chars),
            ("rightChars", right_chars),
            ("firstLine", first_line),
            ("firstLineChars", first_line_chars),
            ("hanging", hanging),
            ("hangingChars", hanging_chars),
        ):
            if value not in (None, 0):
                issues.append(
                    f"body level-{level} heading carries nonzero effective {key} indent at paragraph {paragraph_index}: {text}; "
                    f"style_chain={style_chain_summary}"
                )
        if row.get("numPr") is True:
            issues.append(f"body level-{level} heading must not carry paragraph numbering/list state at paragraph {paragraph_index}: {text}")
        if level == 1:
            if left is not None and abs(left) > 500:
                issues.append(f"body level-1 heading direct left indent is too large at paragraph {row.get('paragraph_index')}: {text}")
            if right is not None and abs(right) > 500:
                issues.append(f"body level-1 heading direct right indent is too large at paragraph {row.get('paragraph_index')}: {text}")
            if first_line not in (None, 0):
                issues.append(f"body level-1 heading must not carry firstLine indent at paragraph {row.get('paragraph_index')}: {text}")
            if alignment == "left":
                issues.append(f"body level-1 heading alignment is left instead of template/body-heading baseline at paragraph {row.get('paragraph_index')}: {text}")
            if spacing.get("line") == "499":
                issues.append(f"body level-1 heading line spacing appears copied from front-matter title at paragraph {row.get('paragraph_index')}: {text}")
        if level in {2, 3, 4}:
            if left is not None and abs(left) > 1000:
                issues.append(f"body level-{level} heading direct left indent is too large at paragraph {row.get('paragraph_index')}: {text}")
            if right is not None and abs(right) > 1000:
                issues.append(f"body level-{level} heading direct right indent is too large at paragraph {row.get('paragraph_index')}: {text}")
    if not rows:
        issues.append("no real body heading rows were available for direct-format audit")
    return issues


def body_heading_template_baseline_issues(
    rows: list[dict[str, object]],
    reference_rows: list[dict[str, object]] | None,
) -> list[str]:
    if not reference_rows:
        return body_heading_direct_format_issues(rows)
    issues: list[str] = []
    reference_by_level: dict[int, dict[str, object]] = {}
    for reference_row in reference_rows:
        try:
            level = int(reference_row.get("level"))
        except (TypeError, ValueError):
            continue
        reference_by_level.setdefault(level, reference_row)

    indent_keys = (
        "left",
        "right",
        "leftChars",
        "rightChars",
        "firstLine",
        "firstLineChars",
        "hanging",
        "hangingChars",
    )
    if not rows:
        return ["no real body heading rows were available for template-baseline audit"]
    for row in rows:
        try:
            level = int(row.get("level"))
        except (TypeError, ValueError):
            issues.append(f"body heading row lacks integer level for template-baseline audit: {row}")
            continue
        text = str(row.get("text", ""))
        paragraph_index = row.get("paragraph_index")
        if row.get("numPr") is True:
            issues.append(f"body level-{level} heading must not carry paragraph numbering/list state at paragraph {paragraph_index}: {text}")
        reference_row = reference_by_level.get(level)
        if reference_row is None:
            issues.extend(body_heading_direct_format_issues([row]))
            continue
        actual_indent = row.get("effective_indent") if isinstance(row.get("effective_indent"), dict) else {}
        reference_indent = reference_row.get("effective_indent") if isinstance(reference_row.get("effective_indent"), dict) else {}
        for key in indent_keys:
            actual_value = str(actual_indent.get(key, "") or "")
            reference_value = str(reference_indent.get(key, "") or "")
            if (
                key == "left"
                and level in {2, 3}
                and actual_value == ""
                and reference_value in {"210", "420"}
            ):
                continue
            if actual_value != reference_value:
                issues.append(
                    f"body level-{level} heading effective {key} indent does not match template baseline at paragraph {paragraph_index}: "
                    f"actual={actual_value or 'blank'} reference={reference_value or 'blank'} text={text}"
                )
        actual_alignment = str(row.get("alignment", "") or "")
        reference_alignment = str(reference_row.get("alignment", "") or "")
        if level == 1 and actual_alignment == "center" and reference_alignment == "":
            continue
        if actual_alignment != reference_alignment:
            issues.append(
                f"body level-{level} heading alignment does not match template baseline at paragraph {paragraph_index}: "
                f"actual={actual_alignment or 'blank'} reference={reference_alignment or 'blank'} text={text}"
            )
    return issues


def reference_body_heading_format_contract(reference_docx: Path | None) -> dict[str, object] | None:
    if reference_docx is None:
        return None
    with zipfile.ZipFile(reference_docx) as zf:
        document_root = load_xml(zf, "word/document.xml")
        if document_root is None:
            raise ValueError(f"word/document.xml is missing or invalid in reference DOCX: {reference_docx}")
        styles_root = load_xml(zf, "word/styles.xml")
        paragraphs = visible_paragraphs(document_root)
        texts = [text_of(p).strip() for p in paragraphs]
    rows = collect_body_heading_format_rows(paragraphs, texts, styles_root)
    return {
        "reference_docx_path": str(reference_docx),
        "reference_docx_sha256": sha256_file(reference_docx),
        "sample_rows": rows[:12],
        "issues": [] if rows else ["reference DOCX has no real body heading rows"],
        "passed": bool(rows),
    }


def is_chinese_numeral_chapter_heading(text: str) -> bool:
    return bool(re.match(r"^\s*\u7b2c\s*[\u96f6\u3007\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u5343\u4e24]+\s*\u7ae0", text or ""))


def is_arabic_chapter_heading_format(text: str) -> bool:
    return bool(re.match(r"^\s*\u7b2c[1-9]\d*\u7ae0(?:[\s\u3000]{1,}|\t)\S", text or ""))


def reference_allows_chinese_chapter_headings(reference_docx: Path | None) -> bool:
    if reference_docx is None:
        return True
    try:
        with zipfile.ZipFile(reference_docx) as zf:
            document_root = load_xml(zf, "word/document.xml")
    except (OSError, zipfile.BadZipFile):
        return False
    if document_root is None:
        return False
    return any(
        is_chinese_numeral_chapter_heading(text_of(paragraph).strip())
        for paragraph in visible_paragraphs(document_root)
    )


def is_reference_title(text: str) -> bool:
    return compact(text) in {"\u53c2\u8003\u6587\u732e", "references", "bibliography"}


def is_acknowledgement_title(text: str) -> bool:
    return compact(text) in {"\u81f4\u8c22", "\u8c22\u8f9e", "acknowledgements", "acknowledgments"}


def is_appendix_title(text: str) -> bool:
    return compact(text) in {"\u9644\u5f55", "appendix"}


def is_strict_acknowledgement_title_text(text: str) -> bool:
    value = (text or "").strip()
    if value in ACKNOWLEDGEMENT_TITLE_EXACT:
        return True
    if compact(value) in {"\u81f4\u8c22", "\u8c22\u8f9e"}:
        return True
    return value.lower() in {"acknowledgements", "acknowledgments"}


def paragraph_has_num_pr(paragraph: ET.Element) -> bool:
    return paragraph.find("./w:pPr/w:numPr", NS) is not None


def load_xml(zf: zipfile.ZipFile, name: str) -> ET.Element | None:
    try:
        return ET.fromstring(zf.read(name))
    except (KeyError, ET.ParseError):
        return None


def style_name_map(styles_root: ET.Element | None) -> dict[str, str]:
    if styles_root is None:
        return {}
    result: dict[str, str] = {}
    for style in styles_root.findall("./w:style", NS):
        style_id = attr(style, "styleId")
        name = attr(style.find("./w:name", NS))
        if style_id:
            result[style_id] = name
    return result


def has_header_footer_body_leak(text: str) -> bool:
    value = normalize_text(text)
    return any(pattern.search(value) for pattern in HEADER_FOOTER_BODY_LEAK_PATTERNS)


def is_static_institutional_header_text(text: str) -> bool:
    value = compact(text)
    if not value:
        return True
    if has_header_footer_body_leak(text):
        return False
    if value.lower() in {"目录", "toc", "tableofcontents"}:
        return False
    return True


def visible_paragraphs(document_root: ET.Element) -> list[ET.Element]:
    body = document_root.find("./w:body", NS)
    if body is None:
        return []
    paragraphs: list[ET.Element] = []
    for child in list(body):
        if child.tag == qn("p"):
            paragraphs.append(child)
        elif child.tag == qn("sdt"):
            content = child.find(".//w:sdtContent", NS)
            if content is not None:
                paragraphs.extend(content.findall(".//w:p", NS))
    return paragraphs


def relationship_targets(root: ET.Element | None) -> dict[str, str]:
    if root is None:
        return {}
    targets: dict[str, str] = {}
    for rel in list(root):
        rid = rel.attrib.get("Id", "")
        target = rel.attrib.get("Target", "")
        if rid:
            targets[rid] = target
    return targets


def part_text_and_fields(zf: zipfile.ZipFile, part_name: str) -> tuple[str, str, int]:
    root = load_xml(zf, f"word/{part_name}")
    if root is None:
        return "", "", 0
    field_instructions = [
        *(node.text or "" for node in root.findall(".//w:instrText", NS)),
        *(node.attrib.get(W + "instr", "") for node in root.findall(".//w:fldSimple", NS)),
    ]
    return (
        "".join(node.text or "" for node in root.findall(".//w:t", NS)),
        " ".join(field_instructions),
        sum(1 for item in field_instructions if "PAGE" in item.upper()),
    )


def instr_text_of_run(run: ET.Element) -> str:
    return " ".join(node.text or "" for node in run.findall(".//w:instrText", NS))


def run_contains_page_field(run: ET.Element) -> bool:
    if "PAGE" in instr_text_of_run(run).upper():
        return True
    return any("PAGE" in node.attrib.get(W + "instr", "").upper() for node in run.findall(".//w:fldSimple", NS))


def run_size_values(run: ET.Element) -> list[int]:
    sizes: list[int] = []
    for size in run.findall(".//w:rPr/w:sz", NS):
        raw = attr(size)
        if re.fullmatch(r"\d+", raw or ""):
            sizes.append(int(raw))
    return sizes


def paragraph_page_field_runs(paragraph: ET.Element) -> list[ET.Element]:
    runs = paragraph.findall(".//w:r", NS)
    result: list[ET.Element] = [run for run in runs if run_contains_page_field(run)]
    if result:
        return result
    for field in paragraph.findall(".//w:fldSimple", NS):
        if "PAGE" not in field.attrib.get(W + "instr", "").upper():
            continue
        result.extend(field.findall(".//w:r", NS))
    return result


def footer_page_field_contract(zf: zipfile.ZipFile, footer_parts: list[str]) -> dict[str, object]:
    rows: list[dict[str, object]] = []
    issues: list[str] = []
    page_field_count = 0
    for part in footer_parts:
        root = load_xml(zf, f"word/{part}")
        if root is None:
            issues.append(f"footer part cannot be parsed: {part}")
            continue
        for paragraph_index, paragraph in enumerate(root.findall(".//w:p", NS), start=1):
            paragraph_page_runs = paragraph_page_field_runs(paragraph)
            if not paragraph_page_runs:
                continue
            page_field_count += len(paragraph_page_runs)
            sizes = sorted({size for run in paragraph_page_runs for size in run_size_values(run)})
            all_paragraph_sizes = sorted(set(run_font_sizes(paragraph)))
            row = {
                "part": part,
                "paragraph_index": paragraph_index,
                "page_field_run_count": len(paragraph_page_runs),
                "field_run_sizes_half_points": sizes,
                "paragraph_sizes_half_points": all_paragraph_sizes,
                "expected_sizes_half_points": sorted(EXPECTED_FOOTER_PAGE_SIZE_HALF_POINTS),
                "text": normalize_text(text_of(paragraph))[:80],
            }
            rows.append(row)
            if not sizes:
                row["font_size_policy"] = "inherited-page-field-run-format-accepted"
            elif not set(sizes).issubset(EXPECTED_FOOTER_PAGE_SIZE_HALF_POINTS):
                issues.append(f"footer PAGE field in {part} paragraph {paragraph_index} uses non-baseline font size(s): {sizes}")
    if page_field_count <= 0:
        issues.append("footer PAGE field is missing from all bound footer parts")
    return {
        "passed": not issues,
        "page_field_count": page_field_count,
        "expected_sizes_half_points": sorted(EXPECTED_FOOTER_PAGE_SIZE_HALF_POINTS),
        "rows": rows,
        "issues": issues,
    }


def header_footer_visual_contract(
    zf: zipfile.ZipFile,
    *,
    header_parts: list[str],
    footer_parts: list[str],
    require_header_rule_surface: bool = False,
) -> dict[str, object]:
    rows: list[dict[str, object]] = []
    issues: list[str] = []
    for kind, parts in (("header", header_parts), ("footer", footer_parts)):
        if not parts:
            issues.append(f"{kind} visual surface has no bound part")
            continue
        for part in parts:
            root = load_xml(zf, f"word/{part}")
            if root is None:
                issues.append(f"{kind} visual surface part cannot be parsed: {part}")
                rows.append({"kind": kind, "part": part, "parsed": False})
                continue
            text = normalize_text(text_of(root))
            fields = " ".join(
                [
                    *(node.text or "" for node in root.findall(".//w:instrText", NS)),
                    *(node.attrib.get(W + "instr", "") for node in root.findall(".//w:fldSimple", NS)),
                ]
            )
            drawing_count = len(root.findall(".//w:drawing", NS))
            pict_count = len(root.findall(".//w:pict", NS))
            vshape_count = len(root.findall(".//v:shape", NS))
            has_rule_surface = bool(drawing_count or pict_count or vshape_count)
            row = {
                "kind": kind,
                "part": part,
                "parsed": True,
                "text": text[:160],
                "field_instruction": normalize_text(fields)[:160],
                "has_page_field": "PAGE" in fields.upper(),
                "has_placeholder": has_unresolved_placeholder_text(text),
                "has_body_heading_leak": any(pattern.search(text) for pattern in HEADER_FOOTER_BODY_LEAK_PATTERNS),
                "drawing_count": drawing_count,
                "pict_count": pict_count,
                "vshape_count": vshape_count,
                "has_rule_surface": has_rule_surface,
            }
            rows.append(row)
            if kind == "header" and require_header_rule_surface and text and not has_rule_surface:
                issues.append(f"header visual surface lacks template rule/line drawing in {part}")
            if row["has_placeholder"]:
                issues.append(f"{kind} visual surface still contains placeholder text in {part}")
            if row["has_body_heading_leak"] and not (kind == "header" and has_rule_surface):
                issues.append(f"{kind} visual surface contains body/TOC helper text leakage in {part}")
    if footer_parts and not any(row.get("kind") == "footer" and row.get("has_page_field") for row in rows):
        issues.append("footer visual surface has no PAGE field in bound footer parts")
    return {"passed": not issues, "rows": rows, "issues": issues}


def iter_relationship_roots(zf: zipfile.ZipFile) -> list[ET.Element]:
    roots: list[ET.Element] = []
    for name in zf.namelist():
        if not name.startswith("word/_rels/") or not name.endswith(".rels"):
            continue
        try:
            roots.append(ET.fromstring(zf.read(name)))
        except ET.ParseError:
            continue
    return roots


def media_targets_for_relationships(relationship_roots: list[ET.Element]) -> set[str]:
    targets: set[str] = set()
    for root in relationship_roots:
        for rel in list(root):
            target = rel.attrib.get("Target", "")
            rel_type = rel.attrib.get("Type", "")
            if "image" not in rel_type and "/media/" not in target and not target.startswith("media/"):
                continue
            normalized = target.replace("\\", "/")
            if normalized.startswith("../"):
                normalized = normalized[3:]
            if not normalized.startswith("word/"):
                normalized = f"word/{normalized.lstrip('/')}"
            targets.add(normalized)
    return targets


def drawing_relationship_ids(element: ET.Element) -> set[str]:
    ids: set[str] = set()
    for blip in element.findall(".//a:blip", NS):
        embed = blip.attrib.get(R + "embed") or blip.attrib.get(R + "link")
        if embed:
            ids.add(embed)
    for image_data in element.findall(".//v:imagedata", NS):
        rid = image_data.attrib.get(R + "id")
        if rid:
            ids.add(rid)
    for node in element.iter():
        for key, value in node.attrib.items():
            if key.startswith(R) and value:
                ids.add(value)
    return ids


def media_parts_for_relationship_ids(rids: set[str], rels: dict[str, str]) -> list[str]:
    result: list[str] = []
    for rid in sorted(rids):
        target = rels.get(rid, "")
        if not target:
            continue
        normalized = target.replace("\\", "/")
        if normalized.startswith("../"):
            normalized = normalized[3:]
        if not normalized.startswith("word/"):
            normalized = f"word/{normalized.lstrip('/')}"
        if "/media/" not in normalized and not normalized.startswith("word/media/"):
            continue
        result.append(normalized)
    return result


def media_part_sha256(zf: zipfile.ZipFile, part_name: str) -> str:
    try:
        return hashlib.sha256(zf.read(part_name)).hexdigest()
    except KeyError:
        return ""


def image_paragraphs_in_zone(
    paragraphs: list[ET.Element],
    rels: dict[str, str],
    *,
    zone_end: int,
    zf: zipfile.ZipFile | None = None,
) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for index, paragraph in enumerate(paragraphs[: max(zone_end, 0)], start=1):
        relationship_ids = drawing_relationship_ids(paragraph)
        media_parts = media_parts_for_relationship_ids(relationship_ids, rels)
        if not media_parts:
            continue
        rows.append(
            {
                "paragraph_index": index,
                "relationship_ids": sorted(relationship_ids),
                "media_parts": media_parts,
                "media_sha256": {
                    part: media_part_sha256(zf, part) if zf is not None else ""
                    for part in media_parts
                },
            }
        )
    return rows


def cover_logo_contract(
    paragraphs: list[ET.Element],
    texts: list[str],
    rels: dict[str, str],
    relationship_roots: list[ET.Element],
    zf: zipfile.ZipFile,
    *,
    reference_logo: dict[str, object] | None = None,
    required: bool = True,
) -> dict[str, object]:
    indices = first_surface_indices(texts)
    front_end = min(
        [item for item in (indices["zh_abstract"], indices["toc"], indices["first_body"]) if item is not None],
        default=len(paragraphs) + 1,
    )
    image_rows = image_paragraphs_in_zone(paragraphs, rels, zone_end=front_end - 1, zf=zf)
    first_logo = image_rows[0] if image_rows else None
    all_media_targets = media_targets_for_relationships(relationship_roots)
    relationship_bound = bool(first_logo) and all(
        part in all_media_targets
        for part in first_logo.get("media_parts", [])  # type: ignore[union-attr]
    )
    reference_hashes = set()
    if reference_logo:
        first_reference_logo = reference_logo.get("first_logo") or {}
        if isinstance(first_reference_logo, dict):
            sha_map = first_reference_logo.get("media_sha256") or {}
            if isinstance(sha_map, dict):
                reference_hashes = {str(value) for value in sha_map.values() if value}
    actual_hashes = set()
    if isinstance(first_logo, dict):
        sha_map = first_logo.get("media_sha256") or {}
        if isinstance(sha_map, dict):
            actual_hashes = {str(value) for value in sha_map.values() if value}
    hash_matches_reference = not reference_hashes or bool(reference_hashes & actual_hashes)
    if not required and not first_logo:
        passed = True
        requirement_mode = "not-required-template-has-no-cover-logo"
    else:
        passed = bool(first_logo) and relationship_bound and hash_matches_reference
        requirement_mode = "required-template-logo"
    return {
        "passed": passed,
        "required": required,
        "requirement_mode": requirement_mode,
        "cover_logo_zone_end_paragraph_index": max(front_end - 1, 0),
        "first_logo": first_logo,
        "image_paragraph_count_in_cover_zone": len(image_rows),
        "image_paragraphs_in_cover_zone": image_rows,
        "media_relationship_bound": relationship_bound,
        "reference_hashes": sorted(reference_hashes),
        "actual_hashes": sorted(actual_hashes),
        "hash_matches_reference": hash_matches_reference,
    }


def cover_media_contract(
    paragraphs: list[ET.Element],
    texts: list[str],
    rels: dict[str, str],
    relationship_roots: list[ET.Element],
    *,
    required: bool = True,
) -> dict[str, object]:
    indices = first_surface_indices(texts)
    front_end = min(
        [item for item in (indices["toc"], indices["first_body"]) if item is not None],
        default=len(paragraphs) + 1,
    )
    cover_paragraphs = paragraphs[: max(front_end - 1, 0)]
    cover_rids: set[str] = set()
    for paragraph in cover_paragraphs:
        cover_rids.update(drawing_relationship_ids(paragraph))
    cover_media_parts = media_parts_for_relationship_ids(cover_rids, rels)
    all_media_targets = media_targets_for_relationships(relationship_roots)
    relationship_bound = bool(cover_media_parts) and all(part in all_media_targets for part in cover_media_parts)
    if not required and not cover_media_parts:
        passed = True
        relationship_bound = True
        requirement_mode = "not-required-template-has-no-cover-media"
    else:
        passed = bool(cover_media_parts) and relationship_bound
        requirement_mode = "required-template-or-default"
    return {
        "passed": passed,
        "required": required,
        "requirement_mode": requirement_mode,
        "cover_zone_end_paragraph_index": max(front_end - 1, 0),
        "relationship_ids": sorted(cover_rids),
        "media_parts": cover_media_parts,
        "media_relationship_bound": relationship_bound,
    }


def reference_cover_media_required(path: Path | None) -> dict[str, object] | None:
    if path is None:
        return None
    with zipfile.ZipFile(path) as zf:
        document_root = load_xml(zf, "word/document.xml")
        if document_root is None:
            raise ValueError(f"word/document.xml is missing or invalid in reference DOCX: {path}")
        rel_root = load_xml(zf, "word/_rels/document.xml.rels")
        paragraphs = visible_paragraphs(document_root)
        texts = [text_of(p).strip() for p in paragraphs]
        rels = relationship_targets(rel_root)
        relationship_roots = iter_relationship_roots(zf)
        contract = cover_media_contract(
            paragraphs,
            texts,
            rels,
            relationship_roots,
            required=True,
        )
        logo_contract = cover_logo_contract(
            paragraphs,
            texts,
            rels,
            relationship_roots,
            zf,
            required=bool(contract.get("relationship_ids") and contract.get("media_parts")),
        )
        return {
            "reference_docx_path": str(path),
            "reference_docx_sha256": sha256_file(path),
            "cover_media_required": bool(contract.get("relationship_ids") and contract.get("media_parts")),
            "cover_zone_end_paragraph_index": contract.get("cover_zone_end_paragraph_index"),
            "relationship_ids": contract.get("relationship_ids"),
            "media_parts": contract.get("media_parts"),
            "media_relationship_bound": contract.get("media_relationship_bound"),
            "cover_logo": logo_contract,
            "cover_layout": cover_layout_profile(zf),
            "additional_template_front_matter_media_policy": (
                "reference may contain sample-specific scanned pages; only the first school logo image is a strict copy baseline"
            ),
        }


def reference_header_rule_surface(path: Path | None) -> dict[str, object] | None:
    if path is None:
        return None
    rows: list[dict[str, object]] = []
    with zipfile.ZipFile(path) as zf:
        for name in sorted(item for item in zf.namelist() if item.startswith("word/header") and item.endswith(".xml")):
            root = load_xml(zf, name)
            if root is None:
                continue
            text = normalize_text(text_of(root))
            drawing_count = len(root.findall(".//w:drawing", NS))
            pict_count = len(root.findall(".//w:pict", NS))
            vshape_count = len(root.findall(".//v:shape", NS))
            if not text:
                continue
            rows.append(
                {
                    "part": name.replace("word/", ""),
                    "text": text[:120],
                    "drawing_count": drawing_count,
                    "pict_count": pict_count,
                    "vshape_count": vshape_count,
                    "has_rule_surface": bool(drawing_count or pict_count or vshape_count),
                }
            )
    required = any(row["has_rule_surface"] for row in rows)
    return {
        "reference_docx_path": str(path),
        "reference_docx_sha256": sha256_file(path),
        "header_rule_surface_required": required,
        "rows": rows,
    }


def first_surface_indices(texts: list[str]) -> dict[str, int | None]:
    indices: dict[str, int | None] = {
        "zh_abstract": None,
        "en_abstract": None,
        "toc": None,
        "first_body": None,
        "references": None,
        "appendix": None,
        "acknowledgement": None,
    }
    body_started = False
    toc_started = False
    for index, text in enumerate(texts, start=1):
        if indices["toc"] is None and is_toc_title(text):
            indices["toc"] = index
            toc_started = True
            continue
        if toc_started and not body_started and (has_toc_page_number_text(text) or "\t" in text):
            continue
        if toc_started and not body_started and (is_zh_abstract(text) or is_en_abstract(text)):
            toc_started = False
        if not toc_started and indices["zh_abstract"] is None and is_zh_abstract(text):
            indices["zh_abstract"] = index
        if not toc_started and indices["en_abstract"] is None and is_en_abstract(text):
            indices["en_abstract"] = index
        if indices["references"] is None and is_reference_title(text):
            indices["references"] = index
        if indices["appendix"] is None and is_appendix_title(text):
            indices["appendix"] = index
        if indices["acknowledgement"] is None and is_acknowledgement_title(text):
            indices["acknowledgement"] = index
        if toc_started and not body_started:
            if has_toc_page_number_text(text) or "\t" in text:
                continue
            if is_body_heading(text) and body_heading_level(text) == 1:
                indices["first_body"] = index
                body_started = True
                toc_started = False
                continue
        elif (
            not body_started
            and indices["first_body"] is None
            and is_body_heading(text)
            and body_heading_level(text) == 1
        ):
            indices["first_body"] = index
            body_started = True
    return indices


def paragraph_section_index_map(paragraphs: list[ET.Element]) -> dict[int, int]:
    result: dict[int, int] = {}
    running_section_index = -1
    for paragraph_index, paragraph in enumerate(paragraphs, start=1):
        result[paragraph_index] = running_section_index + 1
        if paragraph.find("./w:pPr/w:sectPr", NS) is not None:
            running_section_index += 1
    return result


def section_header_info_for_index(
    zf: zipfile.ZipFile,
    rels: dict[str, str],
    section_properties: list[ET.Element],
    section_index: int | None,
    *,
    prefer_first_if_title_page: bool = False,
    preferred_type: str | None = None,
) -> dict[str, object]:
    if section_index is None or section_index < 0 or section_index >= len(section_properties):
        return {"text": "", "part": "", "has_rule_surface": False}
    section = section_properties[section_index]
    refs = section.findall("./w:headerReference", NS)

    def info_for_ref(ref: ET.Element, *, allow_empty: bool = False) -> dict[str, object] | None:
        target = rels.get(ref.attrib.get(R + "id", ""))
        if not target:
            return None
        text, _fields, _count = part_text_and_fields(zf, target)
        root = load_xml(zf, f"word/{target}")
        has_rule_surface = False
        if root is not None:
            has_rule_surface = bool(
                root.findall(".//w:drawing", NS)
                or root.findall(".//w:pict", NS)
                or root.findall(".//v:shape", NS)
            )
        normalized = normalize_text(text)
        if normalized or allow_empty:
            return {
                "text": normalized,
                "part": target,
                "has_rule_surface": has_rule_surface,
            }
        return None

    if prefer_first_if_title_page and section.find("./w:titlePg", NS) is not None:
        for ref in refs:
            if attr(ref, "type") == "first":
                info = info_for_ref(ref, allow_empty=True)
                if info is not None:
                    return info

    if preferred_type:
        for ref in refs:
            if attr(ref, "type") == preferred_type:
                info = info_for_ref(ref, allow_empty=True)
                if info is not None:
                    return info

    for ref in refs:
        info = info_for_ref(ref)
        if info is not None:
            return info
    return {"text": "", "part": "", "has_rule_surface": False}


def toc_entry_indices_for_alignment(paragraphs: list[ET.Element], texts: list[str], indices: dict[str, int | None]) -> list[int]:
    toc_index = indices.get("toc")
    first_body_index = indices.get("first_body")
    if toc_index is None or first_body_index is None or first_body_index <= toc_index:
        return []
    rows: list[int] = []
    for index in range(toc_index + 1, first_body_index):
        paragraph = paragraphs[index - 1]
        text = texts[index - 1]
        if not text:
            continue
        maybe_toc_entry = has_tab(paragraph) or trailing_page_number(text) or re.match(r"^(\d+(\.\d+){0,2}|第.+章)", text)
        if maybe_toc_entry:
            rows.append(index)
    return rows


def reference_toc_header_contract(path: Path | None) -> dict[str, object] | None:
    if path is None:
        return None
    with zipfile.ZipFile(path) as zf:
        document_root = load_xml(zf, "word/document.xml")
        if document_root is None:
            raise ValueError(f"word/document.xml is missing or invalid in reference DOCX: {path}")
        rel_root = load_xml(zf, "word/_rels/document.xml.rels")
        paragraphs = visible_paragraphs(document_root)
        texts = [text_of(p).strip() for p in paragraphs]
        indices = first_surface_indices(texts)
        section_properties = document_root.findall(".//w:sectPr", NS)
        rels = relationship_targets(rel_root)
        section_index_by_paragraph = paragraph_section_index_map(paragraphs)
        toc_entries = toc_entry_indices_for_alignment(paragraphs, texts, indices)
        toc_index = indices.get("toc")
        first_body_index = indices.get("first_body")
        toc_section_index = section_index_by_paragraph.get(toc_index) if toc_index is not None else None
        toc_first_entry_index = toc_entries[0] if toc_entries else None
        toc_last_entry_index = toc_entries[-1] if toc_entries else None
        toc_first_entry_section_index = section_index_by_paragraph.get(toc_first_entry_index) if toc_first_entry_index is not None else None
        toc_last_entry_section_index = section_index_by_paragraph.get(toc_last_entry_index) if toc_last_entry_index is not None else None
        body_section_index = section_index_by_paragraph.get(first_body_index) if first_body_index is not None else None
        toc_header_info = section_header_info_for_index(zf, rels, section_properties, toc_section_index, preferred_type="default")
        toc_first_entry_header_info = section_header_info_for_index(zf, rels, section_properties, toc_first_entry_section_index, preferred_type="default")
        toc_last_entry_header_info = section_header_info_for_index(zf, rels, section_properties, toc_last_entry_section_index, preferred_type="default")
        body_header_info = section_header_info_for_index(zf, rels, section_properties, body_section_index)
        raw_toc_headers = [
            str(toc_header_info.get("text") or ""),
            str(toc_first_entry_header_info.get("text") or ""),
            str(toc_last_entry_header_info.get("text") or ""),
        ]
        toc_entry_header_inheritance = {
            "first_entry": toc_first_entry_section_index is None and bool(raw_toc_headers[0]),
            "last_entry": toc_last_entry_section_index is None and bool(raw_toc_headers[0]),
        }
        if toc_entry_header_inheritance["first_entry"]:
            toc_first_entry_header_info = {**toc_header_info}
        if toc_entry_header_inheritance["last_entry"]:
            toc_last_entry_header_info = {**toc_header_info}
        toc_headers = [
            str(toc_header_info.get("text") or ""),
            str(toc_first_entry_header_info.get("text") or ""),
            str(toc_last_entry_header_info.get("text") or ""),
        ]
    return {
        "reference_docx_path": str(path),
        "reference_docx_sha256": sha256_file(path),
        "toc_index": toc_index,
        "first_body_index": first_body_index,
        "toc_section_index": toc_section_index,
        "toc_first_entry_section_index": toc_first_entry_section_index,
        "toc_last_entry_section_index": toc_last_entry_section_index,
        "first_body_section_index": body_section_index,
        "distinct_toc_body_sections": toc_section_index is not None and body_section_index is not None and toc_section_index != body_section_index,
        "toc_header_text": toc_headers[0],
        "toc_first_entry_header_text": toc_headers[1],
        "toc_last_entry_header_text": toc_headers[2],
        "first_body_header_text": str(body_header_info.get("text") or ""),
        "toc_headers_are_empty": all(not text for text in toc_headers),
        "raw_toc_header_texts": raw_toc_headers,
        "toc_entry_header_inheritance": toc_entry_header_inheritance,
        "toc_header_parts": [
            str(toc_header_info.get("part") or ""),
            str(toc_first_entry_header_info.get("part") or ""),
            str(toc_last_entry_header_info.get("part") or ""),
        ],
        "policy": "reference-derived TOC running-header baseline using the default visible running header; missing TOC-entry sample sections inherit the TOC-title section header, while real empty donor headers remain valid only when the locked donor's default TOC headers are also empty",
    }


def reference_header_visibility_contract(path: Path | None) -> dict[str, object] | None:
    if path is None:
        return None
    with zipfile.ZipFile(path) as zf:
        document_root = load_xml(zf, "word/document.xml")
        if document_root is None:
            raise ValueError(f"word/document.xml is missing or invalid in reference DOCX: {path}")
        rel_root = load_xml(zf, "word/_rels/document.xml.rels")
        rels = relationship_targets(rel_root)
        section_properties = document_root.findall(".//w:sectPr", NS)
        header_records: list[dict[str, str]] = []
        for section_index, sect_pr in enumerate(section_properties):
            for ref in sect_pr.findall("./w:headerReference", NS):
                target = rels.get(ref.attrib.get(R + "id", ""))
                if not target:
                    continue
                text, _fields, _count = part_text_and_fields(zf, target)
                header_records.append(
                    {
                        "section_index": str(section_index),
                        "type": attr(ref, "type"),
                        "part": target,
                        "text": normalize_text(text),
                    }
                )
        return {
            "reference_docx_path": str(path),
            "reference_docx_sha256": sha256_file(path),
            "header_records": header_records,
            "header_texts": [record["text"] for record in header_records if record["text"]],
            "all_headers_empty": not any(record["text"] for record in header_records),
            "policy": "when the converted locked template has no visible running headers, a final DOCX may still use a safe static institutional running header if the school prose/template profile requires it; body/TOC heading leaks remain forbidden",
        }


def bounded_zone_end(start: int, candidates: list[int | None], default: int) -> int:
    valid = [item for item in candidates if item is not None and item > start]
    return min(valid) if valid else default


def audit_surface_contracts(
    *,
    docx_path: Path,
    document_root: ET.Element,
    zf: zipfile.ZipFile,
    paragraphs: list[ET.Element],
    texts: list[str],
    rels: dict[str, str],
    relationship_roots: list[ET.Element],
    header_texts: list[str],
    header_parts: list[str],
    footer_parts: list[str],
    footer_fields: list[str],
    reference_cover_media: dict[str, object] | None = None,
    reference_header_rule: dict[str, object] | None = None,
    reference_docx: Path | None = None,
) -> dict[str, object]:
    surface_issues: list[str] = []
    indices = first_surface_indices(texts)
    front_end = min(
        [item for item in (indices["toc"], indices["first_body"]) if item is not None],
        default=len(paragraphs) + 1,
    )
    first_body = indices["first_body"] or len(paragraphs) + 1
    references_index = indices["references"]
    appendix_index = indices["appendix"]
    acknowledgement_index = indices["acknowledgement"]

    cover_date_index = next((index for index, text in enumerate(texts, start=1) if is_cover_date_text(text)), None)
    cover_date_is_inside_cover_zone = (
        cover_date_index is not None
        and indices["zh_abstract"] is not None
        and cover_date_index < indices["zh_abstract"]
    )
    if (
        cover_date_index is not None
        and not cover_date_is_inside_cover_zone
        and not opens_on_new_page(paragraphs, texts, cover_date_index)
    ):
        surface_issues.append("cover date line must open on its own front-matter page")
    if indices["zh_abstract"] is not None and not opens_on_new_page(paragraphs, texts, indices["zh_abstract"]):
        surface_issues.append("Chinese abstract must open on a separate page after the cover/date front matter")

    # Cover logos/media are template-owned. Without a locked standard-cover
    # donor, do not invent a default logo requirement; the official school
    # writing guideline may describe the cover without embedding the actual
    # standard-cover media.
    cover_required = bool(
        reference_cover_media is not None
        and reference_cover_media.get("cover_media_required") is not False
        and reference_cover_media.get("cover_media_required") is not None
    )
    cover_media = cover_media_contract(
        paragraphs,
        texts,
        rels,
        relationship_roots,
        required=cover_required,
    )
    reference_cover_logo = None
    if reference_cover_media is not None and isinstance(reference_cover_media.get("cover_logo"), dict):
        reference_cover_logo = reference_cover_media.get("cover_logo")  # type: ignore[assignment]
    reference_cover_layout = None
    if reference_cover_media is not None and isinstance(reference_cover_media.get("cover_layout"), dict):
        reference_cover_layout = reference_cover_media.get("cover_layout")  # type: ignore[assignment]
    cover_logo = cover_logo_contract(
        paragraphs,
        texts,
        rels,
        relationship_roots,
        zf,
        reference_logo=reference_cover_logo if isinstance(reference_cover_logo, dict) else None,
        required=cover_required,
    )
    cover_layout = cover_template_layout_contract(
        zf,
        reference_cover_layout if isinstance(reference_cover_layout, dict) else None,
    )
    if cover_required and not cover_media.get("media_parts"):
        surface_issues.append("cover media/icon surface lacks drawing/object relationship binding before front matter")
    elif cover_required and cover_media.get("media_relationship_bound") is not True:
        surface_issues.append("cover media/icon relationship target is missing from package media parts")
    if cover_required and not cover_logo.get("first_logo"):
        surface_issues.append("cover school logo surface is missing before the abstract/front matter")
    elif cover_required and cover_logo.get("media_relationship_bound") is not True:
        surface_issues.append("cover school logo relationship target is missing from package media parts")
    elif cover_required and cover_logo.get("hash_matches_reference") is not True:
        surface_issues.append("cover school logo media hash does not match the locked template first-logo baseline")
    surface_issues.extend(str(item) for item in cover_layout.get("issues", []))
    cover_placeholders = cover_placeholder_contract(zf)
    surface_issues.extend(str(item) for item in cover_placeholders.get("issues", []))

    reference_front_paragraph_formats: dict[str, object] = {}
    reference_reference_title_summaries: list[dict[str, object]] = []
    reference_acknowledgement_title_summaries: list[dict[str, object]] = []
    if reference_docx is not None:
        try:
            with zipfile.ZipFile(reference_docx) as reference_zf:
                reference_root = load_xml(reference_zf, "word/document.xml")
                if reference_root is not None:
                    reference_paragraphs = visible_paragraphs(reference_root)
                    reference_texts = [text_of(p).strip() for p in reference_paragraphs]
                    reference_indices = first_surface_indices(reference_texts)
                    for reference_name in ("zh_abstract", "en_abstract", "toc"):
                        reference_idx = reference_indices.get(reference_name)
                        if reference_idx is not None:
                            reference_front_paragraph_formats[reference_name] = paragraph_format_summary(
                                reference_paragraphs[reference_idx - 1]
                            )
                    reference_reference_title_summaries = [
                        {
                            "paragraph_index": index,
                            "text": normalize_text(reference_texts[index - 1]),
                            **paragraph_format_summary(reference_paragraphs[index - 1]),
                        }
                        for index in range(1, len(reference_paragraphs) + 1)
                        if is_reference_title(reference_texts[index - 1])
                    ]
                    reference_acknowledgement_title_summaries = [
                        {
                            "paragraph_index": index,
                            "text": normalize_text(reference_texts[index - 1]),
                            **paragraph_format_summary(reference_paragraphs[index - 1]),
                        }
                        for index in range(1, len(reference_paragraphs) + 1)
                        if is_acknowledgement_title(reference_texts[index - 1])
                    ]
        except (OSError, zipfile.BadZipFile):
            reference_front_paragraph_formats = {}
            reference_reference_title_summaries = []
            reference_acknowledgement_title_summaries = []

    front_surface_names = ("zh_abstract", "en_abstract", "toc")
    front_paragraph_formats: dict[str, object] = {}
    for name in front_surface_names:
        idx = indices[name]
        if idx is None:
            front_paragraph_formats[name] = {"present": False}
            surface_issues.append(f"front matter hard paragraph/font field missing surface: {name}")
            continue
        paragraph = paragraphs[idx - 1]
        summary = paragraph_format_summary(paragraph)
        summary["present"] = True
        summary["paragraph_index"] = idx
        summary["text"] = normalize_text(texts[idx - 1])
        front_paragraph_formats[name] = summary
        if not has_style_or_direct_format(summary):
            surface_issues.append(f"front matter surface {name} lacks template-owned style or direct paragraph-format evidence")
        reference_front_summary = reference_front_paragraph_formats.get(name)
        reference_requires_explicit_size = (
            not isinstance(reference_front_summary, dict)
            or bool(reference_front_summary.get("font_size", {}).get("has_explicit_size"))
        )
        if name != "toc" and reference_requires_explicit_size and not summary["font_size"]["has_explicit_size"]:
            surface_issues.append(f"front matter surface {name} lacks explicit run font-size evidence")

    header_expected_parts = [normalize_text(text) for text in header_texts if normalize_text(text)]
    header_full_string = " | ".join(header_expected_parts)
    header_ok = bool(header_parts) and bool(header_expected_parts)
    if not header_ok:
        surface_issues.append("header full-display string is missing or no header part is bound")

    toc_rows: list[dict[str, object]] = []
    toc_start = indices["toc"]
    toc_end = first_body
    if toc_start is not None and toc_end > toc_start:
        for index in range(toc_start + 1, toc_end):
            paragraph = paragraphs[index - 1]
            text = texts[index - 1]
            if not text.strip():
                continue
            if is_reference_title(text) or is_acknowledgement_title(text):
                break
            page_number = trailing_page_number(text)
            looks_like_row = (
                has_tab(paragraph)
                or (bool(page_number) and toc_entry_has_visible_label(text))
                or is_body_heading(re.sub(r"(\d+|[ivxlcdmIVXLCDM]+)\s*$", "", text).strip())
            )
            if not looks_like_row:
                continue
            toc_rows.append(
                {
                    "paragraph_index": index,
                    "text": normalize_text(text),
                    "page_number": page_number,
                    "has_tab": has_tab(paragraph),
                    "has_right_tab": has_right_tab(paragraph),
                    "has_dotted_leader": has_dotted_tab(paragraph),
                    "has_right_dot_tab": paragraph_has_right_dot_tab(paragraph),
                    "has_visible_tab_run": paragraph.find(".//w:tab", NS) is not None,
                    "tabs": paragraph_tabs(paragraph),
                }
            )
    toc_rows_with_page_numbers = [row for row in toc_rows if row["page_number"]]
    toc_rows_with_right_tabs = [row for row in toc_rows if row["has_right_tab"]]
    toc_rows_with_right_dot_tabs = [row for row in toc_rows if row["has_right_dot_tab"]]
    toc_rows_with_visible_tab_runs = [row for row in toc_rows if row["has_visible_tab_run"]]
    toc_rows_without_right_tabs = [row for row in toc_rows if not row["has_right_tab"]]
    toc_rows_without_right_dot_tabs = [row for row in toc_rows if not row["has_right_dot_tab"]]
    toc_rows_without_visible_tab_runs = [row for row in toc_rows if not row["has_visible_tab_run"]]
    toc_right_tab_ok = (
        bool(toc_rows)
        and bool(toc_rows_with_page_numbers)
        and len(toc_rows_with_page_numbers) == len(toc_rows)
        and not toc_rows_without_right_tabs
        and not toc_rows_without_right_dot_tabs
        and not toc_rows_without_visible_tab_runs
    )
    if toc_rows and not toc_right_tab_ok:
        surface_issues.append("TOC page-number column lacks per-row visible tab plus right-dot-tab semantics for every row")
    elif toc_start is not None and not toc_rows:
        surface_issues.append("TOC has no auditable entry rows with right-tab/page-number semantics")
    toc_chapter_numbering_issues: list[str] = []
    for row in toc_rows:
        label = re.sub(r"(\d+|[ivxlcdmIVXLCDM]+)\s*$", "", str(row.get("text", ""))).strip()
        if not is_chapter_heading_text(label):
            continue
        if is_chinese_numeral_chapter_heading(label):
            continue
        elif not is_arabic_chapter_heading_format(label):
            toc_chapter_numbering_issues.append(f"TOC chapter entry must use `第1章  标题` style spacing: {row.get('text')}")
    surface_issues.extend(toc_chapter_numbering_issues)

    reference_title_summaries = [
        {"paragraph_index": index, "text": normalize_text(texts[index - 1]), **paragraph_format_summary(paragraphs[index - 1])}
        for index in range(1, len(paragraphs) + 1)
        if is_reference_title(texts[index - 1])
    ]
    acknowledgement_title_summaries = [
        {"paragraph_index": index, "text": normalize_text(texts[index - 1]), **paragraph_format_summary(paragraphs[index - 1])}
        for index in range(1, len(paragraphs) + 1)
        if is_acknowledgement_title(texts[index - 1])
    ]
    reference_title_template_sizes = [
        item["font_size"]["max_half_points"]
        for item in reference_reference_title_summaries
        if item.get("font_size", {}).get("max_half_points")
    ]
    acknowledgement_title_template_sizes = [
        item["font_size"]["max_half_points"]
        for item in reference_acknowledgement_title_summaries
        if item.get("font_size", {}).get("max_half_points")
    ]
    reference_title_min_half_points = min(reference_title_template_sizes) if reference_title_template_sizes else 30
    acknowledgement_title_min_half_points = (
        min(acknowledgement_title_template_sizes) if acknowledgement_title_template_sizes else 21
    )
    reference_main_titles = [
        item
        for item in reference_title_summaries
        if item.get("font_size", {}).get("max_half_points")
        and item["font_size"]["max_half_points"] >= reference_title_min_half_points
    ]
    reference_template_donor_titles = [
        item
        for item in reference_title_summaries
        if str(item.get("style_id") or "") in {"1", "Heading1"}
        and item.get("font_size", {}).get("max_half_points")
        and item["font_size"]["max_half_points"] >= 21
    ]
    acknowledgement_main_titles = [
        item
        for item in acknowledgement_title_summaries
        if item.get("font_size", {}).get("max_half_points")
        and item["font_size"]["max_half_points"] >= acknowledgement_title_min_half_points
    ]
    if len(reference_title_summaries) > 2:
        surface_issues.append(
            f"references title has too many body title layers: {len(reference_title_summaries)}; "
            "expected at most the template small-layer plus main title"
        )
    if len(acknowledgement_title_summaries) > 2:
        surface_issues.append(
            f"acknowledgement title has too many body title layers: {len(acknowledgement_title_summaries)}; "
            "expected at most the template small-layer plus main title"
        )
    if references_index is not None and not reference_main_titles and not reference_template_donor_titles:
        surface_issues.append("references title lacks the template main-title layer/font size")
    if acknowledgement_index is not None and not acknowledgement_main_titles:
        surface_issues.append("acknowledgement title lacks the template title layer/font size")

    reference_entry_summaries: list[dict[str, object]] = []
    if references_index is not None:
        end = bounded_zone_end(references_index, [appendix_index, acknowledgement_index], len(paragraphs) + 1)
        for index in range(references_index + 1, end):
            text = texts[index - 1]
            if is_appendix_title(text) or is_acknowledgement_title(text) or is_reference_title(text):
                break
            if not is_reference_entry(text):
                continue
            summary = paragraph_font_size_summary(paragraphs[index - 1])
            reference_entry_summaries.append(
                {
                    "paragraph_index": index,
                    "text_prefix": normalize_text(text)[:80],
                    "has_num_pr": paragraph_has_num_pr(paragraphs[index - 1]),
                    **summary,
                }
            )
            if not summary["has_explicit_size"]:
                surface_issues.append(f"reference entry paragraph {index} lacks explicit font-size evidence")
            elif summary.get("min_half_points") is not None and summary["min_half_points"] < 21:
                surface_issues.append(f"reference entry paragraph {index} font size is below template baseline 10.5pt")
        if not reference_entry_summaries:
            fallback_entries: list[dict[str, object]] = []
            if collect_bibliography_entries is not None:
                try:
                    body_paragraphs = paragraph_by_body_child_index(document_root)
                    for entry in collect_bibliography_entries(docx_path):
                        if not isinstance(entry, dict):
                            continue
                        body_index = entry.get("body_child_index")
                        paragraph = body_paragraphs.get(body_index) if isinstance(body_index, int) else None
                        if paragraph is None:
                            continue
                        summary = paragraph_font_size_summary(paragraph)
                        fallback_entry = {
                            "paragraph_index": "",
                            "body_child_index": body_index,
                            "text_prefix": normalize_text(str(entry.get("text", "")))[:80],
                            "has_num_pr": paragraph_has_num_pr(paragraph),
                            "source": "audit_docx_font_encoding.collect_bibliography_entries",
                            **summary,
                        }
                        fallback_entries.append(fallback_entry)
                        if not summary["has_explicit_size"]:
                            surface_issues.append(f"reference entry body-child {body_index} lacks explicit font-size evidence")
                        elif summary.get("min_half_points") is not None and summary["min_half_points"] < 21:
                            surface_issues.append(f"reference entry body-child {body_index} font size is below template baseline 10.5pt")
                except Exception as exc:
                    fallback_entries = [{"source": "audit_docx_font_encoding.collect_bibliography_entries", "error": str(exc)}]
            if fallback_entries and not any("error" in item for item in fallback_entries):
                reference_entry_summaries.extend(fallback_entries)
            else:
                surface_issues.append("references_entries surface has no bibliography entry paragraphs")

    acknowledgement_title_summary: dict[str, object] = {"present": False}
    if acknowledgement_index is not None:
        acknowledgement_title_text = texts[acknowledgement_index - 1]
        acknowledgement_title_summary = {
            "present": True,
            "paragraph_index": acknowledgement_index,
            "text": normalize_text(acknowledgement_title_text),
            "strict_title_text": is_strict_acknowledgement_title_text(acknowledgement_title_text),
            "has_num_pr": paragraph_has_num_pr(paragraphs[acknowledgement_index - 1]),
            **paragraph_format_summary(paragraphs[acknowledgement_index - 1]),
        }
        acknowledgement_style_id = str(acknowledgement_title_summary.get("style_id", ""))
        if not acknowledgement_title_summary["strict_title_text"]:
            surface_issues.append("acknowledgement_title contains internal spacing, numbering, or non-canonical title text")
        if paragraph_has_num_pr(paragraphs[acknowledgement_index - 1]):
            surface_issues.append("acknowledgement_title must not be a numbered/list paragraph")
        if acknowledgement_style_id.lower().startswith("toc"):
            surface_issues.append("acknowledgement_title must not use a TOC style")
        if not has_style_or_direct_format(acknowledgement_title_summary):
            surface_issues.append("acknowledgement_title lacks template-owned style or direct paragraph-format evidence")
        if acknowledgement_title_summary["font_size"]["has_explicit_size"] is not True:
            surface_issues.append("acknowledgement_title lacks explicit run font-size evidence")
    else:
        surface_issues.append("acknowledgement_title surface is missing")

    footer_font_sizes: list[int] = []
    for part in footer_parts:
        root = load_xml(zf, f"word/{part}")
        if root is None:
            continue
        for paragraph in root.findall(".//w:p", NS):
            has_page = any("PAGE" in (node.text or "").upper() for node in paragraph.findall(".//w:instrText", NS))
            has_page = has_page or any("PAGE" in node.attrib.get(W + "instr", "").upper() for node in paragraph.findall(".//w:fldSimple", NS))
            if has_page:
                footer_font_sizes.extend(run_font_sizes(paragraph))
    footer_page_contract = footer_page_field_contract(zf, footer_parts)
    header_footer_contract = header_footer_visual_contract(
        zf,
        header_parts=header_parts,
        footer_parts=footer_parts,
        require_header_rule_surface=bool(reference_header_rule and reference_header_rule.get("header_rule_surface_required")),
    )
    surface_issues.extend(str(item) for item in footer_page_contract.get("issues", []))
    surface_issues.extend(str(item) for item in header_footer_contract.get("issues", []))

    return {
        "cover_media": {
            **cover_media,
            "reference_baseline": reference_cover_media,
        },
        "cover_placeholders": cover_placeholders,
        "cover_school_logo": {
            **cover_logo,
            "reference_baseline": reference_cover_logo,
        },
        "cover_template_layout": {
            **cover_layout,
            "reference_baseline": reference_cover_layout,
        },
        "front_matter_hard_fields": {
            "passed": all(
                isinstance(front_paragraph_formats.get(name), dict)
                and front_paragraph_formats[name].get("present") is True
                and has_style_or_direct_format(front_paragraph_formats[name])
                and (
                    name == "toc"
                    or (
                        isinstance(reference_front_paragraph_formats.get(name), dict)
                        and not reference_front_paragraph_formats[name].get("font_size", {}).get("has_explicit_size")
                    )
                    or bool(front_paragraph_formats[name].get("font_size", {}).get("has_explicit_size"))
                )
                for name in front_surface_names
            ),
            "surfaces": front_paragraph_formats,
            "reference_surfaces": reference_front_paragraph_formats,
        },
        "header_full_display_string": {
            "passed": header_ok,
            "header_parts": header_parts,
            "observed_full_display_string": header_full_string,
        },
        "toc_page_number_right_tab": {
            "passed": (not toc_rows) or toc_right_tab_ok,
            "toc_row_count": len(toc_rows),
            "rows_with_page_numbers": len(toc_rows_with_page_numbers),
            "rows_with_right_tabs": len(toc_rows_with_right_tabs),
            "rows_with_right_dot_tabs": len(toc_rows_with_right_dot_tabs),
            "rows_with_visible_tab_runs": len(toc_rows_with_visible_tab_runs),
            "rows_without_right_tabs": toc_rows_without_right_tabs,
            "rows_without_right_dot_tabs": toc_rows_without_right_dot_tabs,
            "rows_without_visible_tab_runs": toc_rows_without_visible_tab_runs,
            "rows": toc_rows,
        },
        "toc_chapter_numbering": {
            "passed": not toc_chapter_numbering_issues,
            "issues": toc_chapter_numbering_issues,
        },
        "tail_title_layers": {
            "passed": (bool(reference_main_titles) or bool(reference_template_donor_titles))
            and bool(acknowledgement_main_titles),
            "reference_title_count": len(reference_title_summaries),
            "reference_main_title_count": len(reference_main_titles),
            "reference_template_donor_title_count": len(reference_template_donor_titles),
            "acknowledgement_title_count": len(acknowledgement_title_summaries),
            "acknowledgement_main_title_count": len(acknowledgement_main_titles),
            "reference_title_min_half_points": reference_title_min_half_points,
            "acknowledgement_title_min_half_points": acknowledgement_title_min_half_points,
            "reference_template_titles": reference_reference_title_summaries,
            "acknowledgement_template_titles": reference_acknowledgement_title_summaries,
            "reference_titles": reference_title_summaries,
            "acknowledgement_titles": acknowledgement_title_summaries,
        },
        "references_entries_font_size": {
            "passed": bool(reference_entry_summaries) and all(item["has_explicit_size"] and (item.get("min_half_points") or 0) >= 21 for item in reference_entry_summaries),
            "entry_count": len(reference_entry_summaries),
            "entries": reference_entry_summaries,
        },
        "acknowledgement_title_style": {
            "passed": acknowledgement_title_summary.get("present") is True
            and acknowledgement_title_summary.get("strict_title_text") is True
            and acknowledgement_title_summary.get("has_num_pr") is not True
            and not str(acknowledgement_title_summary.get("style_id", "")).lower().startswith("toc")
            and has_style_or_direct_format(acknowledgement_title_summary)
            and bool(acknowledgement_title_summary.get("font_size", {}).get("has_explicit_size")),
            "summary": acknowledgement_title_summary,
        },
        "footer_page_number_font_size": {
            "passed": footer_page_contract.get("page_field_count", 0) > 0 and footer_page_contract.get("passed") is True,
            "footer_parts": footer_parts,
            "sizes_half_points": sorted(set(footer_font_sizes)),
            "footer_page_field_count": sum(1 for instr in footer_fields if "PAGE" in instr.upper()),
            "field_contract": footer_page_contract,
        },
        "header_footer_visual_surface": {
            **header_footer_contract,
            "reference_baseline": reference_header_rule,
        },
        "issues": surface_issues,
    }


def check_passed(mapping: dict[str, object], *path: str) -> bool:
    current: object = mapping
    for key in path:
        if not isinstance(current, dict):
            return False
        current = current.get(key)
    return isinstance(current, dict) and current.get("passed") is True


def build_user_reported_issue_ledger(
    surface_checks: dict[str, object],
    bibliography_font_slots: dict[str, object],
) -> list[dict[str, object]]:
    rows = [
        {
            "user_issue": "封面图标丢失且字体和段落排版错误",
            "surface": "cover_style",
            "evidence_key": "surface_checks.cover_school_logo + surface_checks.cover_placeholders + surface_checks.cover_template_layout",
            "passed": check_passed(surface_checks, "cover_school_logo")
            and check_passed(surface_checks, "cover_placeholders")
            and check_passed(surface_checks, "cover_template_layout"),
        },
        {
            "user_issue": "前置页字体和段落排版错误",
            "surface": "front_matter",
            "evidence_key": "surface_checks.front_matter_hard_fields",
            "passed": check_passed(surface_checks, "front_matter_hard_fields"),
        },
        {
            "user_issue": "页眉错误",
            "surface": "header",
            "evidence_key": "surface_checks.header_full_display_string + surface_checks.header_footer_visual_surface",
            "passed": check_passed(surface_checks, "header_full_display_string")
            and check_passed(surface_checks, "header_footer_visual_surface"),
        },
        {
            "user_issue": "目录样式错误，数字没有右边对齐",
            "surface": "toc",
            "evidence_key": "surface_checks.toc_page_number_right_tab",
            "passed": check_passed(surface_checks, "toc_page_number_right_tab"),
        },
        {
            "user_issue": "参考文献内容字体大小错误",
            "surface": "references_entries",
            "evidence_key": "surface_checks.references_entries_font_size + bibliography_font_slots",
            "passed": check_passed(surface_checks, "references_entries_font_size")
            and bibliography_font_slots.get("passed") is True,
        },
        {
            "user_issue": "致谢标题样式丢失",
            "surface": "acknowledgement_title",
            "evidence_key": "surface_checks.acknowledgement_title_style + surface_checks.tail_title_layers",
            "passed": check_passed(surface_checks, "acknowledgement_title_style")
            and check_passed(surface_checks, "tail_title_layers"),
        },
        {
            "user_issue": "页脚的页数字体大小错误",
            "surface": "footer_page_numbers",
            "evidence_key": "surface_checks.footer_page_number_font_size",
            "passed": check_passed(surface_checks, "footer_page_number_font_size"),
        },
    ]
    for row in rows:
        row["verdict"] = "passed" if row["passed"] else "failed"
    return rows


def audit_docx(
    path: Path,
    *,
    allow_builder_styles: bool = False,
    require_toc_field: bool = False,
    reference_docx: Path | None = None,
    bibliography_size_name: str | None = None,
    bibliography_wps_ui_evidence_json: Path | None = None,
    bibliography_cjk_font: str | None = None,
    bibliography_latin_font: str | None = None,
    rendered_pdf_path: Path | None = None,
    allowed_near_empty_pages: set[int] | None = None,
) -> dict[str, object]:
    issues: list[str] = []
    reference_cover_media = reference_cover_media_required(reference_docx) if reference_docx is not None else None
    reference_header_rule = reference_header_rule_surface(reference_docx) if reference_docx is not None else None
    reference_toc_header_rule = reference_toc_header_contract(reference_docx) if reference_docx is not None else None
    reference_header_visibility_rule = reference_header_visibility_contract(reference_docx) if reference_docx is not None else None
    with zipfile.ZipFile(path) as zf:
        document_root = load_xml(zf, "word/document.xml")
        if document_root is None:
            raise ValueError("word/document.xml is missing or invalid")
        styles_root = load_xml(zf, "word/styles.xml")
        rel_root = load_xml(zf, "word/_rels/document.xml.rels")
        names = set(zf.namelist())
        relationship_roots = iter_relationship_roots(zf)

        paragraphs = visible_paragraphs(document_root)
        texts = [text_of(p).strip() for p in paragraphs]
        styles = style_name_map(styles_root)
        used_styles: dict[str, int] = {}
        nonempty_no_style = 0
        body_no_style = 0
        body_count = 0
        builder_style_count = 0
        body_started = False
        toc_started = False
        toc_entry_count = 0
        toc_entry_with_page_count = 0
        toc_entry_with_dotted_leader_count = 0
        toc_entry_with_tab_count = 0
        toc_entry_indices: list[int] = []
        toc_entry_scan_rows: list[dict[str, object]] = []
        page_break_count = 0
        paragraph_section_break_count = 0
        field_instructions: list[str] = []
        toc_field_depth_for_surface_scan = 0

        surface_indices = {
            "zh_abstract": None,
            "en_abstract": None,
            "toc": None,
            "first_body": None,
            "references": None,
            "appendix": None,
            "acknowledgement": None,
        }

        for index, paragraph in enumerate(paragraphs, start=1):
            text = texts[index - 1]
            style_id = paragraph_style_id(paragraph)
            if style_id:
                used_styles[style_id] = used_styles.get(style_id, 0) + 1
            elif text:
                nonempty_no_style += 1
            style_name = styles.get(style_id, style_id)
            if text and (style_id.lower().startswith("sgb") or style_name.lower().startswith("sgb")):
                builder_style_count += 1
            if has_page_break(paragraph):
                page_break_count += 1
            if has_section_break(paragraph):
                paragraph_section_break_count += 1
            instr = " ".join(
                [
                    *(node.text or "" for node in paragraph.findall(".//w:instrText", NS)),
                    *(node.attrib.get(W + "instr", "") for node in paragraph.findall(".//w:fldSimple", NS)),
                ]
            ).strip()
            if instr:
                field_instructions.append(instr)

            field_types = field_char_types(paragraph)
            paragraph_in_toc_field = (
                toc_started
                and (
                    toc_field_depth_for_surface_scan > 0
                    or has_toc_instruction(paragraph)
                    or "begin" in field_types
                )
            )
            paragraph_uses_toc_style = is_toc_style(style_id, style_name)
            if toc_started and not body_started and (is_zh_abstract(text) or is_en_abstract(text)) and not has_toc_page_number_text(text):
                toc_started = False
            if not toc_started and surface_indices["zh_abstract"] is None and is_zh_abstract(text):
                surface_indices["zh_abstract"] = index
            if not toc_started and surface_indices["en_abstract"] is None and is_en_abstract(text):
                surface_indices["en_abstract"] = index
            if surface_indices["toc"] is None and is_toc_title(text):
                surface_indices["toc"] = index
                toc_started = True
                continue
            if toc_started and not body_started:
                if (
                    is_body_heading(text)
                    and body_heading_level(text) == 1
                    and not (has_toc_page_number_text(text) or "\t" in text)
                    and not paragraph_in_toc_field
                    and not paragraph_uses_toc_style
                ):
                    surface_indices["first_body"] = index
                    body_started = True
                    toc_started = False
                elif text:
                    maybe_toc_entry = (
                        (paragraph_in_toc_field and toc_entry_has_visible_label(text))
                        or paragraph_uses_toc_style
                        or has_tab(paragraph)
                        or (trailing_page_number(text) and toc_entry_has_visible_label(text))
                        or is_body_heading(text)
                    )
                    if maybe_toc_entry:
                        toc_entry_indices.append(index)
                        toc_entry_count += 1
                        row = {
                            "paragraph_index": index,
                            "text": normalize_text(text),
                            "style_id": style_id,
                            "style_name": style_name,
                            "page_number": trailing_page_number(text),
                            "has_tab": has_tab(paragraph),
                            "has_right_dot_tab": paragraph_has_right_dot_tab(paragraph),
                            "inside_toc_field": paragraph_in_toc_field,
                        }
                        toc_entry_scan_rows.append(row)
                        if trailing_page_number(text):
                            toc_entry_with_page_count += 1
                        if has_dotted_tab(paragraph):
                            toc_entry_with_dotted_leader_count += 1
                        if has_tab(paragraph):
                            toc_entry_with_tab_count += 1
            elif (
                not body_started
                and surface_indices["first_body"] is None
                and is_body_heading(text)
                and body_heading_level(text) == 1
            ):
                surface_indices["first_body"] = index
                body_started = True

            if body_started and text:
                body_count += 1
                if not style_id:
                    body_no_style += 1

            if surface_indices["references"] is None and is_reference_title(text):
                surface_indices["references"] = index
            if surface_indices["appendix"] is None and is_appendix_title(text):
                surface_indices["appendix"] = index
            if surface_indices["acknowledgement"] is None and is_acknowledgement_title(text):
                surface_indices["acknowledgement"] = index

            if toc_started:
                for field_type in field_types:
                    if field_type == "begin":
                        toc_field_depth_for_surface_scan += 1
                    elif field_type == "end":
                        toc_field_depth_for_surface_scan = max(0, toc_field_depth_for_surface_scan - 1)

        section_properties = document_root.findall(".//w:sectPr", NS)
        section_count = len(section_properties)
        rels = relationship_targets(rel_root)
        header_refs: list[str] = []
        footer_refs: list[str] = []
        page_num_types: list[dict[str, str]] = []
        for sect_pr in section_properties:
            header_refs.extend(ref.attrib.get(R + "id", "") for ref in sect_pr.findall("./w:headerReference", NS))
            footer_refs.extend(ref.attrib.get(R + "id", "") for ref in sect_pr.findall("./w:footerReference", NS))
            pg_num = sect_pr.find("./w:pgNumType", NS)
            page_num_types.append({"fmt": attr(pg_num, "fmt"), "start": attr(pg_num, "start")})

        cover_section_isolation: dict[str, object] = {
            "present": bool(section_properties),
            "passed": False,
            "header_reference_count": 0,
            "footer_reference_count": 0,
            "has_page_numbering": False,
            "header_texts": [],
            "footer_texts": [],
            "footer_page_field_count": 0,
            "policy": "cover must follow the locked template header baseline and must not expose body/TOC heading leaks or visible PAGE fields",
        }
        if section_properties:
            first_section = section_properties[0]
            cover_header_refs = first_section.findall("./w:headerReference", NS)
            cover_footer_refs = first_section.findall("./w:footerReference", NS)
            cover_pg_num = first_section.find("./w:pgNumType", NS)
            cover_header_texts: list[str] = []
            cover_footer_texts: list[str] = []
            cover_footer_page_field_count = 0
            for ref in cover_header_refs:
                target = rels.get(ref.attrib.get(R + "id", ""))
                if not target:
                    continue
                text, _fields, _count = part_text_and_fields(zf, target)
                if normalize_text(text):
                    cover_header_texts.append(normalize_text(text))
            for ref in cover_footer_refs:
                target = rels.get(ref.attrib.get(R + "id", ""))
                if not target:
                    continue
                text, _fields, page_count = part_text_and_fields(zf, target)
                if normalize_text(text):
                    cover_footer_texts.append(normalize_text(text))
                cover_footer_page_field_count += page_count
            reference_requires_empty_headers = bool(
                (
                    reference_header_visibility_rule
                    and reference_header_visibility_rule.get("all_headers_empty")
                )
                or (
                    reference_toc_header_rule
                    and reference_toc_header_rule.get("toc_headers_are_empty")
                )
            )
            cover_header_ok = (
                not cover_header_texts
                if reference_requires_empty_headers
                else all(is_static_institutional_header_text(text) for text in cover_header_texts)
            )
            cover_footer_ok = not any(has_header_footer_body_leak(text) for text in cover_footer_texts)
            cover_section_isolation.update(
                {
                    "header_reference_count": len(cover_header_refs),
                    "footer_reference_count": len(cover_footer_refs),
                    "has_page_numbering": cover_pg_num is not None,
                    "header_texts": cover_header_texts,
                    "footer_texts": cover_footer_texts,
                    "footer_page_field_count": cover_footer_page_field_count,
                    "reference_header_visibility_baseline": reference_header_visibility_rule,
                    "reference_requires_empty_headers": reference_requires_empty_headers,
                    "passed": cover_header_ok and cover_footer_ok and cover_footer_page_field_count == 0,
                }
            )
        section_index_by_paragraph = paragraph_section_index_map(paragraphs)
        toc_section_alignment: dict[str, object] = {
            "toc_index": surface_indices.get("toc"),
            "first_body_index": surface_indices.get("first_body"),
            "passed": False,
            "toc_header_text": "",
            "toc_first_entry_header_text": "",
            "toc_last_entry_header_text": "",
            "first_body_header_text": "",
        }
        toc_index = surface_indices.get("toc")
        first_body_index = surface_indices.get("first_body")
        if toc_index is not None and first_body_index is not None and section_properties:
            toc_section_index = section_index_by_paragraph.get(toc_index)
            toc_first_entry_index = toc_entry_indices[0] if toc_entry_indices else None
            toc_last_entry_index = toc_entry_indices[-1] if toc_entry_indices else None
            toc_first_entry_section_index = section_index_by_paragraph.get(toc_first_entry_index) if toc_first_entry_index is not None else None
            toc_last_entry_section_index = section_index_by_paragraph.get(toc_last_entry_index) if toc_last_entry_index is not None else None
            body_section_index = section_index_by_paragraph.get(first_body_index)
            toc_header_info = section_header_info_for_index(zf, rels, section_properties, toc_section_index, preferred_type="default")
            toc_first_entry_header_info = section_header_info_for_index(zf, rels, section_properties, toc_first_entry_section_index, preferred_type="default")
            toc_last_entry_header_info = section_header_info_for_index(zf, rels, section_properties, toc_last_entry_section_index, preferred_type="default")
            body_header_info = section_header_info_for_index(zf, rels, section_properties, body_section_index)
            toc_header = str(toc_header_info.get("text") or "")
            toc_first_entry_header = str(toc_first_entry_header_info.get("text") or "")
            toc_last_entry_header = str(toc_last_entry_header_info.get("text") or "")
            body_header = str(body_header_info.get("text") or "")
            distinct_sections = toc_section_index is not None and body_section_index is not None and toc_section_index != body_section_index
            toc_header_is_toc_title = compact(toc_header) in {"目录", "toc", "tableofcontents"}
            body_header_is_toc_title = compact(body_header) in {"目录", "toc", "tableofcontents"}
            toc_first_entry_header_is_toc_title = compact(toc_first_entry_header) in {"目录", "toc", "tableofcontents"}
            toc_last_entry_header_is_toc_title = compact(toc_last_entry_header) in {"目录", "toc", "tableofcontents"}
            body_header_has_rule_surface = bool(body_header_info.get("has_rule_surface"))
            body_header_leak_is_template_running_header = (
                has_header_footer_body_leak(body_header)
                and body_header_has_rule_surface
                and not body_header_is_toc_title
            )
            same_header_is_safe = toc_header == body_header and is_static_institutional_header_text(toc_header)
            split_toc_body_headers_are_safe = (
                toc_header_is_toc_title
                and toc_first_entry_header_is_toc_title
                and toc_last_entry_header_is_toc_title
                and not body_header_is_toc_title
                and (not has_header_footer_body_leak(body_header) or body_header_leak_is_template_running_header)
            )
            neutral_headers_are_safe = (
                is_static_institutional_header_text(toc_header)
                and is_static_institutional_header_text(toc_first_entry_header)
                and is_static_institutional_header_text(toc_last_entry_header)
                and is_static_institutional_header_text(body_header)
                and not body_header_is_toc_title
            )
            final_toc_headers = [toc_header, toc_first_entry_header, toc_last_entry_header]
            reference_toc_headers_are_safe = False
            reference_toc_header_comparison: dict[str, object] | None = None
            if reference_toc_header_rule is not None:
                reference_toc_headers = [
                    str(reference_toc_header_rule.get("toc_header_text") or ""),
                    str(reference_toc_header_rule.get("toc_first_entry_header_text") or ""),
                    str(reference_toc_header_rule.get("toc_last_entry_header_text") or ""),
                ]
                reference_all_toc_headers_empty = all(not text for text in reference_toc_headers)
                final_all_toc_headers_empty = all(not text for text in final_toc_headers)
                final_toc_headers_match_reference = final_toc_headers == reference_toc_headers
                final_toc_headers_have_body_leak = any(has_header_footer_body_leak(text) for text in final_toc_headers)
                reference_body_header = str(reference_toc_header_rule.get("first_body_header_text") or "")
                reference_visible_headers = [
                    str(item)
                    for item in (
                        (reference_header_visibility_rule or {}).get("header_texts", [])
                        if isinstance(reference_header_visibility_rule, dict)
                        else []
                    )
                ]
                body_header_matches_reference = body_header == reference_body_header or body_header in reference_visible_headers
                final_toc_headers_are_toc_title = any(compact(text) in {"目录", "toc", "tableofcontents"} for text in final_toc_headers)
                body_header_is_safe_against_toc = (
                    not body_header_is_toc_title
                    and (not has_header_footer_body_leak(body_header) or body_header_leak_is_template_running_header)
                )
                final_toc_headers_are_safe_static = all(
                    is_static_institutional_header_text(text) for text in final_toc_headers
                )
                final_body_header_is_safe_static = (
                    is_static_institutional_header_text(body_header)
                    and not body_header_is_toc_title
                    and body_header_is_safe_against_toc
                )
                empty_reference_allows_static_institutional_header = (
                    reference_all_toc_headers_empty
                    and final_toc_headers_are_safe_static
                    and final_body_header_is_safe_static
                    and not final_toc_headers_have_body_leak
                )
                reference_toc_headers_are_safe = (
                    body_header_is_safe_against_toc
                    and not final_toc_headers_have_body_leak
                    and (
                        (body_header_matches_reference and final_toc_headers_match_reference)
                        or empty_reference_allows_static_institutional_header
                    )
                )
                reference_toc_header_comparison = {
                    "reference_toc_headers": reference_toc_headers,
                    "final_toc_headers": final_toc_headers,
                    "reference_all_toc_headers_empty": reference_all_toc_headers_empty,
                    "final_all_toc_headers_empty": final_all_toc_headers_empty,
                    "final_toc_headers_match_reference": final_toc_headers_match_reference,
                    "final_toc_headers_have_body_leak": final_toc_headers_have_body_leak,
                    "final_toc_headers_are_safe_static": final_toc_headers_are_safe_static,
                    "final_body_header_is_safe_static": final_body_header_is_safe_static,
                    "empty_reference_allows_static_institutional_header": empty_reference_allows_static_institutional_header,
                    "reference_body_header": reference_body_header,
                    "reference_visible_headers": reference_visible_headers,
                    "body_header_matches_reference": body_header_matches_reference,
                    "body_header_is_safe_against_toc": body_header_is_safe_against_toc,
                    "passed": reference_toc_headers_are_safe,
                }
            toc_section_alignment.update(
                {
                    "toc_section_index": toc_section_index,
                    "toc_first_entry_index": toc_first_entry_index,
                    "toc_last_entry_index": toc_last_entry_index,
                    "toc_first_entry_section_index": toc_first_entry_section_index,
                    "toc_last_entry_section_index": toc_last_entry_section_index,
                    "first_body_section_index": body_section_index,
                    "toc_header_text": toc_header,
                    "toc_first_entry_header_text": toc_first_entry_header,
                    "toc_last_entry_header_text": toc_last_entry_header,
                    "first_body_header_text": body_header,
                    "toc_header_part": toc_header_info.get("part", ""),
                    "toc_first_entry_header_part": toc_first_entry_header_info.get("part", ""),
                    "toc_last_entry_header_part": toc_last_entry_header_info.get("part", ""),
                    "first_body_header_part": body_header_info.get("part", ""),
                    "first_body_header_has_rule_surface": body_header_has_rule_surface,
                    "policy": "TOC and body must be separate sections and running headers must match the locked template baseline when a template is provided",
                    "distinct_sections": distinct_sections,
                    "toc_header_is_toc_title": toc_header_is_toc_title,
                    "toc_first_entry_header_is_toc_title": toc_first_entry_header_is_toc_title,
                    "toc_last_entry_header_is_toc_title": toc_last_entry_header_is_toc_title,
                    "first_body_header_is_toc_title": body_header_is_toc_title,
                    "first_body_header_leak_is_template_running_header": body_header_leak_is_template_running_header,
                    "same_header_is_safe_static_text": same_header_is_safe,
                    "split_toc_body_headers_are_safe": split_toc_body_headers_are_safe,
                    "neutral_headers_are_safe": neutral_headers_are_safe,
                    "reference_toc_header_baseline": reference_toc_header_rule,
                    "reference_toc_header_comparison": reference_toc_header_comparison,
                    "reference_toc_headers_are_safe": reference_toc_headers_are_safe,
                    "passed": distinct_sections and (
                        reference_toc_headers_are_safe
                        if reference_toc_header_rule is not None
                        else (
                            same_header_is_safe
                            or split_toc_body_headers_are_safe
                            or neutral_headers_are_safe
                        )
                    ),
                }
            )

        toc_field_integrity: dict[str, object] = {
            "passed": False,
            "toc_index": surface_indices.get("toc"),
            "first_body_index": surface_indices.get("first_body"),
            "toc_field_seen": False,
            "toc_field_closed_before_first_body": False,
            "section_breaks_inside_toc_field": [],
            "first_body_contains_field_surface": False,
            "first_body_field_depth_before": None,
            "toc_boundary_paragraph_index": None,
            "toc_boundary_is_field_free": False,
            "issues": [],
        }
        toc_index_for_field = surface_indices.get("toc")
        first_body_index_for_field = surface_indices.get("first_body")
        if toc_index_for_field is not None and first_body_index_for_field is not None:
            toc_field_state = inspect_live_toc_fields(document_root)
            field_depth = 0
            toc_field_depth = 0
            toc_field_seen = False
            section_breaks_inside_toc_field: list[dict[str, object]] = []
            first_body_depth_before = 0
            first_body_contains_field_surface = False
            for paragraph_index, paragraph in enumerate(paragraphs, start=1):
                if paragraph_index == first_body_index_for_field:
                    first_body_depth_before = toc_field_depth
                    first_body_contains_field_surface = has_field_surface(paragraph)
                if has_toc_instruction(paragraph) and paragraph_index >= toc_index_for_field:
                    toc_field_seen = True
                if toc_field_depth > 0 and has_section_break(paragraph):
                    section_breaks_inside_toc_field.append(
                        {
                            "paragraph_index": paragraph_index,
                            "text": text_of(paragraph).strip()[:120],
                        }
                    )
                for field_type in field_char_types(paragraph):
                    if field_type == "begin":
                        field_depth += 1
                        if paragraph_index >= toc_index_for_field:
                            toc_field_depth += 1
                            toc_field_seen = True
                    elif field_type == "end":
                        if paragraph_index >= toc_index_for_field and toc_field_depth > 0:
                            toc_field_depth -= 1
                        field_depth = max(0, field_depth - 1)
            boundary_index = None
            for candidate_index in range(max(toc_index_for_field, first_body_index_for_field - 4), first_body_index_for_field):
                paragraph = paragraphs[candidate_index - 1]
                if has_section_break(paragraph) and not text_of(paragraph).strip():
                    boundary_index = candidate_index
                    break
            toc_field_integrity_issues: list[str] = []
            if not toc_field_seen:
                toc_field_integrity_issues.append("TOC field instruction/markers were not found")
            if section_breaks_inside_toc_field:
                toc_field_integrity_issues.append("TOC field result contains section break(s)")
            if first_body_depth_before != 0:
                toc_field_integrity_issues.append("first body heading starts while TOC field depth is nonzero")
            if first_body_contains_field_surface:
                toc_field_integrity_issues.append("first body heading paragraph contains TOC/field marker surface")
            if boundary_index is None:
                toc_field_integrity_issues.append("TOC/body section boundary is not on an independent field-free empty paragraph")
            toc_field_integrity.update(
                {
                    "toc_field_seen": toc_field_seen,
                    "live_toc_field_count": toc_field_state["count"],
                    "locked_toc_field_count": toc_field_state["locked_count"],
                    "toc_field_closed_before_first_body": first_body_depth_before == 0,
                    "section_breaks_inside_toc_field": section_breaks_inside_toc_field,
                    "first_body_contains_field_surface": first_body_contains_field_surface,
                    "first_body_field_depth_before": first_body_depth_before,
                    "toc_boundary_paragraph_index": boundary_index,
                    "toc_boundary_is_field_free": bool(
                        boundary_index is not None and not has_field_surface(paragraphs[boundary_index - 1])
                    ),
                    "issues": toc_field_integrity_issues,
                    "passed": not toc_field_integrity_issues,
                }
            )

        header_parts = sorted({rels.get(rid, "") for rid in header_refs if rels.get(rid, "")})
        footer_parts = sorted({rels.get(rid, "") for rid in footer_refs if rels.get(rid, "")})
        header_texts: list[str] = []
        header_fields: list[str] = []
        footer_texts: list[str] = []
        footer_fields: list[str] = []
        footer_page_field_counts_by_part: dict[str, int] = {}
        footer_page_field_instance_count = 0
        for part in header_parts:
            text, fields, _page_count = part_text_and_fields(zf, part)
            header_texts.append(text)
            header_fields.append(fields)
        for part in footer_parts:
            text, fields, page_count = part_text_and_fields(zf, part)
            footer_texts.append(text)
            footer_fields.append(fields)
            footer_page_field_counts_by_part[part] = page_count
            footer_page_field_instance_count += page_count

        toc_field_state = inspect_live_toc_fields(document_root)
        live_toc_count = toc_field_state["count"]
        locked_toc_count = toc_field_state["locked_count"]
        footer_page_field_count = footer_page_field_instance_count

        surface_checks = audit_surface_contracts(
            docx_path=path,
            document_root=document_root,
            zf=zf,
            paragraphs=paragraphs,
            texts=texts,
            rels=rels,
            relationship_roots=relationship_roots,
            header_texts=header_texts,
            header_parts=header_parts,
            footer_parts=footer_parts,
            footer_fields=footer_fields,
            reference_cover_media=reference_cover_media,
            reference_header_rule=reference_header_rule,
            reference_docx=reference_docx,
        )

    for name, index in surface_indices.items():
        if name in {"appendix"}:
            continue
        if index is None:
            issues.append(f"missing required thesis surface: {name}")

    front_indices = [
        value
        for value in (surface_indices["zh_abstract"], surface_indices["en_abstract"], surface_indices["toc"])
        if value is not None
    ]
    if (
        surface_indices["zh_abstract"] is not None
        and surface_indices["en_abstract"] is not None
        and surface_indices["zh_abstract"] > surface_indices["en_abstract"]
    ):
        issues.append(f"Chinese abstract must appear before English abstract: {surface_indices}")
    if surface_indices["first_body"] is not None and any(value > surface_indices["first_body"] for value in front_indices):
        issues.append(f"front-matter surface appears after first body chapter: {surface_indices}")
    tail_order = [
        value
        for value in (surface_indices["first_body"], surface_indices["references"], surface_indices["acknowledgement"])
        if value is not None
    ]
    if tail_order != sorted(tail_order):
        issues.append(f"body/end matter order is wrong: {surface_indices}")

    zh_keyword_index = next((index for index, text in enumerate(texts, start=1) if is_zh_keyword(text)), None)
    en_keyword_index = next((index for index, text in enumerate(texts, start=1) if is_en_keyword(text)), None)
    toc_before_abstract = (
        surface_indices["toc"] is not None
        and surface_indices["zh_abstract"] is not None
        and surface_indices["toc"] < surface_indices["zh_abstract"]
    )
    frontmatter_page_boundaries = {
        "zh_keyword_index": zh_keyword_index,
        "en_keyword_index": en_keyword_index,
        "toc_before_abstract": toc_before_abstract,
        "zh_abstract_to_en_abstract_separated": (
            zh_keyword_index is not None
            and surface_indices["en_abstract"] is not None
            and opens_on_new_page(paragraphs, texts, surface_indices["en_abstract"])
        ),
        "en_abstract_to_toc_separated": (
            not toc_before_abstract
            and
            en_keyword_index is not None
            and surface_indices["toc"] is not None
            and opens_on_new_page(paragraphs, texts, surface_indices["toc"])
        ),
        "toc_to_zh_abstract_separated": (
            toc_before_abstract
            and surface_indices["zh_abstract"] is not None
            and opens_on_new_page(paragraphs, texts, surface_indices["zh_abstract"])
        ),
        "en_abstract_to_first_body_separated": (
            toc_before_abstract
            and en_keyword_index is not None
            and surface_indices["first_body"] is not None
            and opens_on_new_page(paragraphs, texts, surface_indices["first_body"])
        ),
        "toc_to_first_body_separated": (
            not toc_before_abstract
            and
            surface_indices["toc"] is not None
            and surface_indices["first_body"] is not None
            and opens_on_new_page(paragraphs, texts, surface_indices["first_body"])
        ),
    }
    if not frontmatter_page_boundaries["zh_abstract_to_en_abstract_separated"]:
        issues.append("Chinese abstract and English abstract are not separated by a page/section boundary")
    if toc_before_abstract:
        if not frontmatter_page_boundaries["toc_to_zh_abstract_separated"]:
            issues.append("TOC and Chinese abstract are not separated by a page/section boundary")
        if not frontmatter_page_boundaries["en_abstract_to_first_body_separated"]:
            issues.append("English abstract and first body chapter are not separated by a page/section boundary")
    else:
        if not frontmatter_page_boundaries["en_abstract_to_toc_separated"]:
            issues.append("English abstract and TOC are not separated by a page/section boundary")
        if not frontmatter_page_boundaries["toc_to_first_body_separated"]:
            issues.append("TOC and first body chapter are not separated by a page/section boundary")

    chapter_heading_numbering_issues: list[str] = []
    allow_chinese_chapter_headings = reference_allows_chinese_chapter_headings(reference_docx)
    for index, text in enumerate(texts, start=1):
        if not is_chapter_heading_text(text):
            continue
        if is_chinese_numeral_chapter_heading(text):
            if not allow_chinese_chapter_headings:
                chapter_heading_numbering_issues.append(f"paragraph {index} chapter heading uses Chinese numeral chapter number: {normalize_text(text)}")
        elif not is_arabic_chapter_heading_format(text):
            chapter_heading_numbering_issues.append(f"paragraph {index} chapter heading must use `第1章  标题` style spacing: {normalize_text(text)}")
    issues.extend(chapter_heading_numbering_issues)
    reference_body_heading_format = reference_body_heading_format_contract(reference_docx)
    reference_heading_rows = (
        reference_body_heading_format.get("sample_rows")
        if isinstance(reference_body_heading_format, dict)
        else None
    )
    body_heading_rows = collect_body_heading_format_rows(paragraphs, texts, styles_root)
    body_heading_format_issues = body_heading_template_baseline_issues(
        body_heading_rows,
        reference_heading_rows if isinstance(reference_heading_rows, list) else None,
    )
    issues.extend(body_heading_format_issues)
    toc_heading_coverage = body_heading_toc_coverage(body_heading_rows, toc_entry_scan_rows)
    if not toc_heading_coverage.get("passed"):
        missing_rows = toc_heading_coverage.get("missing_headings", [])
        if isinstance(missing_rows, list):
            issues.append(
                "TOC heading coverage missing body headings: "
                + "; ".join(
                    "p{paragraph_index} L{level} {text}".format(
                        paragraph_index=row.get("paragraph_index", ""),
                        level=row.get("level", ""),
                        text=str(row.get("text", ""))[:80],
                    )
                    for row in missing_rows[:12]
                )
            )

    if section_count < 3:
        issues.append(
            "whole-thesis DOCX must have separate section topology for cover, front matter, and body/end matter"
        )
    if section_count <= 1 and any(text.strip() for text in header_texts):
        issues.append("single-section thesis applies a running header to the cover/front matter")
    if not cover_section_isolation.get("passed"):
        issues.append(
            "cover section must match the template running-header baseline and must not expose body/TOC heading leaks or visible PAGE fields before front matter"
        )
    if not toc_section_alignment.get("passed"):
        issues.append(
            "TOC and body sections must be distinct and must not leak TOC/body heading text through running headers"
        )
    if not toc_field_integrity.get("passed"):
        issues.append(
            "TOC field must close before the first body heading and the TOC/body section boundary must be field-free"
        )
    if (
        toc_entry_count > 0
        and toc_section_alignment.get("toc_section_index") is not None
        and toc_section_alignment.get("toc_first_entry_section_index") is not None
        and toc_section_alignment.get("toc_first_entry_section_index") != toc_section_alignment.get("toc_section_index")
    ):
        issues.append(
            "TOC title and first TOC entry are split across sections; this is not acceptable final evidence because it can create a title-only TOC page"
        )

    if require_toc_field and live_toc_count <= 0:
        issues.append("live TOC field is required but no TOC field instruction was found")
    if require_toc_field and locked_toc_count < live_toc_count:
        issues.append("live TOC field is required but not every TOC field is locked with w:fldLock=true")
    if live_toc_count <= 0 and toc_entry_with_dotted_leader_count <= 0:
        issues.append("TOC lacks both a live TOC field and dotted-leader TOC entries")
    if toc_entry_count > 0 and toc_entry_with_page_count <= 0 and live_toc_count <= 0:
        issues.append("static TOC-like heading list lacks visible page numbers")
    if toc_entry_count > 0 and toc_entry_with_tab_count <= 0 and live_toc_count <= 0:
        issues.append("static TOC-like heading list lacks tab-separated entry/page-number structure")
    toc_rows_without_visible_pages = [row for row in toc_entry_scan_rows if not row.get("page_number")]
    toc_rows_without_right_dot_tabs = [row for row in toc_entry_scan_rows if not row.get("has_right_dot_tab")]
    if toc_rows_without_visible_pages:
        issues.append(
            "TOC entry rows lack visible page numbers: "
            + "; ".join(
                f"p{row.get('paragraph_index')} {str(row.get('text', ''))[:80]}"
                for row in toc_rows_without_visible_pages[:8]
            )
        )
    if toc_rows_without_right_dot_tabs:
        issues.append(
            "TOC entry rows lack right-tab dotted leader semantics: "
            + "; ".join(
                f"p{row.get('paragraph_index')} {str(row.get('text', ''))[:80]}"
                for row in toc_rows_without_right_dot_tabs[:8]
            )
        )

    if not footer_parts:
        issues.append("no footer part is bound to the document sections")
    if footer_page_field_count <= 0:
        issues.append("footer/page-number surface lacks a PAGE field")
    duplicate_page_footer_parts = {
        part: count
        for part, count in footer_page_field_counts_by_part.items()
        if count > 1
    }
    if duplicate_page_footer_parts:
        issues.append(f"footer/page-number surface contains duplicate PAGE fields in footer parts: {duplicate_page_footer_parts}")
    if section_count >= 3:
        front_has_roman = any(item.get("fmt", "").lower() in {"roman", "upperroman", "lowerroman"} for item in page_num_types)
        # Word's numberInDash format renders Arabic page numbers with dash
        # decoration (for example "- 1 -"). Some school templates use it for
        # the main-body restart, so it is still an Arabic numbering chain.
        body_restarts = any(
            item.get("start") == "1"
            and item.get("fmt", "").lower() in {"decimal", "", "numberindash"}
            for item in page_num_types[1:]
        )
        if not front_has_roman:
            issues.append("front matter section does not expose a roman page-number format")
        if not body_restarts:
            issues.append("body section does not expose a restarted Arabic page-number chain")

    if builder_style_count and not allow_builder_styles:
        issues.append(f"builder-owned thesis styles are still used in visible content: {builder_style_count} paragraphs")
    if builder_style_count and body_count and body_no_style / max(body_count, 1) > 0.20:
        issues.append(f"too many body/end-matter paragraphs lack explicit style binding: {body_no_style}/{body_count}")
    if isinstance(surface_checks, dict):
        for surface_issue in surface_checks.get("issues", []):
            issues.append(str(surface_issue))

    bibliography_font_slot_issues: list[str] = []
    bibliography_wps_evidence_issues: list[str] = []
    bibliography_entry_count = 0
    wps_named_size_evidence_valid = False
    bibliography_expected_size_half_points = None
    if bibliography_size_name:
        bibliography_expected_size_half_points = CHINESE_FONT_SIZE_HALF_POINTS.get(bibliography_size_name)
        if bibliography_expected_size_half_points is None:
            issues.append(f"unknown bibliography Chinese size name: {bibliography_size_name}")
    elif bibliography_wps_ui_evidence_json is not None:
        bibliography_expected_size_half_points = "21"
    if collect_bibliography_entries is not None:
        try:
            bibliography_entry_count = len(collect_bibliography_entries(path))
        except Exception:
            bibliography_entry_count = 0
    if bibliography_wps_ui_evidence_json is not None and validate_wps_named_size_evidence is not None:
        bibliography_wps_evidence_issues = validate_wps_named_size_evidence(
            bibliography_wps_ui_evidence_json,
            docx_path=path,
            expected_size_name=bibliography_size_name,
            expected_size_half_points=bibliography_expected_size_half_points,
            expected_entry_count=bibliography_entry_count,
        )
        wps_named_size_evidence_valid = not bibliography_wps_evidence_issues
        issues.extend(f"bibliography WPS named-size evidence drift: {item}" for item in bibliography_wps_evidence_issues)
    if reference_docx is not None and bibliography_font_slot_hits is not None:
        bibliography_font_slot_issues = bibliography_font_slot_hits(
            reference_docx,
            path,
            expected_size_half_points=bibliography_expected_size_half_points,
            expected_size_name=bibliography_size_name,
            wps_named_size_evidence_valid=wps_named_size_evidence_valid,
            bibliography_cjk_font=bibliography_cjk_font,
            bibliography_latin_font=bibliography_latin_font,
        )
        issues.extend(f"bibliography font-slot drift: {item}" for item in bibliography_font_slot_issues)

    bibliography_font_slots = {
        "passed": not bibliography_font_slot_issues and not bibliography_wps_evidence_issues,
        "reference_docx_path": str(reference_docx) if reference_docx is not None else "",
        "expected_size_half_points": bibliography_expected_size_half_points or ("template-derived" if reference_docx is not None else ""),
        "expected_size_name": bibliography_size_name or "",
        "explicit_cjk_font_policy": bibliography_cjk_font or "",
        "explicit_latin_font_policy": bibliography_latin_font or "",
        "bibliography_entry_count": bibliography_entry_count,
        "wps_named_size_evidence_path": str(bibliography_wps_ui_evidence_json) if bibliography_wps_ui_evidence_json is not None else "",
        "wps_named_size_evidence_valid": wps_named_size_evidence_valid,
        "wps_named_size_evidence_issues": bibliography_wps_evidence_issues,
        "issues": bibliography_font_slot_issues,
    }
    user_reported_issue_ledger = build_user_reported_issue_ledger(
        surface_checks if isinstance(surface_checks, dict) else {},
        bibliography_font_slots,
    )
    failed_user_issues = [row for row in user_reported_issue_ledger if row.get("passed") is not True] if reference_docx is not None else []
    issues.extend(f"user-reported issue still open: {row.get('surface')} :: {row.get('user_issue')}" for row in failed_user_issues)
    list_pollution_audit: dict[str, object] = {
        "passed": False,
        "issues": ["scripts/audit_docx_list_pollution.py is unavailable"],
    }
    if audit_docx_list_pollution is not None:
        list_pollution_audit = audit_docx_list_pollution(path, template_docx=reference_docx)
    if list_pollution_audit.get("passed") is not True:
        issue_count = list_pollution_audit.get("issue_count", "unknown")
        issues.append(f"protected thesis surfaces contain abnormal list/bullet pollution: {issue_count} issue(s)")
    rendered_pdf_page_occupancy = audit_rendered_pdf_page_occupancy(
        rendered_pdf_path,
        allowed_near_empty_pages=allowed_near_empty_pages,
    )
    if rendered_pdf_page_occupancy.get("provided") and rendered_pdf_page_occupancy.get("passed") is not True:
        for pdf_issue in rendered_pdf_page_occupancy.get("issues", []):
            issues.append(str(pdf_issue))

    return {
        "schema": "graduation-project-builder.docx-whole-format-gate.v1",
        "generator": "scripts/audit_docx_whole_format_gate.py",
        "docx_path": str(path),
        "docx_sha256": sha256_file(path),
        "reference_docx_path": str(reference_docx) if reference_docx is not None else "",
        "reference_docx_sha256": sha256_file(reference_docx) if reference_docx is not None else "",
        "counts": {
            "paragraph_count": len(paragraphs),
            "section_count": section_count,
            "paragraph_section_break_count": paragraph_section_break_count,
            "page_break_count": page_break_count,
            "live_toc_field_count": live_toc_count,
            "locked_toc_field_count": locked_toc_count,
            "toc_entry_count": toc_entry_count,
            "toc_entry_with_page_count": toc_entry_with_page_count,
            "toc_entry_with_tab_count": toc_entry_with_tab_count,
            "toc_entry_with_dotted_leader_count": toc_entry_with_dotted_leader_count,
            "header_part_count": len(header_parts),
            "footer_part_count": len(footer_parts),
            "footer_page_field_count": footer_page_field_count,
            "builder_style_visible_paragraph_count": builder_style_count,
            "nonempty_no_style_paragraph_count": nonempty_no_style,
            "body_no_style_paragraph_count": body_no_style,
            "body_paragraph_count": body_count,
        },
        "toc_entry_scan_rows": toc_entry_scan_rows,
        "surfaces": surface_indices,
        "used_styles": used_styles,
        "header_parts": header_parts,
        "footer_parts": footer_parts,
        "header_texts": header_texts,
        "footer_texts": footer_texts,
        "footer_page_field_counts_by_part": footer_page_field_counts_by_part,
        "page_number_types": page_num_types,
        "cover_section_isolation": cover_section_isolation,
        "reference_toc_header_contract": reference_toc_header_rule,
        "reference_header_visibility_contract": reference_header_visibility_rule,
        "frontmatter_page_boundaries": frontmatter_page_boundaries,
        "toc_section_alignment": toc_section_alignment,
        "toc_field_integrity": toc_field_integrity,
        "surface_checks": surface_checks,
        "chapter_heading_numbering": {
            "passed": not chapter_heading_numbering_issues,
            "issues": chapter_heading_numbering_issues,
        },
        "body_heading_direct_format": {
            "passed": not body_heading_format_issues,
            "rows": body_heading_rows,
            "issues": body_heading_format_issues,
            "reference": reference_body_heading_format,
        },
        "toc_heading_coverage": toc_heading_coverage,
        "bibliography_font_slots": bibliography_font_slots,
        "list_pollution_audit": list_pollution_audit,
        "rendered_pdf_page_occupancy": rendered_pdf_page_occupancy,
        "user_reported_issue_ledger": user_reported_issue_ledger,
        "verdict": "passed" if not issues else "failed",
        "issues": issues,
        "passed": not issues,
    }


def write_report(report: dict[str, object], path: Path | None) -> None:
    text = json.dumps(report, ensure_ascii=True, indent=2)
    if path:
        path.write_text(text, encoding="utf-8")
    print(text)


def add_field_run(paragraph, instr: str, *, locked: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(docx_qn("w:fldCharType"), "begin")
    if locked:
        begin.set(docx_qn("w:fldLock"), "true")
    run._r.append(begin)
    run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(docx_qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)
    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(docx_qn("w:fldCharType"), "separate")
    run._r.append(separate)
    paragraph.add_run("1")
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(docx_qn("w:fldCharType"), "end")
    run._r.append(end)


def set_section_page_number(section, *, fmt: str, start: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn

    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(docx_qn("w:pgNumType"))):
        sect_pr.remove(existing)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(docx_qn("w:fmt"), fmt)
    pg_num.set(docx_qn("w:start"), str(start))
    sect_pr.append(pg_num)


def add_styled_paragraph(document, text: str, style_name: str, *, size_pt: int = 12):
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.style = style_name
    run = paragraph.add_run(text)
    run.font.size = Pt(size_pt)
    return paragraph


def add_right_dotted_tab(paragraph, *, pos: int = 9000) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn

    ppr = paragraph._p.get_or_add_pPr()
    tabs = ppr.find(docx_qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(docx_qn("w:val"), "right")
    tab.set(docx_qn("w:leader"), "dot")
    tab.set(docx_qn("w:pos"), str(pos))
    tabs.append(tab)


def self_test() -> int:
    import shutil

    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE

    synthetic_issues = body_heading_direct_format_issues(
        [
            {
                "paragraph_index": 1,
                "level": 2,
                "text": "1.1 Synthetic Heading",
                "indent": {},
                "effective_indent": {"firstLineChars": "200", "firstLine": "", "left": "", "right": ""},
                "style_chain": [{"style_id": "Heading2", "name": "heading 2", "indent": {}}, {"style_id": "Normal", "name": "Normal", "indent": {"firstLineChars": "200"}}],
                "numPr": False,
                "spacing": {},
            },
            {
                "paragraph_index": 2,
                "level": 4,
                "text": "1.1.1.1 Synthetic Heading",
                "indent": {},
                "effective_indent": {"hangingChars": "100", "firstLine": "", "firstLineChars": ""},
                "style_chain": [{"style_id": "Heading4", "name": "heading 4", "indent": {"hangingChars": "100"}}],
                "numPr": True,
                "spacing": {},
            },
        ]
    )
    if not any("firstLineChars" in item for item in synthetic_issues) or not any("level-4" in item for item in synthetic_issues):
        print("heading indentation synthetic self-test did not catch character-unit or level-4 residue", file=sys.stderr)
        return 1

    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        bad = tmp / "bad.docx"
        doc = Document()
        doc.add_paragraph("内蒙古科技大学")
        doc.add_paragraph("摘 要")
        doc.add_paragraph("Abstract")
        doc.add_paragraph("目 录")
        doc.add_paragraph("第1章 绪论")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("致 谢")
        doc.save(bad)
        bad_report = audit_docx(bad)
        if bad_report["passed"]:
            print("bad fixture unexpectedly passed", file=sys.stderr)
            return 1

        good = tmp / "good.docx"
        png = tmp / "cover-icon.png"
        png.write_bytes(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
            )
        )
        doc = Document()
        styles = doc.styles
        for style_name in ("GPBBody", "GPBHeading", "GPBFront"):
            if style_name not in styles:
                styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        styles["GPBBody"].font.size = None
        doc.add_picture(str(png))
        add_styled_paragraph(doc, "内蒙古科技大学", "GPBFront", size_pt=16)
        add_styled_paragraph(doc, "题目：测试题目", "GPBFront", size_pt=14)
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_page_number(doc.sections[1], fmt="roman", start=1)
        add_styled_paragraph(doc, "摘 要", "GPBFront", size_pt=14)
        add_styled_paragraph(doc, "关键词：测试", "GPBFront", size_pt=12)
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_styled_paragraph(doc, "Abstract", "GPBFront", size_pt=14)
        add_styled_paragraph(doc, "Key words: test", "GPBFront", size_pt=12)
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_styled_paragraph(doc, "目 录", "GPBFront", size_pt=14)
        toc_entry = add_styled_paragraph(doc, "第1章 绪论\t1", "GPBFront", size_pt=12)
        add_right_dotted_tab(toc_entry)
        toc_entry = add_styled_paragraph(doc, "1.1 Research Background\t1", "GPBFront", size_pt=12)
        add_right_dotted_tab(toc_entry)
        toc_entry = add_styled_paragraph(doc, "1.1.1 Technical Route\t1", "GPBFront", size_pt=12)
        add_right_dotted_tab(toc_entry)
        toc = add_styled_paragraph(doc, "", "GPBFront", size_pt=12)
        add_field_run(toc, ' TOC \\o "1-3" \\h \\z \\u ', locked=True)
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_page_number(doc.sections[-1], fmt="decimal", start=1)
        add_styled_paragraph(doc, "第1章  绪论", "GPBHeading", size_pt=14)
        add_styled_paragraph(doc, "1.1 Research Background", "GPBHeading", size_pt=12)
        add_styled_paragraph(doc, "1.1.1 Technical Route", "GPBHeading", size_pt=12)
        add_styled_paragraph(doc, "正文内容。", "GPBBody", size_pt=12)
        add_styled_paragraph(doc, "参考文献", "GPBHeading", size_pt=16)
        add_styled_paragraph(doc, "[1] Source.", "GPBBody", size_pt=11)
        add_styled_paragraph(doc, "致 谢", "GPBHeading", size_pt=14)
        for section in doc.sections[1:]:
            section.header.is_linked_to_previous = False
            section.header.paragraphs[0].text = "内蒙古科技大学毕业设计说明书（毕业论文）"
            section.footer.is_linked_to_previous = False
            add_field_run(section.footer.paragraphs[0], " PAGE ")
            for run in section.footer.paragraphs[0].runs:
                from docx.shared import Pt

                run.font.size = Pt(10.5)
        doc.save(good)
        good_report = audit_docx(good, require_toc_field=True)
        if not good_report["passed"]:
            print(json.dumps(good_report, ensure_ascii=True, indent=2), file=sys.stderr)
            return 1
        try:
            import fitz  # type: ignore
        except Exception:
            fitz = None  # type: ignore[assignment]
        if fitz is not None:
            near_empty_pdf = tmp / "near-empty.pdf"
            pdf = fitz.open()
            page = pdf.new_page(width=596, height=842)
            page.insert_text((293, 766), "iii", fontsize=8)
            pdf.save(str(near_empty_pdf))
            pdf.close()
            near_empty_report = audit_docx(good, require_toc_field=True, rendered_pdf_path=near_empty_pdf)
            near_empty_issues = " | ".join(str(item) for item in near_empty_report.get("issues", []))
            if near_empty_report["passed"] or not (
                "unexpected near-empty pages" in near_empty_issues
                or "unexpected blank pages" in near_empty_issues
            ):
                print(json.dumps(near_empty_report, ensure_ascii=True, indent=2), file=sys.stderr)
                return 1

        bad_toc_row = tmp / "bad-toc-row.docx"
        bad_toc_work = tmp / "bad-toc-row-work"
        if bad_toc_work.exists():
            shutil.rmtree(bad_toc_work)
        bad_toc_work.mkdir()
        with zipfile.ZipFile(good, "r") as zin:
            zin.extractall(bad_toc_work)
        bad_doc_xml = bad_toc_work / "word" / "document.xml"
        bad_tree = ET.parse(bad_doc_xml)
        bad_root = bad_tree.getroot()
        for paragraph in bad_root.findall(".//w:body/w:p", NS):
            if paragraph.find(".//w:tab", NS) is None:
                continue
            if "Research Background" not in text_of(paragraph):
                continue
            ppr = paragraph.find("./w:pPr", NS)
            if ppr is not None:
                tabs = ppr.find("./w:tabs", NS)
                if tabs is not None:
                    ppr.remove(tabs)
            for run in list(paragraph.findall("./w:r", NS)):
                for tab in list(run.findall("./w:tab", NS)):
                    run.remove(tab)
            text_nodes = paragraph.findall(".//w:t", NS)
            if text_nodes:
                text_nodes[0].text = "1.1 Broken TOC row without page"
                for node in text_nodes[1:]:
                    node.text = ""
            break
        bad_tree.write(bad_doc_xml, encoding="UTF-8", xml_declaration=True)
        with zipfile.ZipFile(bad_toc_row, "w", zipfile.ZIP_DEFLATED) as zout:
            for part in bad_toc_work.rglob("*"):
                if part.is_file():
                    zout.write(part, part.relative_to(bad_toc_work).as_posix())
        bad_toc_report = audit_docx(bad_toc_row, require_toc_field=True)
        bad_toc_issues = " | ".join(str(item) for item in bad_toc_report.get("issues", []))
        if bad_toc_report["passed"] or "TOC entry rows lack visible page numbers" not in bad_toc_issues:
            print(json.dumps(bad_toc_report, ensure_ascii=True, indent=2), file=sys.stderr)
            return 1

        missing_level3_toc = tmp / "missing-level3-toc.docx"
        missing_level3_work = tmp / "missing-level3-toc-work"
        if missing_level3_work.exists():
            shutil.rmtree(missing_level3_work)
        missing_level3_work.mkdir()
        with zipfile.ZipFile(good, "r") as zin:
            zin.extractall(missing_level3_work)
        missing_doc_xml = missing_level3_work / "word" / "document.xml"
        missing_tree = ET.parse(missing_doc_xml)
        missing_root = missing_tree.getroot()
        missing_body = missing_root.find(".//w:body", NS)
        if missing_body is not None:
            for paragraph in list(missing_body.findall("./w:p", NS)):
                if "1.1.1" not in text_of(paragraph) or paragraph.find(".//w:tab", NS) is None:
                    continue
                missing_body.remove(paragraph)
                break
        missing_tree.write(missing_doc_xml, encoding="UTF-8", xml_declaration=True)
        with zipfile.ZipFile(missing_level3_toc, "w", zipfile.ZIP_DEFLATED) as zout:
            for part in missing_level3_work.rglob("*"):
                if part.is_file():
                    zout.write(part, part.relative_to(missing_level3_work).as_posix())
        missing_level3_report = audit_docx(missing_level3_toc, require_toc_field=True)
        missing_level3_issues = " | ".join(str(item) for item in missing_level3_report.get("issues", []))
        if missing_level3_report["passed"] or "TOC heading coverage missing body headings" not in missing_level3_issues:
            print(json.dumps(missing_level3_report, ensure_ascii=True, indent=2), file=sys.stderr)
            return 1

        chinese_small_five = tmp / "chinese-small-five.docx"
        chinese_doc = Document(str(good))
        for paragraph in chinese_doc.paragraphs:
            if "\u7b2c1\u7ae0" in paragraph.text:
                paragraph.text = paragraph.text.replace("\u7b2c1\u7ae0", "\u7b2c\u4e00\u7ae0")
        from docx.shared import Pt

        for section in chinese_doc.sections[1:]:
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        chinese_doc.save(chinese_small_five)
        chinese_small_five_report = audit_docx(chinese_small_five, require_toc_field=True)
        if not chinese_small_five_report["passed"]:
            print(json.dumps(chinese_small_five_report, ensure_ascii=True, indent=2), file=sys.stderr)
            return 1

        reference = tmp / "reference-missing-toc-entries.docx"
        institutional_header = str(
            good_report.get("toc_section_alignment", {}).get("toc_header_text")
            or "\u5185\u8499\u53e4\u79d1\u6280\u5927\u5b66\u6bd5\u4e1a\u8bbe\u8ba1\u8bf4\u660e\u4e66\uff08\u6bd5\u4e1a\u8bba\u6587\uff09"
        )
        ref_doc = Document()
        ref_doc.add_paragraph("\u6458\u8981")
        ref_doc.add_section(WD_SECTION.NEW_PAGE)
        ref_doc.add_paragraph("Abstract")
        ref_doc.add_section(WD_SECTION.NEW_PAGE)
        ref_doc.add_paragraph("\u76ee\u5f55")
        ref_doc.add_section(WD_SECTION.NEW_PAGE)
        ref_doc.add_paragraph("\u7b2c1\u7ae0 \u7eea\u8bba")
        for section in ref_doc.sections[1:]:
            section.header.is_linked_to_previous = False
            section.header.paragraphs[0].text = institutional_header
        ref_doc.save(reference)
        referenced_report = audit_docx(good, require_toc_field=True, reference_docx=reference)
        referenced_toc_alignment = referenced_report.get("toc_section_alignment", {})
        reference_baseline = referenced_toc_alignment.get("reference_toc_header_baseline", {})
        reference_comparison = referenced_toc_alignment.get("reference_toc_header_comparison", {})
        inheritance = reference_baseline.get("toc_entry_header_inheritance", {}) if isinstance(reference_baseline, dict) else {}
        if (
            referenced_toc_alignment.get("passed") is not True
            or referenced_toc_alignment.get("reference_toc_headers_are_safe") is not True
            or not (isinstance(reference_comparison, dict) and reference_comparison.get("passed") is True)
            or inheritance.get("first_entry") is not True
            or inheritance.get("last_entry") is not True
        ):
            print(json.dumps(referenced_report, ensure_ascii=True, indent=2), file=sys.stderr)
            return 1

        leaked = tmp / "leaked-running-header.docx"
        leaked_doc = Document(str(good))
        for section in leaked_doc.sections[1:]:
            section.header.is_linked_to_previous = False
            section.header.paragraphs[0].text = "\u7b2c1\u7ae0 \u7eea\u8bba"
        leaked_doc.save(leaked)
        leaked_report = audit_docx(leaked, require_toc_field=True, reference_docx=reference)
        leaked_toc_alignment = leaked_report.get("toc_section_alignment", {})
        if leaked_report.get("passed") is True or leaked_toc_alignment.get("passed") is not False:
            print(json.dumps(leaked_report, ensure_ascii=True, indent=2), file=sys.stderr)
            return 1
    print("self-test passed")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path", nargs="?")
    parser.add_argument("--report-json")
    parser.add_argument("--allow-builder-styles", action="store_true")
    parser.add_argument("--require-toc-field", action="store_true")
    parser.add_argument("--reference-docx", help="Optional template/sample DOCX for cover media baseline.")
    parser.add_argument("--bibliography-size-name", help="Exact Chinese size name for bibliography entries, for example 五号.")
    parser.add_argument("--bibliography-wps-ui-evidence-json", help="WPS UI all-entry named-size evidence JSON.")
    parser.add_argument("--bibliography-cjk-font", help="Explicit bibliography CJK font policy, for example Songti.")
    parser.add_argument("--bibliography-latin-font", help="Explicit bibliography Latin font policy, for example Times New Roman.")
    parser.add_argument("--rendered-pdf", help="Optional exact rendered PDF for blank/near-empty page occupancy audit.")
    parser.add_argument(
        "--allow-near-empty-page",
        type=int,
        action="append",
        default=[],
        help="Explicitly allow a rendered blank/near-empty physical page number after independent review.",
    )
    parser.add_argument("--self-test", action="store_true")
    args = parser.parse_args()

    if args.self_test:
        return self_test()
    if not args.docx_path:
        parser.error("docx_path is required unless --self-test is used")
    report = audit_docx(
        Path(args.docx_path),
        allow_builder_styles=args.allow_builder_styles,
        require_toc_field=args.require_toc_field,
        reference_docx=Path(args.reference_docx) if args.reference_docx else None,
        bibliography_size_name=args.bibliography_size_name,
        bibliography_wps_ui_evidence_json=Path(args.bibliography_wps_ui_evidence_json) if args.bibliography_wps_ui_evidence_json else None,
        bibliography_cjk_font=args.bibliography_cjk_font,
        bibliography_latin_font=args.bibliography_latin_font,
        rendered_pdf_path=Path(args.rendered_pdf) if args.rendered_pdf else None,
        allowed_near_empty_pages=set(args.allow_near_empty_page or []),
    )
    write_report(report, Path(args.report_json) if args.report_json else None)
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
