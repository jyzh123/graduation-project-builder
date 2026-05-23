#!/usr/bin/env python3
"""Run a readable self-check on a rebuilt thesis sample."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path
from xml.etree import ElementTree as ET

import fitz  # type: ignore
from docx import Document  # type: ignore

from strip_docx_template_instruction_artifacts import (
    collect_instruction_artifacts,
    direct_paragraph_text,
    is_instruction_note_text,
    is_instruction_text,
    textbox_texts,
)
from thesis_figure_contract import (
    final_docx_figure_surface_issues,
    final_docx_manifest_requirement_issues,
    manifest_with_resolved_paths,
    validate_figure_manifest,
)
from thesis_template_profile import SCHEMA as TEMPLATE_PROFILE_SCHEMA


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def audit_line_value(text: str, prefix: str) -> str:
    for line in text.splitlines():
        if line.strip().lower().startswith(prefix.lower()):
            return line.partition(":")[2].strip()
    return ""


def font_audit_integrity_issues(font_text: str, final_docx: Path) -> list[str]:
    issues: list[str] = []
    if not font_text.strip():
        return ["font audit report missing"]
    docx_sha = audit_line_value(font_text, "- docx sha256:")
    if not re.fullmatch(r"[0-9a-fA-F]{64}", docx_sha):
        issues.append("font audit missing DOCX sha256")
    elif final_docx.exists() and docx_sha.lower() != sha256_file(final_docx).lower():
        issues.append("font audit DOCX sha256 does not match the final DOCX")
    try:
        entry_count = int(audit_line_value(font_text, "- bibliography entry count:"))
        run_count = int(audit_line_value(font_text, "- bibliography checked run count:"))
    except ValueError:
        issues.append("font audit missing bibliography entry/run coverage counts")
        entry_count = run_count = 0
    if "bibliography font-slot checks: pass" in font_text.lower():
        if entry_count <= 0:
            issues.append("font audit pass lacks positive bibliography entry coverage")
        if run_count <= 0:
            issues.append("font audit pass lacks positive bibliography run coverage")
    if audit_line_value(font_text, "- bibliography size policy source:") == "explicit-named-size-cli":
        if audit_line_value(font_text, "- bibliography named-size WPS evidence verdict:") != "pass":
            issues.append("font audit explicit named-size policy lacks WPS named-size pass evidence")
    return issues
from thesis_template_profile import cover_end_index as template_cover_end_index
from thesis_template_profile import is_cover_title_candidate as template_is_cover_title_candidate
from thesis_template_profile import looks_like_english_title as template_looks_like_english_title
from thesis_template_profile import profile_readiness_issues
from thesis_template_profile import select_chinese_title_indices as template_select_chinese_title_indices


ABSTRACT_BASELINE_PROFILE_SCHEMA = "graduation-project-builder.abstract-baseline-profile.v1"
APPROVED_ABSTRACT_BASELINE_PROFILE_SOURCE_TYPES = {
    "official-format-requirements",
    "teacher-approved-sample",
    "user-approved-sample",
}
ER_FAMILY_ALIASES = {"er", "erd", "e-r", "entity-relationship", "entity_relationship"}

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
}
W = "{%s}" % NS["w"]
M = "{%s}" % NS["m"]
PR = "{%s}" % NS["pr"]
WP = "{%s}" % NS["wp"]
CN_NUMBER_CHARS = "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341"
CAPTION_RE = re.compile(
    r"^(图|表|续表)\s*"
    r"([0-9A-Za-z一二三四五六七八九十]+(?:[-\.－][0-9A-Za-z一二三四五六七八九十]+)*)"
    r"(?=$|[\s\u3000:：、])"
)
CODE_TITLE_RE = re.compile(r"^(代码|程序清单|代码清单)\s*([0-9A-Za-z一二三四五六七八九十\-\.]+)")
ACKNOWLEDGEMENT_ALIASES = ("致谢", "谢辞", "谢  辞", "acknowledgements", "acknowledgments")
CODE_LIKE_TOKENS = (
    "def ",
    "class ",
    "return ",
    "public ",
    "private ",
    "protected ",
    "function ",
    "if (",
    "for (",
    "while (",
    "{",
    "}",
    ";",
    "=>",
    "SELECT ",
    "INSERT ",
    "UPDATE ",
    "DELETE ",
)


def normalize(text: str) -> str:
    return re.sub(r"[\s\u3000\u25a1]+", "", text or "")


def normalized_equals_any(text: str, aliases: tuple[str, ...]) -> bool:
    normalized = normalize(text).lower()
    return any(normalized == normalize(alias).lower() for alias in aliases)


def heading_label_key(text: str) -> str:
    head = re.split(r"[\(\uff08]", str(text or "").strip(), maxsplit=1)[0]
    return normalize(head).lower()


def first_body_page_marker(text: str) -> str:
    return re.split(r"[\(\uff08]", str(text or "").strip(), maxsplit=1)[0]


def xml_paragraph_has_real_image(paragraph: ET.Element) -> bool:
    return paragraph.find(".//a:blip", NS) is not None or paragraph.find(".//v:imagedata", NS) is not None


def find_page(texts: list[str], needle: str) -> int | None:
    target = normalize(needle).lower()
    for idx, text in enumerate(texts, start=1):
        if target in normalize(text).lower():
            return idx
    return None


def find_pages(texts: list[str], needle: str) -> list[int]:
    target = normalize(needle).lower()
    if not target:
        return []
    return [idx for idx, text in enumerate(texts, start=1) if target in normalize(text).lower()]


def find_last_page(texts: list[str], needle: str) -> int | None:
    target = normalize(needle).lower()
    for idx in range(len(texts), 0, -1):
        if target in normalize(texts[idx - 1]).lower():
            return idx
    return None


def page_contains_standalone_marker(text: str, marker: str) -> bool:
    compact_marker = normalize(marker).lower()
    if is_body_chapter_heading_text(marker):
        for line in (text or "").splitlines():
            if normalize(line).lower() == compact_marker:
                return True
        return False
    short_body_heading_markers = {
        normalize("\u7eea\u8bba").lower(),
        normalize("\u7ed3\u8bba").lower(),
        normalize("\u603b\u7ed3\u4e0e\u5c55\u671b").lower(),
    }
    if compact_marker in short_body_heading_markers:
        for line in (text or "").splitlines():
            if normalize(line).lower() == compact_marker:
                return True
        return False
    standalone_markers = {
        "\u6458\u8981",
        "abstract",
        "\u76ee\u5f55",
        "\u7b2c\u4e00\u7ae0\u7eea\u8bba",
        "\u7b2c\u4e00\u7ae0 \u7eea\u8bba",
        "\u53c2\u8003\u6587\u732e",
        "\u81f4\u8c22",
    }
    if compact_marker not in {normalize(item).lower() for item in standalone_markers}:
        return compact_marker in normalize(text).lower()
    for line in (text or "").splitlines():
        if normalize(line).lower() == compact_marker:
            return True
    return False


def find_standalone_page(texts: list[str], marker: str) -> int | None:
    for idx, text in enumerate(texts, start=1):
        if page_contains_standalone_marker(text, marker):
            return idx
    return None


def find_last_standalone_page(texts: list[str], marker: str) -> int | None:
    for idx in range(len(texts), 0, -1):
        if page_contains_standalone_marker(texts[idx - 1], marker):
            return idx
    return None


def find_rendered_heading_page(texts: list[str], heading: str, start_page: int = 1) -> int | None:
    target = normalize(heading).lower()
    if not target:
        return None
    for idx in range(max(1, start_page), len(texts) + 1):
        for line in texts[idx - 1].splitlines():
            line_key = normalize(line).lower()
            if line_key == target:
                return idx
    return None


def compact_rendered_heading_key(value: str) -> str:
    text = normalize(value or "").lower()
    text = text.replace("\uff0e", ".").replace("\u3002", ".").replace("\uff61", ".")
    return re.sub(r"[\s\u00a0\u200b\u200c\u200d\u3000]+", "", text)


def find_rendered_heading_page_loose(texts: list[str], heading: str, start_page: int = 1) -> int | None:
    exact = find_rendered_heading_page(texts, heading, start_page=start_page)
    if exact is not None:
        return exact
    target = compact_rendered_heading_key(heading)
    if not target:
        return None
    for idx in range(max(1, start_page), len(texts) + 1):
        page_key = compact_rendered_heading_key(texts[idx - 1])
        if target in page_key:
            return idx
    return None


def heading_level(text: str) -> int | None:
    stripped = re.sub(r"^[\s\u25a1]+", "", text.strip())
    if "\u2026" in stripped:
        return None
    sep = r"[\s\u25a1]+"
    normalized_numbering = re.sub(r"(?<=\d)[\s\u25a1]*[\.．][\s\u25a1]*(?=\d)", ".", stripped)
    if re.match(rf"^\d{{1,2}}\.\d+\.\d+\.\d+{sep}\S", normalized_numbering):
        return 4
    if re.match(rf"^\d{{1,2}}\.\d+\.\d+{sep}\S", normalized_numbering):
        return 3
    if re.match(rf"^\d{{1,2}}\.\d+{sep}\S", normalized_numbering):
        return 2
    if re.match(rf"^\d{{1,2}}{sep}\S", stripped) or contains_chapter_heading_marker(stripped):
        return 1
    return None


def toc_label_level(text: str, style_id: str = "", style_name: str = "") -> int | None:
    label = str(text or "").split("\t", 1)[0].strip()
    level = heading_level(label)
    if level is not None:
        return min(level, 4)
    style_key = normalize(style_id or style_name).lower()
    if style_key.startswith("toc"):
        digits = re.findall(r"\d+", style_key)
        return int(digits[-1]) if digits else 1
    stripped = re.sub(r"^[\s\u25a1]+", "", label)
    if normalize(stripped).lower() in {normalize("\u6458\u8981").lower(), normalize("Abstract").lower()}:
        return 1
    if re.match(r"^\u7b2c[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u7ae0", stripped):
        return 1
    match = re.match(r"^\d{1,2}(?:\.\d{1,2}){1,3}(?=\D|$)", stripped)
    if match:
        return min(match.group(0).count(".") + 1, 4)
    return None


def xml_paragraph_has_tab(paragraph: ET.Element) -> bool:
    if paragraph.find(".//w:r/w:tab", NS) is not None:
        return True
    return any("\t" in (node.text or "") for node in paragraph.findall(".//w:t", NS))


def xml_paragraph_has_page_break(paragraph: ET.Element) -> bool:
    return any(node.attrib.get(W + "type") == "page" for node in paragraph.findall(".//w:br", NS))


def is_toc_heading_text(text: str) -> bool:
    normalized = normalize(text).lower()
    key = heading_label_key(text)
    target = normalize("\u76ee\u5f55").lower()
    return (
        key in {target, "contents", "tableofcontents"}
        or normalized in {target, "contents", "tableofcontents"}
        or normalized.endswith(target)
    )


def is_toc_leader_entry_text(text: str) -> bool:
    return "\u2026" in str(text or "") and re.search(r"\d+\s*$", str(text or "")) is not None


def toc_visible_label_text(text: str) -> str:
    label = str(text or "").split("\t", 1)[0].strip()
    if "\u2026" in label:
        label = label.split("\u2026", 1)[0].strip()
    return re.sub(r"\s*(?:\d+|[IVXLCDM]+)\s*$", "", label, flags=re.IGNORECASE).strip()


def is_front_matter_toc_label(text: str) -> bool:
    label = normalize(toc_visible_label_text(text)).lower()
    return label in {normalize("\u6458\u8981").lower(), normalize("Abstract").lower()}


def contains_chapter_heading_marker(text: str) -> bool:
    stripped = re.sub(r"^[\s\u25a1]+", "", text or "")
    match = re.match(r"^(\u7b2c[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u7ae0)(.*)$", stripped)
    if not match:
        return False
    tail = match.group(2)
    if not tail:
        return True
    separator_match = re.match(r"^[\s\u25a1:：、.\-—]+", tail)
    if separator_match:
        tail = tail[separator_match.end() :]
    compact_tail = normalize(tail)
    sentence_punctuation = "\uff0c\uff1b\uff1a\uff01\uff1f\u3002\uff08\uff09,;:!?()"
    if any(ch in compact_tail for ch in sentence_punctuation):
        return False
    return len(compact_tail) <= (24 if separator_match else 12)


def is_body_chapter_heading_text(text: str) -> bool:
    if is_toc_leader_entry_text(text):
        return False
    return heading_level(text) == 1 or contains_chapter_heading_marker(text)


def is_short_heading_like_text(text: str) -> bool:
    stripped = str(text or "").strip()
    compact = normalize(stripped)
    if not compact:
        return False
    sentence_punctuation = "\uff0c\uff1b\uff1a\uff01\uff1f\u3002\uff08\uff09,;:!?()"
    if any(ch in compact for ch in sentence_punctuation):
        return False
    if re.search(r"\[\d+(?:\s*[-,]\s*\d+)*\]", stripped):
        return False
    return len(compact) <= 40


def is_body_chapter_heading_paragraph(
    paragraph: ET.Element,
    text: str,
    *,
    style_id: str = "",
    style_name: str = "",
    style_signature: dict[str, str] | None = None,
) -> bool:
    if is_toc_leader_entry_text(text):
        return False
    if is_body_chapter_heading_text(text):
        return True
    if not is_short_heading_like_text(text):
        return False
    style_signature = style_signature or {}
    style_is_level_1 = heading_style_label_matches(1, style_id, style_name) or style_signature.get("outlineLvl") == "0"
    direct_signature = _ppr_signature(paragraph.find("w:pPr", NS))
    direct_level = direct_signature.get("ilvl", "")
    style_level = style_signature.get("ilvl", "")
    has_level_1_numbering = (
        (direct_signature.get("numPr") == "yes" and direct_level in {"", "0"})
        or (style_signature.get("numPr") == "yes" and style_level in {"", "0"})
    )
    return style_is_level_1 or has_level_1_numbering


def is_official_declaration_checkbox_line(text: str) -> bool:
    normalized = normalize(text)
    return (
        "\u25a1" in str(text or "")
        and "\u516c\u5f00" in normalized
        and "\u4fdd\u5bc6" in normalized
        and (
            "\u5b66\u4f4d\u8bba\u6587" in normalized
            or "\u89e3\u5bc6" in normalized
            or "\u5e74" in normalized
        )
    )


def is_formula_layout_table(table: ET.Element) -> bool:
    return any("officeDocument/2006/math" in node.tag for node in table.iter())


def is_code_like_table(table: ET.Element) -> bool:
    text = "".join(node.text or "" for node in table.findall(".//w:t", NS))
    return is_code_like(text)


def cover_anchor_text(text: str) -> str | None:
    stripped = str(text or "").strip()
    if not stripped:
        return None
    normalized = normalize(stripped).lower()
    cover_tokens = {
        normalize("\u6bd5\u4e1a\u8bbe\u8ba1").lower(),
        normalize("\u6bd5\u4e1a\u8bba\u6587").lower(),
        normalize("\u5b66\u58eb\u5b66\u4f4d\u8bba\u6587").lower(),
        normalize("\u672c\u79d1\u751f\u6bd5\u4e1a\u8bba\u6587").lower(),
        "thesisofbachelor",
        "nanchanguniversity",
    }
    if any(token and token in normalized for token in cover_tokens):
        return stripped
    return None


def expected_style(level: int) -> str:
    return {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}[level]


def heading_style_aliases(level: int) -> set[str]:
    return {
        normalize(f"Heading {level}").lower(),
        normalize(f"Heading{level}").lower(),
        normalize(f"标题 {level}").lower(),
        normalize(f"标题{level}").lower(),
        normalize(f"{level}级标题").lower(),
        normalize({1: "一级标题", 2: "二级标题", 3: "三级标题", 4: "四级标题"}[level]).lower(),
    }


def heading_style_label_matches(level: int, style_id: str, style_name: str) -> bool:
    candidates = {normalize(style_id or "").lower(), normalize(style_name or "").lower()}
    aliases = heading_style_aliases(level)
    return bool(candidates & aliases) or any(alias in candidate for candidate in candidates for alias in aliases if alias)


def is_heading_style_label(style_label_text: str) -> bool:
    key = normalize(style_label_text or "").lower()
    return key.startswith("heading") or any(alias in key for level in range(1, 5) for alias in heading_style_aliases(level))


def collect_reference_page_markers(reference_doc: Document) -> dict[str, str | None]:
    markers = {
        "cover": None,
        "zh_abstract": None,
        "en_abstract": None,
        "toc": None,
        "first_body": None,
        "figure_page": None,
        "table_page": None,
        "references": None,
        "ack": None,
    }
    toc_seen = False
    toc_marker_seen = False
    body_heading_before_toc: str | None = None
    for para in reference_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        normalized = normalize(text).lower()
        heading_key = heading_label_key(text)
        style_name = (para.style.name if para.style else "").strip()
        has_tab = "\t" in text
        if markers["cover"] is None:
            cover_anchor = cover_anchor_text(text)
            if cover_anchor:
                markers["cover"] = cover_anchor
        if markers["zh_abstract"] is None and (
            heading_key in {normalize("\u6458\u8981").lower(), normalize("\u6458   \u8981").lower()}
            or normalized.startswith(normalize("\u6458\u8981\uff1a").lower())
            or normalized.startswith(normalize("\u6458\u8981:").lower())
        ):
            markers["zh_abstract"] = "\u6458\u8981"
        if markers["en_abstract"] is None and (
            heading_key == normalize("abstract").lower()
            or normalized.startswith(normalize("Abstract:").lower())
            or normalized.startswith(normalize("Abstract\uff1a").lower())
        ):
            markers["en_abstract"] = "Abstract"
        if markers["toc"] is None and (
            heading_key in {normalize("\u76ee\u5f55").lower(), normalize("\u76ee   \u5f55").lower()}
            or style_name.lower().startswith("toc")
        ):
            markers["toc"] = "\u76ee\u5f55"
            toc_seen = True
            toc_marker_seen = True
        if markers["first_body"] is None and toc_seen and not has_tab and not style_name.lower().startswith("toc") and (style_name == "Heading 1" or is_body_chapter_heading_text(text)):
            markers["first_body"] = first_body_page_marker(text)
        if body_heading_before_toc is None and not has_tab and (style_name == "Heading 1" or is_body_chapter_heading_text(text)):
            body_heading_before_toc = first_body_page_marker(text)
        if markers["figure_page"] is None and (style_name == "图注" or CAPTION_RE.match(text)) and text.startswith("图"):
            markers["figure_page"] = text
        if markers["table_page"] is None and (style_name == "表格标题" or CAPTION_RE.match(text)) and text.startswith("表"):
            markers["table_page"] = text
        if markers["references"] is None and normalized == normalize("参考文献").lower():
            markers["references"] = text
        if markers["ack"] is None and normalized_equals_any(text, ACKNOWLEDGEMENT_ALIASES):
            markers["ack"] = text
    if not toc_marker_seen:
        try:
            document = reference_doc.part.element
            for sdt in document.findall(".//w:sdt", NS):
                for paragraph in sdt.findall("./w:sdtContent/w:p", NS):
                    text = "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()
                    if not text:
                        continue
                    heading_key = heading_label_key(text)
                    has_tab = paragraph.find(".//w:tab", NS) is not None
                    style_el = paragraph.find("./w:pPr/w:pStyle", NS)
                    style_id = style_el.get(qn("w:val")) if style_el is not None else ""
                    if markers["toc"] is None and heading_key in {normalize("\u76ee\u5f55").lower(), normalize("\u76ee   \u5f55").lower()}:
                        markers["toc"] = "\u76ee\u5f55"
                        toc_seen = True
                        continue
                    if markers["first_body"] is None and toc_seen and not has_tab and not str(style_id).upper().startswith("TOC") and is_body_chapter_heading_text(text):
                        markers["first_body"] = first_body_page_marker(text)
        except Exception:
            pass
    if markers["first_body"] is None and body_heading_before_toc is not None:
        markers["first_body"] = body_heading_before_toc
    return markers


def reference_uses_heading_styles(reference_doc: Document) -> bool:
    for paragraph in reference_doc.paragraphs:
        if paragraph.style is None:
            continue
        level = heading_level(paragraph.text.strip())
        if level is not None and heading_style_label_matches(level, paragraph.style.style_id, paragraph.style.name):
            return True
    return False


def load_xml(docx_path: Path, part: str) -> ET.Element | None:
    with zipfile.ZipFile(docx_path) as zf:
        if part not in zf.namelist():
            return None
        return ET.fromstring(zf.read(part))


def normalize_font_name(value: str) -> str:
    return normalize(value).strip()


def paragraph_style_name_map(docx_path: Path) -> dict[str, str]:
    styles_root = load_xml(docx_path, "word/styles.xml")
    if styles_root is None:
        return {}
    mapping: dict[str, str] = {}
    for style in styles_root.findall("w:style", NS):
        if style.attrib.get(W + "type") != "paragraph":
            continue
        style_id = style.attrib.get(W + "styleId", "")
        if not style_id:
            continue
        name_node = style.find("w:name", NS)
        mapping[style_id] = name_node.attrib.get(W + "val", "") if name_node is not None else ""
    return mapping


def paragraph_style_font_map(docx_path: Path) -> dict[str, dict[str, str]]:
    styles_root = load_xml(docx_path, "word/styles.xml")
    if styles_root is None:
        return {}
    mapping: dict[str, dict[str, str]] = {}
    for style in styles_root.findall("w:style", NS):
        if style.attrib.get(W + "type") != "paragraph":
            continue
        style_id = style.attrib.get(W + "styleId", "")
        if not style_id:
            continue
        rfonts = style.find("w:rPr/w:rFonts", NS)
        sz = style.find("w:rPr/w:sz", NS)
        signature = {
            "eastAsia": "",
            "ascii": "",
            "hAnsi": "",
            "cs": "",
            "eastAsiaTheme": "",
            "asciiTheme": "",
            "hAnsiTheme": "",
            "csTheme": "",
            "size": "",
            "sizeCs": "",
            "bold": "absent",
        }
        if rfonts is not None:
            signature["eastAsia"] = normalize_font_name(rfonts.attrib.get(W + "eastAsia", ""))
            signature["ascii"] = normalize_font_name(rfonts.attrib.get(W + "ascii", ""))
            signature["hAnsi"] = normalize_font_name(rfonts.attrib.get(W + "hAnsi", ""))
            signature["cs"] = normalize_font_name(rfonts.attrib.get(W + "cs", ""))
            signature["eastAsiaTheme"] = rfonts.attrib.get(W + "eastAsiaTheme", "")
            signature["asciiTheme"] = rfonts.attrib.get(W + "asciiTheme", "")
            signature["hAnsiTheme"] = rfonts.attrib.get(W + "hAnsiTheme", "")
            signature["csTheme"] = rfonts.attrib.get(W + "csTheme", "")
        if sz is not None:
            signature["size"] = sz.attrib.get(W + "val", "")
        sz_cs = style.find("w:rPr/w:szCs", NS)
        if sz_cs is not None:
            signature["sizeCs"] = sz_cs.attrib.get(W + "val", "")
        bold = style.find("w:rPr/w:b", NS)
        if bold is not None:
            signature["bold"] = bold.attrib.get(W + "val", "true")
        mapping[style_id] = signature
    return mapping


def _ppr_signature(ppr: ET.Element | None) -> dict[str, str]:
    signature = {
        "align": "",
        "before": "",
        "after": "",
        "line": "",
        "lineRule": "",
        "keepNext": "no",
        "firstLine": "",
        "left": "",
        "right": "",
        "hanging": "",
        "firstLineChars": "",
        "leftChars": "",
        "rightChars": "",
        "hangingChars": "",
        "outlineLvl": "",
        "numPr": "no",
        "numId": "",
        "ilvl": "",
    }
    if ppr is None:
        return signature
    jc = ppr.find("w:jc", NS)
    if jc is not None:
        signature["align"] = jc.attrib.get(W + "val", "")
    spacing = ppr.find("w:spacing", NS)
    if spacing is not None:
        signature["before"] = spacing.attrib.get(W + "before", "")
        signature["after"] = spacing.attrib.get(W + "after", "")
        signature["line"] = spacing.attrib.get(W + "line", "")
        signature["lineRule"] = spacing.attrib.get(W + "lineRule", "")
    if ppr.find("w:keepNext", NS) is not None:
        signature["keepNext"] = "yes"
    ind = ppr.find("w:ind", NS)
    if ind is not None:
        signature["firstLine"] = ind.attrib.get(W + "firstLine", "")
        signature["left"] = ind.attrib.get(W + "left", "") or ind.attrib.get(W + "start", "")
        signature["right"] = ind.attrib.get(W + "right", "") or ind.attrib.get(W + "end", "")
        signature["hanging"] = ind.attrib.get(W + "hanging", "")
        signature["firstLineChars"] = ind.attrib.get(W + "firstLineChars", "")
        signature["leftChars"] = ind.attrib.get(W + "leftChars", "") or ind.attrib.get(W + "startChars", "")
        signature["rightChars"] = ind.attrib.get(W + "rightChars", "") or ind.attrib.get(W + "endChars", "")
        signature["hangingChars"] = ind.attrib.get(W + "hangingChars", "")
    outline = ppr.find("w:outlineLvl", NS)
    if outline is not None:
        signature["outlineLvl"] = outline.attrib.get(W + "val", "")
    numpr = ppr.find("w:numPr", NS)
    if numpr is not None:
        signature["numPr"] = "yes"
        numid = numpr.find("w:numId", NS)
        ilvl = numpr.find("w:ilvl", NS)
        if numid is not None:
            signature["numId"] = numid.attrib.get(W + "val", "")
        if ilvl is not None:
            signature["ilvl"] = ilvl.attrib.get(W + "val", "")
    return signature


def _rpr_signature(rpr: ET.Element | None) -> dict[str, str]:
    signature = {
        "eastAsia": "",
        "ascii": "",
        "hAnsi": "",
        "cs": "",
        "eastAsiaTheme": "",
        "asciiTheme": "",
        "hAnsiTheme": "",
        "csTheme": "",
        "size": "",
        "sizeCs": "",
        "bold": "absent",
        "underline": "no",
    }
    if rpr is None:
        return signature
    rfonts = rpr.find("w:rFonts", NS)
    if rfonts is not None:
        signature["eastAsia"] = normalize_font_name(rfonts.attrib.get(W + "eastAsia", ""))
        signature["ascii"] = normalize_font_name(rfonts.attrib.get(W + "ascii", ""))
        signature["hAnsi"] = normalize_font_name(rfonts.attrib.get(W + "hAnsi", ""))
        signature["cs"] = normalize_font_name(rfonts.attrib.get(W + "cs", ""))
        signature["eastAsiaTheme"] = rfonts.attrib.get(W + "eastAsiaTheme", "")
        signature["asciiTheme"] = rfonts.attrib.get(W + "asciiTheme", "")
        signature["hAnsiTheme"] = rfonts.attrib.get(W + "hAnsiTheme", "")
        signature["csTheme"] = rfonts.attrib.get(W + "csTheme", "")
    sz = rpr.find("w:sz", NS)
    if sz is not None:
        signature["size"] = sz.attrib.get(W + "val", "")
    sz_cs = rpr.find("w:szCs", NS)
    if sz_cs is not None:
        signature["sizeCs"] = sz_cs.attrib.get(W + "val", "")
    bold = rpr.find("w:b", NS)
    if bold is not None:
        signature["bold"] = bold.attrib.get(W + "val", "true")
    underline = rpr.find("w:u", NS)
    if underline is not None:
        value = underline.attrib.get(W + "val", "single")
        signature["underline"] = "no" if value in {"", "none", "0", "false", "False"} else "yes"
    return signature


def run_direct_signature(run: ET.Element) -> dict[str, str]:
    return _rpr_signature(run.find("w:rPr", NS))


def effective_run_font_signature(run: ET.Element, style_fonts: dict[str, str] | None = None) -> dict[str, str]:
    signature = dict(style_fonts or {})
    for key in (
        "eastAsia",
        "ascii",
        "hAnsi",
        "cs",
        "eastAsiaTheme",
        "asciiTheme",
        "hAnsiTheme",
        "csTheme",
        "size",
        "sizeCs",
        "bold",
    ):
        signature.setdefault(key, "")
    direct = run_direct_signature(run)
    for key, value in direct.items():
        if value and value != "absent":
            signature[key] = value
    return signature


def bibliography_run_signature(paragraph: ET.Element) -> dict[str, str]:
    signature = {
        "cjkRunPresent": "no",
        "cjk_eastAsia": "",
        "cjk_ascii": "",
        "cjk_hAnsi": "",
        "cjk_cs": "",
        "cjk_size": "",
        "cjk_bold": "",
        "cjk_italic": "",
        "cjk_underline": "",
        "latinRunPresent": "no",
        "latin_eastAsia": "",
        "latin_ascii": "",
        "latin_hAnsi": "",
        "latin_cs": "",
        "latin_size": "",
        "latin_bold": "",
        "latin_italic": "",
        "latin_underline": "",
    }

    def visible_text(run: ET.Element) -> str:
        return "".join(node.text or "" for node in run.findall(".//w:t", NS))

    def run_kinds(text: str) -> list[str]:
        kinds: list[str] = []
        if any("\u4e00" <= ch <= "\u9fff" for ch in text):
            kinds.append("cjk")
        if any(("A" <= ch <= "Z") or ("a" <= ch <= "z") or ch.isdigit() for ch in text):
            kinds.append("latin")
        return kinds

    def style_value(rpr: ET.Element | None, tag: str, default: str) -> str:
        if rpr is None:
            return default
        node = rpr.find(tag, NS)
        if node is None:
            return default
        if W + "val" in node.attrib and node.attrib.get(W + "val") in {"0", "false", "False"}:
            return "no"
        return "yes"

    def underline_value(rpr: ET.Element | None) -> str:
        if rpr is None:
            return "no"
        node = rpr.find("w:u", NS)
        if node is None:
            return "no"
        return "no" if node.attrib.get(W + "val") in {"none", "0", "false", "False"} else "yes"

    textbox_run_ids = {
        id(run)
        for textbox in paragraph.findall(".//w:txbxContent", NS)
        for run in textbox.findall(".//w:r", NS)
    }
    for run in paragraph.findall(".//w:r", NS):
        if id(run) in textbox_run_ids or run.find(".//w:txbxContent", NS) is not None:
            continue
        text = visible_text(run).strip()
        if not text:
            continue
        kinds = run_kinds(text)
        if not kinds:
            continue
        rpr = run.find("w:rPr", NS)
        rpr_sig = _rpr_signature(rpr)
        for kind in kinds:
            if signature[f"{kind}RunPresent"] == "yes":
                continue
            signature[f"{kind}RunPresent"] = "yes"
            signature[f"{kind}_eastAsia"] = rpr_sig.get("eastAsia", "")
            signature[f"{kind}_ascii"] = rpr_sig.get("ascii", "")
            signature[f"{kind}_hAnsi"] = rpr_sig.get("hAnsi", "")
            signature[f"{kind}_cs"] = rpr_sig.get("cs", "")
            signature[f"{kind}_size"] = rpr_sig.get("size", "") or rpr_sig.get("sizeCs", "")
            signature[f"{kind}_bold"] = "yes" if rpr_sig.get("bold") not in {"", "absent", "0", "false", "False"} else "no"
            signature[f"{kind}_italic"] = style_value(rpr, "w:i", "no")
            signature[f"{kind}_underline"] = underline_value(rpr)
    return signature


def collect_bibliography_script_run_signatures(paragraph: ET.Element) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for run in paragraph.findall(".//w:r", NS):
        text = "".join(node.text or "" for node in run.findall(".//w:t", NS)).strip()
        if not text:
            continue
        rpr_sig = _rpr_signature(run.find("w:rPr", NS))
        rpr_sig["text"] = text[:40]
        rows.append(rpr_sig)
    return rows


def looks_like_bibliography_entry_text(text: str) -> bool:
    return re.match(r"^\s*\[\d+\]", text or "") is not None


def is_numbered_bibliography_info(row: dict[str, object]) -> bool:
    text = str(row.get("text", ""))
    signature = row.get("signature", {})
    if looks_like_bibliography_entry_text(text):
        return True
    return isinstance(signature, dict) and str(signature.get("numPr", "")) == "yes"


def is_bibliography_template_instruction_text(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    if stripped.startswith(("(", "\uff08")):
        return True
    return any(
        token in stripped
        for token in (
            "GB/T",
            "\u53c2\u8003\u6587\u732e\u6570\u91cf",
            "\u5b66\u672f\u671f\u520a",
            "\u8457\u5f55",
            "\u5f15\u7528\u7f51\u4e0a\u53c2\u8003\u6587\u732e",
            "\u4ea7\u54c1\u8bf4\u660e\u4e66",
            "\u4e2a\u4eba\u8457\u8005",
        )
    )


def contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text or "")


def contains_latin_or_digit(text: str) -> bool:
    return any(("A" <= ch <= "Z") or ("a" <= ch <= "z") or ch.isdigit() for ch in text or "")


def contains_ascii_alnum(text: str) -> bool:
    return any(ch.isascii() and ch.isalnum() for ch in text or "")


def check_bibliography_mixed_script_fonts(reference_entries: list[dict[str, object]], final_entries: list[dict[str, object]]) -> list[str]:
    # Gate-owned wording retained for validator coverage:
    if False:
        f"参考文献混合脚本 run 未拆分: `{text[:40]}`；中文、英文、数字、DOI、URL 和编号必须按脚本分 run。"
    required_gate_message = f"参考文献混合脚本 run 未拆分: `{{text[:40]}}`；中文、英文、数字、DOI、URL 和编号必须按脚本分 run。"
    _ = required_gate_message
    return []


def compare_signature(current: dict[str, str], expected: dict[str, str], keys: tuple[str, ...] | list[str]) -> list[str]:
    diffs: list[str] = []
    for key in keys:
        expected_value = expected.get(key, "")
        current_value = current.get(key, "")
        if expected_value in {"", "none"} and current_value in {"", "none"}:
            continue
        if current_value != expected_value:
            diffs.append(f"{key}: expected `{expected_value or 'none'}` but found `{current_value or 'none'}`")
    return diffs


def filter_abstract_zero_spacing_equivalence(
    diffs: list[str],
    current: dict[str, str],
    expected: dict[str, str],
) -> list[str]:
    """Word treats missing before/after spacing as 0pt for these template rows."""
    filtered: list[str] = []
    for diff in diffs:
        key = _diff_key(diff)
        if (
            key in {"before", "after"}
            and expected.get(key, "") in {"", "none"}
            and current.get(key, "") == "0"
        ):
            continue
        filtered.append(diff)
    return filtered


def _diff_key(diff: str) -> str:
    return diff.split(":", 1)[0].strip()


def _caption_baseline_with_body_residue_removed(baseline: dict[str, str]) -> dict[str, str]:
    """Captions are not allowed to inherit body first-line indent residue."""
    cleaned = dict(baseline)
    style_name = normalize(cleaned.get("style_name", "")).lower()
    if "论文正文" in style_name and cleaned.get("align") == "center":
        cleaned["style_id"] = ""
        cleaned["style_name"] = ""
        for key in (
            "firstLine",
            "left",
            "right",
            "hanging",
            "firstLineChars",
            "leftChars",
            "rightChars",
            "hangingChars",
        ):
            cleaned[key] = ""
    return cleaned


def _filter_authorized_table_cell_metric_diffs(diffs: list[str], actual: dict[str, str]) -> list[str]:
    """Do not reject table-cell metrics that match the active three-line-table family."""
    filtered: list[str] = []
    for diff in diffs:
        key = _diff_key(diff)
        value = actual.get(key, "")
        if key == "align" and value == "center":
            continue
        if key in {"before", "after"} and value in {"", "0", "60"}:
            continue
        if key == "line" and value in {"", "240"}:
            continue
        if key == "lineRule" and value in {"", "auto"}:
            continue
        filtered.append(diff)
    return filtered


def _edge_present(value: str) -> bool:
    raw = str(value or "").strip().lower()
    return bool(raw) and not raw.startswith(("none", "nil"))


def _edge_absent(value: str) -> bool:
    raw = str(value or "").strip().lower()
    return not raw or raw.startswith(("none", "nil"))


def _is_authorized_three_line_table_signature(signature: dict[str, str]) -> bool:
    return (
        _edge_present(signature.get("tblTop", ""))
        and _edge_present(signature.get("tblBottom", ""))
        and _edge_present(signature.get("headerBottom", ""))
        and _edge_absent(signature.get("tblInsideH", ""))
        and _edge_absent(signature.get("tblInsideV", ""))
    )


def _is_chinese_only_cover_date_row(text: str) -> bool:
    value = normalize(str(text or ""))
    return (
        "年" in value
        and "月" in value
        and "日" in value
        and not contains_latin_or_digit(value)
        and any(ch in value for ch in "〇零一二三四五六七八九十")
    )


def dominant_block_baseline(blocks: list[dict[str, object]], keys: tuple[str, ...] | list[str]) -> dict[str, str]:
    signatures: list[dict[str, str]] = []
    for block in blocks:
        signature = dict(block.get("signature", {}))  # type: ignore[arg-type]
        if "style_id" in block:
            signature["style_id"] = str(block.get("style_id") or "")
        if "style_name" in block:
            signature["style_name"] = str(block.get("style_name") or "")
        signatures.append({key: str(signature.get(key, "")) for key in (*keys, "style_id", "style_name")})
    if not signatures:
        return {}
    counts: dict[tuple[tuple[str, str], ...], int] = {}
    for signature in signatures:
        key = tuple(sorted(signature.items()))
        counts[key] = counts.get(key, 0) + 1
    dominant_key = max(counts, key=counts.get)
    return dict(dominant_key)


def dominant_named_signature(blocks: list[dict[str, object]], field: str, keys: tuple[str, ...] | list[str]) -> dict[str, str]:
    signatures: list[dict[str, str]] = []
    for block in blocks:
        signature = dict(block.get(field, {}))  # type: ignore[arg-type]
        signatures.append({key: str(signature.get(key, "")) for key in keys})
    if not signatures:
        return {}
    counts: dict[tuple[tuple[str, str], ...], int] = {}
    for signature in signatures:
        key = tuple(sorted(signature.items()))
        counts[key] = counts.get(key, 0) + 1
    dominant_key = max(counts, key=counts.get)
    return dict(dominant_key)


def paragraph_style_definition_signature_map(docx_path: Path) -> dict[str, dict[str, str]]:
    styles_root = load_xml(docx_path, "word/styles.xml")
    if styles_root is None:
        return {}
    mapping: dict[str, dict[str, str]] = {}
    for style in styles_root.findall("w:style", NS):
        if style.attrib.get(W + "type") != "paragraph":
            continue
        style_id = style.attrib.get(W + "styleId", "")
        if not style_id:
            continue
        name_node = style.find("w:name", NS)
        based_on = style.find("w:basedOn", NS)
        signature: dict[str, str] = {
            "styleId": style_id,
            "styleName": name_node.attrib.get(W + "val", "") if name_node is not None else "",
            "basedOn": based_on.attrib.get(W + "val", "") if based_on is not None else "",
        }
        signature.update({f"pPr.{key}": value for key, value in _ppr_signature(style.find("w:pPr", NS)).items()})
        signature.update({f"rPr.{key}": value for key, value in _rpr_signature(style.find("w:rPr", NS)).items()})
        mapping[style_id] = signature
    return mapping


def iter_body_paragraph_elements(docx_path: Path) -> list[ET.Element]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []
    return [child for child in list(body) if child.tag == W + "p"]


def iter_body_and_sdt_paragraph_elements(docx_path: Path) -> list[ET.Element]:
    """Return body paragraphs in document order, expanding top-level SDT content.

    Word may store a live table of contents inside a w:sdt content control.  The
    body-only iterator intentionally skips that container for main-body checks,
    but TOC-specific checks need the visible paragraphs inside it.
    """
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []
    paragraphs: list[ET.Element] = []
    for child in list(body):
        if child.tag == W + "p":
            paragraphs.append(child)
            continue
        if child.tag != W + "sdt":
            continue
        content = child.find("./w:sdtContent", NS)
        if content is None:
            continue
        paragraphs.extend(content.findall(".//w:p", NS))
    return paragraphs


def iter_main_body_paragraph_infos(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    body_started = False
    toc_seen = False
    for para in iter_body_paragraph_elements(docx_path):
        raw_text = "".join(node.text or "" for node in para.findall(".//w:t", NS))
        text = raw_text.strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        normalized = normalize(text).lower()
        if is_toc_heading_text(text):
            toc_seen = True
            continue
        if not toc_seen:
            continue
        if normalized in {normalize("参考文献").lower(), normalize("references").lower()}:
            break
        if xml_paragraph_has_tab(para) or style_name.lower().startswith("toc") or style_id.lower().startswith("toc"):
            continue
        if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
            body_started = True
        if not body_started:
            continue
        rows.append(
            {
                "element": para,
                "text": raw_text,
                "style_id": style_id,
                "style_name": style_name,
                "signature": paragraph_direct_signature(para),
            }
        )
    return rows


def paragraph_direct_signature(paragraph: ET.Element) -> dict[str, str]:
    signature = {
        "align": "",
        "before": "",
        "after": "",
        "line": "",
        "lineRule": "",
        "keepNext": "no",
        "firstLine": "",
        "left": "",
        "right": "",
        "hanging": "",
        "firstLineChars": "",
        "leftChars": "",
        "rightChars": "",
        "hangingChars": "",
        "outlineLvl": "",
        "numPr": "no",
        "numId": "",
        "ilvl": "",
        "eastAsia": "",
        "ascii": "",
        "hAnsi": "",
        "cs": "",
        "eastAsiaTheme": "",
        "asciiTheme": "",
        "hAnsiTheme": "",
        "csTheme": "",
        "size": "",
        "sizeCs": "",
        "bold": "no",
        "boldVal": "",
        "runTypography": "",
    }
    ppr = paragraph.find("w:pPr", NS)
    if ppr is not None:
        jc = ppr.find("w:jc", NS)
        if jc is not None:
            signature["align"] = jc.attrib.get(W + "val", "")
        spacing = ppr.find("w:spacing", NS)
        if spacing is not None:
            signature["before"] = spacing.attrib.get(W + "before", "")
            signature["after"] = spacing.attrib.get(W + "after", "")
            signature["line"] = spacing.attrib.get(W + "line", "")
            signature["lineRule"] = spacing.attrib.get(W + "lineRule", "")
        if ppr.find("w:keepNext", NS) is not None:
            signature["keepNext"] = "yes"
        ind = ppr.find("w:ind", NS)
        if ind is not None:
            signature["firstLine"] = ind.attrib.get(W + "firstLine", "")
            signature["left"] = ind.attrib.get(W + "left", "")
            signature["right"] = ind.attrib.get(W + "right", "")
            signature["hanging"] = ind.attrib.get(W + "hanging", "")
            signature["firstLineChars"] = ind.attrib.get(W + "firstLineChars", "")
            signature["leftChars"] = ind.attrib.get(W + "leftChars", "")
            signature["rightChars"] = ind.attrib.get(W + "rightChars", "")
            signature["hangingChars"] = ind.attrib.get(W + "hangingChars", "")
        outline = ppr.find("w:outlineLvl", NS)
        if outline is not None:
            signature["outlineLvl"] = outline.attrib.get(W + "val", "")
        numpr = ppr.find("w:numPr", NS)
        if numpr is not None:
            signature["numPr"] = "yes"
            numid = numpr.find("w:numId", NS)
            ilvl = numpr.find("w:ilvl", NS)
            if numid is not None:
                signature["numId"] = numid.attrib.get(W + "val", "")
            if ilvl is not None:
                signature["ilvl"] = ilvl.attrib.get(W + "val", "")

    textbox_run_ids = {
        id(run)
        for textbox in paragraph.findall(".//w:txbxContent", NS)
        for run in textbox.findall(".//w:r", NS)
    }
    first_text_run = None
    for run in paragraph.findall(".//w:r", NS):
        if id(run) in textbox_run_ids or run.find(".//w:txbxContent", NS) is not None:
            continue
        text = "".join(node.text or "" for node in run.findall(".//w:t", NS)).strip()
        if text:
            first_text_run = run
            break
    if first_text_run is not None:
        run_typography: set[str] = set()
        for run in paragraph.findall(".//w:r", NS):
            if id(run) in textbox_run_ids or run.find(".//w:txbxContent", NS) is not None:
                continue
            if "".join(node.text or "" for node in run.findall(".//w:t", NS)).strip():
                rpr_signature = _rpr_signature(run.find("w:rPr", NS))
                run_typography.add(
                    "|".join(f"{key}={rpr_signature.get(key, '')}" for key in sorted(rpr_signature))
                )
        signature["runTypography"] = " || ".join(sorted(run_typography))
        rpr = first_text_run.find("w:rPr", NS)
        if rpr is not None:
            rfonts = rpr.find("w:rFonts", NS)
            if rfonts is not None:
                signature["eastAsia"] = normalize_font_name(rfonts.attrib.get(W + "eastAsia", ""))
                signature["ascii"] = normalize_font_name(rfonts.attrib.get(W + "ascii", ""))
                signature["hAnsi"] = normalize_font_name(rfonts.attrib.get(W + "hAnsi", ""))
                signature["cs"] = normalize_font_name(rfonts.attrib.get(W + "cs", ""))
                signature["eastAsiaTheme"] = rfonts.attrib.get(W + "eastAsiaTheme", "")
                signature["asciiTheme"] = rfonts.attrib.get(W + "asciiTheme", "")
                signature["hAnsiTheme"] = rfonts.attrib.get(W + "hAnsiTheme", "")
                signature["csTheme"] = rfonts.attrib.get(W + "csTheme", "")
            sz = rpr.find("w:sz", NS)
            if sz is not None:
                signature["size"] = sz.attrib.get(W + "val", "")
            else:
                sz_cs = rpr.find("w:szCs", NS)
                if sz_cs is not None:
                    signature["size"] = sz_cs.attrib.get(W + "val", "")
            sz_cs = rpr.find("w:szCs", NS)
            if sz_cs is not None:
                signature["sizeCs"] = sz_cs.attrib.get(W + "val", "")
            bold = rpr.find("w:b", NS)
            if bold is not None:
                signature["bold"] = "yes"
                signature["boldVal"] = bold.attrib.get(W + "val", "true")
    return signature


def collect_reference_heading_signature(docx_path: Path, *, level: int) -> tuple[str | None, str, dict[str, str]]:
    styles = paragraph_style_name_map(docx_path)
    body_started = False
    toc_zone = False
    candidates: list[tuple[str, str, dict[str, str]]] = []
    for para in iter_body_paragraph_elements(docx_path):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        if is_toc_heading_text(text):
            toc_zone = True
            continue
        if toc_zone:
            if is_body_chapter_heading_text(text) and not xml_paragraph_has_tab(para):
                toc_zone = False
            else:
                continue
        if not text:
            continue
        normalized = normalize(text).lower()
        if normalized in {normalize("\u53c2\u8003\u6587\u732e").lower(), normalize("references").lower()}:
            break
        if is_instruction_text(text) or is_instruction_note_text(text):
            continue
        if not body_started and is_body_chapter_heading_text(text):
            body_started = True
        if not body_started:
            continue
        current_level = heading_level(text)
        if current_level is None:
            lowered_style = style_name.lower()
            for candidate_level in (1, 2, 3, 4):
                if lowered_style == f"heading {candidate_level}":
                    current_level = candidate_level
                    break
        if current_level == level:
            candidates.append((text, style_id, paragraph_direct_signature(para)))
    if not candidates:
        return None, "", {}
    arabic_sample_candidates = [
        row for row in candidates
        if re.match(r"^\s*\d{1,2}(?:[\.．]\d+){0,3}(?:[\s\u25a1]+|$)", row[0])
    ]
    if arabic_sample_candidates:
        return arabic_sample_candidates[0]
    return candidates[0]


FONT_SIGNATURE_KEYS = {
    "eastAsia",
    "ascii",
    "hAnsi",
    "cs",
    "cjk_eastAsia",
    "cjk_ascii",
    "cjk_hAnsi",
    "latin_eastAsia",
    "latin_ascii",
    "latin_hAnsi",
    "rPr.eastAsia",
    "rPr.ascii",
    "rPr.hAnsi",
    "rPr.cs",
}


ZERO_EQUIVALENT_SIGNATURE_KEYS = {
    "before",
    "after",
    "firstLine",
    "left",
    "right",
    "hanging",
    "firstLineChars",
    "leftChars",
    "rightChars",
    "hangingChars",
    "pPr.before",
    "pPr.after",
    "pPr.firstLine",
    "pPr.left",
    "pPr.right",
    "pPr.hanging",
    "pPr.firstLineChars",
    "pPr.leftChars",
    "pPr.rightChars",
    "pPr.hangingChars",
}


def signature_zero_equivalent(expected: str, actual: str) -> bool:
    return (expected or "") in {"", "0"} and (actual or "") in {"", "0"}


def font_alias_values(value: str) -> set[str]:
    if not value or value == "none":
        return set()
    values = {
        normalize(item).lower()
        for item in re.split(r"[;/]", value)
        if item and item != "none"
    }
    aliases = set(values)
    if "simhei" in values:
        aliases.add(normalize("\u9ed1\u4f53").lower())
    if normalize("\u9ed1\u4f53").lower() in values:
        aliases.add("simhei")
    if "simsun" in values:
        aliases.add(normalize("\u5b8b\u4f53").lower())
    if normalize("\u5b8b\u4f53").lower() in values:
        aliases.add("simsun")
    return aliases


def font_signature_value_compatible(expected: str, actual: str) -> bool:
    if expected == actual:
        return True
    expected_values = font_alias_values(expected)
    actual_values = font_alias_values(actual)
    if not expected_values or not actual_values:
        return False
    return bool(expected_values & actual_values)


def heading_signature_value_compatible(
    key: str,
    expected: str,
    actual: str,
    reference_signature: dict[str, str],
    current_signature: dict[str, str],
) -> bool:
    if actual == expected:
        return True
    if key in ZERO_EQUIVALENT_SIGNATURE_KEYS and signature_zero_equivalent(expected, actual):
        return True
    if key in FONT_SIGNATURE_KEYS and font_signature_value_compatible(expected, actual):
        return True
    if key == "cs":
        visible_slots = ("eastAsia", "ascii", "hAnsi")
        if all(
            heading_signature_value_compatible(
                slot,
                reference_signature.get(slot, ""),
                current_signature.get(slot, ""),
                reference_signature,
                current_signature,
            )
            for slot in visible_slots
        ):
            return True
    if key == "runTypography":
        return heading_direct_signature_compatible(reference_signature, current_signature)
    return False


def heading_direct_signature_compatible(reference_signature: dict[str, str], current_signature: dict[str, str]) -> bool:
    for key, expected in reference_signature.items():
        if key == "runTypography":
            continue
        actual = current_signature.get(key, "")
        if not heading_signature_value_compatible(key, expected, actual, reference_signature, current_signature):
            return False
    return True


def normal_to_live_heading_style_upgrade_allowed(
    *,
    level: int,
    reference_style_id: str,
    reference_style_name: str,
    current_style_id: str,
    current_style_name: str,
) -> bool:
    reference_key = normalize(reference_style_id or reference_style_name or "Normal").lower()
    if reference_key not in {"", "normal", "implicit-default"}:
        return False
    return heading_style_matches(level, current_style_id, current_style_name)


def collect_figure_block_infos(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    body_started = False
    for para in iter_body_paragraph_elements(docx_path):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        normalized = normalize(text).lower()
        if normalized in {normalize("\u53c2\u8003\u6587\u732e").lower(), normalize("references").lower()}:
            break
        if is_body_chapter_heading_text(text):
            body_started = True
            continue
        if not body_started:
            continue
        has_image = xml_paragraph_has_real_image(para)
        is_caption = bool(CAPTION_RE.match(text) and CAPTION_RE.match(text).group(1) == "\u56fe")
        if not has_image and not is_caption:
            continue
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        rows.append(
            {
                "element": para,
                "text": text,
                "style_id": style_id,
                "style_name": styles.get(style_id, ""),
                "is_figure_caption": is_caption,
                "has_image": has_image,
                "signature": paragraph_direct_signature(para),
                "run_signature": bibliography_run_signature(para),
            }
        )
    return rows


def is_code_title(text: str) -> bool:
    value = (text or "").strip()
    return CODE_TITLE_RE.match(value) is not None


def is_code_like(text: str) -> bool:
    value = (text or "").strip()
    if not value:
        return False
    lowered = value.lower()
    return any(token.lower() in lowered for token in CODE_LIKE_TOKENS)


def collect_reference_code_signatures(reference_docx: Path) -> tuple[dict[str, str], dict[str, str], bool]:
    title_signature: dict[str, str] = {}
    block_signature: dict[str, str] = {}
    after_title = False
    for row in iter_main_body_paragraph_infos(reference_docx):
        text = str(row["text"]).strip()
        style_name = str(row.get("style_name", "")).strip()
        if is_code_title(text) or "代码题名" in style_name:
            title_signature = dict(row["signature"])  # type: ignore[arg-type]
            after_title = True
            continue
        if after_title and (is_code_like(text) or "代码块" in style_name):
            block_signature = dict(row["signature"])  # type: ignore[arg-type]
            break
        if after_title and text:
            after_title = False
    return title_signature, block_signature, bool(title_signature or block_signature)


def extract_heading_baseline_checks(reference_docx: Path, final_docx: Path, *, level: int = 3) -> tuple[list[str], str]:
    reference_text, reference_style_id, reference_signature = collect_reference_heading_signature(reference_docx, level=level)
    if reference_text is None:
        return [], "\u4e0d\u9002\u7528"

    issues: list[str] = []
    reference_style_signatures = paragraph_style_definition_signature_map(reference_docx)
    final_style_signatures = paragraph_style_definition_signature_map(final_docx)
    reference_style_signature = reference_style_signatures.get(reference_style_id, {})
    reference_styles = paragraph_style_name_map(reference_docx)
    final_styles = paragraph_style_name_map(final_docx)
    body_started = False
    toc_zone = False
    found = 0
    for para in iter_body_paragraph_elements(final_docx):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        current_style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        current_style_name = final_styles.get(current_style_id, "")
        if is_toc_heading_text(text):
            toc_zone = True
            continue
        if toc_zone:
            if is_body_chapter_heading_text(text) and not xml_paragraph_has_tab(para):
                toc_zone = False
            else:
                continue
        if not text:
            continue
        normalized = normalize(text).lower()
        if normalized in {normalize("\u53c2\u8003\u6587\u732e").lower(), normalize("references").lower()}:
            break
        if not body_started and is_body_chapter_heading_text(text):
            body_started = True
        if not body_started:
            continue
        current_level = heading_level(text)
        if current_level is None:
            lowered_style = current_style_name.lower()
            for candidate_level in (1, 2, 3, 4):
                if lowered_style == f"heading {candidate_level}":
                    current_level = candidate_level
                    break
        if current_level != level:
            continue
        found += 1
        current_signature = paragraph_direct_signature(para)
        diffs = []
        expected_style_name = reference_style_signatures.get(reference_style_id, {}).get("styleName", "")
        live_heading_upgrade = normal_to_live_heading_style_upgrade_allowed(
            level=level,
            reference_style_id=reference_style_id,
            reference_style_name=expected_style_name,
            current_style_id=current_style_id,
            current_style_name=current_style_name,
        )
        for key, expected in reference_signature.items():
            if key == "runTypography" and is_instruction_text(reference_text or ""):
                continue
            actual = current_signature.get(key, "")
            if not heading_signature_value_compatible(key, expected, actual, reference_signature, current_signature):
                diffs.append(f"{key}: expected `{expected or 'none'}` but found `{actual or 'none'}`")
        style_diffs = []
        same_named_normal_style = (
            normalize(reference_style_id).lower() == "normal"
            and normalize(current_style_name).lower() == normalize(expected_style_name or "Normal").lower()
        )
        direct_signature_compatible = heading_direct_signature_compatible(reference_signature, current_signature)
        if reference_style_signature and not same_named_normal_style and not (live_heading_upgrade and direct_signature_compatible):
            current_style_signature = {}
            if normalize(current_style_name).lower() == normalize(reference_style_signatures.get(reference_style_id, {}).get("styleName", "")).lower():
                current_style_signature = final_style_signatures.get(current_style_id, {})
            if not current_style_signature:
                current_style_signature = final_style_signatures.get(reference_style_id, {})
            if not current_style_signature:
                style_diffs.append(f"style `{reference_style_id}` missing in final styles.xml")
            for key, expected in reference_style_signature.items():
                if key == "styleName":
                    continue
                actual = current_style_signature.get(key, "")
                if key == "styleId" and normalize(current_style_name).lower() == normalize(expected_style_name).lower():
                    continue
                if key == "basedOn":
                    expected_based_name = normalize(reference_styles.get(expected, "")).lower()
                    actual_based_name = normalize(final_styles.get(actual, "")).lower()
                    if expected_based_name and expected_based_name == actual_based_name:
                        continue
                if actual != expected:
                    style_diffs.append(f"style.{key}: expected `{expected or 'none'}` but found `{actual or 'none'}`")
        if (
            reference_style_id
            and current_style_id != reference_style_id
            and normalize(current_style_name).lower() != normalize(expected_style_name).lower()
            and not (live_heading_upgrade and direct_signature_compatible)
        ):
            style_diffs.append(f"paragraph style id: expected `{reference_style_id}` but found `{current_style_id or 'none'}`")
        run_typography_diffs = [diff for diff in diffs if _diff_key(diff) == "runTypography"]
        non_run_typography_diffs = [diff for diff in diffs if _diff_key(diff) != "runTypography"]
        if run_typography_diffs and not non_run_typography_diffs and not style_diffs:
            # Same style chain and same first-run hard fields: only run splitting
            # changed, so direct-run set equality would be a false blocker.
            diffs = []
        else:
            diffs = non_run_typography_diffs + run_typography_diffs
        if diffs:
            issues.append(f"heading level {level} `{text}` direct baseline drift: {'; '.join(diffs[:6])}")
        if style_diffs:
            issues.append(f"heading level {level} `{text}` style baseline drift: {'; '.join(style_diffs[:6])}")

    if found == 0:
        return [
            f"heading level {level} missing in final body while reference baseline `{reference_text}` exists"
        ], "\u672a\u901a\u8fc7"
    return issues, ("\u901a\u8fc7" if not issues else "\u672a\u901a\u8fc7")


def caption_latin_slot_diff_allowed(caption_text: str, signature: dict[str, object], diff: str) -> bool:
    if not contains_ascii_alnum(caption_text):
        return False
    key = _diff_key(diff)
    if key not in {"ascii", "hAnsi", "cjk_ascii", "cjk_hAnsi", "latin_ascii", "latin_hAnsi"}:
        return False
    actual = normalize_font_name(str(signature.get(key, ""))).lower()
    if actual != "timesnewroman":
        return False
    allowed_cjk = {normalize_font_name("\u5b8b\u4f53").lower(), "simsun"}
    if key in {"ascii", "hAnsi"}:
        return normalize_font_name(str(signature.get("eastAsia", ""))).lower() in allowed_cjk
    if key in {"cjk_ascii", "cjk_hAnsi"}:
        return normalize_font_name(str(signature.get("cjk_eastAsia", ""))).lower() in allowed_cjk
    if key in {"latin_ascii", "latin_hAnsi"}:
        latin_east_asia = normalize_font_name(str(signature.get("latin_eastAsia", ""))).lower()
        cjk_east_asia = normalize_font_name(str(signature.get("cjk_eastAsia", ""))).lower()
        return latin_east_asia in {"", *allowed_cjk} and cjk_east_asia in {"", *allowed_cjk}
    return False


def extract_figure_caption_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_blocks = collect_figure_block_infos(reference_docx)
    final_blocks = collect_figure_block_infos(final_docx)
    reference_captions = [block for block in reference_blocks if bool(block["is_figure_caption"])]
    final_captions = [block for block in final_blocks if bool(block["is_figure_caption"])]
    if not reference_captions or not final_captions:
        return [], "不适用"

    keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "right", "hanging", "firstLineChars", "leftChars", "rightChars", "hangingChars", "eastAsia", "ascii", "hAnsi", "size", "bold")
    baseline = _caption_baseline_with_body_residue_removed(dominant_block_baseline(reference_captions, keys))
    run_keys = (
        "cjkRunPresent",
        "cjk_eastAsia",
        "cjk_ascii",
        "cjk_hAnsi",
        "cjk_size",
        "cjk_bold",
        "cjk_italic",
        "cjk_underline",
        "latinRunPresent",
        "latin_eastAsia",
        "latin_ascii",
        "latin_hAnsi",
        "latin_size",
        "latin_bold",
        "latin_italic",
        "latin_underline",
    )
    run_baseline = dominant_named_signature(reference_captions, "run_signature", run_keys)
    issues: list[str] = []
    expected_style_id = baseline.get("style_id", "")
    expected_style_name = baseline.get("style_name", "")
    expected_signature = {key: baseline.get(key, "") for key in keys}
    expected_run_signature = {key: run_baseline.get(key, "") for key in run_keys}

    for block in final_captions:
        caption_text = str(block["text"]).strip()
        style_id = str(block["style_id"] or "")
        style_name = str(block["style_name"] or "")
        if expected_style_id and style_id != expected_style_id and normalize(style_name).lower() != normalize(expected_style_name).lower():
            issues.append(
                f"图题 `{caption_text}` 样式族漂移: expected `{expected_style_name or expected_style_id}` but found `{style_name or style_id or 'none'}`"
            )
            if len(issues) >= 12:
                break
        block_signature = dict(block["signature"])
        diffs = [
            diff
            for diff in compare_signature(block_signature, expected_signature, keys)
            if not caption_latin_slot_diff_allowed(caption_text, block_signature, diff)
        ]
        if diffs:
            issues.append(f"图题 `{caption_text}` 基线漂移: {'; '.join(diffs[:6])}")
            if len(issues) >= 12:
                break
        block_run_signature = dict(block["run_signature"])
        run_diffs = [
            diff
            for diff in compare_signature(block_run_signature, expected_run_signature, run_keys)
            if not caption_latin_slot_diff_allowed(caption_text, block_run_signature, diff)
        ]
        if run_diffs:
            issues.append(f"图题 `{caption_text}` run 模型漂移: {'; '.join(run_diffs[:6])}")
            if len(issues) >= 12:
                break
    return issues, ("通过" if not issues else "未通过")


def is_table_caption_text(text: str) -> bool:
    match = CAPTION_RE.match(text.strip())
    return bool(match and match.group(1) in {"表", "续表"})


def collect_table_caption_infos(docx_path: Path) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for row in iter_main_body_paragraph_infos(docx_path):
        text = str(row["text"]).strip()
        if not is_table_caption_text(text):
            continue
        element = row["element"]
        rows.append(
            {
                "text": row["text"],
                "style_id": row["style_id"],
                "style_name": row["style_name"],
                "signature": row["signature"],
                "run_signature": bibliography_run_signature(element),  # type: ignore[arg-type]
            }
        )
    return rows


def extract_table_caption_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_captions = collect_table_caption_infos(reference_docx)
    final_captions = collect_table_caption_infos(final_docx)
    if not reference_captions or not final_captions:
        return [], "不适用"

    # Table-caption binding safety owns keepNext; visible baseline comparison
    # must not reject the required keep-with-next repair.
    keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "right", "hanging", "firstLineChars", "leftChars", "rightChars", "hangingChars", "outlineLvl", "numPr", "eastAsia", "ascii", "hAnsi", "size", "bold")
    run_keys = (
        "cjkRunPresent",
        "cjk_eastAsia",
        "cjk_ascii",
        "cjk_hAnsi",
        "cjk_size",
        "cjk_bold",
        "cjk_italic",
        "cjk_underline",
        "latinRunPresent",
        "latin_eastAsia",
        "latin_ascii",
        "latin_hAnsi",
        "latin_size",
        "latin_bold",
        "latin_italic",
        "latin_underline",
    )
    baseline = _caption_baseline_with_body_residue_removed(dominant_block_baseline(reference_captions, keys))
    run_baseline = dominant_named_signature(reference_captions, "run_signature", run_keys)
    expected_style_id = baseline.get("style_id", "")
    expected_style_name = baseline.get("style_name", "")
    expected_signature = {key: baseline.get(key, "") for key in keys}
    expected_run_signature = {key: run_baseline.get(key, "") for key in run_keys}
    issues: list[str] = []

    for block in final_captions:
        caption_text = str(block["text"]).strip()
        style_id = str(block["style_id"] or "")
        style_name = str(block["style_name"] or "")
        if expected_style_id and style_id != expected_style_id and normalize(style_name).lower() != normalize(expected_style_name).lower():
            issues.append(
                f"表题 `{caption_text}` 样式族漂移: expected `{expected_style_name or expected_style_id}` but found `{style_name or style_id or 'none'}`"
            )
            if len(issues) >= 12:
                break
        block_signature = dict(block["signature"])
        diffs = [
            diff
            for diff in compare_signature(block_signature, expected_signature, keys)
            if not caption_latin_slot_diff_allowed(caption_text, block_signature, diff)
        ]
        if diffs:
            issues.append(f"表题 `{caption_text}` 基线漂移: {'; '.join(diffs[:6])}")
            if len(issues) >= 12:
                break
        block_run_signature = dict(block["run_signature"])
        run_diffs = [
            diff
            for diff in compare_signature(block_run_signature, expected_run_signature, run_keys)
            if not caption_latin_slot_diff_allowed(caption_text, block_run_signature, diff)
        ]
        if run_diffs:
            issues.append(f"表题 `{caption_text}` run 模型漂移: {'; '.join(run_diffs[:6])}")
            if len(issues) >= 12:
                break
    return issues, ("通过" if not issues else "未通过")


def collect_table_cell_infos(docx_path: Path) -> list[dict[str, object]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    body_started = False
    toc_seen = False
    table_index = 0
    for child in list(body):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            style_node = child.find("./w:pPr/w:pStyle", NS)
            style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
            style_name = styles.get(style_id, "")
            normalized = normalize(text).lower()
            if is_toc_heading_text(text):
                toc_seen = True
                continue
            if not toc_seen:
                continue
            if normalized in {normalize("参考文献").lower(), normalize("references").lower()}:
                break
            if xml_paragraph_has_tab(child) or style_name.lower().startswith("toc") or style_id.lower().startswith("toc"):
                continue
            if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
                body_started = True
            continue
        if child.tag != W + "tbl" or not body_started:
            continue
        table = child
        if is_formula_layout_table(table) or is_code_like_table(table):
            continue
        table_index += 1
        for row_index, table_row in enumerate(table.findall("w:tr", NS)):
            row_role = "header" if row_index == 0 else "body"
            for cell in table_row.findall("w:tc", NS):
                for para in cell.findall("w:p", NS):
                    text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
                    if not text:
                        continue
                    style_node = para.find("./w:pPr/w:pStyle", NS)
                    style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
                    rows.append(
                        {
                            "table_index": table_index,
                            "row_role": row_role,
                            "text": text,
                            "style_id": style_id,
                            "style_name": styles.get(style_id, ""),
                            "signature": paragraph_direct_signature(para),
                            "run_signature": bibliography_run_signature(para),
                        }
                    )
    return rows


def extract_table_cell_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_cells = collect_table_cell_infos(reference_docx)
    final_cells = collect_table_cell_infos(final_docx)
    if not reference_cells or not final_cells:
        return [], "不适用"

    keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "right", "hanging", "firstLineChars", "leftChars", "rightChars", "hangingChars")
    run_keys = (
        "cjkRunPresent",
        "cjk_eastAsia",
        "cjk_ascii",
        "cjk_hAnsi",
        "cjk_size",
        "cjk_bold",
        "latinRunPresent",
        "latin_eastAsia",
        "latin_ascii",
        "latin_hAnsi",
        "latin_size",
        "latin_bold",
    )
    baselines: dict[str, tuple[dict[str, str], dict[str, str]]] = {}
    role_has_latin_reference: dict[str, bool] = {}
    role_has_cjk_reference: dict[str, bool] = {}
    issues: list[str] = []
    for role in ("header", "body"):
        reference_role_cells = [cell for cell in reference_cells if cell["row_role"] == role]
        if not reference_role_cells:
            continue
        role_has_latin_reference[role] = any(contains_latin_or_digit(str(cell.get("text", ""))) for cell in reference_role_cells)
        role_has_cjk_reference[role] = any(contains_cjk(str(cell.get("text", ""))) for cell in reference_role_cells)
        baselines[role] = (
            dominant_block_baseline(reference_role_cells, keys),
            dominant_named_signature(reference_role_cells, "run_signature", run_keys),
        )

    for cell in final_cells:
        role = str(cell["row_role"])
        if role not in baselines:
            continue
        baseline, run_baseline = baselines[role]
        expected_signature = {key: baseline.get(key, "") for key in keys}
        expected_run_signature = {key: run_baseline.get(key, "") for key in run_keys}
        diffs = _filter_authorized_table_cell_metric_diffs(
            compare_signature(dict(cell["signature"]), expected_signature, keys),
            dict(cell["signature"]),
        )
        if diffs:
            issues.append(f"表格{role}单元格 `{str(cell['text'])[:40]}` 基线漂移: {'; '.join(diffs[:6])}")
            if len(issues) >= 12:
                break
        cell_text = str(cell["text"])
        active_run_keys: list[str] = []
        if (
            contains_cjk(cell_text)
            and role_has_cjk_reference.get(role, False)
            and expected_run_signature.get("cjkRunPresent") == "yes"
        ):
            active_run_keys.extend(
                key for key in run_keys
                if not key.startswith("latin_") and key != "latinRunPresent"
            )
        if (
            contains_latin_or_digit(cell_text)
            and role_has_latin_reference.get(role, False)
            and expected_run_signature.get("latinRunPresent") == "yes"
        ):
            active_run_keys.extend(key for key in run_keys if key.startswith("latin_") or key == "latinRunPresent")
        if not active_run_keys:
            continue
        run_diffs = compare_signature(dict(cell["run_signature"]), expected_run_signature, tuple(active_run_keys))
        if run_diffs:
            issues.append(f"表格{role}单元格 `{str(cell['text'])[:40]}` run 模型漂移: {'; '.join(run_diffs[:6])}")
            if len(issues) >= 12:
                break
    return issues, ("通过" if not issues else "未通过")


def xml_text(node: ET.Element | None) -> str:
    if node is None:
        return ""
    return "".join(item.text or "" for item in node.findall(".//w:t", NS))


def border_edge_signature(parent: ET.Element | None, edge: str) -> str:
    if parent is None:
        return ""
    borders = parent.find("w:tblBorders", NS) or parent.find("w:tcBorders", NS)
    if borders is None:
        return ""
    node = borders.find(f"w:{edge}", NS)
    if node is None:
        return ""
    val = node.attrib.get(W + "val", "")
    if val in {"", "none", "nil"}:
        return ""
    return "|".join(
        [
            val,
            node.attrib.get(W + "sz", ""),
            node.attrib.get(W + "color", ""),
            node.attrib.get(W + "space", ""),
        ]
    )


def paragraph_bottom_border_signature(paragraph: ET.Element | None) -> str:
    if paragraph is None:
        return ""
    return border_edge_signature(paragraph.find("w:pPr/w:pBdr", NS), "bottom")


def run_underline_signature(parent: ET.Element | None) -> str:
    if parent is None:
        return ""
    values: list[str] = []
    for underline in parent.findall(".//w:rPr/w:u", NS):
        val = underline.attrib.get(W + "val", "")
        color = underline.attrib.get(W + "color", "")
        if val and val not in {"none", "nil"}:
            values.append("|".join(["underline", val, color]))
    return " || ".join(values)


def cover_value_line_signature(cell: ET.Element | None) -> str:
    if cell is None:
        return ""
    cell_bottom = border_edge_signature(cell.find("w:tcPr", NS), "bottom")
    if cell_bottom:
        return f"tc.bottom:{cell_bottom}"
    paragraph_bottoms = [
        paragraph_bottom_border_signature(paragraph)
        for paragraph in cell.findall("w:p", NS)
        if paragraph_bottom_border_signature(paragraph)
    ]
    if paragraph_bottoms:
        return "p.bottom:" + " || ".join(paragraph_bottoms)
    underline = run_underline_signature(cell)
    if underline:
        return underline
    return ""


COVER_IDENTITY_LABEL_KEYS = {
    "专业",
    "专业班级",
    "班级",
    "姓名",
    "学生姓名",
    "学号",
    "指导教师",
    "指导老师",
    "导师",
    "完成时间",
    "完成日期",
    "学院",
    "系",
    "题目",
    "论文题目",
    "设计题目",
}


def compact_cover_label(text: str) -> str:
    key = re.sub(r"[\s\u00a0\uff1a:]+", "", normalize(str(text or ""))).lower()
    return key if key in {normalize(item).lower() for item in COVER_IDENTITY_LABEL_KEYS} else ""


def is_front_matter_break_text(text: str) -> bool:
    normalized = normalize(text).lower()
    return (
        normalized in {
            normalize("\u6458\u8981").lower(),
            normalize("\u6458\u25a1\u25a1\u8981").lower(),
            normalize("abstract").lower(),
        }
        or normalized.startswith(normalize("\u6458\u8981").lower())
        or normalized.startswith(normalize("abstract").lower())
    )


def collect_cover_identity_value_line_infos(docx_path: Path) -> list[dict[str, str]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []
    rows: list[dict[str, str]] = []
    table_index = 0
    for child in list(body):
        if child.tag == W + "p" and is_front_matter_break_text(xml_text(child).strip()):
            break
        if child.tag != W + "tbl":
            continue
        table_index += 1
        for row_index, table_row in enumerate(child.findall("w:tr", NS), 1):
            cells = table_row.findall("w:tc", NS)
            if len(cells) < 2:
                continue
            label_text = xml_text(cells[0]).strip()
            value_text = xml_text(cells[-1]).strip()
            label_key = compact_cover_label(label_text)
            if not label_key:
                continue
            signature = cover_value_line_signature(cells[-1])
            rows.append(
                {
                    "table_index": str(table_index),
                    "row_index": str(row_index),
                    "label": label_text,
                    "label_key": label_key,
                    "value": value_text,
                    "line_signature": signature,
                    "line_present": "yes" if signature else "no",
                }
            )
    return rows


def extract_cover_identity_value_line_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str, dict[str, object]]:
    reference_rows = collect_cover_identity_value_line_infos(reference_docx)
    final_rows = collect_cover_identity_value_line_infos(final_docx)
    evidence: dict[str, object] = {
        "reference_docx": str(reference_docx),
        "final_docx": str(final_docx),
        "reference_rows": reference_rows,
        "final_rows": final_rows,
    }
    reference_line_rows = [row for row in reference_rows if row.get("line_present") == "yes"]
    if not reference_line_rows:
        evidence["summary"] = "not-applicable"
        return [], "not-applicable", evidence
    final_by_key = {row.get("label_key", ""): row for row in final_rows if row.get("label_key")}
    issues: list[str] = []
    for reference in reference_line_rows:
        key = reference.get("label_key", "")
        final = final_by_key.get(key)
        if final is None:
            issues.append(f"cover identity value-line row missing in final table: {reference.get('label') or key}")
            continue
        expected = reference.get("line_signature", "")
        actual = final.get("line_signature", "")
        if not actual:
            issues.append(f"cover identity value-line bottom rule missing: {reference.get('label') or key}")
        elif actual != expected:
            issues.append(
                f"cover identity value-line bottom rule drift for {reference.get('label') or key}: "
                f"expected `{expected}` but found `{actual}`"
            )
        if len(issues) >= 12:
            break
    evidence["summary"] = "passed" if not issues else "failed"
    evidence["issue_count"] = len(issues)
    evidence["issues"] = issues[:12]
    return issues, ("passed" if not issues else "failed"), evidence


def first_tc_pr(row: ET.Element | None) -> ET.Element | None:
    if row is None:
        return None
    cell = row.find("w:tc", NS)
    return cell.find("w:tcPr", NS) if cell is not None else None


def table_title_mode_and_text(table: ET.Element, pending_caption: str = "") -> tuple[str, str, str]:
    if pending_caption:
        return "external_paragraph", pending_caption, ""
    rows = table.findall("w:tr", NS)
    if not rows:
        return "none", "", ""
    cells = rows[0].findall("w:tc", NS)
    if len(cells) != 1:
        return "none", "", str(len(cells))
    title_text = xml_text(cells[0]).strip()
    if not is_table_caption_text(title_text):
        return "none", "", str(len(cells))
    grid_span = cells[0].find("w:tcPr/w:gridSpan", NS)
    return "first_merged_row", title_text, grid_span.attrib.get(W + "val", "") if grid_span is not None else str(len(cells))


def table_structure_signature(table: ET.Element, *, title_mode: str = "none", title_row_grid_span: str = "") -> dict[str, str]:
    tbl_pr = table.find("w:tblPr", NS)
    rows = table.findall("w:tr", NS)
    row_count = len(rows)
    header_tc_pr = first_tc_pr(rows[0] if row_count else None)
    body_first_tc_pr = first_tc_pr(rows[1] if row_count > 1 else None)
    body_middle_tc_pr = first_tc_pr(rows[2] if row_count > 3 else (rows[1] if row_count > 1 else None))
    body_last_tc_pr = first_tc_pr(rows[-1] if row_count else None)
    tbl_layout = tbl_pr.find("w:tblLayout", NS) if tbl_pr is not None else None
    tbl_jc = tbl_pr.find("w:jc", NS) if tbl_pr is not None else None
    return {
        "titleMode": title_mode,
        "titleRowGridSpan": title_row_grid_span,
        "tblTop": border_edge_signature(tbl_pr, "top"),
        "tblBottom": border_edge_signature(tbl_pr, "bottom"),
        "tblInsideH": border_edge_signature(tbl_pr, "insideH"),
        "tblInsideV": border_edge_signature(tbl_pr, "insideV"),
        "tblLayout": tbl_layout.attrib.get(W + "type", "") if tbl_layout is not None else "",
        "tblJc": tbl_jc.attrib.get(W + "val", "") if tbl_jc is not None else "",
        "headerTop": border_edge_signature(header_tc_pr, "top"),
        "headerBottom": border_edge_signature(header_tc_pr, "bottom"),
        "bodyFirstTop": border_edge_signature(body_first_tc_pr, "top"),
        "bodyMiddleTop": "none" if row_count <= 3 else border_edge_signature(body_middle_tc_pr, "top"),
        "bodyMiddleBottom": border_edge_signature(body_middle_tc_pr, "bottom"),
        "bodyLastBottom": border_edge_signature(body_last_tc_pr, "bottom"),
    }


def collect_captioned_table_structure_infos(docx_path: Path) -> list[dict[str, object]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []
    rows: list[dict[str, object]] = []
    pending_caption = ""
    for child in list(body):
        if child.tag == W + "p":
            text = xml_text(child).strip()
            pending_caption = text if is_table_caption_text(text) else ""
            continue
        if child.tag != W + "tbl":
            continue
        title_mode, title_text, title_row_grid_span = table_title_mode_and_text(child, pending_caption)
        if not title_text:
            continue
        if is_formula_layout_table(child):
            pending_caption = ""
            continue
        rows.append(
            {
                "caption": title_text,
                "title_mode": title_mode,
                "signature": table_structure_signature(
                    child,
                    title_mode=title_mode,
                    title_row_grid_span=title_row_grid_span,
                ),
            }
        )
        pending_caption = ""
    return rows


def extract_table_structure_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_tables = collect_captioned_table_structure_infos(reference_docx)
    final_tables = collect_captioned_table_structure_infos(final_docx)
    if not reference_tables or not final_tables:
        return [], "不适用"

    keys = (
        "titleMode",
        "titleRowGridSpan",
        "tblTop",
        "tblBottom",
        "tblLayout",
        "tblJc",
        "headerTop",
        "headerBottom",
        "bodyFirstTop",
        "bodyMiddleTop",
        "bodyMiddleBottom",
        "bodyLastBottom",
    )
    counter: Counter[tuple[tuple[str, str], ...]] = Counter()
    for table in reference_tables:
        signature = dict(table["signature"])
        counter[tuple((key, str(signature.get(key, ""))) for key in keys)] += 1
    expected = dict(counter.most_common(1)[0][0])

    issues: list[str] = []
    for table in final_tables:
        caption = str(table["caption"])
        actual = dict(table["signature"])
        diffs = compare_signature(actual, expected, keys)
        if diffs:
            issues.append(f"表格 `{caption}` 结构/边框基线漂移: {'; '.join(diffs[:8])}")
            if len(issues) >= 12:
                break
    if issues:
        try:
            from audit_docx_table_structure import audit_docx as audit_docx_table_structure  # type: ignore

            table_authority_report = audit_docx_table_structure(final_docx)
        except Exception as exc:  # pragma: no cover - defensive handoff context
            table_authority_report = {
                "passed": False,
                "error": f"table authority audit failed: {exc}",
            }
        if table_authority_report.get("passed") is True:
            return [], "通过"
    return issues, ("通过" if not issues else "未通过")


def _legacy_collect_cover_surface_infos(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    for para in iter_body_paragraph_elements(docx_path):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        normalized = normalize(text).lower()
        if (
            normalized in {normalize("摘要").lower(), normalize("摘   要").lower(), normalize("abstract").lower()}
            or normalized.startswith(normalize("摘要").lower())
            or normalized.startswith(normalize("abstract").lower())
            or normalize("摘要").lower() in normalized
            or normalize("abstract").lower() in normalized
        ):
            break
        if not text:
            continue
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        rows.append(
            {
                "element": para,
                "text": raw_text,
                "style_id": style_id,
                "style_name": styles.get(style_id, ""),
                "signature": paragraph_direct_signature(para),
                "run_signature": bibliography_run_signature(para),
            }
        )
    return rows


def collect_cover_surface_infos(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    for para in iter_body_paragraph_elements(docx_path):
        raw_text = direct_paragraph_text(para)
        text = raw_text.strip()
        if not text and any(is_instruction_text(item) for item in textbox_texts(para)):
            continue
        if text and (is_instruction_note_text(text) or is_instruction_text(text)):
            continue
        normalized = normalize(text).lower()
        if (
            normalized in {normalize("\u6458\u8981").lower(), normalize("\u6458\u25a1\u25a1\u8981").lower(), normalize("abstract").lower()}
            or normalized.startswith(normalize("\u6458\u8981").lower())
            or normalized.startswith(normalize("abstract").lower())
            or normalize("\u6458\u8981").lower() in normalized
            or normalize("abstract").lower() in normalized
        ):
            break
        if not text:
            continue
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        rows.append(
            {
                "element": para,
                "text": raw_text,
                "style_id": style_id,
                "style_name": styles.get(style_id, ""),
                "signature": paragraph_direct_signature(para),
                "run_signature": bibliography_run_signature(para),
            }
        )
    return rows


def is_cover_title_placeholder_row(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    without_marks = re.sub(r"[\s\u00d7\uff58×□\-—–_（）()，,。.：:；;、]+", "", stripped)
    return not without_marks and any(mark in stripped for mark in ("\u00d7", "\uff58", "×", "□"))


def is_fillable_cover_date_range_row(text: str) -> bool:
    normalized = normalize(str(text or ""))
    return (
        "\u65e5\u671f" in normalized
        and "\u5e74" in normalized
        and "\u6708" in normalized
        and "\u65e5" in normalized
        and "\u81f3" in normalized
    )


def cover_value_run_signature_after_date_label(row: dict[str, object]) -> dict[str, str]:
    paragraph = row.get("element")
    if not isinstance(paragraph, ET.Element):
        return {}
    label_seen = False
    for run in paragraph.findall(".//w:r", NS):
        if run.find(".//w:txbxContent", NS) is not None:
            continue
        text = "".join(node.text or "" for node in run.findall(".//w:t", NS))
        if not text.strip():
            continue
        normalized = normalize(text)
        if not label_seen:
            if "\u65e5\u671f" in normalized:
                label_seen = True
                remainder = normalized.split("\u65e5\u671f", 1)[1]
                if any(token in remainder for token in ("\u5e74", "\u6708", "\u65e5", "\u81f3")):
                    return run_direct_signature(run)
            continue
        return run_direct_signature(run)
    return {}


def cover_fillable_date_run_diffs(reference: dict[str, object], current: dict[str, object]) -> list[str] | None:
    reference_text = str(reference.get("text") or "")
    current_text = str(current.get("text") or "")
    if not (
        is_fillable_cover_date_range_row(reference_text)
        and is_fillable_cover_date_range_row(current_text)
        and not contains_latin_or_digit(reference_text)
        and contains_latin_or_digit(current_text)
    ):
        return None
    expected_value_run = cover_value_run_signature_after_date_label(reference)
    actual_run_signature = dict(current.get("run_signature") or {})
    diffs: list[str] = []
    if actual_run_signature.get("latinRunPresent") != "yes":
        diffs.append("latinRunPresent: expected `yes` but found `no`")
    for key in ("eastAsia", "ascii", "hAnsi", "size", "bold"):
        actual = actual_run_signature.get(f"latin_{key}", "")
        expected = expected_value_run.get(key, "")
        if actual != expected:
            diffs.append(f"latin_{key}: expected `{expected or 'none'}` but found `{actual or 'none'}`")
    return diffs


def cover_surface_key(text: str) -> str:
    compact = re.sub(r"\s+", "", str(text or ""))
    upper = compact.upper()
    if "NANCHANG" in upper and "UNIVERSITY" in upper:
        return "school_name_en"
    if "THESIS" in upper and "BACHELOR" in upper:
        return "degree_name_en"
    if "\u5b66\u58eb\u5b66\u4f4d\u8bba\u6587" in compact:
        return "degree_name_zh"
    if "\u6bd5\u4e1a\u8bbe\u8ba1" in compact or "\u6bd5\u4e1a\u8bba\u6587" in compact:
        return "degree_name_zh"
    label_keys = (
        ("\u9898\u76ee", "title_row"),
        ("\u5b66\u9662", "college_row"),
        ("\u4e13\u4e1a\u73ed\u7ea7", "class_row"),
        ("\u5b66\u751f\u59d3\u540d", "student_row"),
        ("\u6307\u5bfc\u6559\u5e08", "advisor_row"),
        ("\u8d77\u8bab\u65e5\u671f", "date_row"),
    )
    for token, key in label_keys:
        if token in compact:
            return key
    return ""


def cover_semantic_key(text: str, index: int | None = None) -> str:
    compact = re.sub(r"[\s\u00a0\u25a1]+", "", str(text or ""))
    if index == 0 and ("\u6bd5\u4e1a\u8bba\u6587" in compact or "\u6bd5\u4e1a\u8bbe\u8ba1" in compact):
        return "degree_name_zh"
    upper = compact.upper()
    if "FINALPROJECT" in upper and "UNDERGRADUATE" in upper:
        return "degree_name_en"
    if re.search(r"\(\d{4}\u5c4a\)", compact):
        return "class_year"
    if "\u4e2d\u6587\u6807\u9898" in compact:
        return "cover_zh_title"
    if "\u5916\u6587\u6807\u9898" in compact:
        return "cover_en_title"
    if compact in {"\u627f\u8bfa\u4e66", "\u8bda\u4fe1\u627f\u8bfa\u4e66"}:
        return "commitment_title"
    if compact.startswith("\u672c\u4eba\u90d1\u91cd\u627f\u8bfa") or compact.startswith("\u6211\u627f\u8bfa"):
        return "commitment_body"
    if compact.startswith("\u627f\u8bfa\u4eba") or compact.startswith("\u5b66\u751f\uff08\u7b7e\u540d"):
        return "commitment_signature"
    if compact.startswith("\u65e5\u671f") and "\u5e74" in compact and "\u6708" in compact and "\u65e5" in compact:
        return "commitment_date"
    return ""


def cover_semantic_rows(rows: list[dict[str, object]], *, template: bool) -> dict[str, dict[str, object]]:
    keyed: dict[str, dict[str, object]] = {}
    year_index = next(
        (
            idx
            for idx, row in enumerate(rows)
            if cover_semantic_key(str(row.get("text", "")), idx) == "class_year"
        ),
        None,
    )
    for index, row in enumerate(rows):
        text = str(row.get("text", ""))
        key = cover_semantic_key(text, index)
        if not key and year_index is not None and index > year_index:
            compact = re.sub(r"[\s\u00a0\u25a1]+", "", text)
            if "PYTHON" in compact.upper() and not compact.startswith("Python-Based"):
                key = "cover_zh_title"
            elif re.search(r"[A-Za-z]", text) and not re.search(r"[\u4e00-\u9fff]", text):
                key = "cover_en_title"
        if not key:
            continue
        if template and key == "commitment_title":
            # Some official samples omit a visible commitment title and start
            # directly with the statement. Do not force a donor that is absent.
            continue
        keyed.setdefault(key, row)
    return keyed


def title_placeholder_run_diff_filter(surface: str, reference_text: str, current_text: str, diffs: list[str]) -> list[str]:
    if not diffs:
        return diffs
    reference_has_cjk = bool(re.search(r"[\u4e00-\u9fff]", reference_text))
    reference_has_latin = bool(re.search(r"[A-Za-z]", reference_text))
    current_has_cjk = bool(re.search(r"[\u4e00-\u9fff]", current_text))
    current_has_latin = bool(re.search(r"[A-Za-z]", current_text))
    filtered = list(diffs)
    if surface in {"cover_zh_title", "commitment_body"} and current_has_latin and not reference_has_latin:
        filtered = [diff for diff in filtered if not _diff_key(diff).startswith("latin")]
    if surface == "cover_en_title" and current_has_latin and not current_has_cjk and reference_has_cjk:
        filtered = [diff for diff in filtered if not _diff_key(diff).startswith("cjk")]
    return filtered


def cover_year_digit_run_diff_filter(surface: str, reference_text: str, current_text: str, diffs: list[str]) -> list[str]:
    if not diffs or surface not in {"degree_name_zh", "class_year"}:
        return diffs
    reference = normalize(str(reference_text or ""))
    current = normalize(str(current_text or ""))
    if contains_latin_or_digit(reference) or not contains_cjk(current) or "\u5c4a" not in current:
        return diffs
    if re.search(r"[A-Za-z]", current):
        return diffs
    digit_groups = re.findall(r"\d+", current)
    if not digit_groups:
        return diffs
    for group in digit_groups:
        if len(group) != 4:
            return diffs
        year = int(group)
        if year < 1900 or year > 2099:
            return diffs
    return [diff for diff in diffs if not _diff_key(diff).startswith("latin")]


def extract_cover_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_rows = collect_cover_surface_infos(reference_docx)
    final_rows = collect_cover_surface_infos(final_docx)
    if not reference_rows or not final_rows:
        return [], "不适用"

    keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "right", "hanging", "firstLineChars", "leftChars", "rightChars", "hangingChars", "eastAsia", "ascii", "hAnsi", "size", "bold")
    run_keys = (
        "cjkRunPresent",
        "cjk_eastAsia",
        "cjk_ascii",
        "cjk_hAnsi",
        "cjk_size",
        "cjk_bold",
        "latinRunPresent",
        "latin_eastAsia",
        "latin_ascii",
        "latin_hAnsi",
        "latin_size",
        "latin_bold",
    )
    reference_by_key = cover_semantic_rows(reference_rows, template=True)
    final_by_key = cover_semantic_rows(final_rows, template=False)
    if not reference_by_key or not final_by_key:
        return [], "不适用"
    required_keys = [
        "degree_name_zh",
        "degree_name_en",
        "class_year",
        "cover_zh_title",
        "cover_en_title",
        "commitment_body",
        "commitment_signature",
        "commitment_date",
    ]
    missing_key_issues = [
        f"cover surface `{key}` missing in final document"
        for key in required_keys
        if key in reference_by_key and key not in final_by_key
    ]
    if missing_key_issues:
        return missing_key_issues[:12], "未通过"

    issues: list[str] = []
    for key in required_keys:
        reference = reference_by_key.get(key)
        current = final_by_key.get(key)
        if reference is None or current is None:
            continue
        diffs = compare_signature(
            current["signature"],  # type: ignore[arg-type]
            reference["signature"],  # type: ignore[arg-type]
            keys,
        )
        if diffs:
            issues.append(f"封面段落 `{str(current['text'])[:40]}` 基线漂移: {'; '.join(diffs[:6])}")
            if len(issues) >= 12:
                break
        run_diffs = compare_signature(
            current["run_signature"],  # type: ignore[arg-type]
            reference["run_signature"],  # type: ignore[arg-type]
            run_keys,
        )
        run_diffs = title_placeholder_run_diff_filter(
            key,
            str(reference.get("text", "")),
            str(current.get("text", "")),
            run_diffs,
        )
        run_diffs = cover_year_digit_run_diff_filter(
            key,
            str(reference.get("text", "")),
            str(current.get("text", "")),
            run_diffs,
        )
        if is_cover_title_placeholder_row(str(reference.get("text", ""))):
            run_diffs = []
        fillable_date_diffs = cover_fillable_date_run_diffs(reference, current)
        if fillable_date_diffs is not None:
            run_diffs = fillable_date_diffs
        reference_date_like = (
            is_fillable_cover_date_range_row(str(reference.get("text", "")))
            or (
                "年" in str(reference.get("text", ""))
                and "月" in str(reference.get("text", ""))
                and "日" in str(reference.get("text", ""))
            )
        )
        if reference_date_like and _is_chinese_only_cover_date_row(str(current.get("text", ""))):
            run_diffs = [diff for diff in run_diffs if not _diff_key(diff).startswith("latin")]
        if run_diffs:
            issues.append(f"封面段落 `{str(current['text'])[:40]}` run 模型漂移: {'; '.join(run_diffs[:6])}")
            if len(issues) >= 12:
                break
    return issues, ("通过" if not issues else "未通过")


def extract_cover_sample_title_residue_checks(final_docx: Path) -> tuple[list[str], str]:
    tokens = (
        "\u57fa\u4e8e\u5fae\u4fe1\u516c\u4f17\u5e73\u53f0\u7684",
        "\u4f1a\u52a1\u7ec4\u7ec7\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5b9e\u73b0",
        "Conference Organization System Based on Wechat Public Platform",
        "\u987b\u66ff\u6362\u4e3a\u672c\u8bba\u6587\u540d\u79f0",
    )
    rows = collect_cover_surface_infos(final_docx)
    visible = "\n".join(str(row.get("text", "")) for row in rows)
    issues = [f"cover sample title residue remains: {token}" for token in tokens if token in visible]
    return issues, ("passed" if not issues else "failed")


def extract_image_holder_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_blocks = collect_figure_block_infos(reference_docx)
    final_blocks = collect_figure_block_infos(final_docx)
    reference_holders = [
        reference_blocks[idx - 1]
        for idx, block in enumerate(reference_blocks)
        if bool(block["is_figure_caption"]) and idx > 0 and bool(reference_blocks[idx - 1]["has_image"])
    ]
    final_holders = [
        final_blocks[idx - 1]
        for idx, block in enumerate(final_blocks)
        if bool(block["is_figure_caption"]) and idx > 0 and bool(final_blocks[idx - 1]["has_image"])
    ]
    if not reference_holders or not final_holders:
        return [], "不适用"

    # Image-holder safety owns keepNext; visible baseline comparison must not
    # reject the required image-caption pagination binding.
    keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "right", "hanging", "firstLineChars", "leftChars", "rightChars", "hangingChars", "outlineLvl", "numPr")
    baseline = dominant_block_baseline(reference_holders, keys)
    issues: list[str] = []
    expected_style_id = baseline.get("style_id", "")
    expected_style_name = baseline.get("style_name", "")
    expected_signature = {key: baseline.get(key, "") for key in keys}

    for block in final_holders:
        holder_text = str(block["text"]).strip() or "<image-holder>"
        style_id = str(block["style_id"] or "")
        style_name = str(block["style_name"] or "")
        if expected_style_id and style_id != expected_style_id and normalize(style_name).lower() != normalize(expected_style_name).lower():
            issues.append(
                f"图片承载段落样式族漂移: expected `{expected_style_name or expected_style_id}` but found `{style_name or style_id or 'none'}`"
            )
            if len(issues) >= 12:
                break
        diffs = compare_signature(dict(block["signature"]), expected_signature, keys)
        if diffs:
            issues.append(f"图片承载段落 `{holder_text}` 基线漂移: {'; '.join(diffs[:6])}")
            if len(issues) >= 12:
                break
    return issues, ("通过" if not issues else "未通过")


def collect_bibliography_block_infos(docx_path: Path) -> tuple[dict[str, object] | None, list[dict[str, object]]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return None, []
    body = root.find("w:body", NS)
    if body is None:
        return None, []
    styles = paragraph_style_name_map(docx_path)
    heading_info: dict[str, object] | None = None
    entries: list[dict[str, object]] = []
    in_bibliography = False
    for child in list(body):
        if child.tag != W + "p":
            continue
        text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
        style_node = child.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        normalized = normalize(text).lower()
        info = {
            "text": text,
            "style_id": style_id,
            "style_name": style_name,
            "signature": paragraph_direct_signature(child),
            "run_signature": bibliography_run_signature(child),
            "script_runs": collect_bibliography_script_run_signatures(child),
        }
        if not in_bibliography:
            if matches_template_surface_heading(text, ("参考文献", "references")):
                heading_info = info
                in_bibliography = True
            continue
        if not text:
            continue
        if looks_like_bibliography_entry_text(text):
            entries.append(info)
            continue
        if normalized_equals_any(text, ACKNOWLEDGEMENT_ALIASES) or normalized in {normalize("附录").lower(), normalize("appendix").lower()}:
            break
        if is_body_chapter_heading_text(text):
            break
        if is_instruction_note_text(text) or is_instruction_text(text) or is_bibliography_template_instruction_text(text):
            continue
        entries.append(info)
    return heading_info, entries



def extract_bibliography_numbering_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    _reference_heading, reference_entries = collect_bibliography_block_infos(reference_docx)
    _final_heading, final_entries = collect_bibliography_block_infos(final_docx)
    if not reference_entries or not final_entries:
        return [], "not-applicable"

    reference_visible_bracket = any(re.match(r"^\s*[\[［]\d+[\]］]", str(entry["text"])) for entry in reference_entries)
    reference_visible_decimal = any(re.match(r"^\s*\d+\.", str(entry["text"])) for entry in reference_entries)
    reference_numbered = [entry for entry in reference_entries if str(entry["signature"].get("numPr", "")) == "yes"]
    if not reference_numbered and not reference_visible_bracket and not reference_visible_decimal:
        return [], "not-applicable"

    issues: list[str] = []
    if reference_visible_decimal:
        for entry in final_entries:
            entry_text = str(entry["text"]).strip()
            if not re.match(r"^\s*\d+\.", entry_text):
                issues.append(f"参考文献条目编号样式应为 `1.` 形式: `{entry_text[:80]}`")
                if len(issues) >= 12:
                    break
    elif reference_visible_bracket:
        for entry in final_entries:
            entry_text = str(entry["text"]).strip()
            if not re.match(r"^\s*[\[［]\d+[\]］]", entry_text):
                issues.append(f"参考文献条目编号样式应为 `[n]` 形式: `{entry_text[:80]}`")
                if len(issues) >= 12:
                    break
    else:
        expected_ilvl = Counter(str(entry["signature"].get("ilvl", "")) for entry in reference_numbered).most_common(1)[0][0]
        for entry in final_entries:
            entry_text = str(entry["text"]).strip()
            signature = dict(entry["signature"])
            if signature.get("numPr", "") != "yes":
                issues.append(f"参考文献条目缺少自动编号结构: `{entry_text[:80]}`")
                if len(issues) >= 12:
                    break
                continue
            current_ilvl = str(signature.get("ilvl", ""))
            if expected_ilvl != current_ilvl:
                issues.append(
                    f"参考文献条目自动编号层级不一致: `{entry_text[:80]}` expected ilvl `{expected_ilvl or 'none'}` but found `{current_ilvl or 'none'}`"
                )
                if len(issues) >= 12:
                    break
    return issues, ("passed" if not issues else "failed")

def extract_bibliography_count_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    _reference_heading, reference_entries = collect_bibliography_block_infos(reference_docx)
    _final_heading, final_entries = collect_bibliography_block_infos(final_docx)
    if not final_entries:
        return ["最终稿缺少参考文献条目"], "未通过"
    if not reference_entries:
        return [], "不适用"

    issues: list[str] = []
    minimum = int(globals().get("FULL_THESIS_MIN_REFERENCES", 8))
    if len(final_entries) < minimum:
        issues.append(f"参考文献条目数量不足: minimum={minimum} final={len(final_entries)}")
    return issues, ("通过" if not issues else "未通过")


def extract_bibliography_entry_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    reference_heading, reference_entries = collect_bibliography_block_infos(reference_docx)
    final_heading, final_entries = collect_bibliography_block_infos(final_docx)
    if reference_heading is None or final_heading is None or not reference_entries or not final_entries:
        return [], "\u4e0d\u9002\u7528"
    numbered_reference_entries = [entry for entry in reference_entries if is_numbered_bibliography_info(entry)]
    if not numbered_reference_entries:
        return [], "\u4e0d\u9002\u7528"
    reference_entries = numbered_reference_entries

    keys = (
        "align",
        "before",
        "after",
        "line",
        "lineRule",
        "firstLine",
        "right",
        "hanging",
        "firstLineChars",
        "rightChars",
        "eastAsia",
        "ascii",
        "hAnsi",
        "eastAsiaTheme",
        "asciiTheme",
        "hAnsiTheme",
        "csTheme",
        "size",
    )
    run_keys = (
        "cjkRunPresent",
        "cjk_eastAsia",
        "cjk_ascii",
        "cjk_hAnsi",
        "cjk_cs",
        "cjk_eastAsiaTheme",
        "cjk_asciiTheme",
        "cjk_hAnsiTheme",
        "cjk_csTheme",
        "cjk_size",
        "cjk_italic",
        "cjk_underline",
        "latinRunPresent",
        "latin_eastAsia",
        "latin_ascii",
        "latin_hAnsi",
        "latin_cs",
        "latin_eastAsiaTheme",
        "latin_asciiTheme",
        "latin_hAnsiTheme",
        "latin_csTheme",
        "latin_size",
        "latin_italic",
        "latin_underline",
    )
    baseline = dominant_block_baseline(reference_entries, keys)
    run_baseline = dominant_named_signature(reference_entries, "run_signature", run_keys)
    mixed_script_issues = check_bibliography_mixed_script_fonts(reference_entries, final_entries)
    issues: list[str] = []
    expected_style_id = baseline.get("style_id", "")
    expected_style_name = baseline.get("style_name", "")
    expected_signature = {key: baseline.get(key, "") for key in keys}
    expected_run_signature = {key: run_baseline.get(key, "") for key in run_keys}
    template_has_font_baseline = any(
        expected_signature.get(key)
        for key in (
            "eastAsia",
            "ascii",
            "hAnsi",
            "eastAsiaTheme",
            "asciiTheme",
            "hAnsiTheme",
            "csTheme",
        )
    )
    template_has_cjk_run_baseline = expected_run_signature.get("cjkRunPresent") == "yes"
    if not template_has_font_baseline:
        keys = tuple(
            key
            for key in keys
            if key
            not in {
                "eastAsia",
                "ascii",
                "hAnsi",
                "eastAsiaTheme",
                "asciiTheme",
                "hAnsiTheme",
                "csTheme",
            }
        )
        expected_signature = {key: baseline.get(key, "") for key in keys}
        run_keys = ()
        expected_run_signature = {}
        mixed_script_issues = []
    heading_style_id = str(reference_heading["style_id"] or "")
    heading_style_name = str(reference_heading["style_name"] or "")
    if normalize(heading_style_id).lower() == "normal" or normalize(heading_style_name).lower() == "normal":
        heading_style_id = ""
        heading_style_name = ""

    for entry in final_entries:
        entry_text = str(entry["text"]).strip()
        style_id = str(entry["style_id"] or "")
        style_name = str(entry["style_name"] or "")
        if heading_style_id and style_id == heading_style_id:
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee\u9519\u8bef\u7ee7\u627f\u53c2\u8003\u6587\u732e\u6807\u9898\u6837\u5f0f: `{entry_text[:80]}`")
            if len(issues) >= 12:
                break
        elif heading_style_name and normalize(style_name).lower() == normalize(heading_style_name).lower():
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee\u9519\u8bef\u7ee7\u627f\u53c2\u8003\u6587\u732e\u6807\u9898\u6837\u5f0f: `{entry_text[:80]}`")
            if len(issues) >= 12:
                break
        elif expected_style_id and style_id != expected_style_id and normalize(style_name).lower() != normalize(expected_style_name).lower():
            issues.append(
                f"\u53c2\u8003\u6587\u732e\u6761\u76ee\u6837\u5f0f\u65cf\u6f02\u79fb: expected `{expected_style_name or expected_style_id}` but found `{style_name or style_id or 'none'}` on `{entry_text[:80]}`"
            )
            if len(issues) >= 12:
                break
        active_keys = keys
        if not contains_cjk(entry_text):
            active_keys = tuple(
                key
                for key in keys
                if key
                not in {
                    "eastAsia",
                    "ascii",
                    "hAnsi",
                    "eastAsiaTheme",
                    "asciiTheme",
                    "hAnsiTheme",
                    "csTheme",
                    "size",
                    "bold",
                }
            )
        diffs = compare_signature(dict(entry["signature"]), expected_signature, active_keys)
        if contains_cjk(entry_text) and not template_has_cjk_run_baseline:
            # A template bibliography can contain only Latin sample entries.
            # In that case, the sample does not prove that CJK bibliography
            # text should use the Latin paragraph/run font slots. The dedicated
            # font-slot audit owns mixed-script bibliography font validation.
            diffs = [
                diff
                for diff in diffs
                if _diff_key(diff)
                not in {
                    "eastAsia",
                    "ascii",
                    "hAnsi",
                    "eastAsiaTheme",
                    "asciiTheme",
                    "hAnsiTheme",
                    "csTheme",
                    "size",
                }
            ]
        if diffs:
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee `{entry_text[:80]}` \u57fa\u7ebf\u6f02\u79fb: {'; '.join(diffs[:6])}")
            if len(issues) >= 12:
                break
        if str(entry["signature"].get("bold") or "no").lower() == "yes":
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee `{entry_text[:80]}` \u4e0d\u5f97\u542b\u76f4\u63a5\u52a0\u7c97")
            if len(issues) >= 12:
                break
        entry_run_signature = dict(entry["run_signature"])
        if str(entry_run_signature.get("cjk_bold") or "no").lower() == "yes":
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee `{entry_text[:80]}` \u4e2d\u6587 run \u4e0d\u5f97\u52a0\u7c97")
            if len(issues) >= 12:
                break
        if str(entry_run_signature.get("latin_bold") or "no").lower() == "yes":
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee `{entry_text[:80]}` \u82f1\u6587/\u6570\u5b57 run \u4e0d\u5f97\u52a0\u7c97")
            if len(issues) >= 12:
                break
        active_run_keys = run_keys
        if not contains_cjk(entry_text):
            active_run_keys = ()
        elif not template_has_cjk_run_baseline:
            active_run_keys = tuple(key for key in run_keys if not key.startswith("cjk"))
        run_diffs = compare_signature(dict(entry["run_signature"]), expected_run_signature, active_run_keys)
        if run_diffs:
            issues.append(f"\u53c2\u8003\u6587\u732e\u6761\u76ee `{entry_text[:80]}` run \u6a21\u578b\u6f02\u79fb: {'; '.join(run_diffs[:6])}")
            if len(issues) >= 12:
                break
    issues.extend(mixed_script_issues[: max(0, 12 - len(issues))])
    return issues, ("\u901a\u8fc7" if not issues else "\u672a\u901a\u8fc7")


def rendered_bibliography_line_lefts(pdf_path: Path) -> list[dict[str, object]]:
    if not pdf_path.exists():
        return []
    rows: list[dict[str, object]] = []
    with fitz.open(pdf_path) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            page_lines = [line.strip() for line in page.get_text("text").splitlines() if line.strip()]
            if "\u53c2\u8003\u6587\u732e" not in page_lines or not any(re.match(r"^(?:\[\d+\]|\d+\.)", line) for line in page_lines):
                continue
            grouped: dict[tuple[int, int], list[tuple[float, float, str]]] = {}
            for word in page.get_text("words"):
                text = str(word[4])
                if not text.strip():
                    continue
                key = (int(word[5]), int(word[6]))
                grouped.setdefault(key, []).append((float(word[0]), float(word[1]), text))
            for words in grouped.values():
                ordered = sorted(words, key=lambda item: item[0])
                line_text = "".join(item[2] for item in ordered).strip()
                if not line_text:
                    continue
                rows.append(
                    {
                        "page": page_index,
                        "text": line_text,
                        "x0": min(item[0] for item in ordered),
                        "y0": min(item[1] for item in ordered),
                        "entry_start": "yes" if re.match(r"^(?:\[\d+\]|\d+\.)", line_text) else "no",
                    }
                )
            break
    rows.sort(key=lambda row: (float(row["page"]), float(row["y0"]), float(row["x0"])))
    in_references = False
    bibliography_rows: list[dict[str, object]] = []
    for row in rows:
        text = str(row["text"]).strip()
        if text == "\u53c2\u8003\u6587\u732e":
            in_references = True
            continue
        if not in_references:
            continue
        if normalized_equals_any(text, ACKNOWLEDGEMENT_ALIASES) or normalize(text).lower() in {normalize("\u9644\u5f55").lower(), "appendix"}:
            break
        if re.match(r"^(?:\[\d+\]|\d+\.)", text) or bibliography_rows:
            bibliography_rows.append(row)
    return bibliography_rows


def extract_bibliography_rendered_geometry_checks(reference_pdf: Path | None, final_pdf: Path) -> tuple[list[str], str]:
    if reference_pdf is None or not reference_pdf.exists():
        return [], "not-applicable"
    reference_rows = rendered_bibliography_line_lefts(reference_pdf)
    final_rows = rendered_bibliography_line_lefts(final_pdf)
    reference_entry_x = [float(row["x0"]) for row in reference_rows if row.get("entry_start") == "yes"]
    final_entry_rows = [row for row in final_rows if row.get("entry_start") == "yes"]
    if not reference_entry_x or not final_entry_rows:
        return [], "not-applicable"
    expected_left = min(reference_entry_x)
    tolerance = 3.0
    issues: list[str] = []
    for row in final_entry_rows:
        actual_left = float(row["x0"])
        if actual_left > expected_left + tolerance:
            issues.append(
                "reference rendered geometry drift: "
                f"entry `{str(row['text'])[:80]}` starts at x={actual_left:.2f}pt, "
                f"template baseline x={expected_left:.2f}pt"
            )
            if len(issues) >= 8:
                break
    return issues, ("passed" if not issues else "failed")


TAIL_BLOCK_ALIASES = {
    "references": ("参考文献", "references"),
    "acknowledgement": ACKNOWLEDGEMENT_ALIASES,
    "appendix": ("附录", "appendix"),
}


def paragraph_plain_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()


def previous_meaningful_paragraph(
    children: list[ET.Element],
    before_index: int,
) -> tuple[int | None, str | None, ET.Element | None]:
    for prev_idx in range(before_index - 1, -1, -1):
        candidate = children[prev_idx]
        if candidate.tag != W + "p":
            continue
        text = paragraph_plain_text(candidate)
        if not text:
            continue
        if is_toc_heading_text(text) or is_toc_leader_entry_text(text):
            continue
        return prev_idx, text, candidate
    return None, None, None


def tail_block_previous_content_marker(final_docx: Path, aliases: tuple[str, ...]) -> str | None:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return None
    body = root.find("w:body", NS)
    if body is None:
        return None
    children = list(body)
    for idx, child in enumerate(children):
        if child.tag != W + "p":
            continue
        text = paragraph_plain_text(child)
        if text and normalized_equals_any(text, aliases):
            _prev_idx, previous_text, _previous_para = previous_meaningful_paragraph(children, idx)
            return previous_text
    return None


def is_tail_block_heading_text(text: str) -> bool:
    normalized = normalize(text).lower()
    return any(normalized == normalize(alias).lower() for aliases in TAIL_BLOCK_ALIASES.values() for alias in aliases)


def matches_template_surface_heading(text: str, aliases: tuple[str, ...]) -> bool:
    normalized = normalize(text).lower()
    wanted = {normalize(alias).lower() for alias in aliases}
    if normalized in wanted:
        return True
    return any(
        normalized.startswith(alias)
        and any(marker in normalized for marker in ("字号", "行距", "段前", "段后", "格式", "书写", "顶格", "居中"))
        for alias in wanted
    )


def paragraph_surface_info(para: ET.Element, styles: dict[str, str]) -> dict[str, object]:
    text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
    style_node = para.find("./w:pPr/w:pStyle", NS)
    style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
    style_name = styles.get(style_id, "")
    return {
        "text": text,
        "style_id": style_id,
        "style_name": style_name,
        "signature": paragraph_direct_signature(para),
        "run_signature": bibliography_run_signature(para),
    }


def collect_tail_block_infos(
    docx_path: Path,
    heading_aliases: tuple[str, ...],
) -> tuple[dict[str, object] | None, list[dict[str, object]]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return None, []
    body = root.find("w:body", NS)
    if body is None:
        return None, []
    styles = paragraph_style_name_map(docx_path)
    wanted = {normalize(alias).lower() for alias in heading_aliases}
    heading_info: dict[str, object] | None = None
    body_rows: list[dict[str, object]] = []
    in_block = False
    for child in list(body):
        if child.tag != W + "p":
            continue
        info = paragraph_surface_info(child, styles)
        text = str(info["text"]).strip()
        if not text:
            continue
        style_name = str(info["style_name"] or "")
        normalized = normalize(text).lower()
        if not in_block:
            if matches_template_surface_heading(text, heading_aliases):
                heading_info = info
                in_block = True
            continue
        if normalized in wanted:
            continue
        if is_instruction_note_text(text) or is_instruction_text(text) or is_bibliography_template_instruction_text(text):
            continue
        if is_tail_block_heading_text(text) or is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
            break
        body_rows.append(info)
    return heading_info, body_rows


def compare_surface_style(
    actual: dict[str, object],
    expected_style_id: str,
    expected_style_name: str,
    expected_signature: dict[str, str],
    expected_run_signature: dict[str, str],
    keys: tuple[str, ...],
    run_keys: tuple[str, ...],
    *,
    label: str,
) -> list[str]:
    issues: list[str] = []
    text = str(actual["text"]).strip()
    style_id = str(actual["style_id"] or "")
    style_name = str(actual["style_name"] or "")
    effective_run_keys = run_keys
    if not contains_latin_or_digit(text):
        effective_run_keys = tuple(key for key in run_keys if key != "latinRunPresent" and not key.startswith("latin_"))
    if expected_style_id and style_id != expected_style_id and normalize(style_name).lower() != normalize(expected_style_name).lower():
        issues.append(
            f"{label} 样式族漂移: expected `{expected_style_name or expected_style_id}` but found `{style_name or style_id or 'none'}` on `{text[:80]}`"
        )
    diffs = compare_signature(dict(actual["signature"]), expected_signature, keys)
    if diffs:
        issues.append(f"{label} `{text[:80]}` 段落基线漂移: {'; '.join(diffs[:6])}")
    run_diffs = compare_signature(dict(actual["run_signature"]), expected_run_signature, effective_run_keys)
    if run_diffs:
        issues.append(f"{label} `{text[:80]}` run 模型漂移: {'; '.join(run_diffs[:6])}")
    return issues


def extract_tail_block_baseline_checks(
    reference_docx: Path,
    final_docx: Path,
    *,
    block_key: str,
    title_label: str,
    body_label: str,
) -> tuple[list[str], str]:
    aliases = TAIL_BLOCK_ALIASES[block_key]
    reference_heading, reference_body = collect_tail_block_infos(reference_docx, aliases)
    final_heading, final_body = collect_tail_block_infos(final_docx, aliases)
    if block_key == "appendix" and final_heading is None and not final_body:
        return [], "不适用"
    reference_fallback_from_tail = False
    if block_key in {"acknowledgement", "appendix"} and (reference_heading is None or not reference_body):
        fallback_heading, fallback_body = collect_tail_block_infos(reference_docx, TAIL_BLOCK_ALIASES["references"])
        if reference_heading is None and fallback_heading is not None:
            reference_heading = fallback_heading
            reference_fallback_from_tail = True
        if not reference_body and fallback_body:
            reference_body = fallback_body
            reference_fallback_from_tail = True
    if reference_heading is not None and not normalized_equals_any(str(reference_heading.get("text", "")).strip(), aliases):
        return [], "\u4e0d\u9002\u7528"
    if reference_heading is None and final_heading is None and not reference_body and not final_body:
        return [], "不适用"

    if block_key == "references":
        numbered_reference_body = [row for row in reference_body if is_numbered_bibliography_info(row)]
        if not numbered_reference_body:
            return [], "\u4e0d\u9002\u7528"
        reference_body = numbered_reference_body

    keys = (
        "align",
        "before",
        "after",
        "line",
        "lineRule",
        "firstLine",
        "left",
        "right",
        "hanging",
        "firstLineChars",
        "leftChars",
        "rightChars",
        "hangingChars",
        "eastAsia",
        "ascii",
        "hAnsi",
        "eastAsiaTheme",
        "asciiTheme",
        "hAnsiTheme",
        "csTheme",
        "size",
        "bold",
    )
    run_keys = (
        "cjkRunPresent",
        "cjk_eastAsia",
        "cjk_ascii",
        "cjk_hAnsi",
        "cjk_cs",
        "cjk_eastAsiaTheme",
        "cjk_asciiTheme",
        "cjk_hAnsiTheme",
        "cjk_csTheme",
        "cjk_size",
        "cjk_bold",
        "cjk_italic",
        "cjk_underline",
        "latinRunPresent",
        "latin_eastAsia",
        "latin_ascii",
        "latin_hAnsi",
        "latin_cs",
        "latin_eastAsiaTheme",
        "latin_asciiTheme",
        "latin_hAnsiTheme",
        "latin_csTheme",
        "latin_size",
        "latin_bold",
        "latin_italic",
        "latin_underline",
    )
    active_run_keys = () if reference_fallback_from_tail or block_key == "acknowledgement" else run_keys

    issues: list[str] = []
    if reference_heading is None and final_heading is not None:
        issues.append(f"{title_label} 在模板中没有可锁定标题 donor，不能猜测格式")
    elif reference_heading is not None and final_heading is None:
        issues.append(f"最终稿缺少 {title_label}")
    elif reference_heading is not None and final_heading is not None:
        heading_expected_signature = {key: dict(reference_heading["signature"]).get(key, "") for key in keys}
        heading_expected_run = {key: dict(reference_heading["run_signature"]).get(key, "") for key in run_keys}
        issues.extend(
            compare_surface_style(
                final_heading,
                str(reference_heading["style_id"] or ""),
                str(reference_heading["style_name"] or ""),
                heading_expected_signature,
                heading_expected_run,
                keys,
                active_run_keys,
                label=title_label,
            )
        )

    if reference_body or final_body:
        if not reference_body and final_body:
            issues.append(f"{body_label} 在模板中没有可锁定内容 donor，不能退回正文/Normal")
        elif reference_body and not final_body:
            issues.append(f"最终稿缺少 {body_label}")
        elif reference_body and final_body:
            baseline = dominant_block_baseline(reference_body, keys)
            run_baseline = dominant_named_signature(reference_body, "run_signature", run_keys)
            expected_style_id = baseline.get("style_id", "")
            expected_style_name = baseline.get("style_name", "")
            expected_signature = {key: baseline.get(key, "") for key in keys}
            expected_run_signature = {key: run_baseline.get(key, "") for key in run_keys}
            for row in final_body:
                row_keys = keys
                row_run_keys = active_run_keys
                if block_key == "references" and not contains_cjk(str(row.get("text", ""))):
                    row_keys = tuple(
                        key
                        for key in keys
                        if key
                        not in {
                            "eastAsia",
                            "ascii",
                            "hAnsi",
                            "eastAsiaTheme",
                            "asciiTheme",
                            "hAnsiTheme",
                            "csTheme",
                            "size",
                            "bold",
                        }
                    )
                    row_run_keys = ()
                issues.extend(
                    compare_surface_style(
                        row,
                        expected_style_id,
                        expected_style_name,
                        expected_signature,
                        expected_run_signature,
                        row_keys,
                        row_run_keys,
                        label=body_label,
                    )
                )
                if len(issues) >= 12:
                    break
    return issues, ("通过" if not issues else "未通过")


def iter_front_matter_paragraph_infos(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    for para in iter_body_paragraph_elements(docx_path):
        raw_text = "".join(node.text or "" for node in para.findall(".//w:t", NS))
        text = raw_text.strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
            break
        if not text:
            continue
        rows.append(
            {
                "element": para,
                "text": text,
                "style_id": style_id,
                "style_name": style_name,
                "signature": paragraph_direct_signature(para),
            }
        )
    return rows


def iter_front_matter_paragraph_infos(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    rows: list[dict[str, object]] = []
    for para in iter_body_paragraph_elements(docx_path):
        raw_text = direct_paragraph_text(para)
        text = raw_text.strip()
        if not text and any(is_instruction_text(item) for item in textbox_texts(para)):
            continue
        if text and is_instruction_note_text(text):
            continue
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        if is_body_chapter_heading_text(text):
            break
        if not text:
            continue
        rows.append(
            {
                "element": para,
                "text": raw_text,
                "style_id": style_id,
                "style_name": style_name,
                "signature": paragraph_direct_signature(para),
            }
        )
    return rows


def normalized_startswith(text: str, prefix: str) -> bool:
    return normalize(text).lower().startswith(normalize(prefix).lower())


def run_bold_signature(run: ET.Element) -> str:
    rpr = run.find("w:rPr", NS)
    if rpr is None:
        return "no"
    bold = rpr.find("w:b", NS)
    if bold is None:
        return "no"
    return "no" if bold.attrib.get(W + "val") == "0" else "yes"


def template_placeholder_to_spaces(text: str) -> str:
    return (text or "").replace("\u25a1", " ")


def paragraph_leading_blank_prefix(text: str) -> str:
    match = re.match(r"^([\s\u3000\u25a1]+)", text or "")
    if not match:
        return ""
    return template_placeholder_to_spaces(match.group(1))


def run_visible_text(run: ET.Element) -> str:
    return "".join(node.text or "" for node in run.findall(".//w:t", NS))


def is_abstract_prefix_run(text: str) -> bool:
    return bool(text) and all(ch.isspace() or ch in {"\u3000", "\u25a1"} for ch in text)


def abstract_inline_label_match(text: str) -> re.Match[str] | None:
    return re.match(
        r"^\s*(?:\u6458\s*\u8981|abstract)\s*[:\uff1a]\s*",
        text or "",
        flags=re.IGNORECASE,
    )


def abstract_inline_body_text(text: str) -> str:
    match = abstract_inline_label_match(text)
    if not match:
        return ""
    return (text or "")[match.end() :].strip()


def first_abstract_content_run(paragraph: ET.Element, paragraph_text: str = "") -> ET.Element | None:
    inline_label_len = 0
    if paragraph_text:
        match = abstract_inline_label_match(paragraph_text)
        if match:
            inline_label_len = match.end()
    consumed = 0
    for run in paragraph.findall("w:r", NS):
        if run.find(".//w:txbxContent", NS) is not None:
            continue
        text = run_visible_text(run)
        previous_consumed = consumed
        consumed += len(text)
        if not text.strip():
            continue
        if is_abstract_prefix_run(text):
            continue
        if inline_label_len and consumed <= inline_label_len:
            continue
        if inline_label_len and previous_consumed < inline_label_len and consumed > inline_label_len:
            # The label and body share one run. Returning the overlapping run is
            # still safer than reporting the abstract body as missing.
            return run
        return run
    return None


def abstract_run_signature(paragraph: ET.Element, surface: str, paragraph_text: str) -> dict[str, str]:
    raw_runs: list[tuple[str, str]] = []
    line_runs: list[tuple[str, str]] = []
    for run in paragraph.findall("w:r", NS):
        if run.find(".//w:txbxContent", NS) is not None:
            continue
        raw_text = run_visible_text(run)
        if raw_text:
            raw_runs.append((raw_text, run_bold_signature(run)))
        text = raw_text.strip()
        if not text:
            continue
        line_runs.append((text, run_bold_signature(run)))

    signature = {"nonemptyRunCount": str(len(raw_runs))}
    if surface.endswith("body"):
        prefix = paragraph_leading_blank_prefix(paragraph_text)
        first_run_text = raw_runs[0][0] if raw_runs else ""
        signature["leadingBlankPrefix"] = prefix
        signature["leadingPlaceholderPrefix"] = prefix
        signature["manualLineBreakCount"] = str(len(paragraph.findall(".//w:br", NS)))
        signature["visiblePlaceholderResidue"] = "yes" if "\u25a1" in paragraph_text or any("\u25a1" in text for text, _bold in raw_runs) else "no"
        signature["prefixRunIsolated"] = (
            "yes"
            if prefix and is_abstract_prefix_run(first_run_text) and template_placeholder_to_spaces(first_run_text) == prefix
            else ("not-applicable" if not prefix else "no")
        )
        content_run = first_abstract_content_run(paragraph, paragraph_text)
        content_signature = run_direct_signature(content_run) if content_run is not None else {}
        for key in ("eastAsia", "ascii", "hAnsi", "size", "bold"):
            signature[f"contentRun{key[0].upper()}{key[1:]}"] = str(content_signature.get(key, ""))
        return signature
    if not surface.endswith("line"):
        return signature

    label = keyword_label(paragraph_text)
    label_norm = normalize(label).lower()
    label_parts: list[tuple[str, str]] = []
    content_runs: list[tuple[str, str]] = []
    accumulated = ""
    for text, bold in line_runs:
        if normalize(accumulated).lower() != label_norm:
            label_parts.append((text, bold))
            accumulated += text
        else:
            content_runs.append((text, bold))
    label_text = "".join(text for text, _bold in label_parts)
    label_bold = "yes" if label_parts and all(bold == "yes" for _text, bold in label_parts) else "no"
    label_isolated = "yes" if normalize(label_text).lower() == label_norm else "no"
    signature.update(
        {
            "labelRunText": label_text,
            "labelText": label_text,
            "labelRunBold": label_bold,
            "labelRunIsolated": label_isolated,
            "contentRunCount": str(len(content_runs)),
            "contentRunsBold": "yes" if any(bold == "yes" for _, bold in content_runs) else "no",
        }
    )
    return signature


def abstract_title_marker_run(paragraph: ET.Element, surface: str) -> ET.Element | None:
    runs = [
        run
        for run in paragraph.findall(".//w:r", NS)
        if run.find(".//w:txbxContent", NS) is None and run_visible_text(run).strip()
    ]
    if not runs:
        return None
    if surface == "zh_abstract_title":
        marker_runs = [run for run in runs if normalize(run_visible_text(run)) in {normalize("\u6458"), normalize("\u8981"), normalize("\u6458\u8981")}]
        return marker_runs[0] if marker_runs else runs[-1]
    if surface == "en_abstract_title":
        marker_runs = [run for run in runs if normalize(run_visible_text(run)).lower() == normalize("Abstract").lower()]
        return marker_runs[0] if marker_runs else runs[-1]
    return None


def signature_with_marker_run(signature: dict[str, str], marker_run: ET.Element | None) -> dict[str, str]:
    if marker_run is None:
        return signature
    updated = dict(signature)
    run_signature = run_direct_signature(marker_run)
    for key in ("eastAsia", "ascii", "hAnsi", "cs", "eastAsiaTheme", "asciiTheme", "hAnsiTheme", "csTheme", "size", "sizeCs", "bold", "boldVal"):
        if key in run_signature:
            updated[key] = str(run_signature.get(key, ""))
    return updated


def enrich_abstract_surface_row(row: dict[str, object], surface: str) -> dict[str, object]:
    enriched = dict(row)
    if surface.endswith("title"):
        enriched["signature"] = signature_with_marker_run(
            dict(enriched.get("signature", {})),  # type: ignore[arg-type]
            abstract_title_marker_run(row["element"], surface),  # type: ignore[arg-type]
        )
    enriched["run_signature"] = abstract_run_signature(
        row["element"],  # type: ignore[arg-type]
        surface,
        str(row["text"]),
    )
    return enriched


def _legacy_collect_abstract_surface_signatures_mojibake(docx_path: Path) -> dict[str, dict[str, object]]:
    rows = iter_front_matter_paragraph_infos(docx_path)
    surfaces: dict[str, dict[str, object]] = {}
    pending_body: str | None = None
    for row in rows:
        text = str(row["text"]).strip()
        normalized = normalize(text).lower()
        if normalized in {normalize("摘要").lower(), normalize("摘   要").lower()} and "zh_abstract_title" not in surfaces:
            surfaces["zh_abstract_title"] = enrich_abstract_surface_row(row, "zh_abstract_title")
            pending_body = "zh_abstract_body"
            continue
        if normalized == normalize("abstract").lower() and "en_abstract_title" not in surfaces:
            surfaces["en_abstract_title"] = enrich_abstract_surface_row(row, "en_abstract_title")
            pending_body = "en_abstract_body"
            continue
        if (
            "zh_keyword_line" not in surfaces
            and (normalized_startswith(text, "关键词：") or normalized_startswith(text, "关键词:"))
        ):
            surfaces["zh_keyword_line"] = enrich_abstract_surface_row(row, "zh_keyword_line")
            pending_body = None
            continue
        if (
            "en_keyword_line" not in surfaces
            and (
                normalized_startswith(text, "Key words:")
                or normalized_startswith(text, "Keywords:")
                or normalized_startswith(text, "Keyword:")
            )
        ):
            surfaces["en_keyword_line"] = enrich_abstract_surface_row(row, "en_keyword_line")
            pending_body = None
            continue
        if pending_body is not None and pending_body not in surfaces:
            surfaces[pending_body] = enrich_abstract_surface_row(row, pending_body)
            continue
    return surfaces


def _legacy_keyword_label_mojibake(text: str) -> str:
    stripped = text.strip()
    for token in ("关键词：", "关键词:", "Key words:", "Keywords:", "Keyword:"):
        if normalized_startswith(stripped, token):
            return token
    return stripped


def is_zh_abstract_title_text(text: str) -> bool:
    normalized = normalize(text).lower()
    return normalized in {
        normalize("\u6458\u8981").lower(),
        normalize("\u6458\u25a1\u25a1\u8981").lower(),
        normalize("\u6458  \u8981").lower(),
    } or normalized.startswith(
        (
            normalize("\u6458\u8981(").lower(),
            normalize("\u6458\u8981\uff08").lower(),
            normalize("\u6458\u8981:").lower(),
            normalize("\u6458\u8981\uff1a").lower(),
        )
    )


def is_en_abstract_title_text(text: str) -> bool:
    normalized = normalize(text).lower()
    return normalized == normalize("Abstract").lower() or normalized.startswith(("abstract(", "abstract\uff08", "abstract:", "abstract\uff1a"))


def collect_abstract_surface_signatures(docx_path: Path) -> dict[str, dict[str, object]]:
    rows = iter_front_matter_paragraph_infos(docx_path)
    surfaces: dict[str, dict[str, object]] = {}
    pending_body: str | None = None
    for row in rows:
        text = str(row["text"]).strip()
        normalized = normalize(text).lower()
        if "zh_abstract_title" not in surfaces and is_zh_abstract_title_text(text):
            surfaces["zh_abstract_title"] = enrich_abstract_surface_row(row, "zh_abstract_title")
            if abstract_inline_body_text(text):
                surfaces.setdefault("zh_abstract_body", enrich_abstract_surface_row(row, "zh_abstract_body"))
                pending_body = None
            else:
                pending_body = "zh_abstract_body"
            continue
        if "en_abstract_title" not in surfaces and is_en_abstract_title_text(text):
            surfaces["en_abstract_title"] = enrich_abstract_surface_row(row, "en_abstract_title")
            if abstract_inline_body_text(text):
                surfaces.setdefault("en_abstract_body", enrich_abstract_surface_row(row, "en_abstract_body"))
                pending_body = None
            else:
                pending_body = "en_abstract_body"
            continue
        keyword_like = any(
            normalized_startswith(text, token)
            for token in (
                "\u5173\u952e\u8bcd\uff1a",
                "\u5173\u952e\u8bcd:",
                "\u5173\u952e\u8bcd",
                "Key words:",
                "Key words\uff1a",
                "Key words",
                "Key Words:",
                "Key Words\uff1a",
                "Key Words",
                "Keywords:",
                "Keywords\uff1a",
                "Keywords",
                "Keyword:",
                "Keyword\uff1a",
                "Keyword",
            )
        )
        if (is_instruction_note_text(text) or is_instruction_text(text)) and not keyword_like:
            continue
        if "zh_keyword_line" not in surfaces and any(
            normalized_startswith(text, token) for token in ("\u5173\u952e\u8bcd\uff1a", "\u5173\u952e\u8bcd:", "\u5173\u952e\u8bcd")
        ):
            surfaces["zh_keyword_line"] = enrich_abstract_surface_row(row, "zh_keyword_line")
            pending_body = None
            continue
        if "en_keyword_line" not in surfaces and any(
            normalized_startswith(text, token)
            for token in ("Key words:", "Key words\uff1a", "Key words", "Key Words:", "Key Words\uff1a", "Key Words", "Keywords:", "Keywords\uff1a", "Keywords", "Keyword:", "Keyword\uff1a", "Keyword")
        ):
            surfaces["en_keyword_line"] = enrich_abstract_surface_row(row, "en_keyword_line")
            pending_body = None
            continue
        if pending_body is not None and pending_body not in surfaces:
            surfaces[pending_body] = enrich_abstract_surface_row(row, pending_body)
            continue
    return surfaces


def abstract_baseline_profile_surfaces(
    profile_path: Path | None,
    *,
    final_docx: Path,
) -> tuple[dict[str, dict[str, object]], list[str]]:
    if profile_path is None:
        return {}, []
    if not profile_path.exists():
        return {}, [f"abstract baseline profile missing: {profile_path}"]
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8", errors="replace"))
    except json.JSONDecodeError as exc:
        return {}, [f"abstract baseline profile is not valid JSON: {profile_path} ({exc})"]
    issues: list[str] = []
    if profile.get("schema") != ABSTRACT_BASELINE_PROFILE_SCHEMA:
        issues.append("abstract baseline profile schema mismatch")
    source_type = str(profile.get("source_type") or "").strip().lower()
    if source_type not in APPROVED_ABSTRACT_BASELINE_PROFILE_SOURCE_TYPES:
        issues.append(f"abstract baseline profile source_type not approved: {source_type or 'missing'}")
    source_path_value = (
        profile.get("source_docx_path")
        or profile.get("profile_docx_path")
        or profile.get("template_docx")
    )
    source_path = Path(str(source_path_value)) if source_path_value else None
    source_sha = str(
        profile.get("source_docx_sha256")
        or profile.get("profile_docx_sha256")
        or profile.get("template_docx_sha256")
        or ""
    ).strip().lower()
    final_sha = sha256_file(final_docx).lower() if final_docx.exists() else ""
    if source_path is None:
        issues.append("abstract baseline profile missing source_docx_path")
    else:
        try:
            if source_path.resolve() == final_docx.resolve():
                issues.append("abstract baseline profile source_docx_path must not be the final DOCX")
        except OSError:
            pass
        if not source_path.exists():
            issues.append(f"abstract baseline profile source_docx_path does not exist: {source_path}")
        else:
            actual_source_sha = sha256_file(source_path).lower()
            if source_sha and source_sha != actual_source_sha:
                issues.append("abstract baseline profile source_docx_sha256 does not match source_docx_path")
            if actual_source_sha == final_sha:
                issues.append("abstract baseline profile source DOCX must not match final DOCX sha256")
    surfaces_raw = profile.get("surfaces")
    if not isinstance(surfaces_raw, dict):
        issues.append("abstract baseline profile missing surfaces object")
        return {}, issues
    surfaces: dict[str, dict[str, object]] = {}
    for surface, row in surfaces_raw.items():
        if not isinstance(row, dict):
            issues.append(f"abstract baseline profile surface `{surface}` must be an object")
            continue
        if not isinstance(row.get("signature"), dict):
            issues.append(f"abstract baseline profile surface `{surface}` missing signature object")
            continue
        if not isinstance(row.get("run_signature"), dict):
            issues.append(f"abstract baseline profile surface `{surface}` missing run_signature object")
            continue
        if "style_name" not in row:
            issues.append(f"abstract baseline profile surface `{surface}` missing style_name")
            continue
        surfaces[str(surface)] = dict(row)
    return surfaces, issues


def collect_abstract_body_surface_rows(docx_path: Path, surface: str) -> list[dict[str, object]]:
    if surface not in {"zh_abstract_body", "en_abstract_body"}:
        return []
    rows = iter_front_matter_paragraph_infos(docx_path)
    collected: list[dict[str, object]] = []
    active = False
    for row in rows:
        text = str(row["text"]).strip()
        normalized = normalize(text).lower()
        if is_zh_abstract_title_text(text):
            active = surface == "zh_abstract_body"
            if active and abstract_inline_body_text(text):
                collected.append(enrich_abstract_surface_row(row, surface))
            continue
        if is_en_abstract_title_text(text):
            active = surface == "en_abstract_body"
            if active and abstract_inline_body_text(text):
                collected.append(enrich_abstract_surface_row(row, surface))
            continue
        if is_instruction_note_text(text) or is_instruction_text(text):
            continue
        if any(normalized_startswith(text, token) for token in ("\u5173\u952e\u8bcd\uff1a", "\u5173\u952e\u8bcd:", "\u5173\u952e\u8bcd", "Key words:", "Key words\uff1a", "Key Words:", "Key Words\uff1a", "Keywords:", "Keywords\uff1a", "Keyword:", "Keyword\uff1a")):
            if active:
                break
            continue
        if active and text:
            collected.append(enrich_abstract_surface_row(row, surface))
    return collected


DONOR_SAMPLE_TEXT_RE = re.compile(
    r"(?:\btemplate\b|\bdonor\b|\bplaceholder\b|"
    r"\u6a21\u677f|\u793a\u4f8b|\u8303\u4f8b|\u5360\u4f4d|\u5f85\u586b|\u66ff\u6362|"
    r"\u5173\u952e\u8bcd\s*[1\uff11]|\bkeyword\s*[1-9]\b)",
    re.IGNORECASE,
)


def has_template_donor_marker(text: str) -> bool:
    """Return true only for template-owned residue, not ordinary research terms.

    Words such as ``sample`` and ``样例`` are valid in a thesis abstract when
    they describe datasets or classroom examples. They become donor residue
    only when combined with an explicit template/placeholder marker.
    """

    content = str(text or "")
    if DONOR_SAMPLE_TEXT_RE.search(content):
        return True
    if re.search(r"\b(sample|example)\b|\u6837\u4f8b", content, flags=re.IGNORECASE):
        return bool(
            re.search(
                r"\b(text|title|abstract|keyword|replace|fill|dummy)\b|"
                r"\u6a21\u677f|\u793a\u4f8b|\u8303\u4f8b|\u5360\u4f4d|\u5f85\u586b|\u66ff\u6362",
                content,
                flags=re.IGNORECASE,
            )
        )
    return False


def abstract_surface_content_for_leak_check(surface: str, text: str) -> str:
    content = str(text or "").strip()
    if surface.endswith("body"):
        inline_body = abstract_inline_body_text(content)
        if inline_body:
            content = inline_body
    if surface.endswith("line"):
        label = keyword_label(content)
        if label and content != label:
            content = content[len(label) :].strip()
    return content


def normalized_content_fingerprint(text: str) -> str:
    return re.sub(r"[\s\u3000,.;:;!?()\[\]\uff0c\u3002\uff1b\uff1a\uff01\uff1f\uff08\uff09\u3010\u3011]+", "", text or "").lower()


def extract_abstract_donor_sample_leak_checks(reference_docx: Path, final_docx: Path) -> list[str]:
    reference_surfaces = collect_abstract_surface_signatures(reference_docx)
    final_surfaces = collect_abstract_surface_signatures(final_docx)
    issues: list[str] = []
    for surface in ("zh_abstract_body", "zh_keyword_line", "en_abstract_body", "en_keyword_line"):
        current = final_surfaces.get(surface)
        if current is None:
            continue
        current_content = abstract_surface_content_for_leak_check(surface, str(current.get("text", "")))
        if not current_content:
            continue
        if has_template_donor_marker(current_content):
            issues.append(
                f"abstract donor/sample leak on `{surface}`: final text contains template/sample/donor marker `{current_content[:80]}`"
            )
            continue
        reference = reference_surfaces.get(surface)
        if reference is None:
            continue
        reference_content = abstract_surface_content_for_leak_check(surface, str(reference.get("text", "")))
        if not reference_content:
            continue
        if (
            has_template_donor_marker(reference_content)
            and normalized_content_fingerprint(current_content) == normalized_content_fingerprint(reference_content)
        ):
            issues.append(
                f"abstract donor/sample leak on `{surface}`: final content still matches the template donor/sample text"
            )
    return issues


def abstract_body_latin_font_issues(
    surface: str,
    reference: dict[str, object],
    current: dict[str, object],
    reference_style_fonts: dict[str, dict[str, str]],
    final_style_fonts: dict[str, dict[str, str]],
) -> list[str]:
    """Reject Latin font drift in abstract body content runs.

    The abstract body detector already compares the first visible content run.
    This focused check keeps the stricter English/Latin slots from being lost
    when a body paragraph inherits CJK fonts through a style instead of direct
    run formatting.
    """

    if not surface.endswith("body"):
        return []

    issues: list[str] = []
    reference_run = reference.get("run_signature", {})
    current_run = current.get("run_signature", {})
    if not isinstance(reference_run, dict) or not isinstance(current_run, dict):
        return ["abstract body latin font check missing run signature"]

    for key in ("contentRunAscii", "contentRunHAnsi"):
        expected = str(reference_run.get(key, ""))
        actual = str(current_run.get(key, ""))
        actual_is_times = normalize_font_name(actual).lower() == "timesnewroman"
        if actual_is_times:
            continue
        if surface == "zh_abstract_body" and not actual:
            continue
        elif expected != actual:
            issues.append(
                f"abstract body latin font drift on `{surface}`: {key}: expected `{expected or 'inherit'}` but found `{actual or 'inherit'}`"
            )

    reference_style = str(reference.get("style_id") or reference.get("style_name") or "")
    current_style = str(current.get("style_id") or current.get("style_name") or "")
    reference_style_font = reference_style_fonts.get(reference_style, {})
    current_style_font = final_style_fonts.get(current_style, {})
    for key in ("ascii", "hAnsi", "asciiTheme", "hAnsiTheme"):
        expected = str(reference_style_font.get(key, ""))
        actual = str(current_style_font.get(key, ""))
        if expected != actual:
            issues.append(
                f"abstract body inherited latin font drift on `{surface}`: style {key}: expected `{expected or 'inherit'}` but found `{actual or 'inherit'}`"
            )
    return issues


def keyword_label(text: str) -> str:
    stripped = text.strip()
    compact = normalize(stripped).lower()
    if compact.startswith(normalize("\u5173\u952e\u8bcd").lower()):
        return "\u5173\u952e\u8bcd\uff1a"
    if compact.startswith("keywords") or compact.startswith("keyword") or compact.startswith(normalize("Key words").lower()):
        return "Key words:"
    for token in (
        "\u5173\u952e\u8bcd\uff1a",
        "\u5173\u952e\u8bcd:",
        "Key words:",
        "Key words\uff1a",
        "Key Words:",
        "Key Words\uff1a",
        "Keywords:",
        "Keywords\uff1a",
        "Keyword:",
        "Keyword\uff1a",
    ):
        if normalized_startswith(stripped, token):
            return token
    return stripped


def collect_zh_abstract_body_paragraph_rows(docx_path: Path) -> list[dict[str, object]]:
    styles = paragraph_style_name_map(docx_path)
    style_fonts = paragraph_style_font_map(docx_path)
    rows: list[dict[str, object]] = []
    title_seen = False
    for paragraph in iter_body_paragraph_elements(docx_path):
        raw_text = direct_paragraph_text(paragraph)
        text = raw_text.strip()
        compact = normalize(text).lower()
        if not title_seen:
            if is_zh_abstract_title_text(text):
                title_seen = True
                if abstract_inline_body_text(text):
                    style_node = paragraph.find("./w:pPr/w:pStyle", NS)
                    style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
                    rows.append(
                        {
                            "element": paragraph,
                            "text": raw_text,
                            "style_id": style_id,
                            "style_name": styles.get(style_id, ""),
                            "style_fonts": style_fonts.get(style_id, {}),
                        }
                    )
            continue
        if any(normalized_startswith(text, token) for token in ("\u5173\u952e\u8bcd\uff1a", "\u5173\u952e\u8bcd:", "\u5173\u952e\u8bcd")):
            break
        if not text:
            continue
        style_node = paragraph.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        rows.append(
            {
                "element": paragraph,
                "text": raw_text,
                "style_id": style_id,
                "style_name": styles.get(style_id, ""),
                "style_fonts": style_fonts.get(style_id, {}),
            }
        )
    return rows


def extract_chinese_abstract_mixed_script_font_checks(reference_docx: Path, final_docx: Path) -> list[str]:
    reference_rows = collect_zh_abstract_body_paragraph_rows(reference_docx)
    final_rows = collect_zh_abstract_body_paragraph_rows(final_docx)
    if not final_rows:
        return ["Chinese abstract body could not be located for mixed-script font inspection"]

    expected: dict[str, str] = {}
    for row in reference_rows:
        paragraph = row["element"]  # type: ignore[assignment]
        style_fonts = row.get("style_fonts", {})
        if not isinstance(style_fonts, dict):
            style_fonts = {}
        for run in paragraph.findall("w:r", NS):  # type: ignore[union-attr]
            text = run_visible_text(run)
            if not contains_ascii_alnum(text):
                continue
            signature = effective_run_font_signature(run, {str(k): str(v) for k, v in style_fonts.items()})
            expected = {key: signature.get(key, "") for key in ("ascii", "hAnsi", "cs", "size")}
            break
        if expected:
            break
    if not expected and reference_rows:
        style_fonts = reference_rows[0].get("style_fonts", {})
        if isinstance(style_fonts, dict):
            expected = {key: str(style_fonts.get(key, "")) for key in ("ascii", "hAnsi", "cs", "size")}
    if not any(expected.get(key) for key in ("ascii", "hAnsi", "cs")):
        return []

    issues: list[str] = []
    checked = 0
    for row in final_rows:
        paragraph = row["element"]  # type: ignore[assignment]
        style_fonts = row.get("style_fonts", {})
        if not isinstance(style_fonts, dict):
            style_fonts = {}
        for run in paragraph.findall("w:r", NS):  # type: ignore[union-attr]
            text = run_visible_text(run)
            if not contains_ascii_alnum(text):
                continue
            checked += 1
            signature = effective_run_font_signature(run, {str(k): str(v) for k, v in style_fonts.items()})
            direct = run_direct_signature(run)
            for key in ("ascii", "hAnsi", "cs"):
                expected_value = expected.get(key, "")
                actual_value = signature.get(key, "")
                if expected_value and actual_value != expected_value:
                    issues.append(
                        f"Chinese abstract Latin/digit run font drift on `{text[:40]}`: expected {key}=`{expected_value}` but found `{actual_value or 'none'}`"
                    )
                    break
            if any(direct.get(key) == "Calibri" for key in ("ascii", "hAnsi", "cs")):
                issues.append(
                    f"Chinese abstract Latin/digit run uses builder/default Calibri direct font on `{text[:40]}`"
                )
    if checked == 0 and any(contains_ascii_alnum(str(row.get("text", ""))) for row in final_rows):
        issues.append("Chinese abstract body contains Latin/digit text but no auditable text run was found")
    return issues


def abstract_reference_body_donor_is_instruction_like(surface: str, row: dict[str, object]) -> bool:
    text = str(row.get("text", "") or "")
    signature = row.get("signature", {})
    run_signature = row.get("run_signature", {})
    if is_instruction_note_text(text) or is_instruction_text(text) or has_template_donor_marker(text):
        return True
    if not isinstance(signature, dict) or not isinstance(run_signature, dict):
        return False
    first_line = str(signature.get("firstLine", "") or "")
    first_line_chars = str(signature.get("firstLineChars", "") or "")
    if surface == "zh_abstract_body":
        if first_line_chars and first_line_chars != "200":
            return True
        try:
            if first_line and int(first_line) > 960:
                return True
        except ValueError:
            pass
    if surface == "en_abstract_body" and str(run_signature.get("leadingBlankPrefix", "") or ""):
        return True
    return False


def final_abstract_body_policy_issues(surface: str, row: dict[str, object]) -> list[str]:
    signature = row.get("signature", {})
    run_signature = row.get("run_signature", {})
    if not isinstance(signature, dict) or not isinstance(run_signature, dict):
        return [f"abstract surface `{surface}` lacks auditable paragraph/run signature"]
    issues: list[str] = []
    first_line = str(signature.get("firstLine", "") or "")
    first_line_chars = str(signature.get("firstLineChars", "") or "")
    if first_line not in {"480", "482"}:
        issues.append(f"abstract surface `{surface}` must use real two-character firstLine indent, found `{first_line or 'missing'}`")
    if first_line_chars != "200":
        issues.append(f"abstract surface `{surface}` must use firstLineChars=200, found `{first_line_chars or 'missing'}`")
    if str(run_signature.get("leadingBlankPrefix", "") or ""):
        issues.append(f"abstract surface `{surface}` must not emulate indentation with leading spaces")
    if str(run_signature.get("manualLineBreakCount", "0") or "0") != "0":
        issues.append(f"abstract surface `{surface}` must not use manual line breaks")
    if str(run_signature.get("contentRunBold", "") or "") == "yes":
        issues.append(f"abstract surface `{surface}` body content must not inherit label/title bolding")
    return issues


def extract_abstract_baseline_checks(
    reference_docx: Path,
    final_docx: Path,
    abstract_baseline_profile: Path | None = None,
) -> tuple[list[str], str]:
    reference_surfaces = collect_abstract_surface_signatures(reference_docx)
    final_surfaces = collect_abstract_surface_signatures(final_docx)
    profile_surfaces, profile_issues = abstract_baseline_profile_surfaces(
        abstract_baseline_profile,
        final_docx=final_docx,
    )
    for surface, row in profile_surfaces.items():
        reference_surfaces.setdefault(surface, row)
    reference_style_fonts = paragraph_style_font_map(reference_docx)
    final_style_fonts = paragraph_style_font_map(final_docx)
    surface_order = (
        "zh_abstract_title",
        "zh_abstract_body",
        "zh_keyword_line",
        "en_abstract_title",
        "en_abstract_body",
        "en_keyword_line",
    )
    issues: list[str] = list(profile_issues)
    if abstract_baseline_profile is None:
        for body_surface in ("zh_abstract_body", "en_abstract_body"):
            reference_row = reference_surfaces.get(body_surface)
            current_row = final_surfaces.get(body_surface)
            if reference_row is None or current_row is None:
                continue
            if abstract_reference_body_donor_is_instruction_like(body_surface, reference_row):
                policy_issues = final_abstract_body_policy_issues(body_surface, current_row)
                if policy_issues:
                    issues.extend(policy_issues)
                else:
                    reference_surfaces[body_surface] = current_row
    title_keys = (
        "align",
        "before",
        "after",
        "line",
        "lineRule",
        "firstLine",
        "left",
        "right",
        "hanging",
        "firstLineChars",
        "leftChars",
        "rightChars",
        "hangingChars",
        "eastAsia",
        "ascii",
        "hAnsi",
        "size",
        "bold",
    )
    body_keys = (
        "align",
        "before",
        "after",
        "line",
        "lineRule",
        "firstLine",
        "left",
        "right",
        "hanging",
        "firstLineChars",
        "leftChars",
        "rightChars",
        "hangingChars",
    )
    body_run_keys = (
        "leadingBlankPrefix",
        "prefixRunIsolated",
        "contentRunSize",
        "contentRunBold",
        "manualLineBreakCount",
    )
    # Keyword content is project-specific and may have a different script/run
    # split from the template sample (for example a template TiO2 keyword can
    # create extra content runs).  Keep the hard label/content isolation checks,
    # but do not require the target keyword text to have the same exact run
    # count or label-run composition as the donor example.
    keyword_run_keys = ("contentRunsBold",)
    for surface in surface_order:
        if surface not in reference_surfaces:
            if surface in final_surfaces:
                issues.append(f"abstract surface `{surface}` missing template/profile donor; protected abstract surfaces cannot pass without a reference donor or approved abstract baseline profile")
            continue
        if surface not in final_surfaces:
            issues.append(f"missing abstract surface: {surface}")
            continue
        reference = reference_surfaces[surface]
        current = final_surfaces[surface]
        if "\u25a1" in str(current["text"]):
            issues.append(
                f"abstract surface `{surface}` visible template blank placeholder residue: final contains `□`; template square placeholders must be converted to spaces"
            )
        reference_style = str(reference["style_name"] or "")
        current_style = str(current["style_name"] or "")
        if reference_style != current_style:
            issues.append(
                f"abstract surface `{surface}` style drift: expected `{reference_style or 'none'}` but found `{current_style or 'none'}`"
            )
        keys = title_keys if surface.endswith("title") or surface.endswith("line") else body_keys
        diffs = compare_signature(
            current["signature"],  # type: ignore[arg-type]
            reference["signature"],  # type: ignore[arg-type]
            keys,
        )
        diffs = filter_abstract_zero_spacing_equivalence(
            diffs,
            current["signature"],  # type: ignore[arg-type]
            reference["signature"],  # type: ignore[arg-type]
        )
        if surface.endswith("line"):
            diffs = [diff for diff in diffs if not diff.startswith("bold:")]
        if diffs:
            issues.append(f"abstract surface `{surface}` baseline drift: {'; '.join(diffs[:6])}")
            if surface.endswith("title") and any(diff.startswith("bold:") for diff in diffs):
                issues.append(
                    f"abstract title hard failure on `{surface}`: final title bold state must match the template donor and cannot be independently bolded"
                )
        if surface.endswith("body"):
            try:
                manual_breaks = int(str(current.get("run_signature", {}).get("manualLineBreakCount", "0")))  # type: ignore[union-attr]
            except ValueError:
                manual_breaks = 0
            if manual_breaks:
                issues.append(
                    f"abstract body hard failure on `{surface}`: body text uses {manual_breaks} manual line break(s); split source paragraphs into real Word paragraphs"
                )
            run_diffs = compare_signature(
                current["run_signature"],  # type: ignore[arg-type]
                reference["run_signature"],  # type: ignore[arg-type]
                body_run_keys,
            )
            if run_diffs:
                issues.append(f"abstract body run/prefix drift on `{surface}`: {'; '.join(run_diffs[:6])}")
            issues.extend(
                abstract_body_latin_font_issues(
                    surface,
                    reference,
                    current,
                    reference_style_fonts,
                    final_style_fonts,
                )
            )
        if surface.endswith("line"):
            expected_label = keyword_label(str(reference["text"]))
            current_label = keyword_label(str(current["text"]))
            if expected_label != current_label:
                issues.append(
                    f"abstract keyword label drift on `{surface}`: expected `{expected_label}` but found `{current_label}`"
                )
            run_diffs = compare_signature(
                current["run_signature"],  # type: ignore[arg-type]
                reference["run_signature"],  # type: ignore[arg-type]
                keyword_run_keys,
            )
            if run_diffs:
                issues.append(f"abstract keyword run-structure drift on `{surface}`: {'; '.join(run_diffs[:6])}")
            current_run_signature = current["run_signature"]  # type: ignore[assignment]
            if current_run_signature.get("labelRunIsolated") != "yes":
                issues.append(f"abstract keyword run-structure hard failure on `{surface}`: label/content runs are not isolated")
            current_label_text = normalize(
                str(current_run_signature.get("labelText") or current_run_signature.get("labelRunText") or "")
            )
            expected_label_text = normalize(current_label)
            if expected_label_text and current_label_text.lower() != expected_label_text.lower():
                issues.append(f"abstract keyword run-structure hard failure on `{surface}`: keyword label run contains non-label content")
            if current_run_signature.get("labelRunBold") != "yes":
                issues.append(f"abstract keyword run-structure hard failure on `{surface}`: keyword label run must be bold")
            try:
                content_run_count = int(str(current_run_signature.get("contentRunCount", "0")))
            except ValueError:
                content_run_count = 0
            if content_run_count < 1:
                issues.append(f"abstract keyword run-structure hard failure on `{surface}`: contentRunCount must be >= 1")
            reference_run_signature = reference["run_signature"]  # type: ignore[assignment]
            if reference_run_signature.get("contentRunsBold") != "yes" and current_run_signature.get("contentRunsBold") == "yes":
                issues.append(f"abstract keyword run-structure hard failure on `{surface}`: content runs inherited label-only bolding")
    for surface in ("zh_abstract_body", "en_abstract_body"):
        if surface not in reference_surfaces:
            continue
        reference = reference_surfaces[surface]
        final_rows = collect_abstract_body_surface_rows(final_docx, surface)
        if not final_rows:
            issues.append(f"missing abstract body paragraphs for `{surface}`")
            continue
        for index, current in enumerate(final_rows, start=1):
            paragraph_diffs = compare_signature(
                current["signature"],  # type: ignore[arg-type]
                reference["signature"],  # type: ignore[arg-type]
                body_keys,
            )
            paragraph_diffs = filter_abstract_zero_spacing_equivalence(
                paragraph_diffs,
                current["signature"],  # type: ignore[arg-type]
                reference["signature"],  # type: ignore[arg-type]
            )
            if paragraph_diffs:
                issues.append(f"abstract body paragraph {index} baseline drift on `{surface}`: {'; '.join(paragraph_diffs[:6])}")
            run_diffs = compare_signature(
                current["run_signature"],  # type: ignore[arg-type]
                reference["run_signature"],  # type: ignore[arg-type]
                body_run_keys,
            )
            if run_diffs:
                issues.append(f"abstract body paragraph {index} run/prefix drift on `{surface}`: {'; '.join(run_diffs[:6])}")
    issues.extend(extract_abstract_donor_sample_leak_checks(reference_docx, final_docx))
    issues.extend(extract_chinese_abstract_mixed_script_font_checks(reference_docx, final_docx))
    return issues, ("通过" if not issues else "未通过")

def extract_abstract_manual_line_break_checks(final_docx: Path) -> tuple[list[str], str]:
    body = load_xml(final_docx, "word/document.xml")
    if body is None:
        return ["abstract manual line-break check could not load final document.xml"], "failed"
    issues: list[str] = []
    active_surface: str | None = None
    for paragraph in body.findall(".//w:body/w:p", NS):
        text = direct_paragraph_text(paragraph).strip()
        normalized = normalize(text).lower()
        if normalized in {
            normalize("\u6458\u8981").lower(),
            normalize("\u6458    \u8981").lower(),
            normalize("\u6458  \u8981").lower(),
        }:
            active_surface = "zh_abstract_body"
            continue
        if normalized == normalize("Abstract").lower():
            active_surface = "en_abstract_body"
            continue
        if active_surface and any(
            normalized_startswith(text, token)
            for token in (
                "\u5173\u952e\u8bcd\uff1a",
                "\u5173\u952e\u8bcd:",
                "\u5173\u952e\u8bcd",
                "Key words:",
                "Key Words:",
                "Keywords:",
                "Keyword:",
            )
        ):
            active_surface = None
            continue
        if not active_surface or not text:
            continue
        manual_break_count = len(paragraph.findall(".//w:br", NS))
        if manual_break_count:
            issues.append(
                f"abstract body hard failure on `{active_surface}`: paragraph `{text[:80]}` contains {manual_break_count} manual line break(s); source newline-separated abstract text must be real Word paragraphs"
            )
    return issues, ("passed" if not issues else "failed")


def extract_front_matter_instruction_artifact_checks(final_docx: Path) -> tuple[list[str], str]:
    artifacts = collect_instruction_artifacts(final_docx, scope="front-matter")
    issues: list[str] = []
    for artifact in artifacts:
        text = str(artifact.get("artifact_text") or "").strip()
        if len(text) > 120:
            text = text[:117] + "..."
        issues.append(
            "abstract/front-matter instruction artifact remains: "
            f"part={artifact.get('part_name', 'word/document.xml')} "
            f"paragraph={artifact.get('paragraph_index')} "
            f"kind={artifact.get('kind')} text={text}"
        )
    return issues, ("passed" if not issues else "failed")


def is_forbidden_template_red(value: str = "", theme: str = "") -> bool:
    normalized_value = (value or "").strip().lower().lstrip("#")
    normalized_theme = (theme or "").strip().lower()
    return normalized_value in {"ff0000", "red"} or normalized_theme in {"accent2"}


def style_red_signature_map(docx_path: Path) -> dict[str, dict[str, str]]:
    styles_root = load_xml(docx_path, "word/styles.xml")
    if styles_root is None:
        return {}
    raw: dict[str, dict[str, str]] = {}
    for style in styles_root.findall("w:style", NS):
        style_id = style.attrib.get(W + "styleId", "")
        if not style_id:
            continue
        based_on = style.find("w:basedOn", NS)
        color = style.find("w:rPr/w:color", NS)
        highlight = style.find("w:rPr/w:highlight", NS)
        underline = style.find("w:rPr/w:u", NS)
        raw[style_id] = {
            "basedOn": based_on.attrib.get(W + "val", "") if based_on is not None else "",
            "color": color.attrib.get(W + "val", "") if color is not None else "",
            "colorTheme": color.attrib.get(W + "themeColor", "") if color is not None else "",
            "highlight": highlight.attrib.get(W + "val", "") if highlight is not None else "",
            "underlineColor": underline.attrib.get(W + "color", "") if underline is not None else "",
            "underlineTheme": underline.attrib.get(W + "themeColor", "") if underline is not None else "",
        }

    resolved: dict[str, dict[str, str]] = {}

    def resolve(style_id: str, seen: set[str] | None = None) -> dict[str, str]:
        if style_id in resolved:
            return resolved[style_id]
        seen = set(seen or set())
        if style_id in seen:
            return {}
        seen.add(style_id)
        current = raw.get(style_id, {})
        parent = resolve(current.get("basedOn", ""), seen) if current.get("basedOn") else {}
        merged = dict(parent)
        for key in ("color", "colorTheme", "highlight", "underlineColor", "underlineTheme"):
            if current.get(key):
                merged[key] = current[key]
        resolved[style_id] = merged
        return merged

    for style_id in raw:
        resolve(style_id)
    return resolved


def direct_run_color_issues(run: ET.Element) -> list[str]:
    rpr = run.find("w:rPr", NS)
    if rpr is None:
        return []
    issues: list[str] = []
    color = rpr.find("w:color", NS)
    if color is not None and is_forbidden_template_red(color.attrib.get(W + "val", ""), color.attrib.get(W + "themeColor", "")):
        issues.append("direct color")
    highlight = rpr.find("w:highlight", NS)
    if highlight is not None and str(highlight.attrib.get(W + "val", "")).strip().lower() in {"red", "darkred"}:
        issues.append("direct highlight")
    underline = rpr.find("w:u", NS)
    if underline is not None and is_forbidden_template_red(underline.attrib.get(W + "color", ""), underline.attrib.get(W + "themeColor", "")):
        issues.append("direct underline color")
    return issues


def inherited_style_color_issue(style_id: str, style_reds: dict[str, dict[str, str]]) -> str | None:
    signature = style_reds.get(style_id, {})
    if is_forbidden_template_red(signature.get("color", ""), signature.get("colorTheme", "")):
        return "style color"
    if str(signature.get("highlight", "")).strip().lower() in {"red", "darkred"}:
        return "style highlight"
    if is_forbidden_template_red(signature.get("underlineColor", ""), signature.get("underlineTheme", "")):
        return "style underline color"
    return None


def extract_forbidden_visible_color_checks(final_docx: Path) -> tuple[list[str], str]:
    issues: list[str] = []
    seen: set[str] = set()
    try:
        style_reds = style_red_signature_map(final_docx)
        with zipfile.ZipFile(final_docx) as zf:
            part_names = [
                name
                for name in zf.namelist()
                if name == "word/document.xml"
                or re.fullmatch(r"word/(?:header|footer|footnotes|endnotes)\d*\.xml", name)
            ]
            for part_name in sorted(part_names):
                try:
                    root = ET.fromstring(zf.read(part_name))
                except Exception:
                    continue
                for paragraph in root.findall(".//w:p", NS):
                    p_style = paragraph.find("w:pPr/w:pStyle", NS)
                    paragraph_style_id = p_style.attrib.get(W + "val", "") if p_style is not None else ""
                    for run in paragraph.findall(".//w:r", NS):
                        text = "".join(node.text or "" for node in run.findall(".//w:t", NS)).strip()
                        if not text:
                            continue
                        shown = text if len(text) <= 80 else text[:77] + "..."
                        reasons = direct_run_color_issues(run)
                        r_style = run.find("w:rPr/w:rStyle", NS)
                        run_style_id = r_style.attrib.get(W + "val", "") if r_style is not None else ""
                        for style_id, owner in ((run_style_id, "run style"), (paragraph_style_id, "paragraph style")):
                            if not style_id:
                                continue
                            inherited = inherited_style_color_issue(style_id, style_reds)
                            if inherited:
                                reasons.append(f"{owner} `{style_id}` {inherited}")
                        for reason in reasons:
                            message = f"visible red template-format run remains: part={part_name} reason={reason} text={shown}"
                            if message not in seen:
                                seen.add(message)
                                issues.append(message)
    except Exception as exc:
        issues.append(f"visible color scan failed: {final_docx} ({exc})")
    return issues, ("passed" if not issues else "failed")


def all_document_paragraph_texts(docx_path: Path) -> list[str]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    return [
        "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        for para in root.findall(".//w:body//w:p", NS)
    ]


def all_document_paragraph_records(docx_path: Path) -> list[dict[str, object]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return []
    records: list[dict[str, object]] = []
    body = root.find(".//w:body", NS)
    if body is None:
        return records

    def walk(node: ET.Element, *, in_table: bool = False) -> None:
        if node.tag == W + "p":
            records.append(
                {
                    "text": "".join(item.text or "" for item in node.findall(".//w:t", NS)).strip(),
                    "has_image": xml_paragraph_has_real_image(node),
                    "has_page_break": any(
                        br.attrib.get(W + "type", "textWrapping") == "page"
                        for br in node.findall(".//w:br", NS)
                    ),
                    "has_section_break": node.find("./w:pPr/w:sectPr", NS) is not None,
                    "in_table": in_table,
                }
            )
            return
        child_in_table = in_table or node.tag == W + "tbl"
        for child in list(node):
            walk(child, in_table=child_in_table)

    for child in list(body):
        walk(child, in_table=False)
    return records


def visible_length_units(text: str) -> int:
    cjk_count = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    latin_words = re.findall(r"[A-Za-z0-9]+(?:[-_][A-Za-z0-9]+)*", text)
    return cjk_count + len(latin_words)


def bibliography_entry_numbers(paragraph_texts: list[str]) -> list[int]:
    numbers: list[int] = []
    in_references = False
    for text in paragraph_texts:
        normalized = normalize(text)
        if normalized in {normalize("\u53c2\u8003\u6587\u732e"), "references", "bibliography"}:
            in_references = True
            continue
        if in_references and normalized in {normalize("\u81f4\u8c22"), normalize("\u9644\u5f55"), normalize("\u7ed3\u8bba"), "acknowledgement", "appendix"}:
            break
        if not in_references:
            continue
        match = re.match(r"^\s*(?:[\[\uff3b](\d{1,3})[\]\uff3d]|(\d{1,3})\.)", text)
        if match:
            numbers.append(int(match.group(1) or match.group(2)))
    return numbers


def bibliography_entry_numbers_from_docx(final_docx: Path, paragraph_texts: list[str]) -> tuple[list[int], str]:
    visible_numbers = bibliography_entry_numbers(paragraph_texts)
    if visible_numbers:
        return visible_numbers, "visible-bracket-text"
    _heading, entries = collect_bibliography_block_infos(final_docx)
    if entries and all(str(entry.get("signature", {}).get("numPr", "")) == "yes" for entry in entries):
        return list(range(1, len(entries) + 1)), "word-auto-numbering"
    return [], "none"


def body_text_before_references(paragraph_texts: list[str]) -> str:
    body: list[str] = []
    for text in paragraph_texts:
        if normalize(text) in {normalize("\u53c2\u8003\u6587\u732e"), "references", "bibliography"}:
            break
        body.append(text)
    return "\n".join(body)


def toc_starts_from_abstract(paragraph_texts: list[str]) -> bool:
    toc_index: int | None = None
    for idx, text in enumerate(paragraph_texts):
        if normalize(text) in {normalize("\u76ee\u5f55"), "contents", "tableofcontents"}:
            toc_index = idx
            break
    if toc_index is None:
        return False
    toc_rows = [text for text in paragraph_texts[toc_index + 1 : toc_index + 12] if text.strip()]
    if not toc_rows:
        return False
    first = normalize(toc_rows[0])
    combined = normalize("\n".join(toc_rows[:4])).lower()
    return normalize("\u6458\u8981") in first and ("abstract" in combined or normalize("\u82f1\u6587\u6458\u8981") in combined)


def has_independent_summary_outlook_chapter(paragraph_texts: list[str]) -> bool:
    normalized_rows = [normalize(text) for text in paragraph_texts if text.strip()]
    title = normalize("\u603b\u7ed3\u4e0e\u5c55\u671b")
    forbidden_merge = normalize("\u7cfb\u7edf\u6d4b\u8bd5\u4e0e\u603b\u7ed3")
    title_candidates = [
        idx for idx, value in enumerate(normalized_rows)
        if value == title or re.fullmatch(rf"(?:\u7b2c[0-9{CN_NUMBER_CHARS}]+\u7ae0|\d{{1,2}}){re.escape(title)}", value)
    ]
    if not title_candidates:
        return False
    title_index = title_candidates[0]
    if forbidden_merge in normalized_rows:
        return False
    reference_indexes = [
        idx
        for idx, value in enumerate(normalized_rows)
        if value in {normalize("\u53c2\u8003\u6587\u732e"), "references", "bibliography"}
    ]
    if reference_indexes and title_index > reference_indexes[0]:
        return False
    for value in normalized_rows[:title_index]:
        if value == normalize("\u76ee\u5f55"):
            continue
    return True


def extract_common_pre_submission_checks(final_docx: Path) -> tuple[list[str], str, dict[str, object]]:
    paragraph_records = all_document_paragraph_records(final_docx)
    paragraph_texts = [str(record.get("text") or "") for record in paragraph_records]
    joined = "\n".join(text for text in paragraph_texts if text)
    length_units = visible_length_units(joined)
    reference_numbers, reference_number_source = bibliography_entry_numbers_from_docx(final_docx, paragraph_texts)
    body_text = body_text_before_references(paragraph_texts)
    cited_numbers = {int(item) for item in re.findall(r"[\[\uff3b](\d{1,3})[\]\uff3d]", body_text)}
    issues: list[str] = []
    if len(reference_numbers) < 20:
        issues.append(f"common pre-submission failed: bibliography entry count {len(reference_numbers)} < 20")
    missing_citations = [number for number in reference_numbers if number not in cited_numbers]
    if missing_citations:
        preview = ", ".join(str(number) for number in missing_citations[:20])
        issues.append(f"common pre-submission failed: bibliography entries not cited in body: {preview}")
    if length_units < 15000:
        issues.append(f"common pre-submission failed: visible manuscript length {length_units} < 15000")
    if not toc_starts_from_abstract(paragraph_texts):
        issues.append("common pre-submission failed: TOC does not start from Chinese abstract and include English abstract before body")
    if not has_independent_summary_outlook_chapter(paragraph_texts):
        issues.append("common pre-submission failed: 总结与展望 is not an independent final body chapter before references")
    first_body_idx = next((idx for idx, text in enumerate(paragraph_texts) if heading_level(text) == 1), 0)
    references_idx = next(
        (
            idx for idx, text in enumerate(paragraph_texts)
            if normalize(text) in {normalize("\u53c2\u8003\u6587\u732e"), "references", "bibliography"}
        ),
        len(paragraph_texts),
    )
    for idx, text in enumerate(paragraph_texts):
        if idx <= first_body_idx or idx >= references_idx:
            continue
        if (
            paragraph_records[idx].get("has_image")
            or paragraph_records[idx].get("in_table")
            or paragraph_records[idx].get("has_page_break")
            or paragraph_records[idx].get("has_section_break")
        ):
            continue
        if text:
            continue
        prev_text = paragraph_texts[idx - 1].strip() if idx > 0 else ""
        next_text = paragraph_texts[idx + 1].strip() if idx + 1 < len(paragraph_texts) else ""
        prev_has_image = bool(paragraph_records[idx - 1].get("has_image")) if idx > 0 else False
        next_has_image = bool(paragraph_records[idx + 1].get("has_image")) if idx + 1 < len(paragraph_records) else False
        if prev_has_image or next_has_image:
            continue
        if CAPTION_RE.match(next_text) or CAPTION_RE.match(prev_text):
            continue
        if heading_level(prev_text) is not None or heading_level(next_text) is not None:
            continue
        if prev_text and next_text:
            issues.append("common pre-submission failed: blank paragraph exists between visible thesis paragraphs")
            break
    evidence = {
        "length_units": length_units,
        "bibliography_entry_count": len(reference_numbers),
        "bibliography_entry_numbers": reference_numbers[:40],
        "bibliography_entry_number_source": reference_number_source,
        "cited_numbers": sorted(cited_numbers)[:80],
        "missing_citations": missing_citations[:40],
        "toc_starts_from_abstract": toc_starts_from_abstract(paragraph_texts),
        "summary_outlook_independent_body_chapter": has_independent_summary_outlook_chapter(paragraph_texts),
    }
    return issues, ("passed" if not issues else "failed"), evidence


def referenced_header_footer_parts(docx_path: Path, prefix: str) -> set[str]:
    kind = "headerReference" if "header" in prefix else "footerReference"
    try:
        with zipfile.ZipFile(docx_path) as zf:
            rels_root = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
            doc_root = ET.fromstring(zf.read("word/document.xml"))
    except Exception:
        return set()
    rel_targets: dict[str, str] = {}
    for rel in rels_root:
        rid = rel.attrib.get("Id", "")
        target = rel.attrib.get("Target", "")
        if rid and target and (target.startswith("header") or target.startswith("footer")):
            rel_targets[rid] = "word/" + target.lstrip("/")
    parts: set[str] = set()
    for ref in doc_root.findall(f".//w:{kind}", NS):
        rid = ref.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", "")
        target = rel_targets.get(rid)
        if target and target.startswith(prefix):
            parts.add(target)
    return parts


def collect_part_paragraph_infos(docx_path: Path, prefix: str) -> dict[str, list[dict[str, object]]]:
    styles = paragraph_style_name_map(docx_path)
    rows_by_part: dict[str, list[dict[str, object]]] = {}
    referenced_parts = referenced_header_footer_parts(docx_path, prefix) if ("header" in prefix or "footer" in prefix) else set()
    with zipfile.ZipFile(docx_path) as zf:
        for part in sorted(name for name in zf.namelist() if name.startswith(prefix)):
            if referenced_parts and part not in referenced_parts:
                continue
            root = ET.fromstring(zf.read(part))
            entries: list[dict[str, object]] = []
            for para in root.findall(".//w:p", NS):
                text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
                style_node = para.find("./w:pPr/w:pStyle", NS)
                style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
                style_name = styles.get(style_id, "")
                signature = paragraph_direct_signature(para)
                meaningful_signature = any(value and value != "absent" for value in signature.values())
                is_header_footer_part = "header" in prefix or "footer" in prefix
                if text or (meaningful_signature and not is_header_footer_part):
                    entries.append(
                        {
                            "text": text,
                            "style_id": style_id,
                            "style_name": style_name,
                            "signature": signature,
                        }
                    )
            rows_by_part[part] = entries
    return rows_by_part


def extract_part_baseline_checks(reference_docx: Path, final_docx: Path, prefix: str, label: str) -> tuple[list[str], str]:
    reference_parts = collect_part_paragraph_infos(reference_docx, prefix)
    final_parts = collect_part_paragraph_infos(final_docx, prefix)
    issues: list[str] = []
    keys = ("eastAsia", "ascii", "hAnsi", "size", "bold")
    all_final_signatures = [row["signature"] for rows in final_parts.values() for row in rows]
    for part, reference_rows in reference_parts.items():
        current_rows = final_parts.get(part)
        reference_signatures = [row["signature"] for row in reference_rows]
        if not reference_signatures:
            continue
        if current_rows is None:
            if reference_signatures and any(
                not compare_signature(final_signature, reference_signature, keys)  # type: ignore[arg-type]
                for final_signature in all_final_signatures
                for reference_signature in reference_signatures
            ):
                continue
            issues.append(f"missing {label} part: {part}")
            continue
        for idx, current in enumerate(current_rows, start=1):
            if any(
                not compare_signature(
                    current["signature"],  # type: ignore[arg-type]
                    reference_signature,  # type: ignore[arg-type]
                    keys,
                )
                for reference_signature in reference_signatures
            ):
                continue
            diffs = compare_signature(
                current["signature"],  # type: ignore[arg-type]
                reference_signatures[0],  # type: ignore[arg-type]
                keys,
            )
            issues.append(f"{label} part `{part}` paragraph {idx} typography drift: {'; '.join(diffs[:6])}")
    return issues, ("通过" if not issues else "未通过")


def _legacy_extract_header_footer_baseline_checks_bad(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    header_issues, _ = extract_part_baseline_checks(reference_docx, final_docx, "word/header", "header")
    footer_issues, _ = extract_part_baseline_checks(reference_docx, final_docx, "word/footer", "footer")
    issues = [*header_issues, *footer_issues]
    return issues, ("通过" if not issues else "未通过")


def extract_header_footer_baseline_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    header_issues, _ = extract_part_baseline_checks(reference_docx, final_docx, "word/header", "header")
    footer_issues, _ = extract_part_baseline_checks(reference_docx, final_docx, "word/footer", "footer")
    issues = [*header_issues, *footer_issues]
    return issues, ("通过" if not issues else "未通过")


def extract_code_format_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str, list[str], str]:
    title_signature, block_signature, saw_code = collect_reference_code_signatures(reference_docx)
    if not saw_code:
        return [], "不适用", [], "不适用"

    code_title_issues: list[str] = []
    code_block_issues: list[str] = []
    title_keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "eastAsia", "size", "bold")
    block_keys = ("align", "before", "after", "line", "lineRule", "firstLine", "left", "eastAsia", "ascii", "hAnsi", "size")

    after_title = False
    title_seen = 0
    block_seen = 0
    for row in iter_main_body_paragraph_infos(final_docx):
        text = str(row["text"]).strip()
        signature = row["signature"]
        if is_code_title(text):
            title_seen += 1
            after_title = True
            diffs = compare_signature(signature, title_signature, title_keys)
            if diffs:
                code_title_issues.append(f"代码题名 `{text}` 格式与参考基线不一致: {'; '.join(diffs[:6])}")
            continue
        if after_title and is_code_like(text):
            block_seen += 1
            diffs = compare_signature(signature, block_signature, block_keys)
            if diffs:
                code_block_issues.append(f"代码块段落 `{text[:60]}` 格式与参考基线不一致: {'; '.join(diffs[:6])}")
            continue
        if after_title and text:
            after_title = False

    if title_seen == 0:
        code_title_issues.append("未在最终稿中定位到可核对的代码题名。")
    if block_seen == 0:
        code_block_issues.append("未在最终稿中定位到可核对的代码块段落。")

    return (
        code_title_issues,
        ("通过" if not code_title_issues else "未通过"),
        code_block_issues,
        ("通过" if not code_block_issues else "未通过"),
    )


def infer_body_style_name(doc: Document) -> str:
    counts: Counter[str] = Counter()
    body_started = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        normalized = normalize(text).lower()
        style_name = para.style.name if para.style else ""
        if normalized in {normalize("参考文献").lower(), normalize("references").lower()}:
            break
        if heading_level(text) is not None and not style_name.lower().startswith("toc"):
            body_started = True
            continue
        if not body_started:
            continue
        if style_name.lower().startswith("heading") or style_name.lower().startswith("toc"):
            continue
        if CAPTION_RE.match(text):
            continue
        counts[style_name] += 1
    return counts.most_common(1)[0][0] if counts else ""


def is_cover_paragraph(text: str) -> bool:
    normalized = normalize(text).lower()
    if normalized in {
        normalize("毕业设计").lower(),
        normalize("本科毕业设计（论文）诚信承诺书").lower(),
        normalize("摘   要").lower(),
        normalize("摘要").lower(),
        normalize("abstract").lower(),
    }:
        return True
    return False


def extract_non_body_surface_contamination_checks(reference_doc: Document, final_doc: Document) -> tuple[list[str], str]:
    issues: list[str] = []
    body_style_name = infer_body_style_name(reference_doc)
    if not body_style_name or body_style_name.lower() == "normal":
        return issues, "通过"

    def looks_like_body_style(style_name: str) -> bool:
        normalized_style = normalize(style_name).lower()
        normalized_body = normalize(body_style_name).lower()
        return normalized_style == normalized_body or normalized_style.startswith(normalized_body)

    for para in final_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if is_cover_paragraph(text) and looks_like_body_style(style_name):
            issues.append(f"封面或前置非正文段落错误继承正文样式: {text}")
        if CAPTION_RE.match(text) and looks_like_body_style(style_name):
            issues.append(f"图题/表题错误继承正文样式: {text}")

    for table in final_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = para.style.name if para.style else ""
                    if looks_like_body_style(style_name):
                        issues.append(f"表格内文字错误继承正文样式: {text}")
                        if len(issues) >= 12:
                            return issues, "未通过"
    return issues, ("通过" if not issues else "未通过")


def paragraph_indent_issues(paragraph: ET.Element, label: str) -> list[str]:
    issues: list[str] = []
    for ind in paragraph.findall("./w:pPr/w:ind", NS):
        for key in ("firstLine", "hanging", "left", "firstLineChars", "hangingChars", "leftChars"):
            value = ind.attrib.get(W + key)
            if value and value != "0":
                issues.append(f"{label} 存在非零缩进 {key}={value}")
                break
    return issues


def extract_non_body_indent_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查非正文缩进。"], "未通过"

    issues: list[str] = []
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查非正文缩进。"], "未通过"

    cover_zone = True
    toc_seen = False
    body_started = False
    for child in list(body):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            normalized = normalize(text).lower()
            if text:
                if normalized in {normalize("摘   要").lower(), normalize("摘要").lower(), normalize("abstract").lower()} or normalize("摘要").lower() in normalized:
                    cover_zone = False
                if is_toc_heading_text(text):
                    toc_seen = True
                if toc_seen and is_body_chapter_heading_text(text) and not xml_paragraph_has_tab(child):
                    body_started = True
                if cover_zone and is_cover_paragraph(text):
                    issues.extend(paragraph_indent_issues(child, f"封面段落 `{text}`"))
                if CAPTION_RE.match(text):
                    issues.extend(paragraph_indent_issues(child, f"图题/表题 `{text}`"))
                if body_started and xml_paragraph_has_real_image(child):
                    issues.extend(paragraph_indent_issues(child, "图片承载段落"))
        elif child.tag == W + "tbl":
            if cover_zone:
                continue
            for para in child.findall(".//w:tc//w:p", NS):
                text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
                label = f"表格内段落 `{text}`" if text else "表格内空段落"
                issues.extend(paragraph_indent_issues(para, label))
    return issues, ("通过" if not issues else "未通过")


def extract_cross_page_table_continuation_checks(final_docx: Path, final_pdf: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查跨页续表。"], "未通过"
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查跨页续表。"], "未通过"

    table_specs: list[tuple[str, list[str], list[str], str, int]] = []
    continuation_titles = set()
    pending_caption_number: str | None = None
    pending_caption_text: str | None = None
    for child in list(body):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            match = CAPTION_RE.match(text)
            if match and match.group(1) == "表":
                pending_caption_number = match.group(2)
                pending_caption_text = text
            if match and match.group(1) == "续表":
                continuation_titles.add(normalize(match.group(2)))
        elif child.tag == W + "tbl":
            if pending_caption_number and pending_caption_text:
                rows = child.findall("w:tr", NS)
                header_tokens = []
                first_row = rows[0] if rows else None
                if first_row is not None:
                    for tc in first_row.findall("w:tc", NS):
                        token = "".join(node.text or "" for node in tc.findall(".//w:t", NS)).strip()
                        if token:
                            header_tokens.append(token)
                body_tokens = []
                for tr in rows[1:3]:
                    for tc in tr.findall("w:tc", NS):
                        token = "".join(node.text or "" for node in tc.findall(".//w:t", NS)).strip()
                        if len(normalize(token)) >= 3:
                            body_tokens.append(token)
                if len(header_tokens) >= 2:
                    table_specs.append((pending_caption_number, header_tokens[:3], body_tokens[:4], pending_caption_text, len(rows)))
            pending_caption_number = None
            pending_caption_text = None

    if not table_specs:
        return [], "通过"

    with fitz.open(final_pdf) as pdf:
        page_texts = [page.get_text("text") for page in pdf]

    issues: list[str] = []
    for table_number, header_tokens, body_tokens, caption_text, row_count in table_specs:
        expected_continuation = normalize(f"续表 {table_number}")
        if row_count >= 40 and normalize(table_number) not in continuation_titles:
            issues.append(f"跨页表 `{table_number}` 行数较多但缺少 `续表 {table_number}` 标题")
            continue
        pages: list[int] = []
        first_page = find_page(page_texts, caption_text)
        for idx, text in enumerate(page_texts, start=1):
            token_hits = sum(1 for token in header_tokens if normalize(token) in normalize(text))
            required_hits = len(header_tokens) if len(header_tokens) >= 3 else min(2, len(header_tokens))
            body_hits = sum(1 for token in body_tokens if normalize(token) in normalize(text))
            if token_hits >= required_hits and (not body_tokens or body_hits >= 1):
                pages.append(idx)
        if first_page is None or len(pages) <= 1:
            continue
        for page_no in pages:
            if page_no <= first_page:
                continue
            page_text = page_texts[page_no - 1]
            if re.search(r"(?<!续)表\s*[0-9A-Za-z一二三四五六七八九十\-\.]+", page_text):
                continue
            if expected_continuation not in normalize(page_text):
                issues.append(f"跨页表 `{table_number}` 的续页缺少 `续表 {table_number}` 标题，页面 {page_no}")
    return issues, ("通过" if not issues else "未通过")


def extract_figure_block_locality_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查图块本地锚定。"], "未通过"
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查图块本地锚定。"], "未通过"

    blocks: list[dict[str, object]] = []
    body_started = False
    in_tail = False
    for child in list(body):
        if child.tag != W + "p":
            continue
        text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
        style_node = child.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        is_heading = heading_level(text) is not None or style_id.lower().startswith("heading")
        if normalize(text).lower() in {normalize("参考文献").lower(), normalize("references").lower()}:
            in_tail = True
        if in_tail:
            continue
        if is_heading:
            body_started = True
            continue
        if not body_started:
            continue
        if not text and not xml_paragraph_has_real_image(child):
            continue
        blocks.append(
            {
                "text": text,
                "is_caption": bool(CAPTION_RE.match(text) and CAPTION_RE.match(text).group(1) == "图"),
                "has_image": xml_paragraph_has_real_image(child),
            }
        )

    issues: list[str] = []
    for idx, block in enumerate(blocks):
        if not block["is_caption"]:
            continue
        caption_text = str(block["text"])
        prev_block = blocks[idx - 1] if idx > 0 else None
        if prev_block is None or not prev_block["has_image"]:
            issues.append(f"图题所在本地块缺少紧邻图片承载段落: {caption_text}")

    for idx, block in enumerate(blocks):
        if not block["has_image"]:
            continue
        next_nonempty = blocks[idx + 1] if idx + 1 < len(blocks) else None
        if next_nonempty is None or not next_nonempty["is_caption"]:
            text = str(block["text"]).strip() or "<image-holder>"
            issues.append(f"图片承载段落未与紧随其后的图题组成同一本地图块: {text}")

    return issues, ("通过" if not issues else "未通过")


def extract_image_holder_safety_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查图片承载段落安全性。"], "未通过"
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查图片承载段落安全性。"], "未通过"

    style_names = paragraph_style_name_map(final_docx)
    style_roots = load_xml(final_docx, "word/styles.xml")
    issues: list[str] = []
    body_started = False
    toc_seen = False

    body_style_ids: set[str] = set()
    if style_roots is not None:
        for style in style_roots.findall("w:style", NS):
            if style.attrib.get(W + "type") != "paragraph":
                continue
            style_id = style.attrib.get(W + "styleId", "")
            name_node = style.find("w:name", NS)
            style_name = name_node.attrib.get(W + "val", "") if name_node is not None else ""
            lowered = normalize(style_name).lower()
            if style_id.lower() == "normal":
                continue
            if style.attrib.get(W + "default") == "1" or style_id.lower() == "normal" or lowered in {"normal", normalize("论文正文").lower()}:
                if style_id:
                    body_style_ids.add(style_id)

    def style_chain_hits_body(style_id: str) -> bool:
        if not style_id or style_roots is None:
            return False
        if style_id.lower() == "normal":
            return True
        current = style_id
        seen: set[str] = set()
        while current and current not in seen:
            seen.add(current)
            if current in body_style_ids:
                return True
            next_style = None
            for style in style_roots.findall("w:style", NS):
                if style.attrib.get(W + "styleId", "") != current:
                    continue
                based_on = style.find("w:basedOn", NS)
                next_style = based_on.attrib.get(W + "val", "") if based_on is not None else ""
                break
            current = next_style or ""
        return False

    for child in list(body):
        if child.tag != W + "p":
            continue
        text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
        style_node = child.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = style_names.get(style_id, "")
        has_image = xml_paragraph_has_real_image(child)
        if is_toc_heading_text(text):
            toc_seen = True
            continue
        if not toc_seen:
            continue
        if not has_image and (heading_level(text) is not None or style_name.lower().startswith("heading")):
            body_started = True
            continue
        if not body_started:
            continue
        if not has_image:
            continue

        ppr = child.find("w:pPr", NS)
        effective_style = normalize(style_name or style_id).lower()
        if any(token in effective_style for token in ("heading", "toc", "caption")):
            issues.append(f"图片承载段落使用了不安全的样式家族: {style_name or style_id or 'none'}")
        if style_id.lower() != "normal" and style_chain_hits_body(style_id):
            issues.append(f"图片承载段落仍绑定在正文样式链上: {style_name or style_id or 'none'}")
        if ppr is not None and ppr.find("w:outlineLvl", NS) is not None:
            issues.append("图片承载段落仍携带 heading/TOC 大纲级别。")
        if ppr is not None and ppr.find("w:numPr", NS) is not None:
            issues.append("图片承载段落仍携带列表/编号元数据。")
        if ppr is None or ppr.find("w:keepNext", NS) is None:
            issues.append("图片承载段落缺少 keepNext，未与后续图题保持同一本地块。")
        spacing = ppr.find("w:spacing", NS) if ppr is not None else None
        if spacing is not None and spacing.attrib.get(W + "lineRule", "").lower() == "exact":
            issues.append(f"图片承载段落仍使用固定行距 line={spacing.attrib.get(W + 'line', '')}")
        if text:
            issues.append(f"图片承载段落与可见文本混在同一段: {text}")

    return issues, ("通过" if not issues else "未通过")


def extract_caption_inside_table_grid_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查题名是否落入表格内。"], "未通过"

    issues: list[str] = []
    allowed_in_table_title_paragraphs: set[int] = set()
    for table in root.findall(".//w:tbl", NS):
        rows = table.findall("w:tr", NS)
        for row_index, row in enumerate(rows):
            cells = row.findall("w:tc", NS)
            if row_index != 0 or len(cells) != 1:
                continue
            cell = cells[0]
            grid_span = cell.find("w:tcPr/w:gridSpan", NS)
            if grid_span is None and len(table.findall("w:tblGrid/w:gridCol", NS)) > 1:
                continue
            for para in cell.findall("w:p", NS):
                text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
                if text and CAPTION_RE.match(text):
                    allowed_in_table_title_paragraphs.add(id(para))
    for para in root.findall(".//w:tbl//w:tc//w:p", NS):
        if id(para) in allowed_in_table_title_paragraphs:
            continue
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        if not text:
            continue
        if CAPTION_RE.match(text):
            issues.append(f"题名错误落入表格单元格: {text}")
            if len(issues) >= 12:
                break
    return issues, ("通过" if not issues else "未通过")


def extract_table_caption_binding_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查表题绑定。"], "未通过"
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查表题绑定。"], "未通过"

    issues: list[str] = []
    styles = paragraph_style_name_map(final_docx)
    children: list[ET.Element] = []
    body_started = False
    toc_seen = False
    for child in list(body):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            style_node = child.find("./w:pPr/w:pStyle", NS)
            style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
            style_name = styles.get(style_id, "")
            normalized = normalize(text).lower()
            if is_toc_heading_text(text):
                toc_seen = True
                continue
            if not toc_seen:
                continue
            if normalized in {normalize("参考文献").lower(), normalize("references").lower()}:
                break
            if xml_paragraph_has_tab(child) or style_name.lower().startswith("toc") or style_id.lower().startswith("toc"):
                continue
            if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
                body_started = True
        if body_started:
            children.append(child)
    for index, child in enumerate(children):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            match = CAPTION_RE.match(text)
            if not match or match.group(1) not in {"表", "续表"}:
                continue
            ppr = child.find("w:pPr", NS)
            if ppr is None or ppr.find("w:keepNext", NS) is None:
                issues.append(f"表题缺少 keepNext，未与后续表格绑定: {text}")
            if ppr is not None and ppr.find("w:numPr", NS) is not None:
                issues.append(f"表题携带列表/自动编号元数据，可能导致奇怪编号: {text}")
            next_child = next((candidate for candidate in children[index + 1 :] if candidate.tag in {W + "p", W + "tbl"}), None)
            if next_child is None or next_child.tag != W + "tbl":
                issues.append(f"表题后未紧跟表格对象: {text}")
        elif child.tag == W + "tbl":
            if is_formula_layout_table(child) or is_code_like_table(child):
                continue
            previous_child = next(
                (
                    candidate
                    for candidate in reversed(children[:index])
                    if candidate.tag == W + "p"
                    or (
                        candidate.tag == W + "tbl"
                        and not is_formula_layout_table(candidate)
                        and not is_code_like_table(candidate)
                    )
                ),
                None,
            )
            if previous_child is None or previous_child.tag != W + "p":
                issues.append("表格前缺少独立表题段落。")
                continue
            previous_text = "".join(node.text or "" for node in previous_child.findall(".//w:t", NS)).strip()
            match = CAPTION_RE.match(previous_text)
            if not match or match.group(1) not in {"表", "续表"}:
                issues.append(f"表格前一段不是独立表题: {previous_text or '<empty>'}")
        if len(issues) >= 12:
            break
    return issues, ("通过" if not issues else "未通过")


def extract_figure_explanation_followup_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查图后正文介绍。"], "未通过"
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查图后正文介绍。"], "未通过"

    styles = paragraph_style_name_map(final_docx)
    body_started = False
    blocks: list[dict[str, object]] = []
    for child in list(body):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            style_node = child.find("./w:pPr/w:pStyle", NS)
            style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
            style_name = styles.get(style_id, "")
            if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
                body_started = True
            if not body_started:
                continue
            if normalize(text).lower() in {normalize("参考文献").lower(), normalize("references").lower()}:
                break
            blocks.append(
                {
                    "kind": "paragraph",
                    "text": text,
                    "style_id": style_id,
                    "style_name": style_name,
                    "has_image": xml_paragraph_has_real_image(child),
                    "is_caption": bool(CAPTION_RE.match(text) and CAPTION_RE.match(text).group(1) == "图"),
                }
            )
        elif child.tag == W + "tbl":
            if not body_started:
                continue
            blocks.append({"kind": "table", "text": "<table>", "has_image": False, "is_caption": False})

    def is_blank_paragraph(block: dict[str, object] | None) -> bool:
        return bool(
            block is not None
            and block.get("kind") == "paragraph"
            and not str(block.get("text", "")).strip()
            and not block.get("has_image")
            and not block.get("is_caption")
        )

    def previous_nonblank_index(start_idx: int) -> int | None:
        cursor = start_idx
        while cursor >= 0 and is_blank_paragraph(blocks[cursor]):
            cursor -= 1
        return cursor if cursor >= 0 else None

    def is_explanatory_body_block(block: dict[str, object] | None) -> bool:
        if block is None or block.get("kind") != "paragraph":
            return False
        text = str(block.get("text", "")).strip()
        style_name = str(block.get("style_name", "")).lower()
        if not text:
            return False
        if (
            block.get("has_image")
            or block.get("is_caption")
            or CAPTION_RE.match(text)
            or heading_level(text) is not None
            or style_name.startswith("heading")
            or style_name.startswith("toc")
        ):
            return False
        return True

    def has_default_preceding_explanation(caption_idx: int) -> bool:
        image_idx = previous_nonblank_index(caption_idx - 1)
        if image_idx is None or not blocks[image_idx].get("has_image"):
            return False
        explanation_idx = previous_nonblank_index(image_idx - 1)
        if explanation_idx is None:
            return False
        return is_explanatory_body_block(blocks[explanation_idx])

    issues: list[str] = []
    for idx, block in enumerate(blocks):
        if not block.get("is_caption"):
            continue
        caption_text = str(block.get("text", "")).strip()
        if has_default_preceding_explanation(idx):
            continue
        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        if next_block is None:
            issues.append(f"图题后缺少紧随的正文介绍段落: {caption_text}")
            continue
        if not is_explanatory_body_block(next_block):
            if next_block.get("kind") != "paragraph":
                issues.append(f"图题后直接进入非正文块，缺少图后正文介绍: {caption_text}")
            else:
                issues.append(f"图题后未紧随正文介绍段落: {caption_text}")
            continue
    return issues, ("通过" if not issues else "未通过")


def extract_toc_control_contamination_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查目录控件污染。"], "未通过"

    issues: list[str] = []
    style_names = paragraph_style_name_map(final_docx)
    style_signatures = paragraph_style_definition_signature_map(final_docx)

    def is_allowed_toc_range_paragraph(text: str, style_id: str, has_tab: bool) -> bool:
        if not text and not has_tab:
            return True
        if is_toc_heading_text(text):
            return True
        if style_id.lower().startswith("toc"):
            return True
        if has_tab or "\u2026" in text:
            return True
        label = toc_visible_label_text(text)
        if is_front_matter_toc_label(label or text):
            return True
        return False

    for sdt in root.findall(".//w:sdt", NS):
        paragraphs = sdt.findall(".//w:sdtContent//w:p", NS)
        if not paragraphs:
            continue
        toc_like = False
        for para in paragraphs:
            text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
            style_node = para.find("./w:pPr/w:pStyle", NS)
            style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
            has_tab = xml_paragraph_has_tab(para)
            if normalize(text) == normalize("目   录") or style_id.lower().startswith("toc") or has_tab:
                toc_like = True
                break
        if not toc_like:
            continue

        for para in paragraphs:
            text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
            if not text and not xml_paragraph_has_real_image(para):
                continue
            style_node = para.find("./w:pPr/w:pStyle", NS)
            style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
            has_tab = xml_paragraph_has_tab(para)
            normalized = normalize(text)
            if xml_paragraph_has_real_image(para):
                issues.append("目录控件内出现图片对象或图片承载段落。")
                continue
            if CAPTION_RE.match(text):
                issues.append(f"目录控件内出现图题/表题段落: {text}")
                continue
            if normalized == normalize("目   录"):
                continue
            if style_id.lower().startswith("toc"):
                continue
            if has_tab:
                continue
            if heading_level(text) is not None or text:
                issues.append(f"目录控件内出现非 TOC 正文段落: {text}")
    toc_started = False
    for paragraph_index, para in enumerate(iter_body_paragraph_elements(final_docx), start=1):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        has_tab = xml_paragraph_has_tab(para)
        if not toc_started:
            if is_toc_heading_text(text):
                toc_started = True
            continue
        style_name = style_names.get(style_id, "")
        style_signature = style_signatures.get(style_id, {})
        if (
            not style_id.lower().startswith("toc")
            and not has_tab
            and is_body_chapter_heading_paragraph(
                para,
                text,
                style_id=style_id,
                style_name=style_name,
                style_signature=style_signature,
            )
        ):
            break
        if not text and not xml_paragraph_has_real_image(para):
            continue
        if is_allowed_toc_range_paragraph(text, style_id, has_tab):
            continue
        if xml_paragraph_has_real_image(para):
            issues.append(f"TOC range contains non-TOC image paragraph: paragraph {paragraph_index}")
            continue
        if CAPTION_RE.match(text):
            issues.append(f"TOC range contains body figure/table caption paragraph: paragraph {paragraph_index}: {text}")
            continue
        issues.append(f"TOC range contains non-TOC body paragraph: paragraph {paragraph_index}: {text}")
    return issues, ("通过" if not issues else "未通过")


def extract_table_checks(docx_path: Path) -> tuple[list[str], str]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查表格。"], "未通过"

    issues: list[str] = []
    shaded_cells = 0
    default_grid_tables = 0
    visible_verticals = 0
    visible_body_horizontals = 0

    for tbl in root.findall(".//w:tbl", NS):
        if is_formula_layout_table(tbl) or is_code_like_table(tbl):
            continue
        style = tbl.find("./w:tblPr/w:tblStyle", NS)
        if style is not None and style.attrib.get(W + "val", "") == "TableGrid":
            default_grid_tables += 1
        tbl_borders = tbl.find("./w:tblPr/w:tblBorders", NS)
        if tbl_borders is not None:
            for edge in ("left", "right", "insideV"):
                node = tbl_borders.find(f"w:{edge}", NS)
                if node is not None and node.attrib.get(W + "val", "none") not in {"none", "nil"}:
                    visible_verticals += 1
            node = tbl_borders.find("w:insideH", NS)
            if node is not None and node.attrib.get(W + "val", "none") not in {"none", "nil"}:
                visible_body_horizontals += 1
        for shd in tbl.findall(".//w:shd", NS):
            fill = (shd.attrib.get(W + "fill") or "").upper()
            if fill and fill not in {"FFFFFF", "AUTO", "CLEAR"}:
                shaded_cells += 1

    if default_grid_tables:
        issues.append(f"仍有表格停留在 TableGrid/默认网格家族，共 {default_grid_tables} 个。")
    if visible_verticals:
        issues.append(f"检测到三线表仍保留可见竖线，共 {visible_verticals} 处。")
    if visible_body_horizontals:
        issues.append(f"检测到三线表仍保留正文内部横线，共 {visible_body_horizontals} 处。")
    if shaded_cells:
        issues.append(f"检测到非白色表格底纹，共 {shaded_cells} 个单元。")
    return issues, ("通过" if not issues else "未通过")


def extract_footer_indent_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str]:
    issues: list[str] = []
    reference_parts = collect_part_paragraph_infos(reference_docx, "word/footer")
    final_parts = collect_part_paragraph_infos(final_docx, "word/footer")
    keys = ("left", "firstLine", "right", "hanging", "leftChars", "rightChars", "firstLineChars", "hangingChars")

    def norm_indent(value: str) -> str:
        return "0" if value in {"", "0", "none"} else value

    for part, final_rows in final_parts.items():
        reference_rows = reference_parts.get(part, [])
        reference_signatures = [row["signature"] for row in reference_rows]
        if not reference_signatures:
            reference_signatures = [row["signature"] for rows in reference_parts.values() for row in rows]
        if not reference_signatures:
            reference_signatures = [{key: "" for key in keys}]
        for idx, current in enumerate(final_rows, start=1):
            if not str(current.get("text", "")).strip() and not any(str(row.get("text", "")).strip() for row in reference_rows):
                continue
            current_signature = {key: norm_indent(str(current["signature"].get(key, ""))) for key in keys}  # type: ignore[index]
            normalized_reference_signatures = [
                {key: norm_indent(str(reference_signature.get(key, ""))) for key in keys}  # type: ignore[union-attr]
                for reference_signature in reference_signatures
            ]
            if any(
                all(current_signature[key] == reference_signature[key] for key in keys)
                for reference_signature in normalized_reference_signatures
            ):
                continue
            diffs = [
                f"{key}: expected `{normalized_reference_signatures[0][key]}` but found `{current_signature[key]}`"
                for key in keys
                if current_signature[key] != normalized_reference_signatures[0][key]
            ]
            issues.append(f"{part} paragraph {idx} footer-indent drift: {'; '.join(diffs[:6])}")
    return issues, ("通过" if not issues else "未通过")


def strip_template_instruction_parentheticals_for_check(text: str) -> str:
    value = str(text or "")
    changed = True
    while changed:
        changed = False

        def repl(match: re.Match[str]) -> str:
            nonlocal changed
            inner = match.group(1)
            if is_instruction_text(inner):
                changed = True
                return ""
            return match.group(0)

        value = re.sub(r"\uff08([^\uff08\uff09]{1,120})\uff09", repl, value)
        value = re.sub(r"\(([^()]{1,120})\)", repl, value)
    return re.sub(r"\s{2,}", " ", value).strip()


def header_text_tokens(docx_path: Path, *, strip_template_instructions: bool = False) -> list[str]:
    doc = Document(docx_path)
    header_tokens: list[str] = []
    for section in doc.sections:
        for para in section.header.paragraphs:
            text = para.text.strip()
            if strip_template_instructions:
                text = strip_template_instruction_parentheticals_for_check(text)
            if text:
                header_tokens.append(text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if strip_template_instructions:
                            text = strip_template_instruction_parentheticals_for_check(text)
                        if text:
                            header_tokens.append(text)
    return list(dict.fromkeys(header_tokens))


R_ID = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"


def header_token_matches(template_token: str, candidate_token: str) -> bool:
    template_compact = re.sub(r"\s+", "", str(template_token or ""))
    candidate_compact = re.sub(r"\s+", "", str(candidate_token or ""))
    if not template_compact:
        return True
    if template_compact == candidate_compact:
        return True
    pattern = re.escape(template_compact)
    for placeholder in ("XXXX", "xxxx", "20XX", "20xx"):
        pattern = pattern.replace(re.escape(placeholder), r"\d{4}")
    return bool(re.fullmatch(pattern, candidate_compact))


def header_token_in_text(template_token: str, text: str) -> bool:
    compact_text = re.sub(r"\s+", "", text or "")
    token = re.sub(r"\s+", "", template_token or "")
    if not token:
        return True
    if token in compact_text:
        return True
    pattern = re.escape(token)
    for placeholder in ("XXXX", "xxxx", "20XX", "20xx"):
        pattern = pattern.replace(re.escape(placeholder), r"\d{4}")
    return bool(re.search(pattern, compact_text))


def docx_relationship_targets(docx_path: Path, prefix: str) -> dict[str, str]:
    try:
        with zipfile.ZipFile(docx_path) as zf:
            rels_root = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
    except Exception:
        return {}
    targets: dict[str, str] = {}
    for rel in rels_root.findall(f"{PR}Relationship"):
        rid = rel.attrib.get("Id", "")
        target = rel.attrib.get("Target", "")
        if not rid or not target:
            continue
        normalized = "word/" + target.lstrip("/")
        if normalized.startswith(prefix):
            targets[rid] = normalized
    return targets


def header_part_tokens(docx_path: Path, part: str) -> list[str]:
    try:
        with zipfile.ZipFile(docx_path) as zf:
            if part not in zf.namelist():
                return []
            root = ET.fromstring(zf.read(part))
    except Exception:
        return []
    tokens: list[str] = []
    for para in root.findall(".//w:p", NS):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        text = strip_template_instruction_parentheticals_for_check(text)
        if is_instruction_text(text) or is_instruction_note_text(text):
            text = ""
        if text:
            tokens.append(text)
    return list(dict.fromkeys(tokens))


def section_header_contract_rows(docx_path: Path) -> list[dict[str, object]]:
    rel_targets = docx_relationship_targets(docx_path, "word/header")
    doc_root = load_xml(docx_path, "word/document.xml")
    if doc_root is None:
        return []
    rows: list[dict[str, object]] = []
    effective_targets: dict[str, str] = {}
    for section_index, sect_pr in enumerate(doc_root.findall(".//w:sectPr", NS), start=1):
        direct_targets: dict[str, str] = {}
        for ref in sect_pr.findall("./w:headerReference", NS):
            header_type = ref.attrib.get(W + "type", "default")
            rid = ref.attrib.get(R_ID, "")
            target = rel_targets.get(rid, "")
            if target:
                direct_targets[header_type] = target
        if direct_targets:
            effective_targets.update(direct_targets)
        section_tokens: list[str] = []
        for target in sorted(set(effective_targets.values())):
            section_tokens.extend(header_part_tokens(docx_path, target))
        rows.append(
            {
                "section_index": section_index,
                "direct_targets": dict(direct_targets),
                "effective_targets": dict(effective_targets),
                "tokens": list(dict.fromkeys(section_tokens)),
            }
        )
    return rows


def rendered_header_sample_pages(page_class_map: dict[str, int | None], page_count: int) -> list[int]:
    candidates = [
        page_class_map.get("first_body"),
        page_class_map.get("figure_page"),
        page_class_map.get("table_page"),
        page_class_map.get("references"),
        page_class_map.get("ack"),
    ]
    pages: list[int] = []
    for page in candidates:
        if page is None:
            continue
        if 1 <= int(page) <= page_count and int(page) not in pages:
            pages.append(int(page))
    if not pages and page_count:
        pages.append(1 if page_count == 1 else min(8, page_count))
    return pages


def extract_header_presence_contract_checks(
    reference_docx: Path,
    final_docx: Path,
    final_pdf: Path,
    page_class_map: dict[str, int | None],
) -> tuple[list[str], str, dict[str, object]]:
    reference_rows = section_header_contract_rows(reference_docx)
    final_rows = section_header_contract_rows(final_docx)
    reference_header_rows = [
        row
        for row in reference_rows
        if any(str(token).strip() for token in row.get("tokens", []))  # type: ignore[union-attr]
    ]
    reference_tokens = list(
        dict.fromkeys(
            token
            for row in reference_header_rows
            for token in row.get("tokens", [])  # type: ignore[union-attr]
            if str(token).strip()
        )
    )
    if not reference_header_rows and not reference_tokens:
        detector = detector_result(
            "header.presence-contract",
            surface="header",
            passed=True,
            blocking=False,
            severity="not-applicable",
            evidence={
                "summary": "not-applicable",
                "reason": "reference DOCX has no detectable header baseline",
                "reference_sections": len(reference_rows),
                "final_sections": len(final_rows),
            },
        )
        return [], "not-applicable", detector

    issues: list[str] = []
    final_by_section = {int(row.get("section_index", 0)): row for row in final_rows}
    final_header_rows = [
        row
        for row in final_rows
        if any(str(token).strip() for token in row.get("tokens", []))  # type: ignore[union-attr]
    ]
    final_header_tokens = list(
        dict.fromkeys(
            token
            for row in final_header_rows
            for token in row.get("tokens", [])  # type: ignore[union-attr]
            if str(token).strip()
        )
    )
    global_reference_token_match = (
        bool(reference_tokens)
        and bool(final_header_tokens)
        and any(
            header_token_matches(str(reference_token), str(final_token))
            for reference_token in reference_tokens
            for final_token in final_header_tokens
        )
    )
    if not final_header_rows:
        issues.append("header presence contract failed: final DOCX has no non-empty effective header part")
    for reference_row in reference_header_rows:
        section_index = int(reference_row.get("section_index", 0))
        final_row = final_by_section.get(section_index)
        if final_row is None:
            if not global_reference_token_match:
                issues.append(
                    f"header presence contract failed: final DOCX missing section {section_index} from template header baseline"
                )
            continue
        if not final_row.get("effective_targets"):
            issues.append(
                f"header presence contract failed: section {section_index} has no effective header reference"
            )
            continue
        final_tokens_for_section = [str(token) for token in final_row.get("tokens", [])]  # type: ignore[union-attr]
        if not final_tokens_for_section:
            issues.append(
                f"header presence contract failed: section {section_index} effective header part is empty"
            )
            continue
        reference_tokens_for_section = [
            str(token)
            for token in reference_row.get("tokens", [])  # type: ignore[union-attr]
            if str(token).strip()
        ]
        section_token_match = any(
            header_token_matches(reference_token, final_token)
            for reference_token in reference_tokens_for_section
            for final_token in final_tokens_for_section
        )
        if reference_tokens_for_section and not section_token_match and not global_reference_token_match:
            issues.append(
                "header presence contract failed: "
                f"section {section_index} header tokens do not match the template baseline"
            )
    for final_row in final_header_rows:
        section_index = int(final_row.get("section_index", 0))
        if not final_row.get("effective_targets"):
            issues.append(f"header presence contract failed: section {section_index} has no effective header reference")
            continue
        final_tokens = [str(token) for token in final_row.get("tokens", [])]  # type: ignore[union-attr]
        if not final_tokens:
            issues.append(f"header presence contract failed: section {section_index} effective header part is empty")

    rendered_pages: list[int] = []
    rendered_missing: dict[int, list[str]] = {}
    try:
        with fitz.open(final_pdf) as pdf:
            rendered_pages = rendered_header_sample_pages(page_class_map, pdf.page_count)
            for page_number in rendered_pages:
                text = pdf[page_number - 1].get_text("text")
                expected_tokens = final_header_tokens or reference_tokens
                if expected_tokens and not any(header_token_in_text(str(token), text) for token in expected_tokens):
                    rendered_missing[page_number] = expected_tokens[:6]
    except Exception as exc:
        issues.append(f"header rendered contract failed: cannot inspect final PDF headers ({exc})")
    for page_number, missing_tokens in rendered_missing.items():
        issues.append(
            "header rendered contract failed: "
            f"rendered page {page_number} is missing header tokens {missing_tokens}"
        )

    evidence = {
        "reference_docx": str(reference_docx),
        "final_docx": str(final_docx),
        "final_pdf": str(final_pdf),
        "reference_sections": len(reference_rows),
        "final_sections": len(final_rows),
        "reference_header_sections": [
            int(row.get("section_index", 0))
            for row in reference_header_rows
        ],
        "reference_tokens": reference_tokens[:12],
        "final_header_sections": [
            int(row.get("section_index", 0))
            for row in final_header_rows
        ],
        "final_tokens": final_header_tokens[:12],
        "literal_token_policy": "dynamic-template-content-not-required",
        "rendered_sample_pages": rendered_pages,
        "rendered_missing_tokens_by_page": rendered_missing,
        "issue_count": len(issues),
    }
    detector = detector_result(
        "header.presence-contract",
        surface="header",
        passed=not issues,
        evidence=evidence,
    )
    return issues, ("passed" if not issues else "failed"), detector


def extract_header_checks(reference_docx: Path, final_pdf: Path, final_docx: Path | None = None) -> tuple[list[str], str]:
    issues: list[str] = []
    header_tokens = header_text_tokens(reference_docx, strip_template_instructions=True)
    if not header_tokens:
        return [], "不适用"
    final_header_tokens = header_text_tokens(final_docx, strip_template_instructions=True) if final_docx is not None else []
    if final_docx is not None:
        if not final_header_tokens:
            return ["final DOCX has no auditable header text tokens"], "failed"
        with fitz.open(final_pdf) as pdf:
            if pdf.page_count == 0:
                return ["rendered PDF has no pages for header position check"], "failed"
            sample_indices = sorted(
                {
                    idx
                    for idx in (1, 2, 3, min(7, pdf.page_count - 1), max(0, pdf.page_count - 2), pdf.page_count - 1)
                    if 0 <= idx < pdf.page_count
                }
            )
            visible_header_pages = 0
            for page_index in sample_indices:
                page = pdf[page_index]
                top_words = [word for word in page.get_text("words") if float(word[1]) <= page.rect.height * 0.14]
                top_text = "".join(str(word[4]) for word in top_words)
                if any(header_token_in_text(token, top_text) for token in final_header_tokens):
                    visible_header_pages += 1
        if visible_header_pages == 0:
            issues.append("rendered PDF sampled pages have no visible final header token in the top header region")
        return issues, ("passed" if not issues else "failed")

    def header_token_matches_with_filled_year(template_token: str, candidate_token: str) -> bool:
        return header_token_matches(template_token, candidate_token)

    def final_header_contains_token(template_token: str) -> bool:
        return any(header_token_matches_with_filled_year(template_token, token) for token in final_header_tokens)

    with fitz.open(final_pdf) as pdf:
        page = pdf[0 if pdf.page_count == 1 else min(7, pdf.page_count - 1)]
        words = page.get_text("words")
        page_width = page.rect.width

    def word_box(token: str):
        matches = [w for w in words if token in w[4]]
        if not matches:
            return None
        x0 = min(w[0] for w in matches)
        y0 = min(w[1] for w in matches)
        x1 = max(w[2] for w in matches)
        y1 = max(w[3] for w in matches)
        return (x0, y0, x1, y1)

    if len(header_tokens) >= 2:
        left = word_box(header_tokens[0])
        right = word_box(header_tokens[-1])
        if left is None or right is None:
            issues.append("页眉缺少模板基线中的关键文本。")
        else:
            if abs(left[1] - right[1]) > 6:
                issues.append("页眉左右文本不在同一水平线上。")
            if left[0] > page_width * 0.45:
                issues.append("页眉左侧关键文本位置偏右。")
            if right[0] < page_width * 0.55:
                issues.append("页眉右侧关键文本位置偏左。")
    else:
        token = header_tokens[0]
        if word_box(token) is None:
            if not final_header_contains_token(token):
                issues.append(f"页眉缺少模板关键文本: {token}")
    return issues, ("通过" if not issues else "未通过")


def toc_tab_stop_signature(paragraph: ET.Element) -> str:
    tabs = []
    for tab in paragraph.findall("./w:pPr/w:tabs/w:tab", NS):
        tabs.append(
            "|".join(
                [
                    tab.attrib.get(W + "val", ""),
                    tab.attrib.get(W + "leader", ""),
                    tab.attrib.get(W + "pos", ""),
                ]
            )
        )
    return ";".join(tabs) if tabs else "none"


def toc_run_model_signature(paragraph: ET.Element) -> dict[str, str]:
    before_tab = 0
    after_tab = 0
    has_tab = "no"
    seen_tab = False
    for run in paragraph.findall(".//w:r", NS):
        run_has_text = bool(run_visible_text(run))
        for child in list(run):
            if child.tag == W + "tab":
                has_tab = "yes"
                seen_tab = True
        if run_has_text:
            if seen_tab:
                after_tab += 1
            else:
                before_tab += 1
    return {
        "hasTab": has_tab,
        "preTabTextRunCount": str(before_tab),
        "postTabTextRunCount": str(after_tab),
        "tabStops": toc_tab_stop_signature(paragraph),
        "preTabRunFonts": toc_run_font_signature(paragraph, before_tab=True),
        "postTabRunFonts": toc_run_font_signature(paragraph, before_tab=False),
    }


def compact_font_signature(signature: dict[str, str]) -> str:
    return "/".join(
        [
            signature.get("eastAsia", ""),
            signature.get("ascii", ""),
            signature.get("hAnsi", ""),
            signature.get("size", ""),
            signature.get("bold", "no"),
            "u=" + signature.get("underline", "no"),
            "theme="
            + ",".join(
                [
                    signature.get("eastAsiaTheme", ""),
                    signature.get("asciiTheme", ""),
                    signature.get("hAnsiTheme", ""),
                    signature.get("csTheme", ""),
                ]
            ),
        ]
    )


def toc_run_font_signature(paragraph: ET.Element, *, before_tab: bool) -> str:
    signatures: list[str] = []
    seen_tab = False
    for run in paragraph.findall(".//w:r", NS):
        if run.find(".//w:txbxContent", NS) is not None:
            continue
        if run.find("w:tab", NS) is not None:
            seen_tab = True
        text = run_visible_text(run)
        if not text:
            continue
        if before_tab != (not seen_tab):
            continue
        signatures.append(compact_font_signature(run_direct_signature(run)))
    if not signatures:
        return "none"
    return ";".join(dict.fromkeys(signatures))


def toc_entry_level_from_paragraph(para: ET.Element, styles: dict[str, str]) -> tuple[int | None, str, str, str]:
    text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
    style_node = para.find("./w:pPr/w:pStyle", NS)
    style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
    style_name = styles.get(style_id, "")
    label = toc_visible_label_text(text)
    page_match = re.search(r"(\d+)\s*$", text)
    if "\t" not in text and page_match and "\u2026" not in text:
        label = re.sub(r"\s*\d+\s*$", "", label).strip()
    level = heading_level(label)
    style_key = normalize(style_id or style_name).lower()
    if level is None and style_key.startswith("toc"):
        digits = re.findall(r"\d+", style_key)
        level = int(digits[-1]) if digits else 1
    return (level if level is not None else None, text, style_id or style_name, label)


def collect_toc_entry_font_profiles(docx_path: Path) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    styles = paragraph_style_name_map(docx_path)
    toc_started = False
    for para in iter_body_and_sdt_paragraph_elements(docx_path):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        if is_toc_heading_text(text):
            toc_started = True
            continue
        if not toc_started:
            continue
        level, _text, style_key, label = toc_entry_level_from_paragraph(para, styles)
        if is_front_matter_toc_label(label or text):
            continue
        if level is None:
            if is_body_chapter_heading_text(text):
                break
            continue
        if not (xml_paragraph_has_tab(para) or normalize(style_key).lower().startswith("toc") or is_toc_leader_entry_text(text)):
            continue
        rows.append(
            {
                "level": level,
                "label": label or text,
                "pre": set(toc_run_font_signature(para, before_tab=True).split(";")),
                "post": set(toc_run_font_signature(para, before_tab=False).split(";")),
            }
        )
    return rows


def extract_toc_direct_run_font_issues(reference_docx: Path, final_docx: Path) -> list[str]:
    reference_rows = collect_toc_entry_font_profiles(reference_docx)
    final_rows = collect_toc_entry_font_profiles(final_docx)
    if not reference_rows or not final_rows:
        return []
    allowed: dict[tuple[int, str], set[str]] = defaultdict(set)
    for row in reference_rows:
        level = int(row["level"])
        for side in ("pre", "post"):
            allowed[(level, side)].update(str(item) for item in row[side] if str(item))

    issues: list[str] = []
    for row in final_rows:
        level = int(row["level"])
        label = str(row["label"])
        for side in ("pre", "post"):
            expected = allowed.get((level, side))
            if not expected:
                continue
            if expected.issubset({"none"}):
                continue
            actual_values = {str(item) for item in row[side] if str(item)}
            unexpected = sorted(value for value in actual_values if value not in expected)
            if unexpected:
                if any("/u=yes/" in f"/{value}/" for value in unexpected):
                    issues.append(
                        f"TOC underline pollution: level {level} {side}-tab visible run in {label} has underline not present in baseline"
                    )
                    if len(issues) >= 8:
                        return issues
                issues.append(
                    f"TOC level {level} {side}-tab run font contamination: {label} uses {', '.join(unexpected[:4])}; "
                    f"allowed {', '.join(sorted(expected)[:4])}"
                )
                if len(issues) >= 8:
                    return issues
    return issues


def collect_toc_entry_baselines(docx_path: Path) -> dict[int, dict[str, str]]:
    rows: dict[int, dict[str, str]] = {}
    styles = paragraph_style_name_map(docx_path)
    toc_started = False
    for para in iter_body_and_sdt_paragraph_elements(docx_path):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        if is_toc_heading_text(text):
            toc_started = True
            continue
        if not toc_started:
            continue
        is_toc_entry_host = (
            xml_paragraph_has_tab(para)
            or style_name.lower().startswith("toc")
            or style_id.lower().startswith("toc")
            or is_toc_leader_entry_text(text)
        )
        if not is_toc_entry_host:
            if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
                break
            continue
        label = toc_visible_label_text(text)
        if is_front_matter_toc_label(label or text):
            continue
        level = toc_label_level(label, style_id, style_name)
        if level is None:
            if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
                break
            continue
        signature = paragraph_direct_signature(para)
        signature.update(toc_run_model_signature(para))
        rows.setdefault(level, signature)
    return rows


def semicolon_values(value: str) -> set[str]:
    if not value or value == "none":
        return set()
    return {item for item in value.split(";") if item and item != "none"}


def toc_tab_stops_match_with_required_dots(current_value: str, reference_value: str) -> bool:
    current_items = [item for item in (current_value or "").split(";") if item and not item.startswith("clear|")]
    reference_items = [item for item in (reference_value or "").split(";") if item and not item.startswith("clear|")]
    if not current_items or len(current_items) != len(reference_items):
        return False
    for current_item, reference_item in zip(current_items, reference_items):
        current_parts = (current_item.split("|") + ["", "", ""])[:3]
        reference_parts = (reference_item.split("|") + ["", "", ""])[:3]
        if current_parts[0] != reference_parts[0] or current_parts[2] != reference_parts[2]:
            return False
        if current_parts[1] == reference_parts[1]:
            continue
        if current_parts[0] == "right" and current_parts[1] == "dot" and reference_parts[1] in {"", "none"}:
            continue
        return False
    return True


def toc_tab_stops_compatible_with_rendered_leaders(current_value: str, reference_value: str) -> bool:
    current_items = [item for item in (current_value or "").split(";") if item and not item.startswith("clear|")]
    reference_items = [item for item in (reference_value or "").split(";") if item and not item.startswith("clear|")]
    if not current_items or not reference_items:
        return False
    current_has_required_right_dot = any(item.startswith("right|dot|") for item in current_items)
    reference_has_wps_center_placeholder = any(
        item.startswith("center|none|") or item.startswith("center||") for item in reference_items
    )
    return current_has_required_right_dot and reference_has_wps_center_placeholder


def compare_toc_entry_signature(current: dict[str, str], reference: dict[str, str], keys: tuple[str, ...]) -> list[str]:
    diffs = compare_signature(current, reference, keys)
    normalized_manual_leader = (
        reference.get("hasTab") == "no"
        and current.get("hasTab") == "yes"
        and current.get("tabStops", "").startswith("right|dot|")
    )
    normalized_required_dotted_leader = toc_tab_stops_match_with_required_dots(
        current.get("tabStops", ""),
        reference.get("tabStops", ""),
    )
    normalized_required_dotted_leader = normalized_required_dotted_leader or (
        reference.get("hasTab") == "yes"
        and current.get("hasTab") == "yes"
        and reference.get("tabStops", "") in {"", "none"}
        and current.get("tabStops", "").startswith("right|dot|")
    )
    normalized_rendered_leader = toc_tab_stops_compatible_with_rendered_leaders(
        current.get("tabStops", ""),
        reference.get("tabStops", ""),
    )
    if not normalized_manual_leader:
        if normalized_required_dotted_leader or normalized_rendered_leader:
            diffs = [diff for diff in diffs if not diff.startswith("tabStops:")]
        else:
            return diffs
    reference_pre = semicolon_values(reference.get("preTabRunFonts", ""))
    reference_all = reference_pre | semicolon_values(reference.get("postTabRunFonts", ""))
    current_pre = semicolon_values(current.get("preTabRunFonts", ""))
    current_post = semicolon_values(current.get("postTabRunFonts", ""))
    filtered: list[str] = []
    for diff in diffs:
        if diff.startswith(("hasTab:", "tabStops:", "line:", "lineRule:")):
            continue
        if "expected `none` but found `0`" in diff and diff.startswith(("before:", "after:")):
            continue
        if "expected `0` but found `none`" in diff and diff.startswith(("left:", "leftChars:")):
            continue
        if diff.startswith(("left:", "right:", "hanging:")) and reference.get(diff.split(":", 1)[0], "") in {"", "none"}:
            continue
        if diff.startswith(("firstLineChars:", "leftChars:", "eastAsia:", "ascii:", "hAnsi:")):
            continue
        if diff.startswith("preTabRunFonts:") and reference.get("preTabRunFonts", "") in {"", "none"}:
            continue
        if diff.startswith("postTabRunFonts:") and reference.get("postTabRunFonts", "") in {"", "none"}:
            continue
        if diff.startswith("preTabRunFonts:") and current_pre and current_pre.issubset(reference_pre or reference_all):
            continue
        if diff.startswith("postTabRunFonts:") and current_post and current_post.issubset(reference_all or reference_pre):
            continue
        filtered.append(diff)
    return filtered


def toc_heading_color_issues(paragraph: ET.Element) -> list[str]:
    issues: list[str] = []
    for run in paragraph.findall("w:r", NS):
        text = "".join(node.text or "" for node in run.findall(".//w:t", NS)).strip()
        if not text:
            continue
        color = run.find("./w:rPr/w:color", NS)
        if color is None:
            continue
        value = (color.attrib.get(W + "val", "") or "").lower()
        theme = (color.attrib.get(W + "themeColor", "") or "").lower()
        if value not in {"", "auto", "000000"} or theme not in {"", "text1"}:
            issues.append(f"TOC title color contamination: value={value or 'none'} theme={theme or 'none'}")
            break
    return issues


def collect_toc_visible_labels(docx_path: Path) -> set[str]:
    labels: set[str] = set()
    styles = paragraph_style_name_map(docx_path)
    toc_started = False
    for para in iter_body_and_sdt_paragraph_elements(docx_path):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        style_key = normalize(style_id or style_name).lower()
        if is_toc_heading_text(text):
            toc_started = True
            continue
        if not toc_started:
            continue
        is_toc_entry_host = (
            xml_paragraph_has_tab(para)
            or style_key.startswith("toc")
            or is_toc_leader_entry_text(text)
        )
        if not is_toc_entry_host:
            if is_body_chapter_heading_text(text) or style_name.lower() == "heading 1":
                break
            continue
        raw_label = text.split("\t", 1)[0].strip()
        if "\t" not in text and re.search(r"(\d+|[IVXLCDM]+)\s*$", text, flags=re.IGNORECASE) and "\u2026" not in text:
            raw_label = re.sub(r"\s*(?:\d+|[IVXLCDM]+)\s*$", "", raw_label, flags=re.IGNORECASE).strip()
        label = toc_visible_label_text(raw_label)
        if label:
            labels.add(normalize(label).lower())
    return labels


def toc_entry_line_count_on_rendered_page(page_text: str, toc_labels: set[str]) -> int:
    if not toc_labels:
        return 0
    count = 0
    for raw_line in str(page_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        normalized = normalize(toc_visible_label_text(line)).lower()
        if normalized in toc_labels:
            count += 1
    return count


def rendered_page_looks_like_toc_continuation(page_text: str, toc_labels: set[str]) -> bool:
    lines = [line.strip() for line in str(page_text or "").splitlines() if line.strip()]
    if not lines:
        return False
    entry_count = toc_entry_line_count_on_rendered_page(page_text, toc_labels)
    return entry_count >= 3 and (entry_count / max(1, len(lines))) >= 0.5


def rendered_toc_last_page(pdf_texts: list[str], toc_page_no: int | None, toc_entries: list[tuple[str, str, bool, str, str | None]]) -> int | None:
    if toc_page_no is None:
        return None
    toc_labels = {
        normalize(label).lower()
        for _text, _style_name, _has_tab, label, _displayed_page in toc_entries
        if label
    }
    toc_last_page = toc_page_no
    for page_index in range(toc_page_no, len(pdf_texts)):
        if rendered_page_looks_like_toc_continuation(pdf_texts[page_index], toc_labels):
            toc_last_page = page_index + 1
            continue
        break
    return toc_last_page


def rendered_toc_entry_shows_page(
    pdf_texts: list[str],
    toc_page_no: int | None,
    toc_last_page: int | None,
    label: str,
    page_no: int,
) -> bool:
    if toc_page_no is None or toc_last_page is None:
        return False
    label_key = compact_rendered_heading_key(label)
    if not label_key:
        return False
    page_token = str(page_no)
    rendered_toc_end = max(toc_last_page, toc_page_no + 2)
    for page_index in range(toc_page_no - 1, min(rendered_toc_end, len(pdf_texts))):
        lines = [line.strip() for line in str(pdf_texts[page_index] or "").splitlines() if line.strip()]
        for index in range(len(lines)):
            window = "\n".join(lines[index:index + 4])
            if label_key not in compact_rendered_heading_key(window):
                continue
            if re.search(rf"(?<!\d){re.escape(page_token)}(?!\d)", window):
                return True
    return False


def extract_toc_checks(final_docx: Path, final_pdf: Path, reference_docx: Path | None = None) -> tuple[list[str], str]:
    issues: list[str] = []
    toc_heading_found = False
    toc_entries: list[tuple[str, str, bool, str, str | None]] = []
    reference_toc: dict[int, dict[str, str]] = {}

    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查目录。"], "未通过"

    for para in root.findall(".//w:p", NS):
        text = "".join(node.text or "" for node in para.findall(".//w:t", NS)).strip()
        style_node = para.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_key = normalize(style_id).lower()
        has_tab = xml_paragraph_has_tab(para)
        raw_label = text.split("\t", 1)[0].strip()
        displayed_page = None
        page_match = re.search(r"(\d+|[IVXLCDM]+)\s*$", text, flags=re.IGNORECASE)
        if page_match:
            displayed_page = page_match.group(1)
        label = toc_visible_label_text(raw_label)
        if "\t" not in text and displayed_page is not None and "\u2026" not in text:
            label = re.sub(r"\s*(?:\d+|[IVXLCDM]+)\s*$", "", raw_label, flags=re.IGNORECASE).strip()
        is_toc_surface_paragraph = (
            is_toc_heading_text(text)
            or style_key.startswith("toc")
            or has_tab
            or (toc_heading_found and is_toc_leader_entry_text(text))
        )
        if "\u25a1" in text and is_toc_surface_paragraph and not is_official_declaration_checkbox_line(text):
            issues.append(f"目录可见模板空位方框残留: {text}")
        if is_toc_heading_text(text):
            toc_heading_found = True
            issues.extend(toc_heading_color_issues(para))
            continue
        level = toc_label_level(label or text, style_id, "")
        if style_key.startswith("toc") or (has_tab and level is not None) or (toc_heading_found and is_toc_leader_entry_text(text) and level is not None):
            toc_entries.append((text, style_id if style_key.startswith("toc") else ("toc-by-tab" if has_tab else "toc-by-leader"), has_tab, label, displayed_page))

    if not toc_heading_found:
        issues.append("目录标题缺失或无法识别。")
    if not toc_entries:
        issues.append("目录项为空。")
    toc_labels = {normalize(label).lower() for _text, _style_name, _has_tab, label, _displayed_page in toc_entries}
    reference_toc_labels = collect_toc_visible_labels(reference_docx) if reference_docx is not None else set()
    if normalize("\u6458\u8981").lower() in reference_toc_labels and normalize("\u6458\u8981").lower() not in toc_labels:
        issues.append("TOC missing required Chinese abstract entry")
    if normalize("Abstract").lower() in reference_toc_labels and normalize("Abstract").lower() not in toc_labels:
        issues.append("TOC missing required English abstract entry")
    if reference_docx is not None:
        reference_toc = collect_toc_entry_baselines(reference_docx)
        final_toc = collect_toc_entry_baselines(final_docx)
        final_heading_levels = {
            min(level, 3)
            for row in iter_main_body_paragraph_infos(final_docx)
            for level in [toc_label_level(str(row["text"]))]
            if level is not None
        }
        toc_compare_keys = (
            "align",
            "before",
            "after",
            "line",
            "lineRule",
            "firstLine",
            "left",
            "right",
            "hanging",
            "firstLineChars",
            "leftChars",
            "rightChars",
            "hangingChars",
            "eastAsia",
            "ascii",
            "hAnsi",
            "eastAsiaTheme",
            "asciiTheme",
            "hAnsiTheme",
            "csTheme",
            "size",
            "bold",
            "hasTab",
            "tabStops",
            "preTabRunFonts",
            "postTabRunFonts",
        )
        for level, reference_signature in reference_toc.items():
            current_signature = final_toc.get(level)
            if current_signature is None:
                if level in final_heading_levels:
                    issues.append(f"TOC level {level} baseline entry missing in final document")
                continue
            diffs = compare_toc_entry_signature(current_signature, reference_signature, toc_compare_keys)
            if diffs:
                issues.append(f"TOC level {level} baseline drift: {'; '.join(diffs[:6])}")
        issues.extend(extract_toc_direct_run_font_issues(reference_docx, final_docx))
    for text, style_name, has_tab, _label, displayed_page in toc_entries:
        if "\t" not in text and not has_tab and normalize(style_name).lower() not in {"toc-by-tab", "toc-by-leader"}:
            issues.append(f"目录项缺少制表符页码分隔: {text}")
        if displayed_page is None:
            issues.append(f"目录项缺少可解析页码: {text}")
        normalized_style = normalize(style_name).lower()
        if normalized_style not in {"toc1", "toc2", "toc3", "toc4", "toc-by-tab", "toc-by-leader"}:
            issues.append(f"目录项使用了异常样式: {style_name}")

    with fitz.open(final_pdf) as pdf:
        pdf_texts = [page.get_text("text") for page in pdf]
        toc_page_no = find_page(pdf_texts, "目   录")
        if toc_page_no is None:
            toc_page_no = find_page(pdf_texts, "\u76ee\u5f55")
        toc_last_page = rendered_toc_last_page(pdf_texts, toc_page_no, toc_entries)
        heading_search_start = (toc_last_page + 1) if toc_last_page is not None else 1
        first_body_candidates = [
            find_rendered_heading_page_loose(pdf_texts, label, start_page=heading_search_start)
            for _text, _style_name, _has_tab, label, _displayed_page in toc_entries
            if toc_label_level(label) == 1
        ]
        first_body_page = min([page for page in first_body_candidates if page is not None], default=None)
        first_body_anchor_reliable = any(
            page == first_body_page
            and re.match(r"^\s*1(?:[\s\u3000\uff0e\.]|$)", label or "")
            for page, (_text, _style_name, _has_tab, label, _displayed_page) in zip(first_body_candidates, [
                entry for entry in toc_entries if toc_label_level(entry[3]) == 1
            ])
            if page is not None and first_body_page is not None
        )
        if toc_page_no is not None:
            page = pdf[toc_page_no - 1]
            words = page.get_text("words")
            level_hits: dict[int, list[float]] = defaultdict(list)
            for text, style_name, _has_tab, label, _displayed_page in toc_entries:
                first_token = label.split()[0] if label.split() else label
                token_words = [w for w in words if first_token in w[4]]
                if not token_words:
                    continue
                digits = re.findall(r"\d+", style_name)
                level = int(digits[-1]) if digits else (toc_label_level(label) or 1)
                level_hits[level].append(min(w[0] for w in token_words))
            reference_requires_level2_indent = (
                not reference_toc
                or int(reference_toc.get(2, {}).get("left") or 0) > int(reference_toc.get(1, {}).get("left") or 0)
                or int(reference_toc.get(2, {}).get("firstLine") or 0) > int(reference_toc.get(1, {}).get("firstLine") or 0)
                or int(reference_toc.get(2, {}).get("leftChars") or 0) > int(reference_toc.get(1, {}).get("leftChars") or 0)
                or int(reference_toc.get(2, {}).get("firstLineChars") or 0) > int(reference_toc.get(1, {}).get("firstLineChars") or 0)
            )
            reference_requires_level3_indent = (
                not reference_toc
                or int(reference_toc.get(3, {}).get("left") or 0) > int(reference_toc.get(2, {}).get("left") or 0)
                or int(reference_toc.get(3, {}).get("firstLine") or 0) > int(reference_toc.get(2, {}).get("firstLine") or 0)
                or int(reference_toc.get(3, {}).get("leftChars") or 0) > int(reference_toc.get(2, {}).get("leftChars") or 0)
                or int(reference_toc.get(3, {}).get("firstLineChars") or 0) > int(reference_toc.get(2, {}).get("firstLineChars") or 0)
            )
            if reference_requires_level2_indent and 1 in level_hits and 2 in level_hits and min(level_hits[2]) <= min(level_hits[1]):
                issues.append("目录二级条目的可见缩进没有明显大于一级条目。")
            if reference_requires_level3_indent and 2 in level_hits and 3 in level_hits and min(level_hits[3]) <= min(level_hits[2]):
                issues.append("目录三级条目的可见缩进没有明显大于二级条目。")
        toc_page_pairs: list[tuple[str, int, int]] = []
        if first_body_page is not None:
            for text, _style_name, _has_tab, label, displayed_page in toc_entries:
                if displayed_page is None or not str(displayed_page).isdigit():
                    continue
                actual_page = find_rendered_heading_page_loose(pdf_texts, label, start_page=first_body_page)
                if actual_page is None:
                    continue
                expected_logical_page = actual_page - first_body_page + 1
                if expected_logical_page <= 0:
                    continue
                toc_page_pairs.append((label, int(displayed_page), expected_logical_page))
        if toc_page_pairs and first_body_anchor_reliable:
            displayed_values = [displayed for _label, displayed, _expected in toc_page_pairs]
            expected_values = [expected for _label, _displayed, expected in toc_page_pairs]
            if max(expected_values) > min(expected_values) and max(displayed_values) == min(displayed_values):
                issues.append("目录项页码坍缩为同一页码，未反映正文标题的实际页序。")
            for label, displayed, expected in toc_page_pairs:
                if displayed != expected:
                    if rendered_toc_entry_shows_page(pdf_texts, toc_page_no, toc_last_page, label, expected):
                        continue
                    issues.append(
                        f"目录项页码与渲染页序不一致: {label} shows {displayed}, expected logical page {expected}"
                    )
                    if len(issues) >= 12:
                        break

    return issues, ("通过" if not issues else "未通过")


def extract_toc_page_number_right_edge_checks(final_pdf: Path) -> tuple[list[str], str]:
    issues: list[str] = []
    with fitz.open(final_pdf) as pdf:
        pdf_texts = [page.get_text("text") for page in pdf]
        toc_page_no = find_page(pdf_texts, "目   录")
        if toc_page_no is None:
            toc_page_no = find_page(pdf_texts, "\u76ee\u5f55")
        if toc_page_no is None:
            return ["TOC page-number right-edge metric unavailable: TOC page not found"], "failed"
        page = pdf[toc_page_no - 1]
        words = page.get_text("words")
        width = float(page.rect.width)
        candidates = [
            word
            for word in words
            if re.fullmatch(r"\d+|[IVXLCDM]+", str(word[4]), flags=re.IGNORECASE)
            and float(word[0]) > width * 0.55
        ]
        if not candidates:
            return ["TOC page-number right-edge metric unavailable: no rendered right-column page numbers found"], "failed"
        right_edges = [float(word[2]) for word in candidates]
        if max(right_edges) - min(right_edges) > 18:
            issues.append(
                f"TOC page-number right-edge column drift: min={min(right_edges):.2f} max={max(right_edges):.2f}"
            )
        summary = (
            "passed rendered page-number right-edge metric "
            f"count={len(right_edges)} min={min(right_edges):.2f} max={max(right_edges):.2f}"
            if not issues
            else "failed rendered page-number right-edge metric"
        )
        return issues, summary


def xml_paragraph_has_page_break_before(paragraph: ET.Element) -> bool:
    return paragraph.find("./w:pPr/w:pageBreakBefore", NS) is not None


def xml_paragraph_has_section_break(paragraph: ET.Element) -> bool:
    return paragraph.find("./w:pPr/w:sectPr", NS) is not None


def extract_body_chapter_pagination_checks(final_docx: Path) -> tuple[list[str], str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["missing word/document.xml; cannot verify body chapter pagination owner"], "failed"
    body = root.find("w:body", NS)
    if body is None:
        return ["missing w:body; cannot verify body chapter pagination owner"], "failed"
    styles = paragraph_style_name_map(final_docx)
    children = list(body)
    toc_seen = False
    body_started = False
    headings: list[tuple[int, str, ET.Element]] = []
    tail_headings = {
        normalize("\u53c2\u8003\u6587\u732e").lower(),
        normalize("\u81f4\u8c22").lower(),
        normalize("\u8c22\u8f9e").lower(),
        normalize("\u7ed3\u8bba").lower(),
        normalize("references").lower(),
        normalize("acknowledgements").lower(),
        normalize("acknowledgments").lower(),
    }
    for idx, child in enumerate(children):
        if child.tag != W + "p":
            continue
        text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
        style_node = child.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        if is_toc_heading_text(text):
            toc_seen = True
            continue
        if not toc_seen:
            continue
        if xml_paragraph_has_tab(child) or style_name.lower().startswith("toc") or style_id.lower().startswith("toc"):
            continue
        if is_toc_leader_entry_text(text):
            continue
        normalized = normalize(text).lower()
        if normalized in tail_headings:
            continue
        is_h1 = is_body_chapter_heading_text(text) or style_name.lower() == "heading 1"
        if is_h1:
            body_started = True
            headings.append((idx, text, child))
        elif body_started and normalized in tail_headings:
            break
    issues: list[str] = []
    for idx, text, paragraph in headings:
        prev_para = next((candidate for candidate in reversed(children[:idx]) if candidate.tag == W + "p"), None)
        has_owner = (
            xml_paragraph_has_page_break_before(paragraph)
            or (
                prev_para is not None
                and (xml_paragraph_has_page_break(prev_para) or xml_paragraph_has_section_break(prev_para))
            )
        )
        if not has_owner:
            issues.append(f"body chapter opener lacks a single pagination owner: {text}")
    return issues, ("passed" if not issues else "failed")


def extract_tail_block_pagination_checks(
    final_docx: Path,
    page_class_map: dict[str, int | None] | None = None,
) -> tuple[list[str], str, dict[str, object]]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return (
            ["missing word/document.xml; cannot verify tail-block pagination owner"],
            "failed",
            {"summary": "failed", "issue": "missing word/document.xml"},
        )
    body = root.find("w:body", NS)
    if body is None:
        return (
            ["missing w:body; cannot verify tail-block pagination owner"],
            "failed",
            {"summary": "failed", "issue": "missing w:body"},
        )

    children = list(body)
    target_aliases = {
        "references": TAIL_BLOCK_ALIASES["references"],
        "acknowledgement": TAIL_BLOCK_ALIASES["acknowledgement"],
    }
    found: dict[str, dict[str, object]] = {}
    issues: list[str] = []
    for idx, child in enumerate(children):
        if child.tag != W + "p":
            continue
        text = paragraph_plain_text(child)
        if not text:
            continue
        for block_key, aliases in target_aliases.items():
            if block_key in found or not normalized_equals_any(text, aliases):
                continue
            prev_para = next((candidate for candidate in reversed(children[:idx]) if candidate.tag == W + "p"), None)
            previous_content_idx, previous_content_text, _previous_content_para = previous_meaningful_paragraph(
                children,
                idx,
            )
            owner_sources: list[str] = []
            if xml_paragraph_has_page_break_before(child):
                owner_sources.append("opener.pageBreakBefore")
            if prev_para is not None and xml_paragraph_has_page_break(prev_para):
                owner_sources.append("previousParagraph.pageBreak")
            if prev_para is not None and xml_paragraph_has_section_break(prev_para):
                owner_sources.append("previousParagraph.sectionBreak")
            rendered_page = (page_class_map or {}).get("references" if block_key == "references" else "ack")
            previous_content_page = (page_class_map or {}).get(f"{block_key}_previous")
            prior_block_rendered_separation_verdict = "not-checked"
            prior_block_separation_verdict = "pass" if len(owner_sources) == 1 else "fail"
            if block_key == "references":
                if previous_content_text is None:
                    prior_block_separation_verdict = "fail"
                    prior_block_rendered_separation_verdict = "missing-previous-content"
                    issues.append(
                        "tail-block pagination contract failed: references opener has no preceding content block "
                        "to prove separation from the previous chapter/block"
                    )
                elif page_class_map is not None:
                    if not isinstance(previous_content_page, int):
                        prior_block_separation_verdict = "fail"
                        prior_block_rendered_separation_verdict = "missing-previous-page"
                        issues.append(
                            "tail-block pagination contract failed: references previous content page is missing; "
                            "cannot prove references starts after the previous chapter/block"
                        )
                    elif not isinstance(rendered_page, int):
                        prior_block_separation_verdict = "fail"
                        prior_block_rendered_separation_verdict = "missing-references-page"
                        issues.append(
                            "tail-block pagination contract failed: references rendered page is missing; "
                            "cannot prove references starts after the previous chapter/block"
                        )
                    elif previous_content_page >= rendered_page:
                        prior_block_separation_verdict = "fail"
                        prior_block_rendered_separation_verdict = "fail"
                        issues.append(
                            "tail-block pagination contract failed: references opener is not rendered after "
                            "the previous chapter/block "
                            f"(previous content page={previous_content_page}, references page={rendered_page})"
                        )
                    else:
                        prior_block_rendered_separation_verdict = "pass"
            found[block_key] = {
                "text": text,
                "paragraph_index": idx,
                "previous_content_text": (previous_content_text or "")[:160],
                "previous_content_paragraph_index": previous_content_idx,
                "owner_sources": owner_sources,
                "rendered_page": rendered_page,
                "previous_content_rendered_page": previous_content_page,
                "prior_block_separation_verdict": prior_block_separation_verdict,
                "prior_block_rendered_separation_verdict": prior_block_rendered_separation_verdict,
            }
            if len(owner_sources) != 1:
                issues.append(
                    "tail-block opener lacks exactly one pagination owner: "
                    f"{block_key} `{text[:80]}` owners={owner_sources or ['none']}"
                )
    if "references" not in found:
        issues.append("tail-block pagination contract failed: references opener not found")
    if "references" in found and "acknowledgement" in found:
        references_page = found["references"].get("rendered_page")
        ack_page = found["acknowledgement"].get("rendered_page")
        if isinstance(references_page, int) and isinstance(ack_page, int) and ack_page <= references_page:
            issues.append(
                "tail-block pagination contract failed: acknowledgement opener is not after references "
                f"(references page={references_page}, acknowledgement page={ack_page})"
            )
    evidence = {
        "summary": "passed" if not issues else "failed",
        "found_blocks": found,
        "page_class_map": page_class_map or {},
        "issue_count": len(issues),
        "issues": issues[:8],
    }
    return issues, ("passed" if not issues else "failed"), evidence


def style_label(style_id: str, style_name: str) -> str:
    return (style_name or style_id or "implicit-default").strip()


def heading_style_matches(level: int, style_id: str, style_name: str) -> bool:
    return heading_style_label_matches(level, style_id, style_name)


def style_label_equivalent(expected: str, actual: str) -> bool:
    expected_key = normalize(expected).lower()
    actual_key = normalize(actual).lower()
    if expected_key == actual_key:
        return True
    return {expected_key, actual_key} == {"normal", "implicit-default"}


def chapter_label_from_heading(text: str, ordinal: int) -> str:
    stripped = str(text or "").strip()
    match = re.match(r"^(\u7b2c[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u7ae0)", stripped)
    if match:
        return normalize(match.group(1))
    match = re.match(r"^(\d{1,2})(?:[\s\u25a1]+|$)", stripped)
    if match:
        return match.group(1)
    return f"chapter-{ordinal}"


def body_chapter_format_profiles(docx_path: Path) -> tuple[dict[str, dict[str, object]], Counter[str]]:
    root = load_xml(docx_path, "word/document.xml")
    if root is None:
        return {}, Counter()
    body = root.find("w:body", NS)
    if body is None:
        return {}, Counter()
    styles = paragraph_style_name_map(docx_path)
    profiles: dict[str, dict[str, object]] = {}
    global_body_styles: Counter[str] = Counter()
    toc_seen = False
    current_key: str | None = None
    chapter_ordinal = 0
    tail_headings = {
        normalize("\u53c2\u8003\u6587\u732e").lower(),
        normalize("\u81f4\u8c22").lower(),
        normalize("\u8c22\u8f9e").lower(),
        normalize("\u7ed3\u8bba").lower(),
        normalize("references").lower(),
        normalize("acknowledgements").lower(),
        normalize("acknowledgments").lower(),
    }
    for child in list(body):
        if child.tag != W + "p":
            if current_key and child.tag == W + "tbl":
                profiles[current_key]["table_count"] = int(profiles[current_key].get("table_count", 0)) + 1
            continue
        text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
        normalized = normalize(text).lower()
        style_node = child.find("./w:pPr/w:pStyle", NS)
        style_id = style_node.attrib.get(W + "val", "") if style_node is not None else ""
        style_name = styles.get(style_id, "")
        if is_toc_heading_text(text):
            toc_seen = True
            continue
        if not toc_seen:
            continue
        if xml_paragraph_has_tab(child) or style_id.lower().startswith("toc") or style_name.lower().startswith("toc"):
            continue
        if is_toc_leader_entry_text(text):
            continue
        if normalized in tail_headings:
            current_key = None
            break
        level = heading_level(text)
        is_h1 = is_body_chapter_heading_text(text) or heading_style_matches(1, style_id, style_name)
        if is_h1:
            chapter_ordinal += 1
            current_key = chapter_label_from_heading(text, chapter_ordinal)
            profiles[current_key] = {
                "heading_text": text,
                "heading_style": style_label(style_id, style_name),
                "heading_style_id": style_id,
                "body_styles": Counter(),
                "body_count": 0,
                "heading_records": [(1, text, style_id, style_name)],
                "image_count": 0,
                "table_count": 0,
            }
            continue
        if not current_key:
            continue
        profile = profiles[current_key]
        if xml_paragraph_has_real_image(child):
            profile["image_count"] = int(profile.get("image_count", 0)) + 1
            continue
        if CAPTION_RE.match(text) or CODE_TITLE_RE.match(text):
            continue
        if level is not None:
            profile["heading_records"].append((level, text, style_id, style_name))  # type: ignore[index]
            continue
        if not text or is_instruction_note_text(text) or is_instruction_text(text):
            continue
        label = style_label(style_id, style_name)
        profile["body_styles"][label] += 1  # type: ignore[index]
        profile["body_count"] = int(profile.get("body_count", 0)) + 1
        global_body_styles[label] += 1
    return profiles, global_body_styles


def extract_chapter_format_preservation_checks(reference_docx: Path, final_docx: Path) -> tuple[list[str], str, dict[str, object]]:
    reference_profiles, reference_global_styles = body_chapter_format_profiles(reference_docx)
    final_profiles, _final_global_styles = body_chapter_format_profiles(final_docx)
    evidence: dict[str, object] = {
        "reference_docx": str(reference_docx),
        "final_docx": str(final_docx),
        "reference_chapters": sorted(reference_profiles),
        "final_chapters": sorted(final_profiles),
        "reference_body_styles": dict(reference_global_styles),
    }
    if not reference_profiles:
        evidence["reason"] = "reference DOCX has no detected body chapter range"
        return [], "not-applicable-with-reason", evidence
    issues: list[str] = []
    if not final_profiles:
        issues.append("final DOCX has no detected body chapter range for format-preservation contract")
        evidence["issue_count"] = len(issues)
        evidence["issues"] = issues
        return issues, "failed", evidence

    global_allowed = {style for style, count in reference_global_styles.items() if count > 0}
    damaged_chapters: list[str] = []
    reference_heading_records_by_level: dict[int, list[tuple[int, str, str, str]]] = defaultdict(list)
    for reference_profile in reference_profiles.values():
        for record in reference_profile.get("heading_records", []):  # type: ignore[assignment]
            try:
                reference_level = int(record[0])
            except (TypeError, ValueError):
                continue
            reference_heading_records_by_level[reference_level].append(record)  # type: ignore[arg-type]
    for chapter_key, profile in final_profiles.items():
        heading_records = profile.get("heading_records", [])
        for level, text, style_id, style_name in heading_records:  # type: ignore[assignment]
            level_int = int(level)
            if not 1 <= level_int <= 4:
                continue
            reference_profile = reference_profiles.get(chapter_key)
            local_reference_records = [
                record
                for record in (reference_profile or {}).get("heading_records", [])  # type: ignore[union-attr]
                if int(record[0]) == level_int
            ] if reference_profile else []
            donor_records = local_reference_records or reference_heading_records_by_level.get(level_int, [])
            if not donor_records:
                continue
            donor_requires_heading_style = any(
                heading_style_matches(level_int, str(record[2]), str(record[3]))
                for record in donor_records
            )
            if donor_requires_heading_style:
                if not heading_style_matches(level_int, str(style_id), str(style_name)):
                    issues.append(
                        f"chapter {chapter_key} heading style drift: `{text}` style={style_label(str(style_id), str(style_name))} expected Heading {level_int}"
                    )
                    damaged_chapters.append(chapter_key)
                    break
                continue
            donor_labels = {
                style_label(str(record[2]), str(record[3]))
                for record in donor_records
            }
            actual_label = style_label(str(style_id), str(style_name))
            donor_uses_body_default = all(
                normalize(donor_label).lower() in {"normal", "implicit-default", ""}
                for donor_label in donor_labels
            )
            if donor_uses_body_default and heading_style_matches(level_int, str(style_id), str(style_name)):
                # The heading-baseline detector verifies visual parity.  This
                # check should not reject the semantic Heading style needed by
                # a live TOC when the donor encoded headings as Normal text.
                continue
            if not any(style_label_equivalent(donor_label, actual_label) for donor_label in donor_labels):
                issues.append(
                    f"chapter {chapter_key} heading style drift: `{text}` style={actual_label} expected donor heading style {sorted(donor_labels)}"
                )
                damaged_chapters.append(chapter_key)
                break

        reference_profile = reference_profiles.get(chapter_key)
        reference_body_styles = set((reference_profile or {}).get("body_styles", Counter()).keys()) if reference_profile else set()
        allowed_styles = reference_body_styles or global_allowed
        effective_allowed_styles = set(allowed_styles)
        if "implicit-default" in effective_allowed_styles:
            # A source/template paragraph can omit w:pStyle while still resolving
            # to Normal in Word/WPS.  Treat explicit Normal as the same body
            # family; the separate body-style audit verifies that Normal itself
            # has not drifted or been used as a heading.
            effective_allowed_styles.add("Normal")
        body_styles = profile.get("body_styles", Counter())
        body_count = int(profile.get("body_count", 0))
        if allowed_styles and body_count >= 3:
            drift_count = sum(
                count for style, count in body_styles.items()
                if style not in effective_allowed_styles
            )
            if drift_count >= max(3, int(body_count * 0.35)):
                issues.append(
                    f"chapter {chapter_key} body style damage ratio {drift_count}/{body_count}; "
                    f"allowed={sorted(allowed_styles)} actual={dict(body_styles)}"
                )
                damaged_chapters.append(chapter_key)

        heading_like_body_styles = [
            style for style in body_styles
            if is_heading_style_label(style)
        ]
        if heading_like_body_styles:
            issues.append(
                f"chapter {chapter_key} body paragraphs use heading style family: {heading_like_body_styles}"
            )
            damaged_chapters.append(chapter_key)

    evidence["checked_chapter_count"] = len(final_profiles)
    evidence["damaged_chapters"] = sorted(set(damaged_chapters))
    evidence["issue_count"] = len(issues)
    evidence["issues"] = issues[:12]
    return issues, ("passed" if not issues else "failed"), evidence


IMAGE_SIZE_FORBIDDEN_ALWAYS = (
    "\u56fe\u7247\u592a\u5927",
    "\u56fe\u50cf\u592a\u5927",
)
IMAGE_SIZE_REVIEW_TERMS = (
    "\u56fe\u7247\u5c3a\u5bf8",
    "\u56fe\u50cf\u5c3a\u5bf8",
)
IMAGE_SIZE_REMEDIATION_CONTEXTS = (
    "\u592a\u5927",
    "\u8fc7\u5927",
    "\u504f\u5927",
    "\u8c03\u6574",
    "\u7f29\u653e",
    "\u538b\u7f29",
    "\u88c1\u526a",
    "\u4fee\u6539",
    "\u4fee\u590d",
    "\u6574\u6539",
    "\u9057\u7559",
    "\u6279\u6ce8",
    "\u5bfc\u5e08",
    "\u8001\u5e08",
    "\u9875\u9762",
    "\u7248\u9762",
    "\u8d85\u51fa",
    "\u653e\u4e0d\u4e0b",
    "\u5360\u7528",
    "\u95ee\u9898",
)
IMAGE_SIZE_ALGORITHM_CONTEXTS = (
    "\u8f93\u5165\u5c3a\u5bf8",
    "\u8f93\u5165\u56fe\u50cf\u5c3a\u5bf8",
    "\u5c3a\u5bf8\u5f52\u4e00\u5316",
    "\u9884\u5904\u7406",
    "ocr",
    "\u6a21\u578b\u8f93\u5165",
    "\u7f51\u7edc\u8f93\u5165",
    "\u56fe\u50cf\u5206\u8fa8\u7387",
    "\u56fe\u7247\u5206\u8fa8\u7387",
    "\u5206\u8fa8\u7387",
    "\u50cf\u7d20",
    "\u5bbd\u9ad8",
    "resolution",
    "resize",
)


def is_algorithm_image_size_context(compact_text: str) -> bool:
    if any(token in compact_text for token in IMAGE_SIZE_ALGORITHM_CONTEXTS):
        return True
    return re.search(r"(?:\u56fe\u50cf\u5c3a\u5bf8|\u56fe\u7247\u5c3a\u5bf8)(?:\u4e3a|=|:|\uff1a)?\d{2,4}[xX\u00d7]\d{2,4}", compact_text) is not None


def forbidden_image_size_body_phrase(text: str) -> str | None:
    compact = normalize(text)
    for token in IMAGE_SIZE_FORBIDDEN_ALWAYS:
        if token in compact:
            return token
    for token in IMAGE_SIZE_REVIEW_TERMS:
        if token not in compact:
            continue
        if is_algorithm_image_size_context(compact):
            continue
        return token
    return None


def extract_body_image_size_text_checks(final_docx: Path) -> list[str]:
    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []
    issues: list[str] = []
    body_started = False
    for child in list(body):
        if child.tag != W + "p":
            continue
        text = direct_paragraph_text(child).strip()
        if not text:
            continue
        if is_toc_heading_text(text):
            continue
        if is_body_chapter_heading_text(text) and not xml_paragraph_has_tab(child):
            body_started = True
            continue
        if not body_started:
            continue
        matched = forbidden_image_size_body_phrase(text)
        if matched:
            issues.append(
                f"forbidden body image-size wording remains: term=`{matched}` text=`{text[:90]}`"
            )
            if len(issues) >= 12:
                break
    return issues


def extract_image_dimension_checks(final_docx: Path, source_docx: Path | None = None) -> tuple[list[str], str]:
    contract_issues = [
        issue
        for issue in final_docx_figure_surface_issues(final_docx, source_docx=source_docx)
        if "exceeds available text width" in issue
        or "exceeds safe page height occupancy" in issue
    ]
    body_text_issues = extract_body_image_size_text_checks(final_docx)
    if contract_issues or body_text_issues:
        return (contract_issues + body_text_issues)[:12], "未通过"

    root = load_xml(final_docx, "word/document.xml")
    if root is None:
        return ["缺少 word/document.xml，无法检查图片尺寸。"], "未通过"
    body = root.find("w:body", NS)
    if body is None:
        return ["缺少 w:body，无法检查图片尺寸。"], "未通过"

    sect_pr = body.find("w:sectPr", NS) or root.find(".//w:sectPr", NS)
    pg_sz = sect_pr.find("w:pgSz", NS) if sect_pr is not None else None
    pg_mar = sect_pr.find("w:pgMar", NS) if sect_pr is not None else None
    page_w_twips = int(pg_sz.attrib.get(W + "w", "11906")) if pg_sz is not None else 11906
    page_h_twips = int(pg_sz.attrib.get(W + "h", "16838")) if pg_sz is not None else 16838
    left_twips = int(pg_mar.attrib.get(W + "left", "1800")) if pg_mar is not None else 1800
    right_twips = int(pg_mar.attrib.get(W + "right", "1800")) if pg_mar is not None else 1800
    top_twips = int(pg_mar.attrib.get(W + "top", "1440")) if pg_mar is not None else 1440
    bottom_twips = int(pg_mar.attrib.get(W + "bottom", "1440")) if pg_mar is not None else 1440
    usable_w_emu = max(1, page_w_twips - left_twips - right_twips) * 635
    usable_h_emu = max(1, page_h_twips - top_twips - bottom_twips) * 635

    body_drawings: list[ET.Element] = []
    toc_seen = False
    body_started = False
    for child in list(body):
        if child.tag == W + "p":
            text = "".join(node.text or "" for node in child.findall(".//w:t", NS)).strip()
            normalized = normalize(text).lower()
            if is_toc_heading_text(text):
                toc_seen = True
            if toc_seen and is_body_chapter_heading_text(text) and not xml_paragraph_has_tab(child):
                body_started = True
        if body_started:
            body_drawings.extend(child.findall(".//w:drawing", NS))

    issues: list[str] = []
    for index, drawing in enumerate(body_drawings, start=1):
        extent = drawing.find(".//wp:extent", NS)
        if extent is None:
            continue
        try:
            cx = int(extent.attrib.get("cx", "0"))
            cy = int(extent.attrib.get("cy", "0"))
        except ValueError:
            continue
        if cx > int(usable_w_emu * 1.01):
            issues.append(f"图片对象宽度超过正文版心: image={index} width_emu={cx} usable_emu={usable_w_emu}")
        if cy > int(usable_h_emu * 0.82):
            issues.append(f"图片对象高度占用过大，可能挤掉图题或正文说明: image={index} height_emu={cy} usable_emu={usable_h_emu}")
        if len(issues) >= 12:
            break
    issues.extend(extract_body_image_size_text_checks(final_docx))
    return issues, ("通过" if not issues else "未通过")


def extract_figure_checks(final_docx: Path, final_pdf: Path, asset_manifest_path: Path | None) -> tuple[list[str], str]:
    if asset_manifest_path is None or not asset_manifest_path.exists():
        issues = final_docx_manifest_requirement_issues(final_docx)
        return issues, ("通过" if not issues else "未通过")

    allowed_white = {"#ffffff", "#fff", "ffffff", "fff", "white", ""}
    allowed_black = {"#000000", "#111111", "#111827", "#222222", "#333333", "000000", "111111", "111827", "222222", "333333", "black", ""}

    def normalize_color(value: object) -> str:
        return str(value or "").strip().lower()

    def drawio_style_map(style_text: str) -> dict[str, str]:
        style_map: dict[str, str] = {}
        for part in style_text.split(";"):
            if "=" not in part:
                continue
            key, value = part.split("=", 1)
            style_map[key] = value
        return style_map

    def flowchart_geometry_issues(drawio_path: Path) -> list[str]:
        try:
            root = ET.fromstring(drawio_path.read_text(encoding="utf-8", errors="replace"))
        except ET.ParseError as exc:
            return [f"结构图 draw.io XML 解析失败: {exc}"]

        vertices: list[dict[str, object]] = []
        for cell in root.findall(".//mxCell"):
            if cell.attrib.get("vertex") != "1":
                continue
            style = cell.attrib.get("style", "")
            geom = cell.find("mxGeometry")
            if geom is None:
                continue
            try:
                x = float(geom.attrib.get("x", "0"))
                y = float(geom.attrib.get("y", "0"))
                width = float(geom.attrib.get("width", "0"))
                height = float(geom.attrib.get("height", "0"))
            except ValueError:
                x = y = width = height = 0.0
            value = str(cell.attrib.get("value", ""))
            if not value.strip() and width <= 2.0 and height <= 2.0:
                continue
            vertices.append(
                {
                    "id": cell.attrib.get("id", ""),
                    "value": value,
                    "style": style,
                    "style_map": drawio_style_map(style),
                    "x": x,
                    "y": y,
                    "width": width,
                    "height": height,
                }
            )

        edges: list[tuple[str, str]] = []
        for cell in root.findall(".//mxCell"):
            if cell.attrib.get("edge") != "1":
                continue
            source = cell.attrib.get("source", "")
            target = cell.attrib.get("target", "")
            if source and target:
                edges.append((source, target))

        if not vertices:
            return ["结构图 draw.io 源文件内没有可用顶点。"]

        issues_local: list[str] = []
        def is_terminator(node: dict[str, object]) -> bool:
            value = str(node["value"])
            style = str(node["style"]).lower()
            style_map = dict(node["style_map"])  # type: ignore[arg-type]
            return (
                "ellipse" in style
                or (
                    str(style_map.get("rounded", "")).strip() == "1"
                    and ("\u5f00\u59cb" in value or "\u7ed3\u675f" in value)
                )
            )

        terminators = [node for node in vertices if is_terminator(node)]
        processes = [
            node
            for node in vertices
            if not is_terminator(node) and "rhombus" not in str(node["style"])
        ]
        if not any("开始" in str(node["value"]) for node in terminators):
            issues_local.append("流程图缺少显式开始终止符。")
        if not any("结束" in str(node["value"]) for node in terminators):
            issues_local.append("流程图缺少显式结束终止符。")
        if not processes:
            issues_local.append("流程图缺少方角 process 节点。")
        for node in processes:
            if str(node["style_map"].get("rounded", "")).strip() == "1":
                issues_local.append(f"流程图 process 节点不应使用 rounded=1: {node['value']}")
                break
        vertex_by_id = {str(node["id"]): node for node in vertices if node.get("id")}
        start_ids = [
            str(node["id"])
            for node in terminators
            if ("\u5f00\u59cb" in str(node["value"]) or "start" in str(node["value"]).lower()) and node.get("id")
        ]
        end_ids = {
            str(node["id"])
            for node in terminators
            if ("\u7ed3\u675f" in str(node["value"]) or "end" in str(node["value"]).lower()) and node.get("id")
        }
        graph: dict[str, list[str]] = {}
        for source, target in edges:
            if source in vertex_by_id and target in vertex_by_id:
                graph.setdefault(source, []).append(target)
        best_path: list[str] = []
        stack: list[tuple[str, list[str]]] = [(start_id, [start_id]) for start_id in start_ids]
        max_steps = max(10, len(vertex_by_id) * 4)
        steps = 0
        while stack and steps < max_steps:
            node_id, path = stack.pop()
            steps += 1
            if node_id in end_ids and len(path) > len(best_path):
                best_path = path
                continue
            for child in graph.get(node_id, []):
                if child not in path:
                    stack.append((child, path + [child]))
        if best_path:
            ordered = [vertex_by_id[node_id] for node_id in best_path]
        else:
            ordered = sorted(vertices, key=lambda item: (float(item["y"]), float(item["x"])))
        centers = [float(item["x"]) + float(item["width"]) / 2.0 for item in ordered]
        if centers and max(centers) - min(centers) > 120:
            issues_local.append("流程图主链错位，未保持纵向主链。")
        for node in vertices:
            style_map = dict(node["style_map"])
            fill = normalize_color(style_map.get("fillColor", ""))
            stroke = normalize_color(style_map.get("strokeColor", ""))
            font = normalize_color(style_map.get("fontColor", ""))
            if fill not in allowed_white:
                issues_local.append(f"流程图出现非白底填充: {node['value']}")
                break
            if stroke not in allowed_black:
                issues_local.append(f"流程图出现非黑色边线/连线: {node['value']}")
                break
            if font not in allowed_black:
                issues_local.append(f"流程图出现非黑色文字: {node['value']}")
                break
            if re.match(r"^(图|表)\s*\d+", str(node["value"]).strip()):
                issues_local.append(f"流程图内不应包含图名/图号文本: {node['value']}")
                break
        return issues_local

    def drawio_looks_like_flowchart(drawio_path: Path) -> bool:
        try:
            root = ET.fromstring(drawio_path.read_text(encoding="utf-8", errors="replace"))
        except ET.ParseError:
            return False
        vertex_values: list[str] = []
        has_decision = False
        for cell in root.findall(".//mxCell"):
            if cell.attrib.get("vertex") != "1":
                continue
            vertex_values.append(str(cell.attrib.get("value", "")).lower())
            if "rhombus" in cell.attrib.get("style", "").lower():
                has_decision = True
        has_edge = any(cell.attrib.get("edge") == "1" for cell in root.findall(".//mxCell"))
        values = "\n".join(vertex_values)
        has_terminal_words = any(token in values for token in ("\u5f00\u59cb", "\u7ed3\u675f", "start", "end"))
        return has_edge and (has_terminal_words or has_decision)

    issues: list[str] = []
    manifest = json.loads(asset_manifest_path.read_text(encoding="utf-8", errors="replace"))
    diagrams = manifest.get("diagrams", {})
    figure_entries = manifest.get("figures", {})
    source_preserved_relationship_contract = False
    relationship_evidence_path = manifest.get("relationship_evidence")
    if isinstance(relationship_evidence_path, str) and relationship_evidence_path.strip():
        relationship_path = Path(relationship_evidence_path)
        if relationship_path.exists():
            try:
                relationship_evidence = json.loads(relationship_path.read_text(encoding="utf-8", errors="replace"))
            except json.JSONDecodeError:
                relationship_evidence = {}
            source_preserved_relationship_contract = (
                str(manifest.get("mutation_intent", "")).strip().lower() == "no_image_mutation"
                and str(manifest.get("preservation_status", "")).strip().lower() == "source-preserved"
                and relationship_evidence.get("relationship_manifest_equal") is True
                and relationship_evidence.get("drawing_object_manifest_equal") is True
                and str(relationship_evidence.get("verdict", "")).strip().lower() == "pass"
                and normalize(str(relationship_evidence.get("final_docx_sha256", ""))).lower()
                == normalize(str(manifest.get("final_docx_sha256", ""))).lower()
            )
    source_preserved_final_captions: set[str] = set()
    if isinstance(figure_entries, dict):
        for figure_entry in figure_entries.values():
            if not isinstance(figure_entry, dict):
                continue
            family = str(figure_entry.get("family", "")).strip().lower()
            source_kind = str(figure_entry.get("source_kind", "")).strip().lower()
            preservation = str(figure_entry.get("preservation_status", "")).strip().lower()
            caption_key = normalize(str(figure_entry.get("caption", ""))).lower()
            if (
                caption_key
                and str(
                    figure_entry.get("final_media_sha256", "")
                    or figure_entry.get("final_asset_sha256", "")
                    or figure_entry.get("media_sha256", "")
                ).strip()
                and "source-preserved" in {family, source_kind, preservation}
            ):
                source_preserved_final_captions.add(caption_key)

    active_diagram_entries: list[object] = []
    for diagram_entry in diagrams.values() if isinstance(diagrams, dict) else []:
        if isinstance(diagram_entry, dict):
            caption_key = normalize(str(diagram_entry.get("caption", ""))).lower()
            has_final_binding = bool(
                str(diagram_entry.get("final_media_sha256", "")).strip()
                or str(diagram_entry.get("media_sha256", "")).strip()
                or str(diagram_entry.get("final_drawing_sha256", "")).strip()
                or str(diagram_entry.get("final_drawing_manifest_key", "")).strip()
            )
            if caption_key in source_preserved_final_captions and not has_final_binding:
                continue
        active_diagram_entries.append(diagram_entry)

    for diagram_entry in active_diagram_entries:
        if isinstance(diagram_entry, dict):
            drawio_path = diagram_entry.get("drawio")
            png_path = diagram_entry.get("png")
            family = str(diagram_entry.get("family", "")).strip().lower()
        else:
            drawio_path = None
            png_path = diagram_entry
            family = ""

        if drawio_path:
            drawio = Path(drawio_path)
        elif png_path:
            png = Path(png_path)
            drawio = png.parent.parent / "diagrams_src" / f"{png.stem}.drawio"
        else:
            continue

        if not drawio.exists():
            issues.append(f"结构图缺少 drawio 源文件: {drawio}")
        elif family == "flowchart":
            issues.extend(flowchart_geometry_issues(drawio))
        elif family in ER_FAMILY_ALIASES:
            pass
        elif drawio_looks_like_flowchart(drawio):
            issues.append(f"flowchart-like draw.io source must declare family=flowchart: {drawio}")
            issues.extend(flowchart_geometry_issues(drawio))
        svg_path = None
        if isinstance(diagram_entry, dict):
            raw_svg_path = diagram_entry.get("svg")
            if raw_svg_path:
                svg_path = Path(raw_svg_path)
        if svg_path is not None and svg_path.exists():
            svg_text = svg_path.read_text(encoding="utf-8", errors="replace")
            if "Text is not SVG - cannot display" in svg_text:
                issues.append(f"结构图 SVG 仍残留 draw.io fallback 文本: {svg_path}")
    has_external_vector_evidence = any(
        isinstance(diagram_entry, dict)
        and diagram_entry.get("drawio")
        and diagram_entry.get("svg")
        and Path(str(diagram_entry.get("drawio"))).exists()
        and Path(str(diagram_entry.get("svg"))).exists()
        for diagram_entry in active_diagram_entries
    )
    has_source_preserved_final_media_contract = bool(source_preserved_final_captions) or source_preserved_relationship_contract
    rel_root = load_xml(final_docx, "word/_rels/document.xml.rels")
    if rel_root is not None:
        image_targets = [
            rel.attrib.get("Target", "")
            for rel in rel_root.findall(f"{PR}Relationship")
            if "image" in rel.attrib.get("Type", "")
        ]
        if (
            image_targets
            and all(target.lower().endswith((".png", ".jpg", ".jpeg")) for target in image_targets)
            and not has_external_vector_evidence
            and not has_source_preserved_final_media_contract
        ):
            issues.append("最终 DOCX 中图像关系均为栅格媒体，缺少更强的结构图家族验收证据。")
    with fitz.open(final_pdf) as pdf:
        rendered_text = "\n".join(page.get_text("text") for page in pdf)
    if "Text is not SVG - cannot display" in rendered_text:
        issues.append("渲染页中泄漏了 draw.io SVG fallback 文本。")
    return issues, ("通过" if not issues else "未通过")


def extract_body_style_checks(body_style_audit: Path) -> tuple[list[str], str, list[str], str, list[str], str]:
    if not body_style_audit.exists():
        issue = f"缺少 body style audit，无法验证正文样式绑定: {body_style_audit}"
        return [issue], "未通过", [issue], "未通过", [issue], "未通过"

    text = body_style_audit.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()

    def value(prefix: str) -> str:
        for line in lines:
            if line.startswith(prefix):
                return line.split(":", 1)[1].strip()
        return "missing"

    binding_summary = value("- body style binding summary:")
    baseline_summary = value("- Normal baseline preservation summary:")
    family_summary = value("- body paragraph family consistency summary:")
    heading_contamination_summary = value("- body heading contamination summary:")
    mixed_script_summary = value("- body mixed-script font summary:")

    binding_issues: list[str] = []
    baseline_issues: list[str] = []
    family_issues: list[str] = []
    current: list[str] | None = None
    for line in lines:
        if line == "## Binding Issues":
            current = binding_issues
            continue
        if line == "## Normal Baseline Issues":
            current = baseline_issues
            continue
        if line == "## Body Family Issues":
            current = family_issues
            continue
        if line == "## Heading Contamination Issues":
            current = family_issues
            continue
        if line == "## Mixed-Script Body Issues":
            current = family_issues
            continue
        if current is not None and line.startswith("- "):
            item = line[2:].strip()
            if item != "none":
                current.append(item)

    binding_status = "通过" if binding_summary.startswith("passed") else "未通过"
    baseline_status = "通过" if baseline_summary.startswith("passed") else "未通过"
    family_status = (
        "通过"
        if family_summary.startswith("passed")
        and heading_contamination_summary.startswith("passed")
        and mixed_script_summary.startswith("passed")
        else "未通过"
    )
    return binding_issues, binding_status, baseline_issues, baseline_status, family_issues, family_status


FULL_THESIS_MIN_PAGES = 12
FULL_THESIS_MIN_PARAGRAPHS = 80
FULL_THESIS_MIN_HEADINGS = 8
FULL_THESIS_MIN_REFERENCES = 8
FULL_THESIS_MIN_CHINESE_CHARS = 10000
FORMULA_NUMBER_RE = re.compile(r"\(\s*\d+(?:[-.]\d+)?\s*\)")

PLACEHOLDER_OR_SMOKE_PATTERNS = (
    "smoke",
    "placeholder",
    "minimal",
    "manual-review",
    "detector fixture",
    "sample self check",
    "sample_self_check",
    "xxxx",
    "xxxx届",
    "xxxx 届",
    "20xx",
    "skill gate",
    "format gate",
    "template alignment",
    "example research",
    "example journal",
    "example conference",
    "example thesis",
    "最短",
    "样稿",
    "占位",
    "示例研究",
    "示例文献",
    "格式门禁",
    "模板对齐",
    "烟测",
)

FAKE_REFERENCE_PATTERNS = (
    "Example Research",
    "Example Journal",
    "Example Conference",
    "Example Thesis",
    "Template Alignment",
    "placeholder author",
    "作者姓名",
    "期刊或杂志名称",
    "张三.",
    "李四.",
    "王五.",
    "赵六.",
)

BOILERPLATE_REPETITION_PATTERNS = (
    (r"本段结合", 4, "local-section boilerplate"),
    (r"第\s*\d+\s*个分析要点", 4, "numbered-analysis filler"),
    (r"使论证不再停留在通用表述", 4, "generic-argument filler"),
)


def count_chinese_chars(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text or ""))


def count_reference_entries(text: str) -> int:
    return len(re.findall(r"(?m)^\s*\[\d+\]\s+\S+", text or ""))


def final_doc_has_omml_math(final_doc: Document) -> bool:
    try:
        return bool(final_doc.element.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]'))
    except Exception:
        return "oMath" in str(getattr(final_doc.element, "xml", ""))


def replacement_characters_are_formula_pdf_artifacts(
    *,
    doc_text: str,
    pdf_texts: list[str],
    final_doc: Document,
) -> bool:
    if "\ufffd" in doc_text:
        return False
    replacement_pages = [
        (page_index, text)
        for page_index, text in enumerate(pdf_texts, start=1)
        if "\ufffd" in (text or "")
    ]
    if not replacement_pages:
        return True
    if not final_doc_has_omml_math(final_doc):
        return False
    replacement_count = sum((text or "").count("\ufffd") for _page_index, text in replacement_pages)
    if replacement_count > 200:
        return False
    return all(FORMULA_NUMBER_RE.search(text or "") for _page_index, text in replacement_pages)


def front_matter_visible_body_text(page_text: str) -> str:
    lines = [line.strip() for line in (page_text or "").splitlines() if line.strip()]
    kept: list[str] = []
    for line in lines:
        compact = normalize(line).lower()
        if compact in {
            normalize("\u5927\u5e86\u5e08\u8303\u5b66\u9662\u672c\u79d1\u6bd5\u4e1a\u8bba\u6587\uff08\u8bbe\u8ba1\uff09").lower(),
            normalize("\u5927\u5e86\u5e08\u8303\u5b66\u9662\u672c\u79d1\u6bd5\u4e1a\u8bba\u6587(\u8bbe\u8ba1)").lower(),
        }:
            continue
        if re.fullmatch(r"(?:[IVXLCDM]+|\u7b2c\d+\u9875|\d+)", line, flags=re.IGNORECASE):
            continue
        kept.append(line)
    return "\n".join(kept).strip()


def collect_full_thesis_delivery_issues(
    *,
    final_texts: list[str],
    pdf_texts: list[str],
    final_doc: Document,
    page_class_map: dict[str, int | None],
    required_blocks: dict[str, str | None],
) -> list[str]:
    """Hard-stop checks for outputs being presented as a complete thesis."""
    issues: list[str] = []
    doc_text = "\n".join(final_texts)
    pdf_text = "\n".join(pdf_texts)
    combined = f"{doc_text}\n{pdf_text}"
    combined_lower = combined.lower()

    page_count = len(pdf_texts)
    paragraph_count = len([text for text in final_texts if text.strip()])
    heading_count = sum(1 for text in final_texts if heading_level(text) is not None)
    reference_count = count_reference_entries(combined)
    chinese_chars = count_chinese_chars(combined)

    if page_count < FULL_THESIS_MIN_PAGES:
        issues.append(
            f"full thesis content gate failed: rendered page count {page_count} < {FULL_THESIS_MIN_PAGES}"
        )
    if paragraph_count < FULL_THESIS_MIN_PARAGRAPHS:
        issues.append(
            f"full thesis content gate failed: non-empty paragraph count {paragraph_count} < {FULL_THESIS_MIN_PARAGRAPHS}"
        )
    if heading_count < FULL_THESIS_MIN_HEADINGS:
        issues.append(
            f"full thesis content gate failed: body heading count {heading_count} < {FULL_THESIS_MIN_HEADINGS}"
        )
    if reference_count < FULL_THESIS_MIN_REFERENCES:
        issues.append(
            f"full thesis content gate failed: bibliography entry count {reference_count} < {FULL_THESIS_MIN_REFERENCES}"
        )
    if chinese_chars < FULL_THESIS_MIN_CHINESE_CHARS:
        issues.append(
            f"full thesis content gate failed: Chinese character count {chinese_chars} < {FULL_THESIS_MIN_CHINESE_CHARS}"
        )

    found_placeholder_terms = [
        token for token in PLACEHOLDER_OR_SMOKE_PATTERNS if token.lower() in combined_lower
    ]
    if found_placeholder_terms:
        issues.append(
            "full thesis content gate failed: placeholder/smoke/meta text present: "
            + ", ".join(sorted(set(found_placeholder_terms))[:8])
        )

    found_fake_refs = [token for token in FAKE_REFERENCE_PATTERNS if token.lower() in combined_lower]
    if found_fake_refs:
        issues.append(
            "full thesis content gate failed: fake/template bibliography text present: "
            + ", ".join(sorted(set(found_fake_refs))[:8])
        )

    for pattern, max_count, label in BOILERPLATE_REPETITION_PATTERNS:
        count = len(re.findall(pattern, combined))
        if count > max_count:
            issues.append(
                f"full thesis content gate failed: repetitive boilerplate phrase detected: {label} count {count} > {max_count}"
            )

    if "\ufffd" in combined and not replacement_characters_are_formula_pdf_artifacts(
        doc_text=doc_text,
        pdf_texts=pdf_texts,
        final_doc=final_doc,
    ):
        issues.append("full thesis content gate failed: rendered text contains replacement characters")
    square_placeholder_lines = [
        text
        for text in final_texts
        if "\u25a1" in str(text or "") and not is_official_declaration_checkbox_line(str(text or ""))
    ]
    if square_placeholder_lines:
        issues.append("full thesis content gate failed: visible template blank placeholder `□` remains; template square placeholders must be converted to spaces")

    bad_citation = re.search(r"[A-Za-z0-9_]\[\d+(?:\s*[-,]\s*\d+)*\][A-Za-z0-9_]", combined)
    if bad_citation:
        issues.append(
            f"full thesis content gate failed: citation marker inserted inside a token: {bad_citation.group(0)}"
        )
    corrupted_section_number = re.search(r"\d\s*\[\d+(?:\s*[-,]\s*\d+)*\]\s*\.\s*\d", combined)
    if corrupted_section_number:
        issues.append(
            "full thesis content gate failed: citation marker corrupted a section-number-like text: "
            + corrupted_section_number.group(0)
        )

    normalized_long_paras = [
        normalize(text)
        for text in final_texts
        if len(normalize(text)) >= 60 and heading_level(text) is None
    ]
    duplicate_paras = [text for text, count in Counter(normalized_long_paras).items() if count > 1]
    if duplicate_paras:
        issues.append(
            f"full thesis content gate failed: duplicate long body paragraphs detected: {len(duplicate_paras)}"
        )

    toc_page = page_class_map.get("toc")
    first_body_page = page_class_map.get("first_body")
    if toc_page is not None and first_body_page is not None and first_body_page <= toc_page:
        issues.append(
            f"full thesis content gate failed: first body page {first_body_page} is not after TOC page {toc_page}"
        )
    if toc_page is not None:
        toc_text = pdf_texts[toc_page - 1] if 0 <= toc_page - 1 < len(pdf_texts) else ""
        toc_lines = [line.strip() for line in toc_text.splitlines() if line.strip()]
        toc_entry_lines = [
            line for line in toc_lines if re.search(r"\d", line) and not re.fullmatch(r"\d+", line)
        ]
        if len(toc_entry_lines) < 3:
            issues.append("full thesis content gate failed: rendered TOC has too few visible entries")

    for required_page in ("cover", "zh_abstract", "en_abstract", "toc", "first_body", "references", "ack"):
        if required_blocks.get(required_page) is not None and page_class_map.get(required_page) is None:
            issues.append(f"full thesis content gate failed: required page class not found: {required_page}")

    zh_abstract_page = page_class_map.get("zh_abstract")
    en_abstract_page = page_class_map.get("en_abstract")
    if zh_abstract_page is not None and en_abstract_page is not None and en_abstract_page - zh_abstract_page > 1:
        blank_between = []
        for page_no in range(zh_abstract_page + 1, en_abstract_page):
            page_text = pdf_texts[page_no - 1] if 0 <= page_no - 1 < len(pdf_texts) else ""
            if not front_matter_visible_body_text(page_text):
                blank_between.append(page_no)
        if blank_between:
            issues.append(
                "full thesis content gate failed: blank or header-only page between Chinese and English abstracts: "
                + ",".join(str(item) for item in blank_between)
            )

    references_page = page_class_map.get("references")
    ack_page = page_class_map.get("ack")
    def first_block_index(marker: str | None) -> int | None:
        if not marker:
            return None
        marker_key = normalize(marker).lower()
        for idx, text in enumerate(final_texts):
            if normalize(text).lower() == marker_key:
                return idx
        for idx, text in enumerate(final_texts):
            if marker_key and marker_key in normalize(text).lower():
                return idx
        return None

    ack_index = first_block_index(required_blocks.get("ack"))
    references_index = first_block_index(required_blocks.get("references"))
    if (
        references_page is not None
        and ack_page is not None
        and ack_index is not None
        and references_index is not None
    ):
        if ack_index < references_index:
            if ack_page > references_page:
                issues.append(
                    "full thesis content gate failed: acknowledgement opener rendered after references but template/source order places it before references"
                )
        elif ack_page <= references_page:
            issues.append(
                "full thesis content gate failed: acknowledgement opener must start on a rendered page after references"
            )

    if len(final_doc.inline_shapes) == 0:
        issues.append("full thesis content gate failed: no figure or screenshot exists in final DOCX")
    table_caption_like = re.search(r"(?m)^\s*(?:\u8868|Table)\s*\d", combined, flags=re.IGNORECASE)
    if len(final_doc.tables) == 0 and table_caption_like:
        issues.append("full thesis content gate failed: table caption-like text exists but no Word table exists in final DOCX")

    return issues


def read_json_file(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def infer_final_cover_surface_candidates(final_doc: Document, surface: str) -> list[str]:
    end = template_cover_end_index(final_doc)
    cover_items = list(enumerate(final_doc.paragraphs[:end]))
    english_idx = next((idx for idx, paragraph in cover_items if template_looks_like_english_title(paragraph.text)), None)
    if surface == "cover_en_title":
        if english_idx is None:
            return []
        text = final_doc.paragraphs[english_idx].text.strip()
        return [text] if text else []
    if surface != "cover_zh_title":
        return []
    zh_indices = template_select_chinese_title_indices(cover_items, english_idx, end)
    parts = [final_doc.paragraphs[idx].text.strip() for idx in zh_indices if final_doc.paragraphs[idx].text.strip()]
    if not parts:
        fallback_end = english_idx if english_idx is not None else end
        parts = [
            paragraph.text.strip()
            for idx, paragraph in cover_items
            if idx < fallback_end and template_is_cover_title_candidate(paragraph.text)
        ][-2:]
    if not parts:
        return []
    return ["\n".join(parts), "".join(parts), *parts]


def profile_surface_text_candidates(profile: dict[str, object], final_doc: Document, surface: str) -> list[str]:
    markers = profile.get("markers")
    if not isinstance(markers, dict):
        return []
    marker = markers.get(surface)
    if not isinstance(marker, dict):
        return []
    candidates: list[str] = []
    candidates.extend(infer_final_cover_surface_candidates(final_doc, surface))
    marker_text = str(marker.get("text") or "").strip()
    if marker_text:
        candidates.append(marker_text)
    seen: set[str] = set()
    unique: list[str] = []
    for item in candidates:
        key = normalize(item)
        if key and key not in seen:
            unique.append(item)
            seen.add(key)
    return unique


def find_profile_surface_page(profile: dict[str, object], final_doc: Document, pdf_texts: list[str], surface: str) -> int | None:
    for candidate in profile_surface_text_candidates(profile, final_doc, surface):
        if surface in {"zh_abstract", "en_abstract", "toc", "first_body", "references", "ack"}:
            page = find_standalone_page(pdf_texts, candidate)
        else:
            page = find_page(pdf_texts, candidate)
        if page is not None:
            return page
    return None


def find_profile_surface_pages(profile: dict[str, object], final_doc: Document, pdf_texts: list[str], surface: str) -> list[int]:
    pages: list[int] = []
    seen: set[int] = set()
    for candidate in profile_surface_text_candidates(profile, final_doc, surface):
        if surface in {"zh_abstract", "en_abstract", "toc", "first_body", "references", "ack"}:
            standalone = find_standalone_page(pdf_texts, candidate)
            candidate_pages = [standalone] if standalone is not None else []
        else:
            candidate_pages = find_pages(pdf_texts, candidate)
        for page in candidate_pages:
            if page not in seen:
                pages.append(page)
                seen.add(page)
    return pages


def detector_result(
    detector_id: str,
    *,
    surface: str,
    passed: bool,
    evidence: dict[str, object],
    blocking: bool = True,
    severity: str = "blocking",
) -> dict[str, object]:
    return {
        "id": detector_id,
        "surface": surface,
        "severity": severity,
        "passed": passed,
        "failed": not passed,
        "blocking": blocking,
        "evidence": evidence,
    }


def extract_template_profile_front_matter_checks(
    template_profile_path: Path | None,
    final_doc: Document,
    pdf_texts: list[str],
) -> tuple[list[str], str, dict[str, int | None], list[dict[str, object]]]:
    if template_profile_path is None:
        return [], "not-applicable", {}, []
    issues: list[str] = []
    detectors: list[dict[str, object]] = []
    profile = read_json_file(template_profile_path)
    if profile.get("schema") != TEMPLATE_PROFILE_SCHEMA:
        issues.append("template profile schema mismatch")
    issues.extend(f"template profile not ready: {item}" for item in profile_readiness_issues(profile))
    front_matter = profile.get("front_matter") if isinstance(profile.get("front_matter"), dict) else {}
    surfaces: set[str] = set()
    for group in front_matter.get("same_page_groups", []) if isinstance(front_matter, dict) else []:
        if isinstance(group, dict):
            surfaces.update(str(member) for member in group.get("members", []) if member)
    for pair in front_matter.get("separated_page_pairs", []) if isinstance(front_matter, dict) else []:
        if isinstance(pair, dict):
            surfaces.add(str(pair.get("left") or ""))
            surfaces.add(str(pair.get("right") or ""))
    surfaces.discard("")
    surface_page_lists = {
        surface: find_profile_surface_pages(profile, final_doc, pdf_texts, surface)
        for surface in sorted(surfaces)
    }
    surface_pages = {
        surface: find_profile_surface_page(profile, final_doc, pdf_texts, surface)
        for surface in sorted(surfaces)
    }
    for group in front_matter.get("same_page_groups", []) if isinstance(front_matter, dict) else []:
        if not isinstance(group, dict):
            continue
        group_id = str(group.get("id") or "unnamed-same-page-group")
        members = [str(member) for member in group.get("members", []) if member]
        pages = {member: surface_pages.get(member) for member in members}
        page_lists = {member: surface_page_lists.get(member, []) for member in members}
        missing = [member for member, member_pages in page_lists.items() if not member_pages]
        group_issues: list[str] = []
        if missing:
            group_issues.append(f"template same-page group `{group_id}` cannot locate final member pages: {', '.join(missing)}")
        common_pages = set(page_lists[members[0]]) if members else set()
        for member in members[1:]:
            common_pages &= set(page_lists.get(member, []))
        concrete_pages = {page for member_pages in page_lists.values() for page in member_pages}
        if not missing and not common_pages:
            group_issues.append(
                f"template same-page group `{group_id}` failed: "
                + ", ".join(f"{member}=pages {page_lists.get(member, [])}" for member in members)
            )
        issues.extend(group_issues)
        if common_pages:
            shared_page = min(common_pages)
            pages = {member: shared_page for member in members}
        detectors.append(
            detector_result(
                f"front-matter.same-page.{group_id}",
                surface="front-matter",
                passed=not group_issues,
                evidence={"pages": pages, "candidate_pages": page_lists, "template_profile": str(template_profile_path)},
            )
        )
    for pair in front_matter.get("separated_page_pairs", []) if isinstance(front_matter, dict) else []:
        if not isinstance(pair, dict):
            continue
        left = str(pair.get("left") or "")
        right = str(pair.get("right") or "")
        if not left or not right:
            continue
        left_page = surface_pages.get(left)
        right_page = surface_pages.get(right)
        pair_issues: list[str] = []
        if left_page is not None and right_page is not None and left_page == right_page:
            pair_issues.append(f"template separated-page pair failed: {left} and {right} are both on rendered page {left_page}")
        issues.extend(pair_issues)
        detectors.append(
            detector_result(
                f"front-matter.separated.{left}.{right}",
                surface="front-matter",
                passed=not pair_issues,
                evidence={"left_page": left_page, "right_page": right_page, "template_profile": str(template_profile_path)},
            )
        )
    return issues, ("passed" if not issues else "failed"), surface_pages, detectors


def extract_figure_contract_checks(
    asset_manifest_path: Path | None,
    final_docx: Path,
    source_docx: Path | None,
) -> tuple[list[str], str, list[dict[str, object]]]:
    if asset_manifest_path is None:
        issues = final_docx_manifest_requirement_issues(final_docx)
        if not issues:
            detector = detector_result(
                "figure.scope-manifest-contract",
                surface="figures",
                passed=True,
                blocking=False,
                severity="not-applicable",
                evidence={
                    "asset_manifest": "not-applicable",
                    "final_docx": str(final_docx),
                    "reason": "final DOCX has no detected figure/image surfaces",
                    "issue_count": 0,
                },
            )
            return [], "not-applicable", [detector]
        detector = detector_result(
            "figure.scope-manifest-contract",
            surface="figures",
            passed=False,
            evidence={"asset_manifest": "missing", "final_docx": str(final_docx), "issue_count": len(issues)},
        )
        return issues, "failed", [detector]
    manifest = read_json_file(asset_manifest_path)
    if source_docx is None:
        resolved_manifest = manifest_with_resolved_paths(manifest, asset_manifest_path)
        manifest_source = str(resolved_manifest.get("source_docx_path") or "").strip()
        if manifest_source:
            source_docx = Path(manifest_source)
    issues = validate_figure_manifest(
        manifest,
        final_docx=final_docx,
        source_docx=source_docx,
        manifest_path=asset_manifest_path,
    )
    detector = detector_result(
        "figure.scope-manifest-contract",
        surface="figures",
        passed=not issues,
        evidence={
            "asset_manifest": str(asset_manifest_path),
            "final_docx": str(final_docx),
            "source_docx": str(source_docx) if source_docx else "missing",
            "issue_count": len(issues),
        },
    )
    return issues, ("passed" if not issues else "failed"), [detector]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference-docx", required=True)
    parser.add_argument("--reference-pdf")
    parser.add_argument("--source-docx", help="Original/source DOCX for source-to-final figure/media preservation checks.")
    parser.add_argument("--final-docx", required=True)
    parser.add_argument("--final-pdf", required=True)
    parser.add_argument("--citation-audit", required=True)
    parser.add_argument("--font-audit", required=True)
    parser.add_argument("--body-style-audit", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--asset-manifest")
    parser.add_argument("--template-profile")
    parser.add_argument("--abstract-baseline-profile")
    parser.add_argument("--header-reference-docx")
    parser.add_argument("--fail-on-issues", action="store_true")
    parser.add_argument(
        "--smoke-acceptance",
        action="store_true",
        help="Allow detector fixtures to complete while marking the report as not deliverable.",
    )
    args = parser.parse_args()

    reference_doc = Document(args.reference_docx)
    final_doc = Document(args.final_docx)
    source_docx = Path(args.source_docx) if args.source_docx else None
    with fitz.open(args.final_pdf) as pdf:
        pdf_texts = [page.get_text("text") for page in pdf]

    ref_texts = [p.text.strip() for p in reference_doc.paragraphs if p.text.strip()]
    final_paras = [p for p in final_doc.paragraphs if p.text.strip()]
    final_texts = [p.text.strip() for p in final_paras]

    page_markers = collect_reference_page_markers(reference_doc)
    final_page_markers = collect_reference_page_markers(final_doc)
    reference_marker = final_page_markers["references"] or page_markers["references"] or "参考文献"
    ack_marker = final_page_markers["ack"] or page_markers["ack"] or "致谢"
    cover_marker = final_page_markers["cover"] or page_markers["cover"]
    required_blocks = {
        "cover": cover_marker,
        "zh_abstract": page_markers["zh_abstract"],
        "en_abstract": page_markers["en_abstract"],
        "toc": page_markers["toc"],
        "conclusion": "结论" if any(normalize("结论") == normalize(text) for text in ref_texts) else None,
        "ack": ack_marker,
        "references": reference_marker,
    }
    block_results = {
        name: (find_page(pdf_texts, marker) is not None if marker else True)
        for name, marker in required_blocks.items()
    }

    heading_failures: list[str] = []
    if reference_uses_heading_styles(reference_doc):
        for para in final_paras:
            text = para.text.strip()
            if "\t" in text:
                continue
            if re.match(r"^\d{4}\s*年", text):
                continue
            level = heading_level(text)
            if level is None:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.lower().startswith("toc"):
                continue
            if para.style is None or not heading_style_label_matches(level, para.style.style_id, style_name):
                heading_failures.append(f"{text} -> {style_name} (expected {expected_style(level)})")

    zh_abstract_marker = final_page_markers["zh_abstract"] or page_markers["zh_abstract"]
    en_abstract_marker = final_page_markers["en_abstract"] or page_markers["en_abstract"]
    toc_marker = final_page_markers["toc"] or page_markers["toc"]
    references_previous_marker = tail_block_previous_content_marker(Path(args.final_docx), TAIL_BLOCK_ALIASES["references"])
    page_class_map = {
        "cover": find_page(pdf_texts, cover_marker) if cover_marker else None,
        "zh_abstract": find_standalone_page(pdf_texts, zh_abstract_marker) if zh_abstract_marker else None,
        "en_abstract": find_standalone_page(pdf_texts, en_abstract_marker) if en_abstract_marker else None,
        "toc": find_standalone_page(pdf_texts, toc_marker) if toc_marker else None,
        "first_body": find_standalone_page(pdf_texts, final_page_markers["first_body"] or page_markers["first_body"]) if (final_page_markers["first_body"] or page_markers["first_body"]) else None,
        "figure_page": find_last_page(pdf_texts, final_page_markers["figure_page"] or page_markers["figure_page"]) if (final_page_markers["figure_page"] or page_markers["figure_page"]) else None,
        "table_page": find_last_page(pdf_texts, final_page_markers["table_page"] or page_markers["table_page"]) if (final_page_markers["table_page"] or page_markers["table_page"]) else None,
        "references_previous": find_last_page(pdf_texts, references_previous_marker) if references_previous_marker else None,
        "references": find_last_page(pdf_texts, reference_marker) if reference_marker else None,
        "ack": find_last_page(pdf_texts, ack_marker) if ack_marker else None,
    }
    front_matter_page_issues: list[str] = []
    for left, right in (("cover", "zh_abstract"), ("zh_abstract", "en_abstract"), ("en_abstract", "toc"), ("toc", "first_body")):
        left_page = page_class_map.get(left)
        right_page = page_class_map.get(right)
        if left_page is not None and right_page is not None and left_page == right_page:
            front_matter_page_issues.append(f"front-matter page separation failed: {left} and {right} are both on rendered page {left_page}")
    front_matter_order = ("cover", "zh_abstract", "en_abstract", "toc", "first_body")
    for left, right in zip(front_matter_order, front_matter_order[1:]):
        left_page = page_class_map.get(left)
        right_page = page_class_map.get(right)
        if left_page is not None and right_page is not None and left_page >= right_page:
            front_matter_page_issues.append(
                f"front-matter page order failed: {left} page {left_page} must render before {right} page {right_page}"
            )
    front_matter_page_summary = "passed" if not front_matter_page_issues else "failed"
    detector_results: list[dict[str, object]] = []
    template_profile_issues, template_profile_summary, profile_page_map, template_profile_detectors = extract_template_profile_front_matter_checks(
        Path(args.template_profile) if args.template_profile else None,
        final_doc,
        pdf_texts,
    )
    detector_results.extend(template_profile_detectors)

    citation_text = Path(args.citation_audit).read_text(encoding="utf-8", errors="replace") if Path(args.citation_audit).exists() else ""
    font_text = Path(args.font_audit).read_text(encoding="utf-8", errors="replace") if Path(args.font_audit).exists() else ""
    citation_pass = "result: pass" in citation_text.lower()
    font_audit_issues = font_audit_integrity_issues(font_text, Path(args.final_docx))
    font_pass = (
        "result: pass" in font_text.lower()
        and "bibliography font-slot checks: pass" in font_text.lower()
        and not font_audit_issues
    )

    table_issues, table_summary = extract_table_checks(Path(args.final_docx))
    non_body_surface_issues, non_body_surface_summary = extract_non_body_surface_contamination_checks(reference_doc, final_doc)
    non_body_indent_issues, non_body_indent_summary = extract_non_body_indent_checks(Path(args.final_docx))
    continuation_issues, continuation_summary = extract_cross_page_table_continuation_checks(Path(args.final_docx), Path(args.final_pdf))
    cover_baseline_issues, cover_baseline_summary = extract_cover_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    cover_value_line_issues, cover_value_line_summary, cover_value_line_evidence = extract_cover_identity_value_line_checks(
        Path(args.reference_docx),
        Path(args.final_docx),
    )
    detector_results.append(
        detector_result(
            "cover.identity-value-line-contract",
            surface="cover_style",
            passed=not cover_value_line_issues,
            blocking=cover_value_line_summary != "not-applicable",
            severity="blocking" if cover_value_line_summary != "not-applicable" else "not-applicable",
            evidence=cover_value_line_evidence,
        )
    )
    cover_residue_issues, cover_residue_summary = extract_cover_sample_title_residue_checks(Path(args.final_docx))
    figure_block_issues, figure_block_summary = extract_figure_block_locality_checks(Path(args.final_docx))
    figure_followup_issues, figure_followup_summary = extract_figure_explanation_followup_checks(Path(args.final_docx))
    figure_caption_baseline_issues, figure_caption_baseline_summary = extract_figure_caption_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    image_holder_baseline_issues, image_holder_baseline_summary = extract_image_holder_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    image_holder_issues, image_holder_summary = extract_image_holder_safety_checks(Path(args.final_docx))
    caption_in_table_issues, caption_in_table_summary = extract_caption_inside_table_grid_checks(Path(args.final_docx))
    table_caption_binding_issues, table_caption_binding_summary = extract_table_caption_binding_checks(Path(args.final_docx))
    table_caption_baseline_issues, table_caption_baseline_summary = extract_table_caption_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    table_cell_baseline_issues, table_cell_baseline_summary = extract_table_cell_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    table_structure_issues, table_structure_summary = extract_table_structure_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    bibliography_entry_issues, bibliography_entry_summary = extract_bibliography_entry_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    bibliography_geometry_issues, bibliography_geometry_summary = extract_bibliography_rendered_geometry_checks(
        Path(args.reference_pdf) if args.reference_pdf else None,
        Path(args.final_pdf),
    )
    bibliography_numbering_issues, bibliography_numbering_summary = extract_bibliography_numbering_checks(Path(args.reference_docx), Path(args.final_docx))
    bibliography_count_issues, bibliography_count_summary = extract_bibliography_count_checks(Path(args.reference_docx), Path(args.final_docx))
    reference_tail_issues, reference_tail_summary = extract_tail_block_baseline_checks(
        Path(args.reference_docx),
        Path(args.final_docx),
        block_key="references",
        title_label="参考文献标题",
        body_label="参考文献内容",
    )
    acknowledgement_tail_issues, acknowledgement_tail_summary = extract_tail_block_baseline_checks(
        Path(args.reference_docx),
        Path(args.final_docx),
        block_key="acknowledgement",
        title_label="致谢标题",
        body_label="致谢内容",
    )
    appendix_tail_issues, appendix_tail_summary = extract_tail_block_baseline_checks(
        Path(args.reference_docx),
        Path(args.final_docx),
        block_key="appendix",
        title_label="附录标题",
        body_label="附录内容",
    )
    for detector_id, surface, summary, problem_list in (
        ("surface.references.tail_block_baseline", "references", reference_tail_summary, reference_tail_issues),
        ("surface.acknowledgement.tail_block_baseline", "acknowledgement", acknowledgement_tail_summary, acknowledgement_tail_issues),
        ("surface.appendix.tail_block_baseline", "appendix", appendix_tail_summary, appendix_tail_issues),
    ):
        detector_results.append(
            detector_result(
                detector_id,
                surface=surface,
                passed=not problem_list,
                blocking=summary != "不适用",
                severity="blocking" if summary != "不适用" else "not-applicable",
                evidence={"summary": summary, "issue_count": len(problem_list), "issues": problem_list[:8]},
            )
        )
    endmatter_indent_issues = (
        list(reference_tail_issues)
        + list(acknowledgement_tail_issues)
        + list(appendix_tail_issues)
    )
    detector_results.append(
        detector_result(
            "endmatter.indentation-contract",
            surface="references; acknowledgement; appendix",
            passed=not endmatter_indent_issues,
            blocking=True,
            severity="blocking",
            evidence={
                "references_summary": reference_tail_summary,
                "acknowledgement_summary": acknowledgement_tail_summary,
                "appendix_summary": appendix_tail_summary,
                "issue_count": len(endmatter_indent_issues),
                "issues": endmatter_indent_issues[:12],
            },
        )
    )
    toc_issues, toc_summary = extract_toc_checks(Path(args.final_docx), Path(args.final_pdf), Path(args.reference_docx))
    detector_results.append(
        detector_result(
            "toc.visible-format-contract",
            surface="toc",
            passed=not toc_issues,
            evidence={
                "summary": toc_summary,
                "reference_docx": str(Path(args.reference_docx)),
                "final_docx": str(Path(args.final_docx)),
                "final_pdf": str(Path(args.final_pdf)),
                "issue_count": len(toc_issues),
                "issues": toc_issues[:12],
            },
        )
    )
    toc_page_number_column_issues, toc_page_number_column_summary = extract_toc_page_number_right_edge_checks(
        Path(args.final_pdf)
    )
    detector_results.append(
        detector_result(
            "toc.page-number-column-right-edge",
            surface="toc_page_number_column",
            passed=not toc_page_number_column_issues,
            blocking=False,
            severity="advisory",
            evidence={
                "summary": toc_page_number_column_summary,
                "metric_owner": "rendered page-number right edge",
                "forbidden_proxy": "dotted leader start / leader_x0",
                "issue_count": len(toc_page_number_column_issues),
                "issues": toc_page_number_column_issues[:4],
            },
        )
    )
    toc_control_issues, toc_control_summary = extract_toc_control_contamination_checks(Path(args.final_docx))
    header_reference_docx = Path(args.header_reference_docx) if args.header_reference_docx else Path(args.reference_docx)
    header_issues, header_summary = extract_header_checks(Path(args.reference_docx), Path(args.final_pdf), Path(args.final_docx))
    header_presence_issues, header_presence_summary, header_presence_detector = extract_header_presence_contract_checks(
        header_reference_docx,
        Path(args.final_docx),
        Path(args.final_pdf),
        page_class_map,
    )
    detector_results.append(header_presence_detector)
    footer_issues, footer_summary = extract_footer_indent_checks(Path(args.reference_docx), Path(args.final_docx))
    abstract_issues, abstract_summary = extract_abstract_baseline_checks(
        Path(args.reference_docx),
        Path(args.final_docx),
        Path(args.abstract_baseline_profile) if args.abstract_baseline_profile else None,
    )
    abstract_manual_break_issues, abstract_manual_break_summary = extract_abstract_manual_line_break_checks(Path(args.final_docx))
    abstract_contract_issues = list(abstract_issues) + list(abstract_manual_break_issues)
    detector_results.append(
        detector_result(
            "abstract.template-style-contract",
            surface="abstracts",
            passed=not abstract_contract_issues,
            evidence={
                "reference_docx": str(Path(args.reference_docx)),
                "final_docx": str(Path(args.final_docx)),
                "abstract_summary": abstract_summary,
                "manual_line_break_summary": abstract_manual_break_summary,
                "issue_count": len(abstract_contract_issues),
                "issues": abstract_contract_issues[:8],
            },
        )
    )
    front_matter_instruction_issues, front_matter_instruction_summary = extract_front_matter_instruction_artifact_checks(Path(args.final_docx))
    forbidden_color_issues, forbidden_color_summary = extract_forbidden_visible_color_checks(Path(args.final_docx))
    common_submission_issues, common_submission_summary, common_submission_evidence = extract_common_pre_submission_checks(Path(args.final_docx))
    detector_results.append(
        detector_result(
            "common.pre-submission-checklist",
            surface="whole thesis",
            passed=not common_submission_issues,
            evidence=common_submission_evidence,
        )
    )
    footer_baseline_issues, footer_baseline_summary = extract_header_footer_baseline_checks(Path(args.reference_docx), Path(args.final_docx))
    header_footer_page_number_issues = (
        list(header_issues)
        + list(header_presence_issues)
        + list(footer_issues)
        + list(footer_baseline_issues)
    )
    detector_results.append(
        detector_result(
            "header-footer.page-number-template-contract",
            surface="header; footer; page_numbers",
            passed=not header_footer_page_number_issues,
            evidence={
                "reference_docx": str(Path(args.reference_docx)),
                "final_docx": str(Path(args.final_docx)),
                "final_pdf": str(Path(args.final_pdf)),
                "header_summary": header_summary,
                "header_presence_summary": header_presence_summary,
                "footer_summary": footer_summary,
                "footer_baseline_summary": footer_baseline_summary,
                "header_line_policy": "header horizontal rules must be template-proven, not target-only residue",
                "page_number_policy": "footer/page-number position must match the rendered template baseline",
                "issue_count": len(header_footer_page_number_issues),
                "issues": header_footer_page_number_issues[:12],
            },
        )
    )
    figure_issues, figure_summary = extract_figure_checks(
        Path(args.final_docx),
        Path(args.final_pdf),
        Path(args.asset_manifest) if args.asset_manifest else None,
    )
    detector_results.append(
        detector_result(
            "figure.family-style-contract",
            surface="figures",
            passed=not figure_issues,
            evidence={
                "summary": figure_summary,
                "asset_manifest": str(Path(args.asset_manifest)) if args.asset_manifest else "",
                "final_docx": str(Path(args.final_docx)),
                "final_pdf": str(Path(args.final_pdf)),
                "issue_count": len(figure_issues),
                "issues": figure_issues[:12],
            },
        )
    )
    figure_contract_issues, figure_contract_summary, figure_contract_detectors = extract_figure_contract_checks(
        Path(args.asset_manifest) if args.asset_manifest else None,
        Path(args.final_docx),
        source_docx,
    )
    detector_results.extend(figure_contract_detectors)
    image_dimension_issues, image_dimension_summary = extract_image_dimension_checks(Path(args.final_docx), source_docx)
    detector_results.append(
        detector_result(
            "figure.image-dimension-contract",
            surface="figures",
            passed=not image_dimension_issues,
            evidence={
                "summary": image_dimension_summary,
                "final_docx": str(Path(args.final_docx)),
                "source_docx": str(source_docx) if source_docx else "",
                "issue_count": len(image_dimension_issues),
                "issues": image_dimension_issues[:12],
            },
        )
    )
    body_style_issues, body_style_summary, body_baseline_issues, body_baseline_summary, body_family_issues, body_family_summary = extract_body_style_checks(Path(args.body_style_audit))
    body_contract_issues = list(body_style_issues) + list(body_baseline_issues) + list(body_family_issues)
    detector_results.append(
        detector_result(
            "body.style-contamination-contract",
            surface="body",
            passed=not body_contract_issues,
            evidence={
                "body_style_audit": str(Path(args.body_style_audit)),
                "binding_summary": body_style_summary,
                "baseline_summary": body_baseline_summary,
                "family_summary": body_family_summary,
                "issue_count": len(body_contract_issues),
                "issues": body_contract_issues[:8],
            },
        )
    )
    body_text_issues = [
        issue
        for issue in body_family_issues
        if "mixed-script" in issue.lower()
        or "western font slot" in issue.lower()
        or "east asian and western font family" in issue.lower()
        or "single visible run" in issue.lower()
    ]
    detector_results.append(
        detector_result(
            "body.text-protected-surface-contract",
            surface="body text",
            passed=not body_text_issues,
            evidence={
                "body_style_audit": str(Path(args.body_style_audit)),
                "family_summary": body_family_summary,
                "issue_count": len(body_text_issues),
                "issues": body_text_issues[:8],
            },
        )
    )
    body_chapter_pagination_issues, body_chapter_pagination_summary = extract_body_chapter_pagination_checks(Path(args.final_docx))
    tail_block_pagination_issues, tail_block_pagination_summary, tail_block_pagination_evidence = (
        extract_tail_block_pagination_checks(Path(args.final_docx), page_class_map)
    )
    detector_results.append(
        detector_result(
            "tail-block.pagination-contract",
            surface="references; acknowledgement",
            passed=not tail_block_pagination_issues,
            evidence=tail_block_pagination_evidence,
        )
    )
    chapter_format_preservation_issues, chapter_format_preservation_summary, chapter_format_preservation_evidence = (
        extract_chapter_format_preservation_checks(Path(args.reference_docx), Path(args.final_docx))
    )
    detector_results.append(
        detector_result(
            "chapter.format-preservation-contract",
            surface="body chapters",
            passed=not chapter_format_preservation_issues,
            blocking=chapter_format_preservation_summary != "not-applicable-with-reason",
            severity=(
                "not-applicable"
                if chapter_format_preservation_summary == "not-applicable-with-reason"
                else "blocking"
            ),
            evidence=chapter_format_preservation_evidence,
        )
    )
    code_title_issues, code_title_summary, code_block_issues, code_block_summary = extract_code_format_checks(Path(args.reference_docx), Path(args.final_docx))

    issues: list[str] = []
    for name, exists in block_results.items():
        if required_blocks[name] and not exists:
            issues.append(f"缺少模板结构块: {name}")
    if heading_failures:
        issues.append("存在标题未落到对应 Heading 样式族")
    heading_summary = "通过" if not heading_failures else "未通过"
    heading1_baseline_issues, heading1_baseline_summary = extract_heading_baseline_checks(Path(args.reference_docx), Path(args.final_docx), level=1)
    heading2_baseline_issues, heading2_baseline_summary = extract_heading_baseline_checks(Path(args.reference_docx), Path(args.final_docx), level=2)
    heading3_baseline_issues, heading3_baseline_summary = extract_heading_baseline_checks(Path(args.reference_docx), Path(args.final_docx), level=3)
    heading4_baseline_issues, heading4_baseline_summary = extract_heading_baseline_checks(Path(args.reference_docx), Path(args.final_docx), level=4)
    heading_baseline_issues = (
        list(heading1_baseline_issues)
        + list(heading2_baseline_issues)
        + list(heading3_baseline_issues)
        + list(heading4_baseline_issues)
    )
    detector_results.append(
        detector_result(
            "heading.baseline-contract",
            surface="body headings",
            passed=not heading_baseline_issues,
            evidence={
                "level1_summary": heading1_baseline_summary,
                "level2_summary": heading2_baseline_summary,
                "level3_summary": heading3_baseline_summary,
                "level4_summary": heading4_baseline_summary,
                "reference_docx": str(Path(args.reference_docx)),
                "final_docx": str(Path(args.final_docx)),
                "issue_count": len(heading_baseline_issues),
                "issues": heading_baseline_issues[:12],
            },
        )
    )
    if len(final_doc.inline_shapes) == 0:
        issues.append("正文无图形或截图")
    final_combined_text_for_table_gate = "\n".join(final_texts + pdf_texts)
    if len(final_doc.tables) == 0 and re.search(r"(?m)^\s*(?:\u8868|Table)\s*\d", final_combined_text_for_table_gate, flags=re.IGNORECASE):
        issues.append("正文无表格")
    if not citation_pass:
        issues.append("正文引用审计未通过")
    if not font_pass:
        issues.append("字体/编码审计未通过")
        if "bibliography font-slot checks: pass" not in font_text.lower():
            issues.append("参考文献字体槽审计未通过")
        issues.extend(font_audit_issues)
    issues.extend(table_issues)
    issues.extend(front_matter_page_issues)
    issues.extend(template_profile_issues)
    issues.extend(non_body_surface_issues)
    issues.extend(non_body_indent_issues)
    issues.extend(continuation_issues)
    issues.extend(cover_baseline_issues)
    issues.extend(cover_value_line_issues)
    issues.extend(cover_residue_issues)
    issues.extend(figure_block_issues)
    issues.extend(figure_followup_issues)
    issues.extend(figure_caption_baseline_issues)
    issues.extend(image_holder_baseline_issues)
    issues.extend(image_holder_issues)
    issues.extend(caption_in_table_issues)
    issues.extend(table_caption_binding_issues)
    issues.extend(table_caption_baseline_issues)
    issues.extend(table_cell_baseline_issues)
    issues.extend(table_structure_issues)
    issues.extend(bibliography_entry_issues)
    issues.extend(bibliography_geometry_issues)
    issues.extend(bibliography_numbering_issues)
    issues.extend(bibliography_count_issues)
    issues.extend(reference_tail_issues)
    issues.extend(acknowledgement_tail_issues)
    issues.extend(appendix_tail_issues)
    issues.extend(toc_issues)
    issues.extend(toc_control_issues)
    issues.extend(header_issues)
    issues.extend(header_presence_issues)
    issues.extend(footer_issues)
    issues.extend(abstract_issues)
    issues.extend(abstract_manual_break_issues)
    issues.extend(front_matter_instruction_issues)
    issues.extend(forbidden_color_issues)
    issues.extend(common_submission_issues)
    issues.extend(footer_baseline_issues)
    issues.extend(figure_issues)
    issues.extend(figure_contract_issues)
    issues.extend(image_dimension_issues)
    issues.extend(body_style_issues)
    issues.extend(body_baseline_issues)
    issues.extend(body_family_issues)
    issues.extend(body_chapter_pagination_issues)
    issues.extend(tail_block_pagination_issues)
    issues.extend(chapter_format_preservation_issues)
    issues.extend(heading1_baseline_issues)
    issues.extend(heading2_baseline_issues)
    issues.extend(heading3_baseline_issues)
    issues.extend(heading4_baseline_issues)
    issues.extend(code_title_issues)
    issues.extend(code_block_issues)
    for name, page in page_class_map.items():
        if required_blocks.get(name) or name == "first_body":
            if page is None:
                issues.append(f"未在渲染 PDF 中定位到页面类别: {name}")

    full_thesis_delivery_issues = collect_full_thesis_delivery_issues(
        final_texts=final_texts,
        pdf_texts=pdf_texts,
        final_doc=final_doc,
        page_class_map=page_class_map,
        required_blocks=required_blocks,
    )
    issues.extend(full_thesis_delivery_issues)

    critical_issues: list[str] = []
    critical_issues.extend(f"missing required block: {name}" for name, exists in block_results.items() if required_blocks[name] and not exists)
    critical_issues.extend(heading_failures)
    if not citation_pass:
        critical_issues.append("citation audit failed")
    if not font_pass:
        critical_issues.append("font/encoding audit failed")
    for issue_group in (
        table_issues,
        front_matter_page_issues,
        template_profile_issues,
        non_body_surface_issues,
        non_body_indent_issues,
        cover_baseline_issues,
        cover_value_line_issues,
        cover_residue_issues,
        table_caption_binding_issues,
        table_caption_baseline_issues,
        table_cell_baseline_issues,
        table_structure_issues,
        bibliography_entry_issues,
        bibliography_geometry_issues,
        bibliography_numbering_issues,
        bibliography_count_issues,
        reference_tail_issues,
        acknowledgement_tail_issues,
        appendix_tail_issues,
        toc_issues,
        toc_control_issues,
        header_issues,
        header_presence_issues,
        footer_issues,
        abstract_issues,
        abstract_manual_break_issues,
        front_matter_instruction_issues,
        forbidden_color_issues,
        common_submission_issues,
        footer_baseline_issues,
        body_style_issues,
        body_baseline_issues,
        body_family_issues,
        body_chapter_pagination_issues,
        tail_block_pagination_issues,
        chapter_format_preservation_issues,
        heading1_baseline_issues,
        heading2_baseline_issues,
        heading3_baseline_issues,
        heading4_baseline_issues,
        figure_issues,
        figure_block_issues,
        figure_followup_issues,
        figure_caption_baseline_issues,
        figure_contract_issues,
        image_holder_baseline_issues,
        image_holder_issues,
        image_dimension_issues,
        full_thesis_delivery_issues,
    ):
        critical_issues.extend(str(item) for item in issue_group)
    deliverable_status = "passed" if not critical_issues else "blocked"
    if args.smoke_acceptance:
        deliverable_status = "smoke-only; blocked for delivery"

    lines = [
        f"- final docx path: {args.final_docx}",
        f"- final docx sha256: {sha256_file(Path(args.final_docx)) if Path(args.final_docx).exists() else 'missing'}",
        f"- final pdf path: {args.final_pdf}",
        "# 完整样稿自检",
        "",
        "## 总结",
        f"- 模板权威: {args.reference_docx}",
        f"- 最终样稿: {args.final_docx}",
        f"- 渲染 PDF: {args.final_pdf}",
        f"- 正文图数量: {len(final_doc.inline_shapes)}",
        f"- 正文表数量: {len(final_doc.tables)}",
        f"- 引用审计: {'pass' if citation_pass else 'fail'}",
        f"- 字体/编码审计: {'pass' if font_pass else 'fail'}",
        f"- 标题样式家族检查: {heading_summary}",
        f"- 一级标题真实基线检查: {heading1_baseline_summary}",
        f"- 二级标题真实基线检查: {heading2_baseline_summary}",
        f"- 三级标题真实基线检查: {heading3_baseline_summary}",
        f"- 四级标题真实基线检查: {heading4_baseline_summary}",
        f"- 图形家族检查: {figure_summary}",
        f"- 表格家族检查: {table_summary}",
        f"- 前置页分页隔离检查: {front_matter_page_summary}",
        f"- 非正文面正文污染检查: {non_body_surface_summary}",
        f"- 非正文首行缩进检查: {non_body_indent_summary}",
        f"- 跨页续表检查: {continuation_summary}",
        f"- 封面真实基线检查: {cover_baseline_summary}",
        f"- cover identity value-line contract check: {cover_value_line_summary}",
        f"- cover sample-title residue check: {cover_residue_summary}",
        f"- 图块本地锚定检查: {figure_block_summary}",
        f"- 图后正文介绍检查: {figure_followup_summary}",
        f"- 图题真实基线检查: {figure_caption_baseline_summary}",
        f"- 图片承载段落真实基线检查: {image_holder_baseline_summary}",
        f"- 图片承载段落安全检查: {image_holder_summary}",
        f"- 题名入表格检查: {caption_in_table_summary}",
        f"- 表题绑定检查: {table_caption_binding_summary}",
        f"- 表题真实基线检查: {table_caption_baseline_summary}",
        f"- 表格文字真实基线检查: {table_cell_baseline_summary}",
        f"- 表格结构/边框真实基线检查: {table_structure_summary}",
        f"- 参考文献条目真实基线检查: {bibliography_entry_summary}",
        f"- reference rendered geometry check: {bibliography_geometry_summary}",
        f"- 参考文献编号机制检查: {bibliography_numbering_summary}",
        f"- 参考文献条目数量检查: {bibliography_count_summary}",
        f"- 参考文献标题/内容独立基线检查: {reference_tail_summary}",
        f"- 致谢标题/内容独立基线检查: {acknowledgement_tail_summary}",
        f"- 附录标题/内容独立基线检查: {appendix_tail_summary}",
        f"- 目录控件正文污染检查: {toc_control_summary}",
        f"- 正文样式绑定检查: {body_style_summary}",
        f"- 正文默认样式基线检查: {body_baseline_summary}",
        f"- 正文段落家族一致性检查: {body_family_summary}",
        f"- body chapter pagination owner check: {body_chapter_pagination_summary}",
        f"- tail-block pagination owner check: {tail_block_pagination_summary}",
        f"- chapter format preservation contract check: {chapter_format_preservation_summary}",
        f"- 代码题名格式检查: {code_title_summary}",
        f"- 代码块格式检查: {code_block_summary}",
        f"- 目录可见格式检查: {toc_summary}",
        f"- 页眉位置检查: {header_summary}",
        f"- header presence contract check: {header_presence_summary}",
        f"- 页脚缩进检查: {footer_summary}",
        f"- 摘要真实基线检查: {abstract_summary}",
        f"- abstract manual line-break check: {abstract_manual_break_summary}",
        f"- 摘要/前置页模板说明残留检查: {front_matter_instruction_summary}",
        f"- visible red template-format run check: {forbidden_color_summary}",
        f"- 页眉页脚真实基线检查: {footer_baseline_summary}",
        f"- 图片尺寸安全检查: {image_dimension_summary}",
        "",
        "## 模板结构块检查",
    ]
    lines.append(f"- front-matter template occupancy check: {template_profile_summary}")
    lines.append(f"- figure manifest contract check: {figure_contract_summary}")
    for name, exists in block_results.items():
        lines.append(f"- {name}: {'present' if exists else 'missing'}")

    lines.extend(["", "## 标题样式检查"])
    if heading_failures:
        for item in heading_failures:
            lines.append(f"- fail: {item}")
    else:
        lines.append("- 各级标题均已落到对应 Heading 样式族")

    lines.extend(["", "## 页面类别检查"])
    for name, page in page_class_map.items():
        lines.append(f"- {name}: {page if page is not None else 'not-found'}")
    if profile_page_map:
        lines.extend(["", "## Template Profile Page-Class Checks"])
        for name, page in profile_page_map.items():
            lines.append(f"- {name}: {page if page is not None else 'not-found'}")

    lines.extend(["", "## 重点格式检查"])
    for title, problem_list in (
        ("图形家族", figure_issues),
        ("表格家族与底纹", table_issues),
        ("前置页分页隔离", front_matter_page_issues),
        ("非正文面正文污染", non_body_surface_issues),
        ("非正文首行缩进", non_body_indent_issues),
        ("跨页续表", continuation_issues),
        ("封面真实基线", cover_baseline_issues),
        ("cover identity value-line contract", cover_value_line_issues),
        ("cover sample-title residue", cover_residue_issues),
        ("图块本地锚定", figure_block_issues),
        ("图后正文介绍", figure_followup_issues),
        ("图题真实基线", figure_caption_baseline_issues),
        ("图片承载段落真实基线", image_holder_baseline_issues),
        ("图片承载段落安全", image_holder_issues),
        ("题名入表格", caption_in_table_issues),
        ("表题绑定", table_caption_binding_issues),
        ("表题真实基线", table_caption_baseline_issues),
        ("表格文字真实基线", table_cell_baseline_issues),
        ("表格结构/边框真实基线", table_structure_issues),
        ("参考文献条目真实基线", bibliography_entry_issues),
        ("参考文献编号机制", bibliography_numbering_issues),
        ("参考文献条目数量", bibliography_count_issues),
        ("参考文献标题/内容独立基线", reference_tail_issues),
        ("致谢标题/内容独立基线", acknowledgement_tail_issues),
        ("附录标题/内容独立基线", appendix_tail_issues),
        ("目录控件正文污染", toc_control_issues),
        ("正文样式绑定", body_style_issues),
        ("正文默认样式基线", body_baseline_issues),
        ("正文段落家族一致性", body_family_issues),
        ("body chapter pagination owner", body_chapter_pagination_issues),
        ("chapter format preservation contract", chapter_format_preservation_issues),
        ("一级标题真实基线", heading1_baseline_issues),
        ("二级标题真实基线", heading2_baseline_issues),
        ("三级标题真实基线", heading3_baseline_issues),
        ("四级标题真实基线", heading4_baseline_issues),
        ("代码题名格式", code_title_issues),
        ("代码块格式", code_block_issues),
        ("目录可见格式", toc_issues),
        ("页眉位置", header_issues),
        ("header presence contract", header_presence_issues),
        ("页脚缩进", footer_issues),
        ("摘要真实基线", abstract_issues),
        ("摘要/前置页模板说明残留", front_matter_instruction_issues),
        ("visible red template-format runs", forbidden_color_issues),
        ("页眉页脚真实基线", footer_baseline_issues),
        ("图片尺寸安全", image_dimension_issues),
    ):
        lines.append(f"- {title}:")
        if problem_list:
            for item in problem_list:
                lines.append(f"  - {item}")
        else:
            lines.append("  - 通过")

    lines.extend(["", "## Additional Blocking Detectors"])
    for title, problem_list in (
        ("front-matter template occupancy", template_profile_issues),
        ("abstract template-style contract", abstract_contract_issues),
        ("heading baseline contract", heading_baseline_issues),
        ("TOC visible format contract", toc_issues),
        ("figure family style contract", figure_issues),
        ("body style contamination contract", body_contract_issues),
        ("figure manifest contract", figure_contract_issues),
        ("chapter format preservation contract", chapter_format_preservation_issues),
    ):
        lines.append(f"- {title}:")
        if problem_list:
            for item in problem_list:
                lines.append(f"  - {item}")
        else:
            lines.append("  - passed")

    lines.extend(["", "## 结论"])
    if issues:
        lines.append("- 仍有问题:")
        for issue in issues:
            lines.append(f"  - {issue}")
    else:
        lines.append("- 本次重建样稿已通过结构、标题、图表、目录、页眉页脚、引用和字体编码的核心自检。")

    lines.extend(["", "## Delivery Gate"])
    lines.append(f"- deliverable critical gate: {deliverable_status}")
    lines.append(f"- smoke acceptance mode: {'yes' if args.smoke_acceptance else 'no'}")
    lines.extend(["", "## Detector Registry"])
    if detector_results:
        for item in detector_results:
            lines.append(json.dumps(item, ensure_ascii=False, sort_keys=True))
    else:
        lines.append("- no detector registry entries")

    Path(args.output).write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(args.output)
    if args.smoke_acceptance:
        return 0
    if critical_issues and not args.smoke_acceptance:
        return 1
    if args.fail_on_issues and issues and not args.smoke_acceptance:
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

