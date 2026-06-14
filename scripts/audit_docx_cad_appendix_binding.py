#!/usr/bin/env python3
"""Audit that a thesis appendix embeds final mechanical CAD sheet renders."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "r": R_NS, "a": A_NS, "rel": REL_NS}
SCHEMATIC_SUBSTITUTE_RE = re.compile(
    "|".join(
        [
            r"\u793a\u610f\u56fe",
            r"\u793a\u610f",
            r"\u6982\u5ff5\u56fe",
            r"\u8349\u56fe",
            r"\u7b80\u56fe",
            r"\u6837\u4f8b\u56fe",
            r"\u5360\u4f4d\u56fe",
            r"schematic",
            r"concept(?:ual)?",
            r"placeholder",
            r"mock(?:up)?",
            r"sketch",
            r"sample[- ]figure",
        ]
    ),
    re.IGNORECASE,
)


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def paragraph_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()


def compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def is_reference_heading(text: str) -> bool:
    return compact(text).rstrip(":：") in {"参考文献", "References"}


def is_appendix_heading(text: str) -> bool:
    normalized = compact(text).rstrip(":：")
    return (
        normalized == "附录"
        or bool(re.fullmatch(r"附录[A-ZＡ-Ｚ]?", normalized))
        or normalized.startswith("附录A")
    )


def is_after_appendix_tail_heading(text: str) -> bool:
    return compact(text).rstrip(":：") in {"致谢", "Acknowledgements", "Acknowledgement"}


def read_docx_paragraphs_and_images(docx: Path) -> tuple[list[dict[str, object]], dict[str, dict[str, object]], str]:
    docx_sha = sha256_file(docx)
    with zipfile.ZipFile(docx) as zf:
        document_root = ET.fromstring(zf.read("word/document.xml"))
        rels_root = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
        rid_to_target = {
            rel.attrib.get("Id"): rel.attrib.get("Target", "")
            for rel in rels_root.findall(".//rel:Relationship", NS)
        }
        media_by_target: dict[str, dict[str, object]] = {}
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue
            data = zf.read(name)
            target = name.removeprefix("word/")
            media_by_target[target] = {
                "target": target,
                "sha256": sha256_bytes(data),
                "bytes": len(data),
            }
        paragraphs: list[dict[str, object]] = []
        for index, paragraph in enumerate(document_root.findall(".//w:body/w:p", NS)):
            rids = [node.attrib.get(f"{{{R_NS}}}embed") for node in paragraph.findall(".//a:blip", NS)]
            targets = [rid_to_target.get(rid, "") for rid in rids if rid]
            paragraphs.append(
                {
                    "index": index,
                    "text": paragraph_text(paragraph),
                    "image_targets": targets,
                    "image_hashes": [
                        media_by_target[target]["sha256"]
                        for target in targets
                        if target in media_by_target
                    ],
                }
            )
    return paragraphs, media_by_target, docx_sha


def find_appendix_bounds(paragraphs: list[dict[str, object]]) -> tuple[int | None, int | None]:
    reference_index = None
    for paragraph in paragraphs:
        if is_reference_heading(str(paragraph["text"])):
            reference_index = int(paragraph["index"])
            break
    search_start = reference_index + 1 if reference_index is not None else 0
    appendix_start = None
    for paragraph in paragraphs:
        index = int(paragraph["index"])
        if index < search_start:
            continue
        if is_appendix_heading(str(paragraph["text"])):
            appendix_start = index
            break
    if appendix_start is None:
        return None, None
    appendix_end = len(paragraphs)
    for paragraph in paragraphs:
        index = int(paragraph["index"])
        if index <= appendix_start:
            continue
        if is_after_appendix_tail_heading(str(paragraph["text"])):
            appendix_end = index
            break
    return appendix_start, appendix_end


def collect_cad_pngs(cad_source: Path) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    if cad_source.is_dir():
        for path in sorted(cad_source.rglob("*.png")):
            rows.append(
                {
                    "name": path.name,
                    "path": str(path),
                    "sha256": sha256_file(path),
                    "bytes": path.stat().st_size,
                }
            )
        return rows
    if zipfile.is_zipfile(cad_source):
        with zipfile.ZipFile(cad_source) as zf:
            for info in sorted(zf.infolist(), key=lambda item: item.filename):
                if not info.filename.lower().endswith(".png"):
                    continue
                data = zf.read(info.filename)
                rows.append(
                    {
                        "name": Path(info.filename).name,
                        "path": info.filename,
                        "sha256": sha256_bytes(data),
                        "bytes": len(data),
                    }
                )
        return rows
    raise ValueError(f"CAD source is neither folder nor zip: {cad_source}")


def audit_docx_cad_appendix_binding(
    docx: Path,
    cad_source: Path,
    min_cad_images_in_appendix: int,
    require_all_cad_pngs: bool,
) -> dict[str, object]:
    paragraphs, media_by_target, docx_sha = read_docx_paragraphs_and_images(docx)
    appendix_start, appendix_end = find_appendix_bounds(paragraphs)
    appendix_rows = [
        paragraph
        for paragraph in paragraphs
        if appendix_start is not None and appendix_start <= int(paragraph["index"]) < (appendix_end or len(paragraphs))
    ]
    cad_pngs = collect_cad_pngs(cad_source)
    cad_hash_to_rows: dict[str, list[dict[str, object]]] = {}
    for row in cad_pngs:
        cad_hash_to_rows.setdefault(str(row["sha256"]), []).append(row)
    appendix_image_rows = [row for row in appendix_rows if row["image_hashes"]]
    appendix_hashes = {
        str(value)
        for row in appendix_image_rows
        for value in row["image_hashes"]
    }
    matched_hashes = sorted(hash_value for hash_value in appendix_hashes if hash_value in cad_hash_to_rows)
    matched_cad_rows = [
        cad_row for hash_value in matched_hashes for cad_row in cad_hash_to_rows.get(hash_value, [])
    ]
    missing_cad_rows = [
        row for row in cad_pngs if str(row["sha256"]) not in appendix_hashes
    ]
    schematic_terms = []
    stale_workload_terms = []
    for paragraph in paragraphs:
        text = str(paragraph["text"])
        if SCHEMATIC_SUBSTITUTE_RE.search(text):
            schematic_terms.append({"index": paragraph["index"], "text": text})
        if re.search(r"示意图|样例图|简图", text):
            schematic_terms.append({"index": paragraph["index"], "text": text})
        if re.search(r"三张A0等效工作量图纸|3\.5张A0等效工作量|Three A0-equivalent drawings", text):
            stale_workload_terms.append({"index": paragraph["index"], "text": text})

    issues: list[str] = []
    if appendix_start is None:
        issues.append("formal appendix heading after references was not found")
    if len(cad_pngs) == 0:
        issues.append("CAD source contains no PNG sheet renders for DOCX appendix binding")
    if len(matched_hashes) < min_cad_images_in_appendix:
        issues.append(
            f"appendix embeds {len(matched_hashes)} final CAD sheet render(s), below required minimum {min_cad_images_in_appendix}"
        )
    if require_all_cad_pngs and missing_cad_rows:
        issues.append(
            f"appendix is missing {len(missing_cad_rows)} of {len(cad_pngs)} final CAD PNG sheet render(s)"
        )
    if schematic_terms:
        issues.append("document still contains schematic/placeholder figure wording; final CAD sheets must not be described as schematic/sample/sketch figures")
    if stale_workload_terms:
        issues.append("document still contains stale three-sheet/3.5-A0 workload wording after the final CAD package was revised")

    return {
        "schema": "graduation-project-builder.docx-cad-appendix-binding-audit.v1",
        "docx_path": str(docx),
        "docx_sha256": docx_sha,
        "cad_source_path": str(cad_source),
        "cad_source_sha256": sha256_file(cad_source) if cad_source.is_file() else "",
        "appendix_start_index": appendix_start,
        "appendix_end_index": appendix_end,
        "appendix_paragraph_count": len(appendix_rows),
        "appendix_image_count": len(appendix_image_rows),
        "docx_media_count": len(media_by_target),
        "cad_png_count": len(cad_pngs),
        "matched_cad_png_count": len(matched_hashes),
        "matched_cad_pngs": matched_cad_rows,
        "missing_cad_png_count": len(missing_cad_rows),
        "missing_cad_pngs": missing_cad_rows[:200],
        "min_cad_images_in_appendix": min_cad_images_in_appendix,
        "require_all_cad_pngs": require_all_cad_pngs,
        "schematic_terms": schematic_terms,
        "stale_workload_terms": stale_workload_terms,
        "issues": issues,
        "passed": not issues,
    }


def write_markdown(report: dict[str, object], output: Path) -> None:
    lines = [
        "# DOCX CAD Appendix Binding Audit",
        "",
        f"- docx: `{report['docx_path']}`",
        f"- docx sha256: `{report['docx_sha256']}`",
        f"- CAD source: `{report['cad_source_path']}`",
        f"- appendix image count: {report['appendix_image_count']}",
        f"- CAD PNG count: {report['cad_png_count']}",
        f"- matched CAD PNG count: {report['matched_cad_png_count']}",
        f"- missing CAD PNG count: {report['missing_cad_png_count']}",
        f"- result: {'pass' if report['passed'] else 'fail'}",
        "",
        "## Issues",
    ]
    issues = report.get("issues") or []
    if issues:
        lines.extend(f"- {issue}" for issue in issues)
    else:
        lines.append("- none")
    lines.append("")
    lines.append("## Matched CAD PNGs")
    for row in report.get("matched_cad_pngs", []):
        lines.append(f"- `{row['path']}` sha={str(row['sha256'])[:16]}")
    lines.append("")
    lines.append("## Missing CAD PNGs")
    for row in report.get("missing_cad_pngs", [])[:80]:
        lines.append(f"- `{row['path']}` sha={str(row['sha256'])[:16]}")
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")


def run_self_test() -> int:
    from docx import Document

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        cad = root / "cad"
        cad.mkdir()
        png = cad / "sheet.png"
        png.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?"
            b"\x00\x05\xfe\x02\xfeA\x88\x0f\x9b\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        passing = root / "passing.docx"
        doc = Document()
        doc.add_paragraph("参考文献")
        doc.add_paragraph("1. X[J].")
        doc.add_paragraph("附录")
        doc.add_picture(str(png))
        doc.add_paragraph("附图A.1 CAD图纸")
        doc.save(passing)
        passing_report = audit_docx_cad_appendix_binding(passing, cad, 1, True)
        if not passing_report["passed"]:
            raise AssertionError(passing_report)

        failing = root / "failing.docx"
        doc = Document()
        doc.add_paragraph("参考文献")
        doc.add_paragraph("1. X[J].")
        doc.add_paragraph("附录")
        doc.add_paragraph("只有图纸说明，没有图纸。")
        doc.save(failing)
        failing_report = audit_docx_cad_appendix_binding(failing, cad, 1, True)
        if failing_report["passed"]:
            raise AssertionError("missing CAD image was not rejected")

        schematic = root / "schematic.docx"
        doc = Document()
        doc.add_paragraph("参考文献")
        doc.add_paragraph("1. X[J].")
        doc.add_paragraph("附录")
        doc.add_picture(str(png))
        doc.add_paragraph("Appendix A.1 chain-cylinder schematic placeholder")
        doc.save(schematic)
        schematic_report = audit_docx_cad_appendix_binding(schematic, cad, 1, True)
        if schematic_report["passed"]:
            raise AssertionError("schematic CAD appendix wording was not rejected")
    print("self-test passed")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--cad-source", type=Path)
    parser.add_argument("--report-json", type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--min-cad-images-in-appendix", type=int, default=3)
    parser.add_argument("--require-all-cad-pngs", action="store_true")
    parser.add_argument("--self-test", action="store_true")
    args = parser.parse_args(argv)
    if args.self_test:
        return run_self_test()
    if args.docx is None or args.cad_source is None:
        parser.error("--docx and --cad-source are required unless --self-test is used")
    report = audit_docx_cad_appendix_binding(
        args.docx,
        args.cad_source,
        args.min_cad_images_in_appendix,
        args.require_all_cad_pngs,
    )
    if args.report_json:
        args.report_json.write_text(json.dumps(report, ensure_ascii=True, indent=2), encoding="utf-8")
    if args.report:
        write_markdown(report, args.report)
    else:
        print(json.dumps(report, ensure_ascii=True, indent=2))
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
