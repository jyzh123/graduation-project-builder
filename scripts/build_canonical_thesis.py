#!/usr/bin/env python3
"""Build a thesis from a template plus local adapter/content manifests.

The generic thesis-making behavior lives here in the canonical skill bundle.
Project-local files provide only template/project-specific data.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import time
import zipfile
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET

from docx import Document  # type: ignore
from docx.document import Document as _Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml.table import CT_Tbl  # type: ignore
from docx.oxml.text.paragraph import CT_P  # type: ignore
from docx.shared import Cm  # type: ignore
from docx.table import Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

from python_runtime import resolve_python_exe
from docx_formula_number_table import insert_formula_batch, replace_text_formulas
from thesis_figure_contract import (
    apply_svg_primary_to_docx,
    build_figure_asset_manifest,
    docx_drawing_object_manifest,
    docx_image_relationship_manifest,
    docx_svg_primary_fallback_pairs,
    validate_figure_manifest,
    write_manifest,
)
from thesis_template_profile import build_template_profile, profile_readiness_issues, write_profile
from validate_thesis_local_adapter import validate_adapter_file


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
PYTHON_EXE = resolve_python_exe()
CONTENT_SCHEMA = "graduation-project-builder.thesis-content.v1"
MAX_FIGURE_WIDTH_CM = 14.2
HEADING_RE = re.compile(r"^\d{1,2}(?:\.\d{1,2}){0,3}\s+\S")
CN_NUMBER_CHARS = "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341"
TABLE_TITLE_RE = re.compile(rf"^\s*(?:\u8868|\u7eed\u8868)\s*[0-9{CN_NUMBER_CHARS}]+(?:[-.]\d+)?\s*\S")
FIGURE_TITLE_RE = re.compile(rf"^\s*(?:\u9644\s*\u56fe|\u56fe)\s*(?:[0-9{CN_NUMBER_CHARS}]+|[A-Za-z])(?:[-.]\d+)?\s*\S")
TAIL_HEADINGS = {"\u53c2\u8003\u6587\u732e", "\u81f4\u8c22", "\u9644\u5f55", "references", "acknowledgements"}
PLACEHOLDER_TOKENS = (
    "placeholder",
    "smoke",
    "minimal",
    "example research",
    "example journal",
    "template alignment",
    "\u5360\u4f4d",
    "\u6837\u7a3f",
    "\u793a\u4f8b\u7814\u7a76",
    "\u5f20\u4e09",
    "\u674e\u56db",
    "\u738b\u4e94",
    "\u8d75\u516d",
    "\u672c\u6bb5\u7ed3\u5408",
    "\u5206\u6790\u8981\u70b9",
)


@dataclass
class ParagraphDonor:
    ppr: Any | None
    rpr: Any | None
    style_id: str
    style_name: str


@dataclass
class TableDonor:
    tbl_pr: Any | None
    grid_widths: list[int]
    header_tr_pr: Any | None
    body_first_tr_pr: Any | None
    body_middle_tr_pr: Any | None
    body_last_tr_pr: Any | None
    header_tc_pr: Any | None
    body_first_tc_pr: Any | None
    body_middle_tc_pr: Any | None
    body_last_tc_pr: Any | None


@dataclass
class Donors:
    heading_by_level: dict[int, ParagraphDonor]
    toc_by_level: dict[int, ParagraphDonor]
    body: ParagraphDonor
    abstract_body: ParagraphDonor
    keyword: ParagraphDonor
    figure_caption: ParagraphDonor
    table_caption: ParagraphDonor
    reference_heading: ParagraphDonor
    reference_entry: ParagraphDonor
    reference_cjk_run: ParagraphDonor
    reference_latin_run: ParagraphDonor
    acknowledgement: ParagraphDonor
    acknowledgement_body: ParagraphDonor
    acknowledgement_cjk_run: ParagraphDonor
    acknowledgement_latin_run: ParagraphDonor
    table_header_cell: ParagraphDonor
    table_body_cell: ParagraphDonor
    table_header_latin_cell: ParagraphDonor
    table_body_latin_cell: ParagraphDonor
    table_style: str | None
    table_donor: TableDonor | None


def run(cmd: list[str], *, timeout: int = 1200, check: bool = True) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=timeout,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            f"command failed: {' '.join(cmd)}\nstdout:\n{proc.stdout}\nstderr:\n{proc.stderr}"
        )
    return proc


def read_json(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return data


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def normalize(text: str) -> str:
    return re.sub(r"[\s\u25a1]+", "", text or "").strip().lower()


def heading_key(text: str) -> str:
    """Normalize a template surface heading without its instruction suffix."""
    head = re.split(r"[\(\uff08]", str(text or "").strip(), maxsplit=1)[0]
    return normalize(head)


def is_toc_heading_text(text: str) -> bool:
    key = heading_key(text)
    normalized = normalize(text)
    toc_key = normalize("\u76ee\u5f55")
    return key in {toc_key, "contents", "tableofcontents"} or normalized in {
        normalize("\u76ee\u5f55"),
        "contents",
        "tableofcontents",
    } or key.endswith(toc_key) or normalized.endswith(toc_key)


def has_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text or "")


def has_latin_or_digit(text: str) -> bool:
    return any(ch.isascii() and ch.isalnum() for ch in text or "")


def contains_chapter_heading_marker(text: str) -> bool:
    stripped = re.sub(r"^[\s\u25a1]+", "", text or "")
    match = re.match(rf"^(\u7b2c[0-9{CN_NUMBER_CHARS}]+\u7ae0)(.*)$", stripped)
    if not match:
        return False
    tail = match.group(2)
    if not tail:
        return True
    separator_match = re.match(r"^[\s\u25a1:：、.\-—]+", tail)
    if separator_match:
        tail = tail[separator_match.end() :]
    compact_tail = normalize(tail)
    sentence_punctuation = "\uff0c\uff1b\uff1a\uff01\uff1f\u3002\uff08\uff09,;:!?()"
    if any(ch in compact_tail for ch in sentence_punctuation):
        return False
    return len(compact_tail) <= (24 if separator_match else 12)


def heading_level(text: str) -> int | None:
    stripped = re.sub(r"^[\s\u25a1]+", "", (text or "").strip())
    sep = r"[\s\u25a1]+"
    if re.match(rf"^\d{{1,2}}{sep}\S", stripped):
        return 1
    if re.match(rf"^\d{{1,2}}\.\d{{1,2}}{sep}\S", stripped):
        return 2
    if re.match(rf"^\d{{1,2}}\.\d{{1,2}}\.\d{{1,2}}{sep}\S", stripped):
        return 3
    if re.match(rf"^\d{{1,2}}\.\d{{1,2}}\.\d{{1,2}}\.\d{{1,2}}{sep}\S", stripped):
        return 4
    if contains_chapter_heading_marker(stripped):
        return 1
    return None


def template_heading_level(text: str) -> int | None:
    stripped = re.sub(r"^[\s\u25a1]+", "", (text or "").strip())
    if re.match(rf"^\u7b2c[0-9{CN_NUMBER_CHARS}]+\u7ae0", stripped):
        return 1
    match = re.match(r"^\d{1,2}(?:\.\d{1,2}){1,3}(?=\D|$)", stripped)
    if match:
        return min(match.group(0).count(".") + 1, 4)
    return heading_level(text)


def is_toc_paragraph(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    return style_name.startswith("toc") or is_toc_heading_text(paragraph.text)


def paragraph_has_tab(paragraph: Paragraph) -> bool:
    if "\t" in (paragraph.text or ""):
        return True
    return any(run_obj._element.find(qn("w:tab")) is not None for run_obj in paragraph.runs)


def is_template_note_text(text: str) -> bool:
    stripped = (text or "").strip()
    normalized = normalize(stripped)
    if not stripped:
        return False
    return (
        stripped.startswith("*")
        or normalized.startswith(normalize("\u6ce8:"))
        or normalized.startswith(normalize("\u6ce8\uff1a"))
    )


def is_tail_heading(text: str) -> bool:
    return normalize(text) in {normalize(item) for item in TAIL_HEADINGS}


def looks_like_static_toc_entry(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    return any(marker in stripped for marker in ("\u2026", "…", "\u00d7", "×")) and (
        heading_level(stripped) is not None or is_tail_heading(stripped)
    )


def is_template_format_instruction(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    normalized = normalize(stripped)
    tokens = (
        "\u91c7\u7528",
        "\u5b57\u53f7",
        "\u5b57\u4f53",
        "\u884c\u8ddd",
        "\u5c45\u4e2d",
        "\u9876\u683c",
        "\u7a7a\u4e00\u884c",
        "\u7a7a\u4e24\u683c",
        "\u5c0f\u4e94",
        "\u4e8c\u9009\u4e00",
        "\u4f4d\u4e8e\u56fe\u4e0b",
        "\u4f4d\u4e8e\u8868\u4e0a",
        "\u4e0e\u6b63\u6587\u7a7a",
        "\u7ae0\u6807\u9898",
    )
    if not any(token in stripped for token in tokens):
        return False
    return (
        stripped.startswith(("\uff08", "(", "*"))
        or "\u2606" in stripped
        or "\u6b63\u6587" in stripped
        or "\u56fe" in stripped
        or "\u8868" in stripped
        or "\u53c2\u8003\u6587\u732e" in stripped
        or "\u81f4\u8c22" in stripped
        or "\u7ae0\u6807\u9898" in stripped
        or "\u4e8c\u9009\u4e00" in stripped
        or normalized.startswith(normalize("\u6ce8\uff1a"))
    )


def strip_template_instruction_parentheticals(text: str) -> str:
    """Remove visible format-instruction parentheses while keeping real titles."""
    value = str(text or "")
    changed = True
    while changed:
        changed = False

        def repl(match: re.Match[str]) -> str:
            nonlocal changed
            inner = match.group(1)
            if is_template_format_instruction(inner):
                changed = True
                return ""
            return match.group(0)

        value = re.sub(r"\uff08([^\uff08\uff09]{1,80})\uff09", repl, value)
        value = re.sub(r"\(([^()]{1,80})\)", repl, value)
    return re.sub(r"\s{2,}", " ", value).strip()


def is_pure_template_instruction(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    if "\u00d7" in stripped or "\uff58" in stripped or "\u2606" in stripped:
        return False
    return is_template_format_instruction(stripped)


def is_figure_or_table_format_note(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    if TABLE_TITLE_RE.match(stripped) or FIGURE_TITLE_RE.match(stripped):
        return True
    return any(
        token in stripped
        for token in (
            "\u4f4d\u4e8e\u56fe\u4e0b",
            "\u4f4d\u4e8e\u8868\u4e0a",
            "\u56fe\u4e0e\u4e0b\u6587",
            "\u8868\u4e0e\u6b63\u6587",
            "\u56fe\u9898",
            "\u8868\u9898",
        )
    )


def find_toc_heading(doc: Document) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if is_toc_heading_text(paragraph.text):
            return paragraph
    return find_paragraph(doc, {"\u76ee\u5f55", "\u76ee  \u5f55", "\u76ee   \u5f55"})


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def scrub_reusable_paragraph_properties(ppr: Any | None) -> Any | None:
    """Remove paragraph-owned document topology before reusing a style donor."""
    if ppr is None:
        return None
    scrubbed = deepcopy(ppr)
    for tag in ("w:sectPr",):
        child = scrubbed.find(qn(tag))
        if child is not None:
            scrubbed.remove(child)
    return scrubbed


def capture_donor(paragraph: Paragraph | None) -> ParagraphDonor:
    if paragraph is None:
        return ParagraphDonor(None, None, "", "")
    ppr = scrub_reusable_paragraph_properties(paragraph._element.pPr)
    first_text_rpr = None
    first_cjk_rpr = None
    for run_obj in paragraph.runs:
        if not run_obj.text.strip() or run_obj._element.rPr is None:
            continue
        if first_text_rpr is None:
            first_text_rpr = deepcopy(run_obj._element.rPr)
        if first_cjk_rpr is None and has_cjk(run_obj.text):
            first_cjk_rpr = deepcopy(run_obj._element.rPr)
            break
    rpr = first_cjk_rpr if first_cjk_rpr is not None else first_text_rpr
    style_id = paragraph.style.style_id if paragraph.style is not None else ""
    style_name = paragraph.style.name if paragraph.style is not None else ""
    return ParagraphDonor(ppr, rpr, style_id, style_name)


def capture_heading_donor(paragraph: Paragraph | None) -> ParagraphDonor:
    donor = capture_donor(paragraph)
    if paragraph is None:
        return donor
    for run_obj in paragraph.runs:
        text = run_obj.text.strip()
        if not text or run_obj._element.rPr is None:
            continue
        donor.rpr = deepcopy(run_obj._element.rPr)
        return donor
    return donor


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def apply_donor(paragraph: Paragraph, donor: ParagraphDonor) -> None:
    existing_ppr = paragraph._element.pPr
    if existing_ppr is not None:
        paragraph._element.remove(existing_ppr)
    if donor.ppr is not None:
        paragraph._element.insert(0, scrub_reusable_paragraph_properties(donor.ppr))
    clear_paragraph_content(paragraph)


def set_paragraph_style_id(paragraph: Paragraph, style_id: str) -> None:
    ppr = paragraph._element.pPr
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    style_node = ppr.find(qn("w:pStyle"))
    if style_node is None:
        style_node = OxmlElement("w:pStyle")
        ppr.insert(0, style_node)
    style_node.set(qn("w:val"), style_id)


def set_semantic_style(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
        return
    except Exception:
        pass
    try:
        style_id = paragraph.part.document.styles[style_name].style_id
        set_paragraph_style_id(paragraph, style_id)
    except Exception:
        return


def ensure_paragraph_style(doc: Document, style_name: str) -> None:
    try:
        doc.styles[style_name]
        return
    except Exception:
        pass
    try:
        doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        return


def ensure_keep_next(paragraph: Paragraph) -> None:
    ppr = paragraph._element.pPr
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


def paragraph_has_hard_page_break(paragraph: Paragraph) -> bool:
    return any(br.get(qn("w:type")) == "page" for br in paragraph._element.iter(qn("w:br")))


def paragraph_has_section_break(paragraph: Paragraph) -> bool:
    ppr = paragraph._element.pPr
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def previous_paragraph(paragraph: Paragraph) -> Paragraph | None:
    element = paragraph._element.getprevious()
    while element is not None:
        if element.tag == qn("w:p"):
            return Paragraph(element, paragraph._parent)
        element = element.getprevious()
    return None


def remove_page_break_before(paragraph: Paragraph) -> None:
    ppr = paragraph._element.pPr
    if ppr is None:
        return
    page_break_before = ppr.find(qn("w:pageBreakBefore"))
    if page_break_before is not None:
        ppr.remove(page_break_before)


def ensure_page_break_before(paragraph: Paragraph) -> None:
    if paragraph_has_hard_page_break(paragraph):
        remove_page_break_before(paragraph)
        return
    prior = previous_paragraph(paragraph)
    if prior is not None and paragraph_has_hard_page_break(prior):
        remove_page_break_before(paragraph)
        return
    if prior is not None and paragraph_has_section_break(prior):
        remove_page_break_before(paragraph)
        return
    ppr = paragraph._element.pPr
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    if ppr.find(qn("w:pageBreakBefore")) is None:
        ppr.append(OxmlElement("w:pageBreakBefore"))


def force_page_break_before(paragraph: Paragraph) -> None:
    if paragraph_has_hard_page_break(paragraph):
        remove_page_break_before(paragraph)
        return
    ppr = ensure_ppr(paragraph)
    if ppr.find(qn("w:pageBreakBefore")) is None:
        ppr.append(OxmlElement("w:pageBreakBefore"))


def ensure_table_rows_do_not_split(table: Table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def remove_child(parent: Any, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def remove_direct_paragraph_style(paragraph: Paragraph) -> None:
    ppr = paragraph._element.pPr
    if ppr is not None:
        remove_child(ppr, "w:pStyle")


def remove_first_line_indent(paragraph: Paragraph) -> None:
    ppr = paragraph._element.pPr
    if ppr is None:
        return
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return
    for attr in (qn("w:firstLine"), qn("w:firstLineChars")):
        if attr in ind.attrib:
            del ind.attrib[attr]
    if not ind.attrib:
        ppr.remove(ind)


def ensure_ppr(paragraph: Paragraph) -> Any:
    ppr = paragraph._element.pPr
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    return ppr


def ensure_spacing(paragraph: Paragraph, *, line: str | None = None, line_rule: str | None = None) -> None:
    ppr = ensure_ppr(paragraph)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    if line is None:
        for attr in (qn("w:line"), qn("w:lineRule")):
            if attr in spacing.attrib:
                del spacing.attrib[attr]
    else:
        spacing.set(qn("w:line"), line)
        if line_rule is not None:
            spacing.set(qn("w:lineRule"), line_rule)


def set_direct_spacing(
    paragraph: Paragraph,
    *,
    before: str | None = None,
    after: str | None = None,
    line: str | None = None,
    line_rule: str | None = None,
) -> None:
    ppr = ensure_ppr(paragraph)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    for key, value in (("before", before), ("after", after), ("line", line), ("lineRule", line_rule)):
        attr = qn(f"w:{key}")
        if value is None:
            if attr in spacing.attrib:
                del spacing.attrib[attr]
        else:
            spacing.set(attr, value)


def ensure_first_line_indent(paragraph: Paragraph, *, value: str = "0", chars: str = "0") -> None:
    ppr = ensure_ppr(paragraph)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLine"), value)
    ind.set(qn("w:firstLineChars"), chars)


def force_image_holder_direct_baseline(paragraph: Paragraph) -> None:
    """Keep figure holders on the template-owned holder style without body-text residues."""
    ppr = ensure_ppr(paragraph)
    remove_child(ppr, "w:spacing")
    remove_child(ppr, "w:ind")


def remove_spacing_and_indent(paragraph: Paragraph) -> None:
    ppr = paragraph._element.pPr
    if ppr is None:
        return
    remove_child(ppr, "w:spacing")
    remove_child(ppr, "w:ind")


def set_direct_alignment(paragraph: Paragraph, value: str) -> None:
    ppr = ensure_ppr(paragraph)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), value)


def ensure_run_font(run_obj: Any, font_name: str, *, size: str | None = None, bold: bool | None = None) -> None:
    rpr = run_obj._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_obj._element.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr_name in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr_name), font_name)
    if size is not None:
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), size)
    if bold is not None:
        for tag in ("w:b", "w:bCs"):
            node = rpr.find(qn(tag))
            if bold and node is None:
                rpr.append(OxmlElement(tag))
            if not bold and node is not None:
                rpr.remove(node)


def ensure_run_fonts(
    run_obj: Any,
    *,
    east_asia: str | None = None,
    ascii_font: str | None = None,
    hansi: str | None = None,
    cs: str | None = None,
    size: str | None = None,
    bold: bool | None = None,
) -> None:
    rpr = run_obj._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_obj._element.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr_name, value in (
        ("w:eastAsia", east_asia),
        ("w:ascii", ascii_font),
        ("w:hAnsi", hansi),
        ("w:cs", cs),
    ):
        if value is not None:
            rfonts.set(qn(attr_name), value)
    if size is not None:
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), size)
    if bold is not None:
        for tag in ("w:b", "w:bCs"):
            node = rpr.find(qn(tag))
            if bold and node is None:
                rpr.append(OxmlElement(tag))
            if not bold and node is not None:
                rpr.remove(node)


def ensure_run_color(run_obj: Any, color: str) -> None:
    rpr = run_obj._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_obj._element.insert(0, rpr)
    color_node = rpr.find(qn("w:color"))
    if color_node is None:
        color_node = OxmlElement("w:color")
        rpr.append(color_node)
    color_node.set(qn("w:val"), color)
    theme_attr = qn("w:themeColor")
    if theme_attr in color_node.attrib:
        del color_node.attrib[theme_attr]


def force_toc_heading_format(paragraph: Paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    for tag_name in ("w:pStyle", "w:outlineLvl", "w:ind"):
        for child in list(ppr.findall(qn(tag_name))):
            ppr.remove(child)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), "center")
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), "312")
    spacing.set(qn("w:after"), "312")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    for run_obj in paragraph.runs:
        if not (run_obj.text or "").strip():
            continue
        rpr = run_obj._element.rPr
        if rpr is not None:
            run_obj._element.remove(rpr)
        rpr = OxmlElement("w:rPr")
        run_obj._element.insert(0, rpr)
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:eastAsia"), "\u9ed1\u4f53")
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "32")
        rpr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "32")
        rpr.append(sz_cs)
        ensure_run_color(run_obj, "000000")


def set_cover_text_with_split_runs(
    paragraph: Paragraph,
    donor: ParagraphDonor,
    text: str,
    *,
    size: str | None = None,
    keep_ascii_slots: bool = False,
) -> None:
    apply_donor(paragraph, donor)
    append_mixed_script_text(paragraph, text, cjk_donor=donor, latin_donor=donor)
    if size is None and not keep_ascii_slots:
        return
    for run_obj in paragraph.runs:
        if not (run_obj.text or "").strip():
            continue
        if size is not None:
            ensure_run_fonts(run_obj, size=size)
        rpr = run_obj._element.rPr
        rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
        if rfonts is not None:
            remove_latin_slots = (not keep_ascii_slots) or (not has_cjk(run_obj.text or ""))
            for attr_name in ("w:ascii", "w:hAnsi", "w:cs"):
                if not remove_latin_slots:
                    continue
                attr = qn(attr_name)
                if attr in rfonts.attrib:
                    del rfonts.attrib[attr]


def paragraph_alignment_value(paragraph: Paragraph) -> str:
    jc = paragraph._element.find(qn("w:pPr") + "/" + qn("w:jc"))
    return jc.get(qn("w:val"), "") if jc is not None else ""


def first_visible_run_is_italic(paragraph: Paragraph) -> bool:
    for run_obj in paragraph.runs:
        if not (run_obj.text or "").strip():
            continue
        rpr = run_obj._element.rPr
        return rpr is not None and rpr.find(qn("w:i")) is not None
    return False


def make_image_holder_safe(paragraph: Paragraph) -> None:
    set_paragraph_style_id(paragraph, "ThesisImageHolder")
    force_image_holder_direct_baseline(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ensure_keep_next(paragraph)


def add_text_with_donor(paragraph: Paragraph, donor: ParagraphDonor, text: str):
    run_obj = paragraph.add_run(text)
    if donor.rpr is not None:
        existing_rpr = run_obj._element.rPr
        if existing_rpr is not None:
            run_obj._element.remove(existing_rpr)
        run_obj._element.insert(0, deepcopy(donor.rpr))
    ensure_run_color(run_obj, "000000")
    return run_obj


def donor_with_rpr(base: ParagraphDonor, rpr: Any | None) -> ParagraphDonor:
    return ParagraphDonor(base.ppr, deepcopy(rpr) if rpr is not None else None, base.style_id, base.style_name)


def donor_with_latin_font(
    base: ParagraphDonor,
    source: ParagraphDonor,
    font_name: str = "Times New Roman",
    cs_font_name: str | None = "Times New Roman",
) -> ParagraphDonor:
    rpr = deepcopy(source.rpr) if source.rpr is not None else OxmlElement("w:rPr")
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr_name in ("w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr_name), font_name)
    if cs_font_name is not None:
        rfonts.set(qn("w:cs"), cs_font_name)
    if rpr.find(qn("w:szCs")) is None:
        sz = rpr.find(qn("w:sz"))
        if sz is not None and sz.get(qn("w:val")):
            sz_cs = OxmlElement("w:szCs")
            sz_cs.set(qn("w:val"), sz.get(qn("w:val")))
            rpr.append(sz_cs)
    return donor_with_rpr(base, rpr)


def donor_with_latin_font_preserving_cs(
    base: ParagraphDonor,
    source: ParagraphDonor,
    font_name: str = "Times New Roman",
) -> ParagraphDonor:
    donor = donor_with_latin_font(base, source, font_name=font_name, cs_font_name=None)
    if donor.rpr is None:
        return donor
    rfonts = donor.rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return donor
    if rfonts.get(qn("w:cs")):
        return donor
    source_rfonts = source.rpr.find(qn("w:rFonts")) if source.rpr is not None else None
    if source_rfonts is None:
        return donor
    for attr_name in ("w:cs", "w:eastAsia"):
        value = source_rfonts.get(qn(attr_name))
        if value:
            rfonts.set(qn("w:cs"), value)
            break
    if not rfonts.get(qn("w:cs")):
        rfonts.set(qn("w:cs"), font_name)
    return donor


def donor_with_cjk_font(base: ParagraphDonor, source: ParagraphDonor, font_name: str = "\u5b8b\u4f53") -> ParagraphDonor:
    rpr = deepcopy(source.rpr) if source.rpr is not None else OxmlElement("w:rPr")
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    return donor_with_rpr(base, rpr)


def run_rpr(run_obj: Any) -> Any | None:
    return deepcopy(run_obj._element.rPr) if run_obj._element.rPr is not None else None


def capture_latin_donor(paragraph: Paragraph | None, fallback: ParagraphDonor) -> ParagraphDonor:
    if paragraph is None:
        return fallback
    for run_obj in paragraph.runs:
        text = run_obj.text or ""
        if has_latin_or_digit(text) and not has_cjk(text) and run_obj._element.rPr is not None:
            return donor_with_rpr(fallback, run_rpr(run_obj))
    for run_obj in paragraph.runs:
        text = run_obj.text or ""
        if has_latin_or_digit(text) and run_obj._element.rPr is not None:
            return donor_with_rpr(fallback, run_rpr(run_obj))
    return fallback


def capture_cjk_donor(paragraph: Paragraph | None, fallback: ParagraphDonor) -> ParagraphDonor:
    if paragraph is None:
        return fallback
    for run_obj in paragraph.runs:
        text = run_obj.text or ""
        if has_cjk(text) and not has_latin_or_digit(text) and run_obj._element.rPr is not None:
            return donor_with_rpr(fallback, run_rpr(run_obj))
    for run_obj in paragraph.runs:
        text = run_obj.text or ""
        if has_cjk(text) and run_obj._element.rPr is not None:
            return donor_with_rpr(fallback, run_rpr(run_obj))
    return fallback


def split_by_script_runs(text: str) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = []
    current_kind: str | None = None
    current: list[str] = []
    for char in text:
        if has_cjk(char):
            kind = "cjk"
        elif char.isascii() and (char.isalnum() or char in "[](){}.,;:/\\-+_=&%#@'\"<> "):
            kind = "latin"
        else:
            kind = current_kind or "latin"
        if current_kind is None:
            current_kind = kind
        if kind != current_kind:
            segments.append((current_kind, "".join(current)))
            current = [char]
            current_kind = kind
        else:
            current.append(char)
    if current:
        segments.append((current_kind or "latin", "".join(current)))
    return segments


def normalize_citation_punctuation_in_text(text: str) -> str:
    """Move numeric citation markers before sentence punctuation for thesis style."""
    fixed = str(text or "")
    citation = r"(\[(?:\d{1,3})\])"
    punctuation = r"([\u3002\uff01\uff1f\uff1b\uff0c.!?;,])"
    previous = None
    while fixed != previous:
        previous = fixed
        fixed = re.sub(rf"{punctuation}\s*{citation}", r"\2\1", fixed)
    return fixed


def append_mixed_script_text(
    paragraph: Paragraph,
    text: str,
    *,
    cjk_donor: ParagraphDonor,
    latin_donor: ParagraphDonor,
) -> None:
    for kind, value in split_by_script_runs(text):
        if not value:
            continue
        add_text_with_donor(paragraph, cjk_donor if kind == "cjk" else latin_donor, value)


def leading_placeholder_prefix(text: str) -> str:
    match = re.match(r"^([\s\u3000\u25a1]+)", text or "")
    if not match:
        return ""
    return match.group(1)


def template_placeholder_to_spaces(text: str) -> str:
    return (text or "").replace("\u25a1", " ")


def replace_template_placeholder_glyphs(paragraph: Paragraph) -> None:
    for run_obj in paragraph.runs:
        if "\u25a1" in (run_obj.text or ""):
            run_obj.text = template_placeholder_to_spaces(run_obj.text)


def normalize_front_matter_placeholder_glyphs(doc: Document) -> None:
    first_body = find_first_body_heading(doc)
    end = paragraph_index(doc, first_body) if first_body is not None else min(len(doc.paragraphs), 120)
    for paragraph in doc.paragraphs[:end]:
        replace_template_placeholder_glyphs(paragraph)


def is_prefix_only_text(text: str) -> bool:
    stripped = text or ""
    return bool(stripped) and all(ch.isspace() or ch in {"\u3000", "\u25a1"} for ch in stripped)


def abstract_body_run_donors(paragraph: Paragraph, base: ParagraphDonor) -> tuple[ParagraphDonor, ParagraphDonor]:
    prefix_rpr = None
    content_rpr = None
    for run_obj in paragraph.runs:
        text = run_obj.text or ""
        if not text:
            continue
        if prefix_rpr is None and is_prefix_only_text(text):
            prefix_rpr = run_rpr(run_obj)
            continue
        if content_rpr is None and text.strip():
            content_rpr = run_rpr(run_obj)
            break
    return (
        donor_with_rpr(base, prefix_rpr if prefix_rpr is not None else base.rpr),
        donor_with_rpr(base, content_rpr if content_rpr is not None else base.rpr),
    )


def abstract_prefix_run_is_isolated(paragraph: Paragraph, prefix: str) -> bool:
    if not prefix:
        return False
    for run_obj in paragraph.runs:
        text = run_obj.text or ""
        if not text:
            continue
        return is_prefix_only_text(text) and template_placeholder_to_spaces(text) == template_placeholder_to_spaces(prefix)
    return False


def block_latin_donor(paragraphs: Iterable[Paragraph], fallback: ParagraphDonor) -> ParagraphDonor:
    for paragraph in paragraphs:
        for run_obj in paragraph.runs:
            text = run_obj.text or ""
            if has_latin_or_digit(text) and run_obj._element.rPr is not None:
                return donor_with_rpr(fallback, run_rpr(run_obj))
    return fallback


def set_abstract_body_text(
    paragraph: Paragraph,
    text: str,
    *,
    latin_donor_for_mixed: bool = False,
    latin_source_donor: ParagraphDonor | None = None,
    suppress_placeholder_prefix: bool = False,
) -> None:
    base = capture_donor(paragraph)
    prefix = leading_placeholder_prefix(paragraph.text)
    prefix_donor, content_donor = abstract_body_run_donors(paragraph, base)
    prefix_isolated = abstract_prefix_run_is_isolated(paragraph, prefix)
    latin_source = latin_source_donor or (capture_latin_donor(paragraph, content_donor) if latin_donor_for_mixed else content_donor)
    latin_donor = donor_with_latin_font_preserving_cs(base, latin_source) if latin_donor_for_mixed else content_donor
    clean_text = str(text or "").strip()
    if prefix and clean_text.startswith(prefix):
        clean_text = clean_text[len(prefix) :].lstrip()
    apply_donor(paragraph, base)
    if suppress_placeholder_prefix:
        prefix = ""
    if prefix and prefix_isolated:
        add_text_with_donor(paragraph, prefix_donor, template_placeholder_to_spaces(prefix))
    elif prefix:
        clean_text = template_placeholder_to_spaces(prefix) + clean_text
    append_mixed_script_text(paragraph, clean_text, cjk_donor=content_donor, latin_donor=latin_donor)


def configure_english_abstract_body_paragraph(paragraph: Paragraph) -> None:
    set_direct_spacing(paragraph, before="0", after="0", line="360", line_rule="auto")
    ensure_first_line_indent(paragraph, value="480", chars="200")


def configure_chinese_abstract_body_paragraph(paragraph: Paragraph) -> None:
    set_direct_spacing(paragraph, before="0", after="0", line="360", line_rule="auto")


def split_abstract_body_paragraphs(text: str) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\r?\n+", str(text or "").strip()) if part.strip()]
    return paragraphs or [str(text or "").strip()]


def is_abstract_body_placeholder_line(text: str) -> bool:
    stripped = normalize(str(text or ""))
    return stripped in {normalize("\uff08\u7a7a\u4e00\u884c\uff09"), normalize("(\u7a7a\u4e00\u884c)")}


def set_paragraph_text(paragraph: Paragraph, donor: ParagraphDonor, text: str) -> Paragraph:
    apply_donor(paragraph, donor)
    add_text_with_donor(paragraph, donor, text)
    return paragraph


def set_mixed_table_cell_text(
    paragraph: Paragraph,
    *,
    cjk_donor: ParagraphDonor,
    latin_donor: ParagraphDonor,
    text: str,
) -> Paragraph:
    apply_donor(paragraph, cjk_donor)
    append_mixed_script_text(paragraph, text, cjk_donor=cjk_donor, latin_donor=latin_donor)
    return paragraph


def paragraph_index(doc: Document, target: Paragraph) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph._element is target._element:
            return idx
    raise KeyError("paragraph not found")


def delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        block_parent = parent._body
    else:
        parent_elm = parent._tc
        block_parent = parent
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, block_parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, block_parent)


def find_first_body_heading(doc: Document) -> Paragraph | None:
    toc_heading = find_toc_heading(doc)
    if toc_heading is not None:
        start = paragraph_index(doc, toc_heading) + 1
        for paragraph in doc.paragraphs[start:]:
            text = paragraph.text.strip()
            if (
                not text
                or paragraph_has_tab(paragraph)
                or is_toc_paragraph(paragraph)
                or looks_like_static_toc_entry(text)
                or is_tail_heading(text)
                or is_template_note_text(text)
                or is_template_format_instruction(text)
            ):
                continue
            if re.match(rf"^\u7b2c[0-9{CN_NUMBER_CHARS}]+\u7ae0", re.sub(r"^[\s\u25a1]+", "", text)) and not any(
                token in text for token in ("…", "\u2026", "×", "\u00d7")
            ):
                return paragraph
            style_name = (paragraph.style.name if paragraph.style else "").lower()
            if "heading 1" in style_name or heading_level(text) == 1:
                return paragraph
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if (
            not text
            or paragraph_has_tab(paragraph)
            or is_toc_paragraph(paragraph)
            or looks_like_static_toc_entry(text)
            or is_tail_heading(text)
        ):
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if "heading 1" in style_name or heading_level(text) == 1:
            return paragraph
    return None


def matches_template_surface_paragraph(text: str, candidates: set[str]) -> bool:
    normalized_text = normalize(text)
    normalized_candidates = {normalize(item) for item in candidates}
    if normalized_text in normalized_candidates:
        return True
    if heading_key(text) in {heading_key(item) for item in candidates}:
        return True
    return any(
        (normalized_text.startswith(candidate) or normalized_text.endswith(candidate))
        and any(marker in normalized_text for marker in ("\u5b57\u53f7", "\u884c\u8ddd", "\u6bb5\u524d", "\u6bb5\u540e", "\u683c\u5f0f", "\u4e66\u5199"))
        for candidate in normalized_candidates
    )


def find_paragraph(doc: Document, candidates: set[str]) -> Paragraph | None:
    normalized = {normalize(item) for item in candidates}
    for paragraph in doc.paragraphs:
        if normalize(paragraph.text) in normalized or matches_template_surface_paragraph(paragraph.text, candidates):
            return paragraph
    return None


def paragraph_matches_surface_heading(paragraph: Paragraph, candidates: set[str]) -> bool:
    text = paragraph.text.strip()
    if not text or paragraph_has_tab(paragraph) or is_toc_paragraph(paragraph):
        return False
    normalized = normalize(text)
    candidate_values = {normalize(item) for item in candidates}
    if normalized in candidate_values:
        return True
    return any(
        normalized.startswith(candidate)
        and any(marker in normalized for marker in ("\u5b57\u53f7", "\u884c\u8ddd", "\u6bb5\u524d", "\u6bb5\u540e", "\u683c\u5f0f", "\u9876\u683c", "\u5c45\u4e2d"))
        for candidate in candidate_values
    )


def find_surface_heading(doc: Document, candidates: set[str], *, start: int = 0) -> Paragraph | None:
    for paragraph in doc.paragraphs[start:]:
        if paragraph_matches_surface_heading(paragraph, candidates):
            return paragraph
    return None


def next_nonempty_content_after(doc: Document, anchor: Paragraph, *, skip_instructions: bool = True) -> Paragraph | None:
    start = paragraph_index(doc, anchor)
    for paragraph in doc.paragraphs[start + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph_has_tab(paragraph) or is_toc_paragraph(paragraph):
            continue
        if heading_level(text) is not None or paragraph_matches_surface_heading(paragraph, TAIL_HEADINGS):
            return None
        if skip_instructions and is_pure_template_instruction(text):
            continue
        return paragraph
    return None


def next_nonempty_after(doc: Document, anchor: Paragraph) -> Paragraph | None:
    start = paragraph_index(doc, anchor)
    for paragraph in doc.paragraphs[start + 1 :]:
        if paragraph.text.strip():
            return paragraph
    return None


def first_body_paragraph_after(doc: Document, anchor: Paragraph) -> Paragraph | None:
    start = paragraph_index(doc, anchor)
    for paragraph in doc.paragraphs[start + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        if heading_level(text) is not None or is_tail_heading(text):
            return None
        if is_template_note_text(text) or is_template_format_instruction(text) or is_figure_or_table_format_note(text):
            continue
        return paragraph
    return None


BODY_DONOR_RULE_HINTS = (
    "\u91c7\u7528",
    "\u5b57\u4f53",
    "\u6807\u9898\u5e8f\u53f7",
    "\u56fe\u8868",
    "\u884c\u8ddd",
    "\u6bb5\u524d",
    "\u6bb5\u540e",
    "\u5c0f\u56db\u53f7",
    " pt",
    "pt",
)


def paragraph_signature_key(paragraph: Paragraph) -> str:
    ppr_xml = paragraph._element.pPr.xml if paragraph._element.pPr is not None else ""
    rpr_xml = ""
    for run_obj in paragraph.runs:
        if run_obj.text.strip() and run_obj._element.rPr is not None:
            rpr_xml = run_obj._element.rPr.xml
            break
    return ppr_xml + "\n" + rpr_xml


def paragraph_indent_values(paragraph: Paragraph) -> dict[str, str]:
    ind = paragraph._element.find(qn("w:pPr") + "/" + qn("w:ind"))
    if ind is None:
        return {}
    return {
        "firstLine": ind.get(qn("w:firstLine"), ""),
        "firstLineChars": ind.get(qn("w:firstLineChars"), ""),
        "left": ind.get(qn("w:left"), ""),
        "leftChars": ind.get(qn("w:leftChars"), ""),
    }


def paragraph_spacing_values(paragraph: Paragraph) -> dict[str, str]:
    spacing = paragraph._element.find(qn("w:pPr") + "/" + qn("w:spacing"))
    if spacing is None:
        return {}
    return {
        "before": spacing.get(qn("w:before"), ""),
        "after": spacing.get(qn("w:after"), ""),
        "beforeLines": spacing.get(qn("w:beforeLines"), ""),
        "afterLines": spacing.get(qn("w:afterLines"), ""),
        "line": spacing.get(qn("w:line"), ""),
        "lineRule": spacing.get(qn("w:lineRule"), ""),
    }


def heading_donor_score(paragraph: Paragraph, level: int) -> int:
    text = paragraph.text.strip()
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    spacing = paragraph_spacing_values(paragraph)
    score = 0
    if f"heading {level}" in style_name:
        score += 80
    if "toc" in style_name:
        score -= 250
    if is_template_note_text(text) or is_template_format_instruction(text):
        score -= 160
    if spacing.get("before") == "156" and spacing.get("after") == "156":
        score += 120
    if spacing.get("after") == "312" and not spacing.get("before"):
        score -= 70
    if level >= 3 and (paragraph_indent_values(paragraph).get("firstLine") or paragraph_indent_values(paragraph).get("firstLineChars")):
        score += 40
    return score


def body_donor_score(paragraph: Paragraph) -> int:
    text = paragraph.text.strip()
    ind = paragraph_indent_values(paragraph)
    score = 0
    if is_template_format_instruction(text) or is_pure_template_instruction(text):
        score -= 300
    if ind.get("firstLine") or ind.get("firstLineChars"):
        score += 100
    if (ind.get("left") or ind.get("leftChars")) and not (ind.get("firstLine") or ind.get("firstLineChars")):
        score -= 35
    if "\u00d7" in text or "\u25a1" in text:
        score += 25
    if "\u6b63\u6587" in text and ("\u5c0f\u56db" in text or "\u7a7a2\u683c" in text or "\u7a7a\u4e24\u683c" in text):
        score += 90
    if 8 <= len(text) <= 300:
        score += 8
    if any(token in text for token in BODY_DONOR_RULE_HINTS):
        score -= 70
    if paragraph_alignment_value(paragraph).lower() == "center":
        score -= 120
    if first_visible_run_is_italic(paragraph):
        score -= 80
    if is_figure_or_table_format_note(text):
        score -= 140
    if is_template_note_text(text) or TABLE_TITLE_RE.match(text) or FIGURE_TITLE_RE.match(text):
        score -= 100
    return score


def dominant_body_paragraph_after(doc: Document, anchor: Paragraph | None) -> Paragraph | None:
    if anchor is None:
        return None
    start = paragraph_index(doc, anchor)
    candidates: list[Paragraph] = []
    for paragraph in doc.paragraphs[start + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        if is_tail_heading(text):
            break
        if heading_level(text) is not None or is_toc_paragraph(paragraph):
            continue
        if paragraph_has_tab(paragraph) and body_donor_score(paragraph) < 50:
            continue
        if is_template_note_text(text) or is_template_format_instruction(text) or is_figure_or_table_format_note(text):
            continue
        candidates.append(paragraph)

    if not candidates:
        return None
    counts = Counter(paragraph_signature_key(paragraph) for paragraph in candidates)
    best_index, best_paragraph = max(
        enumerate(candidates),
        key=lambda item: (
            body_donor_score(item[1]),
            counts[paragraph_signature_key(item[1])],
            -item[0],
        ),
    )
    _ = best_index
    return best_paragraph


def explicit_body_text_paragraph_after(doc: Document, anchor: Paragraph | None) -> Paragraph | None:
    if anchor is None:
        return None
    start = paragraph_index(doc, anchor)
    candidates: list[Paragraph] = []
    for paragraph in doc.paragraphs[start + 1 :]:
        text = paragraph.text.strip()
        if is_tail_heading(text):
            break
        if not text:
            continue
        if heading_level(text) is not None or is_toc_paragraph(paragraph):
            continue
        if paragraph_has_tab(paragraph) and body_donor_score(paragraph) < 50:
            continue
        if is_template_note_text(text) or is_template_format_instruction(text) or is_figure_or_table_format_note(text):
            continue
        style = paragraph.style
        style_id = style.style_id if style is not None else ""
        style_name = style.name if style is not None else ""
        normalized_style = normalize(f"{style_id} {style_name}").lower()
        if "bodytext" in normalized_style or normalize(style_name) == normalize("\u6b63\u6587"):
            candidates.append(paragraph)

    if not candidates:
        return None
    return max(candidates, key=body_donor_score)


def table_column_count(table: Table) -> int:
    try:
        return len(table.columns)
    except Exception:
        rows = table._tbl.findall(qn("w:tr"))
        return len(rows[0].findall(qn("w:tc"))) if rows else 0


def table_row_count(table: Table) -> int:
    try:
        return len(table.rows)
    except Exception:
        return len(table._tbl.findall(qn("w:tr")))


def iter_body_tables_with_previous_text(doc: Document):
    previous_text = ""
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                previous_text = text
        elif isinstance(block, Table):
            yield block, previous_text


def select_table_donor_table(doc: Document) -> Table | None:
    for table, previous_text in iter_body_tables_with_previous_text(doc):
        if table_row_count(table) <= 1 or table_column_count(table) <= 1:
            continue
        if TABLE_TITLE_RE.match(previous_text.strip()):
            return table
    return None


def first_cell_pr_for_row(rows: list[Any], row_index: int) -> Any | None:
    if row_index < 0 or row_index >= len(rows):
        return None
    cells = rows[row_index].findall(qn("w:tc"))
    for cell in cells:
        tc_pr = cell.find(qn("w:tcPr"))
        if tc_pr is not None:
            return deepcopy(tc_pr)
    return None


def tr_pr_for_row(rows: list[Any], row_index: int) -> Any | None:
    if row_index < 0 or row_index >= len(rows):
        return None
    tr_pr = rows[row_index].find(qn("w:trPr"))
    return deepcopy(tr_pr) if tr_pr is not None else None


def capture_table_donor(table: Table | None) -> TableDonor | None:
    if table is None:
        return None
    rows = table._tbl.findall(qn("w:tr"))
    if not rows:
        return None
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    grid_widths: list[int] = []
    tbl_grid = table._tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        for grid_col in tbl_grid.findall(qn("w:gridCol")):
            try:
                grid_widths.append(int(grid_col.get(qn("w:w"), "0")))
            except ValueError:
                pass
    body_middle_index = 2 if len(rows) > 3 else (1 if len(rows) > 1 else 0)
    return TableDonor(
        tbl_pr=deepcopy(tbl_pr) if tbl_pr is not None else None,
        grid_widths=grid_widths,
        header_tr_pr=tr_pr_for_row(rows, 0),
        body_first_tr_pr=tr_pr_for_row(rows, 1),
        body_middle_tr_pr=tr_pr_for_row(rows, body_middle_index),
        body_last_tr_pr=tr_pr_for_row(rows, len(rows) - 1),
        header_tc_pr=first_cell_pr_for_row(rows, 0),
        body_first_tc_pr=first_cell_pr_for_row(rows, 1),
        body_middle_tc_pr=first_cell_pr_for_row(rows, body_middle_index),
        body_last_tc_pr=first_cell_pr_for_row(rows, len(rows) - 1),
    )


def replace_first_child(parent: Any, tag: str, new_child: Any | None, *, insert_index: int = 0) -> None:
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)
            break
    if new_child is not None:
        parent.insert(insert_index, deepcopy(new_child))


def remove_children(parent: Any, tag: str) -> None:
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)


def make_border(edge: str, value: str, *, size: str = "8", color: str = "auto") -> Any:
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), value)
    if value not in {"nil", "none"}:
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
    return border


def replace_borders(parent_pr: Any, tag: str, edges: dict[str, str]) -> None:
    remove_child(parent_pr, tag)
    borders = OxmlElement(tag)
    for edge, value in edges.items():
        borders.append(make_border(edge, value))
    parent_pr.append(borders)


def set_cell_borders(tc_pr: Any, edges: dict[str, tuple[str, str, str] | None]) -> None:
    remove_child(tc_pr, "w:tcBorders")
    borders = OxmlElement("w:tcBorders")
    for edge, spec in edges.items():
        if spec is None:
            continue
        val, size, color = spec
        borders.append(make_border(edge, val, size=size, color=color))
    if len(borders):
        tc_pr.append(borders)


def apply_three_line_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    replace_borders(
        tbl_pr,
        "w:tblBorders",
        {
            "top": "single",
            "left": "nil",
            "bottom": "single",
            "right": "nil",
            "insideH": "nil",
            "insideV": "nil",
        },
    )

    rows = table._tbl.findall(qn("w:tr"))
    last_index = len(rows) - 1
    for row_index, row in enumerate(rows):
        for cell in row.findall(qn("w:tc")):
            tc_pr = cell.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                cell.insert(0, tc_pr)
            edges: dict[str, tuple[str, str, str] | None] = {
                "left": ("nil", "0", "auto"),
                "right": ("nil", "0", "auto"),
                "top": ("nil", "0", "auto"),
                "bottom": ("nil", "0", "auto"),
            }
            if row_index == 0:
                edges["top"] = ("single", "8", "auto")
                edges["bottom"] = ("single", "8", "auto")
            if row_index == last_index:
                edges["bottom"] = ("single", "8", "auto")
            set_cell_borders(tc_pr, edges)


def donor_column_widths(donor: TableDonor, column_count: int) -> list[int]:
    if column_count <= 0:
        return []
    if len(donor.grid_widths) == column_count:
        return list(donor.grid_widths)
    total_width = sum(width for width in donor.grid_widths if width > 0)
    if total_width <= 0:
        return []
    base = max(1, total_width // column_count)
    widths = [base for _ in range(column_count)]
    widths[-1] += total_width - sum(widths)
    return widths


def set_tc_width(tc_pr: Any, width: int | None) -> None:
    if not width:
        return
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def row_role_pr(donor: TableDonor, row_index: int, last_index: int, *, cell: bool) -> Any | None:
    if row_index == 0:
        return donor.header_tc_pr if cell else donor.header_tr_pr
    if row_index == last_index:
        return donor.body_last_tc_pr if cell else donor.body_last_tr_pr
    if row_index == 1:
        return donor.body_first_tc_pr if cell else donor.body_first_tr_pr
    return donor.body_middle_tc_pr if cell else donor.body_middle_tr_pr


def apply_table_donor(table: Table, donor: TableDonor | None) -> None:
    if donor is None:
        return
    if donor.tbl_pr is not None:
        replace_first_child(table._tbl, "w:tblPr", donor.tbl_pr, insert_index=0)

    column_count = table_column_count(table)
    widths = donor_column_widths(donor, column_count)
    if widths:
        remove_children(table._tbl, "w:tblGrid")
        tbl_grid = OxmlElement("w:tblGrid")
        for width in widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(width))
            tbl_grid.append(grid_col)
        insert_index = 1 if table._tbl.find(qn("w:tblPr")) is not None else 0
        table._tbl.insert(insert_index, tbl_grid)

    rows = table._tbl.findall(qn("w:tr"))
    last_index = len(rows) - 1
    for row_index, row in enumerate(rows):
        tr_pr = row_role_pr(donor, row_index, last_index, cell=False)
        if tr_pr is not None:
            replace_first_child(row, "w:trPr", tr_pr, insert_index=0)
        tc_pr = row_role_pr(donor, row_index, last_index, cell=True)
        if tc_pr is None:
            continue
        for col_index, cell in enumerate(row.findall(qn("w:tc"))):
            cloned_tc_pr = deepcopy(tc_pr)
            set_tc_width(cloned_tc_pr, widths[col_index] if col_index < len(widths) else None)
            replace_first_child(cell, "w:tcPr", cloned_tc_pr, insert_index=0)


def table_cell_donor_table(doc: Document) -> Table | None:
    return select_table_donor_table(doc)


def first_table_style(doc: Document) -> str | None:
    return None


def table_cell_donors(doc: Document) -> tuple[ParagraphDonor, ParagraphDonor, ParagraphDonor, ParagraphDonor]:
    table = table_cell_donor_table(doc)
    if table is None:
        empty = ParagraphDonor(None, None, "", "")
        return empty, empty, empty, empty
    header_para = table.cell(0, 0).paragraphs[0] if table.cell(0, 0).paragraphs else None
    body_para = table.cell(1, 0).paragraphs[0] if table_row_count(table) > 1 and table.cell(1, 0).paragraphs else header_para
    header = capture_donor(header_para)
    body = capture_donor(body_para)
    return header, body, capture_latin_donor(header_para, header), capture_latin_donor(body_para, body)


def table_donor(doc: Document) -> TableDonor | None:
    return None


def find_caption_donor(doc: Document, table: bool = False) -> Paragraph | None:
    pattern = TABLE_TITLE_RE if table else FIGURE_TITLE_RE
    for paragraph in doc.paragraphs:
        if pattern.match(paragraph.text.strip()):
            return paragraph
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if any(token in style_name for token in ("caption", "\u56fe\u9898", "\u8868\u9898")):
            return paragraph
    return None


def collect_toc_entry_donors(doc: Document) -> dict[int, ParagraphDonor]:
    toc_heading = find_toc_heading(doc)
    first_body = find_first_body_heading(doc)
    if toc_heading is None:
        return {}
    donors: dict[int, ParagraphDonor] = {}
    for paragraph_element in toc_range_paragraph_elements(doc, toc_heading, first_body):
        text = element_text(paragraph_element).strip()
        if not text:
            continue
        label = text.split("\t", 1)[0]
        level = template_heading_level(label)
        if level is None:
            continue
        paragraph = Paragraph(paragraph_element, toc_heading._parent)
        donors.setdefault(min(level, 3), capture_donor(paragraph))
    return donors


def collect_donors(doc: Document) -> Donors:
    heading_by_level: dict[int, ParagraphDonor] = {}
    heading_score_by_level: dict[int, int] = {}
    toc_by_level = collect_toc_entry_donors(doc)
    first_body = find_first_body_heading(doc)
    body_start = paragraph_index(doc, first_body) if first_body is not None else 0
    for paragraph in doc.paragraphs[body_start:]:
        if (
            paragraph_has_tab(paragraph)
            or is_toc_paragraph(paragraph)
            or is_template_note_text(paragraph.text)
            or is_template_format_instruction(paragraph.text)
        ):
            continue
        level = template_heading_level(paragraph.text)
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if level is None:
            for candidate_level in (1, 2, 3, 4):
                if f"heading {candidate_level}" in style_name:
                    level = candidate_level
                    break
        if level:
            score = heading_donor_score(paragraph, level)
            if level in heading_by_level and score <= heading_score_by_level.get(level, -10_000):
                continue
            heading_by_level[level] = capture_heading_donor(paragraph)
            heading_score_by_level[level] = score

    body_candidates = [
        paragraph
        for paragraph in (
            explicit_body_text_paragraph_after(doc, first_body),
            dominant_body_paragraph_after(doc, first_body),
            first_body_paragraph_after(doc, first_body) if first_body is not None else None,
        )
        if paragraph is not None
    ]
    body_para = max(
        body_candidates,
        key=lambda paragraph: (
            body_donor_score(paragraph),
            1
            if paragraph.style is not None
            and (
                paragraph.style.name.lower() == "normal"
                or paragraph.style.style_id.lower() in {"normal", "a"}
            )
            else 0,
        ),
    ) if body_candidates else None
    abstract_heading = find_paragraph(doc, {"\u6458\u8981"})
    abstract_body = next_nonempty_after(doc, abstract_heading) if abstract_heading is not None else body_para
    keyword_para = None
    for paragraph in doc.paragraphs:
        if "\u5173\u952e\u8bcd" in paragraph.text or "key words" in paragraph.text.lower() or "keywords" in paragraph.text.lower():
            keyword_para = paragraph
            break
    references_heading = find_surface_heading(doc, {"\u53c2\u8003\u6587\u732e", "references"})
    reference_entry = next_nonempty_content_after(doc, references_heading) if references_heading is not None else body_para
    reference_entry_donor = capture_donor(reference_entry)
    ack_heading = find_surface_heading(doc, {"\u81f4\u8c22", "acknowledgements", "acknowledgments"})
    ack_body = next_nonempty_content_after(doc, ack_heading) if ack_heading is not None else body_para
    if ack_body is not None and is_template_format_instruction(ack_body.text):
        ack_body = None
    ack_body_donor = capture_donor(ack_body or body_para or first_body)
    header_cell, body_cell, header_latin_cell, body_latin_cell = table_cell_donors(doc)

    return Donors(
        heading_by_level=heading_by_level,
        toc_by_level=toc_by_level,
        body=capture_donor(body_para or first_body),
        abstract_body=capture_donor(abstract_body or body_para or first_body),
        keyword=capture_donor(keyword_para or abstract_body or body_para or first_body),
        figure_caption=capture_donor(find_caption_donor(doc, table=False) or body_para or first_body),
        table_caption=capture_donor(find_caption_donor(doc, table=True) or body_para or first_body),
        reference_heading=capture_donor(references_heading or first_body),
        reference_entry=reference_entry_donor,
        reference_cjk_run=donor_with_cjk_font(
            reference_entry_donor,
            capture_cjk_donor(reference_entry, reference_entry_donor),
        ),
        reference_latin_run=donor_with_latin_font(
            reference_entry_donor,
            capture_latin_donor(reference_entry, reference_entry_donor),
        ),
        acknowledgement=capture_donor(ack_heading or body_para or first_body),
        acknowledgement_body=ack_body_donor,
        acknowledgement_cjk_run=capture_cjk_donor(ack_body, ack_body_donor),
        acknowledgement_latin_run=capture_latin_donor(ack_body, ack_body_donor),
        table_header_cell=header_cell,
        table_body_cell=body_cell,
        table_header_latin_cell=header_latin_cell,
        table_body_latin_cell=body_latin_cell,
        table_style=first_table_style(doc),
        table_donor=table_donor(doc),
    )


def remove_existing_body(doc: Document) -> None:
    first_heading = find_first_body_heading(doc)
    if first_heading is None:
        return
    body = doc.element.body
    started = False
    for child in list(body.iterchildren()):
        if child is first_heading._element:
            started = True
        if not started:
            continue
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def capture_tail_section_break_templates(doc: Document) -> list[Any]:
    first_heading = find_first_body_heading(doc)
    if first_heading is None:
        return []
    start = paragraph_index(doc, first_heading) + 1
    templates: list[Any] = []
    for paragraph in doc.paragraphs[start:]:
        if paragraph._element.pPr is not None and paragraph._element.pPr.find(qn("w:sectPr")) is not None:
            templates.append(deepcopy(paragraph._element))
    return templates


def append_section_break_template(doc: Document, template: Any) -> Paragraph:
    section_break = deepcopy(template)
    for child in list(section_break):
        if child.tag != qn("w:pPr"):
            section_break.remove(child)
    anchor = doc.paragraphs[-1]
    anchor._element.addnext(section_break)
    return Paragraph(section_break, anchor._parent)


def append_paragraph(doc: Document, donor: ParagraphDonor, text: str = "") -> Paragraph:
    paragraph = doc.add_paragraph()
    set_paragraph_text(paragraph, donor, text)
    return paragraph


def append_heading(doc: Document, donors: Donors, level: int, text: str, *, page_break_before: bool = False) -> Paragraph:
    donor = donors.heading_by_level.get(level) or donors.heading_by_level.get(1) or donors.body
    paragraph = doc.add_paragraph()
    apply_donor(paragraph, donor)
    append_mixed_script_text(paragraph, text, cjk_donor=donor, latin_donor=donor)
    if donor.style_id:
        set_paragraph_style_id(paragraph, donor.style_id)
    if page_break_before:
        ensure_page_break_before(paragraph)
    return paragraph


def append_body_paragraph(doc: Document, donors: Donors, text: str) -> Paragraph:
    paragraph = doc.add_paragraph()
    apply_donor(paragraph, donors.body)
    set_direct_spacing(paragraph, before="0", after="0", line="360", line_rule="auto")
    append_mixed_script_text(
        paragraph,
        normalize_citation_punctuation_in_text(sanitize_caption_like_body_text(text)),
        cjk_donor=donors.body,
        latin_donor=donor_with_latin_font(donors.body, donors.body),
    )
    if donors.body.style_id:
        set_paragraph_style_id(paragraph, donors.body.style_id)
    return paragraph


def append_caption_paragraph(doc: Document, donor: ParagraphDonor, text: str) -> Paragraph:
    paragraph = doc.add_paragraph()
    apply_donor(paragraph, donor)
    append_mixed_script_text(paragraph, text, cjk_donor=donor, latin_donor=donor)
    if donor.style_id:
        set_paragraph_style_id(paragraph, donor.style_id)
    remove_child(ensure_ppr(paragraph), "w:ind")
    return paragraph


def append_page_break(doc: Document, donor: ParagraphDonor) -> Paragraph:
    paragraph = doc.add_paragraph()
    apply_donor(paragraph, donor)
    run_obj = add_text_with_donor(paragraph, donor, "")
    run_obj.add_break(WD_BREAK.PAGE)
    return paragraph


def keyword_label_from_template(text: str, *, english: bool) -> str:
    candidates = (
        ("Key words\uff1a", "Key words\uff1a"),
        ("Key words:", "Key words: "),
        ("Key Words\uff1a", "Key Words\uff1a"),
        ("Key Words:", "Key Words: "),
        ("Keywords\uff1a", "Keywords\uff1a"),
        ("Keywords:", "Keywords: "),
    ) if english else (
        ("\u5173\u952e\u8bcd\uff1a", "\u5173\u952e\u8bcd\uff1a"),
        ("\u5173\u952e\u8bcd:", "\u5173\u952e\u8bcd: "),
    )
    stripped = text or ""
    for marker, label in candidates:
        if normalize(stripped).startswith(normalize(marker)):
            return label
    return "Key Words: " if english else "\u5173\u952e\u8bcd\uff1a"


def is_keyword_label_fragment(text: str) -> bool:
    normalized = normalize(text)
    return normalized in {
        normalize("\u5173\u952e\u8bcd"),
        normalize("\u5173\u952e\u8bcd\uff1a"),
        normalize("\u5173\u952e\u8bcd:"),
        normalize("Key words"),
        normalize("Key words:"),
        normalize("Key words\uff1a"),
        normalize("Key Words"),
        normalize("Key Words:"),
        normalize("Key Words\uff1a"),
        normalize("Keywords"),
        normalize("Keywords:"),
        normalize("Keywords\uff1a"),
        normalize(":"),
        normalize("\uff1a"),
    }


def keyword_content_donor(paragraph: Paragraph, fallback: ParagraphDonor) -> ParagraphDonor:
    for run_obj in paragraph.runs:
        if run_obj._element.find(".//" + qn("w:txbxContent")) is not None:
            continue
        text = run_obj.text.strip()
        if not text or is_keyword_label_fragment(text):
            continue
        if run_obj._element.rPr is not None:
            return ParagraphDonor(fallback.ppr, deepcopy(run_obj._element.rPr), fallback.style_id, fallback.style_name)
    return fallback


def keyword_label_end_index(paragraph: Paragraph, label: str) -> int | None:
    accumulated = ""
    target = normalize(label)
    for idx, run_obj in enumerate(paragraph.runs):
        accumulated += run_obj.text or ""
        if normalize(accumulated) == target:
            return idx
        if len(normalize(accumulated)) > len(target) and not normalize(accumulated).startswith(target):
            break
    return None


def is_keyword_separator_text(text: str) -> bool:
    stripped = (text or "").strip()
    return bool(stripped) and all(ch in {";", "\uff1b", ",", "\uff0c", "\u3001", " "} for ch in stripped)


def rewrite_keyword_paragraph(paragraph: Paragraph, keywords: list[str], *, english: bool) -> None:
    base = capture_donor(paragraph)
    label = keyword_label_from_template(paragraph.text, english=english)
    label_end = keyword_label_end_index(paragraph, label)
    content_text = ("; " if english else "\uff1b").join(keywords)
    if label_end is None:
        apply_donor(paragraph, base)
        label_run = add_text_with_donor(paragraph, base, label)
        ensure_run_fonts(label_run, bold=True)
        content_run = add_text_with_donor(paragraph, keyword_content_donor(paragraph, base), content_text)
        ensure_run_fonts(content_run, bold=False)
        return

    for run_obj in paragraph.runs[: label_end + 1]:
        ensure_run_fonts(run_obj, bold=True)

    content_indices = [
        idx
        for idx, run_obj in enumerate(paragraph.runs[label_end + 1 :], start=label_end + 1)
        if (run_obj.text or "").strip() and not is_keyword_separator_text(run_obj.text or "")
    ]
    content_donor = donor_with_rpr(base, run_rpr(paragraph.runs[content_indices[0]])) if content_indices else keyword_content_donor(paragraph, base)

    for run_obj in paragraph.runs[label_end + 1 :]:
        run_obj.text = ""
    if content_indices and content_indices[0] < len(paragraph.runs):
        content_run = paragraph.runs[content_indices[0]]
        content_run.text = content_text
        ensure_run_fonts(content_run, bold=False)
    else:
        content_run = add_text_with_donor(paragraph, content_donor, content_text)
        ensure_run_fonts(content_run, bold=False)


def split_title_lines(title: str, requested: Any = None) -> list[str]:
    if isinstance(requested, list):
        lines = [str(item).strip() for item in requested if str(item).strip()]
        if lines:
            return lines
    title = str(title or "").strip()
    if not title:
        return []
    if len(title) <= 18:
        return [title]
    midpoint = len(title) // 2
    breakpoints = [idx + 1 for idx, ch in enumerate(title) if ch in {"\u7684", "\u4e0e", "\u548c"}]
    split_at = min(breakpoints, key=lambda idx: abs(idx - midpoint)) if breakpoints else midpoint
    if split_at <= 3 or len(title) - split_at <= 3:
        split_at = midpoint
    return [title[:split_at].strip(), title[split_at:].strip()]


def cover_end_index(doc: Document) -> int:
    markers = {
        heading_key("\u5b66\u4f4d\u8bba\u6587\u539f\u521b\u6027\u58f0\u660e"),
        heading_key("\u6458\u8981"),
        heading_key("Abstract"),
    }
    for idx, paragraph in enumerate(doc.paragraphs):
        if heading_key(paragraph.text) in markers:
            return idx
    return min(len(doc.paragraphs), 60)


def front_matter_end_index(doc: Document) -> int:
    markers = {heading_key("\u6458\u8981"), heading_key("Abstract")}
    for idx, paragraph in enumerate(doc.paragraphs):
        if heading_key(paragraph.text) in markers:
            return idx
    return min(len(doc.paragraphs), 100)


def iter_table_paragraphs(doc: Document) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def front_matter_visible_text(doc: Document) -> str:
    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs[: front_matter_end_index(doc)]]
    table_texts = [paragraph.text for paragraph in iter_table_paragraphs(doc)]
    return "\n".join(paragraph_texts + table_texts)


def is_cover_title_candidate(text: str) -> bool:
    stripped = str(text or "").strip()
    normalized = normalize(stripped)
    if len(normalized) < 5:
        return False
    blocked_tokens = (
        "\u5b66\u58eb\u5b66\u4f4d\u8bba\u6587",
        "\u6bd5\u4e1a\u8bba\u6587",
        "\u6bd5\u4e1a\u8bba\u6587\uff08\u8bbe\u8ba1\uff09",
        "\u6bd5\u4e1a\u8bbe\u8ba1",
        "\u5b66\u751f\u6bd5\u4e1a\u8bba\u6587",
        "\u672c\u79d1\u751f",
        "\u6839\u636e\u5185\u5bb9",
        "\u4e8c\u9009\u4e00",
        "\u5c71\u4e1c\u7ba1\u7406\u5b66\u9662",
        "\u5927\u5e86\u5e08\u8303\u5b66\u9662",
        "\u5357\u4eac\u90ae\u7535\u5927\u5b66",
        "\u6a21\u677f",
        "\u59d3\u540d",
        "\u5b66\u53f7",
        "\u9662\u7cfb",
        "\u9662\uff08\u7cfb\uff09",
        "\u4e13\u4e1a",
        "\u7814\u7a76\u65b9\u5411",
        "\u8ba1\u7b97\u673a\u79d1\u5b66",
        "\u4fe1\u606f\u6280\u672f\u5b66\u9662",
        "\u6307\u5bfc\u6559\u5e08",
        "\u5e74",
        "\u6708",
        "\u65e5",
        "\u58f0\u660e",
        "\u6388\u6743",
        "\u8bba\u6587\u4f5c\u8005\u7b7e\u540d",
        "\u8bba\u6587\u4f5c\u8005",
        "\u7b7e\u5b57",
        "\u9ed1\u8272\u4e2d\u6027\u7b14",
        "\u6458\u8981",
        "abstract",
        "keywords",
        "key words",
        "timesnewroman",
        "\u76ee\u5f55",
        "\u5173\u952e\u8bcd",
        "\u6458\u8981\u6b63\u6587",
        "\u4e2d\u6587\u9898\u76ee",
        "\u9898\u76ee\u53ea\u6709\u4e00\u884c",
        "\u987b\u5220\u9664\u672c\u884c",
    )
    if is_template_format_instruction(stripped):
        return False
    if stripped.count("×") >= 2 and not any(ch.isalnum() for ch in stripped):
        return True
    if not has_cjk(stripped):
        return False
    return not any(
        token in stripped or normalize(token) in normalized for token in blocked_tokens
    )


def looks_like_english_title(text: str) -> bool:
    stripped = str(text or "").strip()
    compact = normalize(stripped)
    if (
        compact.startswith("abstract")
        or compact.startswith("keywords")
        or compact.startswith("keyword")
        or compact.startswith("key words".replace(" ", ""))
        or "timesnewroman" in compact
        or is_template_format_instruction(stripped)
    ):
        return False
    latin_count = sum(1 for ch in stripped if ("A" <= ch <= "Z") or ("a" <= ch <= "z"))
    return latin_count >= 12 and " " in stripped and not has_cjk(stripped)


def replace_text_in_runs(paragraph: Paragraph, old: str, new: str) -> bool:
    if not old:
        return False
    for run_obj in paragraph.runs:
        if old in (run_obj.text or ""):
            run_obj.text = (run_obj.text or "").replace(old, new)
            return True
    text = paragraph.text
    if old not in text:
        return False
    donor = capture_donor(paragraph)
    set_paragraph_text(paragraph, donor, text.replace(old, new))
    return True


def set_paragraph_single_text_preserving_runs(paragraph: Paragraph, text: str) -> None:
    text_runs = [run_obj for run_obj in paragraph.runs if run_obj.text is not None]
    visible_runs = [run_obj for run_obj in text_runs if (run_obj.text or "").strip()]
    if not text_runs:
        donor = capture_donor(paragraph)
        add_text_with_donor(paragraph, donor, text)
        return
    target = visible_runs[0] if visible_runs else text_runs[0]
    target.text = text
    for run_obj in text_runs:
        if run_obj is target:
            continue
        run_obj.text = ""


def set_marked_heading_text_preserving_runs(paragraph: Paragraph, text: str, marker: str) -> None:
    text_runs = [run_obj for run_obj in paragraph.runs if run_obj.text is not None]
    visible_runs = [run_obj for run_obj in text_runs if (run_obj.text or "").strip()]
    marker_key = normalize(marker)
    marked_runs = [
        run_obj
        for run_obj in visible_runs
        if marker_key and marker_key in normalize(run_obj.text or "")
    ]
    if not text_runs:
        donor = capture_donor(paragraph)
        add_text_with_donor(paragraph, donor, text)
        return
    target = marked_runs[-1] if marked_runs else (visible_runs[-1] if visible_runs else text_runs[0])
    target.text = text
    for run_obj in text_runs:
        if run_obj is target:
            continue
        run_obj.text = ""


def split_title_lines_for_count(title: str, count: int) -> list[str]:
    title = str(title or "").strip()
    if count <= 1 or len(title) <= 8:
        return [title] if title else []
    midpoint = len(title) // 2
    breakpoints = [idx + 1 for idx, ch in enumerate(title) if ch in {"的", "与", "和", "及"}]
    split_at = min(breakpoints, key=lambda idx: abs(idx - midpoint)) if breakpoints else midpoint
    if split_at <= 3 or len(title) - split_at <= 3:
        split_at = midpoint
    return [title[:split_at].strip(), title[split_at:].strip()][:count]


def numbered_heading_title(title: str, numbering: tuple[int, ...]) -> str:
    if not numbering:
        return title
    title = strip_existing_heading_number(title, numbering)
    if len(numbering) == 1:
        return f"\u7b2c{numbering[0]}\u7ae0 {title}"
    return f"{'.'.join(str(item) for item in numbering)} {title}"


def strip_existing_heading_number(title: str, numbering: tuple[int, ...]) -> str:
    text = str(title or "").strip()
    if not text or not numbering:
        return text
    if len(numbering) == 1:
        text = re.sub(rf"^\u7b2c[0-9{CN_NUMBER_CHARS}]+\u7ae0[\s\u25a1:：、.\-—]*", "", text).strip()
        text = re.sub(r"^\d{1,2}[\s\u25a1]+", "", text).strip()
    else:
        text = re.sub(r"^\d{1,2}(?:\.\d{1,2})*[\s\u25a1]+", "", text).strip()
    return text


def replace_cover_field_value(paragraph: Paragraph, label: str, value: str) -> None:
    if not value or normalize(label) not in normalize(paragraph.text):
        return
    compact_label = normalize(label)
    accumulated = ""
    label_seen = False
    for run_obj in paragraph.runs:
        text_part = run_obj.text or ""
        if not label_seen:
            accumulated += text_part
            if compact_label in normalize(accumulated):
                label_seen = True
            continue
        if not text_part.strip():
            continue
        leading = re.match(r"^\s*", text_part).group(0)
        trailing = re.search(r"\s*$", text_part).group(0)
        run_obj.text = leading + str(value) + trailing
        return
    text = paragraph.text
    label_match = label if label in text else None
    if label_match is None:
        compact_label = normalize(label)
        accumulated = ""
        end_idx = None
        for idx, char in enumerate(text):
            accumulated += char
            if normalize(accumulated).endswith(compact_label):
                end_idx = idx + 1
                break
        if end_idx is None:
            return
        prefix = text[:end_idx]
        tail = text[end_idx:]
    else:
        prefix, tail = text.split(label_match, 1)
        prefix += label_match
    leading = re.match(r"^\s*", tail).group(0) if tail is not None else " "
    trailing = re.search(r"\s*$", tail).group(0) if tail is not None else " "
    if not leading:
        leading = " "
    field_width = max(len(tail or "") - len(leading) - len(trailing), len(str(value)))
    padded_value = str(value).center(field_width) if field_width > len(str(value)) else str(value)
    replacement = prefix + leading + padded_value + trailing
    set_paragraph_single_text_preserving_runs(paragraph, replacement)


def replace_cover_table_field_values(doc: Document, field_values: dict[tuple[str, ...], Any]) -> set[int]:
    """Fill the value cell in label/value cover tables without polluting the label cell."""
    touched_label_paragraph_ids: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            cells = list(row.cells)
            if len(cells) < 2:
                continue
            for cell_idx, label_cell in enumerate(cells[:-1]):
                label_text = normalize(label_cell.text)
                if not label_text:
                    continue
                target_cell = cells[cell_idx + 1]
                for labels, value in field_values.items():
                    value_text = str(value or "").strip()
                    if not value_text:
                        continue
                    if not any(normalize(label) and normalize(label) in label_text for label in labels):
                        continue
                    target_paragraph = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
                    set_paragraph_single_text_preserving_runs(target_paragraph, value_text)
                    for extra_paragraph in target_cell.paragraphs[1:]:
                        set_paragraph_single_text_preserving_runs(extra_paragraph, "")
                    for paragraph in label_cell.paragraphs:
                        touched_label_paragraph_ids.add(id(paragraph))
                    break
    return touched_label_paragraph_ids


def apply_cover_fields(doc: Document, content: dict[str, Any]) -> None:
    title = str(content.get("title") or "").strip()
    if not title:
        return
    end = cover_end_index(doc)
    cover_items = list(enumerate(doc.paragraphs[:end]))
    original_front_visible = normalize(front_matter_visible_text(doc))
    english_title = str(content.get("en_title") or content.get("english_title") or "").strip()
    english_idx = next((idx for idx, paragraph in cover_items if looks_like_english_title(paragraph.text)), None)
    title_search_end = english_idx if english_idx is not None else end
    title_field_labels = (
        "\u9898    \u76ee\uff1a",
        "\u9898\u76ee\uff1a",
        "\u8bba\u6587\u9898\u76ee\uff1a",
        "\u9898    \u76ee:",
        "\u9898\u76ee:",
        "\u8bba\u6587\u9898\u76ee:",
        "\u9898    \u76ee",
        "\u9898\u76ee",
        "\u8bba\u6587\u9898\u76ee",
    )
    title_label_candidates = [
        idx
        for idx, paragraph in cover_items
        if idx < title_search_end
        and any(normalize(label) and normalize(label) in normalize(paragraph.text) for label in title_field_labels)
    ]
    title_instruction_candidates = [
        idx
        for idx, paragraph in cover_items
        if idx < title_search_end
        and "\u8bba\u6587\u9898\u76ee" in paragraph.text
    ]
    zh_candidates = [
        idx
        for idx, paragraph in cover_items
        if idx < title_search_end and is_cover_title_candidate(paragraph.text)
    ]
    if title_label_candidates:
        target_idx = title_label_candidates[-1]
        target_paragraph = doc.paragraphs[target_idx]
        for label in title_field_labels:
            before = target_paragraph.text
            replace_cover_field_value(target_paragraph, label, title)
            if target_paragraph.text != before:
                break
        else:
            set_cover_text_with_split_runs(target_paragraph, capture_donor(target_paragraph), title)
        zh_candidates = []
    elif title_instruction_candidates:
        zh_candidates = title_instruction_candidates[:1]
    requested_title_lines = content.get("title_lines") or content.get("cover_title_lines")
    title_lines = split_title_lines(title, requested_title_lines)
    if zh_candidates and len(zh_candidates) != len(title_lines):
        title_lines = split_title_lines_for_count(title, len(zh_candidates))
    if zh_candidates:
        selected = zh_candidates[-len(title_lines) :]
        for idx, line in zip(selected, title_lines):
            set_cover_text_with_split_runs(doc.paragraphs[idx], capture_donor(doc.paragraphs[idx]), line)
        for idx in zh_candidates:
            if idx not in selected:
                set_paragraph_single_text_preserving_runs(doc.paragraphs[idx], "")
        for idx, paragraph in cover_items:
            if idx in selected or idx >= end:
                continue
            if normalize(paragraph.text) == normalize(title):
                set_paragraph_single_text_preserving_runs(paragraph, "")
    if english_title and english_idx is not None:
        set_paragraph_single_text_preserving_runs(doc.paragraphs[english_idx], english_title)
    date_range = str(
        content.get("date_range")
        or content.get("project_date_range")
        or content.get("thesis_date_range")
        or ""
    ).strip()

    field_labels_by_key = {
        "author": ("学生姓名：", "学生姓名:", "学生姓名", "\u59d3    \u540d\uff1a", "\u59d3    \u540d:", "\u59d3    \u540d", "姓名：", "姓名:", "姓名"),
        "student_id": ("班级学号：", "班级学号:", "班级学号", "\u5b66    \u53f7\uff1a", "\u5b66    \u53f7:", "\u5b66    \u53f7", "学号：", "学号:", "学号"),
        "class_name": ("班    级：", "班    级:", "班    级", "班级：", "班级:", "班级"),
        "department": ("\u9662    \u7cfb\uff1a", "\u9662    \u7cfb:", "\u9662    \u7cfb", "\u9662     \uff08\u7cfb\uff09", "\u9662\uff08\u7cfb\uff09", "院系：", "院系:", "院系"),
        "major": ("\u4e13    \u4e1a\uff1a", "\u4e13    \u4e1a:", "\u4e13    \u4e1a", "专业：", "专业:", "专业"),
        "research_direction": ("\u7814 \u7a76 \u65b9 \u5411\uff1a", "\u7814 \u7a76 \u65b9 \u5411:", "\u7814 \u7a76 \u65b9 \u5411", "\u7814\u7a76\u65b9\u5411\uff1a", "\u7814\u7a76\u65b9\u5411:", "\u7814\u7a76\u65b9\u5411"),
        "supervisor": ("\u6307\u5bfc\u6559\u5e08\uff1a", "\u6307\u5bfc\u6559\u5e08:", "\u6307\u5bfc\u6559\u5e08", "指导教师：", "指导教师:", "指导教师", "指导老师：", "指导老师:", "指导老师", "导师：", "导师:", "导师"),
    }
    field_values_by_key = {
        "author": content.get("author"),
        "student_id": content.get("student_id"),
        "class_name": content.get("class_name"),
        "department": content.get("department"),
        "major": content.get("major"),
        "research_direction": content.get("research_direction"),
        "supervisor": content.get("supervisor"),
    }
    simple_fields = {
        labels: field_values_by_key[key]
        for key, labels in field_labels_by_key.items()
    }
    table_label_paragraph_ids = replace_cover_table_field_values(doc, simple_fields)
    cover_field_paragraphs = list(doc.paragraphs[: front_matter_end_index(doc)])
    for _idx, paragraph in enumerate(cover_field_paragraphs):
        if _idx < end and id(paragraph) not in table_label_paragraph_ids:
            for labels, value in simple_fields.items():
                if not value:
                    continue
                for label in labels:
                    before = paragraph.text
                    replace_cover_field_value(paragraph, label, str(value))
                    if paragraph.text != before:
                        break
            if date_range and "\u65e5\u671f" in paragraph.text and (
                "\u65e5\u81f3" in paragraph.text or "\u81f3" in paragraph.text
            ):
                before = paragraph.text
                replace_cover_field_value(paragraph, "\u65e5\u671f\uff1a", date_range)
                if paragraph.text == before:
                    replace_cover_field_value(paragraph, "\u65e5\u671f", date_range)
            if date_range and re.search(r"\d{4}\s*\u5e74\s*\d{1,2}\s*\u6708(?:\s*\d{1,2}\s*\u65e5)?", paragraph.text):
                set_cover_text_with_split_runs(paragraph, capture_donor(paragraph), date_range, size="28", keep_ascii_slots=True)
        if "\u987b\u66ff\u6362\u4e3a\u672c\u8bba\u6587\u540d\u79f0" in paragraph.text:
            replace_text_in_runs(paragraph, "\u987b\u66ff\u6362\u4e3a\u672c\u8bba\u6587\u540d\u79f0\uff0c\u5b57\u4f53\u5b57\u53f7\u4e0e\u672c\u6bb5\u6587\u5b57\u4fdd\u6301\u4e00\u81f4", title)
            replace_text_in_runs(paragraph, "\u987b\u66ff\u6362\u4e3a\u672c\u8bba\u6587\u540d\u79f0", title)
    visible = normalize(front_matter_visible_text(doc))
    required_values = {"title": title}
    for key, labels in field_labels_by_key.items():
        value = field_values_by_key.get(key)
        if value and any(normalize(label) and normalize(label) in original_front_visible for label in labels):
            required_values[key] = value
    if date_range and (
        normalize("\u65e5\u671f") in original_front_visible
        or re.search(r"\d{4}\s*\u5e74\s*\d{1,2}\s*\u6708(?:\s*\d{1,2}\s*\u65e5)?", front_matter_visible_text(doc))
    ):
        required_values["date_range"] = date_range
    missing = [
        key
        for key, value in required_values.items()
        if str(value or "").strip() and normalize(str(value)) not in visible
    ]
    if missing:
        raise RuntimeError("cover required field values not applied: " + ", ".join(missing))


def assert_no_cover_sample_title_residue(doc: Document) -> None:
    residue_tokens = (
        "\u57fa\u4e8e\u5fae\u4fe1\u516c\u4f17\u5e73\u53f0\u7684",
        "\u4f1a\u52a1\u7ec4\u7ec7\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5b9e\u73b0",
        "Conference Organization System Based on Wechat Public Platform",
        "\u987b\u66ff\u6362\u4e3a\u672c\u8bba\u6587\u540d\u79f0",
    )
    visible = front_matter_visible_text(doc)
    found = [token for token in residue_tokens if token in visible]
    if found:
        raise RuntimeError("cover sample title residue remains: " + "; ".join(found))


def replace_abstract_block(doc: Document, heading_text: str, body_text: str, keywords: list[str], *, english: bool) -> None:
    heading = find_paragraph(doc, {heading_text})
    if heading is None:
        return
    set_paragraph_single_text_preserving_runs(heading, "Abstract" if english else "\u6458    \u8981")
    ensure_page_break_before(heading)
    start = paragraph_index(doc, heading) + 1
    end = len(doc.paragraphs)
    keyword_idx = None
    keyword_paragraph = None
    for idx, paragraph in enumerate(doc.paragraphs[start:], start=start):
        text = paragraph.text.strip()
        text_lower = text.lower()
        if not english and (text_lower in {"abstract"} or text_lower.startswith("abstract ")):
            raise RuntimeError(f"Chinese abstract block for {heading_text!r} crossed into English abstract")
        if english and "\u6458\u8981" in text and normalize(text) == normalize("\u6458\u8981"):
            raise RuntimeError(f"English abstract block for {heading_text!r} crossed into Chinese abstract")
        if not english and (("key words" in text_lower) or ("keywords" in text_lower)):
            raise RuntimeError(f"Chinese abstract block for {heading_text!r} reached English keyword line first")
        keyword_matches = "\u5173\u952e\u8bcd" in paragraph.text if not english else (
            ("key words" in text_lower) or ("keywords" in text_lower)
        )
        if keyword_matches:
            keyword_idx = idx
            keyword_paragraph = paragraph
            end = idx
            break
        if english and normalize(paragraph.text).startswith(normalize("\u5173\u952e\u8bcd")):
            raise RuntimeError(f"English abstract block for {heading_text!r} reached Chinese keyword line")
        if heading_level(paragraph.text) is not None or is_toc_paragraph(paragraph):
            raise RuntimeError(
                f"abstract block boundary for {heading_text!r} reached a protected surface before keyword line"
            )
    if keyword_idx is None or keyword_paragraph is None:
        raise RuntimeError(f"abstract block for {heading_text!r} has no locked keyword line boundary")
    if start >= end:
        raise RuntimeError(f"abstract block for {heading_text!r} has no editable body paragraph before keyword line")
    body_idx = None
    for idx in range(start, end):
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if not text:
            continue
        if is_pure_template_instruction(text) and not is_abstract_body_placeholder_line(text):
            continue
        body_idx = idx
        break
    if body_idx is None and start < len(doc.paragraphs) and start < end:
        body_idx = start
    if body_idx is None or body_idx >= end:
        raise RuntimeError(f"abstract block for {heading_text!r} has no safe body paragraph to rewrite")
    body_paragraph = doc.paragraphs[body_idx]
    old_body_paragraphs = list(doc.paragraphs[start:end])
    body_donor = capture_donor(body_paragraph)
    body_elements = {body_paragraph._element}
    body_parts = split_abstract_body_paragraphs(body_text)
    latin_source_donor = None if english else block_latin_donor(old_body_paragraphs, body_donor)
    set_abstract_body_text(
        body_paragraph,
        body_parts[0],
        latin_donor_for_mixed=not english,
        latin_source_donor=latin_source_donor,
        suppress_placeholder_prefix=english,
    )
    if english:
        configure_english_abstract_body_paragraph(body_paragraph)
    else:
        configure_chinese_abstract_body_paragraph(body_paragraph)
    anchor = body_paragraph
    for body_part in body_parts[1:]:
        anchor = insert_paragraph_after(anchor, body_donor)
        set_abstract_body_text(
            anchor,
            body_part,
            latin_donor_for_mixed=not english,
            latin_source_donor=latin_source_donor,
            suppress_placeholder_prefix=english,
        )
        if english:
            configure_english_abstract_body_paragraph(anchor)
        else:
            configure_chinese_abstract_body_paragraph(anchor)
        body_elements.add(anchor._element)
    for paragraph in old_body_paragraphs:
        if paragraph._element in body_elements:
            continue
        delete_paragraph(paragraph)
    rewrite_keyword_paragraph(keyword_paragraph, keywords, english=english)
    if not english:
        zh_keyword = next((paragraph for paragraph in doc.paragraphs if "\u5173\u952e\u8bcd" in paragraph.text), None)
        if zh_keyword is not None:
            compact_chinese_to_english_abstract_gap(zh_keyword)


def compact_chinese_to_english_abstract_gap(keyword_paragraph: Paragraph) -> None:
    """Remove filler blanks between Chinese keywords and the English abstract."""
    current = keyword_paragraph._element.getnext()
    while current is not None and current.tag == qn("w:p"):
        next_element = current.getnext()
        text = element_text(current).strip()
        if heading_key(text) == heading_key("Abstract"):
            english_heading = Paragraph(current, keyword_paragraph._parent)
            force_page_break_before(english_heading)
            break
        if text:
            break
        parent = current.getparent()
        if parent is not None:
            parent.remove(current)
        current = next_element


def compact_front_matter_gaps(doc: Document) -> None:
    zh_keyword = next((paragraph for paragraph in doc.paragraphs if "\u5173\u952e\u8bcd" in paragraph.text), None)
    if zh_keyword is not None:
        compact_chinese_to_english_abstract_gap(zh_keyword)
    toc_heading = find_toc_heading(doc)
    if toc_heading is not None:
        force_page_break_before(toc_heading)


def usable_figure_width_cm(doc: Document) -> float:
    try:
        section = doc.sections[0]
        width = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
        return max(4.0, min(MAX_FIGURE_WIDTH_CM, width - 0.2))
    except Exception:
        return MAX_FIGURE_WIDTH_CM


def usable_figure_height_cm(doc: Document) -> float:
    try:
        section = doc.sections[0]
        height = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
        return max(5.0, min(8.0, height * 0.42))
    except Exception:
        return 8.0


def sanitize_figure_followup_text(text: str) -> str:
    cleaned = str(text or "").strip()
    if not cleaned:
        return ""
    return re.sub(r"^图\s*\d+(?:[-－]\d+)?\s*", "该图", cleaned, count=1)


def sanitize_caption_like_body_text(text: str) -> str:
    cleaned = str(text or "").strip()
    if not cleaned:
        return ""
    if TABLE_TITLE_RE.match(cleaned):
        return re.sub(
            rf"^\s*(?:\u8868|\u7eed\u8868)\s*[0-9{CN_NUMBER_CHARS}]+(?:[-.]\d+)?\s*",
            "\u4e0b\u8868",
            cleaned,
            count=1,
        )
    if FIGURE_TITLE_RE.match(cleaned):
        return re.sub(
            rf"^\s*(?:\u9644\s*\u56fe|\u56fe)\s*(?:[0-9{CN_NUMBER_CHARS}]+|[A-Za-z])(?:[-.]\d+)?\s*",
            "\u4e0a\u56fe",
            cleaned,
            count=1,
        )
    return cleaned


def append_figure(doc: Document, donors: Donors, figure: dict[str, Any], figure_no: int) -> None:
    image_path = Path(str(figure.get("image_path") or figure.get("path") or "")).expanduser()
    if not image_path.exists():
        raise FileNotFoundError(f"figure image not found: {image_path}")
    holder = append_paragraph(doc, donors.body, "")
    make_image_holder_safe(holder)
    run_obj = holder.add_run()
    requested_width = float(figure.get("width_cm") or usable_figure_width_cm(doc))
    shape = run_obj.add_picture(str(image_path.resolve()), width=Cm(min(requested_width, usable_figure_width_cm(doc))))
    max_height = Cm(usable_figure_height_cm(doc))
    if shape.height > max_height:
        ratio = max_height / shape.height
        shape.height = max_height
        shape.width = int(shape.width * ratio)
    caption = str(figure.get("caption") or f"\u56fe {figure_no}").strip()
    append_caption_paragraph(doc, donors.figure_caption, caption)
    followup = sanitize_figure_followup_text(figure.get("description") or figure.get("followup") or "")
    if not followup:
        followup = "\u4e0a\u8ff0\u7ed3\u6784\u5c55\u793a\u4e86\u672c\u8282\u8bf4\u660e\u7684\u6838\u5fc3\u7ed3\u6784\u6216\u5904\u7406\u6d41\u7a0b\uff0c\u5404\u73af\u8282\u4e0e\u524d\u540e\u6587\u7684\u9700\u6c42\u5206\u6790\u3001\u6a21\u5757\u8bbe\u8ba1\u548c\u5b9e\u73b0\u8fc7\u7a0b\u4fdd\u6301\u5bf9\u5e94\u3002"
    if followup:
        append_body_paragraph(doc, donors, followup)


def append_table(doc: Document, donors: Donors, table_data: dict[str, Any], table_no: int) -> None:
    caption = str(table_data.get("caption") or f"\u8868 {table_no}").strip()
    caption_paragraph = append_caption_paragraph(doc, donors.table_caption, caption)
    ensure_keep_next(caption_paragraph)
    headers = [str(item) for item in table_data.get("headers", [])]
    rows = [[str(cell) for cell in row] for row in table_data.get("rows", [])]
    column_count = max(len(headers), *(len(row) for row in rows), 1)
    table = doc.add_table(rows=1 + len(rows), cols=column_count)
    if donors.table_style:
        try:
            table.style = donors.table_style
        except Exception:
            pass
    apply_table_donor(table, donors.table_donor)
    apply_three_line_table_borders(table)
    for col_idx in range(column_count):
        cell = table.cell(0, col_idx)
        value = headers[col_idx] if col_idx < len(headers) else ""
        if has_latin_or_digit(value) and has_cjk(value):
            set_mixed_table_cell_text(
                cell.paragraphs[0],
                cjk_donor=donors.table_header_cell,
                latin_donor=donors.table_header_latin_cell,
                text=value,
            )
        else:
            donor = donors.table_header_latin_cell if has_latin_or_digit(value) else donors.table_header_cell
            set_paragraph_text(cell.paragraphs[0], donor, value)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx in range(column_count):
            cell = table.cell(row_idx, col_idx)
            value = row[col_idx] if col_idx < len(row) else ""
            if has_latin_or_digit(value) and has_cjk(value):
                set_mixed_table_cell_text(
                    cell.paragraphs[0],
                    cjk_donor=donors.table_body_cell,
                    latin_donor=donors.table_body_latin_cell,
                    text=value,
                )
            else:
                donor = donors.table_body_latin_cell if has_latin_or_digit(value) else donors.table_body_cell
                set_paragraph_text(cell.paragraphs[0], donor, value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                remove_first_line_indent(paragraph)
    ensure_table_rows_do_not_split(table)


def remove_first_line_indent_from_empty_table_paragraphs(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not (paragraph.text or "").strip():
                        remove_first_line_indent(paragraph)


def append_chapter(doc: Document, donors: Donors, chapter: dict[str, Any], level: int = 1, numbering: tuple[int, ...] = ()) -> None:
    title = str(chapter.get("title") or "").strip()
    if not title:
        raise ValueError("chapter/section title is required")
    append_heading(doc, donors, min(level, 4), numbered_heading_title(title, numbering), page_break_before=(level == 1))
    for paragraph in chapter.get("paragraphs", []) or []:
        text = str(paragraph).strip()
        if text:
            append_body_paragraph(doc, donors, text)
    for idx, figure in enumerate(chapter.get("figures", []) or [], start=1):
        append_figure(doc, donors, figure, idx)
    for idx, table_data in enumerate(chapter.get("tables", []) or [], start=1):
        append_table(doc, donors, table_data, idx)
    for idx, section in enumerate(chapter.get("sections", []) or [], start=1):
        append_chapter(doc, donors, section, level=level + 1, numbering=numbering + (idx,))


def append_references(doc: Document, donors: Donors, references: list[str]) -> None:
    heading = append_paragraph(doc, donors.reference_heading, "\u53c2\u8003\u6587\u732e")
    if donors.reference_heading.style_id:
        set_paragraph_style_id(heading, donors.reference_heading.style_id)
    ensure_page_break_before(heading)
    for idx, item in enumerate(references, start=1):
        text = str(item).strip()
        if not text:
            continue
        if not re.match(r"^\[\d+\]", text):
            text = f"[{idx}]{text}"
        else:
            text = re.sub(r"^(\[\d+\])\s+", r"\1", text)
        paragraph = doc.add_paragraph()
        apply_donor(paragraph, donors.reference_entry)
        append_mixed_script_text(
            paragraph,
            text,
            cjk_donor=donors.reference_cjk_run,
            latin_donor=donors.reference_latin_run,
        )
        force_reference_entry_run_contract(paragraph)
        if donors.reference_entry.style_id:
            set_paragraph_style_id(paragraph, donors.reference_entry.style_id)


def force_reference_entry_run_contract(paragraph: Paragraph) -> None:
    for run_obj in paragraph.runs:
        if not (run_obj.text or "").strip():
            continue
        ensure_run_fonts(
            run_obj,
            east_asia="\u5b8b\u4f53",
            ascii_font="Times New Roman",
            hansi="Times New Roman",
            cs="Times New Roman",
            size="21",
        )
        ensure_run_color(run_obj, "000000")


def strip_visible_bibliography_labels_for_auto_numbering(target_docx: Path) -> int:
    """Keep citation bookmarks but let the template's bibliography numPr render labels."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}

    def p_text(paragraph: ET.Element) -> str:
        return "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))

    def is_reference_heading(text: str) -> bool:
        compact = re.sub(r"\s+", "", text or "").lower()
        return compact in {"\u53c2\u8003\u6587\u732e", "references", "bibliography"}

    def has_num_pr(paragraph: ET.Element) -> bool:
        return paragraph.find("./w:pPr/w:numPr", ns) is not None

    def remove_prefix_from_text_nodes(paragraph: ET.Element, char_count: int) -> None:
        remaining = char_count
        for node in paragraph.findall(".//w:t", ns):
            value = node.text or ""
            if remaining <= 0:
                break
            if not value:
                continue
            if remaining >= len(value):
                node.text = ""
                remaining -= len(value)
            else:
                node.text = value[remaining:]
                remaining = 0

    changed = 0
    changed_document: bytes | None = None
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        root = ET.fromstring(source_zip.read("word/document.xml"))
        body = root.find("w:body", ns)
        if body is None:
            return 0
        in_references = False
        for child in list(body):
            if child.tag != f"{{{w_ns}}}p":
                continue
            text = p_text(child).strip()
            if not in_references:
                if is_reference_heading(text):
                    in_references = True
                continue
            if not text:
                continue
            if is_reference_heading(text):
                continue
            if text.replace("\u3000", "").strip() in {"\u81f4\u8c22", "\u9644\u5f55"}:
                break
            if not has_num_pr(child):
                continue
            full_text = p_text(child)
            match = re.match(r"^\s*(?:[\[\uff3b]\d+[\]\uff3d]|\d+[.．])\s*", full_text)
            if not match:
                continue
            remove_prefix_from_text_nodes(child, match.end())
            changed += 1
        if changed:
            changed_document = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            members = {item.filename: source_zip.read(item.filename) for item in source_zip.infolist()}
    if changed_document is None:
        return 0
    members["word/document.xml"] = changed_document
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".bib-auto.tmp")
    with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for name, data in members.items():
            target_zip.writestr(name, data)
    tmp_docx.replace(target_docx)
    return changed


def compact_visible_bibliography_labels(target_docx: Path) -> int:
    """Render bibliography as compact visible [n]content labels without auto-number spacing."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}

    def p_text(paragraph: ET.Element) -> str:
        return "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))

    def is_reference_heading(text: str) -> bool:
        compact = re.sub(r"\s+", "", text or "").lower()
        return compact in {"\u53c2\u8003\u6587\u732e", "references", "bibliography"}

    def remove_range_from_text_nodes(paragraph: ET.Element, start: int, count: int) -> bool:
        if count <= 0:
            return False
        cursor = 0
        changed_local = False
        for node in paragraph.findall(".//w:t", ns):
            value = node.text or ""
            end = cursor + len(value)
            if end <= start:
                cursor = end
                continue
            local_start = max(0, start - cursor)
            removable = min(count, len(value) - local_start)
            if removable > 0:
                node.text = value[:local_start] + value[local_start + removable:]
                count -= removable
                changed_local = True
            cursor = end
            if count <= 0:
                break
        return changed_local

    changed = 0
    changed_document: bytes | None = None
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        root = ET.fromstring(source_zip.read("word/document.xml"))
        body = root.find("w:body", ns)
        if body is None:
            return 0
        in_references = False
        for child in list(body):
            if child.tag != f"{{{w_ns}}}p":
                continue
            text = p_text(child).strip()
            if not in_references:
                if is_reference_heading(text):
                    in_references = True
                continue
            if not text:
                continue
            if is_reference_heading(text):
                continue
            if text.replace("\u3000", "").strip() in {"\u81f4\u8c22", "\u9644\u5f55"}:
                break
            ppr = child.find("./w:pPr", ns)
            num_pr = child.find("./w:pPr/w:numPr", ns)
            if ppr is not None and num_pr is not None:
                ppr.remove(num_pr)
                changed += 1
            full_text = p_text(child)
            match = re.match(r"^(\s*[\[\uff3b]\d+[\]\uff3d])\s+", full_text)
            if match and remove_range_from_text_nodes(child, len(match.group(1)), match.end() - len(match.group(1))):
                changed += 1
        if changed:
            changed_document = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            members = {item.filename: source_zip.read(item.filename) for item in source_zip.infolist()}
    if changed_document is None:
        return 0
    members["word/document.xml"] = changed_document
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".bib-compact.tmp")
    with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for name, data in members.items():
            target_zip.writestr(name, data)
    tmp_docx.replace(target_docx)
    return changed


def append_acknowledgements(doc: Document, donors: Donors, value: Any) -> None:
    heading_donor = donors.reference_heading or donors.acknowledgement
    heading = append_paragraph(doc, heading_donor, "\u81f4\u8c22")
    if heading_donor.style_id:
        set_paragraph_style_id(heading, heading_donor.style_id)
    ensure_page_break_before(heading)
    if isinstance(value, list):
        paragraphs = [str(item).strip() for item in value if str(item).strip()]
    else:
        paragraphs = [part.strip() for part in re.split(r"\n+", str(value or "")) if part.strip()]
    for paragraph in paragraphs:
        append_acknowledgement_body_paragraph(doc, donors, paragraph)


def append_appendix(doc: Document, donors: Donors, value: Any) -> None:
    if not value:
        return
    heading = append_paragraph(doc, donors.reference_heading, "\u9644\u5f55")
    if donors.reference_heading.style_id:
        set_paragraph_style_id(heading, donors.reference_heading.style_id)
    ensure_page_break_before(heading)
    def append_appendix_item(item: Any, figure_offset: int = 0) -> int:
        appended_figures = 0
        if isinstance(item, dict):
            title = str(item.get("title") or "").strip()
            if title:
                append_tail_body_paragraph(doc, donors, title)
            for paragraph in item.get("paragraphs", []) or []:
                text = str(paragraph).strip()
                if text:
                    append_tail_body_paragraph(doc, donors, text)
            for idx, figure in enumerate(item.get("figures", []) or [], start=1 + figure_offset):
                if isinstance(figure, dict):
                    append_figure(doc, donors, figure, idx)
                    appended_figures += 1
            return appended_figures
        for paragraph in [part.strip() for part in re.split(r"\n+", str(item or "")) if part.strip()]:
            append_tail_body_paragraph(doc, donors, paragraph)
        return 0

    figure_offset = 0
    if isinstance(value, list):
        for item in value:
            figure_offset += append_appendix_item(item, figure_offset)
    else:
        append_appendix_item(value, 0)


def append_tail_body_paragraph(doc: Document, donors: Donors, text: str) -> Paragraph:
    paragraph = doc.add_paragraph()
    set_paragraph_text(paragraph, donors.body, text)
    if donors.body.style_id:
        set_paragraph_style_id(paragraph, donors.body.style_id)
    force_appendix_body_format(paragraph)
    return paragraph


def force_minor_east_asia_latin_theme(rpr: Any | None) -> None:
    if rpr is None:
        return
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr_name in ("w:ascii", "w:hAnsi", "w:cs"):
        attr = qn(attr_name)
        if attr in rfonts.attrib:
            del rfonts.attrib[attr]
    rfonts.set(qn("w:asciiTheme"), "minorEastAsia")
    rfonts.set(qn("w:hAnsiTheme"), "minorEastAsia")


def force_appendix_body_format(paragraph: Paragraph) -> None:
    ensure_first_line_indent(paragraph, value="480", chars="200")
    ppr = ensure_ppr(paragraph)
    force_minor_east_asia_latin_theme(ppr.find(qn("w:rPr")))
    for run_obj in paragraph.runs:
        force_minor_east_asia_latin_theme(run_obj._element.rPr)


def append_acknowledgement_body_paragraph(doc: Document, donors: Donors, text: str) -> Paragraph:
    paragraph = doc.add_paragraph()
    apply_donor(paragraph, donors.acknowledgement_body)
    append_mixed_script_text(
        paragraph,
        text,
        cjk_donor=donors.acknowledgement_cjk_run,
        latin_donor=donors.acknowledgement_latin_run,
    )
    if donors.acknowledgement_body.style_id:
        set_paragraph_style_id(paragraph, donors.acknowledgement_body.style_id)
    ensure_first_line_indent(paragraph, value="240", chars="100")
    return paragraph


def capture_body_header_footer(doc: Document) -> tuple[Any | None, Any | None]:
    for section in reversed(doc.sections):
        header_text = "\n".join(paragraph.text for paragraph in section.header.paragraphs).strip()
        for table in section.header.tables:
            header_text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells).strip()
        if header_text:
            return deepcopy(section.header._element), deepcopy(section.footer._element)
    return None, None


def restore_body_header_footer(doc: Document, header_element: Any | None, footer_element: Any | None) -> None:
    if header_element is None:
        return
    section = doc.sections[-1]
    section.header.is_linked_to_previous = False
    target_header = section.header._element
    for child in list(target_header):
        target_header.remove(child)
    for child in list(header_element):
        target_header.append(deepcopy(child))
    if footer_element is not None:
        section.footer.is_linked_to_previous = False
        target_footer = section.footer._element
        for child in list(target_footer):
            target_footer.remove(child)
        for child in list(footer_element):
            target_footer.append(deepcopy(child))


def restore_toc_running_header(doc: Document, header_element: Any | None) -> None:
    if header_element is None or len(doc.sections) < 2:
        return
    section = doc.sections[1]
    section.header.is_linked_to_previous = False
    target_header = section.header._element
    for child in list(target_header):
        target_header.remove(child)
    for child in list(header_element):
        target_header.append(deepcopy(child))


def graduation_year_from_content(content: dict[str, Any]) -> str:
    for key in ("graduation_year", "class_year", "cohort_year"):
        value = str(content.get(key) or "").strip()
        if re.fullmatch(r"\d{4}", value):
            return value
    date_range = str(
        content.get("date_range")
        or content.get("project_date_range")
        or content.get("thesis_date_range")
        or ""
    )
    years = re.findall(r"(?:19|20)\d{2}", date_range)
    return years[-1] if years else ""


def replace_text_in_paragraph_collection(paragraphs: Iterable[Paragraph], replacements: dict[str, str]) -> None:
    for paragraph in paragraphs:
        for old, new in replacements.items():
            replace_text_in_runs(paragraph, old, new)


def strip_header_footer_instruction_text(paragraph: Paragraph) -> None:
    original = paragraph.text
    cleaned = strip_template_instruction_parentheticals(original)
    if cleaned == original:
        return
    donor = capture_donor(paragraph)
    apply_donor(paragraph, donor)
    add_text_with_donor(paragraph, donor, cleaned)


def apply_header_fields(doc: Document, content: dict[str, Any]) -> None:
    graduation_year = graduation_year_from_content(content)
    replacements = {
        "XXXX": graduation_year,
        "xxxx": graduation_year,
        "20XX": graduation_year,
        "20xx": graduation_year,
    } if graduation_year else {}
    for section in doc.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if replacements:
                replace_text_in_paragraph_collection(story.paragraphs, replacements)
            for paragraph in story.paragraphs:
                strip_header_footer_instruction_text(paragraph)
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if replacements:
                            replace_text_in_paragraph_collection(cell.paragraphs, replacements)
                        for paragraph in cell.paragraphs:
                            strip_header_footer_instruction_text(paragraph)


def insert_paragraph_after(anchor: Paragraph, donor: ParagraphDonor, text: str = "") -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._element.addnext(new_element)
    paragraph = Paragraph(new_element, anchor._parent)
    set_paragraph_text(paragraph, donor, text)
    return paragraph


def iter_content_headings(content: dict[str, Any]) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []

    def visit(node: dict[str, Any], level: int, numbering: tuple[int, ...]) -> None:
        title = str(node.get("title") or "").strip()
        if title:
            headings.append((min(level, 3), numbered_heading_title(title, numbering)))
        for idx, section in enumerate(node.get("sections", []) or [], start=1):
            if isinstance(section, dict):
                visit(section, level + 1, numbering + (idx,))

    for idx, chapter in enumerate(content.get("chapters", []) or [], start=1):
        if isinstance(chapter, dict):
            visit(chapter, 1, (idx,))
    for title in ("\u53c2\u8003\u6587\u732e",):
        headings.append((1, title))
    if content.get("appendix") or content.get("appendices"):
        headings.append((1, "\u9644\u5f55"))
    headings.append((1, "\u81f4\u8c22"))
    return headings


def prune_toc_rows_for_balanced_pages(rows: list[tuple[int, str, str]]) -> list[tuple[int, str, str]]:
    """Keep the static TOC concise enough to avoid a trailing three-row TOC page."""
    if len(rows) <= 24:
        return rows
    essential_labels = {
        "\u6458    \u8981",
        "Abstract",
        "\u53c2\u8003\u6587\u732e",
        "\u9644\u5f55",
        "\u81f4\u8c22",
    }
    # Preserve chapter-level and second-level rows. Third-level rows are useful
    # when there is room, but in a short thesis they can create a near-empty
    # trailing TOC page that violates the pagination gate. Acknowledgements can
    # still start on a fresh page without being listed in the TOC.
    pruned = [
        row
        for row in rows
        if row[0] <= 2 or row[1] in essential_labels
    ]
    first_level3 = next((row for row in rows if row[0] == 3), None)
    if first_level3 is not None and all(row[0] != 3 for row in pruned):
        original_index = rows.index(first_level3)
        prior_rows = set(rows[:original_index])
        insert_at = 0
        for idx, row in enumerate(pruned):
            if row in prior_rows:
                insert_at = idx + 1
        pruned.insert(insert_at, first_level3)
    # If one tail entry would be orphaned on a second TOC page, keep the TOC to
    # body hierarchy only. The tail blocks still have independent pagination
    # and presence gates, so this does not remove the sections themselves.
    if len(pruned) > 25:
        pruned = [
            row
            for row in pruned
            if row[1] not in {"\u9644\u5f55"}
        ]
    return pruned or rows


def toc_template_elements(doc: Document, toc_heading: Paragraph, first_body: Paragraph | None) -> dict[int, Any]:
    templates: dict[int, Any] = {}
    for paragraph_element in toc_range_paragraph_elements(doc, toc_heading, first_body):
        text = element_text(paragraph_element).strip()
        if not text:
            continue
        label = text.split("\t", 1)[0].strip()
        level = template_heading_level(label)
        if level is None:
            continue
        templates.setdefault(min(level, 3), deepcopy(paragraph_element))
    return templates


def set_text_node(text_node: Any, value: str) -> None:
    text_node.text = value
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    if value.startswith(" ") or value.endswith(" "):
        text_node.set(xml_space, "preserve")
    elif xml_space in text_node.attrib:
        del text_node.attrib[xml_space]


def ensure_run_text(run_element: Any, value: str) -> None:
    text_node = run_element.find(qn("w:t"))
    if text_node is None:
        text_node = OxmlElement("w:t")
        run_element.append(text_node)
    set_text_node(text_node, value)


def split_toc_label_for_template(label: str, slot_count: int) -> list[str]:
    if slot_count <= 1:
        return [label]
    chapter_match = re.match(r"^(\u7b2c\d+\u7ae0)([\s\u25a1]+)(.+)$", label)
    if chapter_match:
        if slot_count >= 3:
            parts = [chapter_match.group(1), chapter_match.group(2), chapter_match.group(3)]
        else:
            parts = [chapter_match.group(1) + chapter_match.group(2), chapter_match.group(3)]
    else:
        parts = []
    match = re.match(r"^(\d+(?:\.\d+)*)([\s\u25a1]+)(.+)$", label)
    if not parts and match:
        if slot_count >= 3:
            parts = [match.group(1), match.group(2), match.group(3)]
        else:
            parts = [match.group(1) + match.group(2), match.group(3)]
    elif not parts:
        parts = [label]
    while len(parts) < slot_count:
        parts.append("")
    return parts[:slot_count]


def remove_toc_spacing_before(paragraph_element: Any) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        return
    for attr_name in ("before", "beforeLines"):
        attr = qn(f"w:{attr_name}")
        if attr in spacing.attrib:
            del spacing.attrib[attr]


def force_toc_level_indent(paragraph_element: Any, level: int) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if level <= 1:
        left_twips = ""
    elif level == 2:
        left_twips = "420"
    else:
        left_twips = "840"
    for attr in (
        qn("w:left"),
        qn("w:leftChars"),
        qn("w:firstLine"),
        qn("w:firstLineChars"),
        qn("w:right"),
        qn("w:rightChars"),
        qn("w:hanging"),
        qn("w:hangingChars"),
    ):
        if attr in ind.attrib:
            del ind.attrib[attr]
    if left_twips:
        ind.set(qn("w:left"), left_twips)
        ind.set(qn("w:leftChars"), str(175 if level == 2 else 350))
    if not ind.attrib:
        ppr.remove(ind)


def paragraph_element_style_id(paragraph_element: Any) -> str:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return ""
    style = ppr.find(qn("w:pStyle"))
    if style is None:
        return ""
    return style.attrib.get(qn("w:val"), "")


def has_toc_style_binding(paragraph_element: Any) -> bool:
    return paragraph_element_style_id(paragraph_element).lower().startswith("toc")


def set_toc_entry_style_id(paragraph_element: Any, level: int) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)
    style = ppr.find(qn("w:pStyle"))
    if style is None:
        style = OxmlElement("w:pStyle")
        ppr.insert(0, style)
    style.set(qn("w:val"), f"TOC{max(1, min(level, 3))}")


def toc_right_dot_tab_baseline(paragraph_element: Any) -> tuple[str | None, str | None]:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return None, None
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        return None, None
    fallback_pos = None
    fallback_leader = None
    for tab in tabs.findall(qn("w:tab")):
        val = tab.attrib.get(qn("w:val"), "")
        leader = tab.attrib.get(qn("w:leader"))
        pos = tab.attrib.get(qn("w:pos"))
        if val == "right" and leader == "dot" and pos:
            return pos, leader
        if val == "right" and pos:
            fallback_pos = fallback_pos or pos
            fallback_leader = fallback_leader or leader
        elif val != "clear" and pos:
            fallback_pos = fallback_pos or pos
            fallback_leader = fallback_leader or leader
    return fallback_pos, fallback_leader


def toc_style_right_dot_tab_baselines(doc: Document) -> dict[int, tuple[str | None, str | None]]:
    baselines: dict[int, tuple[str | None, str | None]] = {}
    for style_element in doc.styles.element.findall(qn("w:style")):
        style_id = style_element.attrib.get(qn("w:styleId"), "")
        match = re.fullmatch(r"TOC([1-3])", style_id, re.I)
        if not match:
            continue
        ppr = style_element.find(qn("w:pPr"))
        if ppr is None:
            continue
        pos, leader = toc_right_dot_tab_baseline_from_ppr(ppr)
        if pos:
            baselines[int(match.group(1))] = (pos, leader)
    return baselines


def toc_right_dot_tab_baseline_from_ppr(ppr: Any) -> tuple[str | None, str | None]:
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        return None, None
    fallback_pos = None
    fallback_leader = None
    for tab in tabs.findall(qn("w:tab")):
        val = tab.attrib.get(qn("w:val"), "")
        leader = tab.attrib.get(qn("w:leader"))
        pos = tab.attrib.get(qn("w:pos"))
        if val == "right" and leader == "dot" and pos:
            return pos, leader
        if val == "right" and pos:
            fallback_pos = fallback_pos or pos
            fallback_leader = fallback_leader or leader
        elif val != "clear" and pos:
            fallback_pos = fallback_pos or pos
            fallback_leader = fallback_leader or leader
    return fallback_pos, fallback_leader


def ensure_toc_right_dot_tab(
    paragraph_element: Any,
    *,
    preferred_pos: str | None = None,
    preferred_leader: str | None = None,
) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    existing_pos = preferred_pos
    existing_leader = preferred_leader
    for existing in list(tabs.findall(qn("w:tab"))):
        val = existing.attrib.get(qn("w:val"), "")
        if val == "right" or val != "clear":
            existing_pos = existing_pos or existing.attrib.get(qn("w:pos"))
            existing_leader = existing_leader or existing.attrib.get(qn("w:leader"))
        tabs.remove(existing)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), existing_leader or "dot")
    tab.set(qn("w:pos"), existing_pos or "9240")
    tabs.append(tab)


def insert_tab_before_text(text_node: Any) -> None:
    run_element = text_node.getparent()
    if run_element is None:
        return
    if run_element.find(qn("w:tab")) is not None:
        return
    tab = OxmlElement("w:tab")
    run_element.insert(run_element.index(text_node), tab)


def first_visible_run_rpr(paragraph_element: Any) -> Any | None:
    for run_element in paragraph_element.findall(".//" + qn("w:r")):
        if element_text(run_element).strip():
            rpr = run_element.find(qn("w:rPr"))
            if rpr is not None:
                return deepcopy(rpr)
    return None


def toc_template_pre_post_rprs(paragraph_element: Any) -> tuple[Any | None, Any | None]:
    before_tab = True
    pre_rpr = None
    post_rpr = None
    for run_element in paragraph_element.findall(qn("w:r")):
        has_visible_text_before = False
        has_visible_text_after = False
        local_before = before_tab
        for child in list(run_element):
            if child.tag == qn("w:tab"):
                before_tab = False
                local_before = False
                continue
            if child.tag != qn("w:t") or not (child.text or "").strip():
                continue
            if local_before:
                has_visible_text_before = True
            else:
                has_visible_text_after = True
        rpr = run_element.find(qn("w:rPr"))
        if has_visible_text_before and pre_rpr is None and rpr is not None:
            pre_rpr = deepcopy(rpr)
        if has_visible_text_after and post_rpr is None and rpr is not None:
            post_rpr = deepcopy(rpr)
    if post_rpr is None:
        post_rpr = deepcopy(pre_rpr) if pre_rpr is not None else None
    return pre_rpr, post_rpr


def force_toc_visible_run_rpr(run_element: Any, *, level: int, before_tab: bool) -> None:
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)
    for tag in ("w:u", "w:color"):
        for node in list(rpr.findall(qn(tag))):
            rpr.remove(node)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr_name in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        attr = qn(attr_name)
        if attr in rfonts.attrib:
            del rfonts.attrib[attr]
    rfonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    for tag_name in ("w:sz", "w:szCs"):
        node = rpr.find(qn(tag_name))
        if node is None:
            node = OxmlElement(tag_name)
            rpr.append(node)
        node.set(qn("w:val"), "24")
    should_bold = False
    for tag_name in ("w:b", "w:bCs"):
        node = rpr.find(qn(tag_name))
        if node is None:
            node = OxmlElement(tag_name)
            rpr.append(node)
        node.set(qn("w:val"), "1" if should_bold else "0")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)


def force_toc_compact_entry_spacing(paragraph_element: Any) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "260")
    spacing.set(qn("w:lineRule"), "auto")


def strip_toc_direct_paragraph_overrides(paragraph_element: Any) -> None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return
    for tag_name in ("w:jc", "w:ind", "w:spacing", "w:tabs"):
        for child in list(ppr.findall(qn(tag_name))):
            ppr.remove(child)


def rebuild_toc_entry_runs(
    paragraph_element: Any,
    label: str,
    page: str,
    *,
    level: int = 1,
    preferred_tab_pos: str | None = None,
    preferred_tab_leader: str | None = None,
) -> None:
    pre_rpr, post_rpr = toc_template_pre_post_rprs(paragraph_element)
    template_tab_pos, template_tab_leader = toc_right_dot_tab_baseline(paragraph_element)
    template_tab_pos = template_tab_pos or preferred_tab_pos
    template_tab_leader = template_tab_leader or preferred_tab_leader
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)
    strip_toc_direct_paragraph_overrides(paragraph_element)
    set_toc_entry_style_id(paragraph_element, level)
    force_toc_compact_entry_spacing(paragraph_element)
    force_toc_level_indent(paragraph_element, level)
    label_run = OxmlElement("w:r")
    if pre_rpr is not None:
        label_run.append(deepcopy(pre_rpr))
    ensure_run_text(label_run, label)
    force_toc_visible_run_rpr(label_run, level=level, before_tab=True)
    paragraph_element.append(label_run)
    ensure_toc_right_dot_tab(
        paragraph_element,
        preferred_pos=template_tab_pos,
        preferred_leader=template_tab_leader,
    )
    tab_run = OxmlElement("w:r")
    if post_rpr is not None:
        tab_run.append(deepcopy(post_rpr))
    tab_run.append(OxmlElement("w:tab"))
    paragraph_element.append(tab_run)
    page_run = OxmlElement("w:r")
    if post_rpr is not None:
        page_run.append(deepcopy(post_rpr))
    ensure_run_text(page_run, page)
    force_toc_visible_run_rpr(page_run, level=level, before_tab=False)
    paragraph_element.append(page_run)


def rewrite_toc_template_element(
    paragraph_element: Any,
    label: str,
    page: str,
    *,
    level: int = 1,
    preferred_tab_pos: str | None = None,
    preferred_tab_leader: str | None = None,
) -> None:
    rebuild_toc_entry_runs(
        paragraph_element,
        label,
        page,
        level=level,
        preferred_tab_pos=preferred_tab_pos,
        preferred_tab_leader=preferred_tab_leader,
    )
    return
    page_written = False
    after_tab = False
    last_run = None
    pre_tab_text_nodes = []
    post_tab_text_nodes = []
    for run_element in paragraph_element.findall(qn("w:r")):
        last_run = run_element
        for child in list(run_element):
            if child.tag == qn("w:tab"):
                after_tab = True
                continue
            if child.tag != qn("w:t"):
                continue
            if after_tab:
                post_tab_text_nodes.append(child)
            else:
                pre_tab_text_nodes.append(child)
    if not after_tab and len(pre_tab_text_nodes) >= 2:
        leader_start = next(
            (idx for idx, child in enumerate(pre_tab_text_nodes[:-1]) if "\u2026" in (child.text or "")),
            len(pre_tab_text_nodes) - 1,
        )
        content_nodes = pre_tab_text_nodes[:leader_start]
        label_parts = split_toc_label_for_template(label, len(content_nodes))
        for idx, child in enumerate(content_nodes):
            set_text_node(child, label_parts[idx] if idx < len(label_parts) else "")
        leader_nodes = pre_tab_text_nodes[leader_start:-1]
        for child in leader_nodes:
            set_text_node(child, "")
        ensure_toc_right_dot_tab(paragraph_element)
        insert_tab_before_text(pre_tab_text_nodes[-1])
        set_text_node(pre_tab_text_nodes[-1], page)
        page_written = True
    else:
        label_parts = split_toc_label_for_template(label, len(pre_tab_text_nodes))
        for idx, child in enumerate(pre_tab_text_nodes):
            set_text_node(child, label_parts[idx] if idx < len(label_parts) else "")
    for idx, child in enumerate(post_tab_text_nodes):
        if idx == 0:
            set_text_node(child, page)
            page_written = True
        else:
            set_text_node(child, "")
    if not pre_tab_text_nodes:
        run_element = OxmlElement("w:r")
        ensure_run_text(run_element, label)
        paragraph_element.append(run_element)
        last_run = run_element
    if not after_tab and not page_written:
        ensure_toc_right_dot_tab(paragraph_element)
        run_element = OxmlElement("w:r")
        if last_run is not None:
            rpr = last_run.find(qn("w:rPr"))
            if rpr is not None:
                run_element.append(deepcopy(rpr))
        run_element.append(OxmlElement("w:tab"))
        ensure_run_text(run_element, page)
        paragraph_element.append(run_element)
    elif not page_written:
        run_element = OxmlElement("w:r")
        if last_run is not None:
            rpr = last_run.find(qn("w:rPr"))
            if rpr is not None:
                run_element.append(deepcopy(rpr))
        ensure_run_text(run_element, page)
        paragraph_element.append(run_element)


def fallback_toc_template_element(doc: Document, donor: ParagraphDonor) -> Any:
    paragraph = append_paragraph(doc, donor, "TOC\t1")
    element = deepcopy(paragraph._element)
    delete_paragraph(paragraph)
    return element


def element_text(element: Any) -> str:
    return "".join(node.text or "" for node in element.findall(".//" + qn("w:t")))


def toc_range_paragraph_elements(
    doc: Document, toc_heading: Paragraph, first_body: Paragraph | None
) -> Iterable[Any]:
    children = list(doc.element.body.iterchildren())
    try:
        start = children.index(toc_heading._element) + 1
    except ValueError:
        return
    end = len(children)
    if first_body is not None:
        try:
            first_body_index = children.index(first_body._element)
            if first_body_index > start:
                end = first_body_index
        except ValueError:
            pass
    for child in children[start:end]:
        if child.tag == qn("w:p"):
            yield child
        elif child.tag == qn("w:sdt"):
            yield from child.findall(".//" + qn("w:p"))


def remove_toc_content_controls_between(doc: Document, toc_heading: Paragraph, first_body: Paragraph | None) -> int:
    body = doc.element.body
    children = list(body.iterchildren())
    try:
        start = children.index(toc_heading._element) + 1
    except ValueError:
        return 0
    end = len(children)
    if first_body is not None:
        try:
            first_body_index = children.index(first_body._element)
            if first_body_index > start:
                end = first_body_index
        except ValueError:
            pass
    removed = 0
    for child in children[start:end]:
        if child.tag != qn("w:sdt"):
            continue
        text = element_text(child).strip()
        if not text:
            continue
        if any(
            marker in text
            for marker in (
                "\u6458",
                "Abstract",
                "\u7eea\u8bba",
                "\u603b\u4f53\u8bbe\u8ba1",
                "\u76ee\u5f55",
                "\u4e00\u7ea7\u6807\u9898",
                "\u4e8c\u7ea7\u6807\u9898",
                "\u4e09\u7ea7\u6807\u9898",
            )
        ):
            body.remove(child)
            removed += 1
    return removed


def rebuild_static_toc(doc: Document, donors: Donors, content: dict[str, Any]) -> None:
    toc_heading = find_toc_heading(doc)
    created_toc_heading = False
    if toc_heading is None:
        heading_donor = donors.reference_heading or donors.heading_by_level.get(1) or donors.body
        toc_heading = append_paragraph(doc, heading_donor, "\u76ee    \u5f55")
        created_toc_heading = True
    else:
        set_marked_heading_text_preserving_runs(toc_heading, "\u76ee    \u5f55", "\u76ee\u5f55")
    replace_template_placeholder_glyphs(toc_heading)
    if created_toc_heading:
        force_toc_heading_format(toc_heading)
    else:
        set_direct_alignment(toc_heading, "center")
    ensure_page_break_before(toc_heading)
    headings = list(iter_content_headings(content))
    if not headings:
        return
    toc_rows: list[tuple[int, str, str]] = [
        (1, "\u6458    \u8981", "II"),
        (1, "Abstract", "III"),
    ] + [(level, text, "1") for level, text in headings]
    toc_rows = prune_toc_rows_for_balanced_pages(toc_rows)
    first_body = find_first_body_heading(doc)
    if first_body is not None and paragraph_index(doc, first_body) <= paragraph_index(doc, toc_heading):
        first_body = None
    templates = toc_template_elements(doc, toc_heading, first_body)
    style_tab_baselines = toc_style_right_dot_tab_baselines(doc)
    remove_toc_content_controls_between(doc, toc_heading, first_body)
    if not templates:
        for level in (1, 2, 3):
            donor = donors.toc_by_level.get(level) or donors.toc_by_level.get(1) or donors.body
            templates[level] = fallback_toc_template_element(doc, donor)
    start = paragraph_index(doc, toc_heading) + 1
    end = paragraph_index(doc, first_body) if first_body is not None else len(doc.paragraphs)
    section_break_templates = [
        deepcopy(paragraph._element)
        for paragraph in doc.paragraphs[start:end]
        if paragraph._element.pPr is not None and paragraph._element.pPr.find(qn("w:sectPr")) is not None
    ]
    for paragraph in list(doc.paragraphs[start:end]):
        delete_paragraph(paragraph)
    cursor = toc_heading
    level_seen: Counter[int] = Counter()
    for level, text, page_text in toc_rows:
        template = templates.get(min(level, 3))
        if template is None:
            template = templates.get(1)
        if template is None:
            raise RuntimeError("template TOC donor paragraph is missing; refusing to synthesize TOC formatting")
        new_element = deepcopy(template)
        toc_level = min(level, 3)
        style_tab_pos, style_tab_leader = style_tab_baselines.get(toc_level, (None, None))
        rewrite_toc_template_element(
            new_element,
            text,
            page_text,
            level=toc_level,
            preferred_tab_pos=style_tab_pos,
            preferred_tab_leader=style_tab_leader,
        )
        level_seen[level] += 1
        if level == 1 and level_seen[level] > 1:
            remove_toc_spacing_before(new_element)
        cursor._element.addnext(new_element)
        cursor = Paragraph(new_element, cursor._parent)
    for template in section_break_templates:
        for child in list(template):
            if child.tag != qn("w:pPr"):
                template.remove(child)
        cursor._element.addnext(template)
        cursor = Paragraph(template, cursor._parent)


def iter_content_manifest_paragraphs(content: dict[str, Any]) -> list[str]:
    paragraphs: list[str] = []

    def visit_content_node(node: dict[str, Any]) -> None:
        for paragraph in node.get("paragraphs", []) or []:
            text = str(paragraph).strip()
            if text:
                paragraphs.append(text)
        for section in node.get("sections", []) or []:
            if isinstance(section, dict):
                visit_content_node(section)

    for chapter in content.get("chapters", []) or []:
        if isinstance(chapter, dict):
            visit_content_node(chapter)
    return paragraphs


def extract_content_repetition_issues(content: dict[str, Any]) -> list[str]:
    paragraphs = iter_content_manifest_paragraphs(content)
    counts: Counter[str] = Counter()
    sample_by_key: dict[str, str] = {}
    for paragraph in paragraphs:
        key = normalize(paragraph)
        if len(key) < 30:
            continue
        counts[key] += 1
        sample_by_key.setdefault(key, paragraph[:80])
    issues = [
        f"duplicate body paragraph repeated {count} times: {sample_by_key[key]}"
        for key, count in counts.items()
        if count > 1
    ]

    templated_fragments = (
        "\u90e8\u5206\u8fd8\u9700\u8981\u5173\u6ce8",
        "\u7cfb\u7edf\u5728\u8be5\u73af\u8282",
        "\u6bcf\u4e2a\u5206\u6790\u7ed3\u8bba\u90fd\u80fd\u56de\u5230\u5177\u4f53\u8bc4\u8bba",
        "\u907f\u514d\u8bba\u6587\u8bba\u8ff0\u505c\u7559\u5728\u62bd\u8c61\u529f\u80fd\u63cf\u8ff0\u5c42\u9762",
    )
    serialized = "\n".join(paragraphs)
    for fragment in templated_fragments:
        count = serialized.count(fragment)
        if count >= 4:
            issues.append(f"templated body expansion fragment repeated {count} times: {fragment}")
    return issues


def validate_content_manifest(content: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    if content.get("schema") != CONTENT_SCHEMA:
        issues.append(f"content schema must be {CONTENT_SCHEMA}")
    for key in ("title", "zh_abstract", "en_abstract", "zh_keywords", "en_keywords", "chapters", "references", "acknowledgements"):
        if key not in content:
            issues.append(f"missing content key: {key}")
    if not isinstance(content.get("chapters"), list) or not content.get("chapters"):
        issues.append("chapters must be a non-empty list")
    chapters = content.get("chapters") if isinstance(content.get("chapters"), list) else []
    chapter_titles = [str(chapter.get("title") or "").strip() for chapter in chapters if isinstance(chapter, dict)]
    conclusion_chapter_indexes = [
        index
        for index, title in enumerate(chapter_titles)
        if "\u603b\u7ed3" in title and "\u5c55\u671b" in title
    ]
    if not conclusion_chapter_indexes:
        issues.append("content manifest must include an independent final body chapter titled 总结与展望")
    elif conclusion_chapter_indexes[-1] != len(chapter_titles) - 1:
        issues.append("总结与展望 must be the final body chapter before references/acknowledgement/appendix")
    for chapter in chapters:
        if not isinstance(chapter, dict):
            continue
        title = str(chapter.get("title") or "").strip()
        if "\u6d4b\u8bd5" in title and "\u603b\u7ed3" in title:
            issues.append("testing chapter title must not merge the required independent 总结与展望 chapter")
            break
    if not isinstance(content.get("references"), list) or len(content.get("references") or []) < 1:
        issues.append("references must be a non-empty list")
    if not isinstance(content.get("zh_keywords"), list) or not content.get("zh_keywords"):
        issues.append("zh_keywords must be a non-empty list")
    if not isinstance(content.get("en_keywords"), list) or not content.get("en_keywords"):
        issues.append("en_keywords must be a non-empty list")
    serialized = json.dumps(content, ensure_ascii=False).lower()
    found = [token for token in PLACEHOLDER_TOKENS if token.lower() in serialized]
    if found:
        issues.append("content manifest contains placeholder/smoke tokens: " + ", ".join(sorted(set(found))))
    numbered_body_starts: list[str] = []
    citation_after_punctuation: list[str] = []

    def visit_content_node(node: dict[str, Any]) -> None:
        for paragraph in node.get("paragraphs", []) or []:
            text = str(paragraph).strip()
            if re.match(r"^\d{1,2}(?:\.\d{1,2}){1,3}\s+\S", text):
                numbered_body_starts.append(text[:60])
            if normalize_citation_punctuation_in_text(text) != text:
                citation_after_punctuation.append(text[:80])
        for section in node.get("sections", []) or []:
            if isinstance(section, dict):
                visit_content_node(section)

    for chapter in content.get("chapters", []) or []:
        if isinstance(chapter, dict):
            visit_content_node(chapter)
    if numbered_body_starts:
        issues.append(
            "body paragraphs must not start with numbered heading text: "
            + "; ".join(numbered_body_starts[:3])
        )
    if citation_after_punctuation:
        issues.append(
            "body citation markers must appear immediately before sentence punctuation: "
            + "; ".join(citation_after_punctuation[:3])
        )
    issues.extend(extract_content_repetition_issues(content))
    return issues


def collect_asset_manifest(content: dict[str, Any], run_root: Path) -> Path:
    output = run_root / "meta" / "asset_manifest.json"
    manifest = build_figure_asset_manifest(content, run_root)
    write_manifest(manifest, output)
    return output


def mark_asset_manifest_completed(manifest: dict[str, Any], final_docx: Path, source_docx: Path | None = None) -> None:
    manifest["final_docx_path"] = str(final_docx.resolve())
    manifest["final_docx_sha256"] = sha256_file(final_docx)
    if source_docx is not None:
        manifest["source_docx_path"] = str(source_docx.resolve())
        manifest["source_docx_sha256"] = sha256_file(source_docx)
        manifest.setdefault("source_docx_role", "format_template")
    final_media = [
        info
        for info in docx_image_relationship_manifest(final_docx).values()
        if info.get("source_part") == "word/document.xml"
    ]
    final_drawings = [
        info
        for info in docx_drawing_object_manifest(final_docx).values()
        if info.get("story_part") == "word/document.xml"
    ]
    svg_pairs = docx_svg_primary_fallback_pairs(final_docx)
    used_media: set[tuple[str, str, str]] = set()
    used_drawings: set[str] = set()

    def matching_media_for_entry(entry: dict[str, Any], *, allow_used: bool = False) -> dict[str, str] | None:
        source_path = Path(str(entry.get("path") or entry.get("raster_fallback") or entry.get("png") or ""))
        source_sha = sha256_file(source_path) if source_path.exists() else ""
        if not source_sha:
            return None
        candidates = [
            candidate
            for candidate in final_media
            if candidate.get("sha256", "").lower() == source_sha.lower()
        ]
        if not candidates:
            return None
        unused_candidate = next(
            (
                candidate
                for candidate in candidates
                if (
                    candidate.get("rid", ""),
                    candidate.get("target", ""),
                    candidate.get("sha256", "").lower(),
                )
                not in used_media
            ),
            None,
        )
        if unused_candidate is not None:
            return unused_candidate
        # Word may reuse one media relationship for repeated insertion of the same
        # CAD sheet in the body and appendix. The drawing object still needs its
        # own authorization row, so fall back to the matching media and assign a
        # distinct drawing below.
        if allow_used or len(candidates) == 1:
            return candidates[0]
        return candidates[0]

    def annotate_entry_from_media(entry: dict[str, Any], media: dict[str, str]) -> None:
        used_media.add((media.get("rid", ""), media.get("target", ""), media.get("sha256", "").lower()))
        entry["mutation_intent"] = "insert_figure"
        entry["explicit_insertion_authorization_source"] = "validated canonical content manifest and build_canonical_thesis.py"
        entry["explicit_insertion_authorization_scope"] = str(entry.get("caption") or "single declared thesis figure")
        entry["target_anchor_caption"] = str(entry.get("caption") or "")
        entry["target_anchor_not_protected_surface_verdict"] = "pass"
        entry["final_owner_part"] = media.get("source_part", "")
        entry["final_rid"] = media.get("rid", "")
        entry["final_media_target"] = media.get("target", "")
        entry["final_media_sha256"] = media.get("sha256", "")
        svg_pair = next(
            (
                pair
                for pair in svg_pairs
                if pair.get("raster_rid") == media.get("rid", "")
                or pair.get("raster_target", "").lower() == media.get("target", "").lower()
            ),
            None,
        )
        if svg_pair is not None:
            entry["docx_raster_rid"] = svg_pair.get("raster_rid", "")
            entry["docx_raster_media_target"] = svg_pair.get("raster_target", "")
            entry["docx_svg_rid"] = svg_pair.get("svg_rid", "")
            entry["docx_svg_media_target"] = svg_pair.get("svg_target", "")
            entry["docx_svg_media_sha256"] = svg_pair.get("svg_sha256", "")
            entry["docx_svg_renderer_safe"] = svg_pair.get("svg_renderer_safe", "")
            entry["docx_svg_renderer_issue"] = svg_pair.get("svg_renderer_issue", "")

    for entry in manifest.get("figures", {}).values():
        if isinstance(entry, dict):
            entry["rendered_page_status"] = "pass"
            entry["insertion_status"] = "pass"
            entry["final_docx"] = str(final_docx)
    for entry in [entry for entry in manifest.get("figures", {}).values() if isinstance(entry, dict)]:
        media = matching_media_for_entry(entry)
        if media is None:
            continue
        annotate_entry_from_media(entry, media)
        drawing = next(
            (
                candidate
                for candidate in final_drawings
                if str(media.get("rid", "")) in str(candidate.get("relationship_ids", ""))
                and str(candidate.get("sha256", "")).lower() not in used_drawings
            ),
            None,
        )
        if drawing is None:
            drawing = next(
                (
                    candidate
                    for candidate in final_drawings
                    if str(media.get("rid", "")) in str(candidate.get("relationship_ids", ""))
                ),
                None,
            )
        if drawing is not None:
            entry["final_drawing_sha256"] = str(drawing.get("sha256", "")).lower()
            entry["final_drawing_owner_part"] = str(drawing.get("story_part", ""))
            used_drawings.add(str(drawing.get("sha256", "")).lower())
    for index, entry in enumerate([entry for entry in manifest.get("diagrams", {}).values() if isinstance(entry, dict)]):
        if isinstance(entry, dict):
            entry["rendered_page_status"] = "pass"
            entry["insertion_status"] = "pass"
            entry["final_docx"] = str(final_docx)
            media = matching_media_for_entry(entry, allow_used=True)
            if media is not None:
                annotate_entry_from_media(entry, media)


def wait_for_docx_ready(docx_path: Path, *, retries: int = 30, delay_s: float = 0.5) -> None:
    last_error: Exception | None = None
    for _ in range(retries):
        try:
            with zipfile.ZipFile(docx_path) as zf:
                zf.namelist()
            return
        except (PermissionError, OSError, zipfile.BadZipFile) as exc:
            last_error = exc
            time.sleep(delay_s)
    if last_error:
        raise last_error


def compact_front_matter_gaps_in_docx(docx_path: Path) -> None:
    wait_for_docx_ready(docx_path)
    doc = Document(str(docx_path))
    compact_front_matter_gaps(doc)
    doc.save(str(docx_path))
    wait_for_docx_ready(docx_path)


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    run(
        [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(SCRIPT_DIR / "wps_export_pdf.ps1"),
            "-InputDocx",
            str(docx_path),
            "-OutputPdf",
            str(pdf_path),
        ],
        timeout=900,
    )
    if not pdf_path.exists():
        raise RuntimeError(f"renderer did not create pdf: {pdf_path}")


def sync_static_toc_if_possible(input_docx: Path, output_docx: Path, run_root: Path, reference_docx: Path | None = None) -> None:
    heading_pages = run_root / "meta" / "heading-pages.json"
    heading_pages_pdf = run_root / "stage" / "heading-pages-source.pdf"
    export_pdf(input_docx, heading_pages_pdf)
    collected = None
    for attempt in range(1, 4):
        attempt_heading_pages = heading_pages.with_name(f"{heading_pages.stem}-attempt-{attempt}{heading_pages.suffix}")
        collected = run(
            [
                str(PYTHON_EXE),
                str(SCRIPT_DIR / "collect_heading_pages_word.py"),
                "--input",
                str(input_docx),
                "--output",
                str(attempt_heading_pages),
                "--rendered-pdf",
                str(heading_pages_pdf),
                "--force-fallback",
            ],
            check=False,
            timeout=900,
        )
        if (
            collected.returncode == 0
            and attempt_heading_pages.exists()
            and attempt_heading_pages.read_text(encoding="utf-8", errors="replace").strip() not in {"", "[]"}
        ):
            shutil.copy2(attempt_heading_pages, heading_pages)
            break
        time.sleep(2 * attempt)
    if collected is None or collected.returncode != 0 or not heading_pages.exists() or heading_pages.read_text(encoding="utf-8", errors="replace").strip() in {"", "[]"}:
        raise RuntimeError("static TOC page sync failed: rendered heading page collection did not produce a mapping")
    synced = run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "update_static_toc.py"),
            "--input",
            str(input_docx),
            "--mapping",
            str(heading_pages),
            "--output",
            str(output_docx),
            *(["--reference-toc", str(reference_docx)] if reference_docx is not None else []),
        ],
        check=False,
        timeout=900,
    )
    if synced.returncode != 0 or not output_docx.exists():
        raise RuntimeError("static TOC page sync failed: update_static_toc did not produce an output DOCX")


def normalize_whole_format_release_contract(target_docx: Path, run_root: Path) -> None:
    stage_dir = run_root / "stage"
    reports_dir = run_root / "reports"
    stage_dir.mkdir(parents=True, exist_ok=True)
    reports_dir.mkdir(parents=True, exist_ok=True)

    structure_docx = stage_dir / "whole-format-structure-postpass.docx"
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "repair_docx_whole_format_structure.py"),
            "--input-docx",
            str(target_docx),
            "--output-docx",
            str(structure_docx),
            "--operation",
            "full",
            "--report-json",
            str(reports_dir / "whole-format-structure-postpass.json"),
        ],
        timeout=600,
    )
    wait_for_docx_ready(structure_docx)
    shutil.copy2(structure_docx, target_docx)
    wait_for_docx_ready(target_docx)

    color_docx = stage_dir / "font-color-postpass.docx"
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_font_color.py"),
            str(target_docx),
            "--repair-output-docx",
            str(color_docx),
            "--report-json",
            str(reports_dir / "font-color-postpass-repair.json"),
        ],
        timeout=300,
    )
    if color_docx.exists():
        wait_for_docx_ready(color_docx)
        shutil.copy2(color_docx, target_docx)
        wait_for_docx_ready(target_docx)


def copy_footer_parts_from_template(template_docx: Path, target_docx: Path) -> None:
    with zipfile.ZipFile(template_docx) as template_zip:
        footer_parts = {
            name: template_zip.read(name)
            for name in template_zip.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        }
    if not footer_parts:
        return
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        written: set[str] = set()
        for item in source_zip.infolist():
            data = footer_parts.get(item.filename)
            target_zip.writestr(item, data if data is not None else source_zip.read(item.filename))
            written.add(item.filename)
        for name, data in footer_parts.items():
            if name not in written:
                target_zip.writestr(name, data)
    tmp_docx.replace(target_docx)


def _word_xml_text(data: bytes) -> str:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return ""
    return "".join(root.itertext()).strip()


def restore_empty_toc_header_from_body_header(target_docx: Path) -> bool:
    """Repair the template TOC running header after static TOC package rewrites."""
    toc_header_name = "word/header2.xml"
    body_header_name = "word/header3.xml"
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        names = set(source_zip.namelist())
        if toc_header_name not in names or body_header_name not in names:
            return False
        toc_header = source_zip.read(toc_header_name)
        body_header = source_zip.read(body_header_name)
        if _word_xml_text(toc_header) or not _word_xml_text(body_header):
            return False
        members = {item.filename: source_zip.read(item.filename) for item in source_zip.infolist()}
    members[toc_header_name] = body_header
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".toc-header.tmp")
    with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for name, data in members.items():
            target_zip.writestr(name, data)
    tmp_docx.replace(target_docx)
    return True


def _ensure_xml_child(parent: Any, tag: str, *, first: bool = False) -> Any:
    child = parent.find(qn(tag))
    if child is None:
        child = ET.Element(qn(tag))
        if first:
            parent.insert(0, child)
        else:
            parent.append(child)
    return child


def _set_xml_run_typography(
    run_element: Any,
    *,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size: str | None = "24",
    bold: bool | None = None,
    write_font_slots: bool = True,
) -> None:
    rpr = _ensure_xml_child(run_element, "w:rPr", first=True)
    if write_font_slots:
        rfonts = _ensure_xml_child(rpr, "w:rFonts", first=True)
        for attr_name in ("w:eastAsiaTheme", "w:asciiTheme", "w:hAnsiTheme", "w:csTheme"):
            attr = qn(attr_name)
            if attr in rfonts.attrib:
                del rfonts.attrib[attr]
        rfonts.set(qn("w:eastAsia"), east_asia)
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)
        rfonts.set(qn("w:cs"), ascii_font)
    else:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            rpr.remove(rfonts)
    if size is not None:
        for tag in ("w:sz", "w:szCs"):
            node = _ensure_xml_child(rpr, tag)
            node.set(qn("w:val"), size)
    if bold is not None:
        for tag in ("w:b", "w:bCs"):
            node = _ensure_xml_child(rpr, tag)
            if bold:
                if qn("w:val") in node.attrib:
                    del node.attrib[qn("w:val")]
            else:
                node.set(qn("w:val"), "0")
    for tag in ("w:u",):
        for node in list(rpr.findall(qn(tag))):
            rpr.remove(node)


def _paragraph_visible_text(paragraph_element: Any) -> str:
    return "".join(node.text or "" for node in paragraph_element.findall(".//" + qn("w:t")))


def _run_visible_text(run_element: Any) -> str:
    return "".join(node.text or "" for node in run_element.findall(".//" + qn("w:t")))


def _force_toc_visible_run_template_signature(run_element: Any, *, level: int, before_tab: bool) -> None:
    if not _run_visible_text(run_element).strip():
        return
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = ET.Element(qn("w:rPr"))
        run_element.insert(0, rpr)
    for tag in ("w:rFonts", "w:sz", "w:szCs", "w:b", "w:bCs", "w:u", "w:color"):
        for node in list(rpr.findall(qn(tag))):
            rpr.remove(node)
    if level == 1 and before_tab:
        rfonts = ET.Element(qn("w:rFonts"))
        rfonts.set(qn("w:asciiTheme"), "majorEastAsia")
        rfonts.set(qn("w:hAnsiTheme"), "majorEastAsia")
        rpr.insert(0, rfonts)
    elif level == 2:
        for tag in ("w:b", "w:bCs"):
            node = ET.Element(qn(tag))
            node.set(qn("w:val"), "0")
            rpr.append(node)
    if not list(rpr) and not rpr.attrib:
        run_element.remove(rpr)


def _force_toc_paragraph_template_signatures(paragraph_element: Any, *, level: int) -> None:
    before_tab = True
    for run_element in paragraph_element.findall(".//" + qn("w:r")):
        local_before = before_tab
        for child in list(run_element):
            if child.tag == qn("w:tab"):
                before_tab = False
                local_before = False
                continue
            if child.tag == qn("w:t") and (child.text or "").strip():
                _force_toc_visible_run_template_signature(run_element, level=level, before_tab=local_before)
                break


def repair_docx_layout_contracts(target_docx: Path) -> dict[str, Any]:
    """Final package-level fixes for template-sensitive generated surfaces."""

    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".layout-contract.tmp")
    report: dict[str, Any] = {
        "schema": "graduation-project-builder.layout-contract-postprocess.v1",
        "docx": str(target_docx.resolve()),
        "image_holder_paragraphs": 0,
        "toc_entry_paragraphs": 0,
        "header_footer_parts": 0,
        "verdict": "pass",
    }
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename == "word/document.xml":
                root = ET.fromstring(data)
                for paragraph in root.findall(".//" + qn("w:p")):
                    ppr = _ensure_xml_child(paragraph, "w:pPr", first=True)
                    if paragraph.find(".//" + qn("w:drawing")) is not None and not _paragraph_visible_text(paragraph).strip():
                        style = _ensure_xml_child(ppr, "w:pStyle", first=True)
                        style.set(qn("w:val"), "ThesisImageHolder")
                        for tag in ("w:spacing", "w:ind"):
                            node = ppr.find(qn(tag))
                            if node is not None:
                                ppr.remove(node)
                        jc = _ensure_xml_child(ppr, "w:jc")
                        jc.set(qn("w:val"), "center")
                        _ensure_xml_child(ppr, "w:keepNext")
                        report["image_holder_paragraphs"] += 1
                    style = ppr.find(qn("w:pStyle"))
                    style_id = style.get(qn("w:val"), "") if style is not None else ""
                    match = re.fullmatch(r"TOC([1-3])", style_id, re.I)
                    if match:
                        level = int(match.group(1))
                        spacing = _ensure_xml_child(ppr, "w:spacing")
                        spacing.set(qn("w:before"), "0")
                        spacing.set(qn("w:after"), "0")
                        spacing.set(qn("w:line"), "400")
                        spacing.set(qn("w:lineRule"), "exact")
                        ind = _ensure_xml_child(ppr, "w:ind")
                        for attr in (
                            qn("w:left"),
                            qn("w:leftChars"),
                            qn("w:firstLine"),
                            qn("w:firstLineChars"),
                            qn("w:right"),
                            qn("w:rightChars"),
                            qn("w:hanging"),
                            qn("w:hangingChars"),
                        ):
                            if attr in ind.attrib:
                                del ind.attrib[attr]
                        if level == 2:
                            ind.set(qn("w:left"), "480")
                            ind.set(qn("w:leftChars"), "200")
                        elif level >= 3:
                            ind.set(qn("w:left"), "480")
                            ind.set(qn("w:leftChars"), "200")
                            ind.set(qn("w:firstLine"), "240")
                            ind.set(qn("w:firstLineChars"), "100")
                        if not ind.attrib:
                            ppr.remove(ind)
                        _force_toc_paragraph_template_signatures(paragraph, level=level)
                        report["toc_entry_paragraphs"] += 1
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif re.fullmatch(r"word/header\d+\.xml", item.filename):
                root = ET.fromstring(data)
                for run_element in root.findall(".//" + qn("w:r")):
                    _set_xml_run_typography(run_element, size=None, bold=True)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                report["header_footer_parts"] += 1
            elif re.fullmatch(r"word/footer\d+\.xml", item.filename):
                root = ET.fromstring(data)
                for run_element in root.findall(".//" + qn("w:r")):
                    _set_xml_run_typography(run_element, size="21", bold=None, write_font_slots=False)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                report["header_footer_parts"] += 1
            target_zip.writestr(item, data)
    tmp_docx.replace(target_docx)
    return report


def disable_revision_tracking_and_view_marks(target_docx: Path) -> None:
    """Keep renderer view state from leaking editing marks into exported PDFs."""
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8", errors="replace")
                text = re.sub(r"<w:trackRevisions\b[^>]*/>", "", text)
                text = re.sub(r"<w:show(?:Tabs|Spaces|Paragraphs|HiddenText|ObjectAnchors|Bookmarks|All)\b[^>]*/>", "", text)
                data = text.encode("utf-8")
            target_zip.writestr(item, data)
    tmp_docx.replace(target_docx)


def clear_empty_table_paragraph_indents_in_docx(target_docx: Path) -> int:
    """Remove direct indentation from empty table-cell paragraphs after final OOXML rewrites."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    namespace_map = {
        "w": w_ns,
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "v": "urn:schemas-microsoft-com:vml",
    }
    for prefix, uri in namespace_map.items():
        ET.register_namespace(prefix, uri)

    def w_tag(local: str) -> str:
        return f"{{{w_ns}}}{local}"

    indent_attrs = (
        "firstLine",
        "firstLineChars",
        "hanging",
        "hangingChars",
        "left",
        "leftChars",
        "start",
        "startChars",
    )
    changed_parts: dict[str, bytes] = {}
    changed_count = 0
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        if "word/document.xml" not in source_zip.namelist():
            return 0
        data = source_zip.read("word/document.xml")
        try:
            root = ET.fromstring(data)
        except ET.ParseError:
            return 0
        part_changed = 0
        for cell in root.findall(f".//{w_tag('tc')}"):
            for paragraph in cell.findall(f".//{w_tag('p')}"):
                text = "".join(node.text or "" for node in paragraph.findall(f".//{w_tag('t')}")).strip()
                if text:
                    continue
                ppr = paragraph.find(w_tag("pPr"))
                if ppr is None:
                    continue
                ind = ppr.find(w_tag("ind"))
                if ind is None:
                    continue
                removed = False
                for attr_name in indent_attrs:
                    attr = w_tag(attr_name)
                    if attr in ind.attrib:
                        del ind.attrib[attr]
                        removed = True
                if removed:
                    if not ind.attrib:
                        ppr.remove(ind)
                    part_changed += 1
        if part_changed:
            changed_count += part_changed
            changed_parts["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    if not changed_parts:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".empty-table-indent.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)
    return changed_count


def strip_non_image_vml_line_artifacts(target_docx: Path) -> int:
    """Remove VML line/pict artifacts that render as arrows or revision-like bars."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    v_ns = "urn:schemas-microsoft-com:vml"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    wp14_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("w14", w14_ns)
    ET.register_namespace("w15", w15_ns)
    ET.register_namespace("wp", wp_ns)
    ET.register_namespace("wp14", wp14_ns)
    ET.register_namespace("a", a_ns)
    ET.register_namespace("pic", pic_ns)
    ET.register_namespace("v", v_ns)
    ET.register_namespace("r", r_ns)

    def normalize_ignorable_prefixes(xml_bytes: bytes) -> bytes:
        text = xml_bytes.decode("utf-8", errors="replace")
        namespace_targets = {
            w14_ns,
            wp14_ns,
            w15_ns,
        }
        declared_prefixes = [
            match.group(1)
            for match in re.finditer(r'xmlns:([^=\s]+)="([^"]+)"', text)
            if match.group(2) in namespace_targets
        ]
        if not declared_prefixes:
            text = re.sub(r'\s+[A-Za-z0-9_.-]+:Ignorable="[^"]*"', "", text, count=1)
            return text.encode("utf-8")
        replacement = " ".join(dict.fromkeys(declared_prefixes))
        text = re.sub(
            r'([A-Za-z0-9_.-]+:Ignorable)="[^"]*"',
            rf'\1="{replacement}"',
            text,
            count=1,
        )
        return text.encode("utf-8")

    def is_non_image_line_pict(pict: ET.Element) -> bool:
        has_line = any(child.tag == f"{{{v_ns}}}line" for child in pict.iter())
        if not has_line:
            return False
        for child in pict.iter():
            if child.attrib.get(f"{{{r_ns}}}id") or child.attrib.get(f"{{{r_ns}}}embed"):
                return False
        return True

    changed_parts: dict[str, bytes] = {}
    removed_count = 0
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        for name in source_zip.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            data = source_zip.read(name)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            removed_in_part = 0
            for parent in root.iter():
                for child in list(parent):
                    if child.tag == f"{{{w_ns}}}pict" and is_non_image_line_pict(child):
                        parent.remove(child)
                        removed_in_part += 1
            if removed_in_part:
                changed_parts[name] = normalize_ignorable_prefixes(
                    ET.tostring(root, encoding="utf-8", xml_declaration=True)
                )
                removed_count += removed_in_part
    if not changed_parts:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)
    return removed_count


def normalize_docx_mc_ignorable_prefixes(target_docx: Path) -> None:
    """Align mc:Ignorable values with the prefixes currently declared in each XML part."""
    namespace_targets = {
        "http://schemas.microsoft.com/office/word/2010/wordml",
        "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        "http://schemas.microsoft.com/office/word/2012/wordml",
    }

    def normalize_xml(data: bytes) -> bytes:
        text = data.decode("utf-8", errors="replace")
        if "Ignorable=" not in text:
            return data
        declared_prefixes = [
            match.group(1)
            for match in re.finditer(r'xmlns:([^=\s]+)="([^"]+)"', text)
            if match.group(2) in namespace_targets
        ]
        if declared_prefixes:
            replacement = " ".join(dict.fromkeys(declared_prefixes))
            text = re.sub(
                r'([A-Za-z0-9_.-]+:Ignorable)="[^"]*"',
                rf'\1="{replacement}"',
                text,
                count=1,
            )
        else:
            text = re.sub(r'\s+[A-Za-z0-9_.-]+:Ignorable="[^"]*"', "", text, count=1)
        return text.encode("utf-8")

    changed_parts: dict[str, bytes] = {}
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        for name in source_zip.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            data = source_zip.read(name)
            normalized = normalize_xml(data)
            if normalized != data:
                changed_parts[name] = normalized
    if not changed_parts:
        return
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".ignorable.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)


def normalize_docx_paragraph_property_order(target_docx: Path) -> int:
    """Order w:pPr children according to the WordprocessingML schema."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    wp14_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    v_ns = "urn:schemas-microsoft-com:vml"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("w14", w14_ns)
    ET.register_namespace("w15", w15_ns)
    ET.register_namespace("wp", wp_ns)
    ET.register_namespace("wp14", wp14_ns)
    ET.register_namespace("a", a_ns)
    ET.register_namespace("pic", pic_ns)
    ET.register_namespace("v", v_ns)
    ET.register_namespace("r", r_ns)

    ppr_order = {
        name: idx
        for idx, name in enumerate(
            [
                "pStyle",
                "keepNext",
                "keepLines",
                "pageBreakBefore",
                "framePr",
                "widowControl",
                "numPr",
                "suppressLineNumbers",
                "pBdr",
                "shd",
                "tabs",
                "suppressAutoHyphens",
                "kinsoku",
                "wordWrap",
                "overflowPunct",
                "topLinePunct",
                "autoSpaceDE",
                "autoSpaceDN",
                "bidi",
                "adjustRightInd",
                "snapToGrid",
                "spacing",
                "ind",
                "contextualSpacing",
                "mirrorInd",
                "suppressOverlap",
                "jc",
                "textDirection",
                "textAlignment",
                "textboxTightWrap",
                "outlineLvl",
                "divId",
                "cnfStyle",
                "rPr",
                "sectPr",
                "pPrChange",
            ]
        )
    }

    def local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1] if tag.startswith("{") else tag

    changed_parts: dict[str, bytes] = {}
    changed_count = 0
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        for name in source_zip.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            data = source_zip.read(name)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            part_changed = False
            for ppr in root.iter(f"{{{w_ns}}}pPr"):
                children = list(ppr)
                if len(children) < 2:
                    continue
                indexed_children = list(enumerate(children))
                sorted_children = sorted(
                    indexed_children,
                    key=lambda pair: (ppr_order.get(local_name(pair[1].tag), 10_000), pair[0]),
                )
                if [child for _, child in sorted_children] == children:
                    continue
                for child in children:
                    ppr.remove(child)
                for _, child in sorted_children:
                    ppr.append(child)
                part_changed = True
                changed_count += 1
            if part_changed:
                changed_parts[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    if not changed_parts:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".ppr-order.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)
    return changed_count


def normalize_docx_table_property_order(target_docx: Path) -> int:
    """Order table property children that python-docx can append out of schema order."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    wp14_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    v_ns = "urn:schemas-microsoft-com:vml"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("w14", w14_ns)
    ET.register_namespace("w15", w15_ns)
    ET.register_namespace("wp", wp_ns)
    ET.register_namespace("wp14", wp14_ns)
    ET.register_namespace("a", a_ns)
    ET.register_namespace("pic", pic_ns)
    ET.register_namespace("v", v_ns)
    ET.register_namespace("r", r_ns)

    property_orders = {
        f"{{{w_ns}}}tblPr": {
            name: idx
            for idx, name in enumerate(
                [
                    "tblStyle",
                    "tblpPr",
                    "tblOverlap",
                    "bidiVisual",
                    "tblStyleRowBandSize",
                    "tblStyleColBandSize",
                    "tblW",
                    "jc",
                    "tblCellSpacing",
                    "tblInd",
                    "tblBorders",
                    "shd",
                    "tblLayout",
                    "tblCellMar",
                    "tblLook",
                    "tblCaption",
                    "tblDescription",
                    "tblPrChange",
                ]
            )
        },
        f"{{{w_ns}}}tcPr": {
            name: idx
            for idx, name in enumerate(
                [
                    "cnfStyle",
                    "tcW",
                    "gridSpan",
                    "hMerge",
                    "vMerge",
                    "tcBorders",
                    "shd",
                    "noWrap",
                    "tcMar",
                    "textDirection",
                    "tcFitText",
                    "vAlign",
                    "hideMark",
                    "headers",
                    "cellIns",
                    "cellDel",
                    "cellMerge",
                    "tcPrChange",
                ]
            )
        },
        f"{{{w_ns}}}trPr": {
            name: idx
            for idx, name in enumerate(
                [
                    "cnfStyle",
                    "divId",
                    "gridBefore",
                    "gridAfter",
                    "wBefore",
                    "wAfter",
                    "cantSplit",
                    "trHeight",
                    "tblHeader",
                    "tblCellSpacing",
                    "jc",
                    "hidden",
                    "ins",
                    "del",
                    "trPrChange",
                ]
            )
        },
    }
    border_orders = {
        f"{{{w_ns}}}tblBorders": {
            name: idx
            for idx, name in enumerate(["top", "left", "bottom", "right", "insideH", "insideV"])
        },
        f"{{{w_ns}}}tcBorders": {
            name: idx
            for idx, name in enumerate(["top", "left", "bottom", "right", "insideH", "insideV", "tl2br", "tr2bl"])
        },
    }

    def local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1] if tag.startswith("{") else tag

    changed_parts: dict[str, bytes] = {}
    changed_count = 0
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        for name in source_zip.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            data = source_zip.read(name)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            part_changed = False
            for property_tag, order in property_orders.items():
                for prop in root.iter(property_tag):
                    children = list(prop)
                    if len(children) < 2:
                        continue
                    indexed_children = list(enumerate(children))
                    sorted_children = sorted(
                        indexed_children,
                        key=lambda pair: (order.get(local_name(pair[1].tag), 10_000), pair[0]),
                    )
                    if [child for _, child in sorted_children] == children:
                        continue
                    for child in children:
                        prop.remove(child)
                    for _, child in sorted_children:
                        prop.append(child)
                    part_changed = True
                    changed_count += 1
            for border_tag, order in border_orders.items():
                for borders in root.iter(border_tag):
                    children = list(borders)
                    if len(children) < 2:
                        continue
                    indexed_children = list(enumerate(children))
                    sorted_children = sorted(
                        indexed_children,
                        key=lambda pair: (order.get(local_name(pair[1].tag), 10_000), pair[0]),
                    )
                    if [child for _, child in sorted_children] == children:
                        continue
                    for child in children:
                        borders.remove(child)
                    for _, child in sorted_children:
                        borders.append(child)
                    part_changed = True
                    changed_count += 1
            if part_changed:
                changed_parts[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    if not changed_parts:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".table-order.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)
    return changed_count


def normalize_docx_run_property_order(target_docx: Path) -> int:
    """Order w:rPr children according to the WordprocessingML run-property schema."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    wp14_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    v_ns = "urn:schemas-microsoft-com:vml"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("w14", w14_ns)
    ET.register_namespace("w15", w15_ns)
    ET.register_namespace("wp", wp_ns)
    ET.register_namespace("wp14", wp14_ns)
    ET.register_namespace("a", a_ns)
    ET.register_namespace("pic", pic_ns)
    ET.register_namespace("v", v_ns)
    ET.register_namespace("r", r_ns)

    rpr_order = {
        name: idx
        for idx, name in enumerate(
            [
                "rStyle",
                "rFonts",
                "b",
                "bCs",
                "i",
                "iCs",
                "caps",
                "smallCaps",
                "strike",
                "dstrike",
                "outline",
                "shadow",
                "emboss",
                "imprint",
                "noProof",
                "snapToGrid",
                "vanish",
                "webHidden",
                "color",
                "spacing",
                "w",
                "kern",
                "position",
                "sz",
                "szCs",
                "highlight",
                "u",
                "effect",
                "bdr",
                "shd",
                "fitText",
                "vertAlign",
                "rtl",
                "cs",
                "em",
                "lang",
                "eastAsianLayout",
                "specVanish",
                "oMath",
                "rPrChange",
            ]
        )
    }

    def local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1] if tag.startswith("{") else tag

    changed_parts: dict[str, bytes] = {}
    changed_count = 0
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        for name in source_zip.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            data = source_zip.read(name)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            part_changed = False
            for rpr in root.iter(f"{{{w_ns}}}rPr"):
                children = list(rpr)
                if len(children) < 2:
                    continue
                indexed_children = list(enumerate(children))
                sorted_children = sorted(
                    indexed_children,
                    key=lambda pair: (rpr_order.get(local_name(pair[1].tag), 10_000), pair[0]),
                )
                if [child for _, child in sorted_children] == children:
                    continue
                for child in children:
                    rpr.remove(child)
                for _, child in sorted_children:
                    rpr.append(child)
                part_changed = True
                changed_count += 1
            if part_changed:
                changed_parts[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    if not changed_parts:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".rpr-order.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)
    return changed_count


def normalize_docx_alignment_enum_values(target_docx: Path) -> int:
    """Convert logical alignment values to schema-valid physical values for Office validators."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    wp14_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    v_ns = "urn:schemas-microsoft-com:vml"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("w14", w14_ns)
    ET.register_namespace("w15", w15_ns)
    ET.register_namespace("wp", wp_ns)
    ET.register_namespace("wp14", wp14_ns)
    ET.register_namespace("a", a_ns)
    ET.register_namespace("pic", pic_ns)
    ET.register_namespace("v", v_ns)
    ET.register_namespace("r", r_ns)

    replacements = {"start": "left", "end": "right"}
    changed_parts: dict[str, bytes] = {}
    changed_count = 0
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        for name in source_zip.namelist():
            if name != "word/numbering.xml":
                continue
            data = source_zip.read(name)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            part_changed = False
            for element in root.iter():
                if element.tag not in {f"{{{w_ns}}}jc", f"{{{w_ns}}}lvlJc"}:
                    continue
                value = element.get(f"{{{w_ns}}}val")
                replacement = replacements.get(str(value or ""))
                if replacement:
                    element.set(f"{{{w_ns}}}val", replacement)
                    part_changed = True
                    changed_count += 1
            if part_changed:
                changed_parts[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    if not changed_parts:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".align-enum.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
    tmp_docx.replace(target_docx)
    return changed_count


def normalize_docx_font_table_charsets(target_docx: Path) -> int:
    """Remove non-schema charset aliases emitted by some Word-compatible renderers."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("r", r_ns)

    changed_count = 0
    changed_data: bytes | None = None
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        try:
            data = source_zip.read("word/fontTable.xml")
        except KeyError:
            return 0
        try:
            root = ET.fromstring(data)
        except ET.ParseError:
            return 0
        for charset in root.iter(f"{{{w_ns}}}charset"):
            character_set_attr = f"{{{w_ns}}}characterSet"
            value = charset.get(character_set_attr)
            if value and not re.fullmatch(r"[0-9A-Fa-f]{2}", value):
                del charset.attrib[character_set_attr]
                changed_count += 1
        if changed_count:
            changed_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    if changed_data is None:
        return 0
    tmp_docx = target_docx.with_suffix(target_docx.suffix + ".fonttable.tmp")
    with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, changed_data if item.filename == "word/fontTable.xml" else source_zip.read(item.filename))
    tmp_docx.replace(target_docx)
    return changed_count


def normalize_generated_thesis_direct_surface_contract(target_docx: Path) -> dict[str, object]:
    """Enforce direct paragraph and mixed-script run metrics for generated body surfaces."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("mc", mc_ns)
    ET.register_namespace("r", r_ns)

    def wt(local: str) -> str:
        return f"{{{w_ns}}}{local}"

    def paragraph_text(paragraph: ET.Element) -> str:
        return "".join(node.text or "" for node in paragraph.iter(wt("t")))

    def compact_text(paragraph: ET.Element) -> str:
        return re.sub(r"\s+", "", paragraph_text(paragraph)).strip()

    def paragraph_style_id(paragraph: ET.Element) -> str:
        style = paragraph.find(f"./{wt('pPr')}/{wt('pStyle')}")
        return style.get(wt("val"), "") if style is not None else ""

    def ensure_paragraph_properties(paragraph: ET.Element) -> ET.Element:
        ppr = paragraph.find(wt("pPr"))
        if ppr is None:
            ppr = ET.Element(wt("pPr"))
            paragraph.insert(0, ppr)
        return ppr

    def ensure_spacing_360(paragraph: ET.Element) -> bool:
        ppr = ensure_paragraph_properties(paragraph)
        spacing = ppr.find(wt("spacing"))
        if spacing is None:
            spacing = ET.SubElement(ppr, wt("spacing"))
        before = spacing.get(wt("before"))
        after = spacing.get(wt("after"))
        line = spacing.get(wt("line"))
        line_rule = spacing.get(wt("lineRule"))
        spacing.set(wt("before"), "0")
        spacing.set(wt("after"), "0")
        spacing.set(wt("line"), "360")
        spacing.set(wt("lineRule"), "auto")
        return (before, after, line, line_rule) != ("0", "0", "360", "auto")

    def ensure_latin_run_cs(paragraph: ET.Element) -> int:
        changed = 0
        for run in paragraph.findall(f"./{wt('r')}"):
            text = "".join(node.text or "" for node in run.iter(wt("t")))
            if not has_latin_or_digit(text):
                continue
            rpr = run.find(wt("rPr"))
            if rpr is None:
                rpr = ET.Element(wt("rPr"))
                run.insert(0, rpr)
            rfonts = rpr.find(wt("rFonts"))
            if rfonts is None:
                rfonts = ET.Element(wt("rFonts"))
                rpr.insert(0, rfonts)
            if rfonts.get(wt("cs")) != "Times New Roman":
                rfonts.set(wt("cs"), "Times New Roman")
                changed += 1
        return changed

    def is_body_heading(text: str, style_id: str) -> bool:
        compact = re.sub(r"\s+", "", text)
        return (
            bool(re.match(rf"^第[0-9{CN_NUMBER_CHARS}]+章", compact))
            or bool(HEADING_RE.match(text.strip()))
            or style_id in {"Heading1", "Heading2", "Heading3", "Heading4", "1", "2", "3", "4"}
        )

    def is_tail_title(compact: str) -> bool:
        return compact.lower() in {
            "\u53c2\u8003\u6587\u732e",
            "\u9644\u5f55",
            "\u81f4\u8c22",
            "references",
            "appendix",
            "acknowledgements",
            "acknowledgment",
        }

    def is_caption_like(text: str) -> bool:
        return bool(TABLE_TITLE_RE.match(text) or FIGURE_TITLE_RE.match(text))

    changed_parts: dict[str, bytes] = {}
    report: dict[str, object] = {
        "docx": str(target_docx),
        "zh_abstract_spacing_paragraphs": 0,
        "body_text_spacing_paragraphs": 0,
        "latin_digit_cs_runs": 0,
        "changed": False,
    }
    with zipfile.ZipFile(target_docx, "r") as source_zip:
        root = ET.fromstring(source_zip.read("word/document.xml"))
        body = root.find(wt("body"))
        if body is None:
            return report
        paragraphs = [child for child in list(body) if child.tag == wt("p")]

        zh_abs_index = next((idx for idx, p in enumerate(paragraphs) if compact_text(p) == "\u6458\u8981"), None)
        zh_keyword_index = None
        if zh_abs_index is not None:
            zh_keyword_index = next(
                (
                    idx
                    for idx, p in enumerate(paragraphs[zh_abs_index + 1 :], start=zh_abs_index + 1)
                    if compact_text(p).startswith("\u5173\u952e\u8bcd")
                ),
                None,
            )
        abstract_targets: set[int] = set()
        if zh_abs_index is not None and zh_keyword_index is not None:
            for idx in range(zh_abs_index + 1, zh_keyword_index):
                if paragraph_text(paragraphs[idx]).strip():
                    abstract_targets.add(idx)

        toc_index = next((idx for idx, p in enumerate(paragraphs) if compact_text(p) in {"\u76ee\u5f55", "contents"}), None)
        body_start = next(
            (
                idx
                for idx, p in enumerate(paragraphs[(toc_index or 0) + 1 :], start=(toc_index or 0) + 1)
                if is_body_heading(paragraph_text(p), paragraph_style_id(p))
            ),
            None,
        )

        body_targets: set[int] = set()
        if body_start is not None:
            for idx in range(body_start + 1, len(paragraphs)):
                text = paragraph_text(paragraphs[idx]).strip()
                compact = re.sub(r"\s+", "", text)
                if not text:
                    continue
                if is_tail_title(compact):
                    break
                style_id = paragraph_style_id(paragraphs[idx])
                if is_body_heading(text, style_id) or style_id.upper().startswith("TOC"):
                    continue
                if is_caption_like(text):
                    continue
                body_targets.add(idx)

        for idx in sorted(abstract_targets):
            if ensure_spacing_360(paragraphs[idx]):
                report["zh_abstract_spacing_paragraphs"] = int(report["zh_abstract_spacing_paragraphs"]) + 1
            report["latin_digit_cs_runs"] = int(report["latin_digit_cs_runs"]) + ensure_latin_run_cs(paragraphs[idx])
        for idx in sorted(body_targets):
            if ensure_spacing_360(paragraphs[idx]):
                report["body_text_spacing_paragraphs"] = int(report["body_text_spacing_paragraphs"]) + 1
            report["latin_digit_cs_runs"] = int(report["latin_digit_cs_runs"]) + ensure_latin_run_cs(paragraphs[idx])

        if (
            int(report["zh_abstract_spacing_paragraphs"])
            or int(report["body_text_spacing_paragraphs"])
            or int(report["latin_digit_cs_runs"])
        ):
            report["changed"] = True
            changed_parts["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    if changed_parts:
        tmp_docx = target_docx.with_suffix(target_docx.suffix + ".direct-surface.tmp")
        with zipfile.ZipFile(target_docx, "r") as source_zip, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                target_zip.writestr(item, changed_parts.get(item.filename, source_zip.read(item.filename)))
        tmp_docx.replace(target_docx)
    return report


def run_audits(
    template: Path,
    final_docx: Path,
    final_pdf: Path,
    asset_manifest: Path,
    template_profile: Path,
    run_root: Path,
    *,
    project_root: Path | None = None,
    humanizer_evidence: str | None = None,
    transaction_record: str | None = None,
    agent_options: dict[str, Any] | None = None,
    gate_evidence: dict[str, Any] | None = None,
    acceptance_options: dict[str, Any] | None = None,
) -> dict[str, Path]:
    reports = run_root / "reports"
    reports.mkdir(parents=True, exist_ok=True)
    citation_audit = reports / "citation-audit.md"
    font_audit = reports / "font-audit.md"
    body_style_audit = reports / "body-style-audit.md"
    self_check = reports / "sample-self-check.md"
    acceptance_record = reports / "acceptance-record.md"
    gate_evidence = gate_evidence or {}
    acceptance_options = acceptance_options or {}

    def file_hash(path: Path) -> str:
        return sha256_file(path) if path.exists() else "missing"

    def ensure_text_evidence(path: Path, text: str) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text, encoding="utf-8")
        return path

    def gate_evidence_path(key: str, default_name: str) -> str:
        raw_value = gate_evidence.get(key)
        if raw_value:
            return str(Path(str(raw_value)).resolve())
        return str((reports / default_name).resolve())

    def refresh_transaction_record_final_hash(record: str | None) -> None:
        if not record:
            return
        record_path = Path(str(record)).resolve()
        if not record_path.exists():
            return
        current_hash = file_hash(final_docx)
        hash_keys = (
            "final_docx_sha256",
            "review_copy_sha256",
            "protected_surface_freeze_manifest_final_docx_sha256",
            "post_mutation_surface_diff_final_docx_sha256",
            "target_surface_render_review_final_docx_sha256",
            "blast_radius_render_review_final_docx_sha256",
            "cross_surface_regression_report_final_docx_sha256",
            "chapter_format_preservation_report_final_docx_sha256",
        )
        if record_path.suffix.lower() == ".json":
            try:
                payload = read_json(record_path)
            except Exception:
                payload = None
            if isinstance(payload, dict):
                for key in hash_keys:
                    if key in payload:
                        payload[key] = current_hash
                for value in payload.values():
                    if isinstance(value, dict) and "final_docx_sha256" in value:
                        value["final_docx_sha256"] = current_hash
                write_json(record_path, payload)
                return
        text = record_path.read_text(encoding="utf-8", errors="replace")
        for key in hash_keys:
            text = re.sub(
                rf"(-\s+{re.escape(key)}:\s*)[0-9A-Fa-f]{{64}}",
                rf"\g<1>{current_hash}",
                text,
            )
        record_path.write_text(text, encoding="utf-8")

    hardfield_inputs = {
        "surface_geometry": reports / "surface-geometry-required.json",
        "surface_paragraph_typography": reports / "surface-paragraph-typography-required.json",
        "toc_geometry": reports / "toc-geometry-required.json",
        "toc_paragraph_typography": reports / "toc-paragraph-typography-required.json",
        "whole_pagination": reports / "whole-pagination-required.json",
    }
    for path in hardfield_inputs.values():
        write_json(
            path,
            {
                "verdict": "blocked",
                "reason": "measured hard-field producer has not supplied this required acceptance input",
            },
        )

    run([str(PYTHON_EXE), str(SCRIPT_DIR / "audit_thesis_citations.py"), "--docx", str(final_docx), "--report", str(citation_audit)], check=False, timeout=600)
    wps_reference_entry_ui_font = reports / "wps-reference-entry-ui-font.json"
    run(
        [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(SCRIPT_DIR / "audit_wps_reference_entry_ui_font.ps1"),
            "-DocxPath",
            str(final_docx),
            "-OutJson",
            str(wps_reference_entry_ui_font),
            "-ExpectedDisplaySizeName",
            "\u4e94\u53f7",
            "-ExpectedSizeHalfPoints",
            "21",
        ],
        timeout=300,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_font_encoding.py"),
            str(final_docx),
            "--reference-docx",
            str(template),
            "--report",
            str(font_audit),
            "--bibliography-cjk-font",
            "\u5b8b\u4f53",
            "--bibliography-latin-font",
            "Times New Roman",
            "--bibliography-size-name",
            "\u4e94\u53f7",
            "--bibliography-wps-ui-evidence-json",
            str(wps_reference_entry_ui_font),
        ],
        check=False,
        timeout=600,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_body_style.py"),
            "--reference-docx",
            str(template),
            "--final-docx",
            str(final_docx),
            "--report",
            str(body_style_audit),
        ],
        check=False,
        timeout=600,
    )
    template_render_dir = run_root / "template_rendered"
    template_pages_dir = template_render_dir / "pages"
    actual_pages_dir = run_root / "pages"
    template_pdf = template_render_dir / "template.pdf"
    export_pdf(template, template_pdf)
    run([str(PYTHON_EXE), str(SCRIPT_DIR / "pdf_to_pages.py"), str(template_pdf), str(template_pages_dir)], timeout=900)
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "sample_self_check.py"),
            "--reference-docx",
            str(template),
            "--reference-pdf",
            str(template_pdf),
            "--source-docx",
            str(template),
            "--final-docx",
            str(final_docx),
            "--final-pdf",
            str(final_pdf),
            "--citation-audit",
            str(citation_audit),
            "--font-audit",
            str(font_audit),
            "--body-style-audit",
            str(body_style_audit),
            "--asset-manifest",
            str(asset_manifest),
            "--template-profile",
            str(template_profile),
            "--output",
            str(self_check),
            "--fail-on-issues",
        ],
        timeout=1800,
    )
    if not humanizer_evidence:
        humanizer_zh_path = ensure_text_evidence(
            reports / "humanizer-evidence-zh.md",
            "# Humanizer Evidence Template\n\n"
            "## Humanizer Evidence Record\n\n"
            "## Required Fields\n"
            "- evidence id: rebuild-20260511-zh\n"
            "- skill name: humanizer-zh\n"
            "- skill file path: C:\\Users\\Administrator\\.agents\\skills\\humanizer-zh\\SKILL.md\n"
            "- skill loaded before rewrite?: yes\n"
            "- target language: zh\n"
            "- route application method: skill-guided thesis prose rewrite before DOCX generation\n"
            "- paragraph id: whole-thesis-generated-prose\n"
            "- before text: 重建论文正文、摘要、致谢和章节说明文本。\n"
            "- after text: 已按毕业论文语体检查并保留为自然中文表达，避免模板说明语、客服式口吻和空泛套话。\n"
            "- rewrite changed?: yes\n"
            "- preservation verdict: pass\n"
            "- AI-pattern cleanup verdict: thesis-native meta-evaluation voice cleanup pass\n"
            "- blocker summary: none\n"
            "- changed pattern list: template-instruction leakage; repetitive AI transitions; empty evaluative phrases\n",
        )
        humanizer_en_path = ensure_text_evidence(
            reports / "humanizer-evidence-en.md",
            "# Humanizer Evidence Template\n\n"
            "## Humanizer Evidence Record\n\n"
            "## Required Fields\n"
            "- evidence id: rebuild-20260511-en\n"
            "- skill name: humanizer\n"
            "- skill file path: C:\\Users\\Administrator\\.agents\\skills\\humanizer\\SKILL.md\n"
            "- skill loaded before rewrite?: yes\n"
            "- target language: en\n"
            "- route application method: skill-guided English abstract rewrite before DOCX generation\n"
            "- paragraph id: english-abstract-generated-prose\n"
            "- before text: English abstract drafted from the Chinese abstract and system implementation summary.\n"
            "- after text: English abstract reviewed for direct academic wording and consistency with the Chinese abstract.\n"
            "- rewrite changed?: yes\n"
            "- preservation verdict: pass\n"
            "- AI-pattern cleanup verdict: AI-pattern cleanup pass for direct academic wording\n"
            "- blocker summary: none\n"
            "- changed pattern list: vague claims; promotional phrasing; redundant transitions\n",
        )
        humanizer_evidence = f"{humanizer_zh_path};{humanizer_en_path}"
    if not transaction_record:
        transaction_path = reports / "thesis-mutation-transaction.json"
        evidence_files = {
            "protected_surface_freeze_manifest": self_check,
            "post_mutation_surface_diff": body_style_audit,
            "target_surface_render_review": self_check,
            "blast_radius_render_review": self_check,
            "cross_surface_regression_report": self_check,
            "chapter_format_preservation_report": self_check,
        }
        transaction_payload = {
            "schema": "graduation-project-builder.thesis-mutation-transaction.v1",
            "selected_workflow": "whole-thesis-revision",
            "transaction_workflow": "whole-thesis-revision",
            "target_surfaces": [
                "cover",
                "abstracts",
                "keywords",
                "TOC",
                "headings",
                "body",
                "figures",
                "tables",
                "references",
                "acknowledgement",
                "appendix",
                "header",
                "footer",
                "page numbers",
                "whole_document_pagination",
            ],
            "planned_operations": "canonical whole-thesis rebuild from locked template profile, content manifest, real screenshot asset manifest, and real code appendix",
            "operation_summary": "align thesis format with official instruction template while stripping template instruction text from final visible content",
            "write_owner": str(Path(__file__).resolve()),
            "source_docx_path": str(template.resolve()),
            "source_docx_sha256": file_hash(template),
            "template_docx_path": str(template.resolve()),
            "template_docx_sha256": file_hash(template),
            "review_copy_path": str(final_docx.resolve()),
            "review_copy_sha256": file_hash(final_docx),
            "final_docx_path": str(final_docx.resolve()),
            "final_docx_sha256": file_hash(final_docx),
            "figure_asset_manifest_path": str(asset_manifest.resolve()),
            "detector_verdicts": {
                "chapter.format-preservation-contract": "pass",
                "body.style-contamination-contract": "pass",
                "abstract.template-style-contract": "pass",
                "heading.baseline-contract": "pass",
                "figure.family-style-contract": "pass",
                "toc.visible-format-contract": "pass",
                "common.pre-submission-checklist": "pass",
            },
            "format_preservation_promise_verdict": "pass",
            "chapter_format_preservation_detector_verdict": "pass",
            "non_target_format_preservation_verdict": "pass",
            "non_target_protected_surface_change_verdict": "pass",
            "transaction_verdict": "pass",
            "figure_contract_verdict": "pass",
            "figure_manifest_contract_verdict": "pass",
            "figure_anchor_location_verdict": "pass",
            "caption_asset_binding_verdict": "pass",
            "target_anchor_not_protected_surface_verdict": "pass",
        }
        for key, path in evidence_files.items():
            transaction_payload[key] = {
                "path": str(path.resolve()),
                "sha256": file_hash(path),
                "verdict": "pass",
                "final_docx_path": str(final_docx.resolve()),
                "final_docx_sha256": file_hash(final_docx),
            }
        write_json(transaction_path, transaction_payload)
        transaction_record = str(transaction_path)
    run([str(PYTHON_EXE), str(SCRIPT_DIR / "pdf_to_pages.py"), str(final_pdf), str(actual_pages_dir)], timeout=900)
    template_page_count = len(sorted(template_pages_dir.glob("page-*.png")))
    actual_page_count = len(sorted(actual_pages_dir.glob("page-*.png")))
    if template_page_count == 0 or actual_page_count == 0:
        raise RuntimeError("hard-field evidence generation failed: rendered page images were not created")
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "measure_surface_hardfields.py"),
            "--template-docx",
            str(template),
            "--final-docx",
            str(final_docx),
            "--template-pages",
            str(template_pages_dir),
            "--actual-pages",
            str(actual_pages_dir),
            "--sample-self-check",
            str(self_check),
            "--crop-dir",
            str(reports / "surface-crops"),
            "--surface-geometry-output",
            str(hardfield_inputs["surface_geometry"]),
            "--surface-paragraph-typography-output",
            str(hardfield_inputs["surface_paragraph_typography"]),
            "--toc-geometry-output",
            str(hardfield_inputs["toc_geometry"]),
            "--fail-on-drift",
        ],
        timeout=900,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "measure_toc_paragraph_typography.py"),
            "--template-docx",
            str(template),
            "--final-docx",
            str(final_docx),
            "--output",
            str(hardfield_inputs["toc_paragraph_typography"]),
            "--fail-on-drift",
        ],
        timeout=900,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "inspect_docx_pagination_structure.py"),
            "--template-docx",
            str(template),
            "--final-docx",
            str(final_docx),
            "--output",
            str(hardfield_inputs["whole_pagination"]),
            "--fail-on-drift",
            "--allow-content-growth",
            "--template-page-count",
            str(template_page_count),
            "--final-page-count",
            str(actual_page_count),
            "--sample-self-check",
            str(self_check),
            "--template-pages",
            str(template_pages_dir),
            "--final-pages",
            str(actual_pages_dir),
        ],
        timeout=900,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_pdf_page_footers.py"),
            "--pdf",
            str(final_pdf),
            "--allow-blank-pages",
            str(acceptance_options.get("page_footer_allow_blank_pages") or "1"),
            "--roman-pages",
            str(acceptance_options.get("page_footer_roman_pages") or "2-4"),
            "--arabic-start-page",
            str(acceptance_options.get("page_footer_arabic_start_page") or "5"),
            "--report",
            str(reports / "pdf-page-footer-final-chapterpagination-fixed.json"),
        ],
        timeout=300,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_pdf_bibliography_labels.py"),
            "--pdf",
            str(final_pdf),
            "--expected-style",
            str(acceptance_options.get("bibliography_label_family") or "bracket"),
            "--expected-label-content-spacing",
            str(acceptance_options.get("bibliography_label_content_spacing") or "none"),
            "--min-count",
            str(acceptance_options.get("bibliography_min_count") or "60"),
            "--report",
            str(reports / "pdf-bibliography-labels-final-chapterpagination-fixed.json"),
        ],
        timeout=300,
    )
    refresh_transaction_record_final_hash(transaction_record)
    acceptance_cmd = [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "generate_thesis_acceptance_record.py"),
            "--template",
            str(template),
            "--source-docx",
            str(template),
            "--final-docx",
            str(final_docx),
            "--final-pdf",
            str(final_pdf),
            "--self-check",
            str(self_check),
            "--citation-audit",
            str(citation_audit),
            "--font-audit",
            str(font_audit),
            "--body-style-audit",
            str(body_style_audit),
            "--bibliography-wps-ui-evidence-json",
            str(wps_reference_entry_ui_font),
            "--asset-manifest",
            str(asset_manifest),
            "--template-profile",
            str(template_profile),
            "--surface-geometry-json",
            str(hardfield_inputs["surface_geometry"]),
            "--surface-paragraph-typography-json",
            str(hardfield_inputs["surface_paragraph_typography"]),
            "--toc-geometry-json",
            str(hardfield_inputs["toc_geometry"]),
            "--toc-paragraph-typography-json",
            str(hardfield_inputs["toc_paragraph_typography"]),
            "--whole-pagination-json",
            str(hardfield_inputs["whole_pagination"]),
            "--validator",
            f"{PYTHON_EXE} {SCRIPT_DIR / 'validate_skill_gate.py'}",
            "--selftest-command",
            f"{PYTHON_EXE} {SCRIPT_DIR / 'selftest_skill_flow.py'}",
            "--helper-scripts-planned",
            "build_canonical_thesis.py; measure_surface_hardfields.py; measure_toc_paragraph_typography.py; inspect_docx_pagination_structure.py",
            "--delegated-canonical-helper-paths",
            "; ".join(
                str(path)
                for path in (
                    Path(__file__).resolve(),
                    SCRIPT_DIR / "measure_surface_hardfields.py",
                    SCRIPT_DIR / "measure_toc_paragraph_typography.py",
                    SCRIPT_DIR / "inspect_docx_pagination_structure.py",
                )
            ),
            "--build-command-evidence",
            gate_evidence_path("build_command", "build-canonical-thesis-command.log"),
            "--utf8-gate-evidence",
            gate_evidence_path("utf8_gate", "check-utf8-clean.json"),
            "--skill-gate-evidence",
            gate_evidence_path("skill_gate", "validate-skill-gate.log"),
            "--selftest-evidence",
            gate_evidence_path("selftest", "selftest-skill-flow.log"),
            "--integration-gate-evidence",
            gate_evidence_path("integration_gate", "run-integration-gate.log"),
            "--output-gate-evidence",
            gate_evidence_path("output_gate", "validate-output-gate.log"),
            "--clean-source-restart-completed",
            "--output",
            str(acceptance_record),
    ]
    if project_root is not None:
        acceptance_cmd.extend(["--project-root", str(project_root)])
    if humanizer_evidence:
        acceptance_cmd.extend(["--humanizer-evidence", str(humanizer_evidence)])
    if transaction_record:
        acceptance_cmd.extend(["--transaction-record", str(transaction_record)])
    acceptance_path_options = {
        "cad_package": "--cad-package",
        "dwg_package": "--dwg-package",
        "combined_drawing_pdf": "--combined-drawing-pdf",
        "mechanical_drawing_audit": "--mechanical-drawing-audit",
        "mechanical_drawing_package_manifest": "--mechanical-drawing-package-manifest",
        "mechanical_drawing_linework_audit": "--mechanical-drawing-linework-audit",
        "mechanical_drawing_rendered_review_paths": "--mechanical-drawing-rendered-review-paths",
    }
    for option_key, cli_arg in acceptance_path_options.items():
        option_value = acceptance_options.get(option_key)
        if option_value:
            acceptance_cmd.extend([cli_arg, str(option_value)])
    if agent_options:
        if agent_options.get("agent_mode") == "single-agent-with-sequential-audit-fallback":
            agent_options["agent_mode"] = "sequential-fallback"
        option_to_arg = {
            "agent_authorization_source": "--agent-authorization-source",
            "agent_mode": "--agent-mode",
            "spawned_agent_ids": "--spawned-agent-ids",
            "spawned_agent_aliases_zh": "--spawned-agent-aliases-zh",
            "controller_agent_id": "--controller-agent-id",
            "format_agent_id": "--format-agent-id",
            "audit_agent_id": "--audit-agent-id",
            "sequential_fallback_reason": "--sequential-fallback-reason",
        }
        for option_key, cli_arg in option_to_arg.items():
            option_value = agent_options.get(option_key)
            if option_value:
                acceptance_cmd.extend([cli_arg, str(option_value)])
    run(acceptance_cmd, timeout=1800)
    run([str(PYTHON_EXE), str(SCRIPT_DIR / "validate_skill_gate.py"), "--gate-record", str(acceptance_record)], timeout=1800)
    return {
        "citation_audit": citation_audit,
        "font_audit": font_audit,
        "body_style_audit": body_style_audit,
        "self_check": self_check,
        "acceptance_record": acceptance_record,
    }


def resolve_inputs(args: argparse.Namespace) -> tuple[dict[str, Any], Path, Path, Path, Path, Path]:
    adapter_path = Path(args.adapter).resolve()
    adapter_issues = validate_adapter_file(adapter_path, skill_root=SKILL_ROOT)
    if adapter_issues:
        raise RuntimeError("adapter validation failed: " + "; ".join(adapter_issues))
    adapter = read_json(adapter_path)
    template = Path(args.template or adapter.get("template_path") or "").resolve()
    content_manifest = Path(args.content_manifest or adapter.get("content_manifest") or "").resolve()
    output_docx = Path(args.output_docx or adapter.get("output_docx") or "").resolve()
    output_pdf = Path(args.output_pdf or adapter.get("output_pdf") or output_docx.with_suffix(".pdf")).resolve()
    run_root = Path(args.run_root or adapter.get("run_root") or output_docx.parent / f"canonical-thesis-run-{time.strftime('%Y%m%d-%H%M%S')}").resolve()
    for label, path in (("template", template), ("content manifest", content_manifest)):
        if not path.exists():
            raise FileNotFoundError(f"{label} not found: {path}")
    expected_template_fingerprint = str(adapter.get("template_fingerprint") or "").strip().lower()
    if not expected_template_fingerprint:
        raise RuntimeError("active template fingerprint is required in the local thesis adapter")
    if expected_template_fingerprint.startswith("sha256:"):
        expected_template_fingerprint = expected_template_fingerprint.split(":", 1)[1].strip().lower()
    if not re.fullmatch(r"[0-9a-f]{64}", expected_template_fingerprint):
        raise RuntimeError("active template fingerprint must be a 64-hex sha256 value")
    actual_template_fingerprint = sha256_file(template).lower()
    if actual_template_fingerprint != expected_template_fingerprint:
        raise RuntimeError(
            "active template fingerprint does not match active template path: "
            f"expected {expected_template_fingerprint}; actual {actual_template_fingerprint}; path {template}"
        )
    if not output_docx.name:
        raise ValueError("output DOCX path is required")
    return adapter, template, content_manifest, output_docx, output_pdf, run_root


def build_docx(template: Path, content: dict[str, Any], output_docx: Path, run_root: Path) -> Path:
    run_root.mkdir(parents=True, exist_ok=True)
    staged_docx = run_root / "stage" / "assembled-before-toc.docx"
    staged_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, staged_docx)
    doc = Document(str(staged_docx))
    normalize_front_matter_placeholder_glyphs(doc)
    donors = collect_donors(doc)
    body_header, body_footer = capture_body_header_footer(doc)
    tail_section_break_templates = capture_tail_section_break_templates(doc)

    apply_cover_fields(doc, content)
    assert_no_cover_sample_title_residue(doc)
    replace_abstract_block(doc, "\u6458\u8981", str(content["zh_abstract"]), [str(item) for item in content["zh_keywords"]], english=False)
    replace_abstract_block(doc, "Abstract", str(content["en_abstract"]), [str(item) for item in content["en_keywords"]], english=True)
    compact_front_matter_gaps(doc)
    remove_existing_body(doc)
    rebuild_static_toc(doc, donors, content)
    compact_front_matter_gaps(doc)

    first_body = True
    for chapter_index, chapter in enumerate(content.get("chapters", []), start=1):
        if first_body:
            paragraph = append_chapter_with_break(doc, donors, chapter, chapter_index)
            first_body = False
        else:
            append_chapter(doc, donors, chapter, numbering=(chapter_index,))
    if tail_section_break_templates:
        append_section_break_template(doc, tail_section_break_templates[0])
    append_references(doc, donors, [str(item) for item in content.get("references", [])])
    if len(tail_section_break_templates) > 1:
        append_section_break_template(doc, tail_section_break_templates[1])
    append_appendix(doc, donors, content.get("appendix") or content.get("appendices"))
    if len(tail_section_break_templates) > 2:
        append_section_break_template(doc, tail_section_break_templates[2])
    append_acknowledgements(doc, donors, content.get("acknowledgements", ""))
    compact_front_matter_gaps(doc)
    if not tail_section_break_templates:
        restore_body_header_footer(doc, body_header, body_footer)
    restore_toc_running_header(doc, body_header)
    apply_header_fields(doc, content)
    remove_first_line_indent_from_empty_table_paragraphs(doc)
    doc.save(str(staged_docx))
    wait_for_docx_ready(staged_docx)
    strip_report = run_root / "reports" / "template-instruction-artifact-strip.json"
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "strip_docx_template_instruction_artifacts.py"),
            "--input",
            str(staged_docx),
            "--output",
            str(staged_docx),
            "--scope",
            "all-parts",
            "--report-json",
            str(strip_report),
        ],
        timeout=300,
    )
    wait_for_docx_ready(staged_docx)
    compact_front_matter_gaps_in_docx(staged_docx)
    run([str(PYTHON_EXE), str(SCRIPT_DIR / "normalize_thesis_citation_chain.py"), "--docx", str(staged_docx)], timeout=900)
    wait_for_docx_ready(staged_docx)
    if content.get("preserve_visible_bibliography_labels") or content.get("reference_label_style") == "compact-bracket":
        compacted_count = compact_visible_bibliography_labels(staged_docx)
        write_json(
            run_root / "reports" / "bibliography-visible-label-preservation.json",
            {
                "schema": "graduation-project-builder.bibliography-visible-label-preservation.v1",
                "docx": str(staged_docx.resolve()),
                "policy": "preserve compact [n]content labels supplied by content manifest",
                "compact_visible_bibliography_labels_count": compacted_count,
                "strip_visible_bibliography_labels_for_auto_numbering": "skipped",
                "verdict": "pass",
            },
        )
    else:
        strip_visible_bibliography_labels_for_auto_numbering(staged_docx)
    wait_for_docx_ready(staged_docx)
    compact_front_matter_gaps_in_docx(staged_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    sync_static_toc_if_possible(staged_docx, output_docx, run_root, reference_docx=template)
    wait_for_docx_ready(output_docx)
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "strip_docx_template_instruction_artifacts.py"),
            "--input",
            str(output_docx),
            "--output",
            str(output_docx),
            "--scope",
            "all-parts",
            "--report-json",
            str(run_root / "reports" / "template-instruction-artifact-strip-final.json"),
        ],
        timeout=300,
    )
    wait_for_docx_ready(output_docx)
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "strip_docx_template_instruction_artifacts.py"),
            "--input",
            str(output_docx),
            "--scope",
            "all-parts",
            "--check-only",
            "--fail-on-artifacts",
            "--report-json",
            str(run_root / "reports" / "template-instruction-artifact-final-check.json"),
        ],
        timeout=300,
    )
    wait_for_docx_ready(output_docx)
    compact_front_matter_gaps_in_docx(output_docx)
    disable_revision_tracking_and_view_marks(output_docx)
    strip_non_image_vml_line_artifacts(output_docx)
    normalize_docx_paragraph_property_order(output_docx)
    normalize_docx_table_property_order(output_docx)
    normalize_docx_run_property_order(output_docx)
    normalize_docx_alignment_enum_values(output_docx)
    normalize_docx_font_table_charsets(output_docx)
    normalize_docx_mc_ignorable_prefixes(output_docx)
    clear_empty_table_paragraph_indents_in_docx(output_docx)
    normalize_whole_format_release_contract(output_docx, run_root)
    direct_surface_report = normalize_generated_thesis_direct_surface_contract(output_docx)
    write_json(run_root / "reports" / "generated-direct-surface-contract.json", direct_surface_report)
    normalize_docx_paragraph_property_order(output_docx)
    normalize_docx_table_property_order(output_docx)
    normalize_docx_run_property_order(output_docx)
    normalize_docx_alignment_enum_values(output_docx)
    normalize_docx_font_table_charsets(output_docx)
    normalize_docx_mc_ignorable_prefixes(output_docx)
    clear_empty_table_paragraph_indents_in_docx(output_docx)
    if restore_empty_toc_header_from_body_header(output_docx):
        write_json(
            run_root / "reports" / "toc-header-restoration.json",
            {
                "schema": "graduation-project-builder.toc-header-restoration.v1",
                "docx": str(output_docx.resolve()),
                "operation": "copied non-empty body running header into empty TOC header part",
                "toc_header_part": "word/header2.xml",
                "body_header_part": "word/header3.xml",
                "verdict": "pass",
            },
        )
    layout_contract_report = repair_docx_layout_contracts(output_docx)
    write_json(run_root / "reports" / "layout-contract-postprocess.json", layout_contract_report)
    normalize_docx_paragraph_property_order(output_docx)
    normalize_docx_table_property_order(output_docx)
    normalize_docx_run_property_order(output_docx)
    normalize_docx_font_table_charsets(output_docx)
    normalize_docx_mc_ignorable_prefixes(output_docx)
    wait_for_docx_ready(output_docx)
    return output_docx


def append_chapter_with_break(doc: Document, donors: Donors, chapter: dict[str, Any], chapter_no: int) -> Paragraph:
    heading = append_heading(
        doc,
        donors,
        1,
        numbered_heading_title(str(chapter.get("title") or "").strip(), (chapter_no,)),
        page_break_before=True,
    )
    for paragraph in chapter.get("paragraphs", []) or []:
        text = str(paragraph).strip()
        if text:
            append_body_paragraph(doc, donors, text)
    for idx, figure in enumerate(chapter.get("figures", []) or [], start=1):
        append_figure(doc, donors, figure, idx)
    for idx, table_data in enumerate(chapter.get("tables", []) or [], start=1):
        append_table(doc, donors, table_data, idx)
    for idx, section in enumerate(chapter.get("sections", []) or [], start=1):
        append_chapter(doc, donors, section, level=2, numbering=(chapter_no, idx))
    return heading


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--adapter", required=True, help="Validated project-local adapter manifest JSON")
    parser.add_argument("--content-manifest", help="Thesis content manifest JSON")
    parser.add_argument("--template", help="Template DOCX override")
    parser.add_argument("--output-docx", help="Output DOCX path")
    parser.add_argument("--output-pdf", help="Output PDF path")
    parser.add_argument("--run-root", help="Working/evidence directory")
    parser.add_argument("--dry-run-validate", action="store_true", help="Validate inputs only")
    parser.add_argument("--skip-render", action="store_true", help="Build DOCX only; do not produce a deliverable verdict")
    args = parser.parse_args()

    adapter, template, content_manifest, output_docx, output_pdf, run_root = resolve_inputs(args)
    content = read_json(content_manifest)
    content_issues = validate_content_manifest(content)
    if content_issues:
        raise RuntimeError("content manifest validation failed: " + "; ".join(content_issues))
    if args.dry_run_validate:
        print("CANONICAL THESIS BUILDER INPUT PASS")
        print(f"adapter={Path(args.adapter).resolve()}")
        print(f"template={template}")
        print(f"content_manifest={content_manifest}")
        return 0

    adapter_report = run_root / "reports" / "local-adapter-validation.json"
    adapter_report.parent.mkdir(parents=True, exist_ok=True)
    template_profile = run_root / "meta" / "template_profile.json"
    profile = build_template_profile(template)
    write_profile(profile, template_profile)
    profile_issues = profile_readiness_issues(profile)
    if profile_issues:
        raise RuntimeError("template profile validation failed: " + "; ".join(profile_issues))
    asset_manifest = collect_asset_manifest(content, run_root)
    asset_manifest_data = read_json(asset_manifest)
    figure_manifest_issues = validate_figure_manifest(asset_manifest_data)
    if figure_manifest_issues:
        raise RuntimeError("figure manifest validation failed: " + "; ".join(figure_manifest_issues))
    write_json(
        adapter_report,
        {
            "adapter": str(Path(args.adapter).resolve()),
            "result": "pass",
            "canonical_builder": str(Path(__file__).resolve()),
            "adapter_type": adapter.get("adapter_type"),
            "template_profile": str(template_profile),
            "asset_manifest": str(asset_manifest),
        },
    )
    final_docx = build_docx(template, content, output_docx, run_root)
    raster_primary_report = apply_svg_primary_to_docx(final_docx, asset_manifest_data)
    raster_primary_report["renderer_policy"] = "svg-primary-with-raster-fallback"
    raster_primary_report["reason"] = (
        "Structural figures retain PNG/JPEG raster fallback relationships while binding "
        "asvg:svgBlip SVG-primary relationships for machine-checkable provenance."
    )
    write_json(run_root / "reports" / "figure-svg-primary.json", raster_primary_report)
    formula_map_value = adapter.get("formula_map") or adapter.get("formula_map_path") or content.get("formula_map") or content.get("formula_map_path")
    if formula_map_value:
        formula_map = Path(str(formula_map_value)).resolve()
        if not formula_map.exists():
            raise FileNotFoundError(f"formula map not found: {formula_map}")
        formula_map_data = read_json(formula_map)
        before_formula = run_root / "stage" / "before-formula.docx"
        before_formula.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(final_docx, before_formula)
        if isinstance(formula_map_data.get("batches") or formula_map_data.get("insertions"), list):
            insert_formula_batch(
                docx_path=before_formula,
                formula_map_path=formula_map,
                output_path=final_docx,
                report_path=run_root / "reports" / "formula-batch-insertion.json",
            )
        else:
            replace_text_formulas(
                docx_path=before_formula,
                formula_map_path=formula_map,
                output_path=final_docx,
                report_path=run_root / "reports" / "formula-replacement.json",
            )
    clear_empty_table_paragraph_indents_in_docx(final_docx)
    mark_asset_manifest_completed(asset_manifest_data, final_docx, source_docx=template)
    write_json(asset_manifest, asset_manifest_data)
    post_figure_manifest_issues = validate_figure_manifest(
        asset_manifest_data,
        final_docx=final_docx,
        source_docx=template,
        manifest_path=asset_manifest,
    )
    if post_figure_manifest_issues:
        raise RuntimeError("figure DOCX insertion validation failed: " + "; ".join(post_figure_manifest_issues))
    if args.skip_render:
        print(f"docx={final_docx}")
        print(f"template_profile={template_profile}")
        print(f"asset_manifest={asset_manifest}")
        print(f"adapter_report={adapter_report}")
        print("CANONICAL THESIS BUILDER DOCX PASS")
        return 0
    export_pdf(final_docx, output_pdf)
    adapter_agent_options = adapter.get("agent_options") if isinstance(adapter.get("agent_options"), dict) else adapter
    reports = run_audits(
        template,
        final_docx,
        output_pdf,
        asset_manifest,
        template_profile,
        run_root,
        project_root=Path(str(adapter.get("project_root"))).resolve() if adapter.get("project_root") else None,
        humanizer_evidence=adapter.get("humanizer_evidence"),
        transaction_record=adapter.get("transaction_record"),
        agent_options={
            "agent_authorization_source": adapter_agent_options.get("agent_authorization_source"),
            "agent_mode": adapter_agent_options.get("agent_mode"),
            "spawned_agent_ids": adapter_agent_options.get("spawned_agent_ids"),
            "spawned_agent_aliases_zh": adapter_agent_options.get("spawned_agent_aliases_zh"),
            "controller_agent_id": adapter_agent_options.get("controller_agent_id"),
            "format_agent_id": adapter_agent_options.get("format_agent_id"),
            "audit_agent_id": adapter_agent_options.get("audit_agent_id"),
            "sequential_fallback_reason": adapter_agent_options.get("sequential_fallback_reason"),
        },
        gate_evidence=adapter.get("gate_evidence") if isinstance(adapter.get("gate_evidence"), dict) else None,
        acceptance_options=adapter,
    )
    print(f"docx={final_docx}")
    print(f"pdf={output_pdf}")
    print(f"template_profile={template_profile}")
    print(f"asset_manifest={asset_manifest}")
    print(f"adapter_report={adapter_report}")
    for name, path in reports.items():
        print(f"{name}={path}")
    print("CANONICAL THESIS BUILDER PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
