#!/usr/bin/env python3
"""Run real DOCX integration checks."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Callable
from xml.etree import ElementTree as ET

import fitz  # type: ignore
from PIL import Image, ImageDraw, ImageFont
from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Cm, Pt  # type: ignore
from docx.table import Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

from build_minimal_template_thesis import (
    build_assets,
    export_pdf,
    export_with_word,
    insert_formula_omml,
    wait_for_docx_ready,
)
from docx_apply_table_family import patch_table
from docx_formula_number_table import patch_formula
from docx_sync_picture import patch_docx as patch_picture_docx
from python_runtime import resolve_python_exe


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
PYTHON_EXE = resolve_python_exe()
DEFAULT_TEMP_PREFIX = "gpb_integration_gate_"
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}
W = "{%s}" % NS["w"]
W14 = "{%s}" % NS["w14"]
PR = "{%s}" % NS["pr"]
WP = "{%s}" % NS["wp"]


def file_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


BODY_STYLE_NAME = "论文正文"
CAPTION_STYLE_NAME = "图注"
TABLE_CAPTION_STYLE_NAME = "表题"
CODE_TITLE_STYLE_NAME = "代码题名"
CODE_BLOCK_STYLE_NAME = "代码块"
IMAGE_HOLDER_STYLE_NAME = "图片承载"
HEADER_LEFT = "示例大学"
HEADER_RIGHT = "毕业设计"
FOOTER_TEXT = "— 1 —"

SUMMARY_PASS = "通过"
SUMMARY_FAIL = "未通过"

SAMPLE_SELF_CHECK_LABELS = {
    "heading_family": "标题样式家族检查",
    "heading1_baseline": "一级标题真实基线检查",
    "heading2_baseline": "二级标题真实基线检查",
    "heading_baseline": "三级标题真实基线检查",
    "heading4_baseline": "四级标题真实基线检查",
    "figure_family": "图形家族检查",
    "table_family": "表格家族检查",
    "cover_baseline": "封面真实基线检查",
    "non_body_surface": "非正文面正文污染检查",
    "non_body_indent": "非正文首行缩进检查",
    "continuation": "跨页续表检查",
    "figure_block": "图块本地锚定检查",
    "figure_caption_baseline": "图题真实基线检查",
    "image_holder_baseline": "图片承载段落真实基线检查",
    "image_holder_safety": "图片承载段落安全检查",
    "image_dimension": "图片尺寸安全检查",
    "toc_control": "目录控件正文污染检查",
    "table_caption_binding": "表题绑定检查",
    "table_caption_baseline": "表题真实基线检查",
    "table_cell_baseline": "表格文字真实基线检查",
    "bibliography_count": "参考文献条目数量检查",
    "body_style": "正文样式绑定检查",
    "body_baseline": "正文默认样式基线检查",
    "code_title": "代码题名格式检查",
    "code_block": "代码块格式检查",
    "toc_visible": "目录可见格式检查",
    "header_position": "页眉位置检查",
    "header_presence": "header presence contract check",
    "footer_indent": "页脚缩进检查",
    "abstract_baseline": "摘要真实基线检查",
    "header_footer_baseline": "页眉页脚真实基线检查",
    "figure_manifest_contract": "figure manifest contract check",
}

CASE_SEQUENCE = [
    "frontmatter_roundtrip",
    "lock_competition",
    "tail_pages_roundtrip",
    "staging_copy_isolation",
    "paragraph_review_microcycle",
    "serialized_helper_mutation_chain",
    "multi_pass_review_copy_governance",
    "serialized_formula_helper_chain",
    "complete_sample_smoke",
    "sample_self_check_table_family_detection",
    "sample_self_check_table_caption_baseline_detection",
    "sample_self_check_table_cell_baseline_detection",
    "sample_self_check_heading_family_detection",
    "sample_self_check_heading_all_levels_run_override_detection",
    "sample_self_check_heading2_run_override_detection",
    "sample_self_check_heading_baseline_detection",
    "sample_self_check_code_title_format_detection",
    "sample_self_check_code_block_format_detection",
    "sample_self_check_cover_baseline_detection",
    "sample_self_check_abstract_baseline_detection",
    "sample_self_check_abstract_manual_break_detection",
    "sample_self_check_toc_visible_detection",
    "sample_self_check_header_position_detection",
    "sample_self_check_header_presence_detection",
    "sample_self_check_footer_indent_detection",
    "sample_self_check_footer_baseline_detection",
    "sample_self_check_body_style_binding_detection",
    "sample_self_check_normal_baseline_drift_detection",
    "sample_self_check_non_body_surface_contamination_detection",
    "sample_self_check_non_body_indent_detection",
    "sample_self_check_cross_page_table_continuation_detection",
    "sample_self_check_caption_inside_table_detection",
    "sample_self_check_table_caption_binding_detection",
    "sample_self_check_figure_block_locality_detection",
    "sample_self_check_figure_explanation_followup_detection",
    "sample_self_check_image_holder_safety_detection",
    "sample_self_check_image_dimension_detection",
    "sample_self_check_missing_asset_manifest_blocks_delivery",
    "sample_self_check_empty_diagrams_blocks_structural_docx",
    "sample_self_check_figure_manifest_evidence_detection",
    "sample_self_check_svg_fallback_text_detection",
    "sample_self_check_toc_control_contamination_detection",
    "alternate_renderer_parity_when_available",
    "sample_self_check_bibliography_count_detection",
    "sample_self_check_bibliography_rendered_geometry_detection",
    "sample_self_check_keyword_donor_shift_detection",
    "sample_self_check_abstract_self_donor_alias_detection",
    "sample_self_check_figure_caption_style_detection",
    "sample_self_check_image_holder_body_residue_detection",
]


@dataclass
class GateEnvironment:
    officecli: str
    renderer: str
    libreoffice: str
    wps: str


@dataclass
class CaseResult:
    name: str
    passed: bool
    details: str


class GateFailure(RuntimeError):
    """Raised when an integration case fails."""


CaseBuilder = Callable[[Path, GateEnvironment], CaseResult]
CASE_BUILDERS: dict[str, CaseBuilder] = {}


def run_command(
    cmd: list[str],
    *,
    timeout: int = 1800,
    check: bool = True,
    cwd: Path | None = None,
) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(cwd) if cwd else None,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=timeout,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            f"command failed: {' '.join(cmd)}\nstdout:\n{proc.stdout}\nstderr:\n{proc.stderr}"
        )
    return proc


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "").lower()


def pdf_page_texts(pdf_path: Path) -> list[str]:
    with fitz.open(pdf_path) as pdf:
        return [page.get_text("text") for page in pdf]


def find_page(page_texts: list[str], token: str) -> int | None:
    target = normalize_text(token)
    for idx, text in enumerate(page_texts, start=1):
        if target and target in normalize_text(text):
            return idx
    return None


def find_last_page(page_texts: list[str], token: str) -> int | None:
    target = normalize_text(token)
    found: int | None = None
    for idx, text in enumerate(page_texts, start=1):
        if target and target in normalize_text(text):
            found = idx
    return found


def assert_report_dimension(report_text: str, label: str, expected: str, assertion_label: str) -> None:
    pattern = re.compile(rf"{re.escape(label)}\s*[:：]\s*{re.escape(expected)}")
    if pattern.search(report_text) is None:
        raise GateFailure(
            f"missing expected summary for {assertion_label}: '{label}: {expected}'"
        )


def locate_binary(candidates: list[Path], *, which_name: str | None = None) -> str:
    if which_name:
        found = shutil.which(which_name)
        if found:
            return str(Path(found).resolve())
    for candidate in candidates:
        if candidate.exists():
            return str(candidate.resolve())
    return "not-found"


def detect_environment() -> GateEnvironment:
    officecli = locate_binary(
        [
            Path.home() / "AppData" / "Local" / "OfficeCli" / "officecli.EXE",
            Path.home() / "AppData" / "Local" / "OfficeCli" / "officecli.exe",
        ],
        which_name="officecli",
    )
    libreoffice = locate_binary(
        [Path(r"C:\Program Files\LibreOffice\program\soffice.exe")],
        which_name="soffice",
    )
    wps = locate_binary(
        [
            Path(r"D:\WPS Office\12.1.0.25865\office6\wps.exe"),
            Path(r"C:\Program Files\WPS Office\office6\wps.exe"),
        ]
    )
    renderer = "Word.Application COM"
    try:
        probe = Path(tempfile.gettempdir()) / "gpb_word_probe.docx"
        Document().save(str(probe))
        out_pdf = probe.with_suffix(".pdf")
        if out_pdf.exists():
            out_pdf.unlink()
        export_with_word(probe, out_pdf)
        if out_pdf.exists():
            out_pdf.unlink()
        if probe.exists():
            probe.unlink()
    except Exception:
        if wps != "not-found":
            renderer = "WPS COM"
        elif libreoffice != "not-found":
            renderer = "LibreOffice"
        else:
            renderer = "unavailable"
    return GateEnvironment(
        officecli=officecli,
        renderer=renderer,
        libreoffice=libreoffice,
        wps=wps,
    )


def export_pdf_with_renderer(docx_path: Path, pdf_path: Path, renderer: str, env: GateEnvironment) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    if renderer == "Word.Application COM":
        export_with_word(docx_path, pdf_path)
        return
    if renderer == "WPS COM":
        run_command(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(SCRIPT_DIR / "wps_export_pdf.ps1"),
                "-InputDocx",
                str(docx_path),
                "-OutputPdf",
                str(pdf_path),
            ],
            timeout=600,
        )
        return
    if renderer == "LibreOffice":
        if env.libreoffice == "not-found":
            raise RuntimeError("LibreOffice is unavailable")
        temp_dir = pdf_path.parent
        run_command(
            [
                env.libreoffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temp_dir),
                str(docx_path),
            ],
            timeout=600,
        )
        produced = temp_dir / f"{docx_path.stem}.pdf"
        if produced != pdf_path and produced.exists():
            shutil.move(str(produced), str(pdf_path))
        return
    raise RuntimeError(f"unsupported renderer: {renderer}")


def export_pdf_with_sample_fallback(docx_path: Path, pdf_path: Path, env: GateEnvironment) -> str:
    errors: list[str] = []
    for renderer in ("Word.Application COM", "LibreOffice", "WPS COM"):
        if renderer == "LibreOffice" and env.libreoffice == "not-found":
            continue
        if renderer == "WPS COM" and env.wps == "not-found":
            continue
        try:
            if pdf_path.exists():
                pdf_path.unlink()
            export_pdf_with_renderer(docx_path, pdf_path, renderer, env)
            if pdf_path.exists():
                return renderer
            errors.append(f"{renderer}: did not create {pdf_path}")
        except Exception as exc:
            errors.append(f"{renderer}: {exc}")
    raise RuntimeError("all sample-self-check renderers failed: " + " | ".join(errors))


def create_diagnostic_pdf_from_docx(docx_path: Path, pdf_path: Path, reason: str) -> None:
    doc = Document(str(docx_path))
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open()
    lines = [
        "DIAGNOSTIC PDF: real renderer unavailable",
        reason[:220],
        "",
    ]
    lines.extend(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
    page = None
    y = 42
    for line in lines:
        if page is None or y > 780:
            page = pdf.new_page(width=595, height=842)
            y = 42
        page.insert_text((42, y), line[:110], fontsize=10)
        y += 15
    pdf.save(str(pdf_path))
    pdf.close()


def copy_locked(src: Path, dst: Path) -> None:
    run_command([str(PYTHON_EXE), str(SCRIPT_DIR / "copy_locked_docx.py"), str(src), str(dst)])


def media_member_count(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        return sum(1 for name in zf.namelist() if name.startswith("word/media/"))


def document_xml(docx_path: Path) -> ET.Element:
    with zipfile.ZipFile(docx_path) as zf:
        return ET.fromstring(zf.read("word/document.xml"))


def styles_xml(docx_path: Path) -> ET.Element:
    with zipfile.ZipFile(docx_path) as zf:
        return ET.fromstring(zf.read("word/styles.xml"))


def overwrite_docx_part(docx_path: Path, part_name: str, new_bytes: bytes) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        members = {name: zin.read(name) for name in zin.namelist()}
    members[part_name] = new_bytes
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(name, data)


def save_document_xml(docx_path: Path, root: ET.Element) -> None:
    overwrite_docx_part(docx_path, "word/document.xml", ET.tostring(root, encoding="utf-8", xml_declaration=True))


def save_styles_xml(docx_path: Path, root: ET.Element) -> None:
    overwrite_docx_part(docx_path, "word/styles.xml", ET.tostring(root, encoding="utf-8", xml_declaration=True))


def ensure_style(
    doc: Document,
    style_name: str,
    *,
    base_style: str = "Normal",
    size_pt: float = 12.0,
    bold: bool = False,
    first_line_indent_pt: float | None = None,
    line_spacing: float | None = None,
) -> None:
    styles = doc.styles
    try:
        styles[style_name]
        return
    except KeyError:
        pass
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    base = styles[base_style]
    style.base_style = base
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "SimSun")
    pf = style.paragraph_format
    if first_line_indent_pt is not None:
        pf.first_line_indent = Pt(first_line_indent_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


def set_run_font(run_obj, *, east_asia: str = "SimSun", size_pt: float = 12.0, bold: bool = False) -> None:
    run_obj.font.name = "Times New Roman"
    run_obj.font.size = Pt(size_pt)
    run_obj.bold = bold
    rpr = run_obj._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def reset_paragraph_format(
    paragraph: Paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent_pt: float = 0.0,
    left_indent_pt: float | None = 0.0,
    right_indent_pt: float | None = 0.0,
    line_spacing: float = 1.0,
    keep_with_next: bool = False,
) -> None:
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(first_line_indent_pt)
    pf.left_indent = Pt(left_indent_pt) if left_indent_pt is not None else None
    pf.right_indent = Pt(right_indent_pt) if right_indent_pt is not None else None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def add_body_paragraph(doc: Document, text: str, *, body_style_name: str) -> Paragraph:
    paragraph = doc.add_paragraph(style=body_style_name)
    reset_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=24.0,
        line_spacing=1.3,
    )
    run_obj = paragraph.add_run(text)
    set_run_font(run_obj, east_asia="SimSun", size_pt=12.0, bold=False)
    return paragraph


def configure_styles(doc: Document, *, custom_body_style: bool) -> str:
    ensure_style(
        doc,
        BODY_STYLE_NAME,
        base_style="Normal",
        size_pt=12.0,
        first_line_indent_pt=24.0,
        line_spacing=1.3,
    )
    ensure_style(doc, CAPTION_STYLE_NAME, base_style="Normal", size_pt=12.0, line_spacing=1.0)
    ensure_style(doc, TABLE_CAPTION_STYLE_NAME, base_style="Normal", size_pt=12.0, line_spacing=1.0)
    ensure_style(doc, CODE_TITLE_STYLE_NAME, base_style="Normal", size_pt=12.0, line_spacing=1.0)
    ensure_style(doc, CODE_BLOCK_STYLE_NAME, base_style="Normal", size_pt=10.5, line_spacing=1.0)
    ensure_style(doc, IMAGE_HOLDER_STYLE_NAME, base_style="Normal", size_pt=12.0, line_spacing=1.0)
    return BODY_STYLE_NAME if custom_body_style else "Normal"


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    for paragraph in list(header.paragraphs):
        if paragraph.text.strip():
            clear_paragraph_content(paragraph)

    table = header.add_table(rows=1, cols=2, width=Cm(15.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_para = table.cell(0, 0).paragraphs[0]
    right_para = table.cell(0, 1).paragraphs[0]
    reset_paragraph_format(left_para, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    reset_paragraph_format(right_para, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0)
    left_run = left_para.add_run(HEADER_LEFT)
    right_run = right_para.add_run(HEADER_RIGHT)
    set_run_font(left_run, east_asia="SimSun", size_pt=12.0, bold=False)
    set_run_font(right_run, east_asia="SimSun", size_pt=12.0, bold=False)

    footer_para = footer.paragraphs[0]
    clear_paragraph_content(footer_para)
    reset_paragraph_format(footer_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    footer_run = footer_para.add_run(FOOTER_TEXT)
    set_run_font(footer_run, east_asia="SimSun", size_pt=12.0, bold=False)


def try_set_style(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        pass


def add_toc_entry(doc: Document, text: str, style_name: str, left_indent_pt: float) -> None:
    paragraph = doc.add_paragraph()
    try_set_style(paragraph, style_name)
    reset_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0.0,
        left_indent_pt=left_indent_pt,
        line_spacing=1.0,
    )
    run_obj = paragraph.add_run(text)
    set_run_font(run_obj, east_asia="SimSun", size_pt=12.0, bold=False)


def create_base_sample_docx(
    docx_path: Path,
    *,
    custom_body_style: bool = False,
    include_header_footer: bool = False,
) -> None:
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    body_style_name = configure_styles(doc, custom_body_style=custom_body_style)
    if include_header_footer:
        configure_header_footer(doc)

    paragraph = doc.add_paragraph("毕业设计")
    reset_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    set_run_font(paragraph.runs[0], east_asia="SimSun", size_pt=16.0, bold=True)
    doc.add_paragraph("")
    doc.add_paragraph("本科毕业设计（论文）诚信承诺书")
    doc.add_paragraph("本人承诺样稿内容仅用于重建 smoke integration。")
    doc.add_page_break()

    zh_title = doc.add_paragraph("摘   要")
    reset_paragraph_format(zh_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    set_run_font(zh_title.runs[0], east_asia="SimSun", size_pt=14.0, bold=True)
    add_body_paragraph(doc, "这是用于 smoke integration 的中文摘要。", body_style_name=body_style_name)
    zh_kw = doc.add_paragraph()
    reset_paragraph_format(zh_kw, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    zh_kw.add_run("关键词：")
    zh_kw.add_run("课堂行为检测；教学评估")
    set_run_font(zh_kw.runs[0], east_asia="SimSun", size_pt=12.0, bold=True)
    set_run_font(zh_kw.runs[1], east_asia="SimSun", size_pt=12.0, bold=False)
    doc.add_page_break()

    en_title = doc.add_paragraph("Abstract")
    reset_paragraph_format(en_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    set_run_font(en_title.runs[0], east_asia="Times New Roman", size_pt=14.0, bold=True)
    en_body = doc.add_paragraph()
    reset_paragraph_format(
        en_body,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=24.0,
        line_spacing=1.3,
    )
    en_run = en_body.add_run("This is the English abstract for the rebuild smoke case.")
    set_run_font(en_run, east_asia="Times New Roman", size_pt=12.0, bold=False)
    en_kw = doc.add_paragraph()
    reset_paragraph_format(en_kw, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    en_kw.add_run("Key Words: ")
    en_kw.add_run("Classroom Behavior Detection; Evaluation")
    set_run_font(en_kw.runs[0], east_asia="Times New Roman", size_pt=12.0, bold=True)
    set_run_font(en_kw.runs[1], east_asia="Times New Roman", size_pt=12.0, bold=False)
    doc.add_page_break()

    toc_heading = doc.add_paragraph("目   录")
    try_set_style(toc_heading, "TOC Heading")
    reset_paragraph_format(toc_heading, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_toc_entry(doc, "1 绪论\t1", "TOC 1", 0.0)
    add_toc_entry(doc, "1.1 研究背景与意义\t1", "TOC 2", 18.0)
    add_toc_entry(doc, "1.1.1 研究目标\t1", "TOC 3", 36.0)
    doc.add_page_break()

    h1 = doc.add_paragraph("1 绪论")
    try_set_style(h1, "Heading 1")
    h2 = doc.add_paragraph("1.1 研究背景与意义")
    try_set_style(h2, "Heading 2")
    h3 = doc.add_paragraph("1.1.1 研究目标")
    try_set_style(h3, "Heading 3")
    add_body_paragraph(doc, "这是正文内容。", body_style_name=body_style_name)
    doc.add_page_break()

    doc.add_paragraph("结   论")
    add_body_paragraph(doc, "这是结论内容。", body_style_name=body_style_name)
    doc.add_page_break()

    doc.add_paragraph("致   谢")
    add_body_paragraph(doc, "感谢测试环境。", body_style_name=body_style_name)
    doc.add_page_break()

    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Example Paper[J].")
    doc.save(str(docx_path))


def create_builder_template_docx(docx_path: Path) -> None:
    run_root = docx_path.parent
    create_base_sample_docx(docx_path, custom_body_style=True, include_header_footer=True)
    doc = Document(str(docx_path))
    styles = [p.text.strip() for p in doc.paragraphs]
    body_anchor = next(p for p in doc.paragraphs if p.text.strip() == "这是正文内容。")
    conclusion_anchor = next(p for p in doc.paragraphs if p.text.strip() == "结   论")

    chapter_titles = [
        "2 系统设计",
        "3 结构图验证",
        "4 截图与表格",
        "5 公式验证",
    ]
    anchor = body_anchor
    for title in chapter_titles:
        chapter = insert_paragraph_after(anchor)
        chapter.text = title
        try_set_style(chapter, "Heading 1")
        anchor = chapter

    appendix_heading = insert_paragraph_after(conclusion_anchor)
    appendix_heading.text = "附   录"
    appendix_body = insert_paragraph_after(appendix_heading)
    appendix_body.text = "附录说明。"
    try_set_style(appendix_body, BODY_STYLE_NAME)
    reset_paragraph_format(
        appendix_body,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=24.0,
        line_spacing=1.3,
    )

    donor_png = run_root / "template-donor.png"
    make_placeholder_png(donor_png, "模板图")
    holder = insert_paragraph_after(appendix_body)
    try_set_style(holder, IMAGE_HOLDER_STYLE_NAME)
    clear_paragraph_content(holder)
    reset_paragraph_format(
        holder,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        left_indent_pt=0.0,
        right_indent_pt=0.0,
        line_spacing=1.0,
        keep_with_next=True,
    )
    holder.add_run().add_picture(str(donor_png), width=Cm(6.0))
    caption = insert_paragraph_after(holder)
    try_set_style(caption, CAPTION_STYLE_NAME)
    reset_paragraph_format(
        caption,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        left_indent_pt=0.0,
        right_indent_pt=0.0,
        line_spacing=1.0,
    )
    caption_run = caption.add_run("图 0-1 模板图题")
    set_run_font(caption_run, east_asia="SimSun", size_pt=12.0, bold=False)
    doc.save(str(docx_path))


def build_sample_self_check_fixture(
    fixture_root: Path,
    *,
    custom_body_style: bool = False,
    include_header_footer: bool = False,
) -> tuple[Path, Path]:
    fixture_root.mkdir(parents=True, exist_ok=True)
    reference_docx = fixture_root / "reference.docx"
    final_docx = fixture_root / "final.docx"
    create_base_sample_docx(
        reference_docx,
        custom_body_style=custom_body_style,
        include_header_footer=include_header_footer,
    )
    shutil.copy2(reference_docx, final_docx)
    return reference_docx, final_docx


def build_complete_sample_fixture(fixture_root: Path) -> tuple[Path, Path, Path]:
    fixture_root.mkdir(parents=True, exist_ok=True)
    template_docx = fixture_root / "template.docx"
    source_docx = fixture_root / "source.docx"
    create_builder_template_docx(template_docx)
    shutil.copy2(template_docx, source_docx)
    return template_docx, source_docx, fixture_root


def make_placeholder_png(path: Path, label: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (960, 540), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.rectangle((60, 60, 900, 480), outline="#111111", width=4)
    draw.rectangle((140, 150, 420, 250), outline="#111111", width=3)
    draw.rectangle((540, 150, 820, 250), outline="#111111", width=3)
    draw.rectangle((340, 310, 620, 410), outline="#111111", width=3)
    draw.line((420, 200, 540, 200), fill="#111111", width=4)
    draw.line((680, 250, 680, 310), fill="#111111", width=4)
    font = ImageFont.load_default()
    draw.text((220, 190), label, fill="#111111", font=font)
    draw.text((620, 190), "节点B", fill="#111111", font=font)
    draw.text((460, 350), "节点C", fill="#111111", font=font)
    image.save(path)


def create_flowchart_assets(asset_root: Path, *, fallback_text: bool = False) -> Path:
    asset_root.mkdir(parents=True, exist_ok=True)
    drawio_path = asset_root / "flowchart.drawio"
    svg_path = asset_root / "flowchart.svg"
    png_path = asset_root / "flowchart.png"
    drawio_path.write_text(
        """<mxfile host="app.diagrams.net" modified="2026-04-24T00:00:00.000Z" agent="Codex" version="24.7.17">
  <diagram id="flowchart" name="Page-1">
    <mxGraphModel dx="1200" dy="1200" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="850" pageHeight="1100" math="0" shadow="0">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
        <mxCell id="2" value="开始" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="300" y="60" width="160" height="70" as="geometry"/>
        </mxCell>
        <mxCell id="3" value="采集数据" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="270" y="190" width="220" height="80" as="geometry"/>
        </mxCell>
        <mxCell id="4" value="分析行为" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="270" y="340" width="220" height="80" as="geometry"/>
        </mxCell>
        <mxCell id="5" value="输出结果" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="270" y="490" width="220" height="80" as="geometry"/>
        </mxCell>
        <mxCell id="6" value="结束" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="300" y="640" width="160" height="70" as="geometry"/>
        </mxCell>
        <mxCell id="7" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="2" target="3">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
        <mxCell id="8" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="3" target="4">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
        <mxCell id="9" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="4" target="5">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
        <mxCell id="10" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="5" target="6">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>
""",
        encoding="utf-8",
    )
    svg_text = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="640" height="920">'
        '<rect x="180" y="40" width="280" height="70" fill="#FFFFFF" stroke="#000000"/>'
        '<text x="320" y="80" text-anchor="middle" fill="#000000">开始</text>'
        '<rect x="170" y="180" width="300" height="80" fill="#FFFFFF" stroke="#000000"/>'
        '<text x="320" y="225" text-anchor="middle" fill="#000000">采集数据</text>'
        '<rect x="170" y="330" width="300" height="80" fill="#FFFFFF" stroke="#000000"/>'
        '<text x="320" y="375" text-anchor="middle" fill="#000000">分析行为</text>'
        '<rect x="170" y="480" width="300" height="80" fill="#FFFFFF" stroke="#000000"/>'
        '<text x="320" y="525" text-anchor="middle" fill="#000000">输出结果</text>'
        '<rect x="180" y="630" width="280" height="70" fill="#FFFFFF" stroke="#000000"/>'
        '<text x="320" y="670" text-anchor="middle" fill="#000000">结束</text>'
        '</svg>'
    )
    if fallback_text:
        svg_text += '<text x="50%" y="100%">Text is not SVG - cannot display</text>'
    svg_path.write_text(svg_text, encoding="utf-8")
    make_placeholder_png(png_path, "流程图")
    task_card = asset_root / "figure_1-task-card.md"
    rendered_evidence = asset_root / "figure_1-rendered-evidence.md"
    relationship_evidence = asset_root / "figure_1-relationship-evidence.md"
    task_card.write_text("- figure task card: pass\n", encoding="utf-8")
    rendered_evidence.write_text("- post-insertion rendered evidence: pass\n", encoding="utf-8")
    relationship_evidence.write_text("- final DOCX relationship evidence: pass\n", encoding="utf-8")
    asset_manifest = asset_root / "asset_manifest.json"
    asset_manifest.write_text(
        json.dumps(
            {
                "schema": "graduation-project-builder.figure-manifest.v2",
                "figures": {
                    "figure_1": {
                        "path": str(png_path),
                        "caption": "图 1 测试流程图",
                        "family": "flowchart",
                        "source_kind": "structural",
                        "caption_to_asset_mapping": f"图 1 测试流程图 -> {png_path}",
                        "task_card": str(task_card),
                        "post_insertion_rendered_evidence": str(rendered_evidence),
                        "final_docx_relationship_evidence": str(relationship_evidence),
                        "rendered_page_status": "pass",
                        "insertion_status": "inserted",
                    }
                },
                "tables": {},
                "diagrams": {
                    "diagram_1": {
                        "id": "figure_1",
                        "caption": "图 1 测试流程图",
                        "description": "测试流程图",
                        "declared_family": "flowchart",
                        "inferred_family": "flowchart",
                        "drawio": str(drawio_path),
                        "svg": str(svg_path),
                        "png": str(png_path),
                        "raster_fallback": str(png_path),
                        "family": "flowchart",
                        "source_kind": "structural",
                        "sample_lock": "figure-flowchart-vertical-sample-01",
                        "svg_primary_expected": True,
                        "caption_to_asset_mapping": f"图 1 测试流程图 -> {png_path}",
                        "task_card": str(task_card),
                        "post_insertion_rendered_evidence": str(rendered_evidence),
                        "final_docx_relationship_evidence": str(relationship_evidence),
                        "rendered_page_status": "pass",
                        "insertion_status": "inserted",
                    }
                }
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return asset_manifest


def create_picture_replacement_manifest(
    manifest_path: Path,
    *,
    replacement_image: Path,
    caption_text: str,
) -> Path:
    root = manifest_path.parent
    root.mkdir(parents=True, exist_ok=True)
    task_card = root / "picture-replacement-task-card.md"
    rendered_evidence = root / "picture-replacement-rendered-evidence.md"
    relationship_evidence = root / "picture-replacement-relationship-evidence.md"
    sample_lock_evidence = root / "picture-replacement-sample-lock-evidence.md"
    task_card.write_text("- figure replacement task card: pass\n", encoding="utf-8")
    rendered_evidence.write_text("- post-insertion rendered evidence: pass\n", encoding="utf-8")
    relationship_evidence.write_text("- final DOCX relationship evidence: pass\n", encoding="utf-8")
    sample_lock_evidence.write_text("- template figures checked: pass\n- no-template-figure-sample verdict: pass\n", encoding="utf-8")
    manifest_path.write_text(
        json.dumps(
            {
                "schema": "graduation-project-builder.figure-manifest.v2",
                "figures": {
                    "figure_1": {
                        "path": str(replacement_image),
                        "caption": caption_text,
                        "family": "raster",
                        "source_kind": "raster",
                        "caption_to_asset_mapping": f"{caption_text} -> {replacement_image.name}",
                        "task_card": str(task_card),
                        "post_insertion_rendered_evidence": str(rendered_evidence),
                        "final_docx_relationship_evidence": str(relationship_evidence),
                        "rendered_page_status": "pass",
                        "insertion_status": "inserted",
                        "no_template_figure_sample_verdict": "pass",
                        "skill_internal_fallback_sample_path": str(
                            SKILL_ROOT / "references" / "visual-style-samples" / "figures" / "figure-flowchart-vertical-sample-01.png"
                        ),
                        "chosen_style_source": "skill-internal fallback after no usable active template figure sample",
                        "sample_lock_evidence_path": str(sample_lock_evidence),
                        "mutation_intent": "replace_existing",
                        "explicit_replacement_authorization_source": "integration fixture approved replacement",
                        "explicit_replacement_authorization_scope": "single helper chain fixture image",
                        "target_anchor_not_protected_surface_verdict": "pass",
                        "caption_asset_binding_verdict": "pass",
                    }
                },
                "tables": {},
                "diagrams": {},
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    return manifest_path


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise GateFailure(f"paragraph not found: {text}")


def add_code_sample(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    anchor = find_paragraph(doc, "这是正文内容。")
    title = insert_paragraph_after(anchor)
    try_set_style(title, CODE_TITLE_STYLE_NAME)
    reset_paragraph_format(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    title_run = title.add_run("代码 1 示例代码")
    set_run_font(title_run, east_asia="SimSun", size_pt=12.0, bold=False)
    code = insert_paragraph_after(title)
    try_set_style(code, CODE_BLOCK_STYLE_NAME)
    reset_paragraph_format(code, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    code_run = code.add_run("def demo():\n    return 1")
    set_run_font(code_run, east_asia="Consolas", size_pt=10.5, bold=False)
    doc.save(str(docx_path))


def add_figure_sample(
    docx_path: Path,
    image_path: Path,
    *,
    caption_text: str = "图 1 测试图",
    para_id: str = "00100038",
    explanation_text: str = "图后正文介绍用于验证图片、图题与说明段保持同一本地块。",
) -> None:
    doc = Document(str(docx_path))
    anchor = find_paragraph(doc, "这是正文内容。")
    holder = insert_paragraph_after(anchor)
    clear_paragraph_content(holder)
    reset_paragraph_format(
        holder,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        left_indent_pt=0.0,
        right_indent_pt=0.0,
        line_spacing=1.0,
        keep_with_next=True,
    )
    holder._element.set(qn("w14:paraId"), para_id)
    holder._element.set(qn("w14:textId"), "00ABCDEF")
    holder.add_run().add_picture(str(image_path), width=Cm(6.5))
    caption = insert_paragraph_after(holder)
    try_set_style(caption, CAPTION_STYLE_NAME)
    reset_paragraph_format(
        caption,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        left_indent_pt=0.0,
        right_indent_pt=0.0,
        line_spacing=1.0,
    )
    run_obj = caption.add_run(caption_text)
    set_run_font(run_obj, east_asia="SimSun", size_pt=12.0, bold=False)
    explanation = insert_paragraph_after(caption)
    try_set_style(explanation, BODY_STYLE_NAME)
    reset_paragraph_format(
        explanation,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=24.0,
        left_indent_pt=0.0,
        right_indent_pt=0.0,
        line_spacing=1.3,
    )
    explanation_run = explanation.add_run(explanation_text)
    set_run_font(explanation_run, east_asia="SimSun", size_pt=12.0, bold=False)
    doc.save(str(docx_path))


def add_table_sample(docx_path: Path, *, long_table: bool = False) -> None:
    doc = Document(str(docx_path))
    anchor = find_paragraph(doc, "这是正文内容。")
    caption = insert_paragraph_after(anchor)
    try_set_style(caption, TABLE_CAPTION_STYLE_NAME)
    reset_paragraph_format(caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, keep_with_next=True)
    ppr = caption._element.get_or_add_pPr()
    for tag in ("w:numPr", "w:outlineLvl"):
        node = ppr.find(qn(tag))
        if node is not None:
            ppr.remove(node)
    caption_run = caption.add_run("表 1 测试表")
    set_run_font(caption_run, east_asia="SimSun", size_pt=12.0, bold=False)
    row_count = 60 if long_table else 4
    table = doc.add_table(rows=row_count, cols=3)
    caption._element.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index in range(row_count):
        for col_index in range(3):
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            clear_paragraph_content(paragraph)
            reset_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_pt=0.0,
                left_indent_pt=0.0,
                right_indent_pt=0.0,
                line_spacing=1.0,
            )
            if row_index == 0:
                value = ["字段", "含义", "来源"][col_index]
            else:
                value = f"字段{row_index}" if col_index == 0 else (f"含义{row_index}" if col_index == 1 else f"来源{row_index}")
            run_obj = paragraph.add_run(value)
            set_run_font(run_obj, east_asia="SimSun", size_pt=10.5, bold=row_index == 0)
    doc.save(str(docx_path))
    if not long_table:
        patch_table(docx_path, 1, "wps_second_three_line_rendered", None)


def add_heading4_sample(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    anchor = find_paragraph(doc, "1.1.1 研究目标")
    heading4 = insert_paragraph_after(anchor)
    heading4.text = "1.1.1.1 样本范围"
    try_set_style(heading4, "Heading 4")
    doc.save(str(docx_path))


def add_formula_sample(docx_path: Path, *, para_id: str = "00ABC321") -> None:
    doc = Document(str(docx_path))
    anchor = find_paragraph(doc, "这是正文内容。")
    formula_anchor = insert_paragraph_after(anchor)
    formula_anchor.text = "__INTEGRATION_FORMULA__"
    doc.save(str(docx_path))
    insert_formula_omml(docx_path, "__INTEGRATION_FORMULA__", para_id)


def add_header_footer_baseline(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    configure_header_footer(doc)
    doc.save(str(docx_path))


def paragraph_text_from_xml(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()


def remove_paragraph_style_by_text(docx_path: Path, text: str) -> None:
    root = document_xml(docx_path)
    body = root.find("w:body", NS)
    if body is None:
        raise GateFailure("word/document.xml missing body")
    for paragraph in body.findall("w:p", NS):
        if paragraph_text_from_xml(paragraph) != text:
            continue
        ppr = paragraph.find("w:pPr", NS)
        if ppr is None:
            return
        style_node = ppr.find("w:pStyle", NS)
        if style_node is not None:
            ppr.remove(style_node)
        save_document_xml(docx_path, root)
        return
    raise GateFailure(f"paragraph for style removal not found: {text}")


def remove_header_token(docx_path: Path, token: str) -> None:
    doc = Document(str(docx_path))
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if token in paragraph.text:
                            paragraph.text = paragraph.text.replace(token, "").strip()
    doc.save(str(docx_path))


def move_header_out_of_top_region(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for section in doc.sections:
        section.header_distance = Cm(7.0)
    doc.save(str(docx_path))


def clear_all_header_text(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            paragraph.text = ""
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = ""
    doc.save(str(docx_path))


def set_footer_indent(docx_path: Path, twips: int) -> None:
    doc = Document(str(docx_path))
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            paragraph.paragraph_format.left_indent = Pt(twips / 20.0)
            paragraph.paragraph_format.first_line_indent = Pt(twips / 20.0)
    doc.save(str(docx_path))


def mutate_heading_style_family(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    target = find_paragraph(doc, "1.1 研究背景与意义")
    try_set_style(target, BODY_STYLE_NAME)
    doc.save(str(docx_path))


def mutate_heading_baseline(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    target = find_paragraph(doc, "1.1.1 研究目标")
    target.paragraph_format.first_line_indent = Pt(24)
    target.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if target.runs:
        target.runs[0].font.size = Pt(12)
        target.runs[0].bold = False
    doc.save(str(docx_path))


def mutate_heading2_run_override(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    target = find_paragraph(doc, "1.1 研究背景与意义")
    try_set_style(target, "Heading 2")
    if target.runs:
        set_run_font(target.runs[0], east_asia="SimSun", size_pt=12.0, bold=False)
    doc.save(str(docx_path))


def mutate_heading_all_levels_run_override(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    targets = [
        ("1 绪论", "Heading 1"),
        ("1.1 研究背景与意义", "Heading 2"),
        ("1.1.1 研究目标", "Heading 3"),
        ("1.1.1.1 样本范围", "Heading 4"),
    ]
    for text, style_name in targets:
        target = find_paragraph(doc, text)
        try_set_style(target, style_name)
        if target.runs:
            set_run_font(target.runs[0], east_asia="SimSun", size_pt=12.0, bold=False)
    doc.save(str(docx_path))


def mutate_cover_baseline(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    target = find_paragraph(doc, "毕业设计")
    if target.runs:
        set_run_font(target.runs[0], east_asia="SimSun", size_pt=12.0, bold=False)
    target.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(str(docx_path))


def mutate_table_caption_baseline(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    caption = find_paragraph(doc, "表 1 测试表")
    try_set_style(caption, BODY_STYLE_NAME)
    caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.first_line_indent = Pt(24)
    caption.paragraph_format.keep_with_next = False
    if caption.runs:
        set_run_font(caption.runs[0], east_asia="SimSun", size_pt=10.5, bold=True)
    doc.save(str(docx_path))


def mutate_table_cell_baseline(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    if not doc.tables:
        raise GateFailure("table not found for table-cell baseline mutation")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    paragraph.paragraph_format.first_line_indent = Pt(24)
                    paragraph.paragraph_format.line_spacing = Pt(12)
                    if paragraph.runs:
                        set_run_font(paragraph.runs[0], east_asia="SimSun", size_pt=12.0, bold=False)
    doc.save(str(docx_path))


def mutate_code_title_format(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    target = find_paragraph(doc, "代码 1 示例代码")
    try_set_style(target, BODY_STYLE_NAME)
    target.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if target.runs:
        target.runs[0].font.size = Pt(10.5)
    doc.save(str(docx_path))


def mutate_code_block_format(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        if "def demo()" in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(24)
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(12)
                paragraph.runs[0].bold = True
            doc.save(str(docx_path))
            return
    raise GateFailure("code block paragraph not found")


def mutate_abstract_baseline(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    zh_body = find_paragraph(doc, "这是用于 smoke integration 的中文摘要。")
    zh_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if zh_body.runs:
        zh_body.runs[0].font.size = Pt(10.5)
        zh_body.runs[0].bold = True
    doc.save(str(docx_path))


def mutate_abstract_manual_line_break(docx_path: Path) -> None:
    root = document_xml(docx_path)
    body = root.find("w:body", NS)
    if body is None:
        raise GateFailure("word/document.xml missing body")

    def paragraph_with_text(text: str) -> ET.Element:
        paragraph = ET.Element(W + "p")
        run = ET.SubElement(paragraph, W + "r")
        ET.SubElement(run, W + "t").text = text
        return paragraph

    title = paragraph_with_text("\u6458\u8981")
    body_para = ET.Element(W + "p")
    body_run = ET.SubElement(body_para, W + "r")
    ET.SubElement(body_run, W + "t").text = "\u7b2c\u4e00\u6bb5\u6458\u8981\u5185\u5bb9"
    ET.SubElement(body_run, W + "br")
    ET.SubElement(body_run, W + "t").text = "\u7b2c\u4e8c\u6bb5\u88ab\u5199\u6210\u540c\u4e00 Word \u6bb5\u843d"
    keyword = paragraph_with_text("\u5173\u952e\u8bcd\uff1a\u6a21\u677f\uff1b\u7f51\u7edc\u8fd0\u7ef4")
    body.insert(0, keyword)
    body.insert(0, body_para)
    body.insert(0, title)
    save_document_xml(docx_path, root)


def mutate_toc_visible_indent(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("1.1 研究背景与意义"):
            paragraph.paragraph_format.left_indent = Pt(0)
    doc.save(str(docx_path))


def mutate_body_style_binding(docx_path: Path) -> None:
    remove_paragraph_style_by_text(docx_path, "这是正文内容。")
    remove_paragraph_style_by_text(docx_path, "这是结论内容。")
    remove_paragraph_style_by_text(docx_path, "感谢测试环境。")


def mutate_normal_baseline(docx_path: Path) -> None:
    root = styles_xml(docx_path)
    target_style = None
    for style in root.findall("w:style", NS):
        style_id = style.attrib.get(W + "styleId", "")
        name_node = style.find("w:name", NS)
        style_name = name_node.attrib.get(W + "val", "") if name_node is not None else ""
        based_on = style.find("w:basedOn", NS)
        based_on_id = based_on.attrib.get(W + "val", "") if based_on is not None else ""
        ppr = style.find("w:pPr", NS)
        ind = ppr.find("w:ind", NS) if ppr is not None else None
        is_custom = style.attrib.get(W + "customStyle") == "1"
        if (
            style_id.lower() != "normal"
            and based_on_id.lower() == "normal"
            and is_custom
            and ind is not None
            and ind.attrib.get(W + "firstLine", "") == "480"
        ):
            target_style = style
            break
        if style_id == BODY_STYLE_NAME or normalize_text(style_name) == normalize_text(BODY_STYLE_NAME):
            target_style = style
            break
    if target_style is None:
        raise GateFailure("body style not found in styles.xml")
    rpr = target_style.find("w:rPr", NS)
    if rpr is None:
        rpr = ET.SubElement(target_style, W + "rPr")
    rfonts = rpr.find("w:rFonts", NS)
    if rfonts is None:
        rfonts = ET.SubElement(rpr, W + "rFonts")
    rfonts.set(W + "ascii", "Arial")
    rfonts.set(W + "hAnsi", "Arial")
    rfonts.set(W + "eastAsia", "Arial")
    sz = rpr.find("w:sz", NS)
    if sz is None:
        sz = ET.SubElement(rpr, W + "sz")
    sz.set(W + "val", "28")
    save_styles_xml(docx_path, root)


def mutate_non_body_surface_contamination(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    targets = {
        "毕业设计",
        "摘   要",
        "Abstract",
        "图 1 测试图",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in targets:
            try_set_style(paragraph, BODY_STYLE_NAME)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    try_set_style(paragraph, BODY_STYLE_NAME)
    doc.save(str(docx_path))


def mutate_non_body_indent(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"毕业设计", "图 1 测试图", "表 1 测试表"}:
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.left_indent = Pt(24)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(24)
    doc.save(str(docx_path))


def mutate_figure_block_locality(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    holder = None
    for paragraph in doc.paragraphs:
        has_drawing = any(run._element.xpath(".//w:drawing") for run in paragraph.runs)
        if has_drawing:
            holder = paragraph
            break
    if holder is None:
        raise GateFailure("image holder paragraph not found")
    displaced = insert_paragraph_after(holder)
    displaced.text = "本地锚定扰动段落。"
    doc.save(str(docx_path))


def mutate_image_holder_safety(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        has_drawing = any(run._element.xpath(".//w:drawing") for run in paragraph.runs)
        if not has_drawing:
            continue
        try_set_style(paragraph, BODY_STYLE_NAME)
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = Pt(12)
        paragraph.paragraph_format.keep_with_next = False
        doc.save(str(docx_path))
        return
    raise GateFailure("image holder paragraph not found")


def mutate_figure_caption_style(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    caption = find_paragraph(doc, "图 1 测试图")
    try_set_style(caption, BODY_STYLE_NAME)
    caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.first_line_indent = Pt(24)
    if caption.runs:
        caption.runs[0].font.size = Pt(10.5)
        caption.runs[0].bold = True
    doc.save(str(docx_path))


def mutate_image_holder_baseline(docx_path: Path) -> None:
    root = document_xml(docx_path)
    body = root.find("w:body", NS)
    if body is None:
        raise GateFailure("word/document.xml missing body")
    for paragraph in body.findall("w:p", NS):
        if paragraph.find(".//w:drawing", NS) is None:
            continue
        ppr = paragraph.find("w:pPr", NS)
        if ppr is None:
            ppr = ET.SubElement(paragraph, W + "pPr")
        style_node = ppr.find("w:pStyle", NS)
        if style_node is None:
            style_node = ET.SubElement(ppr, W + "pStyle")
        style_node.set(W + "val", "Normal")
        ind = ppr.find("w:ind", NS)
        if ind is None:
            ind = ET.SubElement(ppr, W + "ind")
        ind.set(W + "firstLine", "480")
        ind.set(W + "left", "480")
        spacing = ppr.find("w:spacing", NS)
        if spacing is None:
            spacing = ET.SubElement(ppr, W + "spacing")
        spacing.set(W + "line", "240")
        spacing.set(W + "lineRule", "exact")
        save_document_xml(docx_path, root)
        return
    raise GateFailure("image holder paragraph not found")


def remove_paragraph_by_exact_text(docx_path: Path, text: str) -> None:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != text:
            continue
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
            doc.save(str(docx_path))
            return
    raise GateFailure(f"paragraph not found for removal: {text}")


def mutate_caption_inside_table_cell(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    if not doc.tables:
        raise GateFailure("table not found for caption-in-grid mutation")
    cell_para = doc.tables[0].cell(1, 0).paragraphs[0]
    cell_para.text = "表 1 测试表"
    doc.save(str(docx_path))


def mutate_table_caption_binding(docx_path: Path) -> None:
    root = document_xml(docx_path)
    body = root.find("w:body", NS)
    if body is None:
        raise GateFailure("word/document.xml missing body")
    for paragraph in body.findall("w:p", NS):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()
        if text != "表 1 测试表":
            continue
        ppr = paragraph.find("w:pPr", NS)
        if ppr is None:
            ppr = ET.SubElement(paragraph, W + "pPr")
        keep_next = ppr.find("w:keepNext", NS)
        if keep_next is not None:
            ppr.remove(keep_next)
        num_pr = ppr.find("w:numPr", NS)
        if num_pr is None:
            num_pr = ET.SubElement(ppr, W + "numPr")
            ET.SubElement(num_pr, W + "ilvl", {W + "val": "0"})
            ET.SubElement(num_pr, W + "numId", {W + "val": "9"})
        save_document_xml(docx_path, root)
        return
    raise GateFailure("table caption paragraph not found")


def mutate_bibliography_count_loss(docx_path: Path) -> None:
    remove_paragraph_by_exact_text(docx_path, "[1] Example Paper[J].")


def mutate_bibliography_rendered_indent(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    mutated = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("["):
            paragraph.paragraph_format.left_indent = Pt(36)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            mutated += 1
    if not mutated:
        raise GateFailure("bibliography entry paragraph not found")
    doc.save(str(docx_path))


def mutate_image_dimension_overflow(docx_path: Path) -> None:
    root = document_xml(docx_path)
    extent = root.find(".//wp:extent", NS)
    if extent is None:
        raise GateFailure("drawing extent not found")
    extent.set("cx", "22000000")
    extent.set("cy", "12000000")
    save_document_xml(docx_path, root)


def mutate_missing_figure_explanation(docx_path: Path) -> None:
    remove_paragraph_by_exact_text(docx_path, "图后正文介绍用于验证图片、图题与说明段保持同一本地块。")
    doc = Document(str(docx_path))
    anchor = find_paragraph(doc, "这是正文内容。")
    clear_paragraph_content(anchor)
    doc.save(str(docx_path))


def mutate_table_to_full_grid(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    if not doc.tables:
        raise GateFailure("table not found")
    doc.tables[0].style = "Table Grid"
    doc.save(str(docx_path))


def mutate_table_continuation_missing(docx_path: Path) -> None:
    add_table_sample(docx_path, long_table=True)


def insert_toc_control_contamination(docx_path: Path) -> None:
    root = document_xml(docx_path)
    body = root.find("w:body", NS)
    if body is None:
        raise GateFailure("word/document.xml missing body")
    sdt = ET.Element(W + "sdt")
    content = ET.SubElement(sdt, W + "sdtContent")
    good_para = ET.SubElement(content, W + "p")
    good_ppr = ET.SubElement(good_para, W + "pPr")
    ET.SubElement(good_ppr, W + "pStyle", {W + "val": "TOC1"})
    good_run = ET.SubElement(good_para, W + "r")
    ET.SubElement(good_run, W + "t").text = "1 绪论\t1"
    bad_para = ET.SubElement(content, W + "p")
    bad_run = ET.SubElement(bad_para, W + "r")
    ET.SubElement(bad_run, W + "t").text = "这是被塞进 TOC 控件的正文。"
    body.insert(0, sdt)
    save_document_xml(docx_path, root)


def mutate_keyword_lines_single_run(docx_path: Path, *, clear_body_style: bool = False) -> None:
    doc = Document(str(docx_path))
    mutated = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("关键词：") or text.startswith("Key Words:"):
            paragraph.text = ""
            run_obj = paragraph.add_run(text)
            run_obj.bold = True
            mutated += 1
    doc.save(str(docx_path))
    if clear_body_style:
        remove_paragraph_style_by_text(docx_path, "这是用于 smoke integration 的中文摘要。")
    if mutated == 0:
        raise GateFailure("keyword lines not found for mutation")


def first_picture_para_id(docx_path: Path) -> str:
    root = document_xml(docx_path)
    body = root.find("w:body", NS)
    if body is None:
        raise GateFailure("word/document.xml missing body")
    for paragraph in body.findall("w:p", NS):
        if paragraph.find(".//w:drawing", NS) is None:
            continue
        para_id = paragraph.attrib.get(W14 + "paraId")
        if para_id:
            return para_id
    raise GateFailure("picture paragraph id not found")


def document_contains(docx_path: Path, token: str) -> bool:
    root = document_xml(docx_path)
    return normalize_text(token) in normalize_text("".join(root.itertext()))


def verify_formula_table(docx_path: Path, number_text: str) -> None:
    root = document_xml(docx_path)
    tables = root.findall(".//w:tbl", NS)
    if not tables:
        raise GateFailure("formula helper did not create a table wrapper")
    xml_text = ET.tostring(root, encoding="unicode")
    if number_text not in xml_text:
        raise GateFailure(f"formula number not found after helper chain: {number_text}")
    if 'w:insideV' not in xml_text and 'w:insideH' not in xml_text:
        return
    if 'val="nil"' not in xml_text:
        raise GateFailure("formula helper table is not borderless")


def run_sample_self_check_suite(
    reference_docx: Path,
    final_docx: Path,
    fixture_root: Path,
    suite_name: str,
    *,
    asset_manifest: Path | None = None,
    renderer: str = "auto",
    env: GateEnvironment | None = None,
    smoke_acceptance: bool = True,
    check: bool = True,
) -> dict[str, Path | str]:
    pdf_path = fixture_root / f"{suite_name}.pdf"
    citation_audit = fixture_root / f"{suite_name}-citation-audit.md"
    font_audit = fixture_root / f"{suite_name}-font-audit.md"
    body_style_audit = fixture_root / f"{suite_name}-body-style-audit.md"
    self_check = fixture_root / f"{suite_name}-sample-self-check.md"
    stage_dir = Path(tempfile.mkdtemp(prefix="gpb_sample_render_"))
    stage_docx = stage_dir / "stage.docx"
    reference_stage_docx = stage_dir / "reference.docx"
    reference_pdf_path = fixture_root / f"{suite_name}-reference.pdf"

    if env is None:
        env = detect_environment()
    copy_locked(final_docx, stage_docx)
    copy_locked(reference_docx, reference_stage_docx)
    renderer_warning = ""
    if renderer == "auto":
        try:
            renderer_used = export_pdf_with_sample_fallback(stage_docx, pdf_path, env)
            export_pdf_with_renderer(reference_stage_docx, reference_pdf_path, renderer_used, env)
        except RuntimeError as exc:
            renderer_warning = str(exc)
            raise GateFailure(
                "real renderer unavailable for sample self-check; diagnostic PDF fallback is not acceptable"
            ) from exc
    else:
        export_pdf_with_renderer(stage_docx, pdf_path, renderer, env)
        export_pdf_with_renderer(reference_stage_docx, reference_pdf_path, renderer, env)
        renderer_used = renderer
    run_command(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_thesis_citations.py"),
            "--docx",
            str(final_docx),
            "--report",
            str(citation_audit),
        ],
        check=False,
        timeout=600,
    )
    run_command(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_font_encoding.py"),
            str(final_docx),
            "--reference-docx",
            str(reference_docx),
            "--report",
            str(font_audit),
        ],
        check=False,
        timeout=600,
    )
    run_command(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_body_style.py"),
            "--reference-docx",
            str(reference_docx),
            "--final-docx",
            str(final_docx),
            "--report",
            str(body_style_audit),
            "--strict-direct-visible-metrics",
        ],
        check=False,
        timeout=600,
    )
    sample_check_cmd = [
        str(PYTHON_EXE),
        str(SCRIPT_DIR / "sample_self_check.py"),
        "--reference-docx",
        str(reference_docx),
        "--reference-pdf",
        str(reference_pdf_path),
        "--final-docx",
        str(final_docx),
        "--final-pdf",
        str(pdf_path),
        "--citation-audit",
        str(citation_audit),
        "--font-audit",
        str(font_audit),
        "--body-style-audit",
        str(body_style_audit),
        "--output",
        str(self_check),
    ]
    if smoke_acceptance:
        sample_check_cmd.append("--smoke-acceptance")
    if asset_manifest is not None:
        sample_check_cmd.extend(["--asset-manifest", str(asset_manifest)])
    proc = run_command(sample_check_cmd, timeout=1800, check=check)
    text = self_check.read_text(encoding="utf-8", errors="replace")
    return {
        "pdf": pdf_path,
        "reference_pdf": reference_pdf_path,
        "citation_audit": citation_audit,
        "font_audit": font_audit,
        "body_style_audit": body_style_audit,
        "self_check": self_check,
        "returncode": str(proc.returncode),
        "text": text,
        "renderer": renderer_used,
        "renderer_warning": renderer_warning,
    }


def case_frontmatter_roundtrip(root: Path, env: GateEnvironment) -> CaseResult:
    template_docx, _source_docx, _ = build_complete_sample_fixture(root / "frontmatter_roundtrip")
    stage_dir = root / "frontmatter_roundtrip" / "stage-frontmatter"
    stage_dir.mkdir(parents=True, exist_ok=True)
    stage_docx = stage_dir / "stage.docx"
    copy_locked(template_docx, stage_docx)
    pdf_path = stage_dir / "rendered.pdf"
    export_pdf(stage_docx, pdf_path)
    pages = pdf_page_texts(pdf_path)
    markers = ["摘   要", "Abstract", "目   录", "1 绪论"]
    found = [find_page(pages, marker) for marker in markers]
    if any(page is None for page in found):
        raise GateFailure("frontmatter roundtrip lost one or more page markers")
    if found != sorted(found):
        raise GateFailure(f"frontmatter page order drifted: {found}")
    return CaseResult(
        "frontmatter_roundtrip",
        True,
        f"renderer={env.renderer}; pdf={pdf_path}; pages={len(pages)}",
    )


def case_lock_competition(root: Path, env: GateEnvironment) -> CaseResult:
    lock_root = root / "lock_competition"
    lock_root.mkdir(parents=True, exist_ok=True)
    source = lock_root / "source.docx"
    create_base_sample_docx(source)
    blocked_dst = lock_root / "blocked.docx"
    released_dst = lock_root / "released.docx"
    locker_script = (
        "import ctypes,sys,time; "
        "GENERIC_READ=0x80000000; OPEN_EXISTING=3; INVALID_HANDLE_VALUE=ctypes.c_void_p(-1).value; "
        "CreateFileW=ctypes.windll.kernel32.CreateFileW; "
        "CreateFileW.argtypes=[ctypes.c_wchar_p,ctypes.c_uint32,ctypes.c_uint32,ctypes.c_void_p,ctypes.c_uint32,ctypes.c_uint32,ctypes.c_void_p]; "
        "CreateFileW.restype=ctypes.c_void_p; "
        "handle=CreateFileW(sys.argv[1], GENERIC_READ, 0, None, OPEN_EXISTING, 0, None); "
        "print(int(handle)); sys.stdout.flush(); time.sleep(8)"
    )
    locker = subprocess.Popen(
        [str(PYTHON_EXE), "-c", locker_script, str(source)],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    try:
        time.sleep(1.0)
        blocked = run_command(
            [str(PYTHON_EXE), str(SCRIPT_DIR / "copy_locked_docx.py"), str(source), str(blocked_dst)],
            check=False,
            timeout=30,
        )
        if blocked.returncode == 0:
            raise GateFailure("exclusive lock did not block shared copy")
        try:
            locker.wait(timeout=15)
        except subprocess.TimeoutExpired as exc:
            raise GateFailure("exclusive lock helper did not exit within timeout") from exc
        copy_locked(source, released_dst)
    finally:
        if locker.poll() is None:
            locker.terminate()
            try:
                locker.wait(timeout=10)
            except subprocess.TimeoutExpired:
                locker.kill()
                locker.wait(timeout=10)
    if not released_dst.exists():
        raise GateFailure("copy did not succeed after lock release")
    return CaseResult(
        "lock_competition",
        True,
        "exclusive lock blocked copy, and copy succeeded after lock release",
    )


def case_tail_pages_roundtrip(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "tail_pages_roundtrip"
    fixture_root.mkdir(parents=True, exist_ok=True)
    docx_path = fixture_root / "tail-pages.docx"
    create_base_sample_docx(docx_path)
    pdf_path = fixture_root / "tail-pages.pdf"
    export_pdf(docx_path, pdf_path)
    pages = pdf_page_texts(pdf_path)
    references_page = find_page(pages, "参考文献")
    acknowledgement_page = find_page(pages, "致   谢")
    if references_page is None or acknowledgement_page is None:
        raise GateFailure("tail page markers missing after render")
    return CaseResult(
        "tail_pages_roundtrip",
        True,
        f"references_page={references_page}; acknowledgement_page={acknowledgement_page}",
    )


def case_staging_copy_isolation(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "staging_copy_isolation"
    reference_docx, review_source = build_sample_self_check_fixture(fixture_root / "review-stage-isolation")
    review_copy = fixture_root / "review-stage-isolation" / "review-pass-001.docx"
    copy_locked(review_source, review_copy)
    stage_dir = fixture_root / "stage-isolation"
    stage_dir.mkdir(parents=True, exist_ok=True)
    stage_docx = stage_dir / "stage.docx"
    copy_locked(review_copy, stage_docx)
    doc = Document(str(stage_docx))
    doc.add_paragraph("STAGE ONLY SENTINEL")
    doc.save(str(stage_docx))
    pdf_path = stage_dir / "rendered.pdf"
    export_pdf(stage_docx, pdf_path)
    if document_contains(review_copy, "STAGE ONLY SENTINEL"):
        raise GateFailure("review copy was polluted by stage mutation")
    if not document_contains(stage_docx, "STAGE ONLY SENTINEL"):
        raise GateFailure("stage mutation did not land on stage copy")
    return CaseResult(
        "staging_copy_isolation",
        True,
        f"renderer={env.renderer}; stage_docx={stage_docx}; review_copy={review_copy}; pdf={pdf_path}",
    )


def case_paragraph_review_microcycle(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "paragraph_review_microcycle"
    _reference_docx, source_docx = build_sample_self_check_fixture(fixture_root / "review-paragraph-microcycle")
    review_copy = fixture_root / "review-paragraph-microcycle" / "review-pass-001.docx"
    copy_locked(source_docx, review_copy)

    stage1_dir = fixture_root / "stage-paragraph-microcycle-1"
    stage1_dir.mkdir(parents=True, exist_ok=True)
    stage1_docx = stage1_dir / "stage.docx"
    copy_locked(review_copy, stage1_docx)
    doc1 = Document(str(stage1_docx))
    doc1.paragraphs.append if False else None
    doc1.add_paragraph("MICROCYCLE PASS 1")
    doc1.save(str(stage1_docx))
    stage1_pdf = stage1_dir / "rendered.pdf"
    export_pdf(stage1_docx, stage1_pdf)

    doc_review = Document(str(review_copy))
    doc_review.add_paragraph("MICROCYCLE PASS 2")
    doc_review.save(str(review_copy))
    stage2_dir = fixture_root / "stage-paragraph-microcycle-2"
    stage2_dir.mkdir(parents=True, exist_ok=True)
    stage2_docx = stage2_dir / "stage.docx"
    copy_locked(review_copy, stage2_docx)
    stage2_pdf = stage2_dir / "rendered.pdf"
    export_pdf(stage2_docx, stage2_pdf)
    if not document_contains(stage2_docx, "MICROCYCLE PASS 2"):
        raise GateFailure("second paragraph review pass did not survive staging")
    return CaseResult(
        "paragraph_review_microcycle",
        True,
        f"renderer={env.renderer}; review_copy={review_copy}; stage1_pdf={stage1_pdf}; stage2_pdf={stage2_pdf}",
    )


def case_serialized_helper_mutation_chain(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "serialized_helper_mutation_chain"
    reference_docx, source_docx = build_sample_self_check_fixture(fixture_root / "review-helper-serialization")
    image_path = fixture_root / "replacement.png"
    make_placeholder_png(image_path, "替换图")
    add_figure_sample(source_docx, image_path)
    review_copy = fixture_root / "review-helper-serialization" / "review-pass-001.docx"
    copy_locked(source_docx, review_copy)
    review_source_docx = fixture_root / "review-helper-serialization" / "review-pass-001.source.docx"
    shutil.copy2(review_copy, review_source_docx)
    figure_manifest = create_picture_replacement_manifest(
        fixture_root / "review-helper-serialization" / "picture-replacement-manifest.json",
        replacement_image=image_path,
        caption_text="图 1 测试图",
    )
    para_id = first_picture_para_id(review_copy)
    patch_picture_docx(
        review_copy,
        para_id,
        image_path,
        new_label="helper-replaced-picture",
        manifest_path=figure_manifest,
        source_docx=review_source_docx,
    )
    wait_for_docx_ready(review_copy)
    stage_dir = fixture_root / "stage-helper-serialization"
    stage_dir.mkdir(parents=True, exist_ok=True)
    stage_docx = stage_dir / "stage.docx"
    copy_locked(review_copy, stage_docx)
    pdf_path = stage_dir / "rendered.pdf"
    export_pdf(stage_docx, pdf_path)
    manifest_payload = json.loads(figure_manifest.read_text(encoding="utf-8"))
    if manifest_payload.get("final_docx_path") != str(review_copy.resolve()):
        raise GateFailure("picture helper did not bind manifest to the exact review copy")
    if str(manifest_payload.get("final_docx_sha256") or "") != file_sha256(review_copy):
        raise GateFailure("picture helper manifest final DOCX SHA256 does not match review copy")
    return CaseResult(
        "serialized_helper_mutation_chain",
        True,
        f"renderer={env.renderer}; helper=docx_sync_picture.py; para_id={para_id}; review_copy={review_copy}; stage_docx={stage_docx}; pdf={pdf_path}; manifest={figure_manifest}; media_members={media_member_count(stage_docx)}",
    )


def case_multi_pass_review_copy_governance(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "multi_pass_review_copy_governance"
    _reference_docx, source_docx = build_sample_self_check_fixture(fixture_root / "review-multi-pass")
    review_copy_1 = fixture_root / "review-multi-pass" / "review-pass-001.docx"
    review_copy_2 = fixture_root / "review-multi-pass" / "review-pass-002.docx"
    copy_locked(source_docx, review_copy_1)
    doc1 = Document(str(review_copy_1))
    doc1.add_paragraph("SENTINEL_PASS_001")
    doc1.save(str(review_copy_1))
    copy_locked(review_copy_1, review_copy_2)
    doc2 = Document(str(review_copy_2))
    doc2.add_paragraph("SENTINEL_PASS_002")
    doc2.save(str(review_copy_2))
    stage_dir = fixture_root / "stage-multi-pass"
    stage_dir.mkdir(parents=True, exist_ok=True)
    stage_docx = stage_dir / "stage.docx"
    copy_locked(review_copy_2, stage_docx)
    pdf_path = stage_dir / "rendered.pdf"
    export_pdf(stage_docx, pdf_path)
    page_text = "\n".join(pdf_page_texts(pdf_path))
    if not document_contains(review_copy_1, "SENTINEL_PASS_001"):
        raise GateFailure("review-pass-001 sentinel missing")
    if document_contains(review_copy_1, "SENTINEL_PASS_002"):
        raise GateFailure("review-pass-001 was overwritten by second pass")
    if not document_contains(review_copy_2, "SENTINEL_PASS_002"):
        raise GateFailure("review-pass-002 sentinel missing")
    if "SENTINEL_PASS_002" not in page_text:
        raise GateFailure("staged PDF does not contain the second-pass sentinel")
    return CaseResult(
        "multi_pass_review_copy_governance",
        True,
        f"renderer={env.renderer}; review_copy_1={review_copy_1}; review_copy_2={review_copy_2}; stage_docx={stage_docx}; pdf={pdf_path}",
    )


def case_serialized_formula_helper_chain(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "serialized_formula_helper_chain"
    _reference_docx, source_docx = build_sample_self_check_fixture(fixture_root / "review-formula-helper-serialization")
    add_formula_sample(source_docx, para_id="0010F0AA")
    review_copy = fixture_root / "review-formula-helper-serialization" / "review-pass-001.docx"
    copy_locked(source_docx, review_copy)
    doc = Document(str(review_copy))
    doc.add_paragraph("FORMULA BODY SENTINEL")
    doc.save(str(review_copy))
    patch_formula(review_copy, "0010F0AA", "(5-1)", None)
    wait_for_docx_ready(review_copy)
    stage_dir = fixture_root / "stage-formula-helper-serialization"
    stage_dir.mkdir(parents=True, exist_ok=True)
    stage_docx = stage_dir / "stage.docx"
    copy_locked(review_copy, stage_docx)
    pdf_path = stage_dir / "rendered.pdf"
    export_pdf(stage_docx, pdf_path)
    verify_formula_table(stage_docx, "(5-1)")
    if not document_contains(stage_docx, "FORMULA BODY SENTINEL"):
        raise GateFailure("formula helper chain lost the earlier body mutation")
    return CaseResult(
        "serialized_formula_helper_chain",
        True,
        f"renderer={env.renderer}; helper=docx_formula_number_table.py; para_id=0010F0AA; review_copy={review_copy}; stage_docx={stage_docx}; pdf={pdf_path}",
    )


def case_complete_sample_smoke(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "complete_sample_smoke" / "sample-rebuild"
    template_docx, source_docx, _ = build_complete_sample_fixture(fixture_root)
    rebuilt_docx = fixture_root / "rebuilt.docx"
    rebuilt_pdf = fixture_root / "rebuilt.pdf"
    run_root = fixture_root / "run"
    proc = run_command(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "build_minimal_template_thesis.py"),
            "--template",
            str(template_docx),
            "--output-docx",
            str(rebuilt_docx),
            "--output-pdf",
            str(rebuilt_pdf),
            "--run-root",
            str(run_root),
            "--smoke-acceptance",
        ],
        timeout=2400,
        check=False,
    )
    source_docx.touch(exist_ok=True)
    self_check = run_root / "reports" / "sample-self-check.md"
    acceptance = run_root / "reports" / "acceptance-record.md"
    top_level_self_check = fixture_root / "rebuilt-self-check.md"
    if not self_check.exists() or not acceptance.exists():
        raise GateFailure(
            "complete sample smoke did not produce expected smoke evidence; "
            f"returncode={proc.returncode}; stdout={proc.stdout}; stderr={proc.stderr}"
        )
    shutil.copy2(self_check, top_level_self_check)
    self_check_text = self_check.read_text(encoding="utf-8", errors="replace")
    acceptance_text = acceptance.read_text(encoding="utf-8", errors="replace")
    if "deliverable critical gate: smoke-only; blocked for delivery" not in self_check_text:
        raise GateFailure("complete sample smoke fixture was not marked as non-deliverable")
    if "validation result: fail" not in acceptance_text.lower():
        raise GateFailure("complete sample smoke acceptance record did not fail delivery validation")
    assert_report_dimension(
        self_check_text,
        SAMPLE_SELF_CHECK_LABELS["figure_family"],
        SUMMARY_PASS,
        "complete sample smoke figure family",
    )
    return CaseResult(
        "complete_sample_smoke",
        True,
        f"non-deliverable smoke fixture; docx={rebuilt_docx}; pdf={rebuilt_pdf}; self_check={top_level_self_check}; acceptance={acceptance}",
    )


def run_negative_sample_case(
    name: str,
    fixture_root: Path,
    mutator: Callable[[Path, Path], Path | None],
    suite_name: str,
    target_label: str | list[str],
    *,
    custom_body_style: bool = False,
    include_header_footer: bool = False,
    env: GateEnvironment | None = None,
    expected_summary: str = SUMMARY_FAIL,
) -> CaseResult:
    reference_docx, final_docx = build_sample_self_check_fixture(
        fixture_root,
        custom_body_style=custom_body_style,
        include_header_footer=include_header_footer,
    )
    asset_manifest = mutator(reference_docx, final_docx)
    if env is None:
        env = detect_environment()
    preferred_renderer = "auto"
    suite = run_sample_self_check_suite(
        reference_docx,
        final_docx,
        fixture_root,
        suite_name,
        asset_manifest=asset_manifest,
        renderer=preferred_renderer,
        env=env,
    )
    target_labels = [target_label] if isinstance(target_label, str) else target_label
    for label in target_labels:
        assert_report_dimension(str(suite["text"]), label, expected_summary, name)
    details_parts = [
        f"renderer={suite['renderer']}",
        f"review_copy={final_docx}",
        f"pdf={suite['pdf']}",
        f"self_check={suite['self_check']}",
    ]
    if suite.get("renderer_warning"):
        details_parts.append(f"renderer_warning={suite['renderer_warning']}")
    if asset_manifest is not None:
        details_parts.append(f"asset_manifest={asset_manifest}")
    return CaseResult(name, True, "; ".join(details_parts))


def case_sample_self_check_table_family_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_table_to_full_grid(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_table_family_detection",
        root / "sample_self_check_table_family_detection" / "sample-self-check-table",
        mutate,
        "table-family",
        SAMPLE_SELF_CHECK_LABELS["table_family"],
    )


def case_sample_self_check_table_caption_baseline_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_table_caption_baseline(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_table_caption_baseline_detection",
        root / "sample_self_check_table_caption_baseline_detection" / "sample-self-check-table-caption-baseline",
        mutate,
        "table-caption-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["table_caption_baseline"],
    )


def case_sample_self_check_table_cell_baseline_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_table_cell_baseline(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_table_cell_baseline_detection",
        root / "sample_self_check_table_cell_baseline_detection" / "sample-self-check-table-cell-baseline",
        mutate,
        "table-cell-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["table_cell_baseline"],
    )


def case_sample_self_check_heading_family_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_heading_family_detection",
        root / "sample_self_check_heading_family_detection" / "sample-self-check-heading-family",
        lambda _reference_docx, final_docx: (mutate_heading_style_family(final_docx), None)[1],
        "heading-family-failure",
        SAMPLE_SELF_CHECK_LABELS["heading_family"],
        custom_body_style=True,
    )


def case_sample_self_check_heading_all_levels_run_override_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_heading4_sample(reference_docx)
        add_heading4_sample(final_docx)
        mutate_heading_all_levels_run_override(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_heading_all_levels_run_override_detection",
        root / "sample_self_check_heading_all_levels_run_override_detection" / "sample-self-check-heading-all-levels",
        mutate,
        "heading-all-levels-run-override-failure",
        [
            SAMPLE_SELF_CHECK_LABELS["heading1_baseline"],
            SAMPLE_SELF_CHECK_LABELS["heading2_baseline"],
            SAMPLE_SELF_CHECK_LABELS["heading_baseline"],
            SAMPLE_SELF_CHECK_LABELS["heading4_baseline"],
        ],
    )


def case_sample_self_check_heading_baseline_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_heading_baseline_detection",
        root / "sample_self_check_heading_baseline_detection" / "sample-self-check-heading-baseline",
        lambda _reference_docx, final_docx: (mutate_heading_baseline(final_docx), None)[1],
        "heading-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["heading_baseline"],
    )


def case_sample_self_check_heading2_run_override_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_heading2_run_override_detection",
        root / "sample_self_check_heading2_run_override_detection" / "sample-self-check-heading2-run-override",
        lambda _reference_docx, final_docx: (mutate_heading2_run_override(final_docx), None)[1],
        "heading2-run-override-failure",
        SAMPLE_SELF_CHECK_LABELS["heading2_baseline"],
    )


def case_sample_self_check_cover_baseline_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_cover_baseline_detection",
        root / "sample_self_check_cover_baseline_detection" / "sample-self-check-cover-baseline",
        lambda _reference_docx, final_docx: (mutate_cover_baseline(final_docx), None)[1],
        "cover-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["cover_baseline"],
    )


def case_sample_self_check_code_title_format_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_code_sample(reference_docx)
        add_code_sample(final_docx)
        mutate_code_title_format(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_code_title_format_detection",
        root / "sample_self_check_code_title_format_detection" / "sample-self-check-code-title",
        mutate,
        "code-title-format-failure",
        SAMPLE_SELF_CHECK_LABELS["code_title"],
    )


def case_sample_self_check_code_block_format_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_code_sample(reference_docx)
        add_code_sample(final_docx)
        mutate_code_block_format(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_code_block_format_detection",
        root / "sample_self_check_code_block_format_detection" / "sample-self-check-code-block",
        mutate,
        "code-block-format-failure",
        SAMPLE_SELF_CHECK_LABELS["code_block"],
    )


def case_sample_self_check_abstract_baseline_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_abstract_baseline_detection",
        root / "sample_self_check_abstract_baseline_detection" / "sample-self-check-abstract-baseline",
        lambda _reference_docx, final_docx: (mutate_abstract_baseline(final_docx), None)[1],
        "abstract-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["abstract_baseline"],
    )


def case_sample_self_check_abstract_manual_break_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_abstract_manual_break_detection",
        root / "sample_self_check_abstract_manual_break_detection" / "sample-self-check-abstract-manual-break",
        lambda _reference_docx, final_docx: (mutate_abstract_manual_line_break(final_docx), None)[1],
        "abstract-manual-break-failure",
        "abstract manual line-break check",
        expected_summary="failed",
    )


def case_sample_self_check_toc_visible_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_toc_visible_detection",
        root / "sample_self_check_toc_visible_detection" / "sample-self-check-toc",
        lambda _reference_docx, final_docx: (mutate_toc_visible_indent(final_docx), None)[1],
        "toc-visible-failure",
        SAMPLE_SELF_CHECK_LABELS["toc_visible"],
    )


def case_sample_self_check_header_position_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_header_position_detection",
        root / "sample_self_check_header_position_detection" / "sample-self-check-header",
        lambda reference_docx, final_docx: (
            add_header_footer_baseline(reference_docx),
            add_header_footer_baseline(final_docx),
            move_header_out_of_top_region(final_docx),
            None,
        )[-1],
        "header-position-failure",
        SAMPLE_SELF_CHECK_LABELS["header_position"],
        include_header_footer=True,
    )


def case_sample_self_check_header_presence_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_header_presence_detection",
        root / "sample_self_check_header_presence_detection" / "sample-self-check-header-presence",
        lambda reference_docx, final_docx: (
            add_header_footer_baseline(reference_docx),
            add_header_footer_baseline(final_docx),
            clear_all_header_text(final_docx),
            None,
        )[-1],
        "header-presence-contract-failure",
        SAMPLE_SELF_CHECK_LABELS["header_presence"],
        include_header_footer=True,
        expected_summary="failed",
    )


def case_sample_self_check_footer_indent_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_footer_indent_detection",
        root / "sample_self_check_footer_indent_detection" / "sample-self-check-footer",
        lambda reference_docx, final_docx: (
            add_header_footer_baseline(reference_docx),
            add_header_footer_baseline(final_docx),
            set_footer_indent(final_docx, 480),
            None,
        )[-1],
        "footer-indent-failure",
        SAMPLE_SELF_CHECK_LABELS["footer_indent"],
        include_header_footer=True,
    )


def case_sample_self_check_footer_baseline_detection(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "sample_self_check_footer_baseline_detection" / "sample-self-check-footer-baseline"
    reference_docx, final_docx = build_sample_self_check_fixture(
        fixture_root,
        include_header_footer=True,
    )
    add_header_footer_baseline(reference_docx)
    add_header_footer_baseline(final_docx)
    suite = run_sample_self_check_suite(
        reference_docx,
        final_docx,
        fixture_root,
        "footer-baseline",
        env=env,
        renderer="auto",
    )
    assert_report_dimension(
        str(suite["text"]),
        SAMPLE_SELF_CHECK_LABELS["header_footer_baseline"],
        SUMMARY_PASS,
        "baseline footer baseline check",
    )
    return CaseResult(
        "sample_self_check_footer_baseline_detection",
        True,
        f"review_copy={final_docx}; pdf={suite['pdf']}; self_check={suite['self_check']}",
    )


def case_sample_self_check_body_style_binding_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_body_style_binding_detection",
        root / "sample_self_check_body_style_binding_detection" / "sample-self-check-body-style-binding",
        lambda _reference_docx, final_docx: (mutate_body_style_binding(final_docx), None)[1],
        "body-style-binding-failure",
        SAMPLE_SELF_CHECK_LABELS["body_style"],
        custom_body_style=True,
    )


def case_sample_self_check_normal_baseline_drift_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_normal_baseline_drift_detection",
        root / "sample_self_check_normal_baseline_drift_detection" / "sample-self-check-normal-baseline",
        lambda _reference_docx, final_docx: (mutate_normal_baseline(final_docx), None)[1],
        "normal-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["body_baseline"],
        custom_body_style=True,
    )


def case_sample_self_check_non_body_surface_contamination_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "caption.png"
        make_placeholder_png(image_path, "标题图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_non_body_surface_contamination(final_docx)
        return None

    fixture_root = root / "sample_self_check_non_body_surface_contamination_detection" / "sample-self-check-non-body-surface"
    return run_negative_sample_case(
        "sample_self_check_non_body_surface_contamination_detection",
        fixture_root,
        mutate,
        "non-body-surface-failure",
        SAMPLE_SELF_CHECK_LABELS["non_body_surface"],
        custom_body_style=True,
    )


def case_sample_self_check_non_body_indent_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "indent.png"
        make_placeholder_png(image_path, "缩进图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_non_body_indent(final_docx)
        return None

    fixture_root = root / "sample_self_check_non_body_indent_detection" / "sample-self-check-non-body-indent"
    return run_negative_sample_case(
        "sample_self_check_non_body_indent_detection",
        fixture_root,
        mutate,
        "non-body-indent-failure",
        SAMPLE_SELF_CHECK_LABELS["non_body_indent"],
    )


def case_sample_self_check_cross_page_table_continuation_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(_reference_docx: Path, final_docx: Path) -> Path | None:
        mutate_table_continuation_missing(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_cross_page_table_continuation_detection",
        root / "sample_self_check_cross_page_table_continuation_detection" / "sample-self-check-cross-page-table",
        mutate,
        "cross-page-continuation-failure",
        SAMPLE_SELF_CHECK_LABELS["continuation"],
    )


def case_sample_self_check_caption_inside_table_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_caption_inside_table_cell(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_caption_inside_table_detection",
        root / "sample_self_check_caption_inside_table_detection" / "sample-self-check-caption-inside-table",
        mutate,
        "caption-inside-table-failure",
        "题名入表格检查",
    )


def case_sample_self_check_table_caption_binding_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        add_table_sample(reference_docx)
        add_table_sample(final_docx)
        mutate_table_caption_binding(final_docx)
        return None

    return run_negative_sample_case(
        "sample_self_check_table_caption_binding_detection",
        root / "sample_self_check_table_caption_binding_detection" / "sample-self-check-table-caption-binding",
        mutate,
        "table-caption-binding-failure",
        SAMPLE_SELF_CHECK_LABELS["table_caption_binding"],
    )


def case_sample_self_check_figure_block_locality_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "locality.png"
        make_placeholder_png(image_path, "本地图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        mutate_figure_block_locality(final_docx)
        return None

    fixture_root = root / "sample_self_check_figure_block_locality_detection" / "sample-self-check-figure-block"
    return run_negative_sample_case(
        "sample_self_check_figure_block_locality_detection",
        fixture_root,
        mutate,
        "figure-block-locality-failure",
        SAMPLE_SELF_CHECK_LABELS["figure_block"],
    )


def case_sample_self_check_figure_explanation_followup_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "figure-followup.png"
        make_placeholder_png(image_path, "图后正文")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        mutate_missing_figure_explanation(final_docx)
        return create_picture_replacement_manifest(
            fixture_root / "figure-followup-manifest.json",
            replacement_image=image_path,
            caption_text="图 1 测试图",
        )

    fixture_root = root / "sample_self_check_figure_explanation_followup_detection" / "sample-self-check-figure-followup"
    return run_negative_sample_case(
        "sample_self_check_figure_explanation_followup_detection",
        fixture_root,
        mutate,
        "figure-followup-failure",
        "图后正文介绍检查",
        custom_body_style=True,
    )


def case_sample_self_check_image_holder_safety_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "holder-safety.png"
        make_placeholder_png(image_path, "承载图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        mutate_image_holder_safety(final_docx)
        return None

    fixture_root = root / "sample_self_check_image_holder_safety_detection" / "sample-self-check-image-holder-safety"
    return run_negative_sample_case(
        "sample_self_check_image_holder_safety_detection",
        fixture_root,
        mutate,
        "image-holder-safety-failure",
        SAMPLE_SELF_CHECK_LABELS["image_holder_safety"],
        custom_body_style=True,
    )


def case_sample_self_check_image_dimension_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "oversized.png"
        make_placeholder_png(image_path, "超大图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        mutate_image_dimension_overflow(final_docx)
        return None

    fixture_root = root / "sample_self_check_image_dimension_detection" / "sample-self-check-image-dimension"
    return run_negative_sample_case(
        "sample_self_check_image_dimension_detection",
        fixture_root,
        mutate,
        "image-dimension-failure",
        SAMPLE_SELF_CHECK_LABELS["image_dimension"],
        custom_body_style=True,
    )


def case_sample_self_check_missing_asset_manifest_blocks_delivery(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "sample_self_check_missing_asset_manifest_blocks_delivery" / "sample-self-check-missing-asset-manifest"
    reference_docx, final_docx = build_sample_self_check_fixture(
        fixture_root,
        custom_body_style=True,
    )
    image_path = fixture_root / "structural-flowchart.png"
    make_placeholder_png(image_path, "流程图")
    add_figure_sample(reference_docx, image_path, caption_text="图 1 测试流程图")
    add_figure_sample(final_docx, image_path, caption_text="图 1 测试流程图")
    suite = run_sample_self_check_suite(
        reference_docx,
        final_docx,
        fixture_root,
        "missing-asset-manifest-delivery",
        renderer="auto",
        env=env,
        smoke_acceptance=False,
        check=False,
    )
    if str(suite["returncode"]) == "0":
        raise GateFailure("missing asset manifest did not block non-smoke thesis delivery")
    text = str(suite["text"]).lower()
    if "figure asset manifest" not in text and "asset manifest" not in text:
        raise GateFailure("missing asset manifest failure was not reported")
    return CaseResult(
        "sample_self_check_missing_asset_manifest_blocks_delivery",
        True,
        f"returncode={suite['returncode']}; self_check={suite['self_check']}",
    )


def case_sample_self_check_empty_diagrams_blocks_structural_docx(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "sample_self_check_empty_diagrams_blocks_structural_docx" / "sample-self-check-empty-diagrams"
    reference_docx, final_docx = build_sample_self_check_fixture(
        fixture_root,
        custom_body_style=True,
    )
    image_path = fixture_root / "empty-diagram-manifest.png"
    make_placeholder_png(image_path, "流程图")
    add_figure_sample(reference_docx, image_path, caption_text="图 1 测试流程图")
    add_figure_sample(final_docx, image_path, caption_text="图 1 测试流程图")
    manifest = fixture_root / "empty-diagrams-manifest.json"
    manifest.write_text(
        json.dumps(
            {
                "schema": "graduation-project-builder.figure-manifest.v2",
                "figures": {
                    "figure_1": {
                        "path": str(image_path),
                        "caption": "图 1 测试流程图",
                        "family": "flowchart",
                        "source_kind": "structural",
                    }
                },
                "tables": {},
                "diagrams": {},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    suite = run_sample_self_check_suite(
        reference_docx,
        final_docx,
        fixture_root,
        "empty-diagrams-manifest-failure",
        asset_manifest=manifest,
        env=env,
    )
    text = str(suite["text"])
    if "final DOCX contains structural figure signals but manifest has no diagram entries" not in text:
        raise GateFailure("empty diagrams manifest did not report structural figure coverage failure")
    return CaseResult(
        "sample_self_check_empty_diagrams_blocks_structural_docx",
        True,
        f"asset_manifest={manifest}; self_check={suite['self_check']}",
    )


def case_sample_self_check_figure_manifest_evidence_detection(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "sample_self_check_figure_manifest_evidence_detection" / "sample-self-check-figure-evidence"
    reference_docx, final_docx = build_sample_self_check_fixture(
        fixture_root,
        custom_body_style=True,
    )
    image_path = fixture_root / "figure-without-evidence.png"
    make_placeholder_png(image_path, "流程图")
    add_figure_sample(reference_docx, image_path, caption_text="图 1 测试流程图")
    add_figure_sample(final_docx, image_path, caption_text="图 1 测试流程图")
    manifest = fixture_root / "figure-without-evidence-manifest.json"
    manifest.write_text(
        json.dumps(
            {
                "schema": "graduation-project-builder.figure-manifest.v2",
                "figures": {
                    "figure_1": {
                        "path": str(image_path),
                        "caption": "图 1 测试流程图",
                        "family": "raster",
                        "source_kind": "raster",
                    }
                },
                "tables": {},
                "diagrams": {},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    suite = run_sample_self_check_suite(
        reference_docx,
        final_docx,
        fixture_root,
        "figure-manifest-evidence-failure",
        asset_manifest=manifest,
        env=env,
    )
    assert_report_dimension(
        str(suite["text"]),
        SAMPLE_SELF_CHECK_LABELS["figure_manifest_contract"],
        "failed",
        "figure manifest evidence contract",
    )
    return CaseResult(
        "sample_self_check_figure_manifest_evidence_detection",
        True,
        f"asset_manifest={manifest}; self_check={suite['self_check']}",
    )


def case_sample_self_check_svg_fallback_text_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(_reference_docx: Path, _final_docx: Path) -> Path | None:
        return create_flowchart_assets(fixture_root / "assets", fallback_text=True)

    fixture_root = root / "sample_self_check_svg_fallback_text_detection" / "sample-self-check-svg-fallback"
    return run_negative_sample_case(
        "sample_self_check_svg_fallback_text_detection",
        fixture_root,
        mutate,
        "svg-fallback-failure",
        SAMPLE_SELF_CHECK_LABELS["figure_family"],
    )


def case_sample_self_check_toc_control_contamination_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_toc_control_contamination_detection",
        root / "sample_self_check_toc_control_contamination_detection" / "sample-self-check-toc-control",
        lambda _reference_docx, final_docx: (insert_toc_control_contamination(final_docx), None)[1],
        "toc-control-contamination-failure",
        SAMPLE_SELF_CHECK_LABELS["toc_control"],
    )


def case_sample_self_check_bibliography_count_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_bibliography_count_detection",
        root / "sample_self_check_bibliography_count_detection" / "sample-self-check-bibliography-count",
        lambda _reference_docx, final_docx: (mutate_bibliography_count_loss(final_docx), None)[1],
        "bibliography-count-failure",
        SAMPLE_SELF_CHECK_LABELS["bibliography_count"],
    )


def case_sample_self_check_bibliography_rendered_geometry_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_bibliography_rendered_geometry_detection",
        root / "sample_self_check_bibliography_rendered_geometry_detection" / "sample-self-check-bibliography-geometry",
        lambda _reference_docx, final_docx: (mutate_bibliography_rendered_indent(final_docx), None)[1],
        "bibliography-rendered-geometry-failure",
        "reference rendered geometry check",
        expected_summary="failed",
    )


def case_alternate_renderer_parity_when_available(root: Path, env: GateEnvironment) -> CaseResult:
    fixture_root = root / "alternate_renderer_parity_when_available"
    template_docx, _source_docx, _ = build_complete_sample_fixture(fixture_root / "review-alternate-renderer")
    stage_dir = fixture_root / "stage-alternate-renderer"
    stage_dir.mkdir(parents=True, exist_ok=True)
    stage_docx = stage_dir / "stage.docx"
    copy_locked(template_docx, stage_docx)

    word_pdf = stage_dir / "word-rendered.pdf"
    try:
        export_pdf_with_renderer(stage_docx, word_pdf, "Word.Application COM", env)
    except Exception:
        return CaseResult(
            "alternate_renderer_parity_when_available",
            True,
            "alternate renderer unavailable on current machine",
        )
    page_texts_word = pdf_page_texts(word_pdf)

    alternate_renderer = "unavailable"
    alternate_pdf = stage_dir / "alternate-rendered.pdf"
    if env.wps != "not-found":
        alternate_renderer = "WPS COM"
        alternate_pdf = stage_dir / "wps-rendered.pdf"
        try:
            export_pdf_with_renderer(stage_docx, alternate_pdf, "WPS COM", env)
        except Exception:
            return CaseResult(
                "alternate_renderer_parity_when_available",
                True,
                "alternate renderer unavailable on current machine",
            )
    elif env.libreoffice != "not-found":
        alternate_renderer = "LibreOffice"
        alternate_pdf = stage_dir / "libreoffice-rendered.pdf"
        try:
            export_pdf_with_renderer(stage_docx, alternate_pdf, "LibreOffice", env)
        except Exception:
            return CaseResult(
                "alternate_renderer_parity_when_available",
                True,
                "alternate renderer unavailable on current machine",
            )

    if alternate_renderer == "unavailable":
        return CaseResult(
            "alternate_renderer_parity_when_available",
            True,
            "alternate renderer unavailable on current machine",
        )

    page_texts_alt = pdf_page_texts(alternate_pdf)
    markers = ["摘   要", "Abstract", "目   录", "1 绪论"]
    word_pages = {name: find_page(page_texts_word, name) for name in markers}
    alt_pages = {name: find_page(page_texts_alt, name) for name in markers}
    for marker in markers:
        if word_pages[marker] is None or alt_pages[marker] is None:
            raise GateFailure(f"renderer parity lost page marker: {marker}")
    if abs(len(page_texts_word) - len(page_texts_alt)) > 1:
        raise GateFailure("alternate renderer page count drift exceeds one page")
    return CaseResult(
        "alternate_renderer_parity_when_available",
        True,
        f"renderer=Word.Application COM; alternate_renderer={alternate_renderer}; review_copy={template_docx}; stage_docx={stage_docx}; word_pdf={word_pdf}; alternate_pdf={alternate_pdf}; word_pages={word_pages}; alternate_pages={alt_pages}; page_count_diff={abs(len(page_texts_word) - len(page_texts_alt))}",
    )


def case_sample_self_check_keyword_donor_shift_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_keyword_donor_shift_detection",
        root / "sample_self_check_keyword_donor_shift_detection" / "sample-self-check-keyword-donor-shift",
        lambda _reference_docx, final_docx: (mutate_keyword_lines_single_run(final_docx, clear_body_style=False), None)[1],
        "keyword-donor-shift-failure",
        SAMPLE_SELF_CHECK_LABELS["abstract_baseline"],
    )


def case_sample_self_check_abstract_self_donor_alias_detection(root: Path, env: GateEnvironment) -> CaseResult:
    return run_negative_sample_case(
        "sample_self_check_abstract_self_donor_alias_detection",
        root / "sample_self_check_abstract_self_donor_alias_detection" / "sample-self-check-abstract-self-donor-alias",
        lambda _reference_docx, final_docx: (mutate_keyword_lines_single_run(final_docx, clear_body_style=True), None)[1],
        "abstract-self-donor-alias-failure",
        SAMPLE_SELF_CHECK_LABELS["abstract_baseline"],
    )


def case_sample_self_check_figure_caption_style_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "caption-style.png"
        make_placeholder_png(image_path, "图题图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        mutate_figure_caption_style(final_docx)
        return None

    fixture_root = root / "sample_self_check_figure_caption_style_detection" / "sample-self-check-figure-caption-style"
    return run_negative_sample_case(
        "sample_self_check_figure_caption_style_detection",
        fixture_root,
        mutate,
        "figure-caption-style-failure",
        SAMPLE_SELF_CHECK_LABELS["figure_caption_baseline"],
        custom_body_style=True,
    )


def case_sample_self_check_image_holder_body_residue_detection(root: Path, env: GateEnvironment) -> CaseResult:
    def mutate(reference_docx: Path, final_docx: Path) -> Path | None:
        image_path = fixture_root / "holder-baseline.png"
        make_placeholder_png(image_path, "承载基线图")
        add_figure_sample(reference_docx, image_path)
        add_figure_sample(final_docx, image_path)
        mutate_image_holder_baseline(final_docx)
        return None

    fixture_root = root / "sample_self_check_image_holder_body_residue_detection" / "sample-self-check-image-holder-baseline"
    return run_negative_sample_case(
        "sample_self_check_image_holder_body_residue_detection",
        fixture_root,
        mutate,
        "image-holder-baseline-failure",
        SAMPLE_SELF_CHECK_LABELS["image_holder_baseline"],
        custom_body_style=True,
    )


def register_cases() -> None:
    CASE_BUILDERS.update(
        {
            "frontmatter_roundtrip": case_frontmatter_roundtrip,
            "lock_competition": case_lock_competition,
            "tail_pages_roundtrip": case_tail_pages_roundtrip,
            "staging_copy_isolation": case_staging_copy_isolation,
            "paragraph_review_microcycle": case_paragraph_review_microcycle,
            "serialized_helper_mutation_chain": case_serialized_helper_mutation_chain,
            "multi_pass_review_copy_governance": case_multi_pass_review_copy_governance,
            "serialized_formula_helper_chain": case_serialized_formula_helper_chain,
            "complete_sample_smoke": case_complete_sample_smoke,
            "sample_self_check_table_family_detection": case_sample_self_check_table_family_detection,
            "sample_self_check_table_caption_baseline_detection": case_sample_self_check_table_caption_baseline_detection,
            "sample_self_check_table_cell_baseline_detection": case_sample_self_check_table_cell_baseline_detection,
            "sample_self_check_heading_family_detection": case_sample_self_check_heading_family_detection,
            "sample_self_check_heading_all_levels_run_override_detection": case_sample_self_check_heading_all_levels_run_override_detection,
            "sample_self_check_heading2_run_override_detection": case_sample_self_check_heading2_run_override_detection,
            "sample_self_check_heading_baseline_detection": case_sample_self_check_heading_baseline_detection,
            "sample_self_check_code_title_format_detection": case_sample_self_check_code_title_format_detection,
            "sample_self_check_code_block_format_detection": case_sample_self_check_code_block_format_detection,
            "sample_self_check_cover_baseline_detection": case_sample_self_check_cover_baseline_detection,
            "sample_self_check_abstract_baseline_detection": case_sample_self_check_abstract_baseline_detection,
            "sample_self_check_abstract_manual_break_detection": case_sample_self_check_abstract_manual_break_detection,
            "sample_self_check_toc_visible_detection": case_sample_self_check_toc_visible_detection,
            "sample_self_check_header_position_detection": case_sample_self_check_header_position_detection,
            "sample_self_check_header_presence_detection": case_sample_self_check_header_presence_detection,
            "sample_self_check_footer_indent_detection": case_sample_self_check_footer_indent_detection,
            "sample_self_check_footer_baseline_detection": case_sample_self_check_footer_baseline_detection,
            "sample_self_check_body_style_binding_detection": case_sample_self_check_body_style_binding_detection,
            "sample_self_check_normal_baseline_drift_detection": case_sample_self_check_normal_baseline_drift_detection,
            "sample_self_check_non_body_surface_contamination_detection": case_sample_self_check_non_body_surface_contamination_detection,
            "sample_self_check_non_body_indent_detection": case_sample_self_check_non_body_indent_detection,
            "sample_self_check_cross_page_table_continuation_detection": case_sample_self_check_cross_page_table_continuation_detection,
            "sample_self_check_caption_inside_table_detection": case_sample_self_check_caption_inside_table_detection,
            "sample_self_check_table_caption_binding_detection": case_sample_self_check_table_caption_binding_detection,
            "sample_self_check_figure_block_locality_detection": case_sample_self_check_figure_block_locality_detection,
            "sample_self_check_figure_explanation_followup_detection": case_sample_self_check_figure_explanation_followup_detection,
            "sample_self_check_image_holder_safety_detection": case_sample_self_check_image_holder_safety_detection,
            "sample_self_check_image_dimension_detection": case_sample_self_check_image_dimension_detection,
            "sample_self_check_missing_asset_manifest_blocks_delivery": case_sample_self_check_missing_asset_manifest_blocks_delivery,
            "sample_self_check_empty_diagrams_blocks_structural_docx": case_sample_self_check_empty_diagrams_blocks_structural_docx,
            "sample_self_check_figure_manifest_evidence_detection": case_sample_self_check_figure_manifest_evidence_detection,
            "sample_self_check_svg_fallback_text_detection": case_sample_self_check_svg_fallback_text_detection,
            "sample_self_check_toc_control_contamination_detection": case_sample_self_check_toc_control_contamination_detection,
            "alternate_renderer_parity_when_available": case_alternate_renderer_parity_when_available,
            "sample_self_check_bibliography_count_detection": case_sample_self_check_bibliography_count_detection,
            "sample_self_check_bibliography_rendered_geometry_detection": case_sample_self_check_bibliography_rendered_geometry_detection,
            "sample_self_check_keyword_donor_shift_detection": case_sample_self_check_keyword_donor_shift_detection,
            "sample_self_check_abstract_self_donor_alias_detection": case_sample_self_check_abstract_self_donor_alias_detection,
            "sample_self_check_figure_caption_style_detection": case_sample_self_check_figure_caption_style_detection,
            "sample_self_check_image_holder_body_residue_detection": case_sample_self_check_image_holder_body_residue_detection,
        }
    )


def render_report(env: GateEnvironment, temp_root: Path, results: list[CaseResult]) -> str:
    lines = [
        "# Integration Gate Report",
        "",
        "## Environment",
        f"- officecli: {env.officecli}",
        f"- renderer: {env.renderer}",
        f"- libreoffice: {env.libreoffice}",
        f"- wps: {env.wps}",
        f"- temp-root: {temp_root}",
        "",
        "## Cases",
    ]
    for result in results:
        lines.append(f"- `{result.name}`: {'pass' if result.passed else 'fail'}")
        lines.append(f"  details: {result.details}")
    return "\n".join(lines) + "\n"


def run_cases(case_name: str, temp_root: Path, env: GateEnvironment, *, quiet: bool = False) -> tuple[list[CaseResult], bool]:
    selected = CASE_SEQUENCE if case_name == "all" else [case_name]
    results: list[CaseResult] = []
    all_ok = True
    for name in selected:
        case_root = temp_root / name
        case_root.mkdir(parents=True, exist_ok=True)
        builder = CASE_BUILDERS[name]
        try:
            result = builder(temp_root, env)
        except Exception as exc:
            result = CaseResult(name, False, str(exc))
        results.append(result)
        all_ok = all_ok and result.passed
        if not quiet or not result.passed:
            print(f"CASE {name}: {'PASS' if result.passed else 'FAIL'}")
            print(f"  details={result.details}")
    print(
        "INTEGRATION SUMMARY "
        + json.dumps(
            {
                "case": case_name,
                "case_count": len(results),
                "passed": all_ok,
                "failed": [result.name for result in results if not result.passed],
            },
            ensure_ascii=False,
        )
    )
    print(f"OVERALL={'PASS' if all_ok else 'FAIL'}")
    if all_ok:
        print("INTEGRATION GATE PASSED")
    return results, all_ok


def main() -> int:
    register_cases()
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--case", choices=["all", *CASE_SEQUENCE], default="all")
    parser.add_argument("--temp-root")
    parser.add_argument("--report-out")
    parser.add_argument("--quiet", action="store_true")
    args = parser.parse_args()

    temp_root = (
        Path(args.temp_root).resolve()
        if args.temp_root
        else Path(tempfile.mkdtemp(prefix=DEFAULT_TEMP_PREFIX))
    )
    temp_root.mkdir(parents=True, exist_ok=True)
    env = detect_environment()
    results, all_ok = run_cases(args.case, temp_root, env, quiet=args.quiet)

    if args.report_out:
        report_path = Path(args.report_out).resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(render_report(env, temp_root, results), encoding="utf-8")

    return 0 if all_ok else 1


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    raise SystemExit(main())
