#!/usr/bin/env python3
"""Rebind CAD appendix images in a DOCX to a final CAD PNG source set."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import tempfile
import zipfile
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent

try:
    from audit_docx_cad_appendix_binding import (
        audit_docx_cad_appendix_binding,
        collect_cad_pngs,
        find_appendix_bounds,
        read_docx_paragraphs_and_images,
    )
except ImportError:  # pragma: no cover
    import sys

    sys.path.insert(0, str(SCRIPT_DIR))
    from audit_docx_cad_appendix_binding import (
        audit_docx_cad_appendix_binding,
        collect_cad_pngs,
        find_appendix_bounds,
        read_docx_paragraphs_and_images,
    )


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def media_member_from_target(target: str) -> str:
    target = target.replace("\\", "/").lstrip("/")
    if target.startswith("word/"):
        return target
    return f"word/{target}"


def repair_docx_cad_appendix_binding(
    source_docx: Path,
    cad_source: Path,
    output_docx: Path,
    old_version: str,
    new_version: str,
) -> dict[str, object]:
    paragraphs, media_by_target, source_sha = read_docx_paragraphs_and_images(source_docx)
    appendix_start, appendix_end = find_appendix_bounds(paragraphs)
    if appendix_start is None or appendix_end is None:
        raise ValueError("appendix block not found")

    appendix_rows = [
        row
        for row in paragraphs
        if appendix_start <= int(row["index"]) < appendix_end and row.get("image_targets")
    ]
    cad_pngs = collect_cad_pngs(cad_source)
    if len(appendix_rows) < len(cad_pngs):
        raise ValueError(f"appendix image rows {len(appendix_rows)} < CAD PNG count {len(cad_pngs)}")

    replacements: list[dict[str, object]] = []
    for row, cad_row in zip(appendix_rows, cad_pngs):
        targets = [str(target) for target in row.get("image_targets", [])]
        if not targets:
            continue
        target = targets[0]
        member = media_member_from_target(target)
        cad_path = Path(str(cad_row["path"]))
        if cad_source.is_dir():
            new_bytes = cad_path.read_bytes()
        else:
            with zipfile.ZipFile(cad_source) as zf:
                new_bytes = zf.read(str(cad_row["path"]))
        old_hash = str(media_by_target.get(target, {}).get("sha256", ""))
        replacements.append(
            {
                "paragraph_index": row["index"],
                "target": target,
                "member": member,
                "old_sha256": old_hash,
                "new_source": str(cad_row["path"]),
                "new_sha256": sha256_bytes(new_bytes),
                "new_bytes": len(new_bytes),
                "_payload": new_bytes,
            }
        )

    if len(replacements) < len(cad_pngs):
        raise ValueError(f"replacement count {len(replacements)} < CAD PNG count {len(cad_pngs)}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    tmp_output = output_docx.with_suffix(output_docx.suffix + ".tmp")
    if tmp_output.exists():
        tmp_output.unlink()
    payload_by_member = {str(row["member"]): row["_payload"] for row in replacements}
    with zipfile.ZipFile(source_docx, "r") as src, zipfile.ZipFile(tmp_output, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename in payload_by_member:
                data = payload_by_member[info.filename]
            elif info.filename.startswith("word/") and info.filename.endswith(".xml"):
                data = data.replace(old_version.encode("ascii"), new_version.encode("ascii"))
            dst.writestr(info, data)
    if output_docx.exists():
        output_docx.unlink()
    tmp_output.replace(output_docx)

    clean_replacements = [{key: value for key, value in row.items() if key != "_payload"} for row in replacements]
    final_report = audit_docx_cad_appendix_binding(output_docx, cad_source, len(cad_pngs), True)
    return {
        "schema": "graduation-project-builder.docx-cad-appendix-binding-repair.v1",
        "source_docx": str(source_docx),
        "source_docx_sha256": source_sha,
        "output_docx": str(output_docx),
        "output_docx_sha256": sha256_file(output_docx),
        "cad_source": str(cad_source),
        "old_version": old_version,
        "new_version": new_version,
        "appendix_start_index": appendix_start,
        "appendix_end_index": appendix_end,
        "replacement_count": len(clean_replacements),
        "replacements": clean_replacements,
        "post_repair_audit_passed": bool(final_report.get("passed")),
        "post_repair_audit": final_report,
    }


def run_self_test() -> int:
    from docx import Document

    png_a = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?"
        b"\x00\x05\xfe\x02\xfeA\x88\x0f\x9b\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    png_b = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc`\xf8\xcf\xc0"
        b"\x00\x00\x03\x01\x01\x00\xc9\xfe\x92\xef\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        old_dir = root / "old"
        cad_dir = root / "cad"
        old_dir.mkdir()
        cad_dir.mkdir()
        old_png = old_dir / "old.png"
        new_png = cad_dir / "sheet.png"
        old_png.write_bytes(png_a)
        new_png.write_bytes(png_b)
        source = root / "source.docx"
        output = root / "output.docx"
        doc = Document()
        doc.add_paragraph("References")
        doc.add_paragraph("1. X[J].")
        doc.add_paragraph("\u9644\u5f55A v14")
        doc.add_picture(str(old_png))
        doc.add_paragraph("Figure A.1 sheet v14")
        doc.add_paragraph("Acknowledgement")
        doc.save(source)
        report = repair_docx_cad_appendix_binding(source, cad_dir, output, "v14", "v16")
        if not report["post_repair_audit_passed"]:
            raise AssertionError(report)
        with zipfile.ZipFile(output) as zf:
            merged = b"".join(zf.read(name) for name in zf.namelist() if name.startswith("word/") and name.endswith(".xml"))
        if b"v14" in merged or b"v16" not in merged:
            raise AssertionError("version text replacement failed")
    print("self-test passed")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-docx", type=Path)
    parser.add_argument("--cad-source", type=Path)
    parser.add_argument("--output-docx", type=Path)
    parser.add_argument("--old-version", default="v14")
    parser.add_argument("--new-version", default="v16")
    parser.add_argument("--report-json", type=Path)
    parser.add_argument("--self-test", action="store_true")
    args = parser.parse_args(argv)
    if args.self_test:
        return run_self_test()
    if args.source_docx is None or args.cad_source is None or args.output_docx is None:
        parser.error("--source-docx, --cad-source, and --output-docx are required unless --self-test is used")
    report = repair_docx_cad_appendix_binding(
        args.source_docx,
        args.cad_source,
        args.output_docx,
        args.old_version,
        args.new_version,
    )
    if args.report_json:
        args.report_json.parent.mkdir(parents=True, exist_ok=True)
        args.report_json.write_text(json.dumps(report, ensure_ascii=True, indent=2), encoding="utf-8")
    else:
        print(json.dumps(report, ensure_ascii=True, indent=2))
    return 0 if report["post_repair_audit_passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
