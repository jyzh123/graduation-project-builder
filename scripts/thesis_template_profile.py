#!/usr/bin/env python3
"""Build and check template-derived page-class profile data for thesis runs."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
import hashlib
import json
import re
from pathlib import Path
from typing import Any

import fitz  # type: ignore
from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore


SCHEMA = "graduation-project-builder.template-profile.v1"
SURFACE_LABELS = {
    "cover_zh_title": "cover Chinese thesis title",
    "cover_en_title": "cover English thesis title",
    "cover_anchor": "cover anchor",
    "zh_abstract": "Chinese abstract",
    "en_abstract": "English abstract",
    "toc": "TOC",
    "first_body": "first body chapter",
}


def normalize(text: str) -> str:
    return re.sub(r"[\s\u25a1]+", "", text or "").strip().lower()


def heading_key(text: str) -> str:
    """Normalize a heading label while ignoring template instruction suffixes."""
    head = re.split(r"[\(\uff08]", str(text or "").strip(), maxsplit=1)[0]
    return normalize(head)


def contains_template_instruction(text: str) -> bool:
    compact = normalize(text)
    instruction_tokens = (
        "timesnewroman",
        "\u9ed1\u4f53",
        "\u5b8b\u4f53",
        "\u5c45\u4e2d",
        "\u6bb5\u524d",
        "\u6bb5\u540e",
        "\u884c\u8ddd",
        "\u5b57\u53f7",
        "\u5c0f\u4e8c",
        "\u5c0f\u56db",
        "\u6b63\u6587",
        "\u5173\u952e\u8bcd",
        "\u82f1\u6587\u5173\u952e\u8bcd",
        "\u4e2d\u6587\u5173\u952e\u8bcd",
        "\u6458\u8981\u6b63\u6587",
    )
    return any(token in compact for token in instruction_tokens)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def has_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text or "")


def looks_like_english_title(text: str) -> bool:
    stripped = str(text or "").strip()
    compact = normalize(stripped)
    if (
        compact.startswith("abstract")
        or compact.startswith("keywords")
        or compact.startswith("keyword")
        or compact.startswith("key words".replace(" ", ""))
        or "timesnewroman" in compact
        or "nanchanguniversity" in compact
        or "thesisofbachelor" in compact
        or contains_template_instruction(stripped)
    ):
        return False
    if re.match(r"^\s*(?:\d+\.|[\[\uff3b]\d+[\]\uff3d])", stripped):
        return False
    if re.match(r"^\s*[图表]\s*\d+(?:[.\uff0e]\d+)?", stripped):
        return False
    if re.search(r"[\[\uff3b]\s*[A-Z]\s*[\]\uff3d]", stripped):
        return False
    latin_count = sum(1 for ch in stripped if ("A" <= ch <= "Z") or ("a" <= ch <= "z"))
    return latin_count >= 12 and " " in stripped and not has_cjk(stripped)


def is_date_like(text: str) -> bool:
    compact = normalize(text)
    if not compact:
        return False
    date_chars = "0-9\u3007\u96f6\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341"
    return bool(re.fullmatch(rf"[{date_chars}]{{2,}}\u5e74[{date_chars}]{{1,3}}\u6708[{date_chars}]{{1,3}}\u65e5", compact))


def is_cover_title_candidate(text: str) -> bool:
    stripped = str(text or "").strip()
    compact = normalize(stripped)
    if not stripped or len(compact) < 6 or not has_cjk(stripped) or is_date_like(stripped):
        return False
    if re.match(r"^\s*(?:\d+\.|[\[\uff3b]\d+[\]\uff3d])", stripped):
        return False
    if re.match(r"^\s*[图表]\s*\d+(?:[.\uff0e]\d+)?", stripped):
        return False
    blocked = (
        "\u5b66\u58eb\u5b66\u4f4d\u8bba\u6587",
        "\u6bd5\u4e1a\u8bbe\u8ba1",
        "\u672c\u79d1\u751f\u6bd5\u4e1a\u8bba\u6587",
        "\u5927\u5e86\u5e08\u8303\u5b66\u9662",
        "\u672c\u79d1\u751f",
        "\u5357\u4eac\u90ae\u7535\u5927\u5b66",
        "\u6a21\u677f",
        "\u8bba\u6587\u9898\u76ee",
        "\u59d3\u540d",
        "\u5b66\u53f7",
        "\u9662\u7cfb",
        "\u9662\uff08\u7cfb\uff09",
        "\u9662(\u7cfb)",
        "\u4e13\u4e1a",
        "\u7814\u7a76\u65b9\u5411",
        "\u6307\u5bfc\u6559\u5e08",
        "\u539f\u521b\u6027\u58f0\u660e",
        "\u7248\u6743\u4f7f\u7528\u6388\u6743\u4e66",
        "\u6458\u8981",
        "abstract",
        "keywords",
        "keyword",
        "key words",
        "timesnewroman",
        "\u76ee\u5f55",
        "\u5173\u952e\u8bcd",
        "\u82f1\u6587\u5173\u952e\u8bcd",
        "\u4e2d\u6587\u5173\u952e\u8bcd",
        "\u6458\u8981\u6b63\u6587",
        "\u4e2d\u6587\u9898\u76ee",
        "\u9898\u76ee\u53ea\u6709\u4e00\u884c",
        "\u987b\u5220\u9664\u672c\u884c",
        "\u5177\u4f53\u4e66\u5199\u5f0f\u6837\u5982\u4e0b",
        "\u4e66\u5199\u5f0f\u6837",
        "\u5e74\u6708\u65e5",
        "\u8bba\u6587\u4f5c\u8005\u7b7e\u540d",
        "\u53c2\u8003\u6587\u732e",
        "\u5b66\u4f4d\u8bba\u6587\u683c\u5f0f",
        "\u5e8f\u53f7",
        "\u4f5c\u8005",
        "\u53d1\u8868\u5730",
        "\u5b66\u4f4d\u6388\u4e88\u5355\u4f4d",
        "\u5e74\u5ea6",
        "\u56fe\u7eb8",
        "cad",
        "\u5185\u8499\u53e4\u79d1\u6280\u5927\u5b66",
    )
    return not contains_template_instruction(stripped) and not any(normalize(token) in compact for token in blocked)


def paragraph_has_page_break(paragraph: Any) -> bool:
    return any(br.get(qn("w:type")) == "page" for br in paragraph._element.findall(".//w:br", paragraph._element.nsmap))


def cover_end_index(doc: Document) -> int:
    markers = {
        heading_key("\u5b66\u4f4d\u8bba\u6587\u539f\u521b\u6027\u58f0\u660e"),
        heading_key("\u6458\u8981"),
        heading_key("Abstract"),
        heading_key("ABSTRACT"),
    }
    for idx, paragraph in enumerate(doc.paragraphs):
        if heading_key(paragraph.text) in markers:
            return idx
    return min(len(doc.paragraphs), 80)


def first_page_break_after(doc: Document, start: int, end: int) -> int | None:
    for idx in range(start, min(end, len(doc.paragraphs))):
        if paragraph_has_page_break(doc.paragraphs[idx]):
            return idx
    return None


def collect_pdf_texts(pdf_path: Path | None) -> list[str]:
    if pdf_path is None or not pdf_path.exists():
        return []
    with fitz.open(pdf_path) as pdf:
        return [page.get_text("text") for page in pdf]


def find_page(pdf_texts: list[str], marker: str) -> int | None:
    target = normalize(marker)
    if not target:
        return None
    for idx, text in enumerate(pdf_texts, start=1):
        if target in normalize(text):
            return idx
    return None


def page_contains_standalone_marker(text: str, marker: str) -> bool:
    target = normalize(marker)
    standalone_markers = {
        "\u6458\u8981",
        "Abstract",
        "\u76ee\u5f55",
        "\u7b2c\u4e00\u7ae0\u7eea\u8bba",
        "\u7b2c\u4e00\u7ae0 \u7eea\u8bba",
        "\u53c2\u8003\u6587\u732e",
        "\u81f4\u8c22",
    }
    if target not in {normalize(item) for item in standalone_markers}:
        return target in normalize(text)
    for line in (text or "").splitlines():
        if normalize(line) == target:
            return True
    return False


def find_standalone_page(pdf_texts: list[str], marker: str) -> int | None:
    target = normalize(marker)
    if not target:
        return None
    for idx, text in enumerate(pdf_texts, start=1):
        if page_contains_standalone_marker(text, marker):
            return idx
    return None


def find_any_page(pdf_texts: list[str], markers: list[str]) -> int | None:
    for marker in markers:
        page = find_page(pdf_texts, marker)
        if page is not None:
            return page
    return None


def find_heading_marker(doc: Document, candidates: set[str]) -> tuple[str, int] | None:
    normalized_candidates = {heading_key(item) for item in candidates}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if heading_key(text) in normalized_candidates:
            return text, idx
    return None


def iter_sdt_paragraph_texts(doc: Document) -> list[tuple[str, int]]:
    results: list[tuple[str, int]] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            results.append((text, idx))
    base = len(results)
    for offset, sdt in enumerate(doc.part.element.findall(".//w:sdt", doc.part.element.nsmap)):
        for paragraph in sdt.findall("./w:sdtContent/w:p", doc.part.element.nsmap):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", doc.part.element.nsmap)).strip()
            if text:
                results.append((text, base + offset))
    return results


def find_heading_marker_anywhere(doc: Document, candidates: set[str]) -> tuple[str, int] | None:
    normalized_candidates = {heading_key(item) for item in candidates}
    for text, idx in iter_sdt_paragraph_texts(doc):
        if heading_key(text) in normalized_candidates:
            return text, idx
    return None


def find_heading_or_prefix_marker(doc: Document, candidates: set[str]) -> tuple[str, int] | None:
    """Find front-matter headings that may be inline labels, e.g. `摘  要：...`."""
    normalized_candidates = {heading_key(item) for item in candidates}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        key = heading_key(text)
        if key in normalized_candidates:
            return text, idx
        prefix = text.split(":", 1)[0].split("：", 1)[0]
        if heading_key(prefix) in normalized_candidates:
            return text, idx
    return None


def paragraph_style_name(paragraph: Any) -> str:
    try:
        return str(paragraph.style.name or "")
    except Exception:
        return ""


def looks_like_toc_entry(paragraph: Any) -> bool:
    text = str(paragraph.text or "").strip()
    if not text:
        return False
    style_name = normalize(paragraph_style_name(paragraph))
    if style_name.startswith("toc") or "\u76ee\u5f55" in style_name:
        return True
    # Static or cached TOC rows commonly retain the real heading text followed
    # by a tab/leader and a rendered page token. Those rows must not become
    # the body-opening marker for the template profile.
    if "\t" in text:
        return True
    if re.search(r"(?:\.{3,}|\u2026+|\u00b7{3,})\s*(?:[ivxlcdm]+|\d+)\s*$", text, flags=re.IGNORECASE):
        return True
    return False


def looks_like_body_chapter_marker(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    if re.match(r"^\d{1,2}\s+\S", stripped):
        return True
    return bool(re.match(r"^\u7b2c\s*[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*\u7ae0", stripped))


def find_first_body_marker(doc: Document) -> tuple[str, int] | None:
    toc_found = find_heading_marker_anywhere(doc, {"\u76ee\u5f55", "\u76ee  \u5f55", "\u76ee\u25a1\u25a1\u5f55"})
    toc_idx = toc_found[1] if toc_found else -1
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not looks_like_body_chapter_marker(text):
            continue
        if idx <= toc_idx:
            continue
        if looks_like_toc_entry(paragraph):
            continue
        return text, idx
    # Some templates store visible body markers in content controls. Use this
    # only as a fallback and still reject TOC-looking cached rows.
    for text, idx in iter_sdt_paragraph_texts(doc):
        if idx <= toc_idx or not looks_like_body_chapter_marker(text):
            continue
        if "\t" in text:
            continue
        return text, idx
    return None


def select_chinese_title_indices(cover_items: list[tuple[int, Any]], english_idx: int | None, end: int) -> list[int]:
    title_search_end = english_idx if english_idx is not None else end
    candidates = [
        idx
        for idx, paragraph in cover_items
        if idx < title_search_end and is_cover_title_candidate(paragraph.text)
    ]
    if not candidates:
        return []
    if english_idx is not None:
        near = [idx for idx in candidates if 0 < english_idx - idx <= 6]
        if near:
            return near
    return candidates[-3:]


def looks_like_format_requirements_document(doc: Document) -> bool:
    sample = "\n".join(paragraph.text.strip() for paragraph in doc.paragraphs[:30] if paragraph.text.strip())
    compact = normalize(sample)
    required = (
        "\u64b0\u5199\u4e0e\u88c5\u8ba2\u89c4\u8303",
        "\u6bd5\u4e1a\u8bbe\u8ba1\u8bf4\u660e\u4e66\u6216\u6bd5\u4e1a\u8bba\u6587\u4e3b\u8981\u90e8\u5206",
        "\u7248\u5f0f",
    )
    return any(normalize(token) in compact for token in required)


def build_template_profile(template_docx: Path, template_pdf: Path | None = None) -> dict[str, Any]:
    generated_at = datetime.now(timezone.utc)
    doc = Document(str(template_docx))
    pdf_texts = collect_pdf_texts(template_pdf)
    end = cover_end_index(doc)
    cover_items = list(enumerate(doc.paragraphs[:end]))
    english_idx = next((idx for idx, paragraph in cover_items if looks_like_english_title(paragraph.text)), None)
    zh_indices = select_chinese_title_indices(cover_items, english_idx, end)
    if looks_like_format_requirements_document(doc):
        english_idx = None
        zh_indices = []
    cover_anchor = next(
        (
            p.text.strip()
            for _idx, p in cover_items
            if "\u5b66\u58eb\u5b66\u4f4d\u8bba\u6587" in p.text
            or "\u6bd5\u4e1a\u8bbe\u8ba1" in p.text
            or "\u672c\u79d1\u751f\u6bd5\u4e1a\u8bba\u6587" in p.text
        ),
        "",
    )
    zh_title_markers = [doc.paragraphs[idx].text.strip() for idx in zh_indices if doc.paragraphs[idx].text.strip()]
    en_title_marker = doc.paragraphs[english_idx].text.strip() if english_idx is not None else ""
    markers: dict[str, dict[str, Any]] = {
        "cover_anchor": {
            "label": SURFACE_LABELS["cover_anchor"],
            "text": cover_anchor,
            "paragraph_indices": [],
            "template_page": find_page(pdf_texts, cover_anchor) if cover_anchor else None,
        },
        "cover_zh_title": {
            "label": SURFACE_LABELS["cover_zh_title"],
            "text": "\n".join(zh_title_markers),
            "paragraph_indices": zh_indices,
            "template_page": find_any_page(pdf_texts, ["".join(zh_title_markers), *zh_title_markers]),
        },
        "cover_en_title": {
            "label": SURFACE_LABELS["cover_en_title"],
            "text": en_title_marker,
            "paragraph_indices": [english_idx] if english_idx is not None else [],
            "template_page": find_page(pdf_texts, en_title_marker) if en_title_marker else None,
        },
    }
    heading_specs = {
        "zh_abstract": {"\u6458\u8981", "\u6458  \u8981", "\u6458\u25a1\u25a1\u8981"},
        "en_abstract": {"Abstract", "ABSTRACT"},
        "toc": {"\u76ee\u5f55", "\u76ee  \u5f55", "\u76ee\u25a1\u25a1\u5f55"},
    }
    for key, candidates in heading_specs.items():
        found = find_heading_or_prefix_marker(doc, candidates) if key in {"zh_abstract", "en_abstract"} else find_heading_marker_anywhere(doc, candidates)
        text, idx = found if found else ("", -1)
        page_markers = [text, *candidates] if text else []
        template_page = find_any_page(pdf_texts, page_markers) if page_markers else None
        if key in {"zh_abstract", "en_abstract", "toc"} and text:
            template_page = find_standalone_page(pdf_texts, text) or template_page
        markers[key] = {
            "label": SURFACE_LABELS[key],
            "text": text,
            "paragraph_indices": [idx] if idx >= 0 else [],
            "template_page": template_page,
        }
    body = find_first_body_marker(doc)
    text, idx = body if body else ("", -1)
    markers["first_body"] = {
        "label": SURFACE_LABELS["first_body"],
        "text": text,
        "paragraph_indices": [idx] if idx >= 0 else [],
        "template_page": find_standalone_page(pdf_texts, text) if text else None,
    }

    same_page_groups: list[dict[str, Any]] = []
    zh_page = markers["cover_zh_title"].get("template_page")
    en_page = markers["cover_en_title"].get("template_page")
    if zh_title_markers and en_title_marker:
        page_break_between = None
        if zh_indices and english_idx is not None:
            page_break_between = first_page_break_after(doc, zh_indices[-1], english_idx + 1)
        same_by_pdf = zh_page is not None and en_page is not None and zh_page == en_page
        same_by_docx = page_break_between is None
        if same_by_pdf or same_by_docx:
            same_page_groups.append(
                {
                    "id": "cover-title-block",
                    "members": ["cover_zh_title", "cover_en_title"],
                    "source": "template-rendered-page" if same_by_pdf else "docx-page-break-boundary",
                    "template_page": zh_page if same_by_pdf else None,
                    "paragraph_indices": zh_indices + ([english_idx] if english_idx is not None else []),
                }
            )

    separated_pairs: list[dict[str, Any]] = []
    for left, right in (("cover_zh_title", "zh_abstract"), ("zh_abstract", "en_abstract"), ("en_abstract", "toc"), ("toc", "first_body")):
        left_page = markers.get(left, {}).get("template_page")
        right_page = markers.get(right, {}).get("template_page")
        if left_page is not None and right_page is not None and left_page != right_page:
            separated_pairs.append({"left": left, "right": right, "source": "template-rendered-page"})

    critical_surfaces = ["zh_abstract", "en_abstract", "toc", "first_body"]
    if zh_title_markers:
        critical_surfaces.insert(0, "cover_zh_title")
    elif cover_anchor:
        critical_surfaces.insert(0, "cover_anchor")
    if en_title_marker:
        critical_surfaces.insert(1, "cover_en_title")
    optional_absent = []
    if not zh_title_markers:
        optional_absent.append("cover_zh_title")
    if not en_title_marker:
        optional_absent.append("cover_en_title")
    for key in ("zh_abstract", "en_abstract", "first_body"):
        if not markers.get(key, {}).get("text"):
            optional_absent.append(key)

    return {
        "schema": SCHEMA,
        "generation_stage": "pre-mutation-template-profile-lock",
        "generated_at_utc": generated_at.isoformat().replace("+00:00", "Z"),
        "generated_at_unix": generated_at.timestamp(),
        "generator": "scripts/thesis_template_profile.py",
        "template_docx": str(template_docx),
        "template_fingerprint": sha256_file(template_docx),
        "template_fingerprint_algorithm": "sha256",
        "template_pdf": str(template_pdf) if template_pdf else "",
        "markers": markers,
        "front_matter": {
            "same_page_groups": same_page_groups,
            "separated_page_pairs": separated_pairs,
            "critical_surfaces": critical_surfaces,
            "optional_absent_surfaces": optional_absent,
        },
    }


def profile_readiness_issues(profile: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    if profile.get("schema") != SCHEMA:
        issues.append("template profile schema mismatch")
    markers = profile.get("markers")
    if not isinstance(markers, dict):
        return issues + ["template profile markers missing"]
    optional_absent = set(profile.get("front_matter", {}).get("optional_absent_surfaces", []))
    for key in profile.get("front_matter", {}).get("critical_surfaces", []):
        marker = markers.get(key, {})
        if not marker.get("text") and key not in {"toc", "first_body"} and key not in optional_absent:
            issues.append(f"template profile missing critical marker: {key}")
    first_body = markers.get("first_body", {})
    first_body_text = str(first_body.get("text") or "")
    if "\t" in first_body_text:
        issues.append("template profile first_body marker appears to be a TOC/cache row")
    return issues


def write_profile(profile: dict[str, Any], output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(profile, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True)
    parser.add_argument("--template-pdf")
    parser.add_argument("--output", required=True)
    parser.add_argument("--fail-on-issues", action="store_true")
    args = parser.parse_args()
    profile = build_template_profile(Path(args.template), Path(args.template_pdf) if args.template_pdf else None)
    write_profile(profile, Path(args.output))
    issues = profile_readiness_issues(profile)
    if issues:
        for issue in issues:
            print(issue)
        return 1 if args.fail_on_issues else 0
    print(f"template_profile={Path(args.output).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
